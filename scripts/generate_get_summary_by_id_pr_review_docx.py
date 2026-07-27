"""Generate Word document: PR review for GET summary by id (WBS 3.11.6)."""

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.shared import Pt

OUTPUT = (
    Path(__file__).resolve().parents[1]
    / "docs"
    / "PR_Review_Get_Summary_By_Id_feature_a-fasaa-get-summary-by-id.docx"
)


def heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def para(doc: Document, text: str, bold: bool = False, highlight: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    if highlight:
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def code(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    for i, h in enumerate(headers):
        tbl.rows[0].cells[i].text = h
    for row in rows:
        cells = tbl.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val


def build() -> Document:
    doc = Document()
    title = doc.add_heading("PR Code Review", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    para(
        doc,
        "WBS 3.11.6 — GET summary by id — "
        "feature/a-fasaa-get-summary-by-id → team-ae-develop",
    )
    para(doc, f"Review date: {date.today().isoformat()}")
    para(doc, "Reviewer: AI-assisted code review (Cursor)")
    doc.add_paragraph()

    heading(doc, "Review metadata", 1)
    table(
        doc,
        ["Field", "Value"],
        [
            ["Source branch", "feature/a-fasaa-get-summary-by-id"],
            ["Target branch", "team-ae-develop"],
            ["Commits", "11c26ff — feat(3.11.6): GET /api/v3/summaries/{id} endpoint"],
            ["Scope", "6 files (+275 / −4)"],
        ],
    )

    heading(doc, "1. Change Summary", 1)
    para(
        doc,
        "Adds a read-side endpoint GET /api/v3/summaries/{id} so consumers can fetch a "
        "specific CallSummary row by primary key (the id carried on SUMMARY_CREATED), "
        "instead of only 'latest by callId'. Response shape matches getLatestSummary so "
        "clients can share deserializers. Also adds GlobalExceptionHandler mapping for "
        "MethodArgumentTypeMismatchException → 400 (fixes invalid path/query types that "
        "previously bubbled as 500), with a corresponding ScheduledVisitControllerTest update.",
    )

    heading(doc, "What changed", 2)
    table(
        doc,
        ["File", "Change"],
        [
            [
                "CallSummaryController.java",
                "New controller at /api/v3/summaries; GET /{id} → 200/404",
            ],
            [
                "CallSummaryService.java",
                "getSummaryEntityById / getSummaryById wrapping repository.findById + toResponse",
            ],
            [
                "GlobalExceptionHandler.java",
                "MethodArgumentTypeMismatchException → 400 with parameter/type hint",
            ],
            [
                "CallSummaryControllerTest.java",
                "WebMvcTest: 200 body, 404, 400 non-numeric id",
            ],
            [
                "CallSummaryServiceTest.java",
                "null id, not found, SUCCESS found, ERROR status still returned",
            ],
            [
                "ScheduledVisitControllerTest.java",
                "Invalid date format expectation 500 → 400 (handler side effect)",
            ],
        ],
    )

    heading(doc, "Verdict", 2)
    para(
        doc,
        "Approve with changes requested on authorization. The endpoint and 400 handler are "
        "clean and well-tested, but any authenticated user can currently fetch any summary "
        "by id (IDOR) because there is no patient/caregiver scope check and no role gate. "
        "That is acknowledged in the controller Javadoc as follow-up, but it should be "
        "treated as a merge blocker for any environment with real PHI summaries — or the "
        "endpoint must stay behind a feature flag / internal-only access until RBAC lands.",
        bold=True,
        highlight=True,
    )

    heading(doc, "2. Bug & Risk Analysis", 1)

    heading(doc, "2.1 Strengths", 2)
    bullets(
        doc,
        [
            "Small, focused PR aligned with WBS 3.11.6 and SUMMARY_CREATED correlation by summaryId.",
            "Reuses toResponse for a consistent contract with GET /api/v3/calls/{callId}/summary.",
            "Null-safe service API; ERROR/NO_TRANSCRIPT rows remain readable by id (useful for debug).",
            "Type-mismatch → 400 is a real platform improvement (ScheduledVisit date format case).",
            "Good WebMvcTest + service unit coverage for happy/404/400/null paths.",
        ],
    )

    heading(doc, "2.2 High — IDOR: no patient/caregiver authorization on read", 2)
    para(
        doc,
        "SecurityConfig authenticates /api/** but does not scope by patient. "
        "CallSummaryController.getSummaryById loads any id with no AuthorizationService / "
        "CaregiverPatientLink / patientId check. Any logged-in user who can guess or obtain "
        "a summary id (e.g. from logs, sequential ids, indexing events) can read another "
        "patient's clinical summary JSON.",
        highlight=True,
    )
    code(
        doc,
        """// Recommended sketch (mirror CallController access patterns)
@GetMapping("/{id}")
public ResponseEntity<Map<String, Object>> getSummaryById(@PathVariable Long id) {
    User user = securityUtil.resolveCurrentUser();
    CallSummary summary = callSummaryService.getSummaryEntityById(id)
        .orElse(null);
    if (summary == null) {
        return ResponseEntity.notFound().build();
    }
    authorizationService.requireAccessToPatientCallOrSummary(user, summary);
    // enforce caregiverVisibility (on_consent / hidden) for caregiver callers
    return ResponseEntity.ok(callSummaryService.toResponse(summary));
}""",
    )
    bullets(
        doc,
        [
            "Also return 404 (not 403) for unauthorized ids if you want to avoid existence leaks — "
            "or 403 if product prefers explicit denial; pick one and document it.",
            "CaregiverVisibility on the entity should be enforced for caregiver roles.",
        ],
    )

    heading(doc, "2.3 Medium — Path / API version inconsistency vs WBS text", 2)
    para(
        doc,
        "WBS text cited as GET /api/summaries/{id}; implementation is /api/v3/summaries/{id}. "
        "Javadoc explains alignment with CallController /api/v3/calls — fine if clients are "
        "updated, but confirm Flutter/indexer consumers use the v3 path.",
    )

    heading(doc, "2.4 Medium — Global 400 handler may change other endpoints' contracts", 2)
    para(
        doc,
        "MethodArgumentTypeMismatchException was previously falling into the generic Exception "
        "handler (often 500). Mapping to 400 is correct, but any client that relied on 500 for "
        "bad path types will change behavior (ScheduledVisitControllerTest already updated). "
        "Worth a quick sweep of other WebMvcTests expecting 500 on bad formats.",
    )

    heading(doc, "2.5 Low — Missing trailing newline on new files", 2)
    para(
        doc,
        "CallSummaryController.java and CallSummaryControllerTest.java lack a final newline "
        "(diff shows \\ No newline at end of file). Minor style/POSIX nit.",
    )

    heading(doc, "2.6 Low — No explicit SecurityConfig matcher for /api/v3/summaries", 2)
    para(
        doc,
        "Covered by catch-all .requestMatchers(\"/api/**\").authenticated(). Explicit matcher "
        "optional for clarity next to /api/v3/calls/**.",
    )

    heading(doc, "2.7 Low — Response does not include summary id", 2)
    para(
        doc,
        "toResponse omits the database id. Clients that fetched by id already know it; "
        "including \"id\": summary.getId() would help deep links and caching. Optional.",
    )

    heading(doc, "2.8 No race conditions", 2)
    para(
        doc,
        "Read-only findById + map. No concurrency concerns introduced.",
    )

    heading(doc, "3. Architecture & Style", 1)
    bullets(
        doc,
        [
            "Controller thin; service owns lookup — good separation.",
            "Honest Javadoc about deferred @PreAuthorize / EnableMethodSecurity — transparent, "
            "but does not remove the IDOR risk for merge.",
            "Reusing toResponse avoids contract drift between callId and id reads.",
            "Tests follow existing WebMvcTest + CareconnectTestConfig patterns.",
        ],
    )

    heading(doc, "4. Recommendations", 1)

    heading(doc, "4.1 [High] Enforce patient-scoped access before returning body", 2)
    para(doc, "See §2.2. Minimum viable gate:")
    code(
        doc,
        """// In CallSummaryService or a dedicated authorizer
public void assertCanReadSummary(User user, CallSummary summary) {
    Long patientId = summary.getPatientId();
    if (patientId == null) {
        // historic rows: resolve via call telemetry / deny
        throw new UnauthorizedException("Summary has no patient scope");
    }
    if (user is patient and owns patientId) return;
    if (user is caregiver/family with active link + visibility allows) return;
    if (user is admin) return;
    throw new UnauthorizedException("Not allowed to read this summary");
}""",
    )

    heading(doc, "4.2 [Medium] Add authz tests", 2)
    code(
        doc,
        """@Test
@WithMockUser(roles = "PATIENT")
void getSummaryById_otherPatient_returns403Or404() throws Exception {
    // stub service/auth so patient A cannot read patient B's summary
    mockMvc.perform(get("/api/v3/summaries/{id}", 101L))
        .andExpect(status().isForbidden()); // or isNotFound()
}""",
    )

    heading(doc, "4.3 [Low] Include id in response map", 2)
    code(
        doc,
        """response.put("id", summary.getId());
response.put("patientId", summary.getPatientId()); // only if caller is authorized
response.put("callId", summary.getCallId());""",
    )

    heading(doc, "4.4 [Low] Add final newlines; optional explicit security matcher", 2)
    code(
        doc,
        """.requestMatchers("/api/v3/summaries/**").authenticated()
.requestMatchers("/api/v3/calls/**").authenticated()""",
    )

    heading(doc, "File-level review comments", 1)
    table(
        doc,
        ["File", "Severity", "Comment"],
        [
            [
                "CallSummaryController.java",
                "High",
                "No RBAC/patient scope — IDOR for any authenticated user.",
            ],
            [
                "CallSummaryController.java",
                "Info",
                "Clear WBS/path documentation; deferred PreAuthorize noted.",
            ],
            [
                "CallSummaryService.java",
                "Info",
                "Clean getSummaryById; consider exposing id in toResponse.",
            ],
            [
                "GlobalExceptionHandler.java",
                "Medium",
                "Good 400 mapping; audit other tests expecting 500 on type mismatch.",
            ],
            [
                "CallSummaryControllerTest.java",
                "Low",
                "Add unauthorized cross-patient case once authz lands.",
            ],
            [
                "CallSummaryServiceTest.java",
                "Info",
                "Solid null/not-found/SUCCESS/ERROR coverage.",
            ],
            [
                "ScheduledVisitControllerTest.java",
                "Info",
                "Correctly updated for 400 — validates handler impact.",
            ],
        ],
    )

    heading(doc, "Test plan (recommended before merge)", 1)
    bullets(
        doc,
        [
            "mvnw test -Dtest=CallSummaryControllerTest,CallSummaryServiceTest,"
            "ScheduledVisitControllerTest",
            "Authenticated user A cannot read user B's summary id (after authz fix).",
            "Caregiver with on_consent + no consent cannot read hidden/on_consent summary.",
            "Non-numeric id → 400 JSON body with parameter name.",
            "Confirm Flutter / indexer use /api/v3/summaries/{id} not /api/summaries/{id}.",
        ],
    )

    return doc


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    build().save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    main()
