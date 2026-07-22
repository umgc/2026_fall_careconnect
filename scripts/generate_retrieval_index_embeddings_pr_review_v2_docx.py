"""Generate Word document: PR review v2 for embeddings branch (post lease/park fixes)."""

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.shared import Pt

OUTPUT = (
    Path(__file__).resolve().parents[1]
    / "docs"
    / "PR_Review_Retrieval_Index_Embeddings_feature_a-rvasireddy_v2.docx"
)


def heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def para(doc: Document, text: str, bold: bool = False, highlight: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    if highlight:
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def code(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    for i, h in enumerate(headers):
        tbl.rows[0].cells[i].text = h
    for row in rows:
        cells = tbl.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val


def build() -> Document:
    doc = Document()
    title = doc.add_heading("PR Code Review", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    para(
        doc,
        "Tasks 4.2 + 4.3 — FTS, claim lease, Bedrock Titan embeddings — "
        "feature/a-rvasireddy-retrieval-index-embeddings → team-ae-develop",
    )
    para(doc, f"Review date: {date.today().isoformat()}")
    para(doc, "Reviewer: AI-assisted code review (Cursor)")
    doc.add_paragraph()

    heading(doc, "Review metadata", 1)
    table(
        doc,
        ["Field", "Value"],
        [
            ["Source branch", "feature/a-rvasireddy-retrieval-index-embeddings"],
            ["Target branch", "team-ae-develop"],
            [
                "Committed ahead",
                "c8fc9bf FTS; 071384f lease/FTS filter; 977ac1b Titan embed; "
                "6fe9d5d after-commit embed",
            ],
            [
                "Uncommitted (required)",
                "Visit park keeps claimed_at; claim SQL make_interval; "
                "IndexWorkerTest + IndexingOutboxClaimSqlTest; smoke script",
            ],
            ["Committed scope", "31 files (+1341 / −61)"],
            [
                "Base note",
                "team-ae-develop tip advanced (e.g. #305) — rebase before merge",
            ],
            [
                "Verdict",
                "Approve with changes — commit lease SQL + visit-park fixes first",
            ],
        ],
    )

    heading(doc, "1. Change Summary", 1)
    para(
        doc,
        "Completes the Ask AI retrieval-index write path after IndexWorker (#297). "
        "Task 4.2: search_vector backfill, FullTextSearchService (patient-scoped "
        "plainto_tsquery / ts_rank_cd), record_type filter before LIMIT. Outbox hardenings: "
        "durable claimed_at lease (multi-ECS), visit deferral without burning attempts. "
        "Task 4.3: ChunkEmbeddingService invokes amazon.titan-embed-text-v1 (1536-d), "
        "writes pgvector via updateEmbedding after ingest commits (afterCommit), "
        "best-effort on Bedrock failure, Titan v2 rejected at construction, throttle "
        "retries, contentHash skip retries embed when vectors still NULL. ECS IAM adds "
        "titan-embed-text-v1.",
    )

    heading(doc, "Working-tree fixes (not in HEAD commit yet)", 2)
    bullets(
        doc,
        [
            "releaseClaimWithoutBurn keeps claimed_at (stops visit reclaim every ~15s)",
            "claim query uses make_interval(mins => :leaseMinutes) instead of int || text",
            "Postgres smoke passed for make_interval claim selection",
        ],
        )
    para(
        doc,
        "Do not open/merge the PR until these are committed — committed HEAD still "
        "nulls the lease on visit park and still uses CAST((:leaseMinutes || ' minutes')).",
        highlight=True,
    )

    heading(doc, "2. Bug & Risk Analysis", 1)

    heading(doc, "High (if uncommitted fixes not merged) — visit reclaim spam", 2)
    para(
        doc,
        "On HEAD commit, releaseClaimWithoutBurn sets claimed_at=null. Next poll "
        "(~15s) reclaims the same visit SUMMARY_CREATED forever until Task 1.4 — "
        "log spam and wasted claim work. Working tree keeps the lease; must be committed.",
        highlight=True,
    )

    heading(doc, "High (if uncommitted fixes not merged) — int || text interval", 2)
    para(
        doc,
        "Hibernate often binds :leaseMinutes as integer. Postgres can reject "
        "integer || unknown at runtime. Working tree uses make_interval; smoke against "
        "local Postgres selected the correct claimable rows. Commit before merge.",
        highlight=True,
    )

    heading(doc, "Medium — visit park still retries every lease window", 2)
    para(
        doc,
        "Keeping claimed_at reduces reclaim to ~claim-lease-minutes (default 10), "
        "not forever silence. Until Task 1.4, expect warn logs every lease expiry. "
        "Optional: dedicated parked_until far in the future, or attempt_count parking "
        "marker, or pause visit event_type in the claim query.",
    )

    heading(doc, "Medium — burn-attempt deferrals still clear the lease", 2)
    para(
        doc,
        "recordDeferOrDeadLetter / recordFailure still set claimed_at=null. "
        "patientId-missing deferrals reclaim every poll until maxAttempts. "
        "Often acceptable (progress toward dead-letter), but can amplify under a "
        "burst of bad payloads. Consider leaving a short lease on burn-deferrals too.",
    )

    heading(doc, "Low — afterCommit embed still on poller thread", 2)
    para(
        doc,
        "Bedrock is outside the JDBC ingest TX (good), but still runs synchronously "
        "on the IndexWorker thread after commit (including Thread.sleep throttle "
        "backoff). Large transcript batches can delay subsequent outbox rows in the "
        "same poll cycle. Task 4.4 delivers ChunkEmbeddingBackfillWorker plus optional "
        "SchemaPatchRunner partial index (V2607161317) for NULL-embedding scans.",
    )

    heading(doc, "Low — self-invocation / TX for embedAndPersist", 2)
    para(
        doc,
        "embedAndPersist is @Transactional and called from RetrievalIndexService "
        "(different bean) after commit — proxy applies correctly. Unit tests without "
        "sync fall back to immediate embed. Solid.",
    )

    heading(doc, "What looks solid", 2)
    bullets(
        doc,
        [
            "FTS type filter before LIMIT; SchemaPatchRunner / Flyway parity",
            "Titan v1 locked to EMBEDDING_DIMENSION=1536; v2 rejected at startup",
            "afterCommit embedding; missing-embedding retry on contentHash match",
            "Optional Bedrock client; embedding.enabled=false in tests",
            "Claim lease + SKIP LOCKED multi-ECS design; lease default 10 minutes",
            "Unit/contract coverage for FTS, embed format, IndexWorker paths",
        ],
    )

    heading(doc, "3. Architecture & Style", 1)
    bullets(
        doc,
        [
            "Clean layering: IndexWorker → RetrievalIndexService → "
            "ChunkEmbeddingService / FullTextSearchService",
            "Native SQL kept off the JPA entity for vector/FTS columns — consistent "
            "with portable RetrievalIndexChunk mapping",
            "PR is large (4.2 + lease + 4.3) but one coherent index write-path story",
            "Rebase onto current team-ae-develop before merge (#305 landed after branch)",
        ],
    )

    heading(doc, "4. Recommendations", 1)

    heading(doc, "A. Commit the working-tree lease fixes before PR", 2)
    code(
        doc,
        "// Already present in working tree — must be on the branch tip:\n"
        "// IndexingOutboxRepository\n"
        "OR claimed_at < (NOW() - make_interval(mins => :leaseMinutes))\n"
        "\n"
        "// IndexWorker.releaseClaimWithoutBurn\n"
        "if (row.getClaimedAt() == null) {\n"
        "    row.setClaimedAt(LocalDateTime.now());\n"
        "}\n"
        "// do NOT null claimed_at\n"
        "row.setLastError(truncate(message));\n"
        "outboxRepository.save(row);",
    )

    heading(doc, "B. Optional: longer park for no-burn deferrals", 2)
    code(
        doc,
        "// Park visits for hours, not just claim-lease-minutes:\n"
        "private void releaseClaimWithoutBurn(final IndexingOutboxRow row, final String message) {\n"
        "    row.setClaimedAt(LocalDateTime.now().plusHours(6)); // or config\n"
        "    row.setLastError(truncate(message));\n"
        "    outboxRepository.save(row);\n"
        "}\n"
        "// Claim query already skips claimed_at >= NOW() - lease;\n"
        "// future claimed_at stays unclaimable until that timestamp ages past lease window.\n"
        "// BETTER: claimed_at < NOW() - lease  OR  use parked_until column for clarity.",
    )
    para(
        doc,
        "Note: claiming uses “claimed_at older than lease”, so setting claimed_at into "
        "the future parks indefinitely until NOW() - lease catches up… actually "
        "future claimed_at is NEVER < NOW()-lease, so it parks until claimed_at "
        "itself ages past NOW()-lease, i.e. until clock reaches claimed_at+lease. "
        "Setting claimed_at = now()+6h parks until ~6h+lease. Document carefully "
        "or add parked_until.",
    )

    heading(doc, "C. Soft lease on burn-deferrals (optional)", 2)
    code(
        doc,
        "private void recordDeferOrDeadLetter(...) {\n"
        "    row.setAttemptCount(attempts + 1);\n"
        "    // Keep/refresh lease so bad rows do not reclaim every 15s\n"
        "    row.setClaimedAt(LocalDateTime.now());\n"
        "    ...\n"
        "}",
    )

    heading(doc, "D. Before merge checklist", 2)
    bullets(
        doc,
        [
            "Commit park + make_interval + tests + optional smoke script",
            "Rebase/merge origin/team-ae-develop",
            "Redeploy 03-platform.yaml for Titan v1 IAM before enabling embeds in ECS",
            "Smoke: visit SUMMARY_CREATED parks with non-null claimed_at; next poll skips; "
            "claim query with make_interval returns expected rows",
            "Smoke: ingest call summary → FTS hit + embedding non-null (or NULL if Bedrock down)",
        ],
    )

    heading(doc, "Verdict", 1)
    para(
        doc,
        "Approve with changes. The FTS + Titan embedding design is sound "
        "(after-commit embed, v1 dimension lock, best-effort failure). Blocking items "
        "for merge are committing the visit-park lease keep + make_interval claim SQL "
        "already in the working tree, then rebasing onto current team-ae-develop.",
        bold=True,
    )

    return doc


def main() -> None:
    doc = build()
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    main()
