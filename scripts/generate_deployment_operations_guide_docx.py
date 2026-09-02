"""Generate the Fall 2026 CareConnect Deployment and Operations Guide as a Word document.

Content reflects the deployment performed and verified on 2026-09-02 against AWS
the cohort AWS account (environment "cfdemo"), not the prior cohort's document.
"""

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

NAVY = RGBColor(0x1F, 0x3A, 0x5F)
HDR_FILL = "1F3A5F"


def set_cell_shading(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def h(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = NAVY
    return heading


def p(doc, text, bold=False, italic=False, size=11):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return para


def bullets(doc, items):
    for item in items:
        para = doc.add_paragraph(str(item), style="List Bullet")
        for run in para.runs:
            run.font.size = Pt(11)


def numbered(doc, items):
    for item in items:
        para = doc.add_paragraph(str(item), style="List Number")
        for run in para.runs:
            run.font.size = Pt(11)


def code(doc, text):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    para.paragraph_format.left_indent = Inches(0.25)
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)
    return para


def table(doc, headers, rows, caption=None, widths=None):
    if caption:
        cap = doc.add_paragraph()
        run = cap.add_run(caption)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = NAVY
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    for i, head in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(head)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(hdr[i], HDR_FILL)
    for r, row in enumerate(rows, start=1):
        for c, val in enumerate(row):
            cell = tbl.rows[r].cells[c]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(9.5)
    if widths:
        for r in tbl.rows:
            for i, w in enumerate(widths):
                r.cells[i].width = Inches(w)
    doc.add_paragraph()
    return tbl


def note(doc, label, text):
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.25)
    lab = para.add_run(f"{label}: ")
    lab.bold = True
    lab.font.size = Pt(10.5)
    lab.font.color.rgb = RGBColor(0xA0, 0x40, 0x00)
    body = para.add_run(text)
    body.font.size = Pt(10.5)
    return para


def build():
    doc = Document()
    for style_name, size in (("Normal", 11),):
        doc.styles[style_name].font.size = Pt(size)
        doc.styles[style_name].font.name = "Calibri"

    # ---------------- Title page ----------------
    for _ in range(4):
        doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("CareConnect")
    r.bold = True
    r.font.size = Pt(34)
    r.font.color.rgb = NAVY

    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = s.add_run("Deployment and Operations Guide")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = NAVY

    for line, size, bold in (
        ("Version 3.0", 14, True),
        ("2026 Fall Cohort", 13, False),
        ("University of Maryland Global Campus", 13, False),
        ("SWEN 670 9040 Software Engineering Capstone", 13, False),
        ("Professor Minagar", 13, False),
        ("September 2, 2026", 13, False),
    ):
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(line)
        run.bold = bold
        run.font.size = Pt(size)

    doc.add_page_break()

    # ---------------- Revision history ----------------
    h(doc, "Revision History", 1)
    table(
        doc,
        ["Editor", "Date", "Reason for Change", "Version"],
        [
            ["2026 Summer Cohort", "7/25/2026", "Initial document", "1.0"],
            ["2026 Summer Cohort", "8/4/2026", "Summer cohort final revision", "2.0"],
            [
                "2026 Fall Cohort",
                "9/2/2026",
                "Fall cohort revision. Rewritten against a deployment performed and "
                "verified on the cohort AWS account. Corrects PostgreSQL version, "
                "CI/CD pipeline description, coverage gate mechanism, repository URL, "
                "and Spring profile guidance. Adds the Amplify SPA rewrite requirement.",
                "3.0",
            ],
        ],
        widths=[1.5, 1.0, 3.6, 0.7],
    )

    h(doc, "Project Team", 1)
    p(
        doc,
        "This revision was produced by the Fall 2026 cohort. The project manager and "
        "team leads listed below are the designated code owners for the repository.",
    )
    table(
        doc,
        ["Name", "Role", "GitHub Handle"],
        [
            ["Crystal Schelmety", "Project Manager", "@cschel"],
            ["Beth Kraus", "Team Lead", "@b-kraus"],
            ["Ian Ard", "Team Lead", "@ianscottard"],
            ["Jonathan Barreto", "Team Lead", "@jonathanbarreto37"],
            ["Shayne McPherson", "Team Lead", "@shaynemcp"],
            ["Sivakumar Ganesan", "Team Lead", "@sganesan0414"],
            ["Sung Yoo", "Team Lead", "@SungWook1207"],
        ],
        widths=[2.2, 1.6, 2.4],
    )

    doc.add_page_break()

    # ---------------- 1. Introduction ----------------
    h(doc, "1. Introduction", 1)

    h(doc, "1.1 Purpose", 2)
    p(
        doc,
        "This guide is the practical reference for setting up, building, configuring, and "
        "deploying the CareConnect platform. Where the Technical Design Document captures "
        "architectural intent and the Software Test Plan captures verification strategy, "
        "this guide captures the operational knowledge needed to actually run the system: "
        "source control practices, local environment setup, and deployment to AWS.",
    )
    p(
        doc,
        "Every command, version number, and measured timing in Section 3.5 comes from a full "
        "deployment performed on September 2, 2026 against the cohort AWS account in "
        "us-east-1, using environment name cfdemo. Where a step could not be verified, that "
        "is stated explicitly rather than assumed.",
    )

    h(doc, "1.2 Intended Audience", 2)
    p(
        doc,
        "This guide is aimed at the operators responsible for standing up, configuring, and "
        "sustaining CareConnect environments: system administrators, DevOps and platform "
        "engineers, and developers performing a release. Readers are expected to be "
        "comfortable with source control, command line tooling, and cloud infrastructure "
        "concepts. Product feature design and application programming detail belong to the "
        "Technical Design Document and the Programmers Guide. End users should consult the "
        "Users Guide.",
    )

    h(doc, "1.3 Document Organization", 2)
    p(
        doc,
        "Section 2 covers source control: repository access, structure, pull request rules, "
        "collaboration, and the CI/CD pipelines. Section 3 covers setup: system requirements, "
        "software installation, configuration, compilation, and deployment to AWS. Section 4 "
        "covers troubleshooting. Section 5 is the AI usage statement.",
    )

    h(doc, "1.4 Corrections to the Prior Revision", 2)
    p(
        doc,
        "Version 2.0 was written by the Summer 2026 cohort. Several of its statements no "
        "longer describe this system. They are corrected here so operators do not follow "
        "stale instructions.",
    )
    table(
        doc,
        ["Topic", "Version 2.0 stated", "Fall 2026 actual"],
        [
            [
                "Repository",
                "github.com/umgc/2026_summer_careconnect",
                "github.com/umgc/2026_fall_careconnect",
            ],
            [
                "PostgreSQL",
                "15, with the pgvector/pgvector:pg15 image",
                "17.6 on RDS (EngineVersion in 02-data.yaml)",
            ],
            [
                "CI pipeline",
                "One pipeline with SonarCloud, coverage gate, and axe-core",
                "That pipeline is team-b-ci.yml. The repository-wide gate is "
                "build-and-analyze.yml, a security and quality scan running on Java 21",
            ],
            [
                "Coverage gate",
                "Fixed tiers: EVV and Shift Scheduling at 100%, auth at 95%",
                "A ratcheting no-regression baseline in scripts/coverage-baseline.json. "
                "EVV currently sits at 52.2%",
            ],
            [
                "Spring profile for deployment",
                "Section 7 example used SpringProfile=prod",
                "cfdemo deploys with SpringProfile=dev. The prod profile requires SSM "
                "parameters and a SendGrid key that are not provisioned",
            ],
            [
                "Backend source path",
                "careconnect-backend / careconnect-frontend",
                "backend/core and frontend",
            ],
            [
                "Amplify manual zip deploy",
                "Not mentioned",
                "Requires a SPA rewrite rule or every route except / returns 404",
            ],
        ],
        caption="Table 1. Corrections to version 2.0",
        widths=[1.4, 2.5, 2.9],
    )

    h(doc, "1.5 Definitions, Abbreviations, and Acronyms", 2)
    table(
        doc,
        ["Term", "Definition"],
        [
            ["Amplify", "AWS service hosting the CareConnect Flutter web frontend."],
            ["API Gateway", "AWS HTTP API that routes browser traffic to the backend through a VPC Link."],
            ["ARN", "Amazon Resource Name, the unique identifier format AWS uses for resources."],
            ["CloudFormation", "AWS infrastructure as code service used to provision the four CareConnect stacks."],
            ["Cloud Map", "AWS service discovery registry that maps the ECS task to a private DNS name."],
            ["CODEOWNER", "A designated reviewer whose approval is required before a pull request can merge."],
            ["CORS", "Cross-Origin Resource Sharing, the browser mechanism controlled by CORS_ALLOWED_LIST."],
            ["dart-define", "A Flutter build-time constant. BACKEND_URL is supplied this way, never hardcoded."],
            ["ECR", "Elastic Container Registry, which stores the backend container image."],
            ["ECS Fargate", "AWS serverless container hosting that runs the Spring Boot backend."],
            ["Fat jar", "The self-contained Spring Boot archive produced by the docker Maven profile."],
            ["IGW", "Internet Gateway. CareConnect uses public subnets with an IGW and no NAT Gateway."],
            ["JDBC", "Java Database Connectivity, the API the backend uses to reach PostgreSQL."],
            ["JWT", "JSON Web Token, the authentication credential validated via SECURITY_JWT_SECRET."],
            ["OIDC", "OpenID Connect, used by GitHub Actions to assume an AWS role without long-lived keys."],
            ["RDS", "Amazon Relational Database Service, the managed PostgreSQL host."],
            ["SPA rewrite", "An Amplify custom rule mapping unmatched paths to /index.html so client-side routes resolve."],
            ["SSM", "AWS Systems Manager Parameter Store, read by the prod Spring profile."],
            ["VPC Link", "The API Gateway component that reaches private resources inside the VPC."],
        ],
        caption="Table 2. Definitions and acronyms",
        widths=[1.5, 5.3],
    )

    doc.add_page_break()

    # ---------------- 2. Source Control ----------------
    h(doc, "2. Source Control", 1)

    h(doc, "2.1 GitHub Access", 2)
    p(doc, "The CareConnect repository is hosted under the UMGC organization:")
    code(doc, "https://github.com/umgc/2026_fall_careconnect")
    p(
        doc,
        "Access is administered by the UMGC Capstone Course organization. Contributors who "
        "cannot reach the repository should contact a team lead or the project manager, "
        "Crystal Schelmety, who can escalate to the professor. Repository administration "
        "rights, including branch protection settings, sit with the course staff rather than "
        "with the student team.",
    )

    h(doc, "2.2 Repository Structure", 2)
    table(
        doc,
        ["Path", "Contents"],
        [
            ["backend/core", "Spring Boot backend. Maven wrapper, Dockerfile, application properties."],
            ["frontend", "Flutter application for web, Android, and iOS. Contains amplify.yml."],
            ["ui-integration/mobile-app", "React and Vite UI preview, built into the Flutter web output."],
            ["cloudformation-fargate", "CloudFormation templates, parameter files, and deploy or destroy scripts."],
            [".github/workflows", "GitHub Actions pipelines."],
            ["scripts", "Repository tooling, including the coverage gate and document generators."],
            ["docs", "Project deliverables and guides."],
        ],
        widths=[2.2, 4.6],
    )
    p(doc, "Branch naming follows this scheme:", bold=True)
    table(
        doc,
        ["Branch pattern", "Description"],
        [
            ["main", "Production branch for stable releases and approved hotfixes."],
            ["develop", "Primary integration branch for ongoing feature work."],
            ["team-<team>-develop", "Team-level integration branches staged before syncing to develop."],
            ["feature/*", "Feature branches created from develop or a team develop branch."],
            ["fix/*", "Focused bug-fix branches following the same merge flow."],
            ["chore/*, docs/*", "Maintenance and documentation branches."],
        ],
        widths=[2.0, 4.8],
    )

    h(doc, "2.3 Pull Request Guidelines", 2)
    p(doc, "Approval and merge rules for the Fall 2026 cohort:")
    bullets(
        doc,
        [
            "Pull requests require two approving reviews before merge.",
            "Self-merges are not permitted. The author of a pull request may not merge it.",
            "The branches main, develop, and the team develop branches are protected.",
            "Code owner approval is governed by .github/CODEOWNERS, listing the project "
            "manager and the six team leads.",
        ],
    )
    note(
        doc,
        "Verification note",
        "The protected status of main, develop, and team-e-develop was confirmed directly "
        "through the GitHub API. The detailed protection settings endpoint requires "
        "repository admin rights, which the student team does not hold, so the approval "
        "count and self-merge restriction above are recorded from the team's own "
        "configuration rather than read back from the API.",
    )

    h(doc, "2.4 Collaboration Methods", 2)
    bullets(
        doc,
        [
            "Microsoft Teams for day-to-day messaging, team meetings, and leads meetings.",
            "Email for formal communications.",
            "GitHub for issue tracking, pull requests, and code review.",
        ],
    )

    h(doc, "2.5 CI/CD Pipeline", 2)
    p(
        doc,
        "The repository defines four GitHub Actions workflows. They serve different purposes, "
        "and conflating them is the most common misreading of the prior revision.",
    )
    table(
        doc,
        ["Workflow", "Purpose", "Notable details"],
        [
            [
                "build-and-analyze.yml",
                "Repository-wide security and quality scan",
                "Runs on Java 21 (Temurin). TruffleHog and Gitleaks for secrets, Flutter "
                "Analyze, Checkstyle, PMD, SpotBugs, Semgrep, Pylint, Bandit, HTMLHint, "
                "Stylelint, Trivy filesystem SCA, and OWASP Dependency-Check. Publishes a "
                "quality report artifact.",
            ],
            [
                "team-b-ci.yml",
                "Full build, test, coverage, and accessibility pipeline",
                "Maven and Flutter build, unit tests, JaCoCo and LCOV coverage, SonarCloud "
                "scans for backend and Dart, the coverage gate, Android emulator end-to-end "
                "tests, and an axe-core WCAG 2.1 AA scan.",
            ],
            [
                "backend-app-deploy.yml",
                "Application-only backend redeploy",
                "Rebuilds and redeploys the image against existing infrastructure. Triggered "
                "on push and by workflow_dispatch with environment and image tag inputs.",
            ],
            [
                "backend-full-deploy.yml",
                "Full four-stack deploy",
                "workflow_dispatch only. Inputs for environment, image tag, and whether to "
                "run Maven tests.",
            ],
        ],
        caption="Table 3. GitHub Actions workflows",
        widths=[1.5, 1.8, 3.5],
    )
    note(
        doc,
        "Important",
        "team-b-ci.yml is named for a single team. A branch named "
        "chore/b-generalize-ci-all-teams exists, indicating generalization is in progress but "
        "not yet merged. Until it lands, the repository-wide gate on every pull request is "
        "build-and-analyze.yml, and the coverage and accessibility gates are not applied "
        "uniformly across all teams.",
    )

    p(doc, "Coverage gate", bold=True)
    p(
        doc,
        "The gate at scripts/coverage-gate.sh reads scripts/coverage-baseline.json. That file "
        "is an auto-generated no-regression baseline whose values, in its own words, only "
        "ratchet up. It is not a set of fixed targets. Regenerate with "
        "GATE_MODE=baseline scripts/coverage-gate.sh. Current frontend line baselines:",
    )
    table(
        doc,
        ["Module", "Line baseline"],
        [
            ["lib/features/communication", "100.0"],
            ["lib/features/auth", "95.2"],
            ["lib/features/shift_scheduling", "86.4"],
            ["lib/features/ai", "82.6"],
            ["lib/widgets", "60.8"],
            ["lib/features/evv", "52.2"],
            ["backend", "empty, no baseline recorded"],
        ],
        widths=[3.4, 3.4],
    )

    doc.add_page_break()

    # ---------------- 3. CareConnect Setup ----------------
    h(doc, "3. CareConnect Setup", 1)

    h(doc, "3.1 System Requirements", 2)
    p(
        doc,
        "CareConnect is a full stack application: a Flutter frontend, a Java Spring Boot "
        "backend, a PostgreSQL database, and containerized deployment to AWS. Building the "
        "complete system locally is resource intensive.",
    )
    p(doc, "Hardware and operating system", bold=True)
    bullets(
        doc,
        [
            "A 64-bit multi-core machine. 8 GB RAM is the minimum; 16 GB or more is strongly "
            "recommended when running an emulator alongside the backend and database.",
            "At least 20 GB free disk. The backend fat jar alone is 227 MB and the container "
            "image is 634 MB compressed.",
            "Hardware virtualization enabled in firmware for Docker and the Android emulator.",
            "Windows 10 or 11, macOS on Apple Silicon or Intel, or a mainstream 64-bit Linux "
            "distribution.",
            "Apple Silicon machines must build the backend image for linux/amd64. The deploy "
            "script already passes --platform linux/amd64.",
        ],
    )
    table(
        doc,
        ["Component", "Required", "Verified during the 2026-09-02 deploy"],
        [
            ["JDK", "17 (LTS)", "OpenJDK 17.0.20 (Homebrew)"],
            ["Maven", "Wrapper provided", "maven-3.9.11 via mvnw"],
            ["Spring Boot", "3.4.5 (managed)", "3.4.5"],
            ["Flutter SDK", "3.x stable", "3.41.0"],
            ["Dart SDK", ">=3.5.0 <4.0.0", "Bundled with Flutter"],
            ["PostgreSQL", "17.x on RDS", "17.6, db.t4g.micro"],
            ["Docker Desktop", "Latest stable", "29.7.2, 8 CPU, 4 GB"],
            ["AWS CLI", "v2", "2.36.24"],
            ["Node.js", "20 LTS or newer", "22.14.0, npm 11.19.0"],
            ["Git", "Latest stable", "Verified"],
        ],
        caption="Table 4. Software requirements",
        widths=[1.5, 1.9, 3.4],
    )
    note(
        doc,
        "JDK caution",
        "The backend targets Java 17 and the Dockerfile runtime is eclipse-temurin:17-jre-jammy. "
        "If your shell defaults to a newer JDK, pin JAVA_HOME for the build. On the machine used "
        "for this deploy the default was JDK 23 and JAVA_HOME had to be set explicitly.",
    )

    h(doc, "3.2 Software Installation", 2)
    p(doc, "Install in this order and verify each tool before continuing.")
    table(
        doc,
        ["Tool", "Verify with", "Why it is needed"],
        [
            ["Git", "git --version", "Clone the repository and manage branches."],
            ["JDK 17", "java -version", "Compile and run the Spring Boot backend."],
            ["Flutter", "flutter --version", "Build the frontend. Enable web with flutter config --enable-web."],
            ["Docker Desktop", "docker info", "Build the backend image. The daemon must be running, not just installed."],
            ["AWS CLI v2", "aws --version", "Authenticate and run deployment scripts."],
            ["Node.js", "node --version", "Build the React UI preview and run axe-core tooling."],
            ["Android Studio", "flutter doctor", "Android SDK and emulator, for the mobile build only."],
        ],
        widths=[1.3, 1.7, 3.8],
    )
    p(doc, "AWS profile setup", bold=True)
    p(
        doc,
        "The Fall 2026 account uses AWS IAM Identity Center, so the profile is configured as an "
        "SSO profile rather than with static access keys. The prior revision described this "
        "profile name as a local label for IAM keys, which is no longer accurate.",
    )
    code(
        doc,
        "[profile careconnect-sso]\n"
        "sso_session = careconnect\n"
        "sso_account_id = <your-account-id>\n"
        "sso_role_name = StudentAdminAccess\n"
        "region = us-east-1\n"
        "output = json\n\n"
        "[sso-session careconnect]\n"
        "sso_region = us-east-1\n"
        "sso_registration_scopes = sso:account:access",
    )
    p(doc, "Log in at the start of every working session. The token expires.")
    code(
        doc,
        "export AWS_PROFILE=careconnect-sso\n"
        "aws sso login --profile careconnect-sso\n"
        "aws sts get-caller-identity --profile careconnect-sso",
    )

    h(doc, "3.3 Program Configuration", 2)
    p(
        doc,
        "No environment URL is hardcoded in application source. The frontend receives "
        "BACKEND_URL as a build-time dart-define, and the backend receives its configuration "
        "as ECS environment variables and injected secrets defined by CloudFormation.",
    )
    p(doc, "3.3.1 Frontend build-time configuration", bold=True)
    table(
        doc,
        ["dart-define", "Value used", "Effect"],
        [
            ["BACKEND_URL", "The API Gateway base URL, no trailing slash and no /v1", "Base URL for all API calls."],
            ["APP_DOMAIN", "The Amplify hostname only, without https://", "Used for deep links and OAuth callbacks."],
            ["APP_PORT", "443", "Port component for generated URLs."],
        ],
        widths=[1.4, 2.6, 2.8],
    )
    note(
        doc,
        "Failure mode",
        "If BACKEND_URL is absent at build time the Flutter app falls back to "
        "http://localhost:8080 and the welcome page reports the backend as unhealthy. "
        "Section 3.4.3 gives a way to confirm the value actually reached the bundle.",
    )
    p(doc, "3.3.2 Backend runtime configuration", bold=True)
    p(doc, "04-service.yaml injects these into the ECS task definition:")
    table(
        doc,
        ["Variable", "Source"],
        [
            ["SPRING_PROFILES_ACTIVE", "SpringProfile stack parameter"],
            ["JDBC_URI, DB_USER", "Outputs from the data stack"],
            ["DB_PASSWORD, SECURITY_JWT_SECRET", "Injected from Secrets Manager, not plain environment values"],
            ["APP_FRONTEND_BASE_URL", "FrontendBaseUrl stack parameter"],
            ["CORS_ALLOWED_LIST", "CorsAllowedList stack parameter"],
            ["CARECONNECT_AI_ENABLED, AI_PROVIDER", "Stack parameters, bedrock by default"],
            ["WEBSOCKET_ENABLED, AWS_WEBSOCKET_API_GATEWAY_ENDPOINT", "Currently false and empty. See Section 4.5"],
            ["SERVER_PORT, ENVIRONMENT", "8081 and the environment name"],
        ],
        widths=[3.0, 3.8],
    )
    p(doc, "3.3.3 Choosing a Spring profile", bold=True)
    p(
        doc,
        "The cfdemo parameter file sets SpringProfile=dev. This is deliberate and differs from "
        "the prior revision's example, which used prod.",
    )
    bullets(
        doc,
        [
            "The prod profile registers SsmPropertySourceInitializer and requires "
            "SENDGRID_API_KEY and FROM_EMAIL. Neither the SSM parameters nor a SendGrid key "
            "exist in the Fall 2026 account, so prod is expected to fail at startup.",
            "The dev profile resolves the datasource from ${JDBC_URI:...}, and the service "
            "template injects JDBC_URI, so dev connects to RDS correctly.",
            "The dev profile also runs DevDataLoader, which seeds demo accounts. See the "
            "warning in Section 3.5.6.",
        ],
    )

    h(doc, "3.4 Compiling CareConnect", 2)

    p(doc, "3.4.1 Backend", bold=True)
    p(
        doc,
        "The Dockerfile is a runtime-only stage. It copies a jar that must already exist, so "
        "the jar is built on the host and the docker Maven profile is mandatory. The default "
        "profile does not produce the archive the Dockerfile expects.",
    )
    code(doc, 'COPY target/careconnect-backend-0.0.1-SNAPSHOT.jar app.jar')
    code(
        doc,
        "export JAVA_HOME=/path/to/jdk-17\n"
        "export PATH=\"$JAVA_HOME/bin:$PATH\"\n"
        "cd backend/core\n"
        "./mvnw -B -ntp clean package -Pdocker -DskipTests",
    )
    p(
        doc,
        "Measured result on 2026-09-02: BUILD SUCCESS in 38.3 seconds, producing a 227 MB fat "
        "jar at backend/core/target/careconnect-backend-0.0.1-SNAPSHOT.jar. Warnings about a "
        "deprecated API in WebClientConfig, unchecked operations in CallRecordingService, and "
        "sun.misc.Unsafe in a Bedrock test are expected and non-blocking.",
    )

    p(doc, "3.4.2 Frontend", bold=True)
    p(
        doc,
        "frontend/amplify.yml performs two build steps, not one. It builds the React UI "
        "preview with Vite and copies the output into web/ui-preview, then builds the Flutter "
        "web app. A manual build that skips the first step ships an app whose /ui-preview "
        "route is missing.",
    )
    code(
        doc,
        "cd ui-integration/mobile-app\n"
        "npm ci && npm run build\n\n"
        "cd ../../frontend\n"
        "rm -rf web/ui-preview && mkdir -p web/ui-preview\n"
        "cp -R ../ui-integration/mobile-app/dist/. web/ui-preview/\n\n"
        "flutter pub get\n"
        "flutter build web --release --base-href \"/\" \\\n"
        "  --dart-define=BACKEND_URL=\"$BACKEND_URL\" \\\n"
        "  --dart-define=APP_DOMAIN=\"$FRONTEND_HOST\" \\\n"
        "  --dart-define=APP_PORT=443",
    )
    p(
        doc,
        "Measured result: the Vite build completed in 1.36 seconds over 1604 modules, and the "
        "Flutter web build completed in 83.2 seconds producing 60 MB in frontend/build/web. "
        "Icon tree-shaking reduced MaterialIcons by 96.6 percent and CupertinoIcons by 99.0 "
        "percent, so the --no-tree-shake-icons workaround named in the prior revision was not "
        "needed.",
    )

    p(doc, "3.4.3 Confirming the build-time configuration landed", bold=True)
    p(
        doc,
        "A missing BACKEND_URL produces a silently wrong build. Confirm it is present in the "
        "compiled bundle before deploying:",
    )
    code(doc, 'grep -c "<your-api-id>.execute-api" frontend/build/web/main.dart.js')
    p(
        doc,
        "A count of at least one confirms the value was compiled in. Occurrences of "
        "localhost:8080 will still appear; that is the source-level fallback literal and is "
        "overridden at runtime by the dart-define.",
    )

    h(doc, "3.5 Deployment to AWS", 2)
    p(doc, "Request path and stack order:")
    code(doc, "Browser -> Amplify -> API Gateway -> VPC Link -> Cloud Map -> ECS Fargate -> RDS")
    table(
        doc,
        ["Order", "Stack", "Template", "Owns"],
        [
            ["1", "careconnect-networking-<env>", "01-networking.yaml", "VPC, subnets, IGW, route tables, security groups"],
            ["2", "careconnect-data-<env>", "02-data.yaml", "RDS PostgreSQL 17.6, Secrets Manager entries, parameter group"],
            ["3", "careconnect-platform-<env>", "03-platform.yaml", "ECR repository, shared IAM, log group"],
            ["4", "careconnect-service-<env>", "04-service.yaml", "ECS task and service, Cloud Map, VPC Link, HTTP API"],
        ],
        caption="Table 5. Stack order",
        widths=[0.6, 2.0, 1.5, 2.7],
    )
    note(
        doc,
        "Cost",
        "This environment creates billable resources. The networking template provisions "
        "public subnets with an Internet Gateway and no NAT Gateway, which avoids the usual "
        "largest line item. Remaining recurring cost is roughly 1.90 US dollars per day: "
        "about 0.45 for RDS db.t4g.micro with 20 GB gp3, about 1.18 for one Fargate task at "
        "1024 CPU and 3072 MiB, and about 0.24 for API Gateway and the VPC Link. Left running "
        "for a month this consumes most of the 75 dollar careconnect-monthly-cost budget. Tear "
        "down when finished.",
    )

    p(doc, "3.5.1 Prerequisites", bold=True)
    numbered(
        doc,
        [
            "AWS CLI v2, JDK 17, and Docker Desktop installed, with the Docker daemon running.",
            "Flutter installed if publishing the frontend.",
            "An active SSO session: aws sso login --profile careconnect-sso, then confirm with "
            "aws sts get-caller-identity.",
            "A database master password of at least 8 characters and a JWT secret of at least "
            "32 characters. Do not commit these.",
        ],
    )
    p(
        doc,
        "The checked-in parameter file carries placeholders (REPLACE_ME_DB_PASSWORD and "
        "REPLACE_ME_WITH_A_LONG_RANDOM_JWT_SECRET_VALUE). The deploy script refuses to proceed "
        "if they survive, so supply real values through environment variables. Generating them "
        "and storing them outside the repository with restrictive permissions works well:",
    )
    code(
        doc,
        "umask 077\n"
        "cat > ~/careconnect-cfdemo-secrets.env <<'EOF'\n"
        "export CARECONNECT_DATABASE_MASTER_PASSWORD='<generated>'\n"
        "export CARECONNECT_JWT_SECRET='<generated, 32+ chars>'\n"
        "EOF\n"
        "chmod 600 ~/careconnect-cfdemo-secrets.env",
    )

    p(doc, "3.5.2 Backend deploy", bold=True)
    code(
        doc,
        "source ~/careconnect-cfdemo-secrets.env\n"
        "export AWS_PROFILE=careconnect-sso\n"
        "export JAVA_HOME=/path/to/jdk-17\n"
        "export PATH=\"$JAVA_HOME/bin:$PATH\"\n\n"
        "./cloudformation-fargate/cdeploy_cloudformation.sh \\\n"
        "  --environment cfdemo --profile careconnect-sso",
    )
    p(
        doc,
        "The script deploys the four stacks in order, reads the ECR repository URI from the "
        "platform stack, builds the jar with the docker profile, builds the image with "
        "--platform linux/amd64, pushes it, and finally deploys the service stack with the "
        "resulting image URI. It prints the API endpoint on completion.",
    )
    p(doc, "Measured result on 2026-09-02:")
    code(
        doc,
        "Deployment complete.\n"
        "Environment:   cfdemo\n"
        "Repository:    careconnect-backend-cfdemo\n"
        "Image URI:     <account-id>.dkr.ecr.us-east-1.amazonaws.com/careconnect-backend-cfdemo:cfdemo\n"
        "API Endpoint:  https://<api-id>.execute-api.us-east-1.amazonaws.com\n"
        "Elapsed time:  00:19:54",
    )
    p(
        doc,
        "Record the API endpoint as BACKEND_URL. It has no trailing slash and no /v1 suffix. "
        "Expect the image build to take several minutes on Apple Silicon: the Dockerfile "
        "installs ffmpeg for speaker identification, which pulls 176 packages and about 113 MB "
        "of archives under emulation. This is normal, not a hang.",
    )
    p(doc, "For later code-only updates that do not touch networking or RDS:")
    code(doc, "./cloudformation-fargate/cdeploy_app_only.sh --environment cfdemo --profile careconnect-sso")

    p(doc, "3.5.3 Waiting for the backend to become reachable", bold=True)
    p(
        doc,
        "Immediately after the service stack completes, the endpoint returns HTTP 503 for a "
        "period even though ECS reports the task as running and Cloud Map reports the instance "
        "as healthy. During this deploy the endpoint alternated between 200 and 503 at about "
        "nine seconds, then settled. Roughly ninety seconds later six consecutive requests all "
        "returned 200 in under 250 milliseconds. Wait two to five minutes before concluding "
        "something is wrong.",
    )
    code(doc, "curl -s \"$BACKEND_URL/v1/api/test/health\"")

    p(doc, "3.5.4 Frontend deploy to Amplify", bold=True)
    p(
        doc,
        "Two options exist. Option A connects the repository so Amplify runs amplify.yml on "
        "each push. Option B uploads a locally built zip. Option B was used and verified for "
        "this revision, and it has one requirement the prior guide omits entirely.",
    )
    p(doc, "Create the app and branch first, because APP_DOMAIN is a build-time value:", bold=True)
    code(
        doc,
        "aws amplify create-app --name careconnect-fall2026 --platform WEB\n"
        "aws amplify create-branch --app-id <appId> --branch-name main",
    )
    p(
        doc,
        "The hostname is <branch>.<appId>.amplifyapp.com. Build the frontend using that value "
        "for APP_DOMAIN and the API endpoint for BACKEND_URL, per Section 3.4.2, then package "
        "and deploy. Zip the contents of build/web, not the folder itself:",
    )
    code(
        doc,
        "cd frontend/build/web\n"
        "zip -r -q careconnect-web.zip . -x \"*.DS_Store\"\n\n"
        "aws amplify create-deployment --app-id <appId> --branch-name main\n"
        "curl -X PUT -T careconnect-web.zip -H \"Content-Type: application/zip\" \"<zipUploadUrl>\"\n"
        "aws amplify start-deployment --app-id <appId> --branch-name main --job-id <jobId>",
    )
    p(doc, "The 20 MB upload and deployment completed in a few seconds during this deploy.")

    p(doc, "3.5.5 Required: add the SPA rewrite rule", bold=True)
    p(
        doc,
        "A Git-connected Amplify app normally receives a single-page-application rewrite "
        "through framework detection. A manual zip deploy does not. Without the rule the site "
        "root loads but every client-side route returns 404. This was observed directly: / "
        "returned 200 while /login returned a 301 followed by a 404, and "
        "aws amplify get-app returned an empty customRules list.",
    )
    code(
        doc,
        "aws amplify update-app --app-id <appId> --custom-rules '[\n"
        "  {\n"
        "    \"source\": \"</^[^.]+$|\\\\.(?!(css|gif|ico|jpg|js|png|txt|svg|woff|woff2|ttf|map|json|webp|otf|wasm|bin|html)$)([^.]+$)/>\",\n"
        "    \"target\": \"/index.html\",\n"
        "    \"status\": \"200\"\n"
        "  }\n"
        "]'",
    )
    p(
        doc,
        "After the rule propagates, confirm that a deep route returns 200 and that a static "
        "asset is still served rather than being rewritten:",
    )
    code(
        doc,
        "curl -o /dev/null -w \"%{http_code}\\n\" \"$FRONTEND_URL/login\"        # expect 200\n"
        "curl -o /dev/null -w \"%{http_code}\\n\" \"$FRONTEND_URL/main.dart.js\"  # expect 200",
    )

    p(doc, "3.5.6 Point the backend at the Amplify origin", bold=True)
    p(
        doc,
        "The backend is deployed before the Amplify URL exists, so its CORS allow-list and "
        "frontend base URL must be updated afterward. Capture the existing stack parameters "
        "first so unspecified parameters are not reset to template defaults, then change only "
        "the two that need to change. Hold SpringProfile at dev and keep the current image URI.",
    )
    code(
        doc,
        "aws cloudformation deploy \\\n"
        "  --stack-name careconnect-service-cfdemo \\\n"
        "  --template-file cloudformation-fargate/templates/04-service.yaml \\\n"
        "  --capabilities CAPABILITY_NAMED_IAM --no-fail-on-empty-changeset \\\n"
        "  --parameter-overrides \\\n"
        "    Environment=cfdemo BackendImageUri=<current image URI> SpringProfile=dev \\\n"
        "    CareConnectAiEnabled=true AiProvider=bedrock FromEmail=noreply@careconnect.com \\\n"
        "    WebSocketApiGatewayEndpoint= \\\n"
        "    FrontendBaseUrl=\"$FRONTEND_URL\" \\\n"
        "    \"CorsAllowedList=http://localhost:*,http://127.0.0.1:*,$FRONTEND_URL\" \\\n"
        "    ContainerPort=8081 DesiredCount=1 TaskCpu=1024 TaskMemory=3072 \\\n"
        "    HealthCheckPath=/v1/api/test/health DomainName= HostedZoneId=",
    )
    p(
        doc,
        "The stack update replaces the ECS task automatically. Confirm the new task definition "
        "carries the expected values before smoke testing.",
    )
    note(
        doc,
        "Security warning",
        "Because SpringProfile is dev, DevDataLoader seeds demo accounts on startup, including "
        "an administrator account whose credentials are hardcoded in "
        "backend/core/src/main/java/com/careconnect/config/DevDataLoader.java. "
        "The API Gateway endpoint is "
        "public and unrestricted, so anyone who learns the URL can sign in as an "
        "administrator. Treat any environment deployed this way as a short-lived demo holding "
        "no real data, and destroy it when finished. Do not use SpringProfile=dev for anything "
        "resembling production.",
    )

    p(doc, "3.5.7 Smoke test", bold=True)
    p(doc, "Test the positive and the negative case. A permissive misconfiguration passes the positive test alone.")
    code(
        doc,
        "curl -s \"$BACKEND_URL/v1/api/test/health\"\n"
        "curl -s -H \"Origin: $FRONTEND_URL\" \"$BACKEND_URL/v1/api/test/health\"\n"
        "curl -s -H \"Origin: https://not-allowed.example.com\" \"$BACKEND_URL/v1/api/test/health\"",
    )
    table(
        doc,
        ["Request", "Expected", "Observed 2026-09-02"],
        [
            ["No Origin header", "200 healthy JSON", "200 healthy"],
            ["Origin set to the Amplify URL", "200 healthy JSON", "200 healthy"],
            ["Origin set to an unlisted host", "403 Invalid CORS request", "403 Invalid CORS request"],
        ],
        widths=[2.2, 2.3, 2.3],
    )
    p(doc, "Then verify in a browser:")
    bullets(
        doc,
        [
            "Open FRONTEND_URL. The welcome page should show 'Ready to connect your care!' and "
            "must not display the orange warning reading 'Backend service is not healthy'.",
            "The browser console should log BASE URL, STATUS CODE: 200, and PARSED STATUS: healthy.",
            "Click Continue and confirm the sign-in page renders, which also confirms the SPA "
            "rewrite is working.",
        ],
    )
    note(
        doc,
        "Reading the welcome page correctly",
        "welcome_page.dart is a landing page with buttons, not an auto-advancing splash. "
        "Navigation happens only when a button is pressed. On a short browser window the "
        "buttons fall below the fold, which can look like the application has hung. It has not.",
    )

    p(doc, "3.5.8 Share with testers", bold=True)
    p(
        doc,
        "Testers need only the Amplify URL. They do not need AWS credentials or the backend "
        "URL.",
    )

    p(doc, "3.5.9 Teardown", bold=True)
    code(
        doc,
        "export AWS_PROFILE=careconnect-sso\n"
        "./cloudformation-fargate/cdestroy_cloudformation.sh \\\n"
        "  --environment cfdemo --profile careconnect-sso\n\n"
        "aws amplify delete-app --app-id <appId> --region us-east-1",
    )
    p(
        doc,
        "The destroy script removes stacks in reverse dependency order and empties the ECR "
        "repository before deleting the platform stack. Confirm all four stacks are gone and "
        "that the Amplify app is deleted separately, since it is not managed by CloudFormation.",
    )

    doc.add_page_break()

    # ---------------- 4. Troubleshooting ----------------
    h(doc, "4. Troubleshooting", 1)

    h(doc, "4.1 AWS credentials and deploy scripts", 2)
    table(
        doc,
        ["Symptom", "Likely cause", "Fix"],
        [
            ["Token has expired and refresh failed, ExpiredToken, InvalidClientTokenId",
             "SSO session expired",
             "aws sso login --profile careconnect-sso, then confirm with aws sts get-caller-identity"],
            ["get-caller-identity shows the account root", "Wrong credentials",
             "Use the SSO role profile, not root keys"],
            ["Stale credentials override the profile", "Leftover environment variables",
             "unset AWS_ACCESS_KEY_ID AWS_SECRET_ACCESS_KEY AWS_SESSION_TOKEN"],
            ["Deploy script not found", "Wrong working directory",
             "Run from the repository root, the folder containing cloudformation-fargate"],
            ["Placeholder value detected", "REPLACE_ME values still in the parameter file",
             "Export CARECONNECT_DATABASE_MASTER_PASSWORD and CARECONNECT_JWT_SECRET before running"],
        ],
        caption="Table 6. Credentials and scripts",
        widths=[2.3, 1.8, 2.7],
    )

    h(doc, "4.2 Build failures", 2)
    table(
        doc,
        ["Symptom", "Likely cause", "Fix"],
        [
            ["Docker build fails on COPY target/...jar", "Wrong Maven profile",
             "Build from backend/core with ./mvnw clean package -Pdocker"],
            ["Compilation errors on a newer JDK", "Shell defaults to a newer Java",
             "Pin JAVA_HOME to a JDK 17 installation"],
            ["Cannot connect to the Docker daemon", "Docker Desktop not running",
             "Start Docker Desktop and wait for docker info to succeed"],
            ["Image build appears to hang installing packages", "ffmpeg install under amd64 emulation",
             "Expected on Apple Silicon. 176 packages and about 113 MB. Allow several minutes"],
            ["flutter build web fails on IconData", "Icon tree-shaking",
             "Add --no-tree-shake-icons. Not required as of this revision"],
            ["Flutter package or cache errors", "Stale artifacts",
             "flutter clean, then flutter pub get"],
        ],
        caption="Table 7. Build failures",
        widths=[2.3, 1.8, 2.7],
    )

    h(doc, "4.3 Deployed backend", 2)
    table(
        doc,
        ["Symptom", "Likely cause", "Fix"],
        [
            ["HTTP 503 right after the service stack completes", "Application still starting",
             "Wait two to five minutes. Confirm the ECS service reports a steady state"],
            ["503 persists", "Task cycling or no healthy target",
             "Check ECS service events, then Cloud Map instances for the service"],
            ["Invalid CORS request from the browser", "Amplify origin not in the allow-list",
             "Re-run the Section 3.5.6 stack update with the correct FRONTEND_URL"],
            ["ClusterNotFoundException", "Wrong cluster name",
             "Use careconnect-<env>-cluster, including the -cluster suffix"],
            ["ECR repository already exists", "Repository names are unique per account and region",
             "Use a unique RepositoryName for a parallel environment"],
            ["ECS task health shows UNKNOWN", "No healthCheck in the container definition",
             "Expected. See Section 4.5"],
        ],
        caption="Table 8. Deployed backend",
        widths=[2.3, 1.8, 2.7],
    )

    h(doc, "4.4 Amplify frontend", 2)
    table(
        doc,
        ["Symptom", "Likely cause", "Fix"],
        [
            ["Site root loads but every other route returns 404", "No SPA rewrite rule",
             "Apply the custom rule in Section 3.5.5"],
            ["Blank site after upload", "Zipped the folder rather than its contents",
             "Zip from inside build/web so index.html is at the archive root"],
            ["Welcome page reports the backend unhealthy", "BACKEND_URL missing at build time",
             "Rebuild with the dart-define and verify with the grep in Section 3.4.3"],
            ["/ui-preview is missing", "React preview build step skipped",
             "Build ui-integration/mobile-app and copy dist into web/ui-preview"],
            ["Console shows a 404 for assets/.env", "The app probes for a bundled .env",
             "Benign. It falls back to assets/.env.example"],
        ],
        caption="Table 9. Amplify frontend",
        widths=[2.3, 1.8, 2.7],
    )

    h(doc, "4.5 Known issues in the current deployment path", 2)
    p(
        doc,
        "These were identified during the 2026-09-02 deployment. They are recorded so "
        "operators are not surprised, and so they can be scheduled for repair.",
    )
    numbered(
        doc,
        [
            "Seeded demo credentials. SpringProfile=dev runs DevDataLoader, which creates an "
            "administrator account with a well-known password on a public endpoint. The exact "
            "credentials are in DevDataLoader.java and are not repeated here. This is the "
            "highest-severity item in this list.",
            "The notification WebSocket cannot connect. 04-service.yaml never provisions a "
            "WebSocket API and documents WebSocketApiGatewayEndpoint as optional, while "
            "WEBSOCKET_ENABLED is deployed as false. The frontend nevertheless connects to "
            "wss://<BACKEND_URL>/ws/notifications unconditionally, producing an uncaught "
            "exception in the console. It does not block the flows verified here.",
            "The Amplify SPA rewrite is not applied automatically for manual zip deploys and "
            "must be added by hand, as in Section 3.5.5.",
            "The Dockerfile HEALTHCHECK targets port 8080 while the service listens on 8081. "
            "This is currently inert because ECS does not use an image's built-in HEALTHCHECK, "
            "but it would fail if a container health check were added to the task definition.",
            "Allowed requests return access-control-allow-origin: * rather than echoing the "
            "specific origin. Enforcement still works, since disallowed origins are rejected "
            "with 403, but the wildcard is broader than the configured allow-list.",
            "cfdemo-service.json contains a BackendImageUri pointing at another cohort's "
            "AWS account, and a CorsAllowedList "
            "entry for that cohort's Amplify URL. The deploy script overrides the image URI, "
            "so this is misleading rather than harmful, but both should be genericized.",
            "The application loads assets/.env, which is not shipped, producing a 404 on every "
            "page load before falling back to assets/.env.example.",
        ],
    )

    doc.add_page_break()

    # ---------------- 5. AI Usage Statement ----------------
    h(doc, "5. AI Usage Statement", 1)
    p(
        doc,
        "Generative AI tooling was used in preparing this revision. Claude, running as a "
        "coding assistant with access to the repository and the AWS command line, was used to "
        "survey the codebase and infrastructure templates, execute the deployment described in "
        "Section 3.5, diagnose the issues recorded in Section 4.5, and draft this document.",
    )
    p(
        doc,
        "The deployment itself was real. The stack names, timings, sizes, HTTP status codes, "
        "and console output quoted in this guide were produced by that deployment against AWS "
        "the cohort AWS account on September 2, 2026, and are reproduced from the session record "
        "rather than composed for illustration. Statements that could not be verified are "
        "labelled as such, most notably the branch protection settings in Section 2.3, which "
        "require repository administrator rights the team does not hold.",
    )
    p(
        doc,
        "All decisions with security, cost, or governance consequences were made by the human "
        "team, including the choice to accept the seeded demo credentials for a short-lived "
        "environment, the selection of the deployment branch, and the revision of the "
        "CODEOWNERS file. The team is responsible for the content of this document.",
    )

    out = Path(__file__).resolve().parents[1] / "docs" / "05_Deployment_and_Operations_Guide_Fall2026.docx"
    doc.save(out)
    print(f"wrote {out}")
    return out


if __name__ == "__main__":
    build()
