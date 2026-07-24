"""Generate Word document: LLM provider abstraction, Bedrock/Team A dependency, BAA fallbacks."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX

ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "LLM_Provider_Abstraction_Bedrock_TeamA_BAA_Fallback.docx"


def heading(doc, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def para(doc, text: str, bold: bool = False, highlight: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    if highlight:
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def bullets(doc, items: list[str], highlight_indices: set[int] | None = None) -> None:
    highlight_indices = highlight_indices or set()
    for i, item in enumerate(items):
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        if i in highlight_indices:
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def code(doc, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"


def table(
    doc,
    headers: list[str],
    rows: list[list[str]],
    highlight_cells: set[tuple[int, int]] | None = None,
) -> None:
    highlight_cells = highlight_cells or set()
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for r_idx, row in enumerate(rows, start=1):
        cells = t.add_row().cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = val
            if (r_idx, c_idx) in highlight_cells:
                for p in cells[c_idx].paragraphs:
                    for run in p.runs:
                        run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def build() -> None:
    doc = Document()

    title = doc.add_heading(
        "LLM Provider Abstraction, Bedrock Dependency on Team A, and BAA-Eligible Fallback Options",
        0,
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para(doc, "CareConnect Team E — AI Gateway & Platform Integration Research")
    para(
        doc,
        "Synthesizes Team E SRS v2.0 (§8.2, §8.6), Milestone 2 TDD, Team A video/summary "
        "quickstart, deployment guides, and current backend codebase (post origin/main baseline).",
    )
    para(
        doc,
        "Yellow highlights mark gaps between design intent and current implementation.",
        highlight=True,
    )
    doc.add_paragraph()

    # 1 Executive Summary
    heading(doc, "1. Executive Summary", 1)
    para(
        doc,
        "Team E requires all language-model inference to run on HIPAA-eligible services under an "
        "active Business Associate Agreement (BAA), accessed through a provider abstraction layer "
        "so the primary model (Amazon Bedrock / Claude) can fail over to an alternate BAA-eligible "
        "provider without client changes. DeepSeek is explicitly under review and no longer assumed.",
    )
    para(
        doc,
        "The codebase has a partial abstraction: AIServiceFactory switches Bedrock vs DeepSeek for "
        "the AIService interface, and AIChatServiceConfig supports OpenAI-spec providers via LangChain4j. "
        "Team A’s video-call stack depends directly on BedrockSentimentService (not the factory) for "
        "sentiment, summarization, and call-summary JSON. Automatic BAA-compliant failover, unified "
        "provider telemetry, and consistent safety/context wiring are not yet implemented.",
        highlight=True,
    )

    # 2 Requirement sources
    heading(doc, "2. Requirement & Design Sources", 1)
    table(
        doc,
        ["Source", "Reference", "LLM / BAA relevance"],
        [
            ["SRS §8.2", "Software Interfaces", "Inference only on HIPAA-eligible BAA services"],
            ["SRS §8.6", "Design Constraints", "Abstraction layer; stable versioned APIs"],
            ["SRS §3", "UC-SUM / Ask AI flows", "Bedrock throttling → retry; dev → mock mode"],
            ["SRS §3", "FR-AI-9", "Minimum-necessary context to model provider (HIPAA/GDPR)"],
            ["SRS architecture", "Managed LLM row", "Bedrock/Claude primary; BAA fallback under review"],
            ["TDD", "Generative module pattern", "Retrieve → validate → invoke via abstraction"],
            ["TDD", "Open decisions", "Primary Bedrock/Claude vs DeepSeek; provision BAA first"],
            ["Deployment guide", "AWS selection rationale", "AWS BAA covers broad HIPAA-eligible stack"],
            ["Team A quickstart", "Chime + Bedrock", "Sentiment/summary Bedrock IAM + heuristic fallback"],
        ],
    )

    # 3 Target architecture
    heading(doc, "3. Target Architecture (SRS / TDD)", 1)
    code(
        doc,
        "Client (Flutter / Ask AI / Summaries)\n"
        "    ↓\n"
        "AI Gateway (Team E) — RBAC, consent, retrieval, safety, HITL\n"
        "    ↓\n"
        "LLM Abstraction Layer (stable contract: prompt in → structured JSON/text out)\n"
        "    ├─ Primary: Amazon Bedrock (Claude / Nova) — AWS BAA\n"
        "    ├─ Failover: BAA-eligible alternate (provider TBD; not DeepSeek by default)\n"
        "    └─ Dev-only: deterministic mock (no PHI to external API)\n"
        "    ↓\n"
        "Audit: provider id, model version, tokens, delivery status (FR-AI-10 / REQ-SC-9)",
    )
    bullets(
        doc,
        [
            "Failover behind the abstraction layer; alternate must be HIPAA-eligible under a BAA.",
            "Bedrock ClientError (throttling/unavailable): exponential backoff + retry; queue if persistent.",
            "Bedrock not provisioned (dev): mock mode with deterministic response (SRS UC-SUM alt flow).",
            "Record LLM provider and model version on summary records for traceability.",
            "DeepSeek direct API and OpenRouter are legacy/dev paths — not assumed BAA-compliant for PHI.",
        ],
    )

    # 4 Codebase abstraction layers
    heading(doc, "4. Current Codebase — Provider Abstraction Layers", 1)

    heading(doc, "4.1 Layer A: AIService + AIServiceFactory", 2)
    table(
        doc,
        ["Component", "Role", "Provider switch"],
        [
            ["AIService", "Interface: processChat + conversation stubs", "Implemented by Bedrock, DeepSeek"],
            ["AIServiceFactory", "Runtime selector", "careconnect.ai.provider = bedrock | deepseek"],
            ["BedrockAIChatService", "BedrockRuntimeClient.invokeModel", "@ConditionalOnProperty bedrock"],
            ["DeepSeekService", "OpenAI-spec REST to api.deepseek.com", "@ConditionalOnProperty deepseek"],
            ["LlmExtractionService", "Invoice JSON extraction", "Uses AIServiceFactory (any provider)"],
        ],
    )
    para(
        doc,
        "Configuration (application-dev.properties): careconnect.ai.enabled=false by default; "
        "careconnect.ai.provider=bedrock; careconnect.ai.model=amazon.nova-lite-v1:0. "
        "Claude Sonnet models commented as alternate.",
    )

    heading(doc, "4.2 Layer B: AIChatService (HTTP chat surface)", 2)
    table(
        doc,
        ["Implementation", "Condition", "Behavior"],
        [
            ["MockAIChatService", "provider=mock (matchIfMissing=true)", "Canned responses; full conversation CRUD"],
            ["BedrockAIChatAdapter", "Always @Service", "Delegates to AIServiceFactory → Bedrock path"],
            ["DefaultAIChatService", "provider=deepseek", "Disabled stub — throws UnsupportedOperationException"],
        ],
        highlight_cells={(2, 2), (3, 2)},
    )
    para(
        doc,
        "AIChatController loads only when careconnect.ai.enabled=true. When enabled with "
        "provider=bedrock, MockAIChatService is excluded but BedrockAIChatAdapter has no "
        "@ConditionalOnProperty — potential bean ambiguity if multiple AIChatService beans register.",
        highlight=True,
    )

    heading(doc, "4.3 Layer C: LangChain4j ChatModel (OpenAI-spec providers)", 2)
    para(
        doc,
        "AIChatServiceConfig (careconnect.ai.enabled=true, provider != bedrock) builds "
        "OpenAiChatModel with configurable careconnect.ai.api.url / api.key / model.name. "
        "Supports any OpenAI-compatible endpoint (DeepSeek, OpenRouter, etc.) but is separate "
        "from AIServiceFactory and not wired into BedrockAIChatAdapter today.",
    )

    heading(doc, "4.4 Layer D: Direct Bedrock integrations (bypass factory)", 2)
    table(
        doc,
        ["Service", "Uses", "Abstraction gap"],
        [
            ["BedrockSentimentService", "Direct BedrockRuntimeClient", "Team A path; heuristic fallback only"],
            ["BedrockAIChatService", "Direct invoke via BedrockModelSupport", "No medical context / safety pass"],
            ["BedrockChatModel", "LangChain4j ChatModel impl", "Legacy payload format; not in active factory"],
            ["OpenRouterService", "RestTemplate → openrouter.ai", "Disabled in dev; not BAA path for PHI"],
        ],
        highlight_cells={(1, 2), (2, 2), (4, 2)},
    )

    heading(doc, "4.5 BedrockModelSupport — model routing", 2)
    bullets(
        doc,
        [
            "Approved models: Nova Lite/Pro, Claude 3 Haiku, Claude 3.5 Sonnet, Claude Sonnet 4/4.5/4.6.",
            "Claude Sonnet 4/4.5 mapped to US inference profile IDs (us.anthropic.claude-…).",
            "Unapproved model requests fall back to default with warning.",
            "buildInvokePayload / parseTextResponse handle Nova vs Claude JSON shapes.",
        ],
    )

    # 5 Team A dependency
    heading(doc, "5. Team A Dependency on Bedrock", 1)
    para(
        doc,
        "Team A owns the Chime video-call stack. Bedrock is a hard runtime dependency for "
        "production-quality sentiment and call-summary generation — not routed through Team E’s "
        "planned AI Gateway or AIServiceFactory.",
        bold=True,
    )

    heading(doc, "5.1 Team A Bedrock touchpoints", 2)
    table(
        doc,
        ["Feature", "Service / API", "Bedrock usage", "Fallback"],
        [
            ["Text sentiment", "BedrockSentimentService.analyzeText", "Nova Pro prompt → JSON score", "Local heuristics"],
            ["Voice sentiment", "analyzeVoiceFromChimeMetrics", "Nova / voice model", "Heuristic from Chime metrics"],
            ["Video sentiment", "analyzeVideoFrame", "Nova Pro + image", "Neutral heuristic"],
            ["Combined sentiment", "analyzeCombined", "Weighted merge", "Excludes fallback channels"],
            ["Call summary JSON", "summarizeTranscript / CallSummaryService", "Nova Pro structured summary", "ERROR summary row"],
            ["REST", "CallController sentiment endpoints", "Exposes fallback flag in logs", "Call continues (SENT-007)"],
        ],
    )

    heading(doc, "5.2 Infrastructure & IAM (Team A provisioning)", 2)
    bullets(
        doc,
        [
            "AwsAccessConfig registers BedrockRuntimeClient when careconnect.aws.enabled=true.",
            "CloudFormation (04-service.yaml): bedrock:InvokeModel on amazon.nova-pro-v1:0 and nova-lite-v1:0 only.",
            "Claude inference profiles NOT in IAM resource list — Claude chat may fail in ECS until IAM updated.",
            "Team A quickstart: aws.bedrock.sentiment.model-id (default nova-pro); optional voice model override.",
            "Comprehend DetectSentiment also in IAM — parallel/legacy path; BedrockSentimentService uses Bedrock first.",
        ],
        highlight_indices={2},
    )

    heading(doc, "5.3 Coupling risks for Team E", 2)
    bullets(
        doc,
        [
            "Ask AI upstream summaries depend on call_summary.summary_json produced by Team A Bedrock path.",
            "Medication timeline (FR-AI-11) reads careInstructions[] from that same summary contract.",
            "Team E cannot swap LLM provider for summaries without Team A changing BedrockSentimentService.",
            "Shared AWS BAA covers Bedrock in-account failover (Nova ↔ Claude) without leaving AWS.",
        ],
        highlight_indices={2},
    )

    # 6 BAA fallback options
    heading(doc, "6. BAA-Eligible Fallback Options", 1)
    para(
        doc,
        "SRS: alternate fallback LLM must be HIPAA-eligible under a BAA; DeepSeek is under review "
        "and not the default assumption. Options below are ranked for CareConnect’s AWS-centric architecture.",
    )

    heading(doc, "6.1 Recommended: in-AWS Bedrock model failover (same BAA)", 2)
    table(
        doc,
        ["Fallback", "BAA status", "Fit", "Notes"],
        [
            ["Nova Lite ↔ Nova Pro", "AWS BAA", "High", "Already partially in IAM; cheap ↔ capable tradeoff"],
            ["Claude via inference profile", "AWS BAA", "High", "BedrockModelSupport ready; extend IAM ARNs"],
            ["Cross-region Bedrock", "AWS BAA", "Medium", "DR; adds latency/cost; same abstraction"],
            ["Amazon Comprehend", "AWS BAA", "Low (sentiment only)", "Already in IAM; not generative summaries"],
        ],
    )
    para(
        doc,
        "Best near-term failover: stay within Bedrock — retry/backoff, then switch approved model ID "
        "(e.g., Nova Pro throttled → Claude Sonnet profile). No second vendor BAA required.",
    )

    heading(doc, "6.2 Alternate cloud LLM (second BAA — higher integration cost)", 2)
    table(
        doc,
        ["Provider", "BAA", "Integration", "Team E consideration"],
        [
            ["Azure OpenAI Service", "Microsoft BAA (HIPAA)", "New AIService impl + Azure private endpoint", "Multi-cloud ops burden"],
            ["Google Vertex AI (Gemini)", "Google Cloud BAA", "New AIService impl + VPC-SC", "Multi-cloud ops burden"],
            ["AWS HealthLake / future FM", "AWS BAA", "If/when applicable generative APIs launch", "Watch AWS roadmap"],
        ],
    )

    heading(doc, "6.3 Not recommended for PHI production paths", 2)
    table(
        doc,
        ["Provider", "Why excluded", "Current code reference"],
        [
            ["DeepSeek API (direct)", "No CareConnect BAA assumed; SRS under review", "DeepSeekService, AiSymptomService"],
            ["OpenRouter", "Routes to arbitrary third-party models; no unified BAA", "OpenRouterService"],
            ["MockAIChatService", "Dev/test only; deterministic canned text", "provider=mock default when unset"],
            ["Public OpenAI API", "Requires OpenAI BAA + enterprise agreement", "AIChatServiceConfig OpenAiChatModel"],
        ],
        highlight_cells={(1, 1), (2, 1)},
    )

    heading(doc, "6.4 Non-LLM fallbacks (already implemented — Team A)", 2)
    bullets(
        doc,
        [
            "BedrockSentimentService: heuristic transcript scoring when Bedrock unavailable (fallback=true).",
            "Voice/video: metric-based heuristics when model invoke fails.",
            "SENT-007 test intent: sentiment service down → video call continues.",
            "SRS dev path: deterministic mock summary when Bedrock not provisioned (not fully wired in code).",
        ],
        highlight_indices={3},
    )

    # 7 Config matrix
    heading(doc, "7. Configuration Property Matrix", 1)
    table(
        doc,
        ["Property", "Default (dev)", "Controls"],
        [
            ["careconnect.ai.enabled", "false", "AIChatController + LangChain config gate"],
            ["careconnect.ai.provider", "bedrock", "AIServiceFactory: bedrock | deepseek"],
            ["careconnect.ai.model", "amazon.nova-lite-v1:0", "Bedrock default model ID"],
            ["careconnect.deepseek.enabled", "true (controllers)", "AiSymptom/AiAllergy/DeepSeekController"],
            ["careconnect.openrouter.enabled", "false (dev)", "OpenRouterService bean"],
            ["careconnect.llm.enabled", "false", "LlmExtractionService (invoice)"],
            ["careconnect.aws.enabled", "varies", "BedrockRuntimeClient bean in AwsAccessConfig"],
            ["aws.bedrock.sentiment.model-id", "amazon.nova-pro-v1:0", "Team A sentiment/summary model"],
        ],
        highlight_cells={(4, 1), (4, 2)},
    )
    para(
        doc,
        "GAP: Fragmented flags (ai.enabled, deepseek.enabled, openrouter.enabled, llm.enabled) "
        "do not express a single BAA-aware provider policy or failover order.",
        highlight=True,
    )

    # 8 Failover target design
    heading(doc, "8. Target Failover Design (Recommended)", 1)
    code(
        doc,
        "interface LlmProvider {\n"
        "  LlmResponse invoke(LlmRequest req);  // providerId, modelId, prompt, maxTokens\n"
        "  boolean isAvailable();\n"
        "  BaaScope baaScope();  // AWS_BAA | AZURE_BAA | DEV_MOCK\n"
        "}\n\n"
        "class LlmRouter {\n"
        "  List<LlmProvider> chain;  // ordered: primary Bedrock, secondary Bedrock model, mock\n"
        "  LlmResponse invoke(LlmRequest req) {\n"
        "    for (provider : chain) {\n"
        "      try { return provider.invoke(req); }\n"
        "      catch (ThrottlingException e) { backoff(); continue; }\n"
        "    }\n"
        "    throw new LlmUnavailableException();\n"
        "  }\n"
        "}",
    )
    bullets(
        doc,
        [
            "Unify BedrockAIChatService and BedrockSentimentService behind LlmRouter (or extend AIServiceFactory).",
            "Persist provider + model on call_summaries and chat audit rows (SRS postcondition).",
            "Dev profile: chain ends with MockLlmProvider — no outbound PHI.",
            "Prod profile: chain excludes mock; requires at least one BAA provider available.",
            "Team A sentiment keeps heuristic as last-resort non-LLM fallback (not a BAA substitute).",
        ],
        highlight_indices={0, 1},
    )

    # 9 Gap analysis
    heading(doc, "9. Codebase Gap Analysis", 1)
    table(
        doc,
        ["Gap", "Impact", "Priority"],
        [
            ["No automatic provider failover", "SRS throttling/queue path unmet", "P0"],
            ["BedrockSentimentService bypasses abstraction", "Team E tied to Team A Bedrock code", "P0"],
            ["IAM lacks Claude inference profile ARNs", "Claude chat/summary may fail in ECS", "P0"],
            ["BedrockAIChatService: raw message only", "FR-AI-9 minimization / retrieval not applied", "P1"],
            ["Dual AIChatService beans risk", "Ambiguous injection when enabling AI", "P1"],
            ["DeepSeek/OpenRouter still in tree", "Misconfiguration could send PHI off BAA", "P1"],
            ["PROGRAMMERS_GUIDE lists DeepSeek primary", "Docs drift from Bedrock reality", "P2"],
            ["No provider/model on summary records", "SRS traceability postcondition", "P1"],
        ],
        highlight_cells={
            (1, 0), (1, 1),
            (2, 0), (2, 1),
            (3, 0), (3, 1),
        },
    )

    # 10 Test alignment
    heading(doc, "10. Test & Operational Alignment", 1)
    table(
        doc,
        ["Test / req", "Expectation", "Current state"],
        [
            ["SRS UC-SUM 2a", "Bedrock throttle → backoff/retry", "Not implemented in CallSummaryService"],
            ["SRS UC-SUM 2b", "Dev mock when Bedrock absent", "MockAIChatService exists; summary mock partial"],
            ["SENT-007", "Sentiment down → call continues", "Heuristic fallback implemented"],
            ["FR-AI-9", "Min-necessary context to provider", "Not enforced on Bedrock chat path"],
            ["Sec 8.2", "BAA-only inference in prod", "Config allows non-BAA providers if enabled"],
        ],
        highlight_cells={(1, 2), (2, 2), (4, 2), (5, 2)},
    )

    # 11 Action items
    heading(doc, "11. Recommended Action Items", 1)
    table(
        doc,
        ["Owner", "Action", "Outcome"],
        [
            ["Team A", "Extend IAM for Claude inference-profile ARNs", "Claude failover works in ECS"],
            ["Team A", "Extract shared LlmRouter or delegate sentiment to factory", "Single abstraction for Team E"],
            ["Team E", "Implement Bedrock-primary + in-account model failover", "SRS failover without new BAA"],
            ["Team E", "Block non-BAA providers when profile=prod", "Prevent accidental DeepSeek/OpenRouter PHI"],
            ["Platform", "Confirm second-vendor BAA if leaving AWS Bedrock", "Legal sign-off on fallback choice"],
            ["Docs", "Update PROGRAMMERS_GUIDE AI section to Bedrock + abstraction", "Onboarding accuracy"],
        ],
    )

    # 12 References
    heading(doc, "12. Related Documents & Code References", 1)
    bullets(
        doc,
        [
            "C:\\Users\\ravic\\Downloads\\CareConnect_SRS_Revision 2.0_TEAM E.docx (§8.2, §8.6, Managed LLM row)",
            "C:\\Users\\ravic\\Downloads\\CareConnect_Milestone_2_TDD_TEAM E.docx",
            "docs/guides/TEAM_A_VIDEO_CALL_QUICKSTART.md",
            "docs/guides/DEPLOYMENT_AND_OPERATIONS_GUIDE.md (AWS BAA rationale)",
            "docs/spring2026_guides/AWS_Fargate_Deployment_Guide.md (Bedrock IAM)",
            "backend/core/src/main/java/com/careconnect/ai/AIServiceFactory.java",
            "backend/core/src/main/java/com/careconnect/ai/bedrock/BedrockModelSupport.java",
            "backend/core/src/main/java/com/careconnect/service/BedrockAIChatService.java",
            "backend/core/src/main/java/com/careconnect/service/BedrockSentimentService.java",
            "backend/core/src/main/java/com/careconnect/service/BedrockAIChatAdapter.java",
            "backend/core/src/main/java/com/careconnect/config/AIChatServiceConfig.java",
            "backend/core/src/main/java/com/careconnect/config/AwsAccessConfig.java",
        ],
    )

    heading(doc, "13. Conclusion", 1)
    para(
        doc,
        "CareConnect’s BAA posture centers on AWS: Bedrock (Claude/Nova) is the intended primary "
        "production LLM under the existing AWS Business Associate Agreement. The codebase reflects "
        "partial progress toward SRS’s abstraction layer via AIServiceFactory and BedrockModelSupport, "
        "but Team A’s BedrockSentimentService remains a direct coupling point for summaries and "
        "sentiment that Team E downstream features inherit.",
    )
    para(
        doc,
        "The lowest-risk BAA-eligible fallback is in-account Bedrock model failover (Nova ↔ Claude "
        "inference profiles), not DeepSeek or OpenRouter. A unified LlmRouter with ordered providers, "
        "prod guardrails against non-BAA endpoints, and provider/model audit fields closes the gap "
        "between SRS §8.2/§8.6 and the current implementation.",
        highlight=True,
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    try:
        doc.save(OUTPUT)
        print(f"Created: {OUTPUT}")
    except PermissionError:
        alt = OUTPUT.with_name(f"{OUTPUT.stem}_refresh{OUTPUT.suffix}")
        doc.save(alt)
        print(f"Created (original locked): {alt}")


if __name__ == "__main__":
    build()
