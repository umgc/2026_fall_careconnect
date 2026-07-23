"""Generate Word document: PR code review for HMAC Gmail OAuth branch."""

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.shared import Pt

OUTPUT = (
    Path(__file__).resolve().parents[1]
    / "docs"
    / "PR_Review_HMAC_OAuth_Gmail_Connect_feature_hmac-signed-oauth-state.docx"
)


def heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def para(doc: Document, text: str, bold: bool = False, highlight: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    if highlight:
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def bullets(doc: Document, items: list[str], highlight_indices: set[int] | None = None) -> None:
    highlight_indices = highlight_indices or set()
    for i, item in enumerate(items):
        p = doc.add_paragraph(item, style="List Bullet")
        if i in highlight_indices:
            for run in p.runs:
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def code(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    for i, h in enumerate(headers):
        tbl.rows[0].cells[i].text = h
    for row in rows:
        cells = tbl.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val


def build() -> Document:
    doc = Document()
    title = doc.add_heading("PR Code Review", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    para(doc, "HMAC-signed Gmail OAuth — feature/hmac-signed-oauth-state → team-ae-develop")
    para(doc, f"Review date: {date.today().isoformat()}")
    para(doc, "Reviewer: AI-assisted code review (Cursor)")
    doc.add_paragraph()

    heading(doc, "Review metadata", 1)
    table(
        doc,
        ["Field", "Value"],
        [
            ["Source branch", "feature/hmac-signed-oauth-state-structured-connection-status-gmail-disconnect-flow"],
            ["Target branch", "team-ae-develop"],
            ["Scope", "21 files changed (+1,908 / −540 lines)"],
            ["Commits (feature)", "6 commits (excludes merged CODEOWNERS PR #161)"],
            ["Related PR", "PR #235 — HMAC-signed OAuth state, structured connection status, Gmail disconnect"],
        ],
    )

    # ── 1. Change Summary ─────────────────────────────────────────────────────
    heading(doc, "1. Change Summary", 1)
    para(
        doc,
        "This PR hardens the Gmail OAuth flow used by the USPS Informed Delivery mail agent. "
        "It replaces unsigned OAuth state, unauthenticated browser start links, and boolean "
        "connection checks with a signed two-step OAuth model, structured status DTOs, token "
        "refresh/revoke handling, and a patient-scoped credential API.",
    )

    heading(doc, "Primary goals", 2)
    table(
        doc,
        ["Area", "What changed"],
        [
            [
                "OAuth CSRF protection",
                "OAuthStateSigner — HMAC-SHA256 signed state (600s TTL) and short-lived start tokens (120s TTL)",
            ],
            [
                "Two-step browser flow",
                "JWT GET /v1/api/email-credentials/gmail/connect-url → startToken → public GET /oauth/google/start",
            ],
            [
                "Redirect safety",
                "OAuthRedirectValidator whitelists returnUrl hosts (localhost + configured frontend origins)",
            ],
            [
                "Credential API",
                "EmailCredentialController + EmailCredentialService: status, connect-url, disconnect",
            ],
            [
                "Structured status",
                "EmailConnectionStatus record: CONNECTED, NOT_CONNECTED, NEEDS_RECONNECT with messages",
            ],
            [
                "Token lifecycle",
                "GoogleOAuthService: ensureFreshToken(), revokeIfPossible(), invalidateCredential()",
            ],
            [
                "Digest integration",
                "USPSDigestService calls ensureFreshToken before Gmail fetch",
            ],
            [
                "Security config",
                "Explicit /v1/api/email-credentials/** matcher before catch-all rules",
            ],
            [
                "Frontend (partial)",
                "Gmail connect/disconnect/status uses AuthTokenManager + /v1/api/email-credentials/*",
            ],
            [
                "Tests",
                "EmailOAuthFlowE2ETest (723 lines), unit tests for signer/validator/controllers",
            ],
        ],
    )

    heading(doc, "OAuth lifecycle (after PR)", 2)
    bullets(
        doc,
        [
            "Authenticated client calls GET /v1/api/email-credentials/gmail/connect-url?patientEmail=...&returnUrl=...",
            "EmailCredentialService validates requirePatientAccess, signs 120s startToken bound to patient user id",
            "Browser opens GET /oauth/google/start?startToken=... (public, token-bound)",
            "Backend re-signs full OAuth state and redirects to Google",
            "Google callback hits GET /oauth/google/callback; state verified; tokens exchanged and stored",
            "User redirected to returnUrl or default /usps-test; failures append oauthError query param",
            "Client polls GET /v1/api/email-credentials/status for structured connection state",
            "DELETE /v1/api/email-credentials/gmail revokes and removes stored credential",
        ],
    )

    heading(doc, "Out of scope / tangential", 2)
    bullets(
        doc,
        [
            "CODEOWNERS update (merged from PR #161) — not part of OAuth feature logic",
            "USPS digest/search/clear-cache frontend paths largely unchanged in this branch",
            "No changes to UspsDigestController or USPSController patient auth in this diff",
        ],
    )

    # ── 2. Bug & Risk Analysis ────────────────────────────────────────────────
    heading(doc, "2. Bug & Risk Analysis", 1)

    heading(doc, "Critical", 2)

    heading(doc, "C1 — @RequirePermission(VIEW_ASSIGNED_PATIENTS) blocks patient self-service", 3)
    para(
        doc,
        "EmailCredentialController gates all endpoints with VIEW_ASSIGNED_PATIENTS. Patients only "
        "have VIEW_HEALTH_DATA per RolePermissionService. A patient connecting their own Gmail for "
        "USPS will receive 403 from PermissionAspect before requirePatientAccess runs.",
    )
    para(doc, "E2E tests only exercise ADMIN role — patientUser fixture is created but never tested.", highlight=True)

    heading(doc, "C2 — Frontend OAuth hardening is incomplete (digest/search/cache paths)", 3)
    para(doc, "Gmail connect/disconnect/status were updated, but these methods were not:", bold=True)
    bullets(
        doc,
        [
            "_fetchDigest — still uses unauthenticated Dio(), demo-user fallback, /api/usps/latest?userId=",
            "_searchMail — still uses unauthenticated Dio(), demo-user fallback, /api/usps/search?userId=",
            "_clearCache — still uses unauthenticated Dio(), demo-user fallback, /api/usps/clear-cache?userId=",
        ],
        highlight_indices={0, 1, 2},
    )
    para(
        doc,
        "Impact: Users can connect Gmail via hardened API but still fetch USPS data through legacy "
        "unauthenticated paths — undermining the security model this PR introduces.",
    )

    heading(doc, "High", 2)

    heading(doc, "H1 — GoogleOAuthService.exchange always inserts a new EmailCredential row", 3)
    para(
        doc,
        "exchange() always credRepo.save(new EmailCredential()) without upserting or deleting the "
        "prior row. Reconnect flows can accumulate multiple GMAIL rows per user; "
        "findFirstByUserIdAndProviderOrderByIdDesc returns latest only — orphaned rows remain.",
    )

    heading(doc, "H2 — OAuth error messages exposed in browser redirect URL", 3)
    para(
        doc,
        "buildOAuthErrorRedirect URL-encodes e.getMessage() into oauthError query param. "
        "Internal errors (token exchange, DB, crypto) appear in browser history and referrer logs. "
        "E2E test callback_whenTokenExchangeFails explicitly expects token+exchange+failed in URL.",
    )
    code(doc, "return base + separator + OAUTH_ERROR_PARAM + \"=\" + encodedMessage;")

    heading(doc, "H3 — Start token replay within TTL", 3)
    para(
        doc,
        "Start tokens (120s) are cryptographically verified but not single-use. A captured startToken "
        "can re-initiate OAuth until expiry. Mitigated by short TTL but not equivalent to nonce consumption.",
    )

    heading(doc, "Medium", 2)
    bullets(
        doc,
        [
            "demo-user fallback retained in EmailCredentialService.resolvePatientUser — inconsistent with hardened auth intent",
            "System.out.println / System.err.println in GoogleOAuthService.exchange — noisy, partial clientId logged in production",
            "E2E test mocks EmailCredentialRepository — status/disconnect flows don't hit real DB persistence",
            "ensureFreshToken concurrent refresh race — two parallel digest requests could both refresh",
            "OAuth redirect_uri double-encoding via UriUtils.encode + build(true) — verify against Google OAuth spec",
            "Patient identifier resolution duplicated in EmailCredentialService (no shared UspsPatientResolver in this branch)",
        ],
    )

    heading(doc, "Low", 2)
    bullets(
        doc,
        [
            "invalidateCredential sets expiresAt to Instant.EPOCH — works but semantically odd vs deletion",
            "No PK/unique constraint on (user_id, provider) in EmailCredential entity",
            "patientUser E2E fixture unused — gap in role coverage",
            "Legacy unsigned state format correctly rejected (good) but no migration path for in-flight OAuth sessions during deploy",
        ],
    )

    # ── 3. Architecture & Style ───────────────────────────────────────────────
    heading(doc, "3. Architecture & Style", 1)

    heading(doc, "Strengths", 2)
    bullets(
        doc,
        [
            "OAuthStateSigner is well-structured: separate start vs callback tokens, constant-time compare, TTL enforcement, nonce",
            "OAuthRedirectValidator provides sensible open-redirect protection with localhost dev fallback",
            "EmailCredentialService cleanly separates credential lifecycle from OAuth transport (EmailOAuthController)",
            "EmailConnectionStatus DTO enables reconnect UX without parsing booleans",
            "Removing @RequirePermission from public EmailOAuthController — correct; browser cannot send JWT",
            "SecurityConfig explicit matcher for /v1/api/email-credentials/** fixes ordering bug vs catch-all",
            "EmailOAuthFlowE2ETest covers full connect → start → callback → status → disconnect lifecycle",
            "USPSDigestService integration with ensureFreshToken prevents stale-token Gmail failures",
        ],
    )

    heading(doc, "Weaknesses", 2)
    bullets(
        doc,
        [
            "Partial frontend migration — Gmail paths hardened, digest paths left on legacy unauthenticated API",
            "Permission layer (@RequirePermission) conflicts with requirePatientAccess for non-caregiver roles",
            "GoogleOAuthService mixes transport, persistence, and console logging",
            "E2E tests mock repository while claiming full-context integration — hybrid test style",
            "No upsert pattern for credential storage",
        ],
    )
    para(doc, "Overall: OAuth backend design is a significant security upgrade over unsigned u:|r: state. Merge readiness depends on fixing permission gating and completing frontend auth migration.")

    # ── 4. Recommendations ────────────────────────────────────────────────────
    heading(doc, "4. Recommendations", 1)

    heading(doc, "R1 — Fix permission gate for patient self-service (Critical)", 2)
    para(doc, "Option A (recommended): Remove @RequirePermission from EmailCredentialController; rely on requirePatientAccess only.", bold=True)
    code(
        doc,
        """@GetMapping("/status")
public ResponseEntity<EmailConnectionStatus> getConnectionStatus(...) {
    // No @RequirePermission — requirePatientAccess in service handles all roles
    return ResponseEntity.ok(emailCredentialService.getGmailConnectionStatus(identifier));
}""",
    )
    para(doc, "Option B: Use @RequireAnyPermission({VIEW_ASSIGNED_PATIENTS, VIEW_HEALTH_DATA}) — requires new annotation + PermissionAspect support.")

    heading(doc, "R2 — Complete frontend auth migration (Critical)", 2)
    code(
        doc,
        """Future<void> _fetchDigest() async {
  final patientEmail = _patientQueryValue();
  if (patientEmail == null) { /* show login error */ return; }

  final dio = await _authenticatedDio();
  final url = '$base/v1/api/usps/latest'
      '?patientEmail=${Uri.encodeComponent(patientEmail)}&date=$dateString';
  final resp = await dio.get(url);
  // ...
}""",
    )
    para(doc, "Apply same pattern to _searchMail and _clearCache; remove all demo-user fallbacks.")

    heading(doc, "R3 — Upsert credentials on OAuth exchange (High)", 2)
    code(
        doc,
        """public void exchange(String userId, String code) {
    // ...
    EmailCredential ec = credRepo
        .findFirstByUserIdAndProviderOrderByIdDesc(userId, EmailCredential.Provider.GMAIL)
        .orElseGet(EmailCredential::new);
    ec.setUserId(userId);
    ec.setProvider(EmailCredential.Provider.GMAIL);
    // set tokens, expiresAt ...
    credRepo.save(ec);
}""",
    )

    heading(doc, "R4 — Sanitize OAuth error redirects (High)", 2)
    code(
        doc,
        """private String buildOAuthErrorRedirect(String returnUrl, Exception e) {
    log.warn("Gmail OAuth callback failed", e);
    String base = resolveSuccessRedirect(returnUrl);
    String separator = base.contains("?") ? "&" : "?";
    return base + separator + OAUTH_ERROR_PARAM + "=oauth_failed";
}""",
    )

    heading(doc, "R5 — Replace System.out with structured logging (Medium)", 2)
    code(
        doc,
        """private static final Logger log = LoggerFactory.getLogger(GoogleOAuthService.class);
// log.debug("Starting token exchange for userId={}", userId);""",
    )

    heading(doc, "R6 — Add patient-role E2E test (Medium)", 2)
    code(
        doc,
        """@Test
@DisplayName("patient can obtain connect-url for self")
void connectUrl_patientSelf_returns200() throws Exception {
    mockMvc.perform(get("/v1/api/email-credentials/gmail/connect-url")
            .with(user(PATIENT_EMAIL).roles("PATIENT")))
        .andExpect(status().isOk());
}""",
    )

    heading(doc, "R7 — Optional: single-use start tokens (Low)", 2)
    para(
        doc,
        "Track consumed start-token nonces in a short-TTL cache (Redis or in-memory with TTL) "
        "to prevent replay within the 120s window.",
    )

    # ── Verdict ───────────────────────────────────────────────────────────────
    heading(doc, "Verdict", 1)
    table(
        doc,
        ["Dimension", "Assessment"],
        [
            ["Security intent", "Major improvement — signed state, JWT-gated start, redirect validation"],
            ["OAuth architecture", "Clean separation; well-tested signer/validator"],
            ["API / RBAC", "requirePatientAccess correct; @RequirePermission too restrictive for patients"],
            ["Frontend completeness", "Gmail flow updated; digest/search/cache still legacy — partial"],
            ["Test coverage", "Strong OAuth lifecycle E2E; missing patient-role scenarios"],
            ["Merge readiness", "Conditional — fix R1 + R2 before merge; R3–R4 strongly recommended"],
        ],
    )

    para(
        doc,
        "This PR delivers a materially stronger Gmail OAuth implementation compared to team-ae-develop "
        "(unsigned u:|r: state, permission-gated public /oauth/google/start). The highest-priority "
        "gaps are the permission gate blocking patients and the incomplete frontend migration for "
        "USPS digest endpoints.",
        highlight=True,
    )

    return doc


def main() -> None:
    doc = build()
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    main()
