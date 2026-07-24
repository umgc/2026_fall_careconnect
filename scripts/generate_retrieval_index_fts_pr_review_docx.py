"""Generate Word document: PR review for Tasks 4.1+4.2 FTS coverage branch."""

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.shared import Pt

OUTPUT = (
    Path(__file__).resolve().parents[1]
    / "docs"
    / "PR_Review_Retrieval_Index_FTS_Coverage_feature_a-rvasireddy.docx"
)


def heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def para(doc: Document, text: str, bold: bool = False, highlight: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    if highlight:
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def code(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    for i, h in enumerate(headers):
        tbl.rows[0].cells[i].text = h
    for row in rows:
        cells = tbl.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val


def build() -> Document:
    doc = Document()
    title = doc.add_heading("PR Code Review", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    para(
        doc,
        "Tasks 4.1 + 4.2 (WBS 3.16.3 / 3.12.1) — RetrievalIndexService, IndexWorker, "
        "and FTS coverage — feature/a-rvasireddy-retrieval-index-fts-coverage → team-ae-develop",
    )
    para(doc, f"Review date: {date.today().isoformat()}")
    para(doc, "Reviewer: AI-assisted code review (Cursor)")
    doc.add_paragraph()

    heading(doc, "Review metadata", 1)
    table(
        doc,
        ["Field", "Value"],
        [
            ["Source branch", "feature/a-rvasireddy-retrieval-index-fts-coverage"],
            ["Target branch", "team-ae-develop"],
            [
                "Commits (feature)",
                "dd7fcbd Task 4.1 IndexWorker; ec057e4 TX isolation; c8fc9bf Task 4.2 FTS",
            ],
            ["Scope", "28 files (+2670 / −7)"],
            ["Verdict", "Approve with changes — claim lease race + FTS filter placement"],
        ],
    )

    heading(doc, "1. Change Summary", 1)
    para(
        doc,
        "This PR lands the Ask AI indexing write path and the FTS keyword-read leg. "
        "Task 4.1 (WBS 3.16.3) adds RetrievalIndexService + IndexWorker to consume "
        "indexing_outbox (SUMMARY_CREATED / TRANSCRIPT_INDEXED), chunk summaries and "
        "transcripts into retrieval_index_chunk, with TransactionTemplate-separated "
        "claim / ingest / status updates (EvvOutboxProcessor-style), SKIP LOCKED "
        "claiming, visit-summary deferral until Task 1.4, and max-attempt dead-lettering. "
        "Task 4.2 (WBS 3.12.1 keyword leg) verifies the search_vector trigger, adds "
        "backfill migration V2607121930 + SchemaPatchRunner patch, and exposes "
        "patient-scoped plainto_tsquery search via FullTextSearchService and "
        "RetrievalIndexChunkRepository.searchByPatientIdFullText.",
    )

    heading(doc, "What changed", 2)
    table(
        doc,
        ["Area", "Change"],
        [
            ["IndexWorker", "Scheduled poller; separate TXs; defer/fail/dead-letter accounting"],
            ["RetrievalIndexService", "ingestSummaryCreated / ingestTranscriptIndexed + chunkers"],
            ["FullTextSearchService", "patientId-scoped FTS with limit clamp and optional type filter"],
            ["Repository", "searchByPatientIdFullText, countMissingSearchVector, deleteBySourceRecordId"],
            ["Migrations", "V2607121930 backfill + trigger reaffirm; V2607071921 comments"],
            ["Config", "careconnect.indexing.*; worker disabled in test profile"],
            ["Tests", "Unit/E2E/config/FTS contract tests (~heavy coverage)"],
        ],
    )

    heading(doc, "Verdict", 2)
    para(
        doc,
        "Approve with changes. Prior batch rollback-only bug is fixed. Core indexing and "
        "FTS contracts are solid and well tested. Before merge under multi-task ECS: fix "
        "or document the claim-without-lease race (SKIP LOCKED only holds during the short "
        "claim TX). Also push record-type filtering into the FTS SQL so LIMIT applies after "
        "RBAC type filters, and correct the stale javadoc on transcript deferral.",
        bold=True,
        highlight=True,
    )

    heading(doc, "2. Bug & Risk Analysis", 1)

    heading(doc, "2.1 Strengths", 2)
    bullets(
        doc,
        [
            "Separate claim / ingest / status transactions — no shared rollback-only poison across the batch.",
            "Visit summaries deferred (not falsely marked processed); deferred rows eventually dead-letter.",
            "Empty drafts do not delete existing chunks; contentHash exact equality skip avoids churn.",
            "FTS always filters patient_id (FR-AI-1); native SELECT omits search_vector/embedding (entity-safe).",
            "Backfill + SchemaPatchRunner parity for ECS envs that rely on patches.",
            "Worker disabled under test profile; strong unit/E2E/config coverage.",
            "Overview no longer dumps raw JSON (PHI/noise) — skip empty overview.",
        ],
    )

    heading(doc, "2.2 High — Claim TX releases locks before processing (multi-ECS race)", 2)
    para(
        doc,
        "claimUnprocessedForPolling uses FOR UPDATE SKIP LOCKED inside a short "
        "TransactionTemplate that commits before processRow. After commit, row locks "
        "are released while processed_at is still NULL. A second ECS task can claim "
        "and process the same outbox rows concurrently → duplicate chunk writes / "
        "delete races. SKIP LOCKED only protects during the claim query itself.",
        highlight=True,
    )
    code(
        doc,
        """-- Prefer claim-with-lease inside the claim TX (example):
UPDATE indexing_outbox
SET attempt_count = attempt_count + 1,  -- or claimed_at = now(), claimed_by = :taskId
    last_error = 'claimed'
WHERE id IN (
  SELECT id FROM indexing_outbox
  WHERE processed_at IS NULL
    AND (claimed_at IS NULL OR claimed_at < now() - interval '2 minutes')
  ORDER BY id
  FOR UPDATE SKIP LOCKED
  LIMIT :limit
)
RETURNING *;""",
    )
    bullets(
        doc,
        [
            "MVP single-task ECS: lower severity; still document the constraint.",
            "Alternative: process each claimed id immediately while holding lock "
            "(longer TX — trade latency for safety) or use advisory locks per id.",
        ],
    )

    heading(doc, "2.3 Medium — FTS record-type filter applied after LIMIT", 2)
    para(
        doc,
        "FullTextSearchService.search(... allowedRecordTypes ...) fetches LIMIT rows "
        "then filters in memory. A mixed-type top-k can return fewer than requested "
        "CALL_SUMMARY (etc.) hits, or zero after filter. Push record_type INTO the SQL.",
    )
    code(
        doc,
        """AND (:recordTypesEmpty = true OR record_type IN (:recordTypes))
-- and apply LIMIT after that filter""",
    )

    heading(doc, "2.4 Medium — Visit summaries will dead-letter before Task 1.4", 2)
    para(
        doc,
        "Visit SUMMARY_CREATED throws IndexingDeferredException and burns attempts. "
        "After max-attempts they get processed_at set and will not auto-retry when "
        "visit indexing ships. Operators need a reset query (clear processed_at / "
        "attempt_count for Task 1.4 visit events) or a dedicated defer-without-burn "
        "bucket for known-unimplemented types.",
    )

    heading(doc, "2.5 Medium — Ingest success + markProcessed failure = reprocess", 2)
    para(
        doc,
        "Ingest commits in TX1; markProcessed is TX2. If TX2 fails, chunks exist but "
        "outbox stays unprocessed. Next poll re-ingests (usually safe via contentHash / "
        "delete+rewrite) but can amplify load. Consider same TX for ingest+mark, or "
        "idempotent outbox key (eventId) on chunks.",
    )

    heading(doc, "2.6 Low — Stale javadoc on ingestTranscriptIndexed", 2)
    para(
        doc,
        "Javadoc still says 'Defers (does not burn attempts)' but IndexWorker now "
        "increments attempt_count on IndexingDeferredException.",
    )

    heading(doc, "2.7 Low — Migration test asserts comment text", 2)
    para(
        doc,
        "task42BackfillMigrationPresent asserts contains(\"plainto_tsquery\") which "
        "only appears in a SQL comment, not executable DDL. Prefer asserting "
        "to_tsvector / WHERE search_vector IS NULL only.",
    )

    heading(doc, "2.8 Low — History includes unrelated CODEOWNERS commits", 2)
    para(
        doc,
        "Branch history includes Create/Update CODEOWNERS from merge origin/main. "
        "File is not in the develop...HEAD file list — no functional impact — but "
        "rebase onto team-ae-develop before PR for a cleaner history if possible.",
    )

    heading(doc, "2.9 No nested rollback-only issue (fixed)", 2)
    para(
        doc,
        "Prior review blocker (shared @Transactional poll catching RuntimeException) "
        "is resolved via TransactionTemplate per stage.",
    )

    heading(doc, "3. Architecture & Style", 1)
    bullets(
        doc,
        [
            "Clear separation: worker (transport/polling) vs RetrievalIndexService "
            "(domain ingest) vs FullTextSearchService (query leg for Task 5.1).",
            "Chunkers are focused and testable; SummaryChunker overview/item split is clean.",
            "FTS config constants on RetrievalIndexSchema document english config + query cap.",
            "ConditionalOnProperty + test-profile disable matches Evv/outbox patterns.",
            "Native FTS SQL is the right place for ranking (not in-memory keyword filter).",
            "Nested @Transactional on ingest joins the worker's TransactionTemplate TX "
            "(REQUIRED) — correct; document so future REQUIRES_NEW does not surprise.",
        ],
    )

    heading(doc, "4. Recommendations", 1)

    heading(doc, "4.1 [High] Lease rows in the claim transaction", 2)
    code(
        doc,
        """@Modifying
@Query(value = \"\"\"
    UPDATE indexing_outbox SET last_error = 'CLAIMED:' || CAST(:leaseId AS text)
    WHERE id IN (
      SELECT id FROM indexing_outbox
      WHERE processed_at IS NULL
        AND (last_error IS NULL OR last_error NOT LIKE 'CLAIMED:%'
             OR updated_heuristic...)
      ORDER BY id FOR UPDATE SKIP LOCKED LIMIT :limit
    )
    \"\"\", nativeQuery = true)
int claimBatch(...);

// Or add claimed_at / claimed_by columns via Flyway — cleaner than overloading last_error.""",
    )

    heading(doc, "4.2 [Medium] Push record_type filter into FTS SQL", 2)
    code(
        doc,
        """@Query(value = \"\"\"
    SELECT id, patient_id, record_type, source_record_id, chunk_text,
           chunk_metadata, indexed_at, consent_scope
    FROM retrieval_index_chunk
    WHERE patient_id = :patientId
      AND search_vector @@ plainto_tsquery('english', :query)
      AND (COALESCE(:filterTypes, false) = false
           OR record_type IN (:recordTypes))
    ORDER BY ts_rank_cd(search_vector, plainto_tsquery('english', :query)) DESC,
             indexed_at DESC
    LIMIT :limit
    \"\"\", nativeQuery = true)
List<RetrievalIndexChunk> searchByPatientIdFullText(...);""",
    )

    heading(doc, "4.3 [Medium] Special-case visit deferral without burning budget", 2)
    code(
        doc,
        """} catch (final IndexingDeferredException ex) {
    if (isVisitNotImplemented(ex)) {
        // leave attempt_count unchanged; optional next_retry_at far in future
        log.warn("Deferring visit indexing until Task 1.4: {}", ex.getMessage());
        return;
    }
    recordDeferOrDeadLetter(row, attempts, root.getMessage());
}""",
    )

    heading(doc, "4.4 [Low] Fix javadoc; tighten migration test", 2)
    code(
        doc,
        """// RetrievalIndexService.ingestTranscriptIndexed:
 * Defers when patientId cannot be resolved. IndexWorker burns attempt budget
 * on IndexingDeferredException so deferred rows eventually dead-letter.

// Migration test:
assertThat(sql).contains("WHERE search_vector IS NULL");
assertThat(sql).contains("to_tsvector('english'");
assertThat(sql).doesNotContain("SELECT * FROM");""",
    )

    heading(doc, "File-level review comments", 1)
    table(
        doc,
        ["File", "Severity", "Comment"],
        [
            [
                "IndexWorker.java",
                "High",
                "Claim TX commits before process — multi-worker duplicate risk.",
            ],
            [
                "IndexingOutboxRepository.java",
                "High",
                "SKIP LOCKED alone is not a durable lease.",
            ],
            [
                "FullTextSearchService.java",
                "Medium",
                "Post-LIMIT type filter under-fills results.",
            ],
            [
                "RetrievalIndexService.java",
                "Low",
                "Stale javadoc on deferral attempt burning; visit dead-letter ops risk.",
            ],
            [
                "V2607121930__….sql",
                "Info",
                "Good backfill; SchemaPatchRunner mirror present.",
            ],
            [
                "RetrievalIndexChunkRepository.java",
                "Info",
                "Patient-scoped FTS + column projection is correct.",
            ],
            [
                "*Test.java",
                "Info",
                "Strong coverage; fix plainto_tsquery assertion on backfill SQL.",
            ],
        ],
    )

    heading(doc, "Test plan (recommended before merge)", 1)
    bullets(
        doc,
        [
            "mvnw test -Dtest=IndexWorkerTest,RetrievalIndexServiceTest,IndexingPipelineE2ETest,"
            "FullTextSearchServiceTest,RetrievalIndexFtsCoverageTest,IndexWorkerConfigTest",
            "Two app instances with worker.enabled=true — confirm no duplicate outbox "
            "processing (or accept single-task constraint).",
            "Insert chunk with NULL search_vector; run migration/patch; FTS finds row.",
            "Visit SUMMARY_CREATED — deferred, eventually DLQ; document reset SQL for Task 1.4.",
            "FTS with allowedRecordTypes=CALL_SUMMARY returns up to limit of that type only "
            "(after SQL filter fix).",
        ],
    )

    return doc


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    build().save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    main()
