"""Generate Word document: weekly tasks and hours (Jul 5–12, 2026) with WBS v4 IDs."""

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

OUTPUT = (
    Path(__file__).resolve().parents[1]
    / "docs"
    / "Team_E_Weekly_Tasks_and_Hours_2026-07-05_to_2026-07-12.docx"
)

WBS_SOURCE = r"C:\Users\ravic\Downloads\WBS v4.docx"

# (Category, WBS v4, Team backlog ref, Description, Hours, Status)
ROWS = [
    (
        "Implementation",
        "3.8.1–3.8.4",
        "P0",
        "Ask AI deploy unblock — ECS Bedrock IAM, pgvector RDS, prod profile, "
        "SchemaPatchRunner, WebSocket empty default, secrets/SSM validation (PR #292 merged)",
        8.0,
        "Merged to team-ae-develop",
    ),
    (
        "Implementation",
        "3.16.1",
        "1.5 / 1.6",
        "retrieval_index_chunk Flyway schema — pgvector extension, FTS trigger, GIN index, "
        "SchemaPatchRunner patches, PR #276 review fixes",
        6.0,
        "Merged to team-ae-develop",
    ),
    (
        "Implementation",
        "3.12.1",
        "2.1",
        "RetrievalScopeService — RBAC scope resolution for Ask AI retrieval (PR #249)",
        4.0,
        "Merged (carry-over verification this week)",
    ),
    (
        "Implementation",
        "3.12.1",
        "2.6",
        "Forbidden scope 403 handler + audit logging on unauthorized retrieval access (PR #273)",
        2.0,
        "Merged (carry-over verification this week)",
    ),
    (
        "Implementation",
        "3.16.3",
        "4.1",
        "RetrievalIndexService + IndexWorker — ingest hooks for SUMMARY_CREATED / "
        "TRANSCRIPT_INDEXED outbox events, chunkers, SKIP LOCKED claim, TX isolation fix, "
        "deferral dead-lettering, E2E tests",
        16.0,
        "Branch: feature/a-rvasireddy-retrieval-index-service-index-worker",
    ),
    (
        "Implementation",
        "3.12.1",
        "4.2",
        "FTS keyword leg — search_vector trigger verification, backfill migration, "
        "FullTextSearchService, patient-scoped plainto_tsquery query, contract tests "
        "(commit c8fc9bf)",
        7.0,
        "Branch: feature/a-rvasireddy-retrieval-index-fts-coverage",
    ),
    (
        "PR review",
        "3.8 / 3.16",
        "P0",
        "Ask AI P0 deploy unblock (own PR) — review doc + follow-up fixes",
        2.0,
        "Approve with minor fixes",
    ),
    (
        "PR review",
        "3.16.3",
        "4.1",
        "Retrieval Index Service / IndexWorker (own branch) — detailed review + "
        "blocking TX/deferral fixes implemented",
        3.0,
        "Request changes → fixed",
    ),
    (
        "PR review",
        "3.11.6",
        "—",
        "GET /api/v3/summaries/{id} (feature/a-fasaa-get-summary-by-id) — "
        "IDOR analysis; refresh after authz commits",
        2.5,
        "Approve with changes",
    ),
    (
        "PR review",
        "3.6.7",
        "2.3.3",
        "Admin analytics dashboard (feature/a-dkinchen-admin-analytics-dashboard)",
        1.5,
        "Request changes",
    ),
    (
        "PR review",
        "3.8.1",
        "—",
        "Deploy permissions fix (feature/a-drattray-deploy-permissions-fix)",
        1.0,
        "Approve with IAM changes",
    ),
    (
        "PR review",
        "3.14.1",
        "—",
        "HMAC OAuth Gmail connect (feature/hmac-signed-oauth-state)",
        1.0,
        "Approve minor",
    ),
    (
        "PR review",
        "3.14.4",
        "3.14.4",
        "Textract OCR fallback",
        1.0,
        "Approve minor",
    ),
    (
        "PR review",
        "3.13",
        "—",
        "STML support (feat/team-ae-stml-support)",
        1.0,
        "Reviewed",
    ),
    (
        "PR review",
        "3.3",
        "—",
        "Sentiment clip playback (feature/a-drattray-sentiment)",
        0.75,
        "Reviewed",
    ),
    (
        "PR review",
        "3.6.1–3.6.2",
        "2.2.2",
        "Telemetry GoRouter observer (feature/a-dkinchen-telemetry-expansion)",
        0.75,
        "Reviewed",
    ),
    (
        "PR review",
        "3.14",
        "—",
        "USPS mail agent (feature/usps-mail-agent)",
        0.5,
        "Reviewed",
    ),
    (
        "Design / docs",
        "1.9",
        "TDD",
        "Team E Milestone 2 TDD section outline (docs/Team_E_Milestone_2_TDD_Section_Outline.docx)",
        2.0,
        "Complete",
    ),
    (
        "Design / docs",
        "1.9 / 1.14",
        "Backlog",
        "Team E implementation task backlog + hybrid retrieval / RBAC / upstream pipeline "
        "design docx generators",
        3.0,
        "Complete",
    ),
]


def heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def para(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True


def build() -> Document:
    doc = Document()
    title = doc.add_heading("Weekly Tasks and Hours", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    para(doc, "CareConnect — Team E (Ask AI / Retrieval)")
    para(doc, "Contributor: Ravichandra Vasireddy")
    para(doc, "Period: July 5 – July 12, 2026")
    para(doc, f"Generated: {date.today().isoformat()}")
    para(doc, f"WBS source: {WBS_SOURCE}")
    doc.add_paragraph()

    total = sum(r[4] for r in ROWS)

    heading(doc, "Summary", 1)
    para(doc, f"Total estimated hours: {total:.1f}", bold=True)
    para(
        doc,
        "Task numbers use CareConnect WBS v4 (Project Plan). Team E backlog tier IDs "
        "(e.g. 4.1, 2.1) are shown in the Backlog ref column where applicable. "
        "Hours are scope-based estimates — adjust in Word if your actual time differs.",
    )

    impl_h = sum(r[4] for r in ROWS if r[0] == "Implementation")
    review_h = sum(r[4] for r in ROWS if r[0] == "PR review")
    docs_h = sum(r[4] for r in ROWS if r[0] == "Design / docs")

    doc.add_paragraph(f"Implementation: {impl_h:.1f} h", style="List Bullet")
    doc.add_paragraph(f"PR reviews: {review_h:.1f} h", style="List Bullet")
    doc.add_paragraph(f"Design / documentation: {docs_h:.1f} h", style="List Bullet")

    heading(doc, "Task detail", 1)

    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    headers = [
        "Category",
        "WBS v4",
        "Backlog ref",
        "Description",
        "Hours",
        "Status",
    ]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h

    for category, wbs, backlog, description, hours, status in ROWS:
        row = table.add_row().cells
        row[0].text = category
        row[1].text = wbs
        row[2].text = backlog
        row[3].text = description
        row[4].text = f"{hours:.1f}"
        row[5].text = status

    total_row = table.add_row().cells
    total_row[0].text = ""
    total_row[1].text = ""
    total_row[2].text = ""
    total_row[3].text = "TOTAL"
    total_row[4].text = f"{total:.1f}"
    total_row[5].text = ""

    doc.add_paragraph()
    heading(doc, "WBS v4 mapping notes", 2)
    for item in [
        "3.16.3 — ingest hooks / indexing pipeline (Team backlog 4.1 IndexWorker)",
        "3.12.1 — RBAC-scoped hybrid retrieval FTS + pgvector (Team backlog 4.2 FTS leg, 2.1 scope)",
        "3.16.1 — Flyway index tables including retrieval_index_chunk (Team backlog 1.5/1.6)",
        "3.11.5 — summary persistence + retrieval-index update (upstream to 3.16.3)",
        "3.11.6 — GET summary by id (reviewed PR for Fon)",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    heading(doc, "Branches in progress (not merged)", 2)
    for item in [
        "feature/a-rvasireddy-retrieval-index-service-index-worker (WBS 3.16.3 / backlog 4.1)",
        "feature/a-rvasireddy-retrieval-index-fts-coverage (WBS 3.12.1 / backlog 4.2)",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    heading(doc, "Next week (planned)", 2)
    for item in [
        "Open PR: WBS 3.16.3 + 3.12.1 FTS → team-ae-develop",
        "WBS 3.12.1 — embeddings / pgvector semantic leg (Team backlog 4.3)",
        "WBS 3.12.1 — full hybrid retrieval RRF merge (Team backlog 5.1)",
        "WBS 3.12.3 — POST /api/ai/ask (Team backlog 5.3)",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    for row in table.rows:
        row.cells[3].width = Inches(3.2)

    return doc


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    build().save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    main()
