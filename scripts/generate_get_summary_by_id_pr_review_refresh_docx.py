"""Generate Word document: PR review refresh for GET summary by id after authz fixes."""

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.shared import Pt

OUTPUT = (
    Path(__file__).resolve().parents[1]
    / "docs"
    / "PR_Review_Get_Summary_By_Id_feature_a-fasaa-get-summary-by-id_refresh.docx"
)


def heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def para(doc: Document, text: str, bold: bool = False, highlight: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    if highlight:
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def code(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    for i, h in enumerate(headers):
        tbl.rows[0].cells[i].text = h
    for row in rows:
        cells = tbl.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val


def build() -> Document:
    doc = Document()
    title = doc.add_heading("PR Code Review", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    para(
        doc,
        "WBS 3.11.6 — GET summary by id (post-authz follow-ups) — "
        "feature/a-fasaa-get-summary-by-id → team-ae-develop",
    )
    para(doc, f"Review date: {date.today().isoformat()}")
    para(doc, "Reviewer: AI-assisted code review (Cursor)")
    doc.add_paragraph()

    heading(doc, "Review metadata", 1)
    table(
        doc,
        ["Field", "Value"],
        [
            ["Source branch", "feature/a-fasaa-get-summary-by-id"],
            ["Target branch", "team-ae-develop"],
            [
                "Commits",
                "11c26ff feat endpoint; 8480fe1 @PreAuthorize; "
                "33fd262 four-way access check (Dominique)",
            ],
            ["Scope", "6 files (+455 / −4)"],
            ["Verdict", "Approve with minor changes"],
        ],
    )

    heading(doc, "1. Change Summary", 1)
    para(
        doc,
        "Delivers WBS 3.11.6: GET /api/v3/summaries/{id} so clients correlating "
        "SUMMARY_CREATED can fetch a specific CallSummary by primary key, with the same "
        "response shape as GET /api/v3/calls/{callId}/summary. Follow-up commits close the "
        "prior IDOR finding: @PreAuthorize(CAREGIVER|PATIENT|ADMIN) plus the same four-way "
        "object check as CallController.getCallSummary (admin / telemetry participant / "
        "transcript access / summary owner → else 403). Also maps "
        "MethodArgumentTypeMismatchException → 400 globally.",
    )

    heading(doc, "What changed", 2)
    table(
        doc,
        ["File", "Change"],
        [
            [
                "CallSummaryController.java",
                "New endpoint + PreAuthorize + four-way access + getCurrentUser",
            ],
            [
                "CallSummaryService.java",
                "getSummaryEntityById / getSummaryById",
            ],
            [
                "GlobalExceptionHandler.java",
                "Type mismatch → 400",
            ],
            [
                "CallSummaryControllerTest.java",
                "200×4 auth paths, 403, 404, 400",
            ],
            [
                "CallSummaryServiceTest.java",
                "null / not-found / SUCCESS / ERROR",
            ],
            [
                "ScheduledVisitControllerTest.java",
                "Invalid date 500 → 400",
            ],
        ],
    )

    heading(doc, "Verdict", 2)
    para(
        doc,
        "Approve with minor changes. Prior High IDOR is largely addressed by mirroring "
        "CallController's four-way check with solid controller tests. Remaining items: "
        "stale Javadoc about EnableMethodSecurity (it is already on in SecurityConfig), "
        "FAMILY_MEMBER excluded from @PreAuthorize while sibling CallController allows "
        "them via the four-way check, redundant second DB read, and optional shared "
        "auth helper / caregiverVisibility (parity gap shared with CallController).",
        bold=True,
        highlight=True,
    )

    heading(doc, "2. Bug & Risk Analysis", 1)

    heading(doc, "2.1 Strengths", 2)
    bullets(
        doc,
        [
            "Object-level auth now matches CallController.getCallSummary — closes prior IDOR.",
            "Explicit 403 (not 404) for unauthorized, matching sibling contract.",
            "Controller tests cover all four allow paths plus 403/404/400.",
            "Response shape reuse via toResponse keeps clients interchangeable.",
            "Type-mismatch → 400 is correct HTTP; ScheduledVisit test updated.",
            "No race conditions — read-only path.",
        ],
    )

    heading(doc, "2.2 Medium — Javadoc wrong: @EnableMethodSecurity already enabled", 2)
    para(
        doc,
        "Controller Javadoc says @PreAuthorize is a silent no-op until "
        "feature/bjackson-rbac-infrastructure. SecurityConfig on team-ae-develop already "
        "has @EnableMethodSecurity. Role gate is live (or will be once this merges onto "
        "develop with that config). Update the comment to avoid false confidence that "
        "role checks are inactive.",
        highlight=True,
    )

    heading(doc, "2.3 Medium — FAMILY_MEMBER blocked by @PreAuthorize", 2)
    para(
        doc,
        "hasAnyRole('CAREGIVER','PATIENT','ADMIN') excludes FAMILY_MEMBER. "
        "CallController.getCallSummary has no equivalent PreAuthorize and allows family "
        "via telemetry/transcript participation. A family member with transcript access "
        "can read via /api/v3/calls/{callId}/summary but gets 403 on /api/v3/summaries/{id} "
        "before the four-way check runs. Either add FAMILY_MEMBER to PreAuthorize or "
        "document intentional divergence.",
    )
    code(
        doc,
        "@PreAuthorize(\"hasAnyRole('CAREGIVER', 'PATIENT', 'FAMILY_MEMBER', 'ADMIN')\")",
    )

    heading(doc, "2.4 Low — Double fetch of the same summary", 2)
    para(
        doc,
        "getSummaryEntityById(id) loads the entity for authz, then getSummaryById(id) "
        "loads again for the response. Prefer mapping the already-loaded entity "
        "(expose toResponse or return Map from a single service call after authz).",
    )
    code(
        doc,
        """// After access check:
return ResponseEntity.ok(callSummaryService.toResponse(summary));
// (make toResponse package/public, or add mapToResponse(CallSummary))""",
    )

    heading(doc, "2.5 Low — getCurrentUser NPE if Authentication is null", 2)
    para(
        doc,
        "Mirrors CallController, so parity is fine. Still fragile: authentication.getName() "
        "NPEs if context is empty. Optional null-safe guard.",
    )
    code(
        doc,
        """private User getCurrentUser() {
    final Authentication authentication =
            SecurityContextHolder.getContext().getAuthentication();
    if (authentication == null || authentication.getName() == null) {
        throw new AppException(HttpStatus.UNAUTHORIZED, "User not authenticated");
    }
    return userRepository.findByEmail(authentication.getName())
            .orElseThrow(() -> new AppException(HttpStatus.UNAUTHORIZED, "User not authenticated"));
}""",
    )

    heading(doc, "2.6 Low — caregiverVisibility not enforced (parity with CallController)", 2)
    para(
        doc,
        "Neither this endpoint nor CallController.getCallSummary filters on "
        "caregiverVisibility (on_consent / hidden). Acceptable for parity; track as a "
        "shared follow-up so caregivers cannot read hidden summaries via either path.",
    )

    heading(doc, "2.7 Low — Path version vs WBS; trailing newline", 2)
    bullets(
        doc,
        [
            "WBS /api/summaries/{id} vs implemented /api/v3/summaries/{id} — document for clients.",
            "CallSummaryControllerTest still missing final newline.",
            "Tests do not assert role-level PreAuthorize denial (e.g. wrong role).",
        ],
    )

    heading(doc, "2.8 Residual vs prior review", 2)
    para(
        doc,
        "Prior High IDOR (any authenticated user) is fixed for the four-way paths. "
        "Role gate + object check is the right two-layer design. Remaining gaps are "
        "role-set completeness and DRY, not open IDOR for CAREGIVER/PATIENT strangers.",
    )

    heading(doc, "3. Architecture & Style", 1)
    bullets(
        doc,
        [
            "Two-layer auth (method security + object check) is the right pattern.",
            "Duplicating CallController's four-way logic is correct for parity but will "
            "drift — extract a shared CallSummaryAccessService / helper.",
            "Controller grew dependencies (telemetry, transcript, users) — expected for "
            "inline authz; extraction would slim it again.",
            "Tests are clear and path-oriented; good regression net for authz.",
            "Service layer remains clean and null-safe.",
        ],
    )

    heading(doc, "4. Recommendations", 1)

    heading(doc, "4.1 [Medium] Include FAMILY_MEMBER; fix Javadoc", 2)
    code(
        doc,
        """@PreAuthorize("hasAnyRole('CAREGIVER', 'PATIENT', 'FAMILY_MEMBER', 'ADMIN')")
@GetMapping("/{id}")
public ResponseEntity<Map<String, Object>> getSummaryById(@PathVariable Long id) {
    // ...
}

// Javadoc: remove "silent no-op until EnableMethodSecurity" — already enabled.""",
    )

    heading(doc, "4.2 [Low] Avoid double load; extract shared access check", 2)
    code(
        doc,
        """// Shared helper (used by CallController + CallSummaryController)
public void requireCallSummaryAccess(User user, CallSummary summary) {
    String callId = summary.getCallId();
    boolean allowed =
            user.getRole() == Role.ADMIN
            || isTelemetryParticipant(user.getId(), callId)
            || callTranscriptService.hasTranscriptAccess(callId, user.getId())
            || user.getId().equals(summary.getGeneratedByUserId());
    if (!allowed) {
        throw new AppException(HttpStatus.FORBIDDEN, "Access denied");
    }
}

// Controller:
CallSummary summary = callSummaryService.getSummaryEntityById(id)
        .orElse(null);
if (summary == null) {
    return ResponseEntity.notFound().build();
}
requireCallSummaryAccess(getCurrentUser(), summary);
return ResponseEntity.ok(callSummaryService.toResponse(summary));""",
    )

    heading(doc, "4.3 [Low] Add PreAuthorize denial test", 2)
    code(
        doc,
        """@Test
@WithMockUser(roles = "USER")  // not CAREGIVER/PATIENT/ADMIN
void getSummaryById_wrongRole_returns403() throws Exception {
    mockMvc.perform(get("/api/v3/summaries/{id}", SUMMARY_ID))
            .andExpect(status().isForbidden());
}""",
    )

    heading(doc, "File-level review comments", 1)
    table(
        doc,
        ["File", "Severity", "Comment"],
        [
            [
                "CallSummaryController.java",
                "Medium",
                "Javadoc stale re EnableMethodSecurity; FAMILY_MEMBER missing from PreAuthorize.",
            ],
            [
                "CallSummaryController.java",
                "Low",
                "Double fetch; consider shared access helper with CallController.",
            ],
            [
                "CallSummaryController.java",
                "Info",
                "Four-way check correctly mirrors sibling endpoint — good IDOR fix.",
            ],
            [
                "CallSummaryService.java",
                "Info",
                "Clean; expose toResponse for single-load mapping.",
            ],
            [
                "CallSummaryControllerTest.java",
                "Low",
                "Strong four-path coverage; add role-gate test; trailing newline.",
            ],
            [
                "GlobalExceptionHandler.java",
                "Info",
                "Good 400 mapping.",
            ],
        ],
    )

    heading(doc, "Test plan (recommended before merge)", 1)
    bullets(
        doc,
        [
            "mvnw test -Dtest=CallSummaryControllerTest,CallSummaryServiceTest,"
            "ScheduledVisitControllerTest",
            "Family member with transcript access: call-scoped summary 200 vs by-id "
            "(expect match after FAMILY_MEMBER fix).",
            "Stranger caregiver → 403; owner / admin / telemetry / transcript → 200.",
            "Non-numeric id → 400.",
            "Confirm clients use /api/v3/summaries/{id}.",
        ],
    )

    return doc


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    build().save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    main()
