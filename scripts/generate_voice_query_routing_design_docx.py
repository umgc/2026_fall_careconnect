"""Generate Word document: Voice query routing through the same retrieval pipeline (design)."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX

ROOT = Path(__file__).resolve().parents[1]
OUTPUT = (
    ROOT
    / "docs"
    / "Voice_Query_Routing_Same_Retrieval_Pipeline_Design.docx"
)


def heading(doc, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def para(doc, text: str, bold: bool = False, highlight: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    if highlight:
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def bullets(doc, items: list[str], highlight_indices: set[int] | None = None) -> None:
    highlight_indices = highlight_indices or set()
    for i, item in enumerate(items):
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        if i in highlight_indices:
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def code(doc, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"


def table(
    doc,
    headers,
    rows,
    highlight_rows: set[int] | None = None,
    highlight_cells: set[tuple[int, int]] | None = None,
) -> None:
    highlight_rows = highlight_rows or set()
    highlight_cells = highlight_cells or set()
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for r_idx, row in enumerate(rows, start=1):
        cells = t.add_row().cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = val
            if r_idx in highlight_rows or (r_idx, c_idx) in highlight_cells:
                for p in cells[c_idx].paragraphs:
                    for run in p.runs:
                        run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def build() -> None:
    doc = Document()

    title = doc.add_heading(
        "Voice Query Routing Through the Same Retrieval Pipeline — Design Specification",
        0,
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para(doc, "CareConnect Team E — Ask AI Voice Input (SRS §3.2 / NFR-AI-1)")
    para(
        doc,
        "Implementation design for routing voice-origin patient questions through the identical "
        "SCC-3 Ask AI pipeline as typed text: client-side STT produces query text; "
        "POST /api/ai/ask invokes AiAskOrchestrator with inputModality=VOICE; hybrid retrieval, "
        "safety validation, HITL, citations, and disclaimers are unchanged. Builds on voice/STT "
        "research, API contract, retrieval orchestration, HITL, and medication timeline designs.",
    )
    para(
        doc,
        "Status: DESIGN ONLY — Ask AI chat is text-only (ai_chat_improved.dart); no /api/ai/ask "
        "gateway; VoiceCommandAI exists for symptom/allergy forms only; five STT stacks do not "
        "feed retrieval today.",
        highlight=True,
    )
    doc.add_paragraph()

    # 1 Goals
    heading(doc, "1. Design Goals", 1)
    bullets(
        doc,
        [
            "SRS §3.2: Voice queries use the same retrieval pipeline as text — no parallel AI path.",
            "FR-AI-9: Minimum-necessary — raw audio never sent to backend for Ask AI MVP.",
            "NFR-AI-1: End-to-end ≤ 5 s p95 including client STT + retrieval + LLM (budget STT ≤ 2 s).",
            "Single AiAskOrchestrator entry — inputModality is audit/metadata only, not routing fork.",
            "Identical RBAC, hybrid search, SafetyPipeline, Tier 2 HITL for voice-origin queries.",
            "Locale alignment: STT language matches request.locale (en-US / es-US).",
            "Reuse proven VoiceCommandAI singleShot pattern from symptom/allergy flows.",
            "Distinguish Ask AI voice mic from telehealth call STT (index upstream, not query path).",
        ],
        highlight_indices={0, 1, 3, 7},
    )

    heading(doc, "1.1 Architectural principle", 2)
    code(
        doc,
        "┌─────────────────────────────────────────────────────────────────┐\n"
        "│  VOICE and TEXT converge at POST /api/ai/ask { query: string }  │\n"
        "│  Backend has NO audio ingestion for Ask AI MVP                  │\n"
        "└─────────────────────────────────────────────────────────────────┘\n"
        "\n"
        "Mic → Client STT → normalized query text → AiAskOrchestrator\n"
        "                      (same as typed input)",
    )

    heading(doc, "1.2 Out of scope (this design)", 2)
    bullets(
        doc,
        [
            "Server-side streaming STT (AWS Transcribe real-time for Ask AI queries).",
            "Wake-word always-on Ask AI (Porcupine optional P2 enhancement only).",
            "Telehealth call capture STT — see transcript review doc (upstream index, not query).",
            "Alexa skill → Ask AI bridge (separate channel; optional P3 extension).",
            "TTS read-back of answers (accessibility feature — client TTS, not backend).",
        ],
    )

    # 2 Pipeline position
    heading(doc, "2. End-to-End Flow", 1)
    code(
        doc,
        "┌─────────────── CLIENT ───────────────┐   ┌──────────── BACKEND (unchanged) ────────────┐\n"
        "│ AskAiVoiceInput                      │   │ AiAskController.postAsk()                   │\n"
        "│   └─ VoiceCaptureService.capture()   │   │   └─ AiAskOrchestrator.ask()                │\n"
        "│        speech_to_text (12s dictation)  │   │        1 RetrievalScopeService              │\n"
        "│        → transcript string           │   │        2 RetrievalQueryPlanner              │\n"
        "│   └─ VoiceQueryNormalizer (optional) │   │        3 HybridRetrievalService             │\n"
        "│        trim, collapse whitespace       │   │        4 MedicationTimelineAggregator*      │\n"
        "│ AskAiService.askVoice(transcript)      │   │        5 RetrievalContextAssembler          │\n"
        "│   POST /v1/api/ai/ask                 │──▶│        6 GroundedLlmService                 │\n"
        "│   { query, inputModality: VOICE, ... }│   │        7 SafetyPipeline + HitlService       │\n"
        "└──────────────────────────────────────┘   │        8 AiAskResponseBuilder               │\n"
        "                                            └─────────────────────────────────────────────┘\n"
        "  * when MEDICATION_TIMELINE intent detected — same as text path",
    )

    # 3 Components
    heading(doc, "3. Component Architecture", 1)
    table(
        doc,
        ["Layer", "Component", "Responsibility"],
        [
            ["Frontend", "AskAiVoiceInput", "Mic button + listening UI on Ask AI chat surface"],
            ["Frontend", "VoiceCaptureService", "Wraps speech_to_text; returns transcript + captureMeta"],
            ["Frontend", "VoiceQueryNormalizer", "Client-side trim, whitespace, empty-query guard"],
            ["Frontend", "AskAiService", "HTTP client for POST /api/ai/ask (text + voice)"],
            ["Backend", "AiAskController", "Validates AiAskRequest; no audio multipart"],
            ["Backend", "VoiceQueryPreflight", "Optional server normalize; reject empty post-STT query"],
            ["Backend", "AiAskOrchestrator", "Same pipeline — reads inputModality for audit only"],
            ["Backend", "AiAuditService", "Log inputModality, sttLocale, captureDurationMs"],
        ],
        highlight_cells={
            (1, 1), (2, 1), (3, 1), (4, 1), (6, 1),
        },
    )

    heading(doc, "3.1 What does NOT get new backend services", 2)
    bullets(
        doc,
        [
            "No VoiceRetrievalService — hybrid search runs on query text identically.",
            "No VoiceSafetyPipeline — SafetyPipeline receives same SafetyInput.",
            "No separate HITL queue — held_item.source_surface=ASK_AI for both modalities.",
            "No Bedrock audio model invocation for Ask AI queries.",
        ],
        highlight_indices={0, 1},
    )

    # 4 STT boundary
    heading(doc, "4. Speech-to-Text Boundary", 1)
    para(
        doc,
        "CareConnect has multiple STT stacks (see Voice_Query_Path research doc). Only one applies "
        "to Ask AI voice queries.",
        bold=True,
    )
    table(
        doc,
        ["STT path", "Used for Ask AI voice query?", "Role"],
        [
            ["Flutter speech_to_text (VoiceCommandAI pattern)", "YES — MVP", "Client captures user question"],
            ["AWS Chime / Web Speech (call embed)", "NO", "Indexes call conversation for retrieval"],
            ["AWS Transcribe post-call", "NO", "Indexes diarized call transcript"],
            ["Sherpa ONNX (notetaker)", "NO", "Indexes clinical notes"],
            ["Amazon Alexa ASR", "NO (P3 optional)", "Structured task intents only today"],
        ],
        highlight_rows={1},
    )

    heading(doc, "4.1 VoiceCaptureService specification", 2)
    code(
        doc,
        "class VoiceCaptureService {\n"
        "  Future<VoiceCaptureResult> capture({\n"
        "    required String locale,           // en_US | es_US → speech_to_text localeId\n"
        "    Duration listenFor = 12s,\n"
        "    Duration pauseFor = 2s,\n"
        "  });\n"
        "}\n\n"
        "record VoiceCaptureResult(\n"
        "  String transcript,\n"
        "  bool isFinal,\n"
        "  String sttEngine,                 // speech_to_text\n"
        "  String localeId,\n"
        "  int captureDurationMs,\n"
        "  String? errorCode                 // PERMISSION_DENIED | UNAVAILABLE | TIMEOUT\n"
        ")",
    )
    bullets(
        doc,
        [
            "Reuse VoiceCommandAI listen options: dictation mode, partialResults, autoPunctuation.",
            "onDevice: false (matches existing VoiceCommandAI — platform cloud STT where available).",
            "Porcupine wake word: NOT used for Ask AI MVP — explicit mic tap only.",
            "Web: speech_to_text via browser Web Speech API; Porcupine disabled (existing behavior).",
            "Empty transcript → client error; do not POST to backend.",
        ],
        highlight_indices={2, 4},
    )

    heading(doc, "4.2 VoiceQueryNormalizer", 2)
    code(
        doc,
        "normalize(String raw) {\n"
        "  s = trim(raw);\n"
        "  s = collapseWhitespace(s);\n"
        "  s = stripLeadingFiller(['hey care connect', 'ok', 'um']);  // optional light cleanup\n"
        "  if (s.length < 2) throw EmptyQueryException;\n"
        "  if (s.length > 2000) s = truncateWithEllipsis(s, 2000);     // API maxLength\n"
        "  return s;\n"
        "}",
    )
    para(
        doc,
        "Server-side VoiceQueryPreflight repeats trim + length check — defense against malformed clients.",
    )

    # 5 API contract
    heading(doc, "5. API Contract Integration", 1)
    para(doc, "Extends POST /api/ai/ask request (see API contract design doc). No new endpoint.", bold=True)

    heading(doc, "5.1 Voice-origin request example", 2)
    code(
        doc,
        "POST /v1/api/ai/ask\n"
        "Authorization: Bearer <JWT>\n"
        "Content-Type: application/json\n\n"
        "{\n"
        '  "query": "When did I start lisinopril?",\n'
        '  "patientId": 42,\n'
        '  "sessionId": "a1b2c3d4-e5f6-7890-abcd-ef1234567890",\n'
        '  "inputModality": "VOICE",\n'
        '  "locale": "en-US",\n'
        '  "clientRequestId": "voice-20260620-001",\n'
        '  "voiceCaptureMeta": {\n'
        '    "sttEngine": "speech_to_text",\n'
        '    "sttLocale": "en_US",\n'
        '    "captureDurationMs": 4200,\n'
        '    "platform": "android"\n'
        "  }\n"
        "}",
    )

    heading(doc, "5.2 voiceCaptureMeta (optional extension)", 2)
    table(
        doc,
        ["Field", "Required", "Purpose"],
        [
            ["sttEngine", "No", "Audit: which client STT produced query"],
            ["sttLocale", "No", "Verify locale alignment with request.locale"],
            ["captureDurationMs", "No", "Latency telemetry; NFR-AI-1 budget tracking"],
            ["platform", "No", "ios | android | web — diagnostics only"],
            ["confidence", "No", "If platform exposes STT confidence — log only, not routing"],
        ],
    )
    para(
        doc,
        "voiceCaptureMeta is NEVER used for retrieval ranking or safety tier decisions — audit/telemetry only.",
        highlight=True,
    )

    heading(doc, "5.3 Response — identical to text", 2)
    bullets(
        doc,
        [
            "Same AiAskResponse envelope: answer, citations, disclaimer, escalation, tier, heldItemId.",
            "UI may show transcript bubble above answer ('You asked: …') — client-side display only.",
            "HITL poll URL and medicationTimeline field unchanged.",
            "Optional response field inputModalityEcho: VOICE for client analytics.",
        ],
    )

    heading(doc, "5.4 Error handling", 2)
    table(
        doc,
        ["Condition", "HTTP", "Client behavior"],
        [
            ["Empty query after normalize", "400 EMPTY_QUERY", "Show 'Could not understand — try again'"],
            ["STT permission denied", "(client-only)", "Open settings prompt; no API call"],
            ["403 FORBIDDEN_SCOPE", "403", "Same as text — scope error message"],
            ["Tier 2 HELD", "200/202", "Same poll flow; show held message"],
            ["503 retrieval degraded", "503", "Retry with text fallback suggestion"],
        ],
    )

    # 6 Orchestrator
    heading(doc, "6. AiAskOrchestrator — Modality-Neutral Pipeline", 1)
    code(
        doc,
        "AiAskResponse ask(AiAskRequest req, User caller) {\n"
        "  String query = voiceQueryPreflight.normalize(req.query());\n"
        "  RetrievalScope scope = scopeService.resolve(req.patientId(), caller, req.sourceTypes());\n"
        "  RetrievalPlan plan = queryPlanner.plan(query, scope);\n"
        "  List<RankedChunk> chunks = hybridRetrieval.search(scope, plan, query);\n"
        "  AggregatedMedicationTimeline timeline = timelineAggregator.maybeAggregate(chunks, plan, query);\n"
        "  GroundedContext ctx = contextAssembler.assemble(chunks, timeline, query, req.locale());\n"
        "  GroundedLlmResponse llm = groundedLlm.invoke(ctx, req.locale());\n"
        "  SafetyOutcome safety = safetyPipeline.process(SafetyInput.from(ctx, llm, req));\n"
        "  // inputModality NOT consulted in any stage above\n"
        "  AiAskResponse resp = responseBuilder.build(llm, safety, ctx, timeline);\n"
        "  auditService.log(req, resp, scope, chunks, req.inputModality(), req.voiceCaptureMeta());\n"
        "  return resp;\n"
        "}",
    )

    heading(doc, "6.1 inputModality usage matrix", 2)
    table(
        doc,
        ["Stage", "Uses inputModality?", "Notes"],
        [
            ["RetrievalScopeService", "No", "RBAC identical"],
            ["RetrievalQueryPlanner", "No", "Intent from query text only"],
            ["HybridRetrievalService", "No", "Same embedding + FTS on query string"],
            ["MedicationTimelineAggregator", "No", "Same FR-AI-11 path"],
            ["GroundedLlmService", "No", "Same prompt; optional hint 'user spoke question' in system footer only"],
            ["SafetyPipeline", "No", "Same validators"],
            ["HitlService", "No", "Same hold/release"],
            ["AiAuditService", "Yes", "FR-AI-10 modality field"],
            ["Analytics / metrics", "Yes", "p95 latency split STT vs backend"],
        ],
        highlight_rows={7, 8},
    )

    heading(doc, "6.2 Optional system prompt footer (P1)", 2)
    para(
        doc,
        "When inputModality=VOICE, append to system prompt: 'The user spoke this question; "
        "interpret colloquial phrasing but answer only from retrieved records.' "
        "Does NOT change retrieval or safety rules.",
    )

    # 7 Frontend integration
    heading(doc, "7. Frontend Integration Design", 1)

    heading(doc, "7.1 Ask AI chat UI changes", 2)
    code(
        doc,
        "ai_chat_improved.dart (or new ask_ai_page.dart)\n"
        "  TextField (existing)\n"
        "  IconButton(mic) → AskAiVoiceInput.startCapture()\n"
        "       → VoiceCaptureService.capture(locale: userLocale)\n"
        "       → show interim transcript in input field (optional UX)\n"
        "       → on final: VoiceQueryNormalizer.normalize()\n"
        "       → AskAiService.ask(query, inputModality: VOICE, voiceCaptureMeta: ...)\n"
        "       → render AiAskResponse (same widget as text send)",
    )
    para(
        doc,
        "GAP: ai_chat_improved.dart has no mic button; AIChatService posts to legacy /ai-chat/chat only.",
        highlight=True,
    )

    heading(doc, "7.2 AskAiService (new — replaces AIChatService for records-grounded Ask AI)", 2)
    code(
        doc,
        "class AskAiService {\n"
        "  static Future<AiAskResponse> ask({\n"
        "    required String query,\n"
        "    required int patientId,\n"
        "    InputModality inputModality = InputModality.text,\n"
        "    VoiceCaptureMeta? voiceCaptureMeta,\n"
        "    String locale = 'en-US',\n"
        "    String? sessionId,\n"
        "  });\n"
        "  // POST ${baseUrl}/v1/api/ai/ask\n"
        "}",
    )
    bullets(
        doc,
        [
            "Legacy AIChatService retained for GENERAL_SUPPORT chat until deprecation.",
            "Feature flag: careconnect.ai.ask.enabled gates new service on client + server.",
            "Shared response renderer for text and voice sends.",
        ],
    )

    heading(doc, "7.3 UX states during voice capture", 2)
    table(
        doc,
        ["State", "UI", "Behavior"],
        [
            ["idle", "Mic icon outline", "Tap to start"],
            ["listening", "Pulsing mic + 'Listening…'", "Partial transcript in input field"],
            ["processing", "Spinner", "STT finalizing or API in flight"],
            ["error", "Snackbar", "Permission / timeout / empty transcript"],
            ["held", "Held banner + poll", "Same as text Tier 2 UX"],
        ],
    )

    heading(doc, "7.4 Accessibility", 2)
    bullets(
        doc,
        [
            "Mic button: semantic label 'Ask by voice'; min touch target 48dp.",
            "Optional TTS read-back of answer via flutter_tts — client-side, post-response.",
            "Voice input not required — text always available (REQ-SC-1 parity).",
        ],
    )

    # 8 Telehealth distinction
    heading(doc, "8. Telehealth STT vs Ask AI Voice Query", 1)
    table(
        doc,
        ["Dimension", "Telehealth call STT", "Ask AI voice query"],
        [
            ["Who speaks", "Patient + care team on call", "Patient/caregiver asking AI"],
            ["STT location", "Chime client + AWS Transcribe", "Client speech_to_text only"],
            ["Backend ingest", "POST /calls/{id}/transcript/segments", "POST /api/ai/ask (text only)"],
            ["Persistence", "call_transcript_segments → index", "Query not indexed (audit log only)"],
            ["Retrieval role", "SOURCE content for future answers", "QUERY that triggers retrieval"],
            ["Team ownership", "Team A / telehealth", "Team E Ask AI gateway"],
        ],
        highlight_cells={(3, 1), (3, 2), (4, 1), (4, 2)},
    )
    para(
        doc,
        "Confusing these paths is a common integration error — document explicitly in code comments "
        "on VoiceCaptureService vs CallTranscriptService.",
        highlight=True,
    )

    # 9 Alexa optional
    heading(doc, "9. Alexa & Alternate Voice Channels (Future)", 1)
    bullets(
        doc,
        [
            "Today: AlexaController handles calendar/tasks — no Ask AI retrieval.",
            "P3 option: Alexa natural-language intent → map to AiAskRequest → same orchestrator.",
            "Alexa would set inputModality=VOICE, voiceCaptureMeta.sttEngine=alexa_asr.",
            "RBAC: Alexa linked patient only; same RetrievalScopeService checks.",
        ],
        highlight_indices={0},
    )

    # 10 Security RBAC
    heading(doc, "10. Security, RBAC & Privacy", 1)
    bullets(
        doc,
        [
            "USE_AI_FEATURES permission required — must be granted to PATIENT (Task 2.2 gap today).",
            "Raw audio never uploaded for Ask AI MVP — reduces PHI exposure and BAA surface.",
            "Client STT may use platform cloud (onDevice: false) — document in privacy notice.",
            "Audit log: query text + inputModality + sttEngine; no audio retention.",
            "Caregiver voice query: same consent + RetrievalScopeService as text.",
            "Rate limit: count voice and text requests equally per user/session.",
        ],
        highlight_indices={0, 1, 2},
    )

    # 11 Latency budget
    heading(doc, "11. Latency Budget (NFR-AI-1)", 1)
    table(
        doc,
        ["Segment", "Target p95", "Owner"],
        [
            ["Client STT capture + finalize", "≤ 2000 ms", "Frontend / platform"],
            ["Network RTT", "≤ 200 ms", "Infrastructure"],
            ["Backend orchestration (scope → retrieval → LLM → safety)", "≤ 2500 ms", "Team E"],
            ["Client render", "≤ 300 ms", "Frontend"],
            ["Total end-to-end", "≤ 5000 ms", "Joint SLA"],
        ],
        highlight_rows={4},
    )
    bullets(
        doc,
        [
            "Show interim 'Searching your records…' immediately on API send — not after STT.",
            "If captureDurationMs + backend processingTimeMs > 5000, log SLA breach metric.",
            "Medication timeline queries may use structured prefilter — same budget as text.",
        ],
    )

    # 12 Config
    heading(doc, "12. Configuration", 1)
    table(
        doc,
        ["Property", "Default", "Layer"],
        [
            ["careconnect.ai.ask.enabled", "false", "Backend — gates AiAskController"],
            ["careconnect.ai.ask.voice.enabled", "true", "Backend — reject VOICE if false"],
            ["careconnect.flutter.ask_ai_voice.enabled", "true", "Frontend — show mic button"],
            ["careconnect.ai.ask.max-query-length", "2000", "Backend"],
            ["careconnect.flutter.voice.listen-seconds", "12", "Frontend — VoiceCaptureService"],
            ["careconnect.flutter.voice.locale", "en_US", "Frontend — override from user profile"],
        ],
    )

    # 13 Tests
    heading(doc, "13. Test Alignment", 1)
    table(
        doc,
        ["Test", "Validation"],
        [
            ["SRS §3.2 voice=text pipeline", "Same orchestrator mock; inputModality only in audit"],
            ["TC-E-AI-011 (med timeline)", "Voice query 'when did I start lisinopril' → same response as text"],
            ["TC-E-SC-001", "Voice medication-change query → Tier 2 hold identical to text"],
            ["Unit: VoiceQueryNormalizer", "Empty, filler strip, max length"],
            ["Unit: VoiceQueryPreflight", "Server rejects blank query"],
            ["Integration: AskAiService", "POST with inputModality=VOICE → 200 + citations"],
            ["E2E Flutter", "Mic tap → mock STT → API → answer rendered"],
            ["Audit", "ai_safety_audit_event includes inputModality=VOICE"],
        ],
        highlight_rows={1, 2},
    )

    # 14 Gap analysis
    heading(doc, "14. Codebase Gap Analysis", 1)
    table(
        doc,
        ["Area", "Today", "Design target"],
        [
            ["Ask AI voice UI", "Text input only in ai_chat_improved.dart", "AskAiVoiceInput + mic button"],
            ["Ask AI API", "POST /v1/api/ai-chat/chat — raw Bedrock", "POST /v1/api/ai/ask unified gateway"],
            ["STT for queries", "VoiceCommandAI on forms only", "VoiceCaptureService on Ask AI"],
            ["Pipeline fork", "N/A", "No fork — inputModality metadata only"],
            ["Audit modality", "ChatAuditService — no inputModality", "AiAuditService with VOICE field"],
            ["Permission", "USE_AI_FEATURES missing for PATIENT", "Task 2.2 before voice launch"],
            ["Legacy chat", "AIChatService → MedicalContextService chart dump path", "Deprecated for records Q&A"],
        ],
        highlight_rows={1, 2, 4, 5},
    )

    # 15 Implementation phases
    heading(doc, "15. Implementation Phases", 1)
    table(
        doc,
        ["Phase", "Deliverable", "Depends on"],
        [
            ["P0", "AiAskController + orchestrator (text path)", "Tier 5 backlog 5.3"],
            ["P0", "AskAiService Flutter client (text)", "API contract DTOs"],
            ["P1", "VoiceCaptureService + AskAiVoiceInput mic button", "P0 gateway live"],
            ["P1", "inputModality=VOICE + voiceCaptureMeta audit", "AiAuditService"],
            ["P1", "VoiceQueryNormalizer + server preflight", "—"],
            ["P2", "Locale es-US STT + disclaimer parity", "SRS Spanish NFR"],
            ["P2", "SLA metrics split STT vs backend", "Telemetry"],
            ["P3", "Alexa → AiAskOrchestrator bridge", "Alexa NLU design"],
            ["P3", "Optional TTS read-back", "Accessibility"],
        ],
        highlight_rows={2, 3, 4},
    )
    para(
        doc,
        "Task backlog: 7.1 mic → speech_to_text → text pipeline; 7.3 voice through same SCC-3. "
        "Do NOT ship voice before Tier 5 gateway + Tier 6 safety/HITL complete.",
        highlight=True,
    )

    # 16 Sequence
    heading(doc, "16. Sequence Diagram (Voice Query)", 1)
    code(
        doc,
        "User     AskAiVoiceInput   VoiceCapture   AskAiService   AiAskController   Orchestrator\n"
        " |            |                 |              |               |                |\n"
        " | tap mic    |                 |              |               |                |\n"
        " |----------->| capture()       |              |               |                |\n"
        " |            |---------------->| STT listen   |               |                |\n"
        " |            |<----------------| transcript   |               |                |\n"
        " |            | normalize()     |              |               |                |\n"
        " |            | ask(VOICE)      |              |               |                |\n"
        " |            |-------------------------------->| POST /ask     |                |\n"
        " |            |                 |              |-------------->| ask(req)       |\n"
        " |            |                 |              |               |--------------->|\n"
        " |            |                 |              |               |  (same 8 stages|\n"
        " |            |                 |              |               |   as text)     |\n"
        " |            |                 |              |<-- AiAskResponse               |\n"
        " |<-- answer + citations --------|              |               |                |",
    )

    # 17 Related docs
    heading(doc, "17. Related Documents", 1)
    bullets(
        doc,
        [
            "docs/Voice_Query_Path_and_STT_Framework_Dependencies.docx (STT inventory & research)",
            "docs/POST_api_ai_ask_Request_Response_Contract_Design.docx (inputModality field)",
            "docs/Retrieval_Orchestration_RBAC_Hybrid_Context_LLM_Design.docx (orchestrator stages)",
            "docs/Secondary_Validation_and_Tier2_HITL_Hold_Release_Design.docx (same safety path)",
            "docs/Medication_Timeline_Aggregation_in_Retrieval_Responses_Design.docx (voice uses same aggregator)",
            "docs/Call_Transcript_Retrieval_Review.docx (telehealth STT upstream — not query path)",
            "docs/Team_E_Implementation_Task_Backlog.docx (Tasks 7.1–7.3, 5.3, 2.2)",
            "frontend/lib/features/ai/presentation/pages/voice_command_ai.dart",
            "frontend/lib/widgets/ai_chat_improved.dart",
            "frontend/lib/services/ai_chat_service.dart",
        ],
    )

    heading(doc, "18. Conclusion", 1)
    para(
        doc,
        "Voice query routing is a thin client adapter over the existing Ask AI retrieval pipeline. "
        "Client-side speech_to_text converts spoken questions to text; POST /api/ai/ask with "
        "inputModality=VOICE invokes the same AiAskOrchestrator path as typed queries — hybrid "
        "retrieval, medication timeline aggregation, SafetyPipeline, Tier 2 HITL, citations, "
        "and disclaimers are unchanged. Telehealth and notetaker STT remain upstream index "
        "producers, not alternate query gateways.",
    )
    para(
        doc,
        "Build order: complete Tier 5 /api/ai/ask (text) → Tier 6 safety/HITL → VoiceCaptureService "
        "+ mic UI (Task 7.1) → audit modality (7.3). Never implement a backend audio upload path "
        "for MVP — it duplicates STT ownership and violates FR-AI-9 minimum-necessary.",
        highlight=True,
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    try:
        doc.save(OUTPUT)
        print(f"Created: {OUTPUT}")
    except PermissionError:
        alt = OUTPUT.with_name(f"{OUTPUT.stem}_refresh{OUTPUT.suffix}")
        doc.save(alt)
        print(f"Created (original locked): {alt}")


if __name__ == "__main__":
    build()
