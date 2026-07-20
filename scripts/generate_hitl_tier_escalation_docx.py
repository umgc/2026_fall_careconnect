"""Generate Word document: HITL / Tier 1-2 escalation triggers and hold-release workflow."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX

ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "HITL_Tier1_Tier2_Escalation_and_Hold_Release_Workflow.docx"


def heading(doc, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def para(doc, text: str, bold: bool = False, highlight: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    if highlight:
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def bullets(doc, items: list[str], highlight_indices: set[int] | None = None) -> None:
    highlight_indices = highlight_indices or set()
    for i, item in enumerate(items):
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        if i in highlight_indices:
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def code(doc, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"


def table(
    doc,
    headers: list[str],
    rows: list[list[str]],
    highlight_cells: set[tuple[int, int]] | None = None,
) -> None:
    highlight_cells = highlight_cells or set()
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for r_idx, row in enumerate(rows, start=1):
        cells = t.add_row().cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = val
            if (r_idx, c_idx) in highlight_cells:
                for p in cells[c_idx].paragraphs:
                    for run in p.runs:
                        run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def build() -> None:
    doc = Document()

    title = doc.add_heading(
        "HITL / Tier 1–2 Escalation Triggers and Hold–Release Workflow", 0
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para(doc, "CareConnect Team E — Safety, Consent & Clarity (Section 7)")
    para(
        doc,
        "Synthesizes SRS v2.0, Milestone 2 TDD, Software Test Plan, hybrid retrieval scope, "
        "RBAC research, medication timeline (FR-AI-11), and codebase review.",
    )
    para(
        doc,
        "Yellow highlights mark gaps between Team E design and current implementation.",
        highlight=True,
    )
    doc.add_paragraph()

    # 1 Executive Summary
    heading(doc, "1. Executive Summary", 1)
    para(
        doc,
        "CareConnect’s Safety, Consent & Clarity framework requires a secondary validation pass on "
        "every AI output before delivery. Outputs are classified as Tier 1 (deliver with notice) or "
        "Tier 2 (hold for Human-in-the-Loop review before the care recipient sees them). Tier 2 items "
        "must not reach the patient until a Human Reviewer / Clinician releases or rejects them.",
    )
    para(
        doc,
        "GAP: No tier classifier, HITL queue, hold state, or release workflow exists in the codebase. "
        "BedrockAIChatService returns raw model text with no escalation flag. GuardrailService blocks "
        "some phrases but does not route to HITL. Audit logging does not meet REQ-SC-9 immutable ledger requirements.",
        highlight=True,
    )

    # 2 Requirement sources
    heading(doc, "2. Requirement & Design Sources", 1)
    table(
        doc,
        ["Source", "Reference", "HITL relevance"],
        [
            ["SRS §7", "Safety, Consent & Clarity", "Tier table, secondary validation, audit"],
            ["SRS §3", "UC-AI-6 Escalate High-Risk Query (HITL)", "Use case + Fig 3.14–3.15"],
            ["SRS §3", "FR-AI-5", "Detect high-risk queries; human review before delivery"],
            ["SRS §7", "REQ-SC-4", "Hold Tier 2 for human review before release"],
            ["SRS §7", "REQ-SC-5 / REQ-SC-6", "Confirmation before action; timeout ≠ approval"],
            ["SRS §7", "REQ-SC-9", "Immutable audit of escalation + confirmation events"],
            ["SRS §3", "FR-AI-10", "Log query, response, citations, escalation, delivery status"],
            ["TDD §3.3", "SCC-3 19-step control sequence", "RBAC → consent → retrieval → validate → HITL → audit"],
            ["TDD", "Safety/Consent Layer", "Secondary validation; holds Tier 2 for HITL"],
            ["Test Plan", "TC-E-SC-001", "Medication-change → held for human review"],
            ["Test Plan", "TC-AI-06 / TC-SCC-03", "High-risk query → HITL routing + audit"],
            ["Test Plan", "FT-009", "Tier 2 human review functional test"],
        ],
    )

    # 3 Tier definitions
    heading(doc, "3. Tier 1 vs Tier 2 Definitions", 1)
    table(
        doc,
        ["Tier", "Delivery rule", "User experience", "Reviewer action"],
        [
            [
                "Tier 1",
                "Deliver immediately with mandatory notices",
                "Answer shown with disclaimer + optional confirm-with-provider prompt",
                "None required before delivery",
            ],
            [
                "Tier 2",
                "HOLD — do not deliver to care recipient until human release",
                "“We’re reviewing this before showing it to you.” (TDD copy)",
                "Human Reviewer / Clinician approves, edits, or rejects",
            ],
        ],
        highlight_cells={(2, 1), (2, 3)},
    )
    para(
        doc,
        "REQ-SC-4: The system shall run a secondary validation pass on all AI output before delivery "
        "and shall hold Tier 2 output for human review before it reaches the care recipient.",
        bold=True,
    )

    # 4 Trigger catalog
    heading(doc, "4. Escalation Trigger Catalog (SRS Table 8)", 1)
    para(doc, "Secondary validation classifies output using these working rules:")
    table(
        doc,
        ["Trigger pattern", "Tier", "Required behavior"],
        [
            [
                "General medication mention",
                "Tier 1",
                "Deliver with medical disclaimer + logged escalation prompt (confirm-with-provider)",
            ],
            [
                "Medication-change request / instruction",
                "Tier 2",
                "HOLD; route to human review before care recipient sees it (TC-E-SC-001)",
            ],
            [
                "New diagnosis / condition language",
                "Tier 1",
                "Answer from cited records only; escalation prompt",
            ],
            [
                "Emergency symptom language",
                "Tier 2",
                "HOLD; surface immediate crisis guidance; human review",
            ],
            [
                "Dosage calculation request",
                "Tier 2",
                "Do not compute; direct to healthcare provider; hold if AI draft exists",
            ],
            [
                "Clinical language beyond cited source",
                "Tier 2",
                "Secondary validation flags unsupported meds/dosages/dates (TDD safety layer)",
            ],
            [
                "Missing uncertainty on low-confidence extraction",
                "Tier 1*",
                "Deliver with needsConfirmation / review badge; may upgrade to Tier 2",
            ],
            [
                "No matching records",
                "Tier 1",
                "Clear no-records message; no general medical answer (FR-AI-2 path)",
            ],
            [
                "Unauthorized cross-patient query",
                "Block",
                "403 + audit; no AI output (RBAC, not tiered delivery)",
            ],
        ],
        highlight_cells={
            (2, 1), (2, 2),
            (3, 1), (3, 2),
            (5, 1), (5, 2),
            (6, 1), (6, 2),
        },
    )
    para(
        doc,
        "Additional Tier 2 triggers from alternate flows: summary medication-change items pending "
        "confirmation, FR-AI-11 timeline answers that recommend initiation/termination without cited "
        "source support, and USPS OCR below confidence threshold (manual review per TDD).",
    )

    # 5 Secondary validation
    heading(doc, "5. Secondary Validation Pass (Pre-Tier Classification)", 1)
    para(doc, "Runs on all AI output before tier assignment (SRS §7 + TDD Safety Layer):", bold=True)
    bullets(
        doc,
        [
            "Detect clinical language that goes beyond the cited source passage.",
            "Identify medication names, dosages, or dates not present in cited sources.",
            "Catch missing uncertainty flags on low-confidence extractions (needsConfirmation).",
            "Verify every grounded claim has a citation (FR-AI-2 / REQ-SC-2).",
            "Attach standard records-based disclaimer and non-dismissible AI notice (REQ-SC-1).",
            "Never produce diagnosis/treatment/dosage advice not explicitly in a cited source (REQ-SC-3).",
        ],
    )
    para(
        doc,
        "Failure modes: HTTP 422 when schema/safety validation fails — withhold output rather than "
        "return unsafe content (TDD error contract).",
    )

    # 6 Hold-release workflow
    heading(doc, "6. Hold–Release Workflow (Target State)", 1)

    heading(doc, "6.1 End-to-end flow (Ask AI + Summaries)", 2)
    code(
        doc,
        "1. User query (text/voice) → AI Gateway\n"
        "2. RBAC + consent scope (RetrievalScopeService)\n"
        "3. Hybrid retrieval → assemble minimum-necessary context\n"
        "4. Bedrock inference → draft answer + citations[]\n"
        "5. Secondary validation pass → assign tier (1 or 2)\n"
        "6. IF Tier 1 → deliver with disclaimer + escalation prompt metadata\n"
        "7. IF Tier 2 → persist held_item; return held status to client (no answer body to patient)\n"
        "8. Human Reviewer queue → review / edit / approve / reject\n"
        "9. ON RELEASE → deliver approved text; log reviewer + timestamp\n"
        "10. ON REJECT → notify user with safe fallback; log reason\n"
        "11. Immutable audit entry for query, draft, tier, hold, release/reject, delivery status",
    )

    heading(doc, "6.2 Hold state model (recommended)", 2)
    table(
        doc,
        ["Field", "Purpose"],
        [
            ["held_item_id", "UUID primary key"],
            ["patient_id", "Row-level scope"],
            ["user_id / session_id", "Originating query context"],
            ["source_surface", "ASK_AI | CALL_SUMMARY | VISIT_SUMMARY | USPS | STML"],
            ["tier", "2"],
            ["status", "PENDING_REVIEW | APPROVED | REJECTED | EXPIRED"],
            ["trigger_codes", "e.g. MEDICATION_CHANGE, EMERGENCY_SYMPTOM, DOSAGE_CALC"],
            ["draft_answer", "Withheld AI text (PHI-minimized storage policy)"],
            ["citations_json", "Source excerpts supporting draft"],
            ["reviewer_user_id", "Human who released/rejected"],
            ["reviewed_at", "Release/reject timestamp"],
            ["review_notes", "Optional clinician comment"],
            ["delivery_status", "HELD | DELIVERED | WITHHELD_PERMANENTLY"],
        ],
        highlight_cells={(1, 0), (6, 0), (12, 0)},
    )

    heading(doc, "6.3 Release paths", 2)
    table(
        doc,
        ["Action", "Actor", "Outcome"],
        [
            ["Approve as-is", "Human Reviewer / Clinician", "Deliver draft to care recipient; audit DELIVERED"],
            ["Approve with edits", "Human Reviewer", "Deliver edited text; store both draft and final"],
            ["Reject", "Human Reviewer", "User sees safe message; no clinical draft delivered"],
            ["Escalate to provider", "System (Tier 2 emergency)", "Crisis guidance + provider contact prompt"],
            ["Expire / auto-close", "System policy", "HELD → EXPIRED; user notified to re-ask or contact provider"],
        ],
    )
    para(
        doc,
        "REQ-SC-6: Timeout, app closure, or inaction is NEVER approval. Held items remain held until "
        "explicit reviewer action.",
        highlight=True,
    )

    heading(doc, "6.4 Client UX copy (TDD)", 2)
    bullets(
        doc,
        [
            "Tier 2 held: “We’re reviewing this before showing it to you.”",
            "Tier 1 medication: confirm-with-provider prompt (escalation: confirm-with-provider)",
            "Success response includes: answer, citations[], disclaimer, escalation flag",
            "Held response: status=held, heldItemId, no answer body to patient",
        ],
    )

    # 7 SCC-3
    heading(doc, "7. SCC-3 Control Sequence (19 Steps)", 1)
    para(doc, "TDD guarded request path integrating HITL:")
    bullets(
        doc,
        [
            "1–4. Authenticate JWT; resolve role; apply RBAC patient scope",
            "5–6. Apply consent + REQ-SC-7 source exclusions",
            "7–10. Hybrid retrieval; assemble prompt; invoke Bedrock",
            "11–14. Schema validation; secondary validation; tier assignment",
            "15–16. IF Tier 2 → write held_item; return held — STOP delivery",
            "17–18. IF Tier 1 → attach disclaimer; deliver to client",
            "19. Immutable audit log (query, tier, escalation, delivery status)",
        ],
        highlight_indices={4, 5},
    )

    # 8 Confirmation workflow
    heading(doc, "8. Confirmation Workflow (Separate from HITL Hold)", 1)
    para(
        doc,
        "REQ-SC-5 requires confirmation before any tool call, record write, or side-effecting action. "
        "This is distinct from Tier 2 hold — a Tier 1 answer may still require approve-once / "
        "approve-for-session / decline before calendar/reminder/care-plan writes (FR-AI-8, FR-SUM-4).",
    )
    table(
        doc,
        ["Step", "Behavior"],
        [
            ["Prompt", "Approve once | Approve for session | Decline"],
            ["Persist", "Confirmation/dismissal stored across restarts (REQ-SC-6)"],
            ["Timeout", "No write; no implicit approval (TC-E-SC-002)"],
            ["Audit", "Log confirmation event immutably (REQ-SC-9)"],
        ],
    )

    # 9 API contract
    heading(doc, "9. API Response Contract", 1)
    code(
        doc,
        "// Tier 1 success (200)\n"
        "{\n"
        '  "success": true,\n'
        '  "answer": "...",\n'
        '  "citations": [{ "type": "transcript", "id": "...", "excerpt": "..." }],\n'
        '  "disclaimer": "This is based on your stored records and is not medical advice.",\n'
        '  "escalation": "confirm-with-provider",\n'
        '  "tier": 1,\n'
        '  "deliveryStatus": "DELIVERED"\n'
        "}\n\n"
        "// Tier 2 hold (200 or 202 — held, not an error)\n"
        "{\n"
        '  "success": true,\n'
        '  "held": true,\n'
        '  "heldItemId": "uuid",\n'
        '  "tier": 2,\n'
        '  "message": "We\'re reviewing this before showing it to you.",\n'
        '  "triggerCodes": ["MEDICATION_CHANGE"],\n'
        '  "deliveryStatus": "HELD"\n'
        "}\n\n"
        "// Safety validation failure (422)\n"
        "{\n"
        '  "success": false,\n'
        '  "errorCode": "SAFETY_VALIDATION_FAILED",\n'
        '  "deliveryStatus": "WITHHELD"\n'
        "}",
    )

    heading(doc, "9.1 Reviewer API (planned)", 2)
    bullets(
        doc,
        [
            "GET /api/ai/hitl/queue — list PENDING_REVIEW items (reviewer role only)",
            "GET /api/ai/hitl/{heldItemId} — draft + citations + trigger context",
            "POST /api/ai/hitl/{heldItemId}/release — approve (optional editedAnswer)",
            "POST /api/ai/hitl/{heldItemId}/reject — reject with reason",
            "GET /api/ai/hitl/{heldItemId}/status — patient polls until DELIVERED or REJECTED",
        ],
        highlight_indices={0, 1, 2, 3, 4},
    )

    # 10 Roles
    heading(doc, "10. Roles & RBAC for HITL", 1)
    table(
        doc,
        ["Role", "HITL permissions"],
        [
            ["Care Recipient (Patient)", "Receives Tier 1; never sees Tier 2 draft until released"],
            ["Caregiver", "May submit queries; Tier 2 holds apply; subject to REQ-SC-8 consent"],
            ["Human Reviewer / Clinician", "View queue; release/reject held items (new role or ADMIN)"],
            ["System Administrator", "Configure safety rules, tier thresholds, retention"],
            ["Family Member", "Read-only where permitted; no HITL reviewer access"],
        ],
        highlight_cells={(3, 0), (3, 1)},
    )
    para(
        doc,
        "GAP: Human Reviewer / Clinician role not defined in RolePermissionService today.",
        highlight=True,
    )

    # 11 Test alignment
    heading(doc, "11. Test Plan Alignment", 1)
    table(
        doc,
        ["Test ID", "Scenario", "Pass criteria"],
        [
            ["TC-E-SC-001", "Medication-change query/output", "Held; routed to human review before patient sees it"],
            ["TC-E-SC-002", "Write confirmation timeout", "No write; timeout ≠ approval"],
            ["TC-AI-06", "High-risk / medical query (UC-AI-6)", "Escalated to HITL; audit entry recorded"],
            ["TC-SCC-03", "Safety assigns Tier 2", "Held + routed; user told item is held"],
            ["FT-009", "Tier 2 human review", "High-risk output held before care recipient"],
            ["TC-E-AI-002", "No matching records", "No-records message; no general medical answer"],
        ],
    )

    # 12 Codebase gap
    heading(doc, "12. Codebase Gap Analysis", 1)
    table(
        doc,
        ["Component", "Location", "Current behavior", "HITL gap"],
        [
            [
                "AI chat",
                "BedrockAIChatService",
                "Raw Bedrock response; no tier, no hold, no escalation flag",
                "No secondary validation or HITL routing",
            ],
            [
                "Guardrails",
                "GuardrailService",
                "SSN redaction; blocks forbidden phrases (throws exception)",
                "Binary block — not Tier 1/2 classification or hold queue",
            ],
            [
                "Response sanitization",
                "ResponseSanitizationService",
                "Strips system info / sensitive patterns",
                "Not wired to Bedrock chat path; no tier logic",
            ],
            [
                "Governance",
                "LangChainGovernanceService",
                "Rate limits + message length",
                "No clinical risk classification",
            ],
            [
                "Audit",
                "ChatAuditService",
                "Metadata-only logs (no content, no escalation)",
                "Does not meet REQ-SC-9 immutable ledger for tier/hold/release",
            ],
            [
                "HITL queue",
                "—",
                "Not implemented",
                "No held_item table, API, or reviewer UI",
            ],
            [
                "Summaries",
                "BedrockSentimentService",
                "Simple summary JSON; no confirm/dismiss write gate",
                "FR-SUM-4 / Tier 2 summary items not enforced",
            ],
            [
                "EVV review (analog)",
                "EvvController.review",
                "Human approve/reject for EVV records",
                "Pattern exists but not connected to AI HITL",
            ],
        ],
        highlight_cells={
            (1, 3), (2, 3), (5, 3), (6, 3), (7, 3),
        },
    )
    para(
        doc,
        "Note: careconnect.ai.enabled=false in dev disables AIChatController entirely — HITL cannot "
        "be exercised locally without opt-in (CARECONNECT_AI_ENABLED=true).",
        highlight=True,
    )

    # 13 Implementation roadmap
    heading(doc, "13. Recommended Implementation Roadmap", 1)
    table(
        doc,
        ["Phase", "Deliverable", "Requirements satisfied"],
        [
            ["P0", "SafetyValidationService + TierClassifier (Table 8 rules)", "REQ-SC-4, FR-AI-5"],
            ["P0", "held_item schema + HITL queue API", "UC-AI-6, TC-E-SC-001"],
            ["P0", "Integrate into AI Gateway before Bedrock response delivery", "SCC-3 steps 11–16"],
            ["P1", "Human Reviewer role + release/reject endpoints", "Hold–release workflow"],
            ["P1", "Client held-state UX + polling/WebSocket", "TDD Tier 2 copy"],
            ["P1", "Extend to summary confirm/dismiss + Tier 2 summary items", "FR-SUM-4, TC-E-SUM-002"],
            ["P2", "Immutable audit ledger (write-once) for escalation events", "REQ-SC-9, FR-AI-10"],
            ["P2", "Spanish safety copy review-verified", "SRS NFR localization"],
            ["P3", "Automate TC-AI-06 / TC-SCC-03 integration tests", "Milestone 3 test automation"],
        ],
        highlight_cells={(1, 0), (2, 0), (3, 0)},
    )

    heading(doc, "13.1 TierClassifier pseudocode", 2)
    code(
        doc,
        "classify(draft, citations, queryContext) -> TierResult:\n"
        "  if unauthorized_scope: throw 403\n"
        "  if unsupported_clinical_claim(draft, citations): return Tier2(MEDICATION_CHANGE or UNSUPPORTED_CLAIM)\n"
        "  if matches(query, EMERGENCY_SYMPTOM_PATTERNS): return Tier2(EMERGENCY_SYMPTOM)\n"
        "  if matches(query, DOSAGE_CALC_PATTERNS): return Tier2(DOSAGE_CALC)\n"
        "  if medication_change_language(draft or query): return Tier2(MEDICATION_CHANGE)\n"
        "  if general_medication_mention: return Tier1(escalation=confirm-with-provider)\n"
        "  if no_matching_records: return Tier1(NO_RECORDS_MESSAGE)\n"
        "  return Tier1()",
    )

    # 14 Cross-workstream
    heading(doc, "14. Integration with Other Team E Workstreams", 1)
    bullets(
        doc,
        [
            "Hybrid retrieval: citations[] from retrieval layer feed secondary validation (unsupported-claim detection).",
            "FR-AI-11 medication timeline: dose-change events in answers trigger Tier 2 if not fully cited.",
            "RBAC: RetrievalScopeService runs before validation; caregiver Tier 2 holds respect REQ-SC-8.",
            "Voice query: same tier pipeline for voice→text queries (SRS §3.2).",
            "USPS mail: OCR confidence below threshold → manual review (parallel HITL pattern).",
            "STML Daily Memory Brief: inherits Safety layer; no bypass of Tier 2 holds.",
        ],
    )

    # 15 Related docs
    heading(doc, "15. Related Documents & Code References", 1)
    bullets(
        doc,
        [
            "C:\\Users\\ravic\\Downloads\\CareConnect_SRS_Revision 2.0_TEAM E.docx (§7, Table 8, UC-AI-6)",
            "C:\\Users\\ravic\\Downloads\\CareConnect_Milestone_2_TDD_TEAM E.docx (SCC-3, Safety Layer)",
            "C:\\Users\\ravic\\Downloads\\CareConnect_Milestone_2_Software_Test_Plan (2).docx",
            "docs/Hybrid_Retrieval_Scope_PostgreSQL_FullText_pgvector.docx",
            "docs/RBAC_Scoped_Retrieval_Source_Types.docx",
            "docs/Medication_Timeline_Retrieval_FR-AI-11.docx",
            "backend/core/src/main/java/com/careconnect/service/chat/GuardrailService.java",
            "backend/core/src/main/java/com/careconnect/service/BedrockAIChatService.java",
            "backend/core/src/main/java/com/careconnect/service/ChatAuditService.java",
            "backend/core/src/main/java/com/careconnect/service/security/ResponseSanitizationService.java",
            "backend/core/src/main/java/com/careconnect/controller/EvvController.java (review pattern analog)",
        ],
    )

    heading(doc, "16. Conclusion", 1)
    para(
        doc,
        "HITL is a mandatory gate in Team E’s Safety, Consent & Clarity framework — not an optional "
        "overlay. Tier 1 delivers with disclaimers and logged escalation prompts; Tier 2 must be held "
        "until a Human Reviewer explicitly releases or rejects. The hold–release workflow spans "
        "classification, queue persistence, reviewer APIs, patient-held UX, and immutable audit.",
    )
    para(
        doc,
        "Priority: implement SafetyValidationService + held_item queue before expanding Ask AI or "
        "medication timeline (FR-AI-11) features — otherwise Tier 2 triggers cannot be enforced.",
        highlight=True,
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    try:
        doc.save(OUTPUT)
        print(f"Created: {OUTPUT}")
    except PermissionError:
        alt = OUTPUT.with_name(f"{OUTPUT.stem}_refresh{OUTPUT.suffix}")
        doc.save(alt)
        print(f"Created (original locked): {alt}")


if __name__ == "__main__":
    build()
