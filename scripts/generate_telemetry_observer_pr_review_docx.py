"""Generate Word document: PR code review for TelemetryGoRouterObserver feature branch."""

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.shared import Pt

OUTPUT = (
    Path(__file__).resolve().parents[1]
    / "docs"
    / "PR_Review_Telemetry_GoRouter_Observer_feature_a-dkinchen-telemetry-expansion.docx"
)


def heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def para(doc: Document, text: str, bold: bool = False, highlight: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    if highlight:
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def bullets(doc: Document, items: list[str], highlight_indices: set[int] | None = None) -> None:
    highlight_indices = highlight_indices or set()
    for i, item in enumerate(items):
        p = doc.add_paragraph(item, style="List Bullet")
        if i in highlight_indices:
            for run in p.runs:
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def code(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    for i, h in enumerate(headers):
        tbl.rows[0].cells[i].text = h
    for row in rows:
        cells = tbl.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val


def build() -> Document:
    doc = Document()
    title = doc.add_heading("PR Code Review", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    para(doc, "TelemetryGoRouterObserver — feature/a-dkinchen-telemetry-expansion → team-ae-develop")
    para(doc, f"Review date: {date.today().isoformat()}")
    para(doc, "WBS 2.2.2 — automatic screen_view tracking via GoRouter observer")
    doc.add_paragraph()

    heading(doc, "Review metadata", 1)
    table(
        doc,
        ["Field", "Value"],
        [
            ["Source branch", "feature/a-dkinchen-telemetry-expansion"],
            ["Target branch", "team-ae-develop"],
            ["Scope", "2 files changed (+297 / −1 lines)"],
            ["Feature commits", "7004b11, 13ea6d3"],
            ["Note", "CODEOWNERS .gitattributes commits (2a92894 / 970caf4) net to zero diff"],
        ],
    )

    heading(doc, "1. Change summary", 1)
    para(
        doc,
        "This PR adds automatic navigation telemetry so every GoRouter transition emits a "
        "screen_view event without requiring manual Telemetry.event(...) calls on each page.",
    )

    heading(doc, "1.1 What was added", 2)
    table(
        doc,
        ["Change", "Purpose"],
        [
            ["TelemetryGoRouterObserver", "NavigatorObserver on appRouter — logs screen_view on push/pop/replace/remove"],
            ["_appRouterRef", "Global ref so observer can read GoRouter.state.uri after navigation"],
            ["telemetry_go_router_observer_test.dart", "Widget tests for push, navigation, opt-out, error resilience"],
        ],
    )

    heading(doc, "1.2 Event shape", 2)
    code(
        doc,
        "await Telemetry.event('screen_view', {'screen': location});\n"
        "// location = router.state.uri.toString()  e.g. \"/settings\", \"/login?redirect=...\"",
    )

    heading(doc, "1.3 Integration", 2)
    code(
        doc,
        "final GoRouter appRouter = _appRouterRef = GoRouter(\n"
        "  initialLocation: '/',\n"
        "  observers: [_telemetryGoRouterObserver],\n"
        "  routes: [ ... ],\n"
        ");",
    )

    heading(doc, "2. Bug and risk analysis", 1)

    heading(doc, "2.1 High — URI length vs guardrails (events silently degraded)", 2)
    para(
        doc,
        "TelemetryGuardrails drops string values longer than 64 characters. Many GoRouter locations "
        "exceed 64 chars when they include IDs or query params (e.g. patient-check-in-detail routes, "
        "oauth callback). The screen key is stripped from details; the event may still POST with empty "
        "details — looks successful but carries no screen info.",
        highlight=True,
    )
    code(doc, "if (v is String && v.length > 64) continue;  // telemetry_guardrails.dart")

    heading(doc, "2.2 Medium — Duplicate screen_view events", 2)
    para(doc, "Manual logging already exists in SettingsPage ('settings'), MenuPage ('menu_page'), and MainScreen bottom nav (semantic names). The observer also fires on the same navigation with the full URI.")
    bullets(
        doc,
        [
            "Risk: inflated analytics counts",
            "Risk: inconsistent screen schema (path vs semantic name)",
        ],
    )

    heading(doc, "2.3 Medium — Out-of-order async telemetry (race)", 2)
    para(
        doc,
        "Each observer callback calls unawaited(_logScreenView()) with no debounce or generation token. "
        "Fast navigation A → B → C can emit concurrent requests that complete out of order on the backend.",
    )

    heading(doc, "2.4 Low — Query strings may carry sensitive data", 2)
    para(
        doc,
        "Logging router.state.uri.toString() includes query parameters. Guardrails block PII keys in "
        "property names, not values. Prefer uri.path only or sanitized route names.",
    )

    heading(doc, "2.5 Low — Other issues", 2)
    bullets(
        doc,
        [
            "StateError on router.state is silently swallowed (no release debug signal)",
            "Observer lives in app_router.dart (~1,150 lines) — coupling concern",
            "Inconsistent screen schema: URIs vs existing semantic names (settings, menu_page)",
        ],
    )

    heading(doc, "2.6 Positive / well-handled", 2)
    table(
        doc,
        ["Case", "Handling"],
        [
            ["Telemetry opt-out", "Respects TelemetrySettings — tested"],
            ["POST failure", "Caught; debugPrint; navigation continues — tested"],
            ["Test isolation", "routerProvider injection avoids production _appRouterRef"],
            ["Guardrails whitelist", "screen_view is allowed"],
            ["Non-blocking", "unawaited keeps navigation synchronous"],
        ],
    )

    heading(doc, "3. Architecture and style", 1)

    heading(doc, "3.1 Strengths", 2)
    table(
        doc,
        ["Area", "Assessment"],
        [
            ["Observer pattern", "Standard Flutter/GoRouter approach for navigation analytics"],
            ["Non-blocking", "Telemetry failures do not block UX"],
            ["Testability", "routerProvider hook + MockClient — good pattern"],
            ["Test quality", "Opt-out, failure case, clear setup"],
            ["WBS alignment", "Centralizes screen tracking vs ad-hoc per-page calls"],
        ],
    )

    heading(doc, "3.2 Concerns", 2)
    table(
        doc,
        ["Area", "Assessment"],
        [
            ["Global _appRouterRef", "Mutable singleton set during GoRouter init — fragile if multiple routers"],
            ["Co-location", "Observer belongs in features/telemetry/ not config/router/"],
            ["Incomplete migration", "Automatic tracking added without removing manual screen_view calls"],
            ["Test gaps", "No tests for didPop, URI > 64 chars, duplicate behavior"],
        ],
    )

    para(doc, "Overall: sound approach for WBS 2.2.2; needs guardrails/schema follow-up before relying on analytics data.", bold=True)

    heading(doc, "4. Recommendations", 1)

    heading(doc, "R1 — Log path only (fixes length + query leak) — required before merge", 2)
    para(doc, "Recommended fix:", bold=True, highlight=True)
    code(
        doc,
        "Future<void> _logScreenView() async {\n"
        "  final router = _activeRouter;\n"
        "  if (router == null) return;\n\n"
        "  final Uri uri;\n"
        "  try {\n"
        "    uri = router.state.uri;\n"
        "  } on StateError {\n"
        "    return;\n"
        "  }\n\n"
        "  final screen = uri.path.isEmpty ? '/' : uri.path;\n\n"
        "  try {\n"
        "    await Telemetry.event('screen_view', {'screen': screen});\n"
        "  } catch (e) {\n"
        "    debugPrint('Telemetry logging failed: $e');\n"
        "  }\n"
        "}",
    )

    heading(doc, "R2 — Debounce rapid navigation", 2)
    code(
        doc,
        "Timer? _screenViewDebounce;\n\n"
        "void _scheduleScreenView() {\n"
        "  _screenViewDebounce?.cancel();\n"
        "  _screenViewDebounce = Timer(\n"
        "    const Duration(milliseconds: 100),\n"
        "    () => unawaited(_logScreenView()),\n"
        "  );\n"
        "}",
    )

    heading(doc, "R3 — Move observer to telemetry feature module", 2)
    code(
        doc,
        "// frontend/lib/features/telemetry/telemetry_go_router_observer.dart\n"
        "class TelemetryGoRouterObserver extends NavigatorObserver { ... }\n\n"
        "// app_router.dart\n"
        "final _telemetryGoRouterObserver = TelemetryGoRouterObserver(\n"
        "  routerProvider: () => _appRouterRef,\n"
        ");",
    )

    heading(doc, "R4 — Remove duplicate manual screen_view calls (follow-up)", 2)
    bullets(
        doc,
        [
            "Remove settings_page.dart manual screen_view after observer migration",
            "Remove menu_page.dart initState screen_view",
            "For MainScreen bottom nav: keep button_tap only, or disable observer for shell routes",
            "Document canonical screen format in Team B telemetry spec",
        ],
    )

    heading(doc, "R5 — Extend guardrails for route paths", 2)
    code(
        doc,
        "if (v is String && v.length > 64) {\n"
        "  if (k == 'screen' && v.startsWith('/') && v.length <= 128) {\n"
        "    out[k] = v;\n"
        "  }\n"
        "  continue;\n"
        "}",
    )

    heading(doc, "R6 — Add missing tests", 2)
    bullets(
        doc,
        [
            "Path-only logging strips query parameters",
            "didPop logs correct screen after back navigation",
            "Long path routes — document expected guardrail behavior",
        ],
    )

    heading(doc, "R7 — PR hygiene", 2)
    bullets(
        doc,
        [
            "Squash or drop no-op CODEOWNERS commits before merge if visible in PR history",
            "Include WBS 2.2.2 in PR description",
        ],
    )

    heading(doc, "5. Verdict and merge checklist", 1)
    table(
        doc,
        ["Category", "Rating"],
        [
            ["Feature completeness (WBS 2.2.2)", "Core observer works"],
            ["Test coverage", "Good start (4 widget tests)"],
            ["Analytics correctness", "URI length + duplicates undermine data quality"],
            ["Security / privacy", "Full URI logging — prefer path-only"],
            ["Merge readiness", "Approve with changes — fix R1 before merge"],
        ],
    )

    para(doc, "Suggested PR title:", bold=True)
    para(doc, "feat(telemetry): automatic screen_view tracking via GoRouter observer (WBS 2.2.2)")

    heading(doc, "Blocking before merge", 2)
    bullets(
        doc,
        [
            "Log uri.path instead of full URI (R1)",
            "Verify guardrails do not strip typical route paths (R5 or test)",
        ],
        highlight_indices={0, 1},
    )

    heading(doc, "Recommended follow-up", 2)
    bullets(
        doc,
        [
            "Remove duplicate manual screen_view calls (R4)",
            "Move class to features/telemetry/ (R3)",
            "Add debouncing (R2)",
        ],
    )

    heading(doc, "6. Files changed", 1)
    bullets(
        doc,
        [
            "frontend/lib/config/router/app_router.dart",
            "frontend/test/features/telemetry/telemetry_go_router_observer_test.dart",
        ],
    )

    return doc


def main() -> None:
    doc = build()
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    main()
