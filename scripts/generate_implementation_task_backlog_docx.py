"""Generate Word document: Team E implementation task backlog (all research docs)."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX

ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "Team_E_Implementation_Task_Backlog.docx"

SOURCE_DOCS = [
    "Call_Transcript_Retrieval_Review.docx",
    "Hybrid_Retrieval_Scope_PostgreSQL_FullText_pgvector.docx",
    "Ask_AI_Upstream_Pipeline_Call_Visit_Summaries.docx",
    "RBAC_Scoped_Retrieval_Source_Types.docx",
    "Voice_Query_Path_and_STT_Framework_Dependencies.docx",
    "Medication_Timeline_Retrieval_FR-AI-11.docx",
    "HITL_Tier1_Tier2_Escalation_and_Hold_Release_Workflow.docx",
    "LLM_Provider_Abstraction_Bedrock_TeamA_BAA_Fallback.docx",
]


def heading(doc, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def para(doc, text: str, bold: bool = False, highlight: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    if highlight:
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def bullets(doc, items: list[str], highlight_indices: set[int] | None = None) -> None:
    highlight_indices = highlight_indices or set()
    for i, item in enumerate(items):
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        if i in highlight_indices:
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def code(doc, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"


def task_table(doc, rows: list[list[str]], highlight_ids: set[str] | None = None) -> None:
    highlight_ids = highlight_ids or set()
    headers = ["ID", "Task", "Owner", "Depends on", "Source research doc(s)"]
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for row in rows:
        cells = t.add_row().cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = val
            if row[0] in highlight_ids:
                for p in cells[c_idx].paragraphs:
                    for run in p.runs:
                        run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def build() -> None:
    doc = Document()

    title = doc.add_heading("Team E Implementation Task Backlog", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para(doc, "CareConnect — Ordered by importance and dependencies")
    para(
        doc,
        "Consolidates implementation tasks from all Team E research documents. "
        "Tasks earlier in the document unblock everything below them.",
    )
    para(
        doc,
        "Legend — Owner: A = Team A (video/summaries/Bedrock), E = Team E (Ask AI/safety/retrieval), "
        "P = Platform/infra, SUM = Summaries workstream, FE = Frontend.",
    )
    para(
        doc,
        "Yellow highlights mark critical-path or safety-blocking items.",
        highlight=True,
    )
    doc.add_paragraph()

    # 1 Overview
    heading(doc, "1. Executive Overview", 1)
    para(
        doc,
        "Eight research documents identified gaps between Team E SRS/TDD requirements and the current "
        "codebase (post origin/main baseline). This backlog organizes ~50 tasks into nine tiers. "
        "Tier 0–2 are foundation work; Tier 3–5 build the Ask AI data path; Tier 6–7 add safety and "
        "voice; Tier 8 covers resilience and test automation.",
    )
    bullets(
        doc,
        [
            "Do not ship patient-facing Ask AI without RetrievalScopeService (2.1) and Safety/HITL (6.1–6.3).",
            "Lock unified summary_json (1.1) and retrieval_index_chunk (1.5) before parallel indexing work.",
            "Medication timeline (FR-AI-11) is derived indexing + query planning — not a separate retrieval system.",
            "Voice query reuses the text Ask AI gateway after Tier 5 and Tier 6 are in place.",
            "Team A coupling (summaries, sentiment, IAM, transcripts) is the main external dependency for Team E.",
        ],
        highlight_indices={0, 1, 4},
    )

    heading(doc, "1.1 Source research documents", 2)
    bullets(doc, [f"docs/{name}" for name in SOURCE_DOCS])

    # 2 Tier 0
    heading(doc, "2. Tier 0 — Platform Prerequisites", 1)
    para(doc, "Blocks production AI everywhere. Complete before scaling Ask AI or summary features.", bold=True)
    task_table(
        doc,
        [
            ["0.1", "Extend ECS IAM for Claude inference-profile ARNs (not only Nova Lite/Pro)", "P", "—", "LLM Provider Abstraction"],
            ["0.2", "Unify LLM config: single policy; prod guardrail blocking DeepSeek/OpenRouter for PHI", "E/P", "—", "LLM Provider Abstraction"],
            ["0.3", "Introduce LlmRouter / extend AIServiceFactory: Bedrock primary → in-account failover → dev mock", "E (+A)", "0.1, 0.2", "LLM Provider Abstraction"],
            ["0.4", "Refactor BedrockSentimentService to use shared router (not direct BedrockRuntimeClient)", "A", "0.3", "LLM Provider Abstraction; Ask AI Upstream"],
            ["0.5", "Persist provider + model_version on call_summaries and chat audit rows", "A/E", "0.3", "LLM Provider Abstraction; Ask AI Upstream"],
            ["0.6", "Fix AIChatService bean wiring — conditional @Primary; Bedrock vs mock never collide", "E", "0.2", "LLM Provider Abstraction"],
        ],
        highlight_ids={"0.1", "0.3"},
    )

    # Tier 1
    heading(doc, "3. Tier 1 — Lock Shared Contracts", 1)
    para(doc, "Blocks indexing, medication timeline, and Ask AI retrieval. Schema lock is TDD R4 mitigation.", bold=True)
    task_table(
        doc,
        [
            ["1.1", "Lock unified summary_json schema (careInstructions, conditions, SOAP, citations, episodeType)", "A/SUM", "—", "Ask AI Upstream; Medication Timeline"],
            ["1.2", "Upgrade Bedrock summary prompt/parser to emit full contract (not reduced subset)", "A", "1.1", "Medication Timeline; Ask AI Upstream"],
            ["1.3", "Schema validation on summary write — reject meds missing status/effectiveDate", "A", "1.1", "Medication Timeline"],
            ["1.4", "Create visit_summaries table + VisitSummaryService mirroring call_summaries", "P/SUM", "1.1", "Ask AI Upstream"],
            ["1.5", "Define and migrate retrieval_index_chunk (pgvector + tsvector, patient_id, record_type)", "P/E", "—", "Hybrid Retrieval; RBAC; Medication Timeline"],
            ["1.6", "Enable pgvector extension + embedding dimension (1536) in Flyway", "P", "1.5", "Hybrid Retrieval"],
        ],
        highlight_ids={"1.1", "1.5"},
    )

    # Tier 2
    heading(doc, "4. Tier 2 — RBAC & Retrieval Scope", 1)
    para(doc, "Must precede any retrieval or Ask AI that returns patient data.", bold=True)
    task_table(
        doc,
        [
            ["2.1", "Implement RetrievalScopeService — patient/caregiver/family + consent (REQ-SC-7/8)", "E", "—", "RBAC Scoped Retrieval"],
            ["2.2", "Grant USE_AI_FEATURES to PATIENT; fix symptom permission gaps", "E", "—", "Voice Query; RBAC Scoped Retrieval"],
            ["2.3", "Index all chunks with patient_id; enforce patient filter on every query", "E", "1.5, 2.1", "RBAC Scoped Retrieval"],
            ["2.4", "Add caregiver_visibility on summary rows; exclude restricted chunks for caregivers", "A/E", "1.4, 2.1", "RBAC Scoped Retrieval"],
            ["2.5", "Scope USPS mail by mailbox owner → patient; allow patient retrieval of own mail", "E", "2.1", "RBAC Scoped Retrieval"],
            ["2.6", "403 + audit on cross-patient / unauthorized source access (no AI output)", "E", "2.1", "RBAC Scoped Retrieval; HITL"],
        ],
        highlight_ids={"2.1", "2.6"},
    )

    # Tier 3
    heading(doc, "5. Tier 3 — Transcript & Upstream Pipeline", 1)
    para(doc, "Feeds the retrieval index with complete transcript and summary data.", bold=True)
    task_table(
        doc,
        [
            ["3.1", "Composite DB index on call_transcript_segments (call_id, start_ms, occurred_at)", "P", "—", "Call Transcript Retrieval Review"],
            ["3.2", "Reliable transcript ingest — client segments + post-call Transcribe; fix S3 null in dev", "A", "—", "Call Transcript Retrieval Review"],
            ["3.3", "Summary regen hook when post-call transcription completes", "A", "3.2", "Call Transcript Retrieval; Voice Query"],
            ["3.4", "Emit SUMMARY_CREATED / TRANSCRIPT_INDEXED events (or polling job) for indexer", "A/E", "1.2, 3.2", "Hybrid Retrieval; Ask AI Upstream"],
            ["3.5", "Single chunker for call + visit summaries (same inner JSON, different envelope)", "E", "1.1, 1.4", "Ask AI Upstream; Hybrid Retrieval"],
        ],
    )

    # Tier 4
    heading(doc, "6. Tier 4 — Indexing Pipeline", 1)
    para(doc, "Blocks grounded Ask AI answers with citations.", bold=True)
    task_table(
        doc,
        [
            ["4.1", "RetrievalIndexService — ingest per record type (TRANSCRIPT, SUMMARY, DOCUMENT, NOTE, USPS)", "E", "1.5, 2.3, 3.4", "Hybrid Retrieval; RBAC"],
            ["4.2", "FTS: maintain search_vector via trigger or application on insert/update", "E", "4.1", "Hybrid Retrieval"],
            ["4.3", "Embeddings: batch embed chunk_text via Bedrock Titan or approved embed model", "E", "0.3, 4.1", "Hybrid Retrieval"],
            ["4.4", "Embedding backfill worker + optional SchemaPatchRunner NULL-embedding partial index (no Flyway at deploy)", "E", "4.1–4.3, 3.2", "Hybrid Retrieval; Transcript Review"],
            ["4.5", "Derived index MEDICATION_TIMELINE_EVENT from careInstructions type=medication", "E", "1.2, 4.1", "Medication Timeline FR-AI-11"],
            ["4.6", "Medication name normalization at index time (RxNorm or alias list)", "E", "4.5", "Medication Timeline FR-AI-11"],
        ],
        highlight_ids={"4.1"},
    )

    para(
        doc,
        "Task 4.4 deliverables: ChunkEmbeddingBackfillWorker scheduled poll, "
        "findMissingEmbeddingsForBackfill / countMissingEmbeddingsForBackfill repository queries, "
        "careconnect.embedding.backfill.* properties, and reuse of ChunkEmbeddingService.embedAndPersist. "
        "Optional DBA follow-up (not blocking MVP): SchemaPatchRunner partial index "
        "idx_retrieval_chunk_embedding_null_backfill (reference SQL V2607161317) when NULL-embedding "
        "backlog is large. No Flyway at ECS deploy.",
    )
    para(doc, "Core Team E deliverable — records-grounded Ask AI.", bold=True)
    task_table(
        doc,
        [
            ["5.1", "HybridRetrievalService — scoped prefilter → FTS + vector → RRF → top-k citations", "E", "2.1, 4.1–4.3", "Hybrid Retrieval"],
            ["5.2", "Medication intent planner — prefilter MEDICATION_TIMELINE_EVENT (FR-AI-11)", "E", "4.5, 5.1", "Medication Timeline FR-AI-11"],
            ["5.3", "POST /api/ai/ask — JWT, scope, retrieval, prompt, min-necessary context (FR-AI-9)", "E", "5.1, 2.1", "Hybrid Retrieval; Ask AI Upstream"],
            ["5.4", "Re-wire MedicalContextService into Bedrock path (replace raw-message-only chat)", "E", "5.3", "Ask AI Upstream; LLM Abstraction"],
            ["5.5", "Citation builder — type, id, excerpt, deep link (FR-AI-2)", "E", "5.1", "Hybrid Retrieval; RBAC"],
            ["5.6", "No-records path — explicit message; no general medical answer (FR-AI-2)", "E", "5.1", "Hybrid Retrieval"],
            ["5.7", "Flutter Ask AI screen — citations, disclaimer, confirm-with-provider prompt", "E/FE", "5.3", "Ask AI Upstream; HITL"],
            ["5.8", "Deprecate PatientContextRetrievalService stub; single hybrid path only", "E", "5.1", "Ask AI Upstream"],
        ],
        highlight_ids={"5.3"},
    )

    # Tier 6
    heading(doc, "8. Tier 6 — Safety, Consent & HITL", 1)
    para(
        doc,
        "Must wrap delivery before patient-facing rollout. Tier 2 hold is an SRS hard requirement.",
        bold=True,
    )
    task_table(
        doc,
        [
            ["6.1", "SafetyValidationService + TierClassifier (SRS Table 8 rules)", "E", "—", "HITL Tier 1–2 Escalation"],
            ["6.2", "held_item schema + HITL queue API (PENDING → APPROVED/REJECTED)", "E", "6.1", "HITL Tier 1–2 Escalation"],
            ["6.3", "Integrate SCC-3 steps 11–16 into Ask AI gateway — tier assign → hold or deliver", "E", "5.3, 6.1", "HITL; Hybrid Retrieval"],
            ["6.4", "Human Reviewer role + release/reject endpoints", "E", "6.2", "HITL Tier 1–2 Escalation"],
            ["6.5", "Client held-state UX — review message + poll/WebSocket until release", "E/FE", "6.2, 6.3", "HITL Tier 1–2 Escalation"],
            ["6.6", "Confirmation workflow — approve once / session / decline (REQ-SC-5/6)", "E", "5.3", "HITL Tier 1–2 Escalation"],
            ["6.7", "Extend HITL to summary items — med-change confirm (FR-SUM-4)", "E/A", "6.1, 1.2", "HITL; Medication Timeline"],
            ["6.8", "Immutable audit ledger — query, tier, hold, release, citations (REQ-SC-9)", "E", "6.3", "HITL; RBAC"],
        ],
        highlight_ids={"6.1", "6.2", "6.3"},
    )

    # Tier 7
    heading(doc, "9. Tier 7 — Voice Query & STT", 1)
    para(doc, "After text Ask AI and safety gateway are working.", bold=True)
    task_table(
        doc,
        [
            ["7.1", "Mic on Ask AI → speech_to_text → existing text pipeline", "E/FE", "5.3, 2.2", "Voice Query Path"],
            ["7.2", "Document STT ownership — Chime live vs post-call Transcribe vs client STT", "A/E", "3.2, 3.3", "Voice Query; Transcript Review"],
            ["7.3", "Voice path through same SCC-3 — RBAC → retrieval → safety → HITL", "E", "5.3, 6.3, 7.1", "Voice Query; HITL"],
        ],
    )

    # Tier 8
    heading(doc, "10. Tier 8 — Resilience, Ops & Test Automation", 1)
    task_table(
        doc,
        [
            ["8.1", "Bedrock throttling: exponential backoff + retry queue (SRS UC-SUM 2a)", "A/E", "0.3", "LLM Provider Abstraction"],
            ["8.2", "Dev mock summary when Bedrock not provisioned (SRS UC-SUM 2b)", "A", "0.3", "LLM Provider Abstraction"],
            ["8.3", "Automate TC-E-SC-001, TC-AI-06, TC-SCC-03, FR-AI-11, hybrid retrieval tests", "E", "Tier 5–6", "All test-plan refs"],
            ["8.4", "Update PROGRAMMERS_GUIDE — Bedrock primary, abstraction, retrieval, HITL", "E/Docs", "Tier 0–6", "LLM Provider Abstraction"],
            ["8.5", "Spanish safety copy review-verified (SRS NFR)", "E", "6.3", "HITL Tier 1–2 Escalation"],
            ["8.6", "Monitoring: retrieval latency, index lag, HITL queue depth, provider failover", "P/E", "Tier 4–6", "Hybrid Retrieval; HITL"],
        ],
    )

    # Critical path
    heading(doc, "11. Critical Path — Minimum Viable Ask AI with Safety", 1)
    code(
        doc,
        "0.1–0.3  LLM router + IAM\n"
        "    ↓\n"
        "1.1–1.5  Unified schema + retrieval_index_chunk\n"
        "    ↓\n"
        "2.1        RetrievalScopeService\n"
        "    ↓                    ↓\n"
        "3.2–3.4  Transcript/events     4.1–4.3  Indexing pipeline\n"
        "    ↓                    ↓\n"
        "    └──────────┬─────────┘\n"
        "               ↓\n"
        "5.1–5.6  Hybrid retrieval + POST /api/ai/ask\n"
        "               ↓\n"
        "6.1–6.3  SafetyValidation + TierClassifier + HITL hold\n"
        "               ↓\n"
        "7.x      Voice query (optional extension)\n"
        "5.2      Medication timeline (parallel after 4.5 + 5.1)",
    )

    # Parallel workstreams
    heading(doc, "12. Parallel Workstreams (after Tier 1 schema lock)", 1)
    table(
        doc,
        ["Stream", "Can start after", "Task IDs"],
        [
            ["Team A — Summaries", "1.1", "1.2, 1.3, 1.4, 3.3, 0.4"],
            ["Platform — Index DB", "—", "1.5, 1.6, 3.1"],
            ["Team E — RBAC", "—", "2.1–2.6 (parallel with 1.5)"],
            ["Team E — Indexer", "1.5, 2.3, 3.4", "4.1–4.6"],
            ["Team E — Ask AI", "4.1, 2.1", "5.1–5.8"],
            ["Team E — Safety", "5.3 (design earlier)", "6.1–6.8"],
            ["Voice", "5.3, 6.3", "7.1–7.3"],
        ],
    )

    # Dependency matrix
    heading(doc, "13. Key Dependency Rules", 1)
    table(
        doc,
        ["If you need…", "You must complete first…"],
        [
            ["Grounded Ask AI answers", "1.5, 2.1, 4.1–4.3, 5.1"],
            ["Medication timeline (FR-AI-11)", "1.2, 4.5, 5.1, 5.2"],
            ["Tier 2 hold before patient sees answer", "5.3, 6.1, 6.2, 6.3"],
            ["Voice Ask AI", "5.3, 6.3, 7.1"],
            ["Claude in ECS production", "0.1 (IAM)"],
            ["Visit summary retrieval", "1.4, 3.5, 4.1"],
            ["Caregiver-scoped summaries", "2.4, 2.1"],
            ["BAA-compliant prod inference", "0.2, 0.3 (block non-BAA providers)"],
        ],
    )

    # Milestone mapping
    heading(doc, "14. Suggested Milestone Mapping", 1)
    table(
        doc,
        ["Milestone", "Tiers / tasks", "Outcome"],
        [
            ["M2 completion / hardening", "0.1–0.6, 1.1–1.3", "Bedrock reliable; summary contract locked"],
            ["M3 — Retrieval foundation", "1.5–1.6, 2.1–2.6, 3.1–3.5", "Index schema + scope + upstream events"],
            ["M3 — Ask AI alpha", "4.1–4.4, 5.1–5.6", "Internal /api/ai/ask with citations"],
            ["M3 — Safety gate", "6.1–6.5, 6.8", "Tier 2 hold + reviewer workflow"],
            ["M3 — Feature complete", "4.5–4.6, 5.2, 5.7, 6.6–6.7, 7.1–7.3", "Med timeline, voice, summary confirm"],
            ["M3 — Production ready", "8.1–8.6", "Failover, tests, monitoring, docs"],
        ],
    )

    # Full task index
    heading(doc, "15. Complete Task Index (all IDs)", 1)
    para(doc, "Quick reference — 47 tasks across nine tiers.")
    code(
        doc,
        "Tier 0: 0.1, 0.2, 0.3, 0.4, 0.5, 0.6\n"
        "Tier 1: 1.1, 1.2, 1.3, 1.4, 1.5, 1.6\n"
        "Tier 2: 2.1, 2.2, 2.3, 2.4, 2.5, 2.6\n"
        "Tier 3: 3.1, 3.2, 3.3, 3.4, 3.5\n"
        "Tier 4: 4.1, 4.2, 4.3, 4.4, 4.5, 4.6\n"
        "Tier 5: 5.1, 5.2, 5.3, 5.4, 5.5, 5.6, 5.7, 5.8\n"
        "Tier 6: 6.1, 6.2, 6.3, 6.4, 6.5, 6.6, 6.7, 6.8\n"
        "Tier 7: 7.1, 7.2, 7.3\n"
        "Tier 8: 8.1, 8.2, 8.3, 8.4, 8.5, 8.6",
    )

    heading(doc, "16. Related Documents", 1)
    bullets(
        doc,
        [
            "C:\\Users\\ravic\\Downloads\\CareConnect_SRS_Revision 2.0_TEAM E.docx",
            "C:\\Users\\ravic\\Downloads\\CareConnect_Milestone_2_TDD_TEAM E.docx",
            "C:\\Users\\ravic\\Downloads\\CareConnect_Milestone_2_Software_Test_Plan (2).docx",
        ]
        + [f"docs/{name}" for name in SOURCE_DOCS],
    )

    heading(doc, "17. Conclusion", 1)
    para(
        doc,
        "This backlog is the single ordered implementation plan derived from Team E research. "
        "Schedule Tier 0 and Tier 1 in the same sprint as Team A summary and IAM work. "
        "Do not expose Ask AI to care recipients until Tier 2 scope enforcement and Tier 6 "
        "safety/HITL are integrated into the same gateway path as POST /api/ai/ask.",
        highlight=True,
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    try:
        doc.save(OUTPUT)
        print(f"Created: {OUTPUT}")
    except PermissionError:
        alt = OUTPUT.with_name(f"{OUTPUT.stem}_refresh{OUTPUT.suffix}")
        doc.save(alt)
        print(f"Created (original locked): {alt}")


def table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for row in rows:
        cells = t.add_row().cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = val


if __name__ == "__main__":
    build()
