"""Generate Word document: Storage audit vs shared retrieval_index_chunk schema."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX

ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "Storage_Audit_vs_Shared_Index_Schema.docx"


def heading(doc, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def para(doc, text: str, bold: bool = False, highlight: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    if highlight:
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def bullets(doc, items: list[str], highlight_indices: set[int] | None = None) -> None:
    highlight_indices = highlight_indices or set()
    for i, item in enumerate(items):
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        if i in highlight_indices:
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def code(doc, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"


def table(doc, headers, rows, highlight_rows: set[int] | None = None) -> None:
    highlight_rows = highlight_rows or set()
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for r_idx, row in enumerate(rows, start=1):
        cells = t.add_row().cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = val
            if r_idx in highlight_rows:
                for p in cells[c_idx].paragraphs:
                    for run in p.runs:
                        run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def build() -> None:
    doc = Document()

    title = doc.add_heading(
        "Storage Audit: Transcript / Summary / Document vs Shared Index Schema", 0
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para(doc, "CareConnect Team E — Indexing & Retrieval Foundation")
    para(
        doc,
        "Audits existing PostgreSQL and object-storage tables against the planned "
        "retrieval_index_chunk shared schema (TDD §7.1, Hybrid Retrieval research). "
        "Identifies field gaps, patient-scope resolution paths, and indexer prerequisites.",
    )
    para(
        doc,
        "Yellow highlights mark blocking gaps for hybrid retrieval indexing.",
        highlight=True,
    )
    doc.add_paragraph()

    # 1 Executive summary
    heading(doc, "1. Executive Summary", 1)
    para(
        doc,
        "CareConnect already persists rich upstream content — call transcript segments, "
        "archived transcript JSON in S3, call summaries, patient notes, and user-uploaded "
        "files — but none of it is indexed for hybrid retrieval. The target "
        "retrieval_index_chunk table does not exist in Flyway migrations; pgvector is only "
        "configured in local Docker (pgvector/pgvector:pg15), not in production migrations.",
    )
    para(
        doc,
        "Critical cross-cutting gap: transcript and summary tables are keyed by call_id only — "
        "they lack patient_id, consent_scope, and caregiver_visibility columns required for "
        "RBAC-scoped indexing. Patient scope must be resolved at index time via call telemetry "
        "(CALL_JOIN events), which is fragile when telemetry is incomplete.",
        highlight=True,
    )
    para(
        doc,
        "Summary JSON stored today is a reduced Bedrock shape (headline, keyConcerns, "
        "recommendedActions) — not the unified call/visit contract with careInstructions[], "
        "conditions[], and SOAP block needed for FR-AI-11 medication timeline indexing.",
        highlight=True,
    )

    # 2 Target schema
    heading(doc, "2. Target Shared Index Schema (Reference)", 1)
    code(
        doc,
        "CREATE TABLE retrieval_index_chunk (\n"
        "  id UUID PRIMARY KEY,\n"
        "  patient_id BIGINT NOT NULL,           -- RBAC filter on every query\n"
        "  record_type VARCHAR(40) NOT NULL,    -- TRANSCRIPT | CALL_SUMMARY | ...\n"
        "  source_record_id VARCHAR(120) NOT NULL,\n"
        "  chunk_text TEXT NOT NULL,\n"
        "  chunk_metadata JSONB,                -- speaker, timestamps, citation, deep link\n"
        "  search_vector TSVECTOR,\n"
        "  embedding vector(1536),\n"
        "  indexed_at TIMESTAMPTZ NOT NULL DEFAULT now(),\n"
        "  consent_scope VARCHAR(40)\n"
        ");",
    )
    bullets(
        doc,
        [
            "Every chunk must be patient-scoped before insert (RetrievalScopeService at query time mirrors this at index time).",
            "Source tables may remain call_id–keyed; indexer denormalizes patient_id into chunks.",
            "FTS + embedding live on the index table — not on source OLTP tables (optional denormalized caches allowed).",
        ],
    )

    # 3 Audit matrix overview
    heading(doc, "3. Audit Summary Matrix", 1)
    table(
        doc,
        ["Source", "Table / store", "Exists?", "patient_id on source?", "Index-ready?", "Priority"],
        [
            ["Transcript segments", "call_transcript_segments", "Yes (V61)", "No — call_id only", "Partial", "P0"],
            ["Transcript archive", "call_transcript_archives + S3 JSON", "Yes (V63)", "No", "Partial", "P0"],
            ["Call summary", "call_summaries", "Yes (V61)", "No — call_id only", "Partial", "P0"],
            ["Visit summary", "visit_summaries", "No", "—", "Missing", "P0"],
            ["Summary items", "inside summary_json", "Partial schema", "No", "Blocked by schema", "P0"],
            ["Uploaded documents", "user_files (+ S3/blob)", "Yes", "Yes (nullable)", "Partial — no OCR text", "P1"],
            ["Clinical notes", "patient_note", "Yes", "Yes", "Ready", "P1"],
            ["USPS mail", "usps_digest_cache", "Yes", "No — userId string", "Not index-ready", "P1"],
            ["Med timeline (derived)", "—", "No", "—", "Derived at index", "P1"],
            ["Hybrid index", "retrieval_index_chunk", "No", "—", "Not implemented", "P0"],
        ],
        highlight_rows={2, 3, 4, 5, 10, 11},
    )

    # 4 Transcript storage
    heading(doc, "4. Transcript Storage Audit", 1)

    heading(doc, "4.1 call_transcript_segments (V61)", 2)
    table(
        doc,
        ["Column / aspect", "Current state", "Index schema need", "Gap"],
        [
            ["Primary key", "id BIGSERIAL", "source_record_id in chunk metadata", "OK — use segment id or composite key"],
            ["Scope key", "call_id VARCHAR(120)", "patient_id on every chunk", "GAP — no patient_id; join required"],
            ["Searchable text", "transcript_text TEXT", "chunk_text", "OK — maps directly"],
            ["Structured metadata", "speaker_label, start_ms, end_ms, source, actor_user_id, occurred_at", "chunk_metadata JSONB", "OK — embed in metadata"],
            ["FTS / vector", "None", "search_vector, embedding on index table", "Expected — indexer responsibility"],
            ["Indexes", "call_id, actor_user_id, start_ms (separate)", "patient_id + record_type on index", "GAP — no composite (call_id, start_ms, occurred_at) on source"],
            ["Segment limits", "Max 200/request, 1200 chars/segment (service)", "Chunk size policy for index", "Align chunker with CallTranscriptService limits"],
        ],
        highlight_rows={2, 7},
    )
    bullets(
        doc,
        [
            "Sources: CLIENT_TRANSCRIPT (web/mobile), POST_CALL_TRANSCRIBE (AWS Transcribe), Chime-side capture.",
            "CallTranscriptService merges DB segments with archived S3 segments for read path.",
            "PostCallTranscriptionService writes segments after recording concat — no auto summary regen hook.",
        ],
    )

    heading(doc, "4.2 call_transcript_archives + S3 (V63)", 2)
    table(
        doc,
        ["Column / aspect", "Current state", "Index need", "Gap"],
        [
            ["Metadata", "storage_key, segment_count, transcript_chars, participant_user_ids, sha256", "Re-index on archive event", "OK for provenance"],
            ["Content", "S3 JSON array of segments", "Re-chunk into index", "Indexer must read S3 when DB rows deleted"],
            ["Trigger", "archive when ≥600 segments or ≥120K chars; optional delete DB rows", "TRANSCRIPT_ARCHIVED event", "GAP — no indexing event emitted"],
            ["S3 in dev", "s3StorageService optional (@Autowired required=false)", "Local dev index backfill", "GAP — archive metadata without content in local dev"],
            ["patient_id", "Not stored", "Denormalize at index time", "Resolve via call_id → telemetry → patient"],
        ],
        highlight_rows={4, 5},
    )

    heading(doc, "4.3 Patient resolution path (call_id → patient_id)", 2)
    code(
        doc,
        "CallController.findPatientInCall(callId):\n"
        "  call_telemetry_events WHERE event_type = 'CALL_JOIN'\n"
        "  → actor_user_id → users.role = PATIENT → patient user id\n"
        "\n"
        "Limitations:\n"
        "  - Fails if patient never joined or telemetry missing\n"
        "  - Uses user id as patient id (Patient entity may differ)\n"
        "  - No persisted patient_id on transcript/summary rows for audit",
    )
    para(
        doc,
        "Recommendation: add patient_id (nullable) to call_summaries and optionally "
        "call_transcript_archives at write time; backfill from telemetry. Indexer should "
        "skip chunks when patient_id cannot be resolved.",
        highlight=True,
    )

    # 5 Summary storage
    heading(doc, "5. Call / Visit Summary Storage Audit", 1)

    heading(doc, "5.1 call_summaries (V61)", 2)
    table(
        doc,
        ["Column / aspect", "Current state", "Unified contract need", "Gap"],
        [
            ["Payload", "summary_json TEXT", "Full unified JSON (SOAP, careInstructions, conditions, …)", "GAP — reduced Bedrock output today"],
            ["Scope", "call_id only", "patient_id + episodeType", "GAP — no patient_id column"],
            ["Status", "status VARCHAR(24)", "READY | ERROR | PENDING", "OK"],
            ["Provenance", "generated_by_user_id, generated_at, transcript_segment_count", "provider + model_version", "GAP — LLM provider/model not stored"],
            ["Visibility", "None", "caregiver_visibility per RBAC research", "GAP — not on row"],
            ["Visit parity", "N/A", "visit_summaries identical shape", "GAP — table not created"],
        ],
        highlight_rows={2, 3, 6, 7},
    )
    para(doc, "Current Bedrock summary JSON shape (BedrockSentimentService.summarizeTranscript):", bold=True)
    code(
        doc,
        '{\n'
        '  "headline": "...",\n'
        '  "overallAssessment": "...",\n'
        '  "keyConcerns": ["..."],\n'
        '  "recommendedActions": ["..."],\n'
        '  "followUpQuestions": ["..."]\n'
        '}',
    )
    para(doc, "Target unified contract (Ask AI upstream research) additionally requires:", bold=True)
    bullets(
        doc,
        [
            "actionItems[], appointments[], careInstructions[] (with status, effectiveDate, type)",
            "conditions[], clinicalObservations, SOAP block, riskLevel, urgencyBanner",
            "summaryConfidence, safety fields, episodeType envelope for visit vs call",
        ],
        highlight_indices={0, 1},
    )

    heading(doc, "5.2 Index chunking strategy for summaries", 2)
    table(
        doc,
        ["record_type", "Source field", "Chunk approach", "Blocked?"],
        [
            ["CALL_SUMMARY", "headline + narrative fields", "1–2 chunks per summary", "Partial — missing fields"],
            ["SUMMARY_ITEM", "actionItems, appointments, careInstructions, conditions", "1 chunk per array item + citation metadata", "Yes — arrays not in current JSON"],
            ["MEDICATION_TIMELINE_EVENT", "careInstructions type=medication", "Derived explode at index", "Yes — depends on unified schema"],
        ],
        highlight_rows={3, 4},
    )

    # 6 Document storage
    heading(doc, "6. Document & File Storage Audit", 1)

    heading(doc, "6.1 user_files", 2)
    table(
        doc,
        ["Column / aspect", "Current state", "Index need", "Gap"],
        [
            ["patient_id", "Long, nullable", "Required on every DOCUMENT chunk", "Partial — nullable; must enforce on medical categories"],
            ["file_category", "Enum: MEDICAL_RECORD, LAB_RESULT, CLINICAL_NOTE, …", "record_type + consent_scope hint", "OK"],
            ["Content", "file_data BLOB or s3_path", "chunk_text from OCR/extraction", "GAP — no persisted extracted text column"],
            ["Search", "None", "FTS + embedding on index", "DocumentProcessingService extracts on demand only"],
            ["MIME / size", "content_type, file_size", "chunk_metadata", "OK"],
        ],
        highlight_rows={4, 5},
    )
    bullets(
        doc,
        [
            "FileManagementService → S3StorageService or DatabaseStorageService.",
            "MedicalContextService / AI chat can extract text at request time — not durable for retrieval index.",
            "Invoice documents (invoice tables, Textract) are a separate path — out of Ask AI scope unless explicitly included.",
        ],
    )
    para(
        doc,
        "Recommendation: add document_text_extracted TEXT or document_chunks JSONB on user_files "
        "(or side table) populated by Textract/Tika pipeline; emit DOCUMENT_OCR_COMPLETE for indexer.",
        highlight=True,
    )

    heading(doc, "6.2 patient_note", 2)
    table(
        doc,
        ["Column", "Current", "Index mapping", "Gap"],
        [
            ["patient_id", "NOT NULL", "patient_id on chunk", "Ready"],
            ["note", "TEXT", "chunk_text", "Ready"],
            ["ai_summary", "TEXT nullable", "Optional second chunk or metadata", "OK"],
            ["author / consent", "Not on row", "chunk_metadata.author, consent_scope", "GAP — add author_user_id if needed for caregiver scope"],
        ],
        highlight_rows={4},
    )

    # 7 USPS
    heading(doc, "7. USPS Mail Storage Audit", 1)
    table(
        doc,
        ["Aspect", "Current (usps_digest_cache)", "Index need (FR-USPS-4)", "Gap"],
        [
            ["Key", "userId (String — auth subject)", "patient_id", "GAP — mailbox owner ≠ patient without mapping"],
            ["Payload", "payloadJson LOB (full digest)", "Per mail piece chunks", "GAP — not normalized rows"],
            ["TTL", "expiresAt", "Re-index on ingest", "Cache-oriented, not durable search store"],
            ["Search today", "USPSDigestService.search() in cached JSON", "Hybrid FTS + vector", "In-memory scan — not scalable"],
            ["OCR text", "Inside digest packages", "visible_text in chunk_metadata", "Needs piece-level explode like summary items"],
        ],
        highlight_rows={1, 2, 3, 5},
    )

    # 8 Supporting stores
    heading(doc, "8. Supporting Storage (Not Primary Index Sources)", 1)
    table(
        doc,
        ["Store", "Role", "Index relevance"],
        [
            ["call_recordings", "S3 artifact paths, transcription_status", "Upstream for POST_CALL_TRANSCRIBE — not directly indexed"],
            ["call_telemetry_events", "CALL_JOIN, sentiment — patient resolution", "Metadata only — not chunked for Ask AI"],
            ["chat_conversations / chat_messages", "AI chat history", "Out of FR-AI-2 records-grounded scope unless policy adds"],
            ["patient_medication", "Structured med list", "Complement to timeline events from summaries — optional future source"],
        ],
    )

    # 9 retrieval_index_chunk status
    heading(doc, "9. retrieval_index_chunk — Implementation Status", 1)
    bullets(
        doc,
        [
            "No Flyway migration creates retrieval_index_chunk or CREATE EXTENSION vector in app migrations.",
            "pgvector available in backend/core/pg_docker/docker-compose.yml for local dev only.",
            "No RetrievalIndexService, embed job, or SUMMARY_CREATED / TRANSCRIPT_INDEXED consumers in codebase.",
            "PatientContextRetrievalService is a stub — not reading any of the audited tables for hybrid search.",
        ],
        highlight_indices={0, 1, 2},
    )

    # 10 Field mapping reference
    heading(doc, "10. Source-to-Index Field Mapping Reference", 1)
    table(
        doc,
        ["record_type", "source_record_id", "patient_id source", "chunk_text source", "chunk_metadata"],
        [
            ["TRANSCRIPT", "segment.id or callId:startMs", "call_id → telemetry → patient", "transcript_text", "speaker, start_ms, end_ms, source, call_id"],
            ["CALL_SUMMARY", "call_summaries.id", "call_id → telemetry (or new column)", "headline + narrative + JSON excerpts", "call_id, generated_at, status"],
            ["VISIT_SUMMARY", "visit_summaries.id", "visit.patient_id", "Same as call", "episodeType=visit, visit_id"],
            ["SUMMARY_ITEM", "summaryId:itemType:itemId", "From parent summary", "Item text + citation excerpt", "sourceTurnId, type, status, effectiveDate"],
            ["DOCUMENT", "user_files.id", "user_files.patient_id", "OCR/extracted text", "mime, category, s3_path, upload date"],
            ["CLINICAL_NOTE", "patient_note.id", "patient_note.patient_id", "note body", "created_at, author if added"],
            ["USPS_MAIL", "digestId:pieceId", "userId → patient mapping", "visible_text, sender", "delivery_date, tier, image flag"],
            ["MEDICATION_TIMELINE_EVENT", "derived key", "From parent summary patient", "Event description", "medication, status, effectiveDate, citations"],
        ],
        highlight_rows={2, 3, 4, 8},
    )

    # 11 Remediation plan
    heading(doc, "11. Remediation Plan (Ordered by Dependency)", 1)
    table(
        doc,
        ["Step", "Storage / schema change", "Unblocks"],
        [
            ["1", "Flyway: CREATE EXTENSION vector; create retrieval_index_chunk", "All indexing"],
            ["2", "Add patient_id to call_summaries; resolve at summary write", "Scoped summary chunks"],
            ["3", "Upgrade summary_json to unified contract + validation", "SUMMARY_ITEM + MEDICATION_TIMELINE_EVENT chunks"],
            ["4", "Create visit_summaries table (mirror call_summaries)", "Visit episode indexing"],
            ["5", "Add caregiver_visibility to summary rows", "RBAC-scoped summary retrieval"],
            ["6", "Emit SUMMARY_CREATED / TRANSCRIPT_INDEXED events (or poll job)", "Indexer triggers"],
            ["7", "Indexer: transcript segments + archive S3 fallback", "TRANSCRIPT chunks"],
            ["8", "Indexer: summary + item explode + med timeline derive", "FR-AI-11, Ask AI citations"],
            ["9", "document_text_extracted sidecar on user_files + OCR pipeline hook", "DOCUMENT chunks"],
            ["10", "Normalize USPS to patient-scoped mail_piece store or map userId→patient", "USPS_MAIL chunks"],
            ["11", "Composite index on call_transcript_segments (call_id, start_ms, occurred_at)", "Efficient backfill"],
            ["12", "Store llm_provider + model_version on call_summaries", "SRS traceability"],
        ],
        highlight_rows={1, 2, 3, 6, 9, 10},
    )

    # 12 Cross-doc links
    heading(doc, "12. Related Research Documents", 1)
    bullets(
        doc,
        [
            "docs/Hybrid_Retrieval_Scope_PostgreSQL_FullText_pgvector.docx — target index design",
            "docs/Call_Transcript_Retrieval_Review.docx — transcript read/write surfaces",
            "docs/Ask_AI_Upstream_Pipeline_Call_Visit_Summaries.docx — unified summary contract",
            "docs/RBAC_Scoped_Retrieval_Source_Types.docx — scope rules per record type",
            "docs/Medication_Timeline_Retrieval_FR-AI-11.docx — derived MEDICATION_TIMELINE_EVENT",
            "docs/Team_E_Implementation_Task_Backlog.docx — tasks 1.5, 3.x, 4.x",
        ],
    )

    heading(doc, "13. Code & Migration References", 1)
    bullets(
        doc,
        [
            "backend/core/src/main/resources/db/migration/V61__create_call_transcript_and_summary_tables.sql",
            "backend/core/src/main/resources/db/migration/V63__create_call_transcript_archive_table.sql",
            "backend/core/src/main/java/com/careconnect/model/CallTranscriptSegment.java",
            "backend/core/src/main/java/com/careconnect/model/CallSummary.java",
            "backend/core/src/main/java/com/careconnect/service/CallTranscriptService.java",
            "backend/core/src/main/java/com/careconnect/service/CallTranscriptArchiveService.java",
            "backend/core/src/main/java/com/careconnect/service/CallSummaryService.java",
            "backend/core/src/main/java/com/careconnect/service/BedrockSentimentService.java",
            "backend/core/src/main/java/com/careconnect/model/UserFile.java",
            "backend/core/src/main/java/com/careconnect/model/PatientNote.java",
            "backend/core/src/main/java/com/careconnect/model/USPSDigestCache.java",
            "backend/core/pg_docker/docker-compose.yml (pgvector image)",
        ],
    )

    heading(doc, "14. Conclusion", 1)
    para(
        doc,
        "Existing OLTP storage is adequate as the system of record for transcripts, summaries, "
        "notes, and files — but it is not index-ready for Team E hybrid retrieval. The shared "
        "retrieval_index_chunk layer is entirely missing; source tables lack patient_id and "
        "summary schema completeness; document and USPS content require extraction/normalization "
        "pipelines before chunking.",
    )
    para(
        doc,
        "Priority sequence: (1) create index table + pgvector migration, (2) denormalize patient_id "
        "onto summaries and fix unified summary_json, (3) build indexer with transcript + summary "
        "handlers, (4) add document OCR persistence and USPS patient mapping.",
        highlight=True,
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    try:
        doc.save(OUTPUT)
        print(f"Created: {OUTPUT}")
    except PermissionError:
        alt = OUTPUT.with_name(f"{OUTPUT.stem}_refresh{OUTPUT.suffix}")
        doc.save(alt)
        print(f"Created (original locked): {alt}")


if __name__ == "__main__":
    build()
