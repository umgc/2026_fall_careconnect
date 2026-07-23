"""Generate Word document for Call Transcript Retrieval Review."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx_codebase_revision import REVISION_BULLETS, REVISION_LABEL, save_document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

OUTPUT = Path(__file__).resolve().parents[1] / "docs" / "Call_Transcript_Retrieval_Review.docx"


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_para(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_code(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val


def build() -> None:
    doc = Document()

    title = doc.add_heading("Call Transcript Stack Review", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_para(
        doc,
        "Inherited retrieval surfaces, indexed record types, and upstream transcript dependencies",
    )
    add_para(doc, "CareConnect — backend/core and frontend call transcript architecture")
    add_para(doc, "Generated from codebase review")
    p = doc.add_paragraph()
    r = p.add_run(REVISION_LABEL)
    r.font.highlight_color = WD_COLOR_INDEX.YELLOW
    add_bullets(doc, REVISION_BULLETS)

    doc.add_paragraph()

    add_heading(doc, "Architecture Overview", 1)
    add_para(
        doc,
        "Transcript data flows through three upstream producers into two indexed stores, "
        "then is read back through several layered retrieval surfaces.",
    )

    add_heading(doc, "Data flow (text diagram)", 2)
    add_bullets(
        doc,
        [
            "Upstream producers: Chime embed (web/mobile) → chime-transcript / speech-recognition; "
            "PostCallTranscriptionService → POST_CALL_TRANSCRIBE; ChimeService.startMeetingTranscription (AWS-side only).",
            "Write surface: POST /api/v3/calls/{callId}/transcript/segments → CallTranscriptService.recordSegments.",
            "Indexed stores: call_transcript_segments (live DB); call_transcript_archives + S3 JSON (cold); call_summaries (derived).",
            "Read surfaces: GET /transcript/segments; GET /summary; CallSummaryService → BedrockSentimentService; "
            "PostCallTelemetrySummaryScreen (frontend).",
        ],
    )

    add_heading(doc, "1. Inherited Retrieval Surfaces", 1)
    add_para(
        doc,
        "These are the read paths that compose live DB rows, archived S3 payloads, or both.",
        bold=True,
    )

    add_heading(doc, "Central merge layer: CallTranscriptService", 2)
    add_para(doc, "All downstream readers funnel through one service that merges live and archived segments.")
    add_code(
        doc,
        "getSegmentsForCall(callId):\n"
        "  dbSegments = segmentRepository.findByCallIdOrderByStartMsAscOccurredAtAsc(callId)\n"
        "  archivedSegments = archiveService.getArchivedSegments(callId)\n"
        "  if archived empty → return dbSegments\n"
        "  if db empty → return archivedSegments\n"
        "  else → mergeSegments(archived, db) with dedup + chronological sort",
    )

    add_heading(doc, "Strengths", 3)
    add_bullets(
        doc,
        [
            "Single merge point; tests cover DB-only, archive-only, and merged cases.",
            "Summary text capped at 16,000 chars via buildTranscriptTextForSummary.",
            "Dedup + chronological sort before returning merged results.",
        ],
    )

    add_heading(doc, "Risks", 3)
    add_bullets(
        doc,
        [
            "Dedup key omits actorUserId and occurredAt — two distinct utterances with the same "
            "speaker/text/timing could collapse incorrectly during partial-archive transitions.",
            "Archive reads silently return empty when S3 is unavailable (getArchivedSegments checks "
            "s3StorageService == null). In dev/local, archived metadata can exist while content appears missing.",
            "countSegments() has three branches (DB count, archive count, full merge). Correct but easy "
            "to drift from getSegmentsForCall() if changed independently.",
        ],
    )

    add_heading(doc, "HTTP retrieval surfaces (CallController)", 2)
    add_table(
        doc,
        ["Endpoint", "Auth model", "Transcript dependency"],
        [
            [
                "GET /{callId}/transcript/segments",
                "admin OR telemetry participant OR hasTranscriptAccess",
                "getSegmentsForCall()",
            ],
            [
                "GET /{callId}/summary",
                "admin OR telemetry OR transcript access OR summary owner",
                "buildTranscriptTextForSummary() + lazy regen on NO_TRANSCRIPT",
            ],
            [
                "GET /{callId}/transcription/debug",
                "authenticated only (no participant check)",
                "Chime in-memory state, not DB",
            ],
            [
                "DELETE /{callId}/telemetry (dev)",
                "dev/local only",
                "purgeForCall()",
            ],
        ],
    )
    doc.add_paragraph()
    add_bullets(
        doc,
        [
            "Access checks are reimplemented per endpoint rather than centralized.",
            "GET /transcription/debug has no participant check — any logged-in user can probe any callId's "
            "Chime transcription state.",
            "POST vs GET auth differs: POST requires isCallParticipant; GET also allows telemetry-based access.",
        ],
    )

    add_heading(doc, "Downstream consumers", 2)
    add_bullets(
        doc,
        [
            "CallSummaryService — reads transcript via buildTranscriptTextForSummary, sentiment via "
            "CallTelemetryService, then BedrockSentimentService.summarizeTranscript. Triggers archival after persist.",
            "BedrockSentimentService.summarizeTranscript — downstream only; never writes segments.",
            "Frontend PostCallTelemetrySummaryScreen — parallel fetches telemetry, summary, transcript segments, "
            "and recording; highlights transcript lines against sentiment timestamps.",
        ],
    )

    add_heading(doc, "Auditable inheritance", 2)
    add_para(
        doc,
        "CallTranscriptSegment, CallTranscriptArchive, and CallSummary extend Auditable (created_at, updated_at). "
        "Archive payloads store occurredAt separately in S3 JSON, so timeline semantics depend on which layer you read from.",
    )

    add_heading(doc, "2. Indexed Record Types", 1)

    add_heading(doc, "CallTranscriptSegment (live rows)", 2)
    add_para(doc, "Table: call_transcript_segments")
    add_bullets(
        doc,
        [
            "Indexes: idx_call_transcript_call_id (call_id); idx_call_transcript_actor (actor_user_id); "
            "idx_call_transcript_start_ms (start_ms).",
            "Indexes match query patterns: by call, access check by actor, chronological ordering by start_ms.",
            "Gap: no composite index on (call_id, start_ms, occurred_at) despite repository sort using all three.",
        ],
    )

    add_heading(doc, "CallTranscriptArchive (cold storage metadata)", 2)
    add_para(doc, "Table: call_transcript_archives")
    add_bullets(
        doc,
        [
            "Indexes: idx_call_transcript_archive_call_id; idx_call_transcript_archive_archived_at.",
            "Payload stored as ArchivedTranscriptSegment JSON in S3.",
            "Access control uses participant_user_ids, derived only from segments with non-null actorUserId.",
        ],
    )

    add_heading(doc, "CallSummary (derived artifact)", 2)
    add_bullets(
        doc,
        [
            "Indexed on call_id and generated_at.",
            "Stores transcript_segment_count and status (SUCCESS, NO_TRANSCRIPT, ERROR).",
            "Not a transcript source — downstream index of summary generation outcomes.",
        ],
    )

    add_heading(doc, "CallRecording.transcription_status", 2)
    add_para(
        doc,
        "Tracks async post-call Transcribe lifecycle (PROCESSING, COMPLETE, FAILED). "
        "Related but separate from segment storage.",
    )

    add_heading(doc, "Schema duplication", 2)
    add_para(
        doc,
        "V52/V61 and V53/V63 are identical duplicate Flyway migrations. Harmless today (IF NOT EXISTS), "
        "but confusing for schema history and future edits.",
    )

    add_heading(doc, "3. Upstream Transcript Dependencies", 1)

    add_heading(doc, "Producer A — Live client capture (primary path)", 2)
    add_para(doc, "Chain: chime_meeting_embed_web.dart → video_call_service.dart → POST /transcript/segments")
    add_bullets(
        doc,
        [
            "Sources: chime-transcript (preferred) or speech-recognition (fallback).",
            "Backend normalizes to CLIENT_TRANSCRIPT when source is blank.",
            "Segments buffered and flushed one-at-a-time with retry on end-call.",
            "actorUserId is set from the authenticated poster — enables hasTranscriptAccess.",
            "Dependency: requires JWT + active call state; flush skips when not in-call unless overridden at end-call.",
        ],
    )

    add_heading(doc, "Producer B — Post-call AWS Transcribe (async)", 2)
    add_para(
        doc,
        "Chain: CallRecordingService → PostCallTranscriptionService.transcribeAndCleanup → "
        "recordSegments(callId, null, segments)",
    )
    add_bullets(
        doc,
        [
            "Source: POST_CALL_TRANSCRIBE.",
            "actorUserId is always null — access relies on telemetry participation or archive participant lists.",
            "Speaker labels from buildSpeakerRoleMap() counting CALL_JOIN telemetry, not actual user IDs.",
            "Runs after maybeGenerateAndStoreCallSummary() on call end — first summary often sees NO_TRANSCRIPT.",
            "Mitigation: GET /summary regenerates when status is NO_TRANSCRIPT but segments now exist.",
            "Gap: no automatic summary regen when post-call transcription completes — only on next summary fetch.",
        ],
    )

    add_heading(doc, "Producer C — Chime backend transcription", 2)
    add_para(
        doc,
        "ChimeService.startMeetingTranscription() starts AWS-side transcription but does not write to "
        "CallTranscriptService. Operational/debug infrastructure only. Client-side Chime SDK capture (Producer A) "
        "is the actual ingestion path.",
    )

    add_heading(doc, "Dependency graph at call end", 2)
    add_code(
        doc,
        "if (shouldEndMeeting) {\n"
        "  maybeRecordFinalOverallSentiment(...);\n"
        "  maybeGenerateAndStoreCallSummary(...);  // uses transcript available NOW\n"
        "  callRecordingService.stopRecording(callId);  // triggers post-call Transcribe LATER\n"
        "  chimeService.endMeeting(callId);\n"
        "}",
    )
    add_bullets(
        doc,
        [
            "Summary generated from whatever transcript exists at end-call time.",
            "Recording stop triggers post-call Transcribe asynchronously.",
            "Live client flush may still be in flight.",
            "Explains NO_TRANSCRIPT → lazy-regen pattern; post-call UI should not assume summary is final immediately.",
        ],
    )

    add_heading(doc, "4. Recent Codebase Changes (Participant Handling)", 1)
    add_para(
        doc,
        "Merged PR #89 (feature/a-drattray-video-call-participant-handling) improves call-end "
        "and multi-party behavior relevant to transcript flush timing and summary generation.",
    )
    add_bullets(
        doc,
        [
            "Frontend VideoCallService maintains _participantUserIds; endCall() POSTs participantUserIds "
            "in JSON body alongside otherPartyId and call context metadata.",
            "CallController.endCall() calls mergeParticipantUserIds() on active roster and notify set — "
            "shouldEndMeeting when ≤1 remaining participant after merge.",
            "When meeting continues: sends participant-left WebSocket to remaining participants with "
            "remainingParticipantCount; when meeting ends: call-ended to notifyIds (active + pending invitees).",
            "endCall still runs maybeGenerateAndStoreCallSummary before stopRecording — transcript flush "
            "at end uses respectInCallState:false with maxAttempts:3 (unchanged intent, improved participant accuracy).",
            "ChimeService.joinMeeting() returns cached attendee credentials for same (callId, userId) — "
            "idempotent re-join; fixes duplicate externalUserId AWS errors on double-tap join.",
            "Conference endpoints: eligible-invitees and invite-participant; pending CONFERENCE_INVITE targets "
            "included in end-call notifications.",
        ],
    )

    add_heading(doc, "Findings Summary", 1)
    add_table(
        doc,
        ["Area", "Severity", "Finding"],
        [
            [
                "Retrieval",
                "Medium",
                "Archive content invisible when S3 is disabled, even if archive rows exist",
            ],
            [
                "Retrieval",
                "Medium",
                "Merge dedup key may drop legitimate segments during DB/archive overlap",
            ],
            [
                "Retrieval",
                "Low",
                "Access-control logic duplicated across 3+ endpoints with slight rule differences",
            ],
            [
                "Security",
                "Medium",
                "GET /transcription/debug lacks participant authorization",
            ],
            [
                "Indexed types",
                "Low",
                "Duplicate Flyway migrations (V52/V61, V53/V63)",
            ],
            [
                "Upstream",
                "Medium",
                "Post-call Transcribe completes after summary; no push regen on COMPLETE",
            ],
            [
                "Upstream",
                "Medium",
                "Post-call segments have actorUserId=null, weakening archive-based ACL",
            ],
            [
                "Upstream",
                "Low",
                "Two Chime transcription paths (SDK client vs backend start) — only client path persists",
            ],
        ],
    )

    add_heading(doc, "Recommendations", 1)
    add_bullets(
        doc,
        [
            "Centralize transcript authorization in one helper used by all read/write/debug endpoints.",
            "On post-call transcription COMPLETE, trigger generateAndStoreSummary (or emit an event the summary service listens to).",
            "Map Transcribe speakers to telemetry user IDs when storing post-call segments, so actorUserId and archive participant lists stay accurate.",
            "Extend merge dedup key with actorUserId and occurredAt (or segment ID when present).",
            "Remove or consolidate duplicate migrations to a single canonical version per table.",
            "Add participant check to /transcription/debug, or restrict to admin/dev profile.",
        ],
    )

    add_heading(doc, "Conclusion", 1)
    add_para(
        doc,
        "The design is coherent: one write service, one merge read service, clear separation between live segments, "
        "cold archive, and derived summaries. Main fragility is timing and authorization around async post-call "
        "transcription, plus archive retrieval depending on S3 availability.",
    )

    add_heading(doc, "Key source files", 1)
    add_bullets(
        doc,
        [
            "backend/core/src/main/java/com/careconnect/service/CallTranscriptService.java",
            "backend/core/src/main/java/com/careconnect/service/CallTranscriptArchiveService.java",
            "backend/core/src/main/java/com/careconnect/service/CallSummaryService.java",
            "backend/core/src/main/java/com/careconnect/service/PostCallTranscriptionService.java",
            "backend/core/src/main/java/com/careconnect/controller/CallController.java",
            "backend/core/src/main/java/com/careconnect/model/CallTranscriptSegment.java",
            "backend/core/src/main/java/com/careconnect/model/CallTranscriptArchive.java",
            "backend/core/src/main/java/com/careconnect/model/CallSummary.java",
            "frontend/lib/services/video_call_service.dart (participantUserIds on end)",
            "backend/core/src/main/java/com/careconnect/service/ChimeService.java (cached join creds)",
            "frontend/lib/widgets/chime_meeting_embed_web.dart",
            "frontend/lib/widgets/post_call_telemetry_summary_screen.dart",
        ],
    )

    save_document(doc, OUTPUT)


if __name__ == "__main__":
    build()
