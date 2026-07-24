"""Generate Word document: POST /api/ai/ask request/response contract design."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX

ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "POST_api_ai_ask_Request_Response_Contract_Design.docx"


def heading(doc, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def para(doc, text: str, bold: bool = False, highlight: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    if highlight:
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def bullets(doc, items: list[str], highlight_indices: set[int] | None = None) -> None:
    highlight_indices = highlight_indices or set()
    for i, item in enumerate(items):
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        if i in highlight_indices:
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def code(doc, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"


def table(doc, headers, rows, highlight_rows: set[int] | None = None) -> None:
    highlight_rows = highlight_rows or set()
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for r_idx, row in enumerate(rows, start=1):
        cells = t.add_row().cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = val
            if r_idx in highlight_rows:
                for p in cells[c_idx].paragraphs:
                    for run in p.runs:
                        run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def build() -> None:
    doc = Document()

    title = doc.add_heading(
        "POST /api/ai/ask — Request/Response Contract Design", 0
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para(doc, "CareConnect Team E — Ask AI Gateway API")
    para(
        doc,
        "Design specification for the records-grounded Ask AI endpoint: hybrid retrieval, "
        "citations, mandatory disclaimers, Tier 1/2 escalation flags, and HITL hold responses. "
        "Synthesizes SRS/TDD, HITL research, hybrid retrieval, RBAC scope, and codebase review.",
    )
    para(
        doc,
        "Status: NOT IMPLEMENTED — existing /v1/api/ai-chat/chat returns raw Bedrock text without "
        "retrieval, citations, disclaimer, or escalation metadata.",
        highlight=True,
    )
    doc.add_paragraph()

    # 1 Overview
    heading(doc, "1. Purpose & Scope", 1)
    bullets(
        doc,
        [
            "Single gateway for natural-language questions over authorized patient records (FR-AI-1, FR-AI-2).",
            "Runs SCC-3 path: RBAC → consent → hybrid retrieval → Bedrock → safety validation → HITL → audit.",
            "Distinct from legacy POST /v1/api/ai-chat/chat (conversation-oriented, no RAG).",
            "Supports text and voice-origin queries (inputModality) through the same contract.",
            "Does not perform side-effecting writes without separate confirmation flow (REQ-SC-5).",
        ],
    )

    heading(doc, "1.1 Requirement traceability", 2)
    table(
        doc,
        ["Requirement", "Contract element"],
        [
            ["FR-AI-1", "403 when patient out of scope; sourceTypes filter"],
            ["FR-AI-2", "citations[] required on every grounded answer"],
            ["FR-AI-3", "disclaimer + aiNotice (records-based, not medical advice)"],
            ["FR-AI-5 / REQ-SC-4", "tier, escalation, held response for Tier 2"],
            ["FR-AI-9", "retrievalMeta.chunksUsed — minimum-necessary context only"],
            ["FR-AI-10 / REQ-SC-9", "auditId on every response; server-side immutable log"],
            ["REQ-SC-1", "Non-dismissible AI notice in response + UI contract"],
            ["REQ-SC-6", "confirmation block separate from answer delivery"],
        ],
    )

    # 2 Endpoint
    heading(doc, "2. Endpoint Definition", 1)
    table(
        doc,
        ["Property", "Value"],
        [
            ["Method", "POST"],
            ["Path", "/api/ai/ask"],
            ["Version", "v1 (prefix /v1 optional — align with /v1/api/* convention)"],
            ["Auth", "Bearer JWT (required)"],
            ["Content-Type", "application/json"],
            ["Permission", "USE_AI_FEATURES + RetrievalScopeService patient access"],
            ["Rate limit", "LangChainGovernanceService / gateway policy"],
            ["SLA", "≤ 5 s p95 end-to-end (NFR-AI-1)"],
        ],
    )
    para(
        doc,
        "Recommended full path: POST /v1/api/ai/ask — parallel to existing /v1/api/ai-chat/* but "
        "implemented on new AiAskController gated by careconnect.ai.ask.enabled.",
        highlight=True,
    )

    # 3 Request
    heading(doc, "3. Request Contract", 1)

    heading(doc, "3.1 Request body (JSON Schema summary)", 2)
    code(
        doc,
        "{\n"
        '  "$schema": "https://json-schema.org/draft/2020-12/schema",\n'
        '  "type": "object",\n'
        '  "required": ["query", "patientId"],\n'
        '  "properties": {\n'
        '    "query": { "type": "string", "minLength": 1, "maxLength": 2000 },\n'
        '    "patientId": { "type": "integer", "format": "int64" },\n'
        '    "sessionId": { "type": "string", "format": "uuid" },\n'
        '    "conversationId": { "type": "string", "format": "uuid" },\n'
        '    "inputModality": { "enum": ["TEXT", "VOICE"], "default": "TEXT" },\n'
        '    "locale": { "type": "string", "pattern": "^(en|es)(-US)?$", "default": "en-US" },\n'
        '    "sourceTypes": {\n'
        '      "type": "array",\n'
        '      "items": { "enum": ["TRANSCRIPT","CALL_SUMMARY","VISIT_SUMMARY","DOCUMENT",\n'
        '                           "CLINICAL_NOTE","USPS_MAIL","MEDICATION_TIMELINE"] }\n'
        '    },\n'
        '    "episodeType": { "enum": ["call", "visit", "any"], "default": "any" },\n'
        '    "clientRequestId": { "type": "string", "maxLength": 64 },\n'
        '    "includeDebugRetrieval": { "type": "boolean", "default": false }\n'
        "  },\n"
        '  "additionalProperties": false\n'
        "}",
    )

    heading(doc, "3.2 Field semantics", 2)
    table(
        doc,
        ["Field", "Required", "Description"],
        [
            ["query", "Yes", "User question in natural language (post-STT text when inputModality=VOICE)"],
            ["patientId", "Yes", "Patient whose records may be retrieved; must pass RetrievalScopeService"],
            ["sessionId", "No", "Stable Ask AI session for audit + optional multi-turn context (not chat DB)"],
            ["conversationId", "No", "Optional link to ai-chat conversation if UI merges surfaces"],
            ["inputModality", "No", "TEXT default; VOICE when query originated from speech_to_text"],
            ["locale", "No", "Disclaimer and safety copy language (SRS Spanish NFR)"],
            ["sourceTypes", "No", "Subset of permitted types; server intersects with RBAC + REQ-SC-7 exclusions"],
            ["episodeType", "No", "Filter summaries/transcripts to call vs visit episodes"],
            ["clientRequestId", "No", "Idempotency key for retry-safe clients"],
            ["includeDebugRetrieval", "No", "Dev/admin only: return retrieval scores in response"],
        ],
    )

    heading(doc, "3.3 Security rules (request)", 2)
    bullets(
        doc,
        [
            "Do NOT accept userId in body — derive caller from JWT sub / UserDetails.",
            "Reject patientId the caller cannot access with 403 + audit (no partial answer).",
            "Sanitize query via InputSanitizationService before retrieval and LLM.",
            "Log query metadata (hash, length, sessionId) — avoid storing raw PHI in audit if policy requires.",
        ],
        highlight_indices={0, 1},
    )

    heading(doc, "3.4 Example request", 2)
    code(
        doc,
        "POST /v1/api/ai/ask HTTP/1.1\n"
        "Authorization: Bearer eyJ...\n"
        "Content-Type: application/json\n\n"
        "{\n"
        '  "query": "When was lisinopril started and has the dose changed?",\n'
        '  "patientId": 42,\n'
        '  "sessionId": "f47ac10b-58cc-4372-a567-0e02b2c3d479",\n'
        '  "inputModality": "TEXT",\n'
        '  "locale": "en-US",\n'
        '  "episodeType": "any"\n'
        "}",
    )

    # 4 Response envelope
    heading(doc, "4. Response Envelope (Common Fields)", 1)
    code(
        doc,
        "// Present on all JSON responses (success and error)\n"
        "{\n"
        '  "success": boolean,\n'
        '  "requestId": "uuid",           // server correlation id\n'
        '  "auditId": "uuid",             // immutable audit ledger key (FR-AI-10)\n'
        '  "sessionId": "uuid | null",    // echoed from request or assigned\n'
        '  "timestamp": "2026-06-28T21:15:00Z",\n'
        '  "deliveryStatus": "DELIVERED | HELD | WITHHELD | NO_RECORDS"\n'
        "}",
    )

    # 5 Success responses
    heading(doc, "5. Success Responses", 1)

    heading(doc, "5.1 Tier 1 — Delivered answer (HTTP 200)", 2)
    code(
        doc,
        "{\n"
        '  "success": true,\n'
        '  "requestId": "a1b2c3d4-...",\n'
        '  "auditId": "e5f6...",\n'
        '  "sessionId": "f47ac10b-...",\n'
        '  "timestamp": "2026-06-28T21:15:00Z",\n'
        '  "deliveryStatus": "DELIVERED",\n'
        '  "tier": 1,\n'
        '  "held": false,\n'
        '  "answer": {\n'
        '    "text": "Based on your records, lisinopril was started on 2026-03-12...",\n'
        '    "format": "markdown"\n'
        "  },\n"
        '  "citations": [ /* AiCitation[] — see §6 */ ],\n'
        '  "disclaimer": {\n'
        '    "text": "This answer is based on your stored health records and is not medical advice.",\n'
        '    "locale": "en-US",\n'
        '    "aiNoticeRequired": true,\n'
        '    "recordsBasedFraming": true\n'
        "  },\n"
        '  "escalation": {\n'
        '    "level": "none | confirm-with-provider | contact-emergency-services",\n'
        '    "prompt": null,\n'
        '    "triggerCodes": []\n'
        "  },\n"
        '  "confirmation": {\n'
        '    "required": false,\n'
        '    "actionType": null,\n'
        '    "prompt": null\n'
        "  },\n"
        '  "retrievalMeta": {\n'
        '    "chunksRetrieved": 8,\n'
        '    "chunksUsed": 4,\n'
        '    "retrievalLatencyMs": 320,\n'
        '    "inferenceLatencyMs": 1100,\n'
        '    "model": { "provider": "bedrock", "modelId": "amazon.nova-lite-v1:0" }\n'
        "  }\n"
        "}",
    )

    heading(doc, "5.2 Tier 1 — No matching records (HTTP 200)", 2)
    code(
        doc,
        "{\n"
        '  "success": true,\n'
        '  "deliveryStatus": "NO_RECORDS",\n'
        '  "tier": 1,\n'
        '  "held": false,\n'
        '  "answer": {\n'
        '    "text": "I could not find any records that answer that question.",\n'
        '    "format": "plain"\n'
        "  },\n"
        '  "citations": [],\n'
        '  "disclaimer": { "text": "...", "aiNoticeRequired": true, "recordsBasedFraming": true },\n'
        '  "escalation": {\n'
        '    "level": "confirm-with-provider",\n'
        '    "prompt": "Would you like to note this question for your care team?",\n'
        '    "triggerCodes": ["NO_MATCHING_RECORDS"]\n'
        "  }\n"
        "}",
    )
    para(doc, "Must NOT return general medical knowledge when NO_RECORDS (FR-AI-2 path).", bold=True)

    heading(doc, "5.3 Tier 2 — Held for human review (HTTP 200 or 202)", 2)
    code(
        doc,
        "{\n"
        '  "success": true,\n'
        '  "deliveryStatus": "HELD",\n'
        '  "tier": 2,\n'
        '  "held": true,\n'
        '  "heldItemId": "8f14e45f-ceea-467a-9a26-386477b17c7a",\n'
        '  "answer": null,\n'
        '  "citations": [],\n'
        '  "message": "We\'re reviewing this before showing it to you.",\n'
        '  "disclaimer": { "text": "...", "aiNoticeRequired": true, "recordsBasedFraming": true },\n'
        '  "escalation": {\n'
        '    "level": "held-for-review",\n'
        '    "prompt": null,\n'
        '    "triggerCodes": ["MEDICATION_CHANGE", "EMERGENCY_SYMPTOM"]\n'
        "  },\n"
        '  "pollUrl": "/v1/api/ai/hitl/{heldItemId}/status"\n'
        "}",
    )
    para(
        doc,
        "Care recipient must NOT receive answer.text or draft citations until Human Reviewer releases. "
        "HTTP 202 Accepted optional when response is explicitly async.",
        highlight=True,
    )

    heading(doc, "5.4 Tier 1 with escalation prompt (medication mention)", 2)
    code(
        doc,
        '  "escalation": {\n'
        '    "level": "confirm-with-provider",\n'
        '    "prompt": "Medication information should be confirmed with your healthcare provider.",\n'
        '    "triggerCodes": ["GENERAL_MEDICATION_MENTION"]\n'
        "  }",
    )

    # 6 Citations
    heading(doc, "6. Citation Object (AiCitation)", 1)
    code(
        doc,
        "{\n"
        '  "citationId": "uuid",\n'
        '  "recordType": "TRANSCRIPT | CALL_SUMMARY | VISIT_SUMMARY | SUMMARY_ITEM |\n'
        '                  "MEDICATION_TIMELINE_EVENT | DOCUMENT | CLINICAL_NOTE | USPS_MAIL",\n'
        '  "sourceRecordId": "string",     // e.g. call_summaries.id, segment id, user_files.id\n'
        '  "chunkId": "uuid",              // retrieval_index_chunk.id when indexed\n'
        '  "title": "Call summary — 2026-06-20",\n'
        '  "excerpt": "Started lisinopril 10mg daily on 2026-03-12...",\n'
        '  "occurredAt": "2026-03-12T14:00:00Z",\n'
        '  "deepLink": "/calls/{callId}/summary#item-{itemId}",\n'
        '  "confidence": 0.92,\n'
        '  "metadata": {\n'
        '    "callId": "...",\n'
        '    "episodeType": "call",\n'
        '    "itemId": "ci-001",\n'
        '    "sourceTurnId": "turn-42",\n'
        '    "speaker": "Dr. Smith",\n'
        '    "startMs": 125000\n'
        "  }\n"
        "}",
    )
    table(
        doc,
        ["recordType", "sourceRecordId", "deepLink pattern"],
        [
            ["TRANSCRIPT", "call_transcript_segments.id", "/calls/{callId}/transcript?t={startMs}"],
            ["CALL_SUMMARY", "call_summaries.id", "/calls/{callId}/summary"],
            ["VISIT_SUMMARY", "visit_summaries.id", "/visits/{visitId}/summary"],
            ["SUMMARY_ITEM", "{summaryId}:{itemId}", "/calls/{callId}/summary#item-{itemId}"],
            ["MEDICATION_TIMELINE_EVENT", "{summaryId}:{itemId}", "Same as summary item"],
            ["DOCUMENT", "user_files.id", "/files/{fileId}"],
            ["CLINICAL_NOTE", "patient_note.id", "/patients/{patientId}/notes/{noteId}"],
            ["USPS_MAIL", "{digestId}:{pieceId}", "/mail/{pieceId}"],
        ],
    )
    bullets(
        doc,
        [
            "Every factual sentence in answer.text must map to ≥1 citation excerpt (validation gate).",
            "Order citations by relevance rank (same order as retrieval merge).",
            "UI: tappable citation opens deepLink; show recordType badge.",
            "Deduplicate by chunkId when multiple sentences cite same source.",
        ],
    )

    # 7 Disclaimer
    heading(doc, "7. Disclaimer & AI Notice Contract", 1)
    table(
        doc,
        ["Field", "Required", "UI behavior"],
        [
            ["disclaimer.text", "Yes", "Always visible with answer; not collapsible by default (REQ-SC-1)"],
            ["disclaimer.aiNoticeRequired", "Yes", "When true, client must show persistent AI-generated label"],
            ["disclaimer.recordsBasedFraming", "Yes", "When true, headline states answer is from records"],
            ["disclaimer.locale", "Yes", "Select en-US / es-US copy server-side"],
        ],
    )
    code(
        doc,
        "// English (default)\n"
        '"This answer is based on your stored health records and is not medical advice."\n\n'
        "// Spanish (SRS NFR)\n"
        '"Esta respuesta se basa en sus registros de salud almacenados y no constituye asesoramiento médico."',
    )

    # 8 Escalation
    heading(doc, "8. Escalation Flags & Tier Model", 1)
    table(
        doc,
        ["escalation.level", "tier", "HTTP", "answer body", "Client action"],
        [
            ["none", "1", "200", "Delivered", "Show disclaimer only"],
            ["confirm-with-provider", "1", "200", "Delivered", "Show escalation.prompt banner + optional CTA"],
            ["contact-emergency-services", "1 or 2", "200/held", "Crisis copy only if Tier 1", "Show emergency UI; may also hold (Tier 2)"],
            ["held-for-review", "2", "200/202", "null", "Show message; poll pollUrl"],
        ],
        highlight_rows={4},
    )
    table(
        doc,
        ["triggerCode", "Typical tier", "Source"],
        [
            ["NO_MATCHING_RECORDS", "1", "Empty hybrid retrieval"],
            ["GENERAL_MEDICATION_MENTION", "1", "SRS Table 8"],
            ["MEDICATION_CHANGE", "2", "TC-E-SC-001"],
            ["EMERGENCY_SYMPTOM", "2", "SRS Table 8"],
            ["DOSAGE_CALC", "2", "SRS Table 8"],
            ["UNSUPPORTED_CLINICAL_CLAIM", "2", "Secondary validation"],
        ],
    )

    # 9 Errors
    heading(doc, "9. Error Responses", 1)
    table(
        doc,
        ["HTTP", "errorCode", "deliveryStatus", "When"],
        [
            ["400", "INVALID_REQUEST", "WITHHELD", "Validation failure (missing query, bad locale)"],
            ["401", "UNAUTHORIZED", "WITHHELD", "Missing/invalid JWT"],
            ["403", "FORBIDDEN_SCOPE", "WITHHELD", "Patient not in caller scope / consent denied"],
            ["422", "SAFETY_VALIDATION_FAILED", "WITHHELD", "Schema or secondary validation failed"],
            ["429", "RATE_LIMITED", "WITHHELD", "Governance rate limit"],
            ["503", "RETRIEVAL_UNAVAILABLE", "WITHHELD", "Index/DB unavailable — no ungrounded fallback"],
        ],
        highlight_rows={5, 6},
    )
    code(
        doc,
        "{\n"
        '  "success": false,\n'
        '  "requestId": "...",\n'
        '  "auditId": "...",\n'
        '  "deliveryStatus": "WITHHELD",\n'
        '  "error": {\n'
        '    "code": "SAFETY_VALIDATION_FAILED",\n'
        '    "message": "We could not verify this answer against your records.",\n'
        '    "details": []\n'
        "  }\n"
        "}",
    )

    # 10 Gateway flow
    heading(doc, "10. Server Processing Flow", 1)
    code(
        doc,
        "POST /api/ai/ask\n"
        "  1. Authenticate JWT → resolve User + Role\n"
        "  2. Validate request body\n"
        "  3. RetrievalScopeService.assertCanAsk(caller, patientId, sourceTypes)\n"
        "  4. HybridRetrievalService.retrieve(query, scope) → chunks[]\n"
        "  5. IF chunks empty → build NO_RECORDS response (§5.2)\n"
        "  6. AiAskService.assemblePrompt(query, chunks) → Bedrock via LlmRouter\n"
        "  7. Parse structured LLM output { answerText, citationRefs[] }\n"
        "  8. Map citationRefs → AiCitation[] from chunk metadata\n"
        "  9. SafetyValidationService.validate(answer, citations, query)\n"
        " 10. TierClassifier.assignTier → Tier 1 or 2\n"
        " 11. IF Tier 2 → persist held_item; return §5.3 (no answer to patient)\n"
        " 12. Attach disclaimer + escalation for Tier 1\n"
        " 13. AuditService.writeImmutable(auditId, …)\n"
        " 14. Return §5.1",
    )

    # 11 DTO mapping
    heading(doc, "11. Proposed Java DTOs (Backend)", 1)
    code(
        doc,
        "// com.careconnect.dto.ai.AiAskRequest\n"
        "record AiAskRequest(\n"
        "    @NotBlank @Size(max=2000) String query,\n"
        "    @NotNull Long patientId,\n"
        "    UUID sessionId,\n"
        "    UUID conversationId,\n"
        "    InputModality inputModality,\n"
        "    String locale,\n"
        "    List<RetrievalRecordType> sourceTypes,\n"
        "    EpisodeTypeFilter episodeType,\n"
        "    String clientRequestId,\n"
        "    boolean includeDebugRetrieval\n"
        ") {}\n\n"
        "// com.careconnect.dto.ai.AiAskResponse\n"
        "record AiAskResponse(\n"
        "    boolean success,\n"
        "    UUID requestId,\n"
        "    UUID auditId,\n"
        "    UUID sessionId,\n"
        "    Instant timestamp,\n"
        "    DeliveryStatus deliveryStatus,\n"
        "    int tier,\n"
        "    boolean held,\n"
        "    UUID heldItemId,\n"
        "    AiAnswerBlock answer,\n"
        "    List<AiCitation> citations,\n"
        "    AiDisclaimer disclaimer,\n"
        "    AiEscalation escalation,\n"
        "    AiConfirmationHint confirmation,\n"
        "    AiRetrievalMeta retrievalMeta,\n"
        "    String message,\n"
        "    String pollUrl,\n"
        "    AiErrorBlock error\n"
        ") {}",
    )
    para(
        doc,
        "Do not extend legacy ChatRequest/ChatResponse — they lack citations, tier, and disclaimer "
        "structures and are tied to conversation persistence.",
        highlight=True,
    )

    # 12 Comparison
    heading(doc, "12. Comparison: Legacy vs Ask AI Gateway", 1)
    table(
        doc,
        ["Aspect", "POST /v1/api/ai-chat/chat", "POST /v1/api/ai/ask (designed)"],
        [
            ["Purpose", "Multi-turn chat persistence", "Records-grounded Q&A"],
            ["Retrieval", "None (raw Bedrock)", "Hybrid FTS + pgvector"],
            ["citations[]", "Missing", "Required"],
            ["disclaimer", "Missing", "Required object"],
            ["escalation / tier", "Missing", "Required"],
            ["HITL hold", "Missing", "heldItemId + pollUrl"],
            ["patientId in body", "Yes + userId", "patientId only; user from JWT"],
            ["Controller", "AIChatController", "AiAskController (new)"],
        ],
        highlight_rows={2, 3, 4, 5, 6, 7},
    )

    # 13 Flutter
    heading(doc, "13. Client (Flutter) Integration Notes", 1)
    bullets(
        doc,
        [
            "New AiAskService calling POST /v1/api/ai/ask — separate from ai_chat_service.dart.",
            "Render disclaimer.aiNoticeRequired as persistent banner (SRS Ask AI screen mock).",
            "Map citations to tappable chips → deepLink navigation.",
            "When held=true, show message and poll pollUrl every 5s until DELIVERED or REJECTED.",
            "When escalation.level=confirm-with-provider, show non-blocking banner with CTA.",
            "Voice: set inputModality=VOICE after speech_to_text; same response handling.",
        ],
    )

    # 14 Tests
    heading(doc, "14. Test Plan Mapping", 1)
    table(
        doc,
        ["Test", "Assert on contract"],
        [
            ["TC-E-AI-001", "200 + citations[] non-empty for matched records"],
            ["TC-E-AI-002", "Every citation has recordType, excerpt, sourceRecordId"],
            ["TC-E-AI-003", "disclaimer.text present; aiNoticeRequired=true"],
            ["TC-E-SC-001", "Tier 2: held=true, answer=null, heldItemId set"],
            ["TC-SCC-03", "triggerCodes populated; deliveryStatus=HELD"],
            ["TC-E-AI-011", "Medication timeline answer cites MEDICATION_TIMELINE_EVENT types"],
            ["No-records path", "deliveryStatus=NO_RECORDS; citations=[]; no clinical advice"],
        ],
    )

    # 15 Related
    heading(doc, "15. Related Documents & References", 1)
    bullets(
        doc,
        [
            "docs/HITL_Tier1_Tier2_Escalation_and_Hold_Release_Workflow.docx",
            "docs/Hybrid_Retrieval_Scope_PostgreSQL_FullText_pgvector.docx",
            "docs/RBAC_Scoped_Retrieval_Source_Types.docx",
            "docs/pgvector_Embedding_Strategy_Summaries_Mail_Documents.docx",
            "docs/Team_E_Implementation_Task_Backlog.docx (task 5.3)",
            "backend/core/src/main/java/com/careconnect/dto/ChatRequest.java (legacy — do not reuse)",
            "backend/core/src/main/java/com/careconnect/controller/AIChatController.java",
        ],
    )

    heading(doc, "16. Conclusion", 1)
    para(
        doc,
        "POST /api/ai/ask is the Team E gateway contract: a structured JSON API that always returns "
        "delivery status, mandatory disclaimer metadata, citation provenance, and escalation/tier "
        "signals. Tier 2 responses intentionally omit the answer body until HITL release. "
        "Implementation should use new DTOs and AiAskController rather than extending the legacy chat path.",
        highlight=True,
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    try:
        doc.save(OUTPUT)
        print(f"Created: {OUTPUT}")
    except PermissionError:
        alt = OUTPUT.with_name(f"{OUTPUT.stem}_refresh{OUTPUT.suffix}")
        doc.save(alt)
        print(f"Created (original locked): {alt}")


if __name__ == "__main__":
    build()
