"""Generate Word document: Audit logging design for Ask AI queries, responses, citations, escalations."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX

ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "Ask_AI_Audit_Logging_Design.docx"


def heading(doc, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def para(doc, text: str, bold: bool = False, highlight: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    if highlight:
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def bullets(doc, items: list[str], highlight_indices: set[int] | None = None) -> None:
    highlight_indices = highlight_indices or set()
    for i, item in enumerate(items):
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        if i in highlight_indices:
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def code(doc, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"


def table(
    doc,
    headers,
    rows,
    highlight_rows: set[int] | None = None,
    highlight_cells: set[tuple[int, int]] | None = None,
) -> None:
    highlight_rows = highlight_rows or set()
    highlight_cells = highlight_cells or set()
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for r_idx, row in enumerate(rows, start=1):
        cells = t.add_row().cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = val
            if r_idx in highlight_rows or (r_idx, c_idx) in highlight_cells:
                for p in cells[c_idx].paragraphs:
                    for run in p.runs:
                        run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def build() -> None:
    doc = Document()

    title = doc.add_heading(
        "Ask AI Audit Logging — Queries, Responses, Citations & Escalations",
        0,
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para(doc, "CareConnect Team E — FR-AI-10 / REQ-SC-9 Implementation Design")
    para(
        doc,
        "Design specification for immutable audit logging across the SCC-3 Ask AI pipeline: "
        "every query, retrieval scope, grounded response, citation set, tier assignment, "
        "escalation metadata, HITL hold/release, and delivery outcome. Unifies scattered "
        "ChatAuditService (SLF4J metadata only) and SecurityAuditService (governance logs) "
        "into a persisted append-only ledger aligned with ai_held_item and SafetyPipeline.",
    )
    para(
        doc,
        "Status: DESIGN ONLY — ChatAuditService writes structured logs only (no DB, no citations, "
        "no tier/escalation). No ai_ask_audit_* tables. SecurityAuditService is SLF4J-only. "
        "EvvAuditEvent provides the closest persisted append-only pattern in codebase.",
        highlight=True,
    )
    doc.add_paragraph()

    # 1 Goals
    heading(doc, "1. Design Goals", 1)
    bullets(
        doc,
        [
            "FR-AI-10: Log every Ask AI query, response, citations, escalation, and delivery status.",
            "REQ-SC-9: Immutable, append-only audit — no UPDATE/DELETE on audit rows from application.",
            "Single auditId returned on every API response (success, error, held) for correlation.",
            "Event-sourced lifecycle: query → scope → retrieval → LLM → validation → tier → delivery/HITL.",
            "Citation audit: chunk IDs, record types, source record IDs — not full chunk text by default.",
            "Escalation audit: level, triggerCodes[], tier — supports TC-AI-06 / TC-SCC-03.",
            "403/422 paths audited with no AI output leaked — scope denial fully traceable.",
            "Admin/reviewer read APIs with RBAC — separate from patient-facing Ask AI response.",
            "PHI minimization: configurable query/answer storage (hash-only vs encrypted full text).",
        ],
        highlight_indices={0, 1, 2, 4, 5},
    )

    heading(doc, "1.1 Position in pipeline", 2)
    code(
        doc,
        "POST /api/ai/ask\n"
        "  [0] AiAuditService.startRequest()     → auditId, requestId assigned\n"
        "  [1] scope resolve / deny              → SCOPE_GRANTED | SCOPE_DENIED event\n"
        "  [2] hybrid retrieval                  → RETRIEVAL_COMPLETED event + chunk refs\n"
        "  [3] LLM inference                     → LLM_COMPLETED event + model metadata\n"
        "  [4] SafetyPipeline                    → VALIDATION_COMPLETED + findings\n"
        "  [5] TierClassifier                    → TIER_ASSIGNED event\n"
        "  [6] HitlService (if Tier 2)           → HITL_HELD event (links heldItemId)\n"
        "  [7] Response build                    → DELIVERED | NO_RECORDS | HELD snapshot\n"
        "  [8] AiAuditService.finalizeRecord()   → immutable summary row + return auditId",
    )

    # 2 Architecture
    heading(doc, "2. Component Architecture", 1)
    table(
        doc,
        ["Component", "Package", "Responsibility"],
        [
            ["AiAuditService", "service.ai.audit", "Orchestrates audit lifecycle; public API for pipeline"],
            ["AiAskAuditEventWriter", "service.ai.audit", "Append-only inserts to ai_ask_audit_event"],
            ["AiAskAuditRecordWriter", "service.ai.audit", "Insert immutable summary on request completion"],
            ["AiCitationAuditMapper", "service.ai.audit", "Serialize citations[] → audit-safe JSON"],
            ["AiEscalationAuditMapper", "service.ai.audit", "Serialize escalation + tier + triggerCodes"],
            ["AiQueryRedactionPolicy", "service.ai.audit", "Hash or encrypt query/answer per config"],
            ["AiAuditReadService", "service.ai.audit", "Admin/reviewer query APIs"],
            ["AiSafetyAuditService", "service.ai.safety", "Delegates HITL sub-events to same audit_id (merged)"],
        ],
        highlight_cells={
            (1, 0), (2, 0), (3, 0), (4, 0), (5, 0),
        },
    )

    heading(doc, "2.1 Relationship to existing audit code", 2)
    table(
        doc,
        ["Existing", "Disposition"],
        [
            ["ChatAuditService", "Retain for legacy /ai-chat/chat; NOT used for /api/ai/ask"],
            ["SecurityAuditService", "Retain for governance violations; cross-reference auditId when available"],
            ["LangChainGovernanceService", "Emit GOVERNANCE_* events linked to auditId on rate limit"],
            ["EvvAuditEvent + AuditLogger", "Pattern reference — append-only JSONB details"],
            ["ai_safety_audit_event (HITL design)", "Merge into unified ai_ask_audit_event schema"],
            ["InputSanitizationService", "Log SANITIZATION_APPLIED event when query modified pre-retrieval"],
        ],
        highlight_rows={1, 4},
    )

    # 3 Data model
    heading(doc, "3. Database Schema (Flyway)", 1)

    heading(doc, "3.1 ai_ask_audit_record — immutable completion snapshot", 2)
    code(
        doc,
        "CREATE TABLE ai_ask_audit_record (\n"
        "  audit_id UUID PRIMARY KEY,\n"
        "  request_id UUID NOT NULL UNIQUE,\n"
        "  session_id UUID,\n"
        "  client_request_id VARCHAR(64),\n"
        "  patient_id BIGINT NOT NULL,\n"
        "  caller_user_id BIGINT NOT NULL,\n"
        "  caller_role VARCHAR(32) NOT NULL,\n"
        "  input_modality VARCHAR(8) NOT NULL DEFAULT 'TEXT',\n"
        "  locale VARCHAR(10) NOT NULL DEFAULT 'en-US',\n"
        "  -- Query (PHI policy)\n"
        "  query_text_hash VARCHAR(64) NOT NULL,\n"
        "  query_length INT NOT NULL,\n"
        "  query_text_encrypted BYTEA,              -- null when hash-only mode\n"
        "  -- Outcome\n"
        "  delivery_status VARCHAR(24) NOT NULL,\n"
        "  tier SMALLINT NOT NULL,\n"
        "  held BOOLEAN NOT NULL DEFAULT false,\n"
        "  held_item_id UUID,\n"
        "  error_code VARCHAR(40),\n"
        "  -- Response snapshot (PHI policy)\n"
        "  answer_text_hash VARCHAR(64),\n"
        "  answer_length INT,\n"
        "  answer_text_encrypted BYTEA,\n"
        "  -- Citations & escalation (audit-safe JSON)\n"
        "  citations_json JSONB NOT NULL DEFAULT '[]',\n"
        "  escalation_json JSONB NOT NULL DEFAULT '{}',\n"
        "  trigger_codes JSONB NOT NULL DEFAULT '[]',\n"
        "  validation_findings_json JSONB,\n"
        "  -- Retrieval & model metadata\n"
        "  retrieval_meta_json JSONB NOT NULL DEFAULT '{}',\n"
        "  scope_json JSONB NOT NULL DEFAULT '{}',\n"
        "  model_provider VARCHAR(32),\n"
        "  model_id VARCHAR(128),\n"
        "  -- Timing\n"
        "  total_latency_ms INT,\n"
        "  created_at TIMESTAMPTZ NOT NULL DEFAULT now()\n"
        ");\n"
        "-- INSERT only from application; no UPDATE trigger blocks mutations\n"
        "CREATE INDEX idx_audit_patient_created ON ai_ask_audit_record(patient_id, created_at DESC);\n"
        "CREATE INDEX idx_audit_caller_created ON ai_ask_audit_record(caller_user_id, created_at DESC);\n"
        "CREATE INDEX idx_audit_session ON ai_ask_audit_record(session_id);",
    )

    heading(doc, "3.2 ai_ask_audit_event — append-only lifecycle", 2)
    code(
        doc,
        "CREATE TABLE ai_ask_audit_event (\n"
        "  id UUID PRIMARY KEY DEFAULT gen_random_uuid(),\n"
        "  audit_id UUID NOT NULL,\n"
        "  event_type VARCHAR(48) NOT NULL,\n"
        "  event_sequence INT NOT NULL,\n"
        "  actor_user_id BIGINT,\n"
        "  payload_json JSONB NOT NULL DEFAULT '{}',\n"
        "  created_at TIMESTAMPTZ NOT NULL DEFAULT now(),\n"
        "  UNIQUE (audit_id, event_sequence)\n"
        ");\n"
        "CREATE INDEX idx_audit_event_audit ON ai_ask_audit_event(audit_id, event_sequence);\n\n"
        "-- DB enforcement (PostgreSQL):\n"
        "CREATE RULE ai_ask_audit_event_no_update AS ON UPDATE TO ai_ask_audit_event DO INSTEAD NOTHING;\n"
        "CREATE RULE ai_ask_audit_event_no_delete AS ON DELETE TO ai_ask_audit_event DO INSTEAD NOTHING;\n"
        "CREATE RULE ai_ask_audit_record_no_update AS ON UPDATE TO ai_ask_audit_record DO INSTEAD NOTHING;\n"
        "CREATE RULE ai_ask_audit_record_no_delete AS ON DELETE TO ai_ask_audit_record DO INSTEAD NOTHING;",
    )
    para(
        doc,
        "HITL release/reject appends HITL_RELEASED / HITL_REJECTED events to the same audit_id "
        "chain; ai_held_item.audit_id FK references ai_ask_audit_record.audit_id.",
        highlight=True,
    )

    # 4 Event types
    heading(doc, "4. Audit Event Types & Payloads", 1)
    table(
        doc,
        ["event_type", "When", "Key payload fields"],
        [
            ["REQUEST_STARTED", "Gateway entry", "requestId, inputModality, voiceCaptureMeta, clientRequestId"],
            ["SCOPE_GRANTED", "RetrievalScopeService OK", "allowedSourceTypes[], consentFlags"],
            ["SCOPE_DENIED", "403 FORBIDDEN_SCOPE", "denialReason, requestedPatientId"],
            ["SANITIZATION_APPLIED", "Query modified", "originalHash, sanitizedHash, violationType"],
            ["RETRIEVAL_COMPLETED", "Hybrid search done", "chunksRetrieved, chunkIds[], intent, latencyMs"],
            ["NO_RECORDS", "Empty retrieval", "shortCircuit: true"],
            ["LLM_COMPLETED", "GroundedLlmService", "provider, modelId, inferenceLatencyMs, citationRefCount"],
            ["VALIDATION_COMPLETED", "SafetyPipeline", "findings[], decision: DELIVER|HOLD|BLOCK"],
            ["TIER_ASSIGNED", "TierClassifier", "tier, triggerCodes[], escalationLevel"],
            ["CITATIONS_ASSEMBLED", "CitationAssembler", "citationIds[], recordTypes[], sourceRecordIds[]"],
            ["DELIVERED", "Tier 1 to patient", "deliveryStatus, answerHash, disclaimerLocale"],
            ["HELD", "Tier 2 hold", "heldItemId, triggerCodes[]"],
            ["HITL_RELEASED", "Reviewer approve", "reviewerUserId, edited: boolean, finalAnswerHash"],
            ["HITL_REJECTED", "Reviewer reject", "reviewerUserId, reasonCode"],
            ["HITL_EXPIRED", "TTL job", "expiresAt"],
            ["ERROR", "4xx/5xx terminal", "errorCode, httpStatus, deliveryStatus: WITHHELD"],
            ["GOVERNANCE_BLOCKED", "Rate limit", "governanceAction from LangChainGovernanceService"],
        ],
        highlight_rows={3, 8, 10, 11, 12, 13},
    )

    # 5 Query audit
    heading(doc, "5. Query Audit", 1)
    table(
        doc,
        ["Field", "Stored where", "Policy"],
        [
            ["query (raw text)", "query_text_encrypted OR omitted", "Config: HASH_ONLY (default prod) vs ENCRYPTED_FULL"],
            ["query_text_hash", "ai_ask_audit_record", "SHA-256(query + patientId + salt) — always stored"],
            ["query_length", "ai_ask_audit_record", "Always stored"],
            ["inputModality", "ai_ask_audit_record", "TEXT | VOICE"],
            ["voiceCaptureMeta", "REQUEST_STARTED payload", "sttEngine, captureDurationMs — no audio"],
            ["sessionId", "ai_ask_audit_record", "Multi-turn correlation"],
            ["sanitized flag", "SANITIZATION_APPLIED event", "When InputSanitizationService modifies query"],
        ],
        highlight_cells={(1, 2), (2, 2)},
    )
    bullets(
        doc,
        [
            "Never log raw JWT, API keys, or full IP — use anonymizedIp like ChatAuditService if needed.",
            "403 scope denial: log query hash + denial reason; do NOT log retrieved chunks.",
            "Duplicate clientRequestId: idempotent — return existing auditId, append DUPLICATE_REQUEST event.",
        ],
    )

    # 6 Response audit
    heading(doc, "6. Response Audit", 1)
    table(
        doc,
        ["deliveryStatus", "answer audit", "Patient received"],
        [
            ["DELIVERED", "answer_hash + length + optional encrypted full text", "Full answer + citations"],
            ["NO_RECORDS", "safe no-records message hash", "No-records template only"],
            ["HELD", "draft_answer on ai_held_item only — NOT in patient-visible audit record until release", "Held message only"],
            ["WITHHELD", "error message hash; no draft", "Safe error copy"],
            ["WITHHELD_PERMANENTLY", "HITL_REJECTED final state", "Rejection fallback message"],
        ],
        highlight_rows={2, 3},
    )
    code(
        doc,
        "// Tier 2 flow — two-phase audit record:\n"
        "// Phase 1 (HELD): ai_ask_audit_record with delivery_status=HELD, answer fields null\n"
        "// Phase 2 (HITL_RELEASED): NEW append-only events + optional ai_ask_audit_delivery\n"
        "//   supplement row OR DELIVERED event payload with final_answer_hash + citations_json\n"
        "\n"
        "// Recommended: keep single ai_ask_audit_record immutable at HELD time;\n"
        "//   add ai_ask_audit_delivery_supplement table (append-only) for post-release snapshot",
    )

    heading(doc, "6.1 ai_ask_audit_delivery_supplement (HITL completion)", 2)
    code(
        doc,
        "CREATE TABLE ai_ask_audit_delivery_supplement (\n"
        "  id UUID PRIMARY KEY,\n"
        "  audit_id UUID NOT NULL REFERENCES ai_ask_audit_record(audit_id),\n"
        "  delivery_status VARCHAR(24) NOT NULL,\n"
        "  final_answer_hash VARCHAR(64),\n"
        "  citations_json JSONB NOT NULL,\n"
        "  reviewer_user_id BIGINT,\n"
        "  reviewed_at TIMESTAMPTZ NOT NULL,\n"
        "  created_at TIMESTAMPTZ NOT NULL DEFAULT now()\n"
        ");  -- INSERT only; satisfies REQ-SC-9 immutability of original record",
    )

    # 7 Citations audit
    heading(doc, "7. Citation Audit", 1)
    para(
        doc,
        "citations_json on ai_ask_audit_record stores an audit-safe subset of AiCitation — "
        "sufficient for FR-AI-10 traceability without duplicating full chunk text.",
        bold=True,
    )
    code(
        doc,
        "[\n"
        "  {\n"
        '    "citationId": "uuid",\n'
        '    "recordType": "MEDICATION_TIMELINE_EVENT",\n'
        '    "sourceRecordId": "summary-123:item-456",\n'
        '    "chunkId": "uuid",\n'
        '    "excerptHash": "sha256...",\n'
        '    "excerptLength": 142,\n'
        '    "occurredAt": "2026-06-28",\n'
        '    "deepLink": "/calls/789/summary#item-456",\n'
        '    "metadata": { "itemId": "item-456", "episodeType": "call" }\n'
        "  }\n"
        "]",
    )
    bullets(
        doc,
        [
            "CITATIONS_ASSEMBLED event lists citationIds[] + chunkIds[] for retrieval correlation.",
            "RETRIEVAL_COMPLETED event lists all chunkIds considered (including unused chunks) when debug enabled.",
            "Validation audit: map each citationId to answer sentence refs in VALIDATION_COMPLETED payload.",
            "Mismatched citation (validation failure): log finding UNSUPPORTED_CLAIM with citationId refs.",
        ],
        highlight_indices={3},
    )

    # 8 Escalation audit
    heading(doc, "8. Escalation & Tier Audit", 1)
    code(
        doc,
        "escalation_json: {\n"
        '  "level": "confirm-with-provider",\n'
        '  "prompt": "Medication information should be confirmed...",\n'
        '  "triggerCodes": ["GENERAL_MEDICATION_MENTION"]\n'
        "}\n\n"
        "trigger_codes (top-level duplicate for indexing): [\"GENERAL_MEDICATION_MENTION\"]\n"
        "tier: 1 | 2",
    )
    table(
        doc,
        ["triggerCode", "Typical tier", "Audit significance"],
        [
            ["EMERGENCY_SYMPTOM", "2", "TC-AI-06 — HITL or crisis copy"],
            ["MEDICATION_CHANGE", "2", "TC-E-SC-001 — hold before delivery"],
            ["UNSUPPORTED_CLAIM", "2 or BLOCK", "Citation validation failure"],
            ["NO_MATCHING_RECORDS", "1", "Escalation without HITL"],
            ["GENERAL_MEDICATION_MENTION", "1", "confirm-with-provider prompt"],
            ["DOSAGE_CALC", "2", "Safety hold"],
        ],
        highlight_rows={1, 2},
    )
    bullets(
        doc,
        [
            "TIER_ASSIGNED event is authoritative — escalation_json on record mirrors final state at response time.",
            "Tier 2: escalation.level=held-for-review; triggerCodes copied to ai_held_item.trigger_codes.",
            "Confirmation flows (REQ-SC-5): separate CONFIRMATION_REQUESTED / CONFIRMATION_COMPLETED events.",
        ],
    )

    # 9 Service API
    heading(doc, "9. AiAuditService API", 1)
    code(
        doc,
        "public interface AiAuditService {\n"
        "  AuditContext startRequest(AiAskRequest req, User caller, UUID requestId);\n"
        "  void appendEvent(UUID auditId, AuditEventType type, Map<String,Object> payload, Long actorUserId);\n"
        "  void finalizeRecord(AuditContext ctx, AiAskResponse response, AuditFinalizeOptions opts);\n"
        "  void finalizeError(AuditContext ctx, AiErrorBlock error, int httpStatus);\n"
        "  Optional<UUID> findByClientRequestId(String clientRequestId);\n"
        "}\n\n"
        "record AuditContext(\n"
        "  UUID auditId,\n"
        "  UUID requestId,\n"
        "  int nextEventSequence,\n"
        "  Instant startedAt\n"
        ")",
    )

    heading(doc, "9.1 Orchestrator integration", 2)
    code(
        doc,
        "AuditContext audit = auditService.startRequest(req, caller, requestId);\n"
        "try {\n"
        "  scope = scopeService.resolve(...);\n"
        "  auditService.appendEvent(audit.auditId(), SCOPE_GRANTED, scopePayload, caller.getId());\n"
        "  chunks = hybridRetrieval.search(...);\n"
        "  auditService.appendEvent(audit.auditId(), RETRIEVAL_COMPLETED, chunkPayload, ...);\n"
        "  // ... LLM, safety, tier ...\n"
        "  response = responseBuilder.build(...);\n"
        "  auditService.finalizeRecord(audit, response, opts);\n"
        "  return response.withAuditId(audit.auditId());\n"
        "} catch (ForbiddenScopeException e) {\n"
        "  auditService.appendEvent(audit.auditId(), SCOPE_DENIED, ...);\n"
        "  auditService.finalizeError(audit, FORBIDDEN_SCOPE, 403);\n"
        "  throw e;\n"
        "}",
    )

    # 10 Read APIs
    heading(doc, "10. Audit Read APIs (Admin / Compliance)", 1)
    table(
        doc,
        ["Method", "Path", "Permission", "Returns"],
        [
            ["GET", "/v1/api/ai/audit/{auditId}", "VIEW_AI_AUDIT or ADMIN", "Record + event chain"],
            ["GET", "/v1/api/ai/audit/patient/{patientId}", "VIEW_AI_AUDIT + patient scope", "Paginated records"],
            ["GET", "/v1/api/ai/audit/session/{sessionId}", "VIEW_AI_AUDIT or self-session", "Session timeline"],
            ["GET", "/v1/api/ai/audit/held/{heldItemId}", "REVIEW_AI_HELD_ITEMS", "Audit + HITL events"],
        ],
        highlight_rows={1, 3},
    )
    bullets(
        doc,
        [
            "Patient role: may view own audit summary (delivery status, timestamp) — NOT draft held content.",
            "Human Reviewer: full event chain + validation findings for held items.",
            "Admin: export CSV/JSON for compliance; hash-only mode returns no decryptable query text.",
            "Never expose query_text_encrypted to client APIs without AUDIT_DECRYPT permission + KMS.",
        ],
        highlight_indices={0, 3},
    )

    # 11 PHI policy
    heading(doc, "11. PHI Minimization & Retention", 1)
    table(
        doc,
        ["Property", "Default", "Effect"],
        [
            ["careconnect.ai.audit.storage-mode", "HASH_ONLY", "query/answer: hash + length only"],
            ["careconnect.ai.audit.storage-mode", "ENCRYPTED_FULL", "AES-256/KMS on query_text_encrypted fields"],
            ["careconnect.ai.audit.retention-days", "2555", "~7 years healthcare default"],
            ["careconnect.ai.audit.log-chunk-ids", "true", "Always log chunk UUIDs"],
            ["careconnect.ai.audit.log-excerpt-hash", "true", "SHA-256 of citation excerpt"],
            ["careconnect.ai.audit.log-full-excerpt", "false", "Dev only — full excerpt in payload"],
        ],
        highlight_rows={1, 5},
    )
    bullets(
        doc,
        [
            "Scheduled purge job archives to cold storage before DELETE (DB rules block app DELETE).",
            "Cross-reference LLM provider audit (Task 0.5): model_id on ai_ask_audit_record.",
            "Align with ChatAuditService hashed user IDs OR store caller_user_id for authorized audit access.",
        ],
    )

    # 12 Comparison
    heading(doc, "12. Gap Analysis vs Current Code", 1)
    table(
        doc,
        ["Capability", "ChatAuditService today", "AiAuditService design"],
        [
            ["Persistence", "SLF4J only", "PostgreSQL ai_ask_audit_* tables"],
            ["Query content", "Length only", "Hash + optional encrypted full text"],
            ["Citations", "Not logged", "citations_json + CITATIONS_ASSEMBLED event"],
            ["Escalation / tier", "Not logged", "escalation_json + TIER_ASSIGNED event"],
            ["HITL lifecycle", "Not logged", "HELD → RELEASED/REJECTED event chain"],
            ["403 scope denial", "Not logged", "SCOPE_DENIED + immutable record"],
            ["auditId in API", "Not returned", "UUID on every AiAskResponse"],
            ["Immutability", "N/A", "DB rules — no UPDATE/DELETE"],
            ["FR-AI-10 compliance", "Partial metadata", "Full requirement coverage"],
        ],
        highlight_rows={1, 2, 3, 4, 5, 6, 8},
    )

    # 13 Tests
    heading(doc, "13. Test Alignment", 1)
    table(
        doc,
        ["Test", "Audit validation"],
        [
            ["TC-AI-06", "High-risk query → TIER_ASSIGNED tier=2 + HELD event + auditId in response"],
            ["TC-SCC-03", "Full event chain REQUEST_STARTED through DELIVERED or HELD"],
            ["TC-E-SC-001", "MEDICATION_CHANGE in trigger_codes on HELD record"],
            ["TC-E-SC-002", "HITL_EXPIRED event; no DELIVERED without HITL_RELEASED"],
            ["Unit", "Append-only: UPDATE on ai_ask_audit_event fails at DB"],
            ["Unit", "Idempotent clientRequestId returns same auditId"],
            ["Integration", "403 path: SCOPE_DENIED + finalizeError; no citations_json populated"],
            ["FR-AI-10", "Every successful /api/ai/ask response includes auditId matching DB row"],
        ],
        highlight_rows={1, 2, 6},
    )

    # 14 Implementation phases
    heading(doc, "14. Implementation Phases", 1)
    table(
        doc,
        ["Phase", "Deliverable", "Depends on"],
        [
            ["P0", "Flyway ai_ask_audit_record + ai_ask_audit_event + DB immutability rules", "—"],
            ["P0", "AiAuditService.startRequest / appendEvent / finalizeRecord", "AiAskController"],
            ["P0", "Wire orchestrator stages 0–8; auditId on AiAskResponse", "Retrieval orchestration"],
            ["P1", "citations_json + escalation_json mappers", "CitationAssembler"],
            ["P1", "SCOPE_DENIED + ERROR paths; 403/422 audit", "RetrievalScopeService"],
            ["P1", "Merge HITL events; ai_ask_audit_delivery_supplement", "HitlService"],
            ["P2", "Admin read APIs + VIEW_AI_AUDIT permission", "RBAC"],
            ["P2", "HASH_ONLY vs ENCRYPTED_FULL config + KMS", "REQ-SC-9 compliance review"],
            ["P2", "Retention archive job", "Ops"],
        ],
        highlight_rows={1, 2, 5, 6},
    )
    para(
        doc,
        "Task backlog 6.8: Immutable audit ledger — query, tier, hold, release, citations. "
        "Build with Tier 6 safety/HITL (P0) — audit must capture hold/release before patient rollout.",
        highlight=True,
    )

    # 15 Event sequence example
    heading(doc, "15. Example Event Sequence (Tier 2 Hold)", 1)
    code(
        doc,
        "audit_id: 8f14e45f-ceea-467a-9a26-386477b17c7a\n"
        "seq  type                    payload summary\n"
        "---  ----------------------  ------------------------------------------\n"
        " 1   REQUEST_STARTED         inputModality=VOICE, queryHash=abc...\n"
        " 2   SCOPE_GRANTED           allowedSourceTypes=[TRANSCRIPT,CALL_SUMMARY,...]\n"
        " 3   RETRIEVAL_COMPLETED     chunkIds=[c1,c2,c3], intent=MEDICATION_TIMELINE\n"
        " 4   LLM_COMPLETED           modelId=amazon.nova-lite-v1:0, latency=980ms\n"
        " 5   VALIDATION_COMPLETED    findings=[MEDICATION_CHANGE], decision=HOLD_TIER2\n"
        " 6   TIER_ASSIGNED           tier=2, triggerCodes=[MEDICATION_CHANGE]\n"
        " 7   CITATIONS_ASSEMBLED     citationIds=[cit1,cit2]\n"
        " 8   HELD                    heldItemId=..., deliveryStatus=HELD\n"
        "     (ai_ask_audit_record inserted: tier=2, held=true, citations_json=[...])\n"
        " 9   HITL_RELEASED           reviewerUserId=55, edited=false  [days later]\n"
        "     (ai_ask_audit_delivery_supplement: DELIVERED, final_answer_hash=...)",
    )

    # 16 Related docs
    heading(doc, "16. Related Documents", 1)
    bullets(
        doc,
        [
            "docs/POST_api_ai_ask_Request_Response_Contract_Design.docx (auditId response field)",
            "docs/Secondary_Validation_and_Tier2_HITL_Hold_Release_Design.docx (ai_safety_audit_event merge)",
            "docs/Retrieval_Orchestration_RBAC_Hybrid_Context_LLM_Design.docx (stage 6f AiAuditService)",
            "docs/HITL_Tier1_Tier2_Escalation_and_Hold_Release_Workflow.docx (REQ-SC-9 requirements)",
            "docs/Voice_Query_Routing_Same_Retrieval_Pipeline_Design.docx (inputModality audit)",
            "docs/LLM_Provider_Abstraction_Bedrock_TeamA_BAA_Fallback.docx (model provider on audit row)",
            "docs/Team_E_Implementation_Task_Backlog.docx (Task 6.8, 0.5, 2.6)",
            "backend/core/src/main/java/com/careconnect/service/ChatAuditService.java",
            "backend/core/src/main/java/com/careconnect/service/evv/AuditLogger.java",
            "backend/core/src/main/java/com/careconnect/model/evv/EvvAuditEvent.java",
        ],
    )

    heading(doc, "17. Conclusion", 1)
    para(
        doc,
        "Ask AI audit logging is implemented as AiAuditService writing to an immutable two-table "
        "ledger: ai_ask_audit_record captures the completion snapshot (query hash, delivery status, "
        "tier, citations_json, escalation_json, retrieval/model metadata); ai_ask_audit_event "
        "captures the append-only SCC-3 lifecycle. HITL completion uses ai_ask_audit_delivery_supplement "
        "without mutating the original record. Every API response returns auditId for FR-AI-10 "
        "correlation; REQ-SC-9 is enforced at the database layer with no-update/no-delete rules.",
    )
    para(
        doc,
        "Replace ChatAuditService for /api/ai/ask — do not extend it. Ship P0 audit with the Ask AI "
        "gateway and HITL hold path; citations and escalation fields are mandatory before production.",
        highlight=True,
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    try:
        doc.save(OUTPUT)
        print(f"Created: {OUTPUT}")
    except PermissionError:
        alt = OUTPUT.with_name(f"{OUTPUT.stem}_refresh{OUTPUT.suffix}")
        doc.save(alt)
        print(f"Created (original locked): {alt}")


if __name__ == "__main__":
    build()
