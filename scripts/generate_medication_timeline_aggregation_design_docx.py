"""Generate Word document: Medication timeline aggregation in retrieval responses (FR-AI-11 design)."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX

ROOT = Path(__file__).resolve().parents[1]
OUTPUT = (
    ROOT
    / "docs"
    / "Medication_Timeline_Aggregation_in_Retrieval_Responses_Design.docx"
)


def heading(doc, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def para(doc, text: str, bold: bool = False, highlight: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    if highlight:
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def bullets(doc, items: list[str], highlight_indices: set[int] | None = None) -> None:
    highlight_indices = highlight_indices or set()
    for i, item in enumerate(items):
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        if i in highlight_indices:
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def code(doc, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"


def table(
    doc,
    headers,
    rows,
    highlight_rows: set[int] | None = None,
    highlight_cells: set[tuple[int, int]] | None = None,
) -> None:
    highlight_rows = highlight_rows or set()
    highlight_cells = highlight_cells or set()
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for r_idx, row in enumerate(rows, start=1):
        cells = t.add_row().cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = val
            if r_idx in highlight_rows or (r_idx, c_idx) in highlight_cells:
                for p in cells[c_idx].paragraphs:
                    for run in p.runs:
                        run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def build() -> None:
    doc = Document()

    title = doc.add_heading(
        "Medication Timeline Aggregation in Retrieval Responses — Design Specification",
        0,
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para(doc, "CareConnect Team E — FR-AI-11 Implementation Design")
    para(
        doc,
        "Specifies how hybrid retrieval results are aggregated into a chronological, cited "
        "medication timeline and surfaced in POST /api/ai/ask responses — both as structured "
        "medicationTimeline payload and as LLM context for narrative answers. Builds on "
        "medication timeline requirements research, retrieval orchestration, hybrid index schema, "
        "API contract, and HITL safety design.",
    )
    para(
        doc,
        "Status: DESIGN ONLY — no MedicationTimelineAggregator, MEDICATION_TIMELINE_EVENT index, "
        "or medicationTimeline response field exists. MedicalContextService lists active med names "
        "only; no event history or timeline assembly.",
        highlight=True,
    )
    doc.add_paragraph()

    # 1 Goals
    heading(doc, "1. Design Goals", 1)
    bullets(
        doc,
        [
            "FR-AI-11: Present initiation, change, pause, and termination events — not a single mention.",
            "Every timeline event maps to ≥1 AiCitation; structured field mirrors narrative answer.",
            "Prominently surface discontinuations and dose changes in response lead + event flags.",
            "Merge structured careInstructions events with transcript/FTS supplements under dedup rules.",
            "Entity resolution: brand/generic/colloquial terms resolve to one normalized medication key.",
            "Optional current-state footer from patient_medication when aligned with retrieved events.",
            "Feed SafetyPipeline: dose-change / unsupported-claim validation uses aggregated events.",
            "Single aggregation path for text, voice, and future STML cross-session queries.",
        ],
        highlight_indices={0, 2, 6},
    )

    heading(doc, "1.1 Position in Ask AI pipeline", 2)
    code(
        doc,
        "RetrievalScopeService → RetrievalQueryPlanner (MEDICATION_TIMELINE intent)\n"
        "        → HybridRetrievalService → RankedChunk[]\n"
        "        → MedicationTimelineAggregator.aggregate()   ← THIS DESIGN\n"
        "        → RetrievalContextAssembler (sorted timeline blocks)\n"
        "        → GroundedLlmService (narrative + citationRefs)\n"
        "        → AiAskResponseBuilder (+ medicationTimeline field)\n"
        "        → SafetyPipeline (unsupported med claims vs aggregated events)",
    )
    para(
        doc,
        "Aggregation runs after hybrid search (Stage 3) and before context assembly (Stage 4). "
        "For non-timeline intents, aggregator is skipped; optional lightweight med mention "
        "extraction may still run for escalation hints only.",
    )

    # 2 Components
    heading(doc, "2. Component Architecture", 1)
    table(
        doc,
        ["Component", "Package", "Responsibility"],
        [
            ["MedicationTimelineAggregator", "service.ai.retrieval.timeline", "Orchestrates extract → resolve → dedup → sort → prominence"],
            ["TimelineEventExtractor", "service.ai.retrieval.timeline", "RankedChunk → MedicationTimelineEvent candidates"],
            ["MedicationEntityResolver", "service.ai.retrieval.timeline", "Query + chunk text → normalized medication keys"],
            ["TimelineEventDeduplicator", "service.ai.retrieval.timeline", "Collapse duplicates; prefer structured sources"],
            ["TimelineProminenceRanker", "service.ai.retrieval.timeline", "Flag HIGH prominence discontinuations/changes"],
            ["CurrentMedicationSnapshotService", "service.ai.retrieval.timeline", "patient_medication active list footer"],
            ["MedicationTimelineResponseBuilder", "service.ai.ask", "AggregatedTimeline → API medicationTimeline JSON"],
            ["RetrievalQueryPlanner", "service.ai.retrieval", "Detect MEDICATION_TIMELINE intent (existing design)"],
        ],
        highlight_cells={
            (1, 0), (2, 0), (3, 0), (4, 0), (5, 0), (6, 0),
        },
    )

    heading(doc, "2.1 Relationship to existing code", 2)
    table(
        doc,
        ["Existing", "Design disposition"],
        [
            ["MedicalContextService.addMedicationsContext()", "Not used for Ask AI timeline; optional snapshot footer only"],
            ["MedicationService / patient_medication", "Current-state footer; not primary event source"],
            ["PatientContextRetrievalService", "Replaced by hybrid retrieval + aggregator"],
            ["CallSummary.summary_json", "Upstream: SummaryChunker must emit MEDICATION_TIMELINE_EVENT rows"],
            ["BedrockAIChatService", "Bypassed — timeline via AiAskOrchestrator path"],
        ],
        highlight_rows={1, 2, 4},
    )

    # 3 Domain models
    heading(doc, "3. Core Domain Models", 1)
    code(
        doc,
        "enum MedicationEventType { STARTED, DISCONTINUED, CHANGED, PAUSED, ACTIVE, MENTION }\n"
        "enum DatePrecision { DAY, MONTH, UNKNOWN }\n"
        "enum TimelineProminence { HIGH, NORMAL }\n\n"
        "record MedicationTimelineEvent(\n"
        "    UUID eventId,\n"
        "    String medicationNameDisplay,\n"
        "    String medicationNameNormalized,   // e.g. lisinopril\n"
        "    MedicationEventType eventType,\n"
        "    LocalDate effectiveDate,\n"
        "    DatePrecision effectiveDatePrecision,\n"
        "    String doseFrom,\n"
        "    String doseTo,\n"
        "    String displayText,\n"
        "    boolean needsConfirmation,\n"
        "    double confidence,\n"
        "    TimelineProminence prominence,\n"
        "    List<String> contextRefs,          // C1..Cn from GroundedContext\n"
        "    UUID primaryChunkId,\n"
        "    RetrievalRecordType sourceRecordType,\n"
        "    String sourceEpisodeId,\n"
        "    String sourceItemId,\n"
        "    String sourceTurnId\n"
        ")\n\n"
        "record AggregatedMedicationTimeline(\n"
        "    QueryIntent intent,                // MEDICATION_TIMELINE\n"
        "    MedicationTimelineScope scope,     // SINGLE | MULTI | CLASS\n"
        "    List<String> resolvedMedicationKeys,\n"
        "    List<MedicationTimelineEvent> events,\n"
        "    List<MedicationTimelineEvent> prominentEvents,  // subset for lead paragraph\n"
        "    CurrentMedicationSnapshot currentState,\n"
        "    AggregationMeta meta\n"
        ")",
    )

    heading(doc, "3.1 Scope classification", 2)
    table(
        doc,
        ["Scope", "Detection", "Aggregation behavior"],
        [
            ["SINGLE", "One resolved medication key from query", "Filter events to that key (+ aliases)"],
            ["MULTI", "Explicit list ('aspirin and lisinopril')", "Union per-key timelines; group in response"],
            ["CLASS", "'blood pressure medications', 'heart meds'", "Expand via alias class table + vector hits"],
            ["GENERAL_MED", "Vague 'my pills' without entity", "All medication-type events in date window (cap 20)"],
        ],
        highlight_rows={3},
    )

    # 4 Input sources
    heading(doc, "4. Input Sources & Chunk Mapping", 1)
    table(
        doc,
        ["recordType", "Metadata fields used", "Event extraction"],
        [
            ["MEDICATION_TIMELINE_EVENT", "medication_name_normalized, event_type, effective_date, dose_from/to, source_item_id", "Primary — 1:1 event per chunk"],
            ["SUMMARY_CARE_INSTRUCTION", "type=medication, status, effectiveDate, itemId", "Map status → eventType"],
            ["CALL_SUMMARY / VISIT_SUMMARY", "episodeId, occurredAt", "NLP fallback if no structured events in chunk"],
            ["TRANSCRIPT", "sourceTurnId, occurred_at", "MENTION events only; lower confidence"],
            ["MEDICATION (OLTP snapshot)", "start_date, end_date, is_active", "CurrentState footer — not timeline events"],
        ],
        highlight_cells={(1, 0), (1, 2), (4, 2)},
    )
    para(
        doc,
        "Indexer prerequisite (Task 4.5): explode careInstructions[] where type=medication into "
        "MEDICATION_TIMELINE_EVENT rows at summary index time. See hybrid retrieval scope doc.",
        highlight=True,
    )

    heading(doc, "4.1 careInstructions → eventType mapping", 2)
    table(
        doc,
        ["careInstructions.status", "MedicationEventType", "Notes"],
        [
            ["started", "STARTED", "Initiation or restart"],
            ["discontinued", "DISCONTINUED", "Always prominence=HIGH"],
            ["changed", "CHANGED", "Parse dose_from/dose_to from text"],
            ["paused", "PAUSED", "Temporary hold"],
            ["active", "ACTIVE", "Ongoing without explicit start event"],
            ["(missing)", "MENTION", "Fallback for unstructured SOAP strings"],
        ],
        highlight_rows={2, 6},
    )

    # 5 Aggregation algorithm
    heading(doc, "5. Aggregation Algorithm", 1)
    para(
        doc,
        "MedicationTimelineAggregator.aggregate(AggregationRequest) executes six deterministic "
        "steps. No LLM call in aggregation — keeps audit reproducibility.",
        bold=True,
    )

    heading(doc, "5.1 Step-by-step pipeline", 2)
    code(
        doc,
        "1. RESOLVE  — MedicationEntityResolver.resolve(query, rankedChunks)\n"
        "             → resolvedKeys[], scope (SINGLE|MULTI|CLASS)\n"
        "2. EXTRACT  — TimelineEventExtractor.fromChunks(chunks, resolvedKeys)\n"
        "             → List<MedicationTimelineEvent> candidates\n"
        "3. FILTER   — Keep events matching resolvedKeys (or class expansion)\n"
        "4. DEDUP    — TimelineEventDeduplicator.dedup(candidates)\n"
        "5. SORT     — effectiveDate ASC; UNKNOWN dates last; tie-break confidence DESC\n"
        "6. PROMINENCE — TimelineProminenceRanker.select(events)\n"
        "             → prominentEvents[] (discontinued + most recent CHANGED)\n"
        "7. SNAPSHOT — CurrentMedicationSnapshotService.build(patientId, resolvedKeys)\n"
        "8. META     — counts, warnings (gaps, conflicting dates)",
    )

    heading(doc, "5.2 Deduplication rules", 2)
    table(
        doc,
        ["Rule", "Behavior"],
        [
            ["Same sourceItemId", "Keep single event; prefer MEDICATION_TIMELINE_EVENT over SUMMARY chunk"],
            ["Same effectiveDate + normalized name + eventType", "Merge citations; keep higher confidence"],
            ["Structured vs transcript MENTION", "Drop MENTION if structured event same date ±1 day"],
            ["Conflicting dose on same date", "Keep both with needsConfirmation=true; SafetyPipeline Tier 2"],
            ["Duplicate across call + visit same episode day", "Merge episode citations into one event"],
        ],
        highlight_rows={3, 4},
    )

    heading(doc, "5.3 Prominence selection (FR-AI-11 UX)", 2)
    bullets(
        doc,
        [
            "All DISCONTINUED events → prominence=HIGH; included in prominentEvents.",
            "Most recent CHANGED event per medication key → prominence=HIGH.",
            "STARTED events → NORMAL unless query asks 'when did I start'.",
            "LLM system prompt: lead paragraph must reference prominentEvents first.",
            "UI: render prominentEvents in alert/callout strip above chronological list.",
        ],
        highlight_indices={0, 1, 4},
    )

    heading(doc, "5.4 Date handling", 2)
    bullets(
        doc,
        [
            "effective_date from chunk_metadata is ISO-8601 date; primary sort key.",
            "patient_medication start_date/end_date are VARCHAR — parse with lenient formatter; mark UNKNOWN on failure.",
            "Transcript occurred_at → date-only for MENTION events; precision=DAY.",
            "Events without date: sort last; needsConfirmation=true; excluded from 'when did I start' direct answers unless only hit.",
        ],
    )

    # 6 Entity resolution
    heading(doc, "6. Medication Entity Resolution", 1)
    code(
        doc,
        "record MedicationAliasEntry(\n"
        "    String normalizedKey,      // rxnorm concept or canonical generic\n"
        "    List<String> brandNames,\n"
        "    List<String> colloquialTerms,\n"
        "    List<String> classTags       // antihypertensive, diuretic, ...\n"
        ")\n\n"
        "// Resolution order:\n"
        "// 1. Exact match on medication_name_normalized in retrieved chunks\n"
        "// 2. Alias table lookup (Prinivil → lisinopril)\n"
        "// 3. FTS token overlap on chunk_text\n"
        "// 4. Vector nearest neighbor on medication class embedding (optional P2)\n"
        "// 5. Ambiguous → return candidate list; API may include clarificationPrompt",
    )
    table(
        doc,
        ["Query example", "Resolved key(s)", "Scope"],
        [
            ["When did I start lisinopril?", "lisinopril", "SINGLE"],
            ["Was my water pill stopped?", "furosemide", "SINGLE"],
            ["History of blood pressure meds", "class:antihypertensive", "CLASS"],
            ["Aspirin and lisinopril timeline", "aspirin, lisinopril", "MULTI"],
        ],
    )
    para(
        doc,
        "GAP: No MedicationEntityResolver or alias table in codebase. Task 4.6 covers index-time normalization.",
        highlight=True,
    )

    # 7 Response contract
    heading(doc, "7. API Response — medicationTimeline Field", 1)
    para(
        doc,
        "Extension to POST /api/ai/ask success response (Tier 1 DELIVERED). Present only when "
        "RetrievalQueryPlanner intent=MEDICATION_TIMELINE or when ≥1 medication event aggregated.",
    )
    code(
        doc,
        "{\n"
        '  "success": true,\n'
        '  "deliveryStatus": "DELIVERED",\n'
        '  "tier": 1,\n'
        '  "answer": { "text": "...", "format": "markdown" },\n'
        '  "medicationTimeline": {\n'
        '    "scope": "SINGLE",\n'
        '    "medicationKeys": ["lisinopril"],\n'
        '    "prominentEvents": [ { ...MedicationTimelineEventDto... } ],\n'
        '    "events": [ { ... chronologically sorted ... } ],\n'
        '    "currentState": {\n'
        '      "medications": [\n'
        '        { "name": "Lisinopril", "dosage": "20 mg", "frequency": "daily", "isActive": true }\n'
        '      ],\n'
        '      "source": "patient_medication",\n'
        '      "asOf": "2026-06-20T12:00:00Z"\n'
        '    },\n'
        '    "aggregationMeta": {\n'
        '      "eventsRetrieved": 8,\n'
        '      "eventsAfterDedup": 3,\n'
        '      "structuredCount": 3,\n'
        '      "mentionSupplementCount": 0,\n'
        '      "hasConflicts": false,\n'
        '      "clarificationPrompt": null\n'
        '    }\n'
        "  },\n"
        '  "citations": [ ... AiCitation with recordType MEDICATION_TIMELINE_EVENT ... ]\n'
        "}",
    )

    heading(doc, "7.1 MedicationTimelineEventDto (API)", 2)
    table(
        doc,
        ["Field", "Type", "Required", "Description"],
        [
            ["eventId", "UUID", "Yes", "Stable id for UI list keys"],
            ["medicationName", "string", "Yes", "Display name"],
            ["medicationKey", "string", "Yes", "Normalized key for grouping"],
            ["eventType", "enum", "Yes", "STARTED | DISCONTINUED | CHANGED | PAUSED | ACTIVE | MENTION"],
            ["effectiveDate", "date|null", "No", "ISO date; null if unknown"],
            ["doseFrom / doseTo", "string|null", "No", "For CHANGED events"],
            ["displayText", "string", "Yes", "Human-readable one-liner from records"],
            ["needsConfirmation", "boolean", "Yes", "FR-AI-6 uncertainty flag"],
            ["prominence", "HIGH|NORMAL", "Yes", "UI callout routing"],
            ["citationIds", "UUID[]", "Yes", "Links to top-level citations array"],
        ],
        highlight_rows={9},
    )

    heading(doc, "7.2 Citation linkage", 2)
    bullets(
        doc,
        [
            "Each event.citationIds[] maps to AiCitation.citationId where recordType=MEDICATION_TIMELINE_EVENT.",
            "AiCitation.metadata carries itemId, sourceTurnId, episodeType for deep links (API contract §6).",
            "answer.text sentences must cite same citationIds as corresponding timeline events (SafetyPipeline V2).",
            "Deduplicate citations at response level when multiple events share one source item.",
        ],
        highlight_indices={2},
    )

    heading(doc, "7.3 NO_RECORDS and partial results", 2)
    table(
        doc,
        ["Scenario", "medicationTimeline", "answer"],
        [
            ["Zero chunks after hybrid search", "omit field", "NO_RECORDS path — no timeline"],
            ["Chunks but no structured events", "events=[]; mentionSupplementCount>0", "Narrative from transcript only; escalation confirm-with-provider"],
            ["Entity ambiguous", "clarificationPrompt set", "Ask user to pick medication"],
            ["Events only with needsConfirmation", "events present; all flagged", "Records-based hedged language"],
        ],
        highlight_rows={1, 2},
    )

    # 8 LLM integration
    heading(doc, "8. LLM Context & Narrative Generation", 1)
    heading(doc, "8.1 Context assembly with aggregated timeline", 2)
    code(
        doc,
        "<medication_timeline aggregated=\"true\" scope=\"SINGLE\" medication=\"lisinopril\">\n"
        "  <event ref=\"C1\" type=\"DISCONTINUED\" date=\"2026-06-27\" prominence=\"HIGH\">\n"
        "    Aspirin 81 mg daily discontinued per cardiology guidance.\n"
        "  </event>\n"
        "  <event ref=\"C2\" type=\"STARTED\" date=\"2026-06-28\" prominence=\"NORMAL\">\n"
        "    Lisinopril increased to 20 mg daily.\n"
        "  </event>\n"
        "</medication_timeline>\n"
        "<other_retrieved_records>...</other_retrieved_records>",
    )
    bullets(
        doc,
        [
            "Timeline XML block built from AggregatedMedicationTimeline — not re-parsed from raw chunks.",
            "Context refs C1..Cn assigned before LLM call; event.contextRefs map back for citation assembly.",
            "Token budget: timeline block gets priority slice (up to 40% of retrieved context tokens).",
            "Sort guaranteed by aggregator — assembler does not re-sort.",
        ],
    )

    heading(doc, "8.2 LLM output schema (timeline queries)", 2)
    code(
        doc,
        "{\n"
        '  "answerText": "Your records show lisinopril was started on 2026-06-28 at 20 mg daily. '
        'Aspirin was discontinued on 2026-06-27.",\n'
        '  "citationRefs": ["C1", "C2"],\n'
        '  "timelineSummary": {\n'
        '    "leadSentence": "Aspirin was discontinued on 2026-06-27; lisinopril was started 2026-06-28.",\n'
        '    "prominentEventRefs": ["C1"]\n'
        "  }\n"
        "}",
    )
    para(
        doc,
        "AiAskResponseBuilder merges LLM timelineSummary.leadSentence into answer.text if present; "
        "validates citationRefs cover all prominentEvents.",
    )

    heading(doc, "8.3 Voice / locale", 2)
    bullets(
        doc,
        [
            "Same aggregator path for voice queries after STT normalization.",
            "displayText and answer.text generated in request.locale (en-US / es-US).",
            "medicationTimeline.eventType enums stay English; UI maps to localized labels.",
        ],
    )

    # 9 Safety integration
    heading(doc, "9. Safety & HITL Integration", 1)
    table(
        doc,
        ["Check", "Uses aggregation", "Outcome"],
        [
            ["UnsupportedClaimValidator", "Union of event displayText + cited excerpts", "Tier 2 if draft dose not in events"],
            ["OutputPatternValidator", "prominentEvents CHANGED/DISCONTINUED", "Tier 2 MEDICATION_CHANGE"],
            ["CitationCoverageValidator", "events[].citationIds vs answer refs", "BLOCK if event uncited"],
            ["Summary confirm HITL", "Same aggregator on careInstruction item", "Hold if CHANGED without citation"],
        ],
        highlight_rows={1, 2},
    )
    para(
        doc,
        "Aggregated timeline is passed to SafetyPipeline as structured input — enables deterministic "
        "validation without re-parsing LLM prose.",
        highlight=True,
    )

    # 10 Indexer
    heading(doc, "10. Index-Time Dependencies", 1)
    code(
        doc,
        "// SummaryChunker.explodeMedicationEvents(summaryJson, patientId, episodeId)\n"
        "for (item : careInstructions where type==medication) {\n"
        "  insert retrieval_index_chunk(\n"
        "    record_type = MEDICATION_TIMELINE_EVENT,\n"
        "    chunk_text = formatEventDisplay(item),\n"
        "    chunk_metadata = {\n"
        "      medication_name_normalized: normalize(item.text),\n"
        "      event_type: mapStatus(item.status),\n"
        "      effective_date: item.effectiveDate,\n"
        "      dose_from, dose_to: parseDoseChange(item.text),\n"
        "      source_item_id: item.itemId,\n"
        "      source_turn_id: item.sourceTurnId,\n"
        "      needs_confirmation: item.needsConfirmation,\n"
        "      confidence: item.confidence,\n"
        "      episode_id, episode_type\n"
        "    },\n"
        "    embedding = embed(chunk_text)\n"
        "  );\n"
        "}",
    )
    bullets(
        doc,
        [
            "Blocked until unified summary_json contract live (Ask AI upstream doc).",
            "B-tree index on (patient_id, (chunk_metadata->>'effective_date')).",
            "GIN index on medication_name_normalized + search_vector for hybrid arms.",
        ],
        highlight_indices={0},
    )

    # 11 Config
    heading(doc, "11. Configuration", 1)
    table(
        doc,
        ["Property", "Default", "Purpose"],
        [
            ["careconnect.ai.timeline.enabled", "true", "Master switch with Ask AI"],
            ["careconnect.ai.timeline.max-events", "25", "Cap events in response"],
            ["careconnect.ai.timeline.include-current-state", "true", "patient_medication footer"],
            ["careconnect.ai.timeline.mention-confidence-floor", "0.6", "Drop low-confidence MENTION events"],
            ["careconnect.ai.timeline.class-expansion-limit", "10", "Max keys for CLASS scope"],
            ["careconnect.ai.timeline.date-window-years", "10", "Default SQL prefilter window"],
        ],
    )

    # 12 Tests
    heading(doc, "12. Test Alignment", 1)
    table(
        doc,
        ["Test case", "Validation"],
        [
            ["TC-E-AI-011", "Multi-event timeline; discontinuation prominent; citations per event"],
            ["TC-E-AI-011 alt", "Multi-entity CLASS query returns grouped events"],
            ["TC-E-SC-001", "Dose-change answer without full citation → Tier 2 hold"],
            ["Unit: dedup", "Same itemId from call+visit → single event"],
            ["Unit: sort", "Unknown dates last; ascending effectiveDate"],
            ["Unit: entity", "water pill → furosemide via alias table"],
            ["Contract", "medicationTimeline present iff intent=MEDICATION_TIMELINE"],
        ],
        highlight_rows={1, 2},
    )

    # 13 Gap analysis
    heading(doc, "13. Codebase Gap Analysis", 1)
    table(
        doc,
        ["Area", "Today", "Design target"],
        [
            ["Timeline aggregation", "None", "MedicationTimelineAggregator service"],
            ["Structured med events in index", "None", "MEDICATION_TIMELINE_EVENT rows"],
            ["API medicationTimeline field", "Not in contract implementation", "MedicationTimelineResponseBuilder"],
            ["Entity resolution", "None", "Alias table + resolver"],
            ["Prominence UX", "None", "prominentEvents + UI callout contract"],
            ["Dedup across sources", "None", "TimelineEventDeduplicator rules"],
            ["Safety integration", "GuardrailService phrases only", "SafetyPipeline uses AggregatedTimeline"],
            ["Current med snapshot", "Active names in MedicalContextService", "CurrentMedicationSnapshotService footer"],
        ],
        highlight_rows={1, 2, 3, 4, 5},
    )

    # 14 Implementation phases
    heading(doc, "14. Implementation Phases", 1)
    table(
        doc,
        ["Phase", "Deliverable", "Depends on"],
        [
            ["P0", "Domain models + TimelineEventExtractor + unit tests", "4.5 indexer schema"],
            ["P0", "MedicationTimelineAggregator (dedup, sort, prominence)", "HybridRetrievalService"],
            ["P0", "Context assembler timeline XML block", "Aggregator output"],
            ["P1", "medicationTimeline API field + citation linkage", "API contract DTOs"],
            ["P1", "MedicationEntityResolver + alias table", "Task 4.6"],
            ["P1", "RetrievalQueryPlanner MEDICATION_TIMELINE intent wiring", "Task 5.2"],
            ["P2", "CurrentMedicationSnapshotService footer", "MedicationRepository"],
            ["P2", "CLASS scope + clarificationPrompt flow", "Alias class tags"],
            ["P2", "SafetyPipeline structured timeline input", "Secondary validation design"],
        ],
        highlight_rows={1, 2, 4},
    )

    # 15 Related docs
    heading(doc, "15. Related Documents", 1)
    bullets(
        doc,
        [
            "docs/Medication_Timeline_Retrieval_FR-AI-11.docx (requirements & gap research)",
            "docs/Retrieval_Orchestration_RBAC_Hybrid_Context_LLM_Design.docx (Stage 2–4 pipeline)",
            "docs/Hybrid_Retrieval_Scope_PostgreSQL_FullText_pgvector.docx (MEDICATION_TIMELINE_EVENT index)",
            "docs/POST_api_ai_ask_Request_Response_Contract_Design.docx (citations + response envelope)",
            "docs/Secondary_Validation_and_Tier2_HITL_Hold_Release_Design.docx (medication safety)",
            "docs/Ask_AI_Upstream_Pipeline_Call_Visit_Summaries.docx (careInstructions contract)",
            "docs/Team_E_Implementation_Task_Backlog.docx (Tasks 4.5, 4.6, 5.2)",
            "backend/core/src/main/java/com/careconnect/service/MedicalContextService.java",
            "backend/core/src/main/java/com/careconnect/model/Medication.java",
        ],
    )

    heading(doc, "16. Conclusion", 1)
    para(
        doc,
        "Medication timeline aggregation is a deterministic post-retrieval stage that transforms "
        "RankedChunk[] into AggregatedMedicationTimeline — sorted, deduplicated, prominence-ranked "
        "events with citation linkage. It feeds both the LLM context block and a structured "
        "medicationTimeline field in the Ask AI response, satisfying FR-AI-11 without relying on "
        "the LLM to invent chronology.",
    )
    para(
        doc,
        "Build order: MEDICATION_TIMELINE_EVENT indexer (4.5) → aggregator core (P0) → API field "
        "(P1) → entity resolver (4.6) → SafetyPipeline integration. Do not ship timeline UX without "
        "structured index rows — transcript-only aggregation fails TC-E-AI-011.",
        highlight=True,
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    try:
        doc.save(OUTPUT)
        print(f"Created: {OUTPUT}")
    except PermissionError:
        alt = OUTPUT.with_name(f"{OUTPUT.stem}_refresh{OUTPUT.suffix}")
        doc.save(alt)
        print(f"Created (original locked): {alt}")


if __name__ == "__main__":
    build()
