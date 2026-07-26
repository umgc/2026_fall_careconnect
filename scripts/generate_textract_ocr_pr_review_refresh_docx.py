"""Generate Word document: refreshed PR code review for Task 3.14.4 Textract OCR fallback."""

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.shared import Pt

OUTPUT = (
    Path(__file__).resolve().parents[1]
    / "docs"
    / "PR_Review_Textract_OCR_Fallback_Task_3.14.4_refresh.docx"
)


def heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def para(doc: Document, text: str, bold: bool = False, highlight: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    if highlight:
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def bullets(doc: Document, items: list[str], highlight_indices: set[int] | None = None) -> None:
    highlight_indices = highlight_indices or set()
    for i, item in enumerate(items):
        p = doc.add_paragraph(item, style="List Bullet")
        if i in highlight_indices:
            for run in p.runs:
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def code(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    for i, h in enumerate(headers):
        tbl.rows[0].cells[i].text = h
    for row in rows:
        cells = tbl.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val


def build() -> Document:
    doc = Document()
    title = doc.add_heading("PR Code Review", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    para(
        doc,
        "Task 3.14.4 — Textract OCR fallback — "
        "feature/3.14.4-textract-ocr-fallback → team-ae-develop",
    )
    para(doc, f"Review date: {date.today().isoformat()}")
    para(doc, "Reviewer: AI-assisted code review (Cursor)")
    para(doc, "Document: refreshed review against current branch tip (ce64ef1)")
    doc.add_paragraph()

    heading(doc, "Review metadata", 1)
    table(
        doc,
        ["Field", "Value"],
        [
            ["Source branch", "feature/3.14.4-textract-ocr-fallback"],
            ["Target branch", "team-ae-develop"],
            ["Scope", "7 files changed (+463 / −55 lines)"],
            ["Commits", "cbac2e5 — Implement Textract OCR fallback (Task 3.14.4)"],
            ["Commits", "ce64ef1 — Add E2E test for full digest parse"],
            ["Authors", "YgPadawan"],
        ],
    )

    # ── 1. Change Summary ─────────────────────────────────────────────────────
    heading(doc, "1. Change Summary", 1)
    para(
        doc,
        "This PR implements Task 3.14.4: when USPS Informed Delivery digest HTML provides "
        "insufficient mailpiece metadata (missing sender and/or generic summary text such as "
        "'mail', 'image', or 'campaign'), GmailParser falls back to AWS Textract OCR on inline "
        "data: URL mailpiece images. OCR extracts a sender name and optional summary line from "
        "the top ~15% of the image via MailpieceOcrService.extractMailpieceMetadata().",
    )

    heading(doc, "What changed", 2)
    table(
        doc,
        ["File / area", "Change"],
        [
            [
                "GmailParser.java",
                "Replaces unconditional inferSenderFromImage() with "
                "enrichWithOcrFallbackWhenMetadataInsufficient(); OCR runs only when sender is "
                "blank OR summary is generic; decodes data: URLs and calls MailpieceOcrService. "
                "Adds package-private setMailpieceOcrService() test seam.",
            ],
            [
                "MailpieceOcrService.java",
                "Adds extractMailpieceMetadata() returning sender + optional summaryLine; "
                "refactors collectTopRegionLines(); extractTopLeftLabel() delegates to new method.",
            ],
            [
                "MailpieceOcrResult.java",
                "New record: (sender, summaryLine).",
            ],
            [
                "GmailParserOcrFallbackTest.java",
                "Unit tests: sender fill, skip when metadata present, summary enrichment, "
                "CID campaign path, skip non-data URLs.",
            ],
            [
                "GmailParserOcrFallbackE2ETest.java",
                "E2E: real parser + real OCR service, mocked TextractClient, 3-mailpiece HTML fixture.",
            ],
            [
                "MailpieceOcrServiceTest.java",
                "New tests for extractMailpieceMetadata sender+summary extraction.",
            ],
            [
                "gmail-digest-ocr-fallback.html",
                "Fixture: 3 mailpieces (no metadata, generic alt, complete metadata).",
            ],
        ],
    )

    heading(doc, "Behavioral shift vs. base branch", 2)
    bullets(
        doc,
        [
            "Before: OCR (extractTopLeftLabel) ran whenever HTML sender extraction failed, "
            "regardless of summary quality.",
            "After: OCR runs when isMetadataInsufficientForOcr() is true — blank sender OR "
            "generic/blank summary. This can increase Textract calls when sender is present "
            "but summary is generic (intentional for summary enrichment).",
            "After: OCR can enrich summary (subject) from a second top-region line, not just sender.",
            "After: OCR is limited to data: inline images; remote https:// thumbnails are not processed.",
        ],
    )

    heading(doc, "Verdict", 2)
    para(
        doc,
        "Approve with minor changes. The change is focused, well-tested, and correctly gates "
        "expensive Textract calls behind insufficient-metadata detection. No blocking correctness "
        "bugs. Address logging of OCR-extracted mail content and E2E call-count assertion before "
        "or shortly after merge.",
        bold=True,
    )

    # ── 2. Bug & Risk Analysis ────────────────────────────────────────────────
    heading(doc, "2. Bug & Risk Analysis", 1)

    heading(doc, "2.1 Strengths", 2)
    bullets(
        doc,
        [
            "Graceful degradation: @Autowired(required=false) + "
            "@ConditionalOnProperty(careconnect.aws.enabled=true) means OCR is skipped when "
            "AWS/Textract is unavailable.",
            "Textract failures are caught and logged; parser returns original metadata unchanged.",
            "CID inline images resolved via payload.inlineCidData() before OCR (campaign path tested).",
            "E2E fixture validates the important negative case: complete HTML metadata must not "
            "be overridden by OCR.",
            "Three parse paths updated consistently: structured mailpiece, standalone img, "
            "and campaign table.",
            "No race conditions: parsing is synchronous and stateless per call; "
            "MailpieceOcrService is a stateless Spring bean.",
        ],
    )

    heading(doc, "2.2 Medium — HTTPS / non-inline images skip OCR", 2)
    para(
        doc,
        "runOcrOnImage() only decodes data: URLs. Mailpieces whose thumbnails are hosted at "
        "https:// URLs will never receive OCR even when metadata is insufficient.",
        highlight=True,
    )
    bullets(
        doc,
        [
            "GmailParserOcrFallbackTest.toDomain_ocrSkipped_forNonDataImageUrls() documents "
            "this limitation as intentional for this PR.",
            "Risk: production digests with remote-only images retain null/generic metadata.",
            "Mitigation: bounded HTTP fetch (timeout, max bytes, allowlisted hosts) or "
            "pre-download during Gmail sync — track as follow-up unless Task 3.14.4 is "
            "explicitly inline-only.",
        ],
    )

    heading(doc, "2.3 Medium — OCR gate can increase Textract volume", 2)
    para(
        doc,
        "isMetadataInsufficientForOcr uses OR: blank sender OR generic summary. Previously OCR "
        "ran only when sender was blank. Digests with known senders but alt/summary of 'mail' / "
        "'image' will now call Textract for summary enrichment — higher AWS cost and sync latency.",
        highlight=True,
    )
    code(
        doc,
        """private boolean isMetadataInsufficientForOcr(String sender, String summary) {
    return isBlank(sender) || isSummaryMetadataInsufficient(summary);
}""",
    )
    bullets(
        doc,
        [
            "Each insufficient mailpiece triggers a synchronous detectDocumentText during toDomain().",
            "No caching by image hash; no concurrency limit or batch API.",
            "Acceptable for Task 3.14.4 if product wants summary enrichment; document the cost tradeoff.",
        ],
    )

    heading(doc, "2.4 Medium — INFO logs may expose mail content", 2)
    para(
        doc,
        "MailpieceOcrService logs detected sender and summary at INFO on every successful OCR. "
        "USPS mailpiece text can include medical, financial, or other sensitive identifiers.",
        highlight=True,
    )
    code(
        doc,
        """log.info("Mailpiece OCR detected sender '{}'{}", sender,
        summaryLine != null ? " with summary hint '" + summaryLine + "'" : "");""",
    )
    bullets(
        doc,
        [
            "Prefer log.debug, or log only success/failure + byte length without OCR text.",
            "Aligns with PHI/PII hygiene expectations for CareConnect.",
        ],
    )

    heading(doc, "2.5 Low — OCR result discarded when sender line fails validation", 2)
    para(
        doc,
        "extractMailpieceMetadata() returns Optional.empty() if no line passes looksLikeSender(), "
        "even when a valid summaryLine was detected. When HTML sender is present but summary is "
        "generic, OCR must still find a sender-like line to return any enrichment — otherwise "
        "the Textract call is wasted.",
    )
    code(
        doc,
        """String sender = topRegionLines.stream()
        .filter(this::looksLikeSender)
        .findFirst()
        .orElse(null);
if (sender == null) {
    return Optional.empty();  // summaryLine is lost
}""",
    )

    heading(doc, "2.6 Low — Wrong but present sender is never corrected", 2)
    para(
        doc,
        "If HTML provides a non-blank sender that is incorrect and summary is adequate, OCR will "
        "not run. Pre-existing class of parsing error; not introduced by this PR.",
    )

    heading(doc, "2.7 Low — Summary fallback may duplicate sender", 2)
    para(
        doc,
        "resolveSummaryWithOcrFallback() uses firstNonBlank(ocr.summaryLine(), ocr.sender()). "
        "When Textract finds only one usable line, subject may equal sender — acceptable but "
        "redundant in the UI.",
    )

    heading(doc, "2.8 Low — looksLikeSummary is less strict than looksLikeSender", 2)
    para(
        doc,
        "looksLikeSummary does not filter generic tokens 'mail', 'campaign', or 'image' that "
        "looksLikeSender rejects. An OCR summary line equal to 'mail' could slip through.",
    )

    heading(doc, "2.9 Low — E2E test does not assert Textract call count", 2)
    para(
        doc,
        "GmailParserOcrFallbackE2ETest verifies output correctness but uses lenient() stubbing "
        "and does not verify detectDocumentText was called exactly twice (mailpieces 1 and 2). "
        "A regression that over-invokes Textract for mailpiece 3 would not be caught.",
    )

    heading(doc, "2.10 Low — No max image size guard before Textract", 2)
    para(
        doc,
        "AWS DetectDocumentText sync API rejects oversized payloads (~5 MB). Oversized data: "
        "URLs will fail inside the catch(Exception) path. A pre-check avoids a wasted API call "
        "and clearer metrics.",
    )

    heading(doc, "2.11 Informational — System.out in GmailParser decode path", 2)
    para(
        doc,
        "decodeDataUrlImage() logs decode failures with System.out.println. Matches existing "
        "GmailParser debug patterns but is noisy; prefer SLF4J for the new OCR path at least.",
    )

    # ── 3. Architecture & Style ───────────────────────────────────────────────
    heading(doc, "3. Architecture & Style", 1)

    heading(doc, "3.1 Design patterns", 2)
    bullets(
        doc,
        [
            "Optional service injection cleanly supports AWS-off dev environments.",
            "MailpieceOcrResult record is a clear immutable DTO between OCR service and parser.",
            "extractTopLeftLabel() preserved as a thin delegate — backward compatible.",
            "Private MailpieceFields record groups sender/summary through the enrichment pipeline.",
            "Test seam setMailpieceOcrService() avoids Spring context in unit tests — good pattern.",
            "Conditional OCR is a classic fallback / enrichment strategy with clear predicates.",
        ],
    )

    heading(doc, "3.2 Code organization", 2)
    bullets(
        doc,
        [
            "OCR trigger logic lives in GmailParser; heuristics (looksLikeSender/Summary) live in "
            "MailpieceOcrService — reasonable separation of concerns.",
            "collectTopRegionLines() improves readability vs. the previous inline stream.",
            "Campaign path assigns summary = fields.summary() then later normalizeSummary(), "
            "while toMailPiece() calls normalizeSummary(fields.summary(), sender) inline — "
            "functionally equivalent, slightly inconsistent ordering.",
        ],
    )

    heading(doc, "3.3 Test quality", 2)
    bullets(
        doc,
        [
            "Unit tests cover positive, negative, CID, and non-data URL cases.",
            "E2E uses real GmailParser + real MailpieceOcrService with mocked Textract boundary — "
            "excellent integration coverage.",
            "HTML fixture documents intent per mailpiece with comments.",
            "MailpieceOcrServiceTest extended without breaking existing extractTopLeftLabel coverage.",
        ],
    )

    heading(doc, "3.4 Cleanliness", 2)
    para(
        doc,
        "Overall clean and reviewable. Main style nits: System.out in parser, INFO logging of "
        "OCR text, and the all-or-nothing Optional when sender validation fails.",
    )

    # ── 4. Recommendations ────────────────────────────────────────────────────
    heading(doc, "4. Recommendations", 1)

    heading(doc, "4.1 [Medium] Stop logging OCR-extracted mail text at INFO", 2)
    para(doc, "Avoid writing sender/summary content into application logs:")
    code(
        doc,
        """// Prefer:
log.debug("Mailpiece OCR succeeded (summaryPresent={})", summaryLine != null);

// Or metrics-only:
log.info("Mailpiece OCR succeeded (bytes={}, lines={})",
        imageBytes.length, topRegionLines.size());""",
    )

    heading(doc, "4.2 [Medium] Assert Textract invocation count in E2E test", 2)
    para(doc, "Lock the negative path so complete-metadata mailpieces never call Textract:")
    code(
        doc,
        """import static org.mockito.Mockito.times;
import static org.mockito.Mockito.verify;

// after parser.toDomain(payload):
verify(textractClient, times(2))
        .detectDocumentText(any(DetectDocumentTextRequest.class));""",
    )

    heading(doc, "4.3 [Medium] Align looksLikeSummary with looksLikeSender generic filters", 2)
    code(
        doc,
        """private boolean looksLikeSummary(String text) {
    if (text == null || text.isBlank() || text.length() < 3) {
        return false;
    }
    String lower = text.strip().toLowerCase(Locale.ROOT);
    if (lower.startsWith("learn more") || lower.contains("click") || lower.contains("visit")) {
        return false;
    }
    if (lower.contains("ridealong") || lower.contains("ride along")) {
        return false;
    }
    if (lower.equals("campaign") || lower.equals("mail") || lower.equals("image")) {
        return false;
    }
    return lower.matches(".*[a-z].*");
}""",
    )

    heading(doc, "4.4 [Low] Return partial OCR result when only summary is usable", 2)
    para(
        doc,
        "For sender-present / summary-generic cases, allow summary enrichment even if OCR "
        "sender line fails:",
    )
    code(
        doc,
        """String sender = topRegionLines.stream()
        .filter(this::looksLikeSender)
        .findFirst()
        .orElse(null);

String summaryLine = topRegionLines.stream()
        .filter(this::looksLikeSummary)
        .filter(line -> sender == null || !line.equalsIgnoreCase(sender))
        .findFirst()
        .orElse(null);

if (sender == null && summaryLine == null) {
    return Optional.empty();
}
return Optional.of(new MailpieceOcrResult(sender, summaryLine));""",
    )
    para(
        doc,
        "Update enrichWithOcrFallbackWhenMetadataInsufficient / resolveSummaryWithOcrFallback "
        "to tolerate null ocr.sender() when only summaryLine is present.",
    )

    heading(doc, "4.5 [Low] Replace System.out.println with structured logging", 2)
    code(
        doc,
        """// In GmailParser decodeDataUrlImage catch:
log.debug("Failed to decode inline image for OCR: {}", ex.getMessage());""",
    )

    heading(doc, "4.6 [Low] Guard against oversized inline images before Textract", 2)
    code(
        doc,
        """private static final int MAX_OCR_IMAGE_BYTES = 5 * 1024 * 1024; // 5 MB

private Optional<MailpieceOcrResult> runOcrOnImage(String imageSrc) {
    if (mailpieceOcrService == null || isBlank(imageSrc)) {
        return Optional.empty();
    }
    return decodeDataUrlImage(imageSrc)
            .filter(decoded -> decoded.bytes().length <= MAX_OCR_IMAGE_BYTES)
            .flatMap(decoded -> mailpieceOcrService.extractMailpieceMetadata(
                    decoded.bytes(), decoded.metadata()));
}""",
    )

    heading(doc, "4.7 [Optional] Narrow OCR when only summary is weak", 2)
    para(
        doc,
        "If cost becomes an issue, consider running OCR only when sender is blank, and using a "
        "cheaper/local heuristic for summary — or cache by SHA-256 of image bytes.",
    )
    code(
        doc,
        """// Example: only OCR when sender missing (summary enrichment becomes best-effort)
private boolean isMetadataInsufficientForOcr(String sender, String summary) {
    return isBlank(sender)
            || (isBlank(sender) && isSummaryMetadataInsufficient(summary));
}

// Better cost-aware variant that still enriches summary when sender blank OR
// when summary generic AND an inline image exists — keep current OR, but cache:
// ConcurrentHashMap / Caffeine keyed by sha256(imageBytes) -> MailpieceOcrResult""",
    )

    heading(doc, "4.8 [Future] Remote image OCR support", 2)
    para(
        doc,
        "If production digests rely on https:// thumbnails, add a bounded HTTP fetch before "
        "extractMailpieceMetadata(). Out of scope unless Task 3.14.4 explicitly requires it.",
    )

    # ── File-level review comments ────────────────────────────────────────────
    heading(doc, "File-level review comments", 1)
    table(
        doc,
        ["File", "Severity", "Comment"],
        [
            [
                "GmailParser.java",
                "Info",
                "Good refactor from inferSenderFromImage to gated "
                "enrichWithOcrFallbackWhenMetadataInsufficient across all three parse paths.",
            ],
            [
                "GmailParser.java",
                "Medium",
                "data:-only OCR leaves https thumbnail mailpieces unenriched.",
            ],
            [
                "GmailParser.java",
                "Medium",
                "OR gate (blank sender || generic summary) can increase Textract volume vs base.",
            ],
            [
                "GmailParser.java",
                "Low",
                "System.out.println for decode failures; prefer SLF4J.",
            ],
            [
                "MailpieceOcrService.java",
                "Medium",
                "log.info includes OCR sender/summary text — prefer debug or non-content metrics.",
            ],
            [
                "MailpieceOcrService.java",
                "Low",
                "All-or-nothing Optional when sender line missing discards valid summaryLine.",
            ],
            [
                "MailpieceOcrService.java",
                "Low",
                "looksLikeSummary should reject generic tokens like looksLikeSender.",
            ],
            [
                "MailpieceOcrResult.java",
                "Info",
                "Clean record; document that sender may become nullable if partial results are allowed.",
            ],
            [
                "GmailParserOcrFallbackE2ETest.java",
                "Low",
                "Add verify(textractClient, times(2)) to lock negative path for mailpiece 3.",
            ],
            [
                "GmailParserOcrFallbackTest.java",
                "Info",
                "Solid unit coverage including CID campaign scenario.",
            ],
            [
                "gmail-digest-ocr-fallback.html",
                "Info",
                "Well-documented fixture; mailpiece 3 correctly has complete metadata despite "
                "generic-looking alt text.",
            ],
        ],
    )

    heading(doc, "Test plan (recommended before merge)", 1)
    bullets(
        doc,
        [
            "Run: mvnw test -Dtest=GmailParserOcrFallbackTest,GmailParserOcrFallbackE2ETest,"
            "MailpieceOcrServiceTest",
            "Verify careconnect.aws.enabled=true in staging with real Textract on a sample digest.",
            "Confirm digest parse latency acceptable with 5–10 OCR-eligible mailpieces.",
            "Validate mailpieces with https-only thumbnails behave as expected "
            "(no OCR, metadata unchanged).",
            "Spot-check logs: ensure OCR-extracted sender/summary are not written at INFO in staging.",
        ],
    )

    return doc


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    build().save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    main()
