"""Generate Word document: PR code review for HMAC Gmail OAuth branch."""

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.shared import Pt

OUTPUT = (
    Path(__file__).resolve().parents[1]
    / "docs"
    / "PR_Review_HMAC_OAuth_Gmail_Connect_feature_hmac-signed-oauth-state_refresh.docx"
)


def heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def para(doc: Document, text: str, bold: bool = False, highlight: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    if highlight:
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def bullets(doc: Document, items: list[str], highlight_indices: set[int] | None = None) -> None:
    highlight_indices = highlight_indices or set()
    for i, item in enumerate(items):
        p = doc.add_paragraph(item, style="List Bullet")
        if i in highlight_indices:
            for run in p.runs:
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def code(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    for i, h in enumerate(headers):
        tbl.rows[0].cells[i].text = h
    for row in rows:
        cells = tbl.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val


def build() -> Document:
    doc = Document()
    title = doc.add_heading("PR Code Review", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    para(
        doc,
        "HMAC-signed OAuth state / Gmail connect-disconnect — "
        "feature/hmac-signed-oauth-state-structured-connection-status-gmail-disconnect-flow → team-ae-develop",
    )
    para(doc, f"Review date: {date.today().isoformat()}")
    para(doc, "Reviewer: AI-assisted code review (Cursor)")
    doc.add_paragraph()

    heading(doc, "Review metadata", 1)
    table(
        doc,
        ["Field", "Value"],
        [
            ["Source branch", "feature/hmac-signed-oauth-state-structured-connection-status-gmail-disconnect-flow"],
            ["Target branch", "team-ae-develop"],
            ["Scope", "~22 files, +2201 / −511 (includes merges; feature commits below)"],
            ["Commits (feature)", "fb730fa harden Gmail OAuth (signed state, refresh, reconnect UX)"],
            ["Commits (feature)", "9d31c07 PR review: error redirect, param helper, security matcher"],
            ["Commits (feature)", "94a2469 patient-scoped signed start token"],
            ["Commits (feature)", "299bc39 E2E tests"],
            ["Commits (feature)", "f9d7858 Apply PR review fixes R1–R4 + patient-role coverage"],
            ["Also in history", "CODEOWNERS commits — confirm whether intentional for this PR"],
        ],
    )

    # ── 1. Change Summary ─────────────────────────────────────────────────────
    heading(doc, "1. Change Summary", 1)
    para(
        doc,
        "This PR hardens the Gmail OAuth connect flow used by USPS Informed Delivery. "
        "It replaces unsigned/open OAuth start with a JWT-authenticated connect-url that issues a "
        "short-lived HMAC start token, signs OAuth state for the Google callback, validates return "
        "URL hosts, upserts credentials on exchange, refreshes/revokes tokens, and exposes structured "
        "connection status plus disconnect. Flutter usps_test_screen is updated to use authenticated APIs.",
    )

    heading(doc, "Flow", 2)
    code(
        doc,
        "Authenticated: GET /v1/api/email-credentials/gmail/connect-url\n"
        "  → HMAC startToken (TTL 120s, patient userId + returnUrl)\n"
        "Public browser: GET /oauth/google/start?startToken=...\n"
        "  → verify startToken → sign OAuth state (TTL 600s) → redirect Google\n"
        "Public: GET /oauth/google/callback?code&state\n"
        "  → verify state → exchange code → upsert EmailCredential → redirect frontend\n"
        "Authenticated: GET .../status | DELETE .../gmail",
    )

    heading(doc, "What changed", 2)
    table(
        doc,
        ["Area", "Change"],
        [
            ["OAuthStateSigner", "HMAC-SHA256 signed state + start tokens; constant-time compare; start≠callback"],
            ["OAuthRedirectValidator", "Allowlist return URL hosts; localhost/127.0.0.1; sanitize/resolve"],
            ["EmailOAuthController", "startToken-gated /oauth/google/start; oauthError query on failure"],
            ["EmailCredentialController/Service", "status, disconnect, connect-url; requirePatientAccess"],
            ["EmailConnectionStatus", "CONNECTED / NOT_CONNECTED / NEEDS_RECONNECT (+ unused EXPIRED)"],
            ["GoogleOAuthService", "upsert credential; ensureFreshToken boolean; revokeIfPossible; invalidate"],
            ["USPSDigestService", "refresh token before Gmail fetch"],
            ["SecurityConfig", "explicit /v1/api/email-credentials/** matcher; /oauth/** remains public"],
            ["Tests", "unit + large EmailOAuthFlowE2ETest; patient-role coverage in R1–R4"],
            ["Flutter usps_test_screen", "authenticated connect/status/disconnect"],
        ],
    )

    heading(doc, "Verdict", 2)
    para(
        doc,
        "Approve with minor changes. Security design (authenticated start token → signed state → "
        "host allowlist → patient access checks) is solid and well tested. Before merge: remove or "
        "gate System.out OAuth logging, consider one-time start-token consumption, fix EmailOAuthController "
        "formatting, drop unrelated CODEOWNERS if accidental, and either use or remove STATUS_EXPIRED.",
        bold=True,
    )

    # ── 2. Bug & Risk Analysis ────────────────────────────────────────────────
    heading(doc, "2. Bug & Risk Analysis", 1)

    heading(doc, "2.1 Strengths", 2)
    bullets(
        doc,
        [
            "Start token cannot be used as callback state (prefix separation) — verified by unit test.",
            "HMAC key derived via SHA-256 of email.crypto.secret; signatures compared in constant time.",
            "connect-url requires authentication + requirePatientAccess before minting startToken.",
            "Return URLs sanitized against configured frontend hosts (open-redirect mitigation).",
            "Credential upsert on exchange avoids duplicate Gmail rows on reconnect.",
            "ensureFreshToken + invalidate + revokeIfPossible improve reconnect UX and USPS digest reliability.",
            "Error redirects use oauthError=oauth_failed without leaking exception text to the query string.",
            "Strong E2E coverage of happy path, bad token, patient self-service, and disconnect.",
            "R1–R4 already addressed prior review (patient self-service, upsert, sanitize errors, Flutter auth).",
        ],
    )

    heading(doc, "2.2 Medium — Start tokens / state are reusable until TTL", 2)
    para(
        doc,
        "Nonce is embedded in the payload but never recorded server-side. A stolen startToken "
        "(120s) or state (600s) can be replayed until expiry. For startToken this means anyone with "
        "the URL can initiate Google OAuth for that patient userId; for state, a replayed callback "
        "code is usually single-use at Google, but state binding is weaker than a server-side jti store.",
        highlight=True,
    )
    code(
        doc,
        """// Optional: one-time start token store (Redis/DB)
public ParsedOAuthState verifyStartToken(String startToken) {
    ParsedOAuthState parsed = verifySignedToken(startToken, START_TOKEN_PREFIX);
    String jti = extractNonce(startToken); // or hash of token
    if (!nonceStore.tryConsume(jti, START_TOKEN_TTL_SECONDS)) {
        throw new IllegalArgumentException("Invalid start token: already used");
    }
    return parsed;
}""",
    )

    heading(doc, "2.3 Medium — localhost always allowed as return host", 2)
    para(
        doc,
        "OAuthRedirectValidator.isAllowedHost always accepts localhost and 127.0.0.1. "
        "In production, a crafted returnUrl of http://localhost:... inside a signed token "
        "(if an attacker obtains a startToken for a victim) redirects the victim browser to localhost "
        "after OAuth. Prefer allowing localhost only when spring profile is dev.",
        highlight=True,
    )
    code(
        doc,
        """if ("localhost".equals(host) || "127.0.0.1".equals(host)) {
    return allowLocalhost; // false when spring.profiles.active contains prod
}""",
    )

    heading(doc, "2.4 Medium — System.out.println in GoogleOAuthService", 2)
    para(
        doc,
        "Token exchange logs clientId (partial), redirectUri, and flow steps via System.out. "
        "Noisy in production and can leak configuration. Replace with SLF4J at DEBUG and never log tokens.",
    )

    heading(doc, "2.5 Low — STATUS_EXPIRED defined but unused", 2)
    para(
        doc,
        "EmailConnectionStatus.STATUS_EXPIRED is never returned; expired tokens map to NEEDS_RECONNECT. "
        "Either wire EXPIRED for soft-expiry UX or remove the constant to avoid API confusion.",
    )

    heading(doc, "2.6 Low — demo-user identifier falls back to currentUser", 2)
    para(
        doc,
        "resolvePatientUser treats blank or \"demo-user\" as the authenticated user. Convenient for "
        "legacy USPS demos; document clearly so caregivers do not accidentally operate on themselves "
        "when the client sends demo-user.",
    )

    heading(doc, "2.7 Low — CODEOWNERS included in branch diff", 2)
    para(
        doc,
        "Diff vs team-ae-develop includes .github/CODEOWNERS changes from commits 9926fc9/7f04cb1. "
        "If those are not part of this feature, drop them from the PR to keep review focused.",
    )

    heading(doc, "2.8 Low — EmailOAuthController formatting", 2)
    para(
        doc,
        "File contains excessive blank lines between every statement (likely a formatter glitch in f9d7858). "
        "Reformat to match project Checkstyle / normal Java style before merge.",
    )

    heading(doc, "2.9 Informational — Public /oauth/**", 2)
    para(
        doc,
        "SecurityConfig keeps /oauth/** permitAll, which is required for Google redirect and external "
        "browser start. Security relies on HMAC tokens — correct pattern if tokens stay short-lived and secret is strong.",
    )

    heading(doc, "2.10 No concurrency bugs in signer", 2)
    para(
        doc,
        "OAuthStateSigner is stateless aside from SecureRandom; Mac is created per call. "
        "EmailCredential upsert is per-user; concurrent reconnects may race on save — acceptable; "
        "consider unique (userId, provider) constraint long-term.",
    )

    # ── 3. Architecture & Style ───────────────────────────────────────────────
    heading(doc, "3. Architecture & Style", 1)

    heading(doc, "3.1 Design patterns", 2)
    bullets(
        doc,
        [
            "Separation of Signer / RedirectValidator / CredentialService / OAuthService is clean.",
            "Authenticated mint + public redeem is a standard pattern for mobile/external-browser OAuth.",
            "Structured EmailConnectionStatus record improves reconnect UX over boolean flags.",
            "requirePatientAccess centralizes RBAC instead of ad-hoc role checks (R1–R4 improvement).",
        ],
    )

    heading(doc, "3.2 Test quality", 2)
    bullets(
        doc,
        [
            "OAuthStateSignerTest covers round-trip, tamper, and start≠state.",
            "EmailOAuthFlowE2ETest is thorough (full MockMvc + real signer).",
            "EmailCredentialServiceTest added for patient-role paths.",
            "Consider adding: expired startToken, disallowed returnUrl host, revoke failure still deletes.",
        ],
    )

    heading(doc, "3.3 Style nits", 2)
    bullets(
        doc,
        [
            "Prefer @Slf4j over System.out in GoogleOAuthService.",
            "Normalize EmailOAuthController whitespace.",
            "google.oauth docs updated — good operator hygiene.",
        ],
    )

    # ── 4. Recommendations ────────────────────────────────────────────────────
    heading(doc, "4. Recommendations", 1)

    heading(doc, "4.1 [Medium] Replace System.out with structured logging", 2)
    code(
        doc,
        """@Slf4j
public class GoogleOAuthService {
  public void exchange(String userId, String code) {
    log.debug("Starting Gmail token exchange for userId={}", userId);
    // never log access_token / refresh_token
    ...
  }
}""",
    )

    heading(doc, "4.2 [Medium] Restrict localhost return URLs outside dev", 2)
    code(
        doc,
        """public OAuthRedirectValidator(
        @Value("${google.oauth.frontend-url:http://localhost}") String googleFrontendUrl,
        @Value("${frontend.base-url:http://localhost:3000}") String appFrontendUrl,
        @Value("${google.oauth.allowed-return-hosts:}") String extraHosts,
        @Value("${google.oauth.allow-localhost-return:true}") boolean allowLocalhost) {
  this.allowLocalhost = allowLocalhost;
  ...
}""",
    )

    heading(doc, "4.3 [Low] One-time start token (optional hardening)", 2)
    para(
        doc,
        "Store consumed nonce/jti in a short TTL cache when verifyStartToken succeeds. "
        "Reject reuse. Same pattern can apply to callback state if desired.",
    )

    heading(doc, "4.4 [Low] Use or remove STATUS_EXPIRED", 2)
    code(
        doc,
        """// Option A — soft expiry without forcing reconnect UI:
if (credential.getExpiresAt() != null
        && credential.getExpiresAt().isBefore(Instant.now())
        && hasRefreshToken(credential)) {
  return EmailConnectionStatus.expired(provider, credential.getExpiresAt());
}
// Option B — delete unused STATUS_EXPIRED constant""",
    )

    heading(doc, "4.5 [Low] Reformat EmailOAuthController", 2)
    para(doc, "Run project formatter / remove double blank lines introduced in f9d7858.")

    heading(doc, "4.6 Pre-merge checklist", 2)
    table(
        doc,
        ["Step", "Expected"],
        [
            ["Confirm CODEOWNERS belongs in this PR", "Drop if unrelated"],
            ["Replace System.out in GoogleOAuthService", "SLF4J DEBUG only"],
            ["Reformat EmailOAuthController", "Normal Java spacing"],
            ["Run EmailOAuth* + EmailCredential* + E2E tests", "All green"],
            ["Manual: connect-url → Google → callback → status CONNECTED", "Works"],
            ["Manual: disconnect → status NOT_CONNECTED; digest fails closed", "Works"],
        ],
    )

    heading(doc, "Summary table — findings by severity", 2)
    table(
        doc,
        ["Severity", "Finding", "Action"],
        [
            ["Medium", "Start token / state replay until TTL (no server jti store)", "4.3 optional"],
            ["Medium", "localhost always allowed for returnUrl", "4.2"],
            ["Medium", "System.out OAuth diagnostics", "4.1"],
            ["Low", "STATUS_EXPIRED unused", "4.4"],
            ["Low", "demo-user → currentUser fallback", "Document"],
            ["Low", "CODEOWNERS may be out of scope", "4.6"],
            ["Low", "EmailOAuthController double-spacing", "4.5"],
        ],
    )

    return doc


def main() -> None:
    doc = build()
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    main()
