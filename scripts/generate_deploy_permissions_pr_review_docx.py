"""Generate Word document: PR code review for deploy-permissions-fix branch."""

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.shared import Pt

OUTPUT = (
    Path(__file__).resolve().parents[1]
    / "docs"
    / "PR_Review_Deploy_Permissions_Fix_feature_a-drattray-deploy-permissions-fix.docx"
)


def heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def para(doc: Document, text: str, bold: bool = False, highlight: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    if highlight:
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def bullets(doc: Document, items: list[str], highlight_indices: set[int] | None = None) -> None:
    highlight_indices = highlight_indices or set()
    for i, item in enumerate(items):
        p = doc.add_paragraph(item, style="List Bullet")
        if i in highlight_indices:
            for run in p.runs:
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def code(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    for i, h in enumerate(headers):
        tbl.rows[0].cells[i].text = h
    for row in rows:
        cells = tbl.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val


def build() -> Document:
    doc = Document()
    title = doc.add_heading("PR Code Review", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    para(
        doc,
        "Deploy permissions fix — "
        "feature/a-drattray-deploy-permissions-fix → team-ae-develop",
    )
    para(doc, f"Review date: {date.today().isoformat()}")
    para(doc, "Reviewer: AI-assisted code review (Cursor)")
    doc.add_paragraph()

    heading(doc, "Review metadata", 1)
    table(
        doc,
        ["Field", "Value"],
        [
            ["Source branch", "feature/a-drattray-deploy-permissions-fix"],
            ["Target branch", "team-ae-develop"],
            [
                "3-dot scope (full)",
                "211 files (+32,125 / −6,551) — includes unrelated Team B/CI/EVV history",
            ],
            [
                "Intentional deploy scope",
                "5 cloudformation-fargate files (+783 / −70)",
            ],
            ["Merge-base", "40ce1ca (old — branch not cleanly based on team-ae tip)"],
            ["Key commits", "b9bf68e / 92af88e — Add missing ECS role permissions"],
            ["Key commits", "e0f5709 / b5824eb — linux/amd64 + temp cleanup"],
            ["Key commits", "ddc95df — SSM env-scope + KmsDecryptViaSsm (review follow-up)"],
            ["Authors", "iamniquey (+ merge noise from develop)"],
        ],
    )

    # ── 1. Change Summary ─────────────────────────────────────────────────────
    heading(doc, "1. Change Summary", 1)
    para(
        doc,
        "Intent: unblock ECS Fargate features that fail when careconnect-{env}-ecsTaskRole "
        "has no runtime policy. The PR attaches an inline task-role policy for Chime "
        "meetings/media pipelines, Kinesis Video, S3 recordings/uploads, Bedrock, Transcribe, "
        "Textract, SES/SNS, SSM (env-scoped), KMS decrypt via SSM, Chime media-pipeline SLR "
        "create, and sts:GetCallerIdentity. Deploy scripts force docker build "
        "--platform linux/amd64 and clean stale careconnect-*-data-*.json temp files. "
        "README / DEPLOY_2026_SUMMER.md document the role surface and Amplify CORS redeploy pitfalls.",
    )
    para(
        doc,
        "IMPORTANT: vs team-ae-develop the 3-dot diff is ~211 files. Reviewers should treat "
        "only cloudformation-fargate IAM/scripts/docs as the product of this PR and require "
        "a rebase/cherry-pick onto current team-ae-develop before merge.",
        bold=True,
        highlight=True,
    )

    heading(doc, "What changed (intentional)", 2)
    table(
        doc,
        ["File / area", "Change"],
        [
            [
                "templates/03-platform.yaml",
                "EcsTaskRole: empty → careconnect-${Environment}-ecs-task-app. "
                "ddc95df narrows SSM to /careconnect/${Environment}/* and adds KmsDecryptViaSsm.",
            ],
            [
                "cdeploy_cloudformation.ps1 / .sh",
                "docker build --platform linux/amd64; startup cleanup of stale "
                "careconnect-*-data-*.json under temp dir.",
            ],
            [
                "README.md / DEPLOY_2026_SUMMER.md",
                "ECS task-role permissions section; redeploy matrix; Amplify CORS §7/§11.",
            ],
        ],
    )

    heading(doc, "Verdict", 2)
    para(
        doc,
        "Request changes (process) + Approve with nits (IAM content). Do not merge the "
        "211-file 3-dot delta as-is. After a clean cherry-pick of the five cloudformation "
        "files onto team-ae-develop, the IAM direction is correct; re-pin Bedrock to the "
        "develop model allow-list and treat CreateBucket/SES * as follow-ups.",
        bold=True,
        highlight=True,
    )

    # ── 2. Bug & Risk Analysis ────────────────────────────────────────────────
    heading(doc, "2. Bug & Risk Analysis", 1)

    heading(doc, "2.1 High — PR surface vs merge-base", 2)
    para(
        doc,
        "merge-base with team-ae-develop is 40ce1ca. HEAD includes Team B CI, Sonar, EVV "
        "tests, frontend, and other develop history that team-ae-develop already diverged from. "
        "Merging as-is risks duplicate/conflicting commits and review blindness. Race is "
        "process/git: not a runtime race in IAM YAML itself.",
        highlight=True,
    )
    bullets(
        doc,
        [
            "Cherry-pick b9bf68e, e0f5709, ddc95df (+ doc commits) onto fresh branch from "
            "team-ae-develop, or manually apply the five-file intentional delta.",
            "Open the PR against that cleaned branch only.",
        ],
    )

    heading(doc, "2.2 High — Bedrock foundation-model/* regresses team-ae-develop", 2)
    para(
        doc,
        "team-ae-develop already pins Bedrock to named foundation models (Nova, Titan embed, "
        "Claude variants) plus inference-profile/*. This branch replaces that with "
        "foundation-model/*, allowing any FM in the region if the role is compromised or "
        "misconfigured. That is a least-privilege regression relative to the Ask AI P0 path.",
        highlight=True,
    )

    heading(doc, "2.3 Medium — S3 CreateBucket + PutBucketPolicy on task role", 2)
    para(
        doc,
        "Prefix-scoped buckets (careconnect-recordings-*, careconnect-uploads-*) are good, "
        "but CreateBucket + PutBucketPolicy let a compromised task create matching buckets "
        "and attach attacker-controlled policies. Prefer CFN-provisioned buckets and object "
        "CRUD only after first-run provisioning is no longer required.",
        highlight=True,
    )

    heading(doc, "2.4 Medium — SES/SNS Resource '*'", 2)
    para(
        doc,
        "Notifications can SendEmail / Publish anywhere in the account. Scope to verified "
        "identities and careconnect-${Environment}-* topics when ARNs are known.",
    )

    heading(doc, "2.5 Medium — KMS decrypt key/* via SSM", 2)
    para(
        doc,
        "KmsDecryptViaSsm correctly uses kms:ViaService=ssm… and fixes SecureString reads "
        "(good follow-up in ddc95df). Resource key/* still allows any account CMK when "
        "called through SSM. Prefer the alias/aws/ssm key ARN or a dedicated CareConnect CMK.",
    )

    heading(doc, "2.6 Low — SSM env narrow vs hardcoded SsmConfig /careconnect/prod/", 2)
    para(
        doc,
        "IAM now matches SsmPropertySourceInitializer (/careconnect/{ENVIRONMENT}/). "
        "SsmConfig still hardcodes /careconnect/prod/. On cfdemo/staging those GetParameter "
        "calls get AccessDenied and fall back to env vars — safer cross-env isolation, but "
        "a dead code path unless Environment=prod.",
    )

    heading(doc, "2.7 Low — Chime / Kinesis edge cases", 2)
    bullets(
        doc,
        [
            "Chime Resource '*' is expected (API largely non-resource-scoped).",
            "Media Insights *Configuration* APIs are account-level; prefer ops-time once.",
            "kinesisvideo:ListStreams may need Resource '*' separately from stream/*.",
            "CreateServiceLinkedRole is conditioned correctly; prefer one-time admin create.",
        ],
    )

    heading(doc, "2.8 Strengths", 2)
    bullets(
        doc,
        [
            "Correctly targets EcsTaskRole (runtime), not execution role (image/secrets).",
            "SSM narrowed to ${Environment} (ddc95df) — improvement over early commit.",
            "KmsDecryptViaSsm addresses SecureString decrypt gap.",
            "docker --platform linux/amd64 prevents Apple Silicon → Fargate exec-format failures.",
            "Stale temp param cleanup avoids interrupted-deploy footguns.",
            "Docs document Amplify CORS wipe after app-only redeploy — high operator value.",
        ],
    )

    # ── 3. Architecture & Style ───────────────────────────────────────────────
    heading(doc, "3. Architecture & Style", 1)
    bullets(
        doc,
        [
            "Pattern: least-privilege ECS task role as CloudFormation-owned IAM — correct.",
            "Sid-named statements, README pointer, env-aware RoleName — clean and reviewable.",
            "Inline policy on the role is fine for demo scale; managed policy would help if "
            "the statement list grows further.",
            "Deploy docs (redeploy matrix, platform-only path) match how students actually "
            "iterate — strong operational design.",
            "Branch hygiene is the main architecture smell: feature branch should not carry "
            "unrelated develop history into team-ae-develop.",
        ],
    )

    # ── 4. Recommendations ────────────────────────────────────────────────────
    heading(doc, "4. Recommendations", 1)

    heading(doc, "4.1 Rebase / cherry-pick before merge", 2)
    code(
        doc,
        """git fetch origin
git checkout -b feature/a-drattray-deploy-permissions-fix-clean origin/team-ae-develop
# cherry-pick IAM + scripts + docs commits, resolve only cloudformation-fargate conflicts
git cherry-pick b9bf68e e0f5709 ddc95df
# or apply the five-file intentional delta manually""",
    )

    heading(doc, "4.2 Re-pin Bedrock to develop allow-list", 2)
    code(
        doc,
        """- Sid: BedrockInvoke
  Effect: Allow
  Action:
    - bedrock:InvokeModel
    - bedrock:InvokeModelWithResponseStream
  Resource:
    - !Sub arn:aws:bedrock:${AWS::Region}::foundation-model/amazon.nova-pro-v1:0
    - !Sub arn:aws:bedrock:${AWS::Region}::foundation-model/amazon.nova-lite-v1:0
    - !Sub arn:aws:bedrock:${AWS::Region}::foundation-model/amazon.titan-embed-text-v1
    - !Sub arn:aws:bedrock:${AWS::Region}::foundation-model/amazon.titan-embed-text-v2:0
    - !Sub arn:aws:bedrock:${AWS::Region}::foundation-model/anthropic.claude-3-haiku-20240307-v1:0
    - !Sub arn:aws:bedrock:${AWS::Region}::foundation-model/anthropic.claude-sonnet-4-5-20250929-v1:0
    # keep remaining models from team-ae-develop BedrockInvokeModel list
    - !Sub arn:aws:bedrock:${AWS::Region}:${AWS::AccountId}:inference-profile/*""",
    )

    heading(doc, "4.3 Scope SES / SNS", 2)
    code(
        doc,
        """- Sid: Notifications
  Effect: Allow
  Action:
    - ses:SendEmail
    - ses:SendRawEmail
  Resource:
    - !Sub arn:aws:ses:${AWS::Region}:${AWS::AccountId}:identity/*
- Sid: SnsPublish
  Effect: Allow
  Action: [sns:Publish]
  Resource: !Sub arn:aws:sns:${AWS::Region}:${AWS::AccountId}:careconnect-${Environment}-*""",
    )

    heading(doc, "4.4 Drop bucket-admin after buckets exist", 2)
    code(
        doc,
        """Action:
  # remove once CFN owns recordings/uploads buckets:
  # - s3:CreateBucket
  # - s3:PutBucketPolicy
  - s3:ListBucket
  - s3:GetObject
  - s3:PutObject
  - s3:DeleteObject""",
    )

    heading(doc, "4.5 Align SsmConfig with ENVIRONMENT", 2)
    code(
        doc,
        """// SsmConfig.java — match SsmPropertySourceInitializer
private String prefix() {
  String env = System.getenv().getOrDefault("ENVIRONMENT", "prod");
  return "/careconnect/" + env.trim() + "/";
}""",
    )

    heading(doc, "Merge checklist", 2)
    bullets(
        doc,
        [
            "Clean branch: only cloudformation-fargate intentional files vs team-ae-develop.",
            "Platform stack deploy + force-new-deployment on ECS service.",
            "Smoke: Chime meeting create, SecureString SSM read, Bedrock invoke, docker push "
            "from arm64 host if used.",
            "Confirm Amplify CORS still set after any app-only redeploy (§7/§11).",
        ],
    )

    return doc


def main() -> None:
    doc = build()
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    main()
