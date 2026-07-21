"""
Build a Team E → shared TDD revision packet document.

Source (content): CareConnect_Milestone_2_TDD_TEAM E (1).docx
Target (edit destination): 01_Technical_Design_Document.docx
"""

from __future__ import annotations

import re
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor

TEAM_E_TDD = Path(r"C:\Users\ravic\Downloads\CareConnect_Milestone_2_TDD_TEAM E (1).docx")
TARGET_TDD = Path(r"C:\Users\ravic\Downloads\01_Technical_Design_Document.docx")
OUTPUT = (
    Path(__file__).resolve().parents[1]
    / "docs"
    / "TDD_Revision_Packets_TeamE_into_01_Technical_Design_Document.docx"
)

# (packet_id, action, team_e_keys, dest, notes)
PACKETS: list[dict] = [
    {
        "id": "RH",
        "title": "Revision History",
        "action": "EDIT",
        "dest": "Revision History (front matter)",
        "team_e": ["Revision History"],
        "guidance": "Add a Team E contribution row (Ask AI / summaries / USPS / STML / safety) "
        "with version, date, and author. Keep existing rows intact.",
    },
    {
        "id": "1",
        "title": "1. Introduction",
        "action": "EDIT",
        "dest": "§1 Introduction (umbrella)",
        "team_e": ["1. Introduction", "1.1 Purpose", "1.3 Document Scope", "1.4 Definitions, Acronyms, and Abbreviations", "1.5 References"],
        "guidance": "Merge Team E purpose/scope language into shared intro only where Team E modules "
        "are missing. Prefer additive bullets over rewriting Team A narrative.",
    },
    {
        "id": "1.2",
        "title": "1.2 Intended Audience → target 1.6",
        "action": "EDIT",
        "dest": "§1.6 Intended Audience",
        "team_e": ["1.2 Intended Audience"],
        "guidance": "Ensure audience list includes patient, caregiver, family, clinician/care-team, "
        "and engineering/ops readers for Ask AI and summary workflows.",
    },
    {
        "id": "2",
        "title": "2. System Overview",
        "action": "EDIT",
        "dest": "§2 System Architecture Overview (umbrella)",
        "team_e": ["2. System Overview", "2.1 System Description", "2.2 System Objectives", "2.4 Operating Environment"],
        "guidance": "Add Team E objectives (grounded retrieval, cited summaries, mail ingest, STML, safety) "
        "into §2.1 goals / overview without duplicating §6 feature write-ups.",
    },
    {
        "id": "2.3",
        "title": "2.3 System Users → target 2.6",
        "action": "EDIT",
        "dest": "§2.6 System Users",
        "team_e": ["2.3 System Users"],
        "guidance": "Merge Team E user/needs/design-impact table into §2.6. Emphasize consent-scoped "
        "caregiver access and STML-first patient UX.",
    },
    {
        "id": "3",
        "title": "3. Architecture Diagram",
        "action": "EDIT / ADD figures",
        "dest": "§2.2 Architecture Diagram; §2.3 Component Responsibilities; §2.5 Workflows; "
        "deployment view → §7.1 AWS Infrastructure Overview",
        "team_e": [
            "3. Architecture Diagram",
            "3.1 High-Level Architecture",
            "3.2 Component Responsibilities",
            "3.3 Data Flow or Workflow View",
            "3.3.1 Use-Case Diagrams Hosted Here",
            "3.4 Module Architecture",
            "3.5 Deployment View",
        ],
        "guidance": "Place AI Gateway / Summaries / Retrieval / IndexWorker / USPS / STML components "
        "under §2.3.6 AI and Communication Services. Put ECS/Bedrock/pgvector deployment notes in §7.1. "
        "Copy Team E architecture figures as List of Figures entries.",
    },
    {
        "id": "4",
        "title": "4. Technology Stack",
        "action": "EDIT",
        "dest": "§3 Technology Stack — decisions/risks → §3.1–§3.2",
        "team_e": ["4. Technology Stack", "4.1 Selected Technologies", "4.2 Technology Decisions and Assumptions", "4.3 Technology Risks"],
        "guidance": "Fold Bedrock, pgvector, FTS/hybrid retrieval, outbox/IndexWorker, Flutter Ask AI "
        "decisions into §3.1 and risks (throttle, lag, cost, model availability) into §3.2.",
    },
    {
        "id": "5.1",
        "title": "5.1 Call/visit summaries",
        "action": "EDIT",
        "dest": "§6.2 Video Call and Transcription Pipeline Design; §6.3 AI and Bedrock Integration Design",
        "team_e": ["5.1 Call or Visit Summaries", "5.1.1 Implemented Use Cases"],
        "guidance": "Add Team E summary confirm / structured items / indexing outbox hooks adjacent to "
        "existing call-summary / SOAP content in §6.2–§6.3.",
    },
    {
        "id": "5.2",
        "title": "5.2 Ask AI",
        "action": "ADD",
        "dest": "§6.12 AI-Assisted Retrieval (Ask AI)  [NEW — after §6.11]",
        "team_e": ["5.2 AI-Assisted Retrieval (Ask AI)", "5.2.1 Implemented Use Cases"],
        "guidance": "Create new §6.12 covering hybrid retrieval (FTS + pgvector), RBAC scope, citations, "
        "HITL/safety gates, voice path notes, and dependency on §5.11 API.",
    },
    {
        "id": "5.3",
        "title": "5.3 USPS Mail Agent",
        "action": "ADD",
        "dest": "§6.13 USPS Informed Delivery Mail Agent  [NEW]",
        "team_e": ["5.3 USPS Informed Delivery Mail Agent", "5.3.1 Implemented Use Cases"],
        "guidance": "New feature subsection: Informed Delivery sync, OCR/Textract fallback, mailpiece "
        "indexing, and RBAC visibility. Cross-link §4.9 MailPiece entity.",
    },
    {
        "id": "5.4",
        "title": "5.4 STML",
        "action": "ADD",
        "dest": "§6.14 Short-Term Memory (STML) Support  [NEW]",
        "team_e": ["5.4 Short-Term Memory (STML) Support", "5.4.1 Implemented Use Cases"],
        "guidance": "Describe STML UX, reminder flows, and backend hooks as implemented by Team E.",
    },
    {
        "id": "5.5",
        "title": "5.5 Safety, Consent & Clarity",
        "action": "ADD",
        "dest": "§6.15 Safety, Consent & Clarity  [NEW]",
        "team_e": ["5.5 Safety, Consent & Clarity"],
        "guidance": "Hold/release, secondary validation, consent visibility, clarity/disclaimer patterns "
        "for patient-facing Ask AI and summaries.",
    },
    {
        "id": "5.6",
        "title": "5.6 Cross-team deps",
        "action": "EDIT",
        "dest": "§6.5 Cross-Team Functional Dependencies",
        "team_e": ["5.6 Cross-Team Functional Dependencies"],
        "guidance": "Extend dependency register with Team E ↔ Team A items (Bedrock IAM, Cognito/JWT, "
        "Chime transcripts, Amplify CORS, ECS task role).",
    },
    {
        "id": "6.2.1",
        "title": "6.2.1 Ask AI endpoint",
        "action": "ADD",
        "dest": "§5.11 Ask AI Endpoint (POST /api/ai/ask)  [NEW — after §5.10]",
        "team_e": ["6. API Design Specification", "6.1 API Conventions and Standards", "6.2 Endpoint Summary", "6.2.1 POST /api/ai/ask", "6.3 Request and Response Examples", "6.4 Error Handling and Status Codes"],
        "guidance": "New API subsection: auth, RBAC scope, request/response contract, citation shape, "
        "error codes. Keep conventions aligned with §5.2.",
    },
    {
        "id": "6.2.2",
        "title": "6.2.2 Summary confirm",
        "action": "EDIT",
        "dest": "§5.5 Transcription and Summary Endpoints (Planned)",
        "team_e": ["6.2.2 POST /api/summaries/{id}/items/{itemId}/confirm"],
        "guidance": "Promote confirm endpoint from Planned to specified; include path params and "
        "success/error status codes from Team E §6.2.2 / §6.4.",
    },
    {
        "id": "7.3.2",
        "title": "7.3.2 MailPiece",
        "action": "ADD",
        "dest": "§4.9 Mail / USPS Entities (MailPiece)  [NEW — after §4.8]",
        "team_e": ["7. Data Model Specification", "7.1 Data Modeling Approach and Assumptions", "7.2 Data Models by Assigned Module", "7.3 Data Dictionary Alignment", "7.3.1 Summary Table", "7.3.2 MailPiece Table"],
        "guidance": "Add MailPiece data dictionary and note retrieval_index_chunk / indexing_outbox "
        "touchpoints if present in Team E ERD/tables.",
    },
    {
        "id": "8",
        "title": "8. UI / UX Design Specification",
        "action": "EDIT / ADD",
        "dest": "§13.1–§13.4 UX/UI Design Specification; accessibility → §11.7",
        "team_e": [
            "8. UI / UX Design Specification",
            "8.1 UI Design Approach and Principles",
            "8.2 Screens and User Workflows Affected",
            "8.3 Wireframes, Mockups, or Screen Notes",
            "8.4 Accessibility Requirements",
            "8.5 Display Text and Localization Impact",
        ],
        "guidance": "Add Ask AI / Summaries / Mail / STML screens under §13.2–§13.3. If §13.4 Display "
        "Language is missing, add it. Route accessibility requirements primarily to §11.7.",
    },
    {
        "id": "9",
        "title": "9. Security Considerations",
        "action": "EDIT",
        "dest": "§8 Security Design; logging/monitoring → §8.8; security testing → §11.6",
        "team_e": [
            "9. Security Considerations",
            "9.1 Data Privacy",
            "9.2 Authentication and Authorization",
            "9.3 Secure Data Transmission",
            "9.4 API Security",
            "9.5 Data Storage and Backup",
            "9.6 Logging and Monitoring",
            "9.7 External Service or AI Tool Security",
            "9.8 Security Testing",
        ],
        "guidance": "Map privacy/RBAC/API/storage/AI-tool notes into §8.*; put AuditLogging / Ask AI "
        "audit events in §8.8; put AI-tool security testing in §11.6.",
    },
    {
        "id": "10",
        "title": "10. Error Handling",
        "action": "EDIT",
        "dest": "§10 Error Handling",
        "team_e": ["10. Error Handling"],
        "guidance": "Add Ask AI / indexing / Bedrock degradation and outbox dead-letter behaviors.",
    },
    {
        "id": "11",
        "title": "11. Performance & Scalability",
        "action": "EDIT",
        "dest": "§9 Performance and Scalability; bottleneck mitigation → §9.4",
        "team_e": ["11. Performance & Scalability", "11.1 Objectives", "11.2 Scalability Strategy", "11.3 Performance Monitoring", "11.4 Bottleneck Mitigation"],
        "guidance": "Incorporate embedding latency, Bedrock throttle, hybrid retrieval cost, and "
        "outbox lease parking into §9.4.",
    },
    {
        "id": "12",
        "title": "12. Testing Strategy",
        "action": "ADD",
        "dest": "§11.20 Team E Feature-Specific Test Cases; §11.21 Team E Test Automation  [NEW]",
        "team_e": ["12. Testing Strategy", "12.1 Testing Overview", "12.2 Testing Levels", "12.3 Feature-Specific Test Cases", "12.4 Test Automation"],
        "guidance": "Do not overwrite Team B coverage sections. Append Team E cases for Ask AI, "
        "summaries confirm, USPS, STML, and safety/HITL.",
    },
    {
        "id": "13",
        "title": "13. Deployment Strategy",
        "action": "EDIT",
        "dest": "§7.8 Deployment Environments; §7.9 Rollback Procedure",
        "team_e": ["13. Deployment Strategy", "13.1 Deployment Environments", "13.2 Infrastructure or Hosting", "13.3 CI/CD Pipeline", "13.4 Rollback Procedure"],
        "guidance": "Add Team E notes for platform IAM (Bedrock/SSM), app-only vs platform redeploy, "
        "and Amplify CORS re-apply after app-only deploys into §7.8–§7.9.",
    },
    {
        "id": "14",
        "title": "14. Known Issues & Limitations",
        "action": "ADD",
        "dest": "§12.9 Team E Known Issues & Limitations  [NEW — after §12.8] with §12.9.1–§12.9.5",
        "team_e": [
            "14. Known Issues & Limitations",
            "14.1 Technical Limitations",
            "14.2 Known Issues",
            "14.3 Exclusions",
            "14.4 Future Enhancements",
            "14.5 Cross-Team Resolution Notes",
        ],
        "guidance": "Mirror Team E §14.1–14.5 as §12.9.1–12.9.5 (technical limitations, known issues, "
        "exclusions, future enhancements, cross-team resolution).",
    },
]


def set_cell_shading(cell, fill_hex: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def para(doc: Document, text: str, *, bold: bool = False, italic: bool = False, highlight: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if highlight:
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def codeish(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def normalize_heading(text: str) -> str:
    t = re.sub(r"\s+", " ", (text or "").strip())
    t = t.replace("—", "-").replace("–", "-")
    # drop trailing odd chars
    t = re.sub(r"[^\w\s./{}()&+\-:#]+$", "", t)
    return t


def match_key(heading_text: str, keys: list[str]) -> str | None:
    h = normalize_heading(heading_text).lower()
    for key in keys:
        k = normalize_heading(key).lower()
        if h == k or h.startswith(k) or k.startswith(h[: min(40, len(h))]):
            return key
        # fuzzy: key without section words
        if k.split(" ", 1)[-1] in h and len(k) > 8:
            return key
    return None


def extract_team_e_sections(path: Path) -> dict[str, dict]:
    """Return {canonical_or_raw_heading: {paras, tables_as_text}}."""
    d = Document(str(path))
    sections: dict[str, dict] = {}
    order: list[str] = []
    current: str | None = None

    def ensure(name: str) -> dict:
        if name not in sections:
            sections[name] = {"paras": [], "tables": []}
            order.append(name)
        return sections[name]

    # Walk body elements in order (paragraphs and tables)
    body = d.element.body
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    for child in body.iterchildren():
        if isinstance(child, CT_P):
            p = Paragraph(child, d)
            style = p.style.name if p.style else ""
            text = (p.text or "").strip()
            if not text:
                continue
            is_heading = style.startswith("Heading") or style.startswith("Title")
            numbered = bool(re.match(r"^(Revision History|\d+(\.\d+)*)\b", text))
            if is_heading or (numbered and len(text) < 140 and style.startswith("Heading")):
                current = text
                ensure(current)
                continue
            if re.match(r"^(Revision History|\d+(\.\d+)*\s+\S)", text) and len(text) < 140 and style == "Normal":
                # e.g. "11. Performance & Scalability" stored as Normal
                current = text
                ensure(current)
                continue
            if current is None:
                current = "_preamble"
                ensure(current)
            sections[current]["paras"].append(text)
        elif isinstance(child, CT_Tbl):
            table = Table(child, d)
            if current is None:
                current = "_preamble"
                ensure(current)
            rows_txt = []
            for row in table.rows:
                cells = [(c.text or "").strip().replace("\n", " | ") for c in row.cells]
                # de-dup repeated merge cells roughly
                dedup = []
                for c in cells:
                    if not dedup or dedup[-1] != c:
                        dedup.append(c)
                rows_txt.append(" || ".join(dedup))
            sections[current]["tables"].append(rows_txt)

    sections["__order__"] = order  # type: ignore[assignment]
    return sections


def find_section_blob(sections: dict, keys: list[str]) -> list[tuple[str, dict]]:
    """Return matching (heading, content) pairs for packet keys."""
    order = sections.get("__order__", [])
    if not isinstance(order, list):
        order = []
    found: list[tuple[str, dict]] = []
    used: set[str] = set()
    for key in keys:
        k_norm = normalize_heading(key).lower()
        best = None
        for h in order:
            if h in used or h not in sections or not isinstance(sections[h], dict):
                continue
            hn = normalize_heading(h).lower()
            if hn == k_norm or hn.startswith(k_norm):
                best = h
                break
            # POST /api/ai/ask style
            key_tail = k_norm
            for prefix in ("6.2.1 ", "6.2.2 ", "7.3.1 ", "7.3.2 "):
                if k_norm.startswith(prefix.strip()) or prefix.strip() in k_norm:
                    pass
            if "post /api/ai/ask" in k_norm and "ai/ask" in hn:
                best = h
                break
            if "mailpiece" in k_norm and "mailpiece" in hn:
                best = h
                break
            if "confirm" in k_norm and "confirm" in hn and "summar" in hn:
                best = h
                break
            if "summary table" in k_norm and "summary table" in hn:
                best = h
                break
        if best is None:
            for h in order:
                if h in used or h not in sections or not isinstance(sections[h], dict):
                    continue
                hn = normalize_heading(h).lower()
                # number prefix match e.g. 5.2
                m = re.match(r"^(\d+(?:\.\d+)*)\b", k_norm)
                if m and hn.startswith(m.group(1) + " "):
                    best = h
                    break
        if best:
            used.add(best)
            found.append((best, sections[best]))
    return found


def action_color(action: str) -> str:
    a = action.upper()
    if a.startswith("ADD"):
        return "C6EFCE"
    if a.startswith("EDIT"):
        return "FFF2CC"
    return "DDEBF7"


def build() -> Document:
    sections = extract_team_e_sections(TEAM_E_TDD)

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)

    title = doc.add_heading(
        "TDD Revision Packets — Team E → 01_Technical_Design_Document", 0
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    para(doc, f"Prepared: {date.today().isoformat()}")
    para(doc, f"Content source: {TEAM_E_TDD.name}")
    para(doc, f"Edit target: {TARGET_TDD.name}")
    para(
        doc,
        "Use this document as the cutsheet while updating the shared Team A Technical Design "
        "Document. Each packet lists where to paste/adapt Team E material, the edit action "
        "(EDIT existing subsection vs ADD new subsection), editorial guidance, and extracted "
        "source text/tables.",
        italic=True,
    )

    heading(doc, "How to use", 1)
    bullets(
        doc,
        [
            "Yellow highlight / EDIT = update an existing target heading.",
            "Green / ADD = insert a new subsection number that does not yet exist in 01_TDD "
            "(confirmed missing: §4.9, §5.11, §6.12–§6.15, §11.20–§11.21, §12.9).",
            "Copy figures manually from the Team E .docx (paths called out in architecture packets).",
            "Prefer additive merges: do not delete other teams’ content in shared sections (§2, §3, §6.5, §8, §10).",
        ],
    )

    heading(doc, "Master mapping table", 1)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, label in enumerate(["#", "Team E source", "Action", "Destination in 01_TDD", "Notes"]):
        hdr[i].text = label
        set_cell_shading(hdr[i], "1F4E79")
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(255, 255, 255)
                r.bold = True

    for pkt in PACKETS:
        row = table.add_row().cells
        row[0].text = pkt["id"]
        row[1].text = pkt["title"]
        row[2].text = pkt["action"]
        set_cell_shading(row[2], action_color(pkt["action"]))
        row[3].text = pkt["dest"]
        row[4].text = pkt["guidance"][:180] + ("…" if len(pkt["guidance"]) > 180 else "")

    doc.add_page_break()
    heading(doc, "Revision packets (detail)", 1)

    for pkt in PACKETS:
        heading(doc, f"Packet {pkt['id']} — {pkt['title']}", 2)
        para(doc, f"Action: {pkt['action']}", bold=True, highlight=pkt["action"].upper().startswith("EDIT"))
        para(doc, f"Destination: {pkt['dest']}", bold=True)
        para(doc, "Guidance:")
        para(doc, pkt["guidance"])
        para(doc, "Team E source headings:")
        bullets(doc, pkt["team_e"])

        matches = find_section_blob(sections, pkt["team_e"])
        if not matches:
            para(doc, "No extractable paragraph/table content found for these headings "
                "(section may be figure-only or title-only). Open the Team E TDD and copy "
                "the named section manually.", italic=True)
            continue

        for hname, content in matches:
            heading(doc, f"Extract — {hname}", 3)
            for ptxt in content.get("paras", [])[:12]:
                para(doc, ptxt)
            if len(content.get("paras", [])) > 12:
                para(doc, f"… ({len(content['paras']) - 12} additional paragraphs omitted — see source)", italic=True)
            for ti, rows in enumerate(content.get("tables", [])[:3], start=1):
                para(doc, f"Table extract {ti}:", bold=True)
                # render as compact monospace lines (avoids huge nested tables)
                preview = rows[:12]
                codeish(doc, "\n".join(preview))
                if len(rows) > 12:
                    para(doc, f"… ({len(rows) - 12} additional rows omitted)", italic=True)
            if len(content.get("tables", [])) > 3:
                para(
                    doc,
                    f"… ({len(content['tables']) - 3} additional tables under this heading — copy from source)",
                    italic=True,
                )

    doc.add_page_break()
    heading(doc, "Suggested new heading skeleton for 01_TDD", 1)
    para(
        doc,
        "Paste these headings into the shared TDD where marked ADD. Body text comes from the packets above.",
    )
    codeish(
        doc,
        """4.9 Mail / USPS Entities (MailPiece)
5.11 Ask AI Endpoint (POST /api/ai/ask)
6.12 AI-Assisted Retrieval (Ask AI)
6.13 USPS Informed Delivery Mail Agent
6.14 Short-Term Memory (STML) Support
6.15 Safety, Consent & Clarity
11.20 Team E Feature-Specific Test Cases
11.21 Team E Test Automation
12.9 Team E Known Issues & Limitations
  12.9.1 Technical Limitations
  12.9.2 Known Issues
  12.9.3 Exclusions
  12.9.4 Future Enhancements
  12.9.5 Cross-Team Resolution Notes
13.4 Display Language and Localization Impact   # add if missing under §13""",
    )

    heading(doc, "Checklist before merge review", 1)
    bullets(
        doc,
        [
            "Revision History has a Team E Milestone 2 integration row.",
            "§6.12–§6.15 and §5.11 exist and cross-link each other.",
            "§4.9 MailPiece links to §6.13.",
            "§6.5 dependency register includes Team E rows.",
            "§8.8 mentions Ask AI audit logging; §11.6 mentions AI-tool security tests.",
            "§12.9.1–12.9.5 populated from Team E §14.1–14.5.",
            "List of Tables / List of Figures updated for any copied diagrams.",
        ],
    )

    return doc


def main() -> None:
    if not TEAM_E_TDD.exists():
        raise SystemExit(f"Missing source: {TEAM_E_TDD}")
    doc = build()
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    # also drop a copy next to the target for convenience
    desktop_copy = TARGET_TDD.parent / OUTPUT.name
    doc.save(desktop_copy)
    print(f"Wrote {OUTPUT}")
    print(f"Wrote {desktop_copy}")


if __name__ == "__main__":
    main()
