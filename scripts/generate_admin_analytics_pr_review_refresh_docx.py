"""Generate Word document: refreshed PR review for admin analytics dashboard."""

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.shared import Pt

OUTPUT = (
    Path(__file__).resolve().parents[1]
    / "docs"
    / "PR_Review_Admin_Analytics_Dashboard_feature_a-dkinchen-admin-analytics-dashboard_refresh.docx"
)


def heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def para(doc: Document, text: str, bold: bool = False, highlight: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    if highlight:
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def code(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    for i, h in enumerate(headers):
        tbl.rows[0].cells[i].text = h
    for row in rows:
        cells = tbl.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val


def build() -> Document:
    doc = Document()
    title = doc.add_heading("PR Code Review", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    para(
        doc,
        "Admin Analytics Dashboard — "
        "feature/a-dkinchen-admin-analytics-dashboard → team-ae-develop",
    )
    para(doc, f"Review date: {date.today().isoformat()}")
    para(doc, "Reviewer: AI-assisted code review (Cursor)")
    para(
        doc,
        "Divergence at review time: 89 commits behind / 3 ahead of team-ae-develop. "
        "Rebase is required before merge.",
        highlight=True,
    )
    doc.add_paragraph()

    heading(doc, "Review metadata", 1)
    table(
        doc,
        ["Field", "Value"],
        [
            ["Source branch", "feature/a-dkinchen-admin-analytics-dashboard"],
            ["Target branch", "team-ae-develop"],
            ["Feature commits", "8278da8 — backend endpoint + tests (WBS 2.3.3)"],
            ["Feature commits", "7a63b89 — Flutter UI + 26 tests (WBS 2.3.3 Phase 2 / 2.3.4)"],
            ["Scope (approx.)", "~40 files in three-dot view; rebase will clarify true delta"],
            ["Behind / ahead", "89 behind / 3 ahead"],
        ],
    )

    # ── 1 ────────────────────────────────────────────────────────────────────
    heading(doc, "1. Change Summary", 1)
    para(
        doc,
        "This PR adds an admin-only anonymous product telemetry analytics dashboard. "
        "Backend exposes GET /v1/api/admin/analytics/summary with dual RBAC "
        "(SecurityConfig ADMIN matcher + AuthorizationService.requireAdmin). "
        "It aggregates totals, distinct sessions, top event names, top features, "
        "sync metrics, and errors-by-endpoint. Flyway V75 adds telemetry_events.session_id "
        "plus a partial index. Flutter adds an admin-gated route/drawer, dashboard page, "
        "KPI/charts/cards, API client, models, and widget/unit tests.",
    )

    heading(doc, "What changed (intentional)", 2)
    table(
        doc,
        ["Area", "Change"],
        [
            ["AdminAnalyticsController", "Summary endpoint; resolves user; requireAdmin; range params"],
            ["AdminAnalyticsService", "Aggregation orchestration; range clamp; SAFE_BUCKET sanitize"],
            ["TelemetryEventRepository", "Native windowed aggregation queries + projections"],
            ["V75 migration", "session_id column + idx_telemetry_events_session_id_time"],
            ["SecurityConfig", "Authorize /v1/api/admin/analytics/** for ADMIN"],
            ["Flutter dashboard", "Page, widgets, API, model, router, drawer, tests"],
        ],
    )

    heading(doc, "Three-dot noise", 2)
    para(
        doc,
        "CallSummary / V74-extend / Summary* DTOs appearing in some three-dot views are "
        "merge-base artifacts from an older merge, not intentional analytics work. "
        "Rebase onto current team-ae-develop cleans the PR picture. Do not treat those "
        "as part of this feature review.",
    )

    heading(doc, "Verdict", 2)
    para(
        doc,
        "Request changes. Analytics design and RBAC are solid and tests cover the happy "
        "path, but the branch is 89 commits behind (unsafe to merge as-is) and V75 is "
        "Flyway-only with no SchemaPatchRunner mirror (prod ECS disables Flyway). "
        "Also harden bigint projections and sync JSON casts before production.",
        bold=True,
        highlight=True,
    )

    # ── 2 ────────────────────────────────────────────────────────────────────
    heading(doc, "2. Bug & Risk Analysis", 1)

    heading(doc, "2.1 Strengths", 2)
    bullets(
        doc,
        [
            "Dual admin gate: SecurityConfig.hasRole(ADMIN) + authorizationService.requireAdmin.",
            "Parameterized native SQL — no string-concat injection risk.",
            "Anonymous aggregates only (no PII/PHI fields in DTOs).",
            "SAFE_BUCKET sanitization for event/feature/endpoint labels.",
            "Flutter: admin-only route + drawer; 403 mapped to clear message; defensive _asInt.",
            "Strong unit/widget test counts for controller, service, and Flutter UI.",
        ],
    )

    heading(doc, "2.2 High — Branch 89 commits behind team-ae-develop", 2)
    para(
        doc,
        "git rev-list --left-right --count team-ae-develop...HEAD reports 89 behind / 3 ahead. "
        "Merging without rebase risks regressing Ask AI indexing/retrieval, telemetry expansions, "
        "and other develop work. Rebase (or merge develop) and re-run tests before merge.",
        highlight=True,
    )
    code(
        doc,
        """git fetch origin
git rebase origin/team-ae-develop
# resolve conflicts; re-run backend + Flutter tests""",
    )

    heading(doc, "2.3 High — V75 missing from SchemaPatchRunner", 2)
    para(
        doc,
        "Prod/ECS schema path uses SchemaPatchRunner + Hibernate ddl-auto; Flyway is disabled. "
        "SchemaPatchRunner has no V75 patch. Hibernate may add session_id, but will not create "
        "the partial index idx_telemetry_events_session_id_time. COUNT(DISTINCT session_id) "
        "over up to 90 days can then degrade badly.",
        highlight=True,
    )
    code(
        doc,
        """// SchemaPatchRunner — mirror V75
applyPatch(
    "V75a – add session_id to telemetry_events",
    "ALTER TABLE telemetry_events ADD COLUMN IF NOT EXISTS session_id VARCHAR(64)"
);
applyPatch(
    "V75b – index telemetry_events(session_id, event_time)",
    "CREATE INDEX IF NOT EXISTS idx_telemetry_events_session_id_time "
        + "ON telemetry_events (session_id, event_time DESC) "
        + "WHERE session_id IS NOT NULL"
);""",
    )

    heading(doc, "2.4 Medium — Native projection long vs PostgreSQL bigint", 2)
    para(
        doc,
        "Projections use primitive long getCount()/getAttempted(). PG COUNT/SUM often return "
        "BigInteger, causing ClassCastException at runtime. Unit tests mock projections and "
        "will not catch this.",
        highlight=True,
    )
    code(
        doc,
        """public interface EventNameCountProjection {
  String getEventName();
  Number getCount();  // not long
}

// mapping
new EventNameCountDTO(eventName, projection.getCount().longValue());""",
    )

    heading(doc, "2.5 Medium — Unsafe ::bigint cast on sync JSON fields", 2)
    para(
        doc,
        "sumSyncCompletedBetween casts details->>'attempted'|succeeded|failed to bigint. "
        "One non-numeric value can 500 the entire summary endpoint.",
    )
    code(
        doc,
        """COALESCE(SUM(
  CASE WHEN (details->>'attempted') ~ '^-?[0-9]+$'
       THEN (details->>'attempted')::bigint ELSE 0 END
), 0) AS attempted""",
    )

    heading(doc, "2.6 Medium — Chatty aggregation (8+ queries)", 2)
    para(
        doc,
        "getSummary issues separate queries for totals, sessions, event names, features, "
        "three sync event counts, sync sums, and errors. Acceptable initially; plan FILTER/"
        "conditional aggregates as telemetry_events grows. JSON details->>'…' filters also "
        "lack expression indexes.",
    )

    heading(doc, "2.7 Medium — Flutter full-page loading wipe on refresh", 2)
    para(
        doc,
        "Each _fetchSummary sets _loading = true, replacing the whole scaffold with a spinner "
        "and dropping filter chips until reload completes. Prefer soft refresh when prior "
        "summary exists.",
    )

    heading(doc, "2.8 Low — Duration.toDays() > 90 truncates", 2)
    para(
        doc,
        "Windows just over 90 days (e.g. 90d + hours) still pass validateRange because "
        "toDays() truncates. Prefer window.compareTo(Duration.ofDays(MAX_DAYS)) > 0.",
    )

    heading(doc, "2.9 Low — Clock hardcoded", 2)
    para(
        doc,
        "private final Clock clock = Clock.systemUTC() inside @RequiredArgsConstructor service "
        "prevents injecting a fixed Clock in tests. Inject Clock as a bean/constructor arg.",
    )

    heading(doc, "2.10 Low — Generic Flutter error string", 2)
    para(doc, "_error = 'Error: $e' can surface raw exception text to admins.")

    # ── 3 ────────────────────────────────────────────────────────────────────
    heading(doc, "3. Architecture & Style", 1)

    heading(doc, "3.1 Design patterns", 2)
    bullets(
        doc,
        [
            "Clear controller → service → repository layering with DTOs.",
            "Dual authz matches existing admin controller patterns.",
            "Inclusive-exclusive time window [from, to) is correct for analytics.",
            "Flutter feature folder (data/models/pages/widgets) is clean.",
        ],
    )

    heading(doc, "3.2 Cleanliness", 2)
    bullets(
        doc,
        [
            "Sanitize-before-emit is good hygiene for dashboard labels.",
            "Keep SchemaPatchRunner in lockstep with every Flyway file (repo convention).",
            "After rebase, confirm client telemetry event names (feature_use, sync_*, "
            "endpoint buckets) still align with develop guardrails.",
        ],
    )

    heading(doc, "3.3 Test quality", 2)
    bullets(
        doc,
        [
            "Backend unit tests cover controller authz path and service aggregation with mocks.",
            "Flutter widget/model tests are extensive (~26).",
            "Missing: @DataJpaTest for native SQL / bigint mapping; MVC security integration test.",
        ],
    )

    # ── 4 ────────────────────────────────────────────────────────────────────
    heading(doc, "4. Recommendations", 1)

    heading(doc, "4.1 [High] Rebase onto team-ae-develop", 2)
    code(
        doc,
        """git fetch origin
git rebase origin/team-ae-develop
# resolve conflicts; re-run:
#   backend: AdminAnalyticsControllerTest,AdminAnalyticsServiceTest
#   frontend: flutter test test/features/admin_analytics""",
    )

    heading(doc, "4.2 [High] Mirror V75 in SchemaPatchRunner", 2)
    para(doc, "See §2.3 code block.")

    heading(doc, "4.3 [Medium] Use Number projections + safe casts", 2)
    para(doc, "See §2.4 and §2.5 code blocks. Apply to all count/sum projections.")

    heading(doc, "4.4 [Medium] Soft refresh in Flutter", 2)
    code(
        doc,
        """Future<void> _fetchSummary() async {
  setState(() {
    _error = null;
    if (_summary == null) _loading = true;
  });
  try {
    final summary = await _api.fetchSummary(...);
    if (!mounted) return;
    setState(() => _summary = summary);
  } catch (e) {
    if (!mounted) return;
    setState(() => _error = 'Unable to load analytics. Please retry.');
  } finally {
    if (mounted) setState(() => _loading = false);
  }
}""",
    )

    heading(doc, "4.5 [Low] Inject Clock", 2)
    code(
        doc,
        """@Service
@RequiredArgsConstructor
public class AdminAnalyticsService {
  private final TelemetryEventRepository telemetryEventRepository;
  private final Clock clock;  // provide Clock.systemUTC() @Bean
}""",
    )

    heading(doc, "4.6 [Low] Stricter max window check", 2)
    code(
        doc,
        """if (Duration.between(from, to).compareTo(Duration.ofDays(MAX_DAYS)) > 0) {
  throw new AppException(HttpStatus.BAD_REQUEST,
      "Date range cannot exceed " + MAX_DAYS + " days");
}""",
    )

    # ── File-level ───────────────────────────────────────────────────────────
    heading(doc, "File-level review comments", 1)
    table(
        doc,
        ["File", "Severity", "Comment"],
        [
            ["AdminAnalyticsController.java", "Info", "Solid dual-authz + range delegation."],
            ["AdminAnalyticsService.java", "Medium", "Inject Clock; tighten max-window; consider query batching."],
            ["TelemetryEventRepository.java", "Medium", "Guard ::bigint casts; Number projections; index depends on V75."],
            ["*Projection.java", "Medium", "Change long → Number for native aggregates."],
            ["SecurityConfig.java", "Info", "Correct ADMIN matcher for analytics routes."],
            ["V75__….sql", "High", "Good SQL but incomplete without SchemaPatchRunner mirror."],
            ["SchemaPatchRunner.java", "High", "Missing V75 — blocking for ECS prod."],
            ["AdminAnalytics*Test.java", "Low", "Add DataJpaTest for native SQL/casts after projection fix."],
            ["admin_analytics_dashboard_page.dart", "Medium", "Soft-refresh; avoid full-page wipe."],
            ["admin_analytics_api.dart", "Info", "Sensible timeout / 403 handling."],
            ["admin_analytics_summary_model.dart", "Info", "Defensive parsing — good."],
            ["app_router.dart / role_based_drawer.dart", "Info", "Admin-only UX aligned with backend."],
        ],
    )

    heading(doc, "Test plan (recommended before merge)", 1)
    bullets(
        doc,
        [
            "Rebase onto origin/team-ae-develop and resolve conflicts.",
            "Add SchemaPatchRunner V75a/V75b; verify on a Flyway-disabled profile boot.",
            "mvnw test -Dtest=AdminAnalyticsControllerTest,AdminAnalyticsServiceTest",
            "Prefer a @DataJpaTest hitting real PG COUNT/SUM mapping after Number change.",
            "flutter test for admin_analytics_* tests.",
            "Manual: admin JWT → summary 200; non-admin → 403; days=91 → 400.",
            "Manual: seed sync_completed with non-numeric details.attempted → endpoint must not 500.",
        ],
    )

    return doc


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    build().save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    main()
