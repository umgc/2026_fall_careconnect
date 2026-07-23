"""Generate Word document: Secondary validation + Tier 2 HITL hold/release design."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX

ROOT = Path(__file__).resolve().parents[1]
OUTPUT = (
    ROOT
    / "docs"
    / "Secondary_Validation_and_Tier2_HITL_Hold_Release_Design.docx"
)


def heading(doc, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def para(doc, text: str, bold: bool = False, highlight: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    if highlight:
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def bullets(doc, items: list[str], highlight_indices: set[int] | None = None) -> None:
    highlight_indices = highlight_indices or set()
    for i, item in enumerate(items):
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        if i in highlight_indices:
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def code(doc, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"


def table(
    doc,
    headers,
    rows,
    highlight_rows: set[int] | None = None,
    highlight_cells: set[tuple[int, int]] | None = None,
) -> None:
    highlight_rows = highlight_rows or set()
    highlight_cells = highlight_cells or set()
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for r_idx, row in enumerate(rows, start=1):
        cells = t.add_row().cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = val
            if r_idx in highlight_rows or (r_idx, c_idx) in highlight_cells:
                for p in cells[c_idx].paragraphs:
                    for run in p.runs:
                        run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def build() -> None:
    doc = Document()

    title = doc.add_heading(
        "Secondary Validation & Tier 2 HITL Hold/Release — Design Specification", 0
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para(doc, "CareConnect Team E — Safety, Consent & Clarity (SCC-3 / REQ-SC-4)")
    para(
        doc,
        "Implementation design for the post-LLM safety gate: secondary validation pass, "
        "Tier 1/2 classification, held_item persistence, Human Reviewer release/reject workflow, "
        "and patient polling. Integrates with AiAskOrchestrator, POST /api/ai/ask contract, "
        "and retrieval orchestration pipeline.",
    )
    para(
        doc,
        "Status: DESIGN ONLY — GuardrailService blocks phrases only; no TierClassifier, "
        "held_item table, or HitlService in codebase. EVV review (EvvService.review) provides "
        "an approve/reject pattern analog.",
        highlight=True,
    )
    doc.add_paragraph()

    # 1 Overview
    heading(doc, "1. Design Goals", 1)
    bullets(
        doc,
        [
            "REQ-SC-4: Secondary validation on all AI output before delivery; Tier 2 held until human release.",
            "FR-AI-5: Detect high-risk content; route to HITL before care recipient sees it.",
            "REQ-SC-9 / FR-AI-10: Immutable audit of validation results, tier, hold, release, delivery.",
            "REQ-SC-6: Timeout and inaction never equal approval for held items or write confirmations.",
            "Fail closed: validation failure → 422 WITHHELD or Tier 2 hold — never silent downgrade to Tier 1.",
            "Single SafetyPipeline invoked from Ask AI, summary confirm flows, and STML (future).",
        ],
        highlight_indices={4, 5},
    )

    heading(doc, "1.1 Position in orchestration pipeline", 2)
    code(
        doc,
        "... HybridRetrieval → ContextAssembly → GroundedLlmService\n"
        "                              │\n"
        "                              ▼\n"
        "                    SafetyPipeline.process(SafetyInput)\n"
        "                              │\n"
        "              ┌───────────────┼───────────────┐\n"
        "              ▼               ▼               ▼\n"
        "         BLOCK (422)     Tier 1 DELIVER   Tier 2 HOLD\n"
        "              │               │               │\n"
        "              ▼               ▼               ▼\n"
        "         AiAskResponse   AiAskResponse   HitlService.createHold\n"
        "         WITHHELD        DELIVERED       → poll / release APIs",
    )

    # 2 Components
    heading(doc, "2. Component Architecture", 1)
    table(
        doc,
        ["Component", "Package", "Responsibility"],
        [
            ["SafetyPipeline", "service.ai.safety", "Orchestrates validators → tier decision"],
            ["SecondaryValidationService", "service.ai.safety", "Citation + claim + schema checks"],
            ["TierClassifier", "service.ai.safety", "SRS Table 8 rules → tier + triggerCodes"],
            ["HitlService", "service.ai.hitl", "held_item CRUD, release, reject, expire"],
            ["HitlReviewController", "controller", "Reviewer REST APIs + queue"],
            ["HitlStatusController", "controller", "Patient poll GET status"],
            ["AiSafetyAuditService", "service.ai.safety", "Immutable append-only audit events"],
            ["DisclaimerService", "service.ai.safety", "Locale disclaimer attachment (Tier 1)"],
        ],
        highlight_cells={
            (1, 0), (2, 0), (3, 0), (4, 0),
        },
    )

    heading(doc, "2.1 Relationship to existing code", 2)
    table(
        doc,
        ["Existing", "Design disposition"],
        [
            ["GuardrailService", "Fold forbidden-phrase checks into ResponseContentValidator; stop throwing for tier routing"],
            ["ResponseSanitizationService", "Run as pre-delivery sanitizer on Tier 1 only; Tier 2 sanitizes on release"],
            ["EvvService.review()", "Pattern for approve/reject + audit — not shared implementation"],
            ["ChatAuditService", "Superseded for Ask AI by AiSafetyAuditService with tier/hold fields"],
            ["BedrockAIChatService", "No safety pass today — bypassed by GroundedLlmService + SafetyPipeline"],
        ],
        highlight_rows={1, 5},
    )

    # 3 SafetyInput/Output
    heading(doc, "3. Core Domain Models", 1)
    code(
        doc,
        "record SafetyInput(\n"
        "    String query,\n"
        "    String draftAnswerText,\n"
        "    List<AiCitation> citations,\n"
        "    List<RankedChunk> retrievedChunks,\n"
        "    Long patientId,\n"
        "    Long callerUserId,\n"
        "    UUID sessionId,\n"
        "    UUID auditId,\n"
        "    SourceSurface sourceSurface,   // ASK_AI, CALL_SUMMARY, ...\n"
        "    String locale\n"
        ")\n\n"
        "record SafetyOutcome(\n"
        "    SafetyDecision decision,       // DELIVER_TIER1 | HOLD_TIER2 | BLOCK\n"
        "    int tier,\n"
        "    List<String> triggerCodes,\n"
        "    List<ValidationFinding> findings,\n"
        "    EscalationHint escalation,\n"
        "    UUID heldItemId               // set when HOLD_TIER2\n"
        ")",
    )

    # 4 Secondary validation
    heading(doc, "4. Secondary Validation Pass", 1)
    para(
        doc,
        "SecondaryValidationService runs an ordered chain of validators. Each produces "
        "ValidationFinding(severity, code, message). Any CRITICAL finding → BLOCK (422). "
        "WARN findings feed TierClassifier.",
    )

    heading(doc, "4.1 Validator chain (execution order)", 2)
    table(
        doc,
        ["#", "Validator", "Checks", "On failure"],
        [
            ["V1", "SchemaValidator", "draftAnswerText + citationRefs parseable; required fields", "BLOCK"],
            ["V2", "CitationCoverageValidator", "Every factual claim has ≥1 citation ref (FR-AI-2)", "BLOCK or Tier 2"],
            ["V3", "UnsupportedClaimValidator", "Med/dose/date not in cited excerpts", "Tier 2 UNSUPPORTED_CLAIM"],
            ["V4", "QueryPatternValidator", "Emergency symptom, dosage calc patterns in query", "Tier 2 trigger"],
            ["V5", "OutputPatternValidator", "Medication-change language in draft", "Tier 2 MEDICATION_CHANGE"],
            ["V6", "UncertaintyValidator", "Low-confidence chunks used without needsConfirmation", "Tier 1 WARN"],
            ["V7", "ResponseContentValidator", "GuardrailService forbidden phrases; ResponseSanitization", "BLOCK"],
            ["V8", "RecordsOnlyValidator", "No general medical advice when citations empty", "BLOCK if not NO_RECORDS path"],
        ],
        highlight_rows={3, 4, 5},
    )

    heading(doc, "4.2 UnsupportedClaimValidator algorithm", 2)
    code(
        doc,
        "For each medication/dose/date entity extracted from draftAnswerText:\n"
        "  union = all cited chunk_text + chunk_metadata excerpts\n"
        "  IF entity not substring-normalized-match in union:\n"
        "     finding(CRITICAL, UNSUPPORTED_CLAIM, entity)\n"
        "IF any dosage arithmetic in draft (regex + NLP light):\n"
        "     finding(CRITICAL, DOSAGE_CALC)",
    )

    heading(doc, "4.3 ValidationFinding severity", 2)
    table(
        doc,
        ["Severity", "Effect"],
        [
            ["INFO", "Audit only"],
            ["WARN", "TierClassifier input; may add confirm-with-provider escalation"],
            ["CRITICAL", "BLOCK → HTTP 422 SAFETY_VALIDATION_FAILED (unless policy maps to Tier 2 hold)"],
        ],
    )

    # 5 Tier classifier
    heading(doc, "5. TierClassifier Design", 1)
    para(doc, "Rule engine evaluated after secondary validation. First matching Tier 2 rule wins hold.", bold=True)

    heading(doc, "5.1 Tier 2 rules (hold — SRS Table 8)", 2)
    table(
        doc,
        ["Priority", "Condition", "triggerCode"],
        [
            ["1", "Query matches EMERGENCY_SYMPTOM_PATTERNS", "EMERGENCY_SYMPTOM"],
            ["2", "Query matches DOSAGE_CALC_PATTERNS", "DOSAGE_CALC"],
            ["3", "Draft or query matches MEDICATION_CHANGE_PATTERNS", "MEDICATION_CHANGE"],
            ["4", "Finding UNSUPPORTED_CLAIM (unsupported med/dose/date)", "UNSUPPORTED_CLAIM"],
            ["5", "Summary item needsConfirmation=true AND type=medication", "SUMMARY_MED_PENDING"],
            ["6", "Policy flag forceReviewForCaregiverQueries", "CAREGIVER_POLICY_HOLD"],
        ],
        highlight_rows={1, 2, 3, 4},
    )

    heading(doc, "5.2 Tier 1 rules (deliver with escalation metadata)", 2)
    table(
        doc,
        ["Condition", "escalation.level", "triggerCode"],
        [
            ["General medication mention in answer", "confirm-with-provider", "GENERAL_MEDICATION_MENTION"],
            ["Empty retrieval / no-records path", "confirm-with-provider", "NO_MATCHING_RECORDS"],
            ["WARN uncertainty on cited chunks", "none (needsConfirmation badge in UI)", "LOW_CONFIDENCE_EXTRACTION"],
            ["Default safe grounded answer", "none", "[]"],
        ],
    )

    heading(doc, "5.3 TierClassifier pseudocode", 2)
    code(
        doc,
        "TierResult classify(SafetyInput input, List<ValidationFinding> findings) {\n"
        "  if (findings.any(CRITICAL && policy.block)) return BLOCK;\n"
        "  for (rule : TIER2_RULES_IN_PRIORITY_ORDER)\n"
        "    if (rule.matches(input, findings))\n"
        "      return TierResult(2, rule.triggerCode, HOLD_TIER2);\n"
        "  EscalationHint esc = TIER1_ESCALATION_RULES.firstMatch(input, findings);\n"
        "  return TierResult(1, esc, DELIVER_TIER1);\n"
        "}",
    )

    # 6 Hold release workflow
    heading(doc, "6. Tier 2 Hold/Release Workflow", 1)

    heading(doc, "6.1 State machine", 2)
    code(
        doc,
        "                    createHold()\n"
        "   [none] ──────────────────────────► PENDING_REVIEW\n"
        "                                           │\n"
        "           ┌───────────────────────────────┼────────────────────────┐\n"
        "           │ release(asIs)               │ release(edited)         │ reject()\n"
        "           ▼                               ▼                         ▼\n"
        "      APPROVED_AS_IS                  APPROVED_EDITED              REJECTED\n"
        "           │                               │                         │\n"
        "           └─────────── deliverToPatient() ┴─────────────────────────┘\n"
        "                                           │\n"
        "                                           ▼\n"
        "                                    DELIVERED / WITHHELD_PERMANENTLY\n"
        "\n"
        "  expireJob(): PENDING_REVIEW ──(TTL exceeded)──► EXPIRED",
    )
    para(
        doc,
        "REQ-SC-6: No transition from PENDING_REVIEW without explicit reviewer_user_id on release/reject.",
        highlight=True,
    )

    heading(doc, "6.2 Database schema (Flyway)", 2)
    code(
        doc,
        "CREATE TABLE ai_held_item (\n"
        "  id UUID PRIMARY KEY,\n"
        "  patient_id BIGINT NOT NULL,\n"
        "  requester_user_id BIGINT NOT NULL,\n"
        "  session_id UUID,\n"
        "  audit_id UUID NOT NULL,\n"
        "  source_surface VARCHAR(32) NOT NULL,\n"
        "  status VARCHAR(24) NOT NULL,\n"
        "  tier SMALLINT NOT NULL DEFAULT 2,\n"
        "  trigger_codes JSONB NOT NULL,\n"
        "  query_text_hash VARCHAR(64),\n"
        "  draft_answer TEXT NOT NULL,\n"
        "  final_answer TEXT,\n"
        "  citations_json JSONB NOT NULL,\n"
        "  validation_findings_json JSONB,\n"
        "  reviewer_user_id BIGINT,\n"
        "  reviewed_at TIMESTAMPTZ,\n"
        "  review_notes VARCHAR(500),\n"
        "  delivery_status VARCHAR(24) NOT NULL,\n"
        "  expires_at TIMESTAMPTZ,\n"
        "  created_at TIMESTAMPTZ NOT NULL DEFAULT now(),\n"
        "  updated_at TIMESTAMPTZ NOT NULL DEFAULT now()\n"
        ");\n"
        "CREATE INDEX idx_held_patient_status ON ai_held_item(patient_id, status);\n"
        "CREATE INDEX idx_held_pending ON ai_held_item(status) WHERE status = 'PENDING_REVIEW';\n\n"
        "CREATE TABLE ai_safety_audit_event (\n"
        "  id UUID PRIMARY KEY,\n"
        "  audit_id UUID NOT NULL,\n"
        "  held_item_id UUID,\n"
        "  event_type VARCHAR(40) NOT NULL,\n"
        "  actor_user_id BIGINT,\n"
        "  payload_json JSONB,\n"
        "  created_at TIMESTAMPTZ NOT NULL DEFAULT now()\n"
        ");  -- append-only; no UPDATE/DELETE",
    )

    heading(doc, "6.3 HitlService operations", 2)
    table(
        doc,
        ["Method", "Behavior", "Audit event"],
        [
            ["createHold(SafetyInput, SafetyOutcome)", "Insert PENDING_REVIEW; return heldItemId", "HITL_HELD"],
            ["release(id, reviewer, editedAnswer?)", "Set APPROVED_*; final_answer; deliver", "HITL_RELEASED"],
            ["reject(id, reviewer, reason)", "Set REJECTED; delivery WITHHELD_PERMANENTLY", "HITL_REJECTED"],
            ["getStatus(id, requester)", "Patient/caregiver poll; redact draft until DELIVERED", "—"],
            ["listQueue(reviewer, filters)", "PENDING_REVIEW items for reviewer role", "—"],
            ["expireStale()", "Scheduled job: PENDING → EXPIRED after policy TTL", "HITL_EXPIRED"],
        ],
        highlight_rows={1, 2, 3},
    )

    heading(doc, "6.4 Patient visibility rules", 2)
    bullets(
        doc,
        [
            "PENDING_REVIEW: patient sees held message only — no draft_answer, no citations content.",
            "After APPROVED_*: patient poll returns final_answer + citations (sanitized).",
            "REJECTED: patient sees safe fallback message; optional link to contact provider.",
            "Caregiver who submitted query may see held status but not draft unless reviewer role.",
            "Reviewer sees full draft + citations + trigger_codes + validation findings.",
        ],
        highlight_indices={0, 4},
    )

    # 7 REST API
    heading(doc, "7. REST API Design", 1)

    heading(doc, "7.1 Patient / caller APIs", 2)
    table(
        doc,
        ["Method", "Path", "Description"],
        [
            ["POST", "/v1/api/ai/ask", "Returns held=true + heldItemId when Tier 2 (see API contract doc)"],
            ["GET", "/v1/api/ai/hitl/{heldItemId}/status", "Poll until DELIVERED | REJECTED | EXPIRED"],
        ],
    )
    code(
        doc,
        "// GET .../status — patient-visible\n"
        "{\n"
        '  "heldItemId": "uuid",\n'
        '  "status": "PENDING_REVIEW | DELIVERED | REJECTED | EXPIRED",\n'
        '  "deliveryStatus": "HELD | DELIVERED | WITHHELD_PERMANENTLY",\n'
        '  "message": "We\'re reviewing this before showing it to you.",\n'
        '  "answer": null,                    // until DELIVERED\n'
        '  "citations": [],\n'
        '  "disclaimer": { ... }\n'
        "}",
    )

    heading(doc, "7.2 Reviewer APIs (Human Reviewer / Clinician role)", 2)
    table(
        doc,
        ["Method", "Path", "Body", "Response"],
        [
            ["GET", "/v1/api/ai/hitl/queue", "—", "List<HeldItemSummary>"],
            ["GET", "/v1/api/ai/hitl/{id}", "—", "HeldItemDetail (draft + citations + findings)"],
            ["POST", "/v1/api/ai/hitl/{id}/release", "{ approve: true, editedAnswer?: string, notes?: string }", "HeldItemDetail"],
            ["POST", "/v1/api/ai/hitl/{id}/reject", "{ reason: string }", "HeldItemDetail"],
        ],
        highlight_rows={2, 3},
    )
    para(
        doc,
        "Model after EvvReviewRequest (approve + comment) but separate DTOs: HitlReleaseRequest, HitlRejectRequest.",
    )

    heading(doc, "7.3 WebSocket optional enhancement", 2)
    para(
        doc,
        "Optional: push heldItemId status change on existing WebSocket channel when reviewer "
        "releases — reduces poll latency. MVP: HTTP poll every 5s acceptable.",
    )

    # 8 Integration
    heading(doc, "8. Integration Points", 1)

    heading(doc, "8.1 AiAskOrchestrator post-processor", 2)
    code(
        doc,
        "SafetyOutcome outcome = safetyPipeline.process(SafetyInput.from(ctx, llm, req));\n"
        "switch (outcome.decision()) {\n"
        "  case BLOCK -> throw SafetyValidationException → 422;\n"
        "  case HOLD_TIER2 -> {\n"
        "    UUID heldId = hitlService.createHold(...);\n"
        "    return AiAskResponse.held(heldId, outcome.triggerCodes());\n"
        "  }\n"
        "  case DELIVER_TIER1 -> {\n"
        "    disclaimerService.attach(...);\n"
        "    return AiAskResponse.delivered(llm, citations, outcome.escalation());\n"
        "  }\n"
        "}",
    )

    heading(doc, "8.2 Summary confirm flow (FR-SUM-4)", 2)
    bullets(
        doc,
        [
            "When user attempts confirm on summary careInstruction with needsConfirmation=true:",
            "Run SafetyPipeline on item text + citations before write.",
            "Tier 2 → hold summary action; show same held UX pattern.",
            "Separate source_surface=CALL_SUMMARY on held_item row.",
        ],
    )

    heading(doc, "8.3 SCC-3 step mapping", 2)
    table(
        doc,
        ["SCC-3 step", "Design component"],
        [
            ["11 Schema validation", "SchemaValidator (V1)"],
            ["12 Secondary validation", "SecondaryValidationService (V2–V8)"],
            ["13 Tier assignment", "TierClassifier"],
            ["15–16 Tier 2 hold", "HitlService.createHold"],
            ["17–18 Tier 1 deliver", "DisclaimerService + AiAskResponse"],
            ["19 Audit", "AiSafetyAuditService"],
        ],
    )

    # 9 Config
    heading(doc, "9. Configuration", 1)
    table(
        doc,
        ["Property", "Default", "Purpose"],
        [
            ["careconnect.ai.safety.enabled", "true", "Master switch when Ask AI enabled"],
            ["careconnect.ai.hitl.ttl-hours", "72", "PENDING_REVIEW → EXPIRED"],
            ["careconnect.ai.hitl.reviewer-roles", "HUMAN_REVIEWER,CLINICIAN,ADMIN", "Queue access"],
            ["careconnect.ai.safety.block-on-unsupported-claim", "true", "422 vs Tier 2 hold policy"],
            ["careconnect.ai.safety.emergency-tier2", "true", "Emergency → hold + crisis copy"],
        ],
    )

    # 10 Security
    heading(doc, "10. Security & RBAC", 1)
    bullets(
        doc,
        [
            "New permission: REVIEW_AI_HELD_ITEMS for Human Reviewer role.",
            "Queue list scoped: reviewers see all pending OR org-scoped by patient link (policy TBD).",
            "Patient poll: heldItemId + requester must match requester_user_id or patient scope.",
            "draft_answer encrypted at rest (AES-256/KMS) optional Phase 2.",
            "Audit events append-only; no application UPDATE on ai_safety_audit_event.",
        ],
        highlight_indices={0, 4},
    )

    # 11 Failure modes
    heading(doc, "11. Failure Modes", 1)
    table(
        doc,
        ["Scenario", "Behavior"],
        [
            ["Validator throws", "Treat as BLOCK; log error; 422 WITHHELD"],
            ["Hitl DB write fails after Tier 2 classify", "Do NOT deliver draft; 503; retry idempotent on auditId"],
            ["Reviewer release while patient poll active", "Optimistic lock on status; second release → 409"],
            ["Duplicate createHold same auditId", "Idempotent return existing heldItemId"],
        ],
        highlight_rows={1, 2},
    )

    # 12 Gap analysis
    heading(doc, "12. Codebase Gap Analysis", 1)
    table(
        doc,
        ["Area", "Today", "Design target"],
        [
            ["Secondary validation", "None on Bedrock path", "SafetyPipeline + 8 validators"],
            ["Tier classification", "None", "TierClassifier rule engine"],
            ["HITL queue", "None", "ai_held_item + HitlService"],
            ["Reviewer APIs", "None", "HitlReviewController"],
            ["Patient poll", "None", "GET /hitl/{id}/status"],
            ["Immutable audit", "ChatAuditService metadata only", "ai_safety_audit_event"],
            ["Human Reviewer role", "Not in RolePermissionService", "New role + permission"],
            ["Analog", "EvvService.review approve/reject", "Reuse pattern, not code"],
        ],
        highlight_rows={1, 2, 3, 4, 5, 6},
    )

    # 13 Implementation phases
    heading(doc, "13. Implementation Phases", 1)
    table(
        doc,
        ["Phase", "Deliverable", "Tests"],
        [
            ["P0", "Flyway ai_held_item + ai_safety_audit_event; SafetyPipeline skeleton", "Unit: state machine"],
            ["P0", "SecondaryValidationService V1–V5; TierClassifier Tier 2 rules", "TC-E-SC-001, TC-SCC-03"],
            ["P0", "HitlService createHold + AiAskOrchestrator integration", "Held response contract"],
            ["P1", "Reviewer REST + queue UI; release/reject", "TC-AI-06"],
            ["P1", "Patient poll + optional WebSocket", "TC-E-SC-002 timeout semantics"],
            ["P2", "Summary confirm HITL; expire job; Spanish copy", "TC-E-SUM-002"],
            ["P2", "KMS encryption for draft_answer; append-only audit enforcement", "REQ-SC-9 compliance"],
        ],
        highlight_rows={1, 2, 3},
    )

    # 14 Related docs
    heading(doc, "14. Related Documents", 1)
    bullets(
        doc,
        [
            "docs/HITL_Tier1_Tier2_Escalation_and_Hold_Release_Workflow.docx (requirements research)",
            "docs/POST_api_ai_ask_Request_Response_Contract_Design.docx",
            "docs/Retrieval_Orchestration_RBAC_Hybrid_Context_LLM_Design.docx",
            "docs/Team_E_Implementation_Task_Backlog.docx (Tier 6 tasks)",
            "backend/core/src/main/java/com/careconnect/service/chat/GuardrailService.java",
            "backend/core/src/main/java/com/careconnect/service/evv/EvvService.java (review analog)",
            "backend/core/src/main/java/com/careconnect/controller/EvvController.java",
        ],
    )

    heading(doc, "15. Conclusion", 1)
    para(
        doc,
        "Secondary validation and Tier 2 HITL are implemented as SafetyPipeline — a composable "
        "validator chain plus TierClassifier rule engine — followed by HitlService persistence "
        "and explicit Human Reviewer release/reject APIs. Tier 2 never delivers draft content to "
        "the care recipient until a reviewer action transitions held_item to APPROVED and "
        "delivery_status to DELIVERED.",
    )
    para(
        doc,
        "Build order: database schema → SecondaryValidationService + TierClassifier → HitlService "
        "→ wire AiAskOrchestrator → reviewer APIs → patient poll. Do not enable patient-facing "
        "Ask AI without P0 hold path complete.",
        highlight=True,
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    try:
        doc.save(OUTPUT)
        print(f"Created: {OUTPUT}")
    except PermissionError:
        alt = OUTPUT.with_name(f"{OUTPUT.stem}_refresh{OUTPUT.suffix}")
        doc.save(alt)
        print(f"Created (original locked): {alt}")


if __name__ == "__main__":
    build()
