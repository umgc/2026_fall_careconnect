"""Generate Word document: pgvector embedding strategy for summaries, mail, documents."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX

ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "pgvector_Embedding_Strategy_Summaries_Mail_Documents.docx"


def heading(doc, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def para(doc, text: str, bold: bool = False, highlight: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    if highlight:
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def bullets(doc, items: list[str], highlight_indices: set[int] | None = None) -> None:
    highlight_indices = highlight_indices or set()
    for i, item in enumerate(items):
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        if i in highlight_indices:
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def code(doc, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"


def table(doc, headers, rows, highlight_rows: set[int] | None = None) -> None:
    highlight_rows = highlight_rows or set()
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for r_idx, row in enumerate(rows, start=1):
        cells = t.add_row().cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = val
            if r_idx in highlight_rows:
                for p in cells[c_idx].paragraphs:
                    for run in p.runs:
                        run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def build() -> None:
    doc = Document()

    title = doc.add_heading(
        "pgvector Embedding Strategy: Summaries, Mail Pieces, and Document Chunks",
        0,
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para(doc, "CareConnect Team E — Hybrid Retrieval (PostgreSQL FTS + pgvector)")
    para(
        doc,
        "Defines how to generate, store, and query embeddings for call/visit summaries, "
        "USPS mail pieces, and uploaded document chunks in retrieval_index_chunk. "
        "Synthesizes TDD §7.1, Hybrid Retrieval scope, Storage Audit, RBAC research, and codebase review.",
    )
    para(
        doc,
        "Yellow highlights mark decisions still open or gaps in current code.",
        highlight=True,
    )
    doc.add_paragraph()

    # 1 Executive summary
    heading(doc, "1. Executive Summary", 1)
    para(
        doc,
        "Team E hybrid retrieval stores all searchable content in a single "
        "retrieval_index_chunk table: chunk_text + tsvector (FTS) + pgvector embedding. "
        "Embeddings must be produced by a HIPAA-eligible provider under the AWS BAA — "
        "Amazon Bedrock Titan Text Embeddings is the recommended default, accessed through "
        "the same LlmRouter / abstraction layer as chat inference.",
    )
    para(
        doc,
        "No embedding pipeline exists in the codebase today. pgvector is enabled only in "
        "local Docker (pgvector/pgvector:pg15 + CREATE EXTENSION vector). "
        "PatientContextRetrievalService uses in-memory substring match — not vectors.",
        highlight=True,
    )
    bullets(
        doc,
        [
            "One embedding model, one vector dimension, one table — all record types share the same schema.",
            "Chunking strategy differs by source; embedding input text is a normalized template per type.",
            "Query-time: embed user question → patient-scoped vector search → merge with FTS via RRF.",
            "Async index refresh ≤ 5 minutes after source write (NFR-AI-3).",
        ],
    )

    # 2 Target storage
    heading(doc, "2. Target Storage (retrieval_index_chunk)", 1)
    code(
        doc,
        "CREATE EXTENSION IF NOT EXISTS vector;\n\n"
        "CREATE TABLE retrieval_index_chunk (\n"
        "  id UUID PRIMARY KEY DEFAULT gen_random_uuid(),\n"
        "  patient_id BIGINT NOT NULL,\n"
        "  record_type VARCHAR(40) NOT NULL,\n"
        "  source_record_id VARCHAR(120) NOT NULL,\n"
        "  chunk_index INT NOT NULL DEFAULT 0,\n"
        "  chunk_text TEXT NOT NULL,\n"
        "  chunk_metadata JSONB,\n"
        "  content_hash VARCHAR(64),          -- SHA-256 of chunk_text for idempotent re-embed\n"
        "  search_vector TSVECTOR,\n"
        "  embedding vector(1024),             -- lock to chosen Bedrock Titan v2 dimension\n"
        "  embed_model_id VARCHAR(80),\n"
        "  indexed_at TIMESTAMPTZ NOT NULL DEFAULT now(),\n"
        "  consent_scope VARCHAR(40),\n"
        "  UNIQUE (source_record_id, record_type, chunk_index)\n"
        ");\n\n"
        "CREATE INDEX idx_retrieval_patient_type ON retrieval_index_chunk(patient_id, record_type);\n"
        "CREATE INDEX idx_retrieval_fts ON retrieval_index_chunk USING GIN(search_vector);\n"
        "CREATE INDEX idx_retrieval_vector ON retrieval_index_chunk\n"
        "  USING hnsw (embedding vector_cosine_ops) WITH (m = 16, ef_construction = 64);",
    )
    para(
        doc,
        "Note: prior research docs used vector(1536) (OpenAI ada / Titan v1 size). "
        "Recommend standardizing on amazon.titan-embed-text-v2:0 with dimensions=1024 "
        "unless Team A already provisioned v1 — dimension must match across index and query embedder.",
        highlight=True,
    )

    # 3 Embedding model
    heading(doc, "3. Embedding Model Selection", 1)
    table(
        doc,
        ["Option", "Model ID", "Dimensions", "BAA", "Recommendation"],
        [
            ["Primary (recommended)", "amazon.titan-embed-text-v2:0", "256 / 512 / 1024 (pick one)", "AWS BAA", "Default for all chunk types"],
            ["Alternate", "amazon.titan-embed-text-v1", "1536", "AWS BAA", "Legacy; matches old 1536 schema drafts"],
            ["Alternate", "cohere.embed-english-v3", "1024", "AWS BAA via Bedrock", "If multilingual mail OCR needed later"],
            ["Excluded for PHI", "OpenAI text-embedding-3-* direct API", "varies", "Separate BAA required", "Do not use in prod without enterprise BAA"],
        ],
        highlight_rows={1, 4},
    )
    bullets(
        doc,
        [
            "Use BedrockRuntimeClient InvokeModel (same IAM path as Nova/Claude chat) — add embed model ARNs to ECS task role.",
            "Store embed_model_id on each chunk row for re-indexing when model changes.",
            "Query embeddings must use the identical model + dimension as index embeddings.",
            "Batch embed: group up to 25 chunks per Bedrock call where API supports input array (reduce latency/cost).",
            "Cache query embeddings per sessionId + normalized query hash (TTL 5–15 min) — TDD mentions embedding cache.",
        ],
    )

    # 4 pgvector index strategy
    heading(doc, "4. pgvector Index Strategy", 1)

    heading(doc, "4.1 Index type: HNSW vs IVFFlat", 2)
    table(
        doc,
        ["Index", "When to use", "CareConnect fit"],
        [
            ["HNSW (recommended)", "Production; moderate data per patient; pgvector ≥ 0.5", "Best default for multi-tenant patient-scoped search"],
            ["IVFFlat", "Very large corpus; batch-built after bulk load", "Fallback if HNSW memory constrained on Aurora"],
            ["No global index only", "—", "Always combine patient_id filter with vector query"],
        ],
    )
    para(doc, "Query pattern (always scope before similarity):", bold=True)
    code(
        doc,
        "SELECT id, chunk_text, chunk_metadata,\n"
        "       1 - (embedding <=> :query_embedding) AS cosine_sim\n"
        "FROM retrieval_index_chunk\n"
        "WHERE patient_id = :patient_id\n"
        "  AND record_type = ANY(:allowed_types)\n"
        "ORDER BY embedding <=> :query_embedding\n"
        "LIMIT 20;",
    )
    para(
        doc,
        "Task 4.4 backfill is an application worker that selects rows where embedding IS NULL "
        "and re-invokes Bedrock via ChunkEmbeddingService. Optional DBA follow-up: "
        "SchemaPatchRunner patch V2607161317 creates partial index "
        "idx_retrieval_chunk_embedding_null_backfill on (indexed_at, id) for embeddable NULL rows — "
        "reference SQL in db/migration; not applied via Flyway at ECS deploy. After bulk backfill: "
        "ANALYZE retrieval_index_chunk; tune ivfflat lists on idx_retrieval_chunk_embedding.",
    )
    bullets(
        doc,
        [
            "Partial indexes optional: CREATE INDEX … WHERE record_type = 'USPS_MAIL' for mail-only API.",
            "After bulk backfill: ANALYZE retrieval_index_chunk; tune ef_search (e.g. 40) at session level for HNSW.",
            "Cosine distance (<=>) with vector_cosine_ops — normalize Bedrock Titan outputs if not pre-normalized.",
        ],
    )

    heading(doc, "4.2 Hybrid merge at query time", 2)
    para(
        doc,
        "Parallel FTS (ts_rank_cd on search_vector) and vector top-K; merge with Reciprocal Rank Fusion (RRF, k=60):",
    )
    code(
        doc,
        "score_rrf(d) = sum over rank lists: 1 / (k + rank_i(d))\n"
        "Final top-K fed to Ask AI prompt assembly (FR-AI-9 minimum-necessary context).",
    )
    para(
        doc,
        "USPS GET /api/mail (FR-USPS-4): weight FTS higher for exact sender/tracking matches; "
        "vector higher for natural-language queries (“package from pharmacy”).",
    )

    # 5 Summary embeddings
    heading(doc, "5. Summary Embedding Strategy", 1)

    heading(doc, "5.1 Current vs target source data", 2)
    table(
        doc,
        ["Aspect", "Codebase today", "Embedding impact"],
        [
            ["Storage", "call_summaries.summary_json TEXT", "Parse JSON before chunking"],
            ["Schema", "Reduced Bedrock shape (headline, keyConcerns, …)", "GAP — unified contract needed for item-level chunks"],
            ["patient_id", "Not on row — resolve via call_id → telemetry", "Must set on every chunk at index time"],
            ["Trigger", "CallSummaryService.generateAndStoreSummary only", "GAP — no INDEX job after save"],
        ],
        highlight_rows={3, 4},
    )

    heading(doc, "5.2 Summary chunk levels", 2)
    table(
        doc,
        ["record_type", "Granularity", "embed_input template", "FTS fields"],
        [
            [
                "CALL_SUMMARY",
                "1 chunk per summary",
                "Call summary\\nHeadline: {headline}\\nAssessment: {overallAssessment}\\nRisk: {riskLevel}",
                "headline, narrative, riskLevel",
            ],
            [
                "VISIT_SUMMARY",
                "1 chunk per visit summary",
                "Visit summary\\nEpisode: visit\\nHeadline: … (same inner JSON as call)",
                "Same as call",
            ],
            [
                "SUMMARY_ITEM",
                "1 chunk per array item",
                "Care instruction | Medication: {name} | Status: {status} | Date: {effectiveDate}\\n{text}",
                "item text, type, status",
            ],
            [
                "MEDICATION_TIMELINE_EVENT",
                "Derived from careInstructions type=medication",
                "Medication event | {name} | {status} | {effectiveDate}\\n{citation excerpt}",
                "medication name, status (structured prefilter first — FR-AI-11)",
            ],
        ],
        highlight_rows={3, 4},
    )
    bullets(
        doc,
        [
            "Do not embed raw full summary_json blob — embed human-readable normalized text per chunk.",
            "Include episodeType and call_id/visit_id in chunk_metadata, not in embed text (reduces token noise).",
            "On summary regen: delete chunks WHERE source_record_id = summaryId; re-insert (or upsert by content_hash).",
            "Store caregiver_visibility in chunk_metadata for RBAC filter at query time.",
        ],
    )

    # 6 Mail embeddings
    heading(doc, "6. USPS Mail Piece Embedding Strategy", 1)

    heading(doc, "6.1 Current mail storage & search", 2)
    table(
        doc,
        ["Aspect", "Codebase (USPSDigestService)", "Embedding strategy"],
        [
            ["Store", "usps_digest_cache.payloadJson (full digest, 24h TTL)", "Explode to per-piece chunks at index time"],
            ["Scope key", "userId (String)", "Map userId → patient_id before insert"],
            ["Search today", "Keyword contains on sender, subject, id only", "FTS + vector over richer embed text"],
            ["OCR", "MailpieceOcrService — sender label from thumbnail via Textract", "Append OCR sender to embed text; GAP — no full visible_text column yet"],
            ["Packages", "PackageItem: trackingNumber, sender, expectedDate", "Separate record_type USPS_PACKAGE"],
        ],
        highlight_rows={1, 3, 4},
    )

    heading(doc, "6.2 Mail embed_input templates", 2)
    code(
        doc,
        "# MailPiece (record_type = USPS_MAIL)\n"
        "USPS informed delivery mail piece\n"
        "Sender: {sender}\n"
        "Subject: {subject}\n"
        "Received: {receivedAt}\n"
        "Digest date: {digestDate}\n"
        "Visible text: {ocrVisibleText or empty}\n\n"
        "# PackageItem (record_type = USPS_PACKAGE)\n"
        "USPS package delivery\n"
        "Sender: {sender}\n"
        "Tracking: {trackingNumber}\n"
        "Expected delivery: {expectedDeliveryDate}",
    )
    bullets(
        doc,
        [
            "One embedding per mail piece / package — do not embed entire digest JSON as single vector.",
            "source_record_id: {userId}:{digestDate}:{piece.id} or stable piece UUID.",
            "Thumbnail images: do not embed image bytes; optional future multimodal model is out of MVP scope.",
            "Hybrid mail search: structured filter on delivery_date range + RRF(FTS, vector).",
            "Consent: mail indexing subject to REQ-SC-7 user exclusion of USPS source type.",
        ],
        highlight_indices={2},
    )

    # 7 Document embeddings
    heading(doc, "7. Document Chunk Embedding Strategy", 1)

    heading(doc, "7.1 Current document handling", 2)
    table(
        doc,
        ["Component", "Behavior", "Gap for embeddings"],
        [
            ["user_files", "BLOB or S3; patient_id nullable; FileCategory enum", "Need patient_id enforced for medical categories"],
            ["DocumentProcessingService", "Extract text: PDF/DOC/DOCX/Tika; max 15,000 chars total", "No sliding-window chunker; no persisted extract"],
            ["TextractService", "Invoice OCR path", "Separate from user_files — not unified indexer yet"],
            ["AI chat attach", "On-demand extract via MedicalContextService", "Not durable for retrieval index"],
        ],
        highlight_rows={2, 3},
    )

    heading(doc, "7.2 Document chunking parameters", 2)
    table(
        doc,
        ["Parameter", "Recommended value", "Rationale"],
        [
            ["Target chunk size", "~512 tokens (~2,000 characters)", "Fits Titan embed input; aligns with clinical paragraph scale"],
            ["Overlap", "10–15% (~200 characters)", "Preserves context across lab result tables / page breaks"],
            ["Max chunks per file", "200 (cap)", "Prevents runaway index on huge PDFs; log truncation"],
            ["record_type", "DOCUMENT", "Distinct from CLINICAL_NOTE (patient_note — single chunk, no split)"],
            ["chunk_metadata", "file_id, chunk_index, page_hint, mime, category", "Citation deep link to FileController download"],
        ],
    )
    para(doc, "embed_input template per document chunk:", bold=True)
    code(
        doc,
        "Medical document ({file_category})\n"
        "File: {original_filename}\n"
        "Uploaded: {uploaded_at}\n"
        "---\n"
        "{chunk_text}",
    )
    bullets(
        doc,
        [
            "Pipeline: upload → extract full text (Tika/Textract) → persist document_extracted_text sidecar → split → embed → index.",
            "Re-index on file replace: content_hash change triggers chunk delete + rebuild.",
            "Lab results / prescriptions: boost FTS weight on structured tokens (dates, mg, panel names) via tsvector weights.",
        ],
        highlight_indices={0},
    )

    # 8 Pipeline
    heading(doc, "8. Embedding Pipeline Architecture", 1)
    code(
        doc,
        "Source write (summary / mail ingest / document OCR complete)\n"
        "    → outbox event { type, sourceId, patientId }\n"
        "    → IndexWorker: Chunker → chunk_text[]\n"
        "    → EmbeddingService (Bedrock Titan v2)\n"
        "    → Upsert retrieval_index_chunk (FTS trigger + embedding column)\n"
        "    → indexed_at; alarm if now - source.updated_at > 5 min",
    )
    table(
        doc,
        ["Event", "Chunker", "Priority"],
        [
            ["SUMMARY_CREATED", "SummaryChunker (levels 5.2)", "P0"],
            ["TRANSCRIPT_INDEXED", "TranscriptSegmentChunker (1 seg = 1 chunk)", "P0"],
            ["DOCUMENT_OCR_COMPLETE", "DocumentSlidingWindowChunker", "P1"],
            ["USPS_DIGEST_INGESTED", "MailPieceChunker", "P1"],
            ["CONSENT_REVOKED / DELETE", "Delete chunks by patient_id + record_type", "P0"],
        ],
    )
    bullets(
        doc,
        [
            "Idempotency: UNIQUE(source_record_id, record_type, chunk_index) + content_hash skip if unchanged.",
            "Failure: retry with exponential backoff; dead-letter queue; do not block source OLTP commit.",
            "FR-AI-9: embed only chunk_text destined for retrieval — not full patient record dumps.",
            "Audit: log embed job patient_id, record_type, chunk count — not embedding vector values.",
        ],
    )

    # 9 Query embedding
    heading(doc, "9. Query-Time Embedding", 1)
    bullets(
        doc,
        [
            "Ask AI (/api/ai/ask): embed sanitized user query once per request; reuse for all record_type searches.",
            "Medication timeline (FR-AI-11): run structured JSON/ SQL prefilter on MEDICATION_TIMELINE_EVENT before vector search.",
            "Mail search (/api/mail): embed natural-language query; combine with keyword branch for tracking numbers.",
            "Voice query: embed text after STT — same model path as text Ask AI.",
        ],
    )

    # 10 Codebase gaps
    heading(doc, "10. Codebase Gap Analysis", 1)
    table(
        doc,
        ["Gap", "Impact", "Remediation"],
        [
            ["No retrieval_index_chunk migration", "Cannot store vectors", "Flyway V-next + pgvector on Aurora"],
            ["No EmbeddingService / Bedrock embed client", "No vectors generated", "New service + IAM for Titan embed ARN"],
            ["No chunker implementations", "No index input", "SummaryChunker, MailPieceChunker, DocumentChunker"],
            ["Summary JSON incomplete", "Item-level embeddings blocked", "Unified schema (Ask AI upstream doc)"],
            ["USPS cache not durable / no patient_id", "Mail vectors not scoped", "Normalize mail_piece table + patient map"],
            ["Document text not persisted", "Re-extract on every index", "document_extracted_text sidecar on user_files"],
            ["PatientContextRetrievalService stub", "Misleading placeholder", "Replace with HybridRetrievalService"],
            ["Hybrid doc used vector(1536); no dimension lock in code", "Schema drift risk", "Pick Titan v2 1024 and update all docs"],
        ],
        highlight_rows={1, 2, 3, 4, 5, 8},
    )

    # 11 Implementation checklist
    heading(doc, "11. Implementation Checklist (Ordered)", 1)
    table(
        doc,
        ["Step", "Deliverable", "Record types"],
        [
            ["1", "Lock embed model + dimension; Flyway retrieval_index_chunk + HNSW index", "All"],
            ["2", "EmbeddingService → Bedrock Titan v2; batch + content_hash idempotency", "All"],
            ["3", "FTS trigger: to_tsvector('english', chunk_text)", "All"],
            ["4", "SummaryChunker + SUMMARY_CREATED consumer", "CALL_SUMMARY, SUMMARY_ITEM, MEDICATION_TIMELINE_EVENT"],
            ["5", "HybridRetrievalService RRF merge + query embed", "Ask AI"],
            ["6", "MailPieceChunker + userId→patient map + USPS_DIGEST consumer", "USPS_MAIL, USPS_PACKAGE"],
            ["7", "Document OCR sidecar + sliding-window chunker", "DOCUMENT"],
            ["8", "ChunkEmbeddingBackfillWorker + optional SchemaPatchRunner partial index (V2607161317)", "All"],
            ["9", "Integration tests with fixed-dimension mock vectors", "TC-E-AI-*, TC-E-USPS-*"],
        ],
    )

    # 12 Related docs
    heading(doc, "12. Related Documents & Code References", 1)
    bullets(
        doc,
        [
            "docs/Hybrid_Retrieval_Scope_PostgreSQL_FullText_pgvector.docx",
            "docs/Storage_Audit_vs_Shared_Index_Schema.docx",
            "docs/Ask_AI_Upstream_Pipeline_Call_Visit_Summaries.docx",
            "docs/RBAC_Scoped_Retrieval_Source_Types.docx",
            "docs/Medication_Timeline_Retrieval_FR-AI-11.docx",
            "docs/LLM_Provider_Abstraction_Bedrock_TeamA_BAA_Fallback.docx",
            "docs/Team_E_Implementation_Task_Backlog.docx (tasks 1.5, 4.2–4.4)",
            "docs/Task_4_4_Embedding_Backfill_Scope.docx",
            "backend/core/pg_docker/docker-compose.yml — pgvector/pg15 image",
            "backend/core/pg_docker/init-scripts/01-init-database.sh — CREATE EXTENSION vector",
            "backend/core/src/main/java/com/careconnect/service/CallSummaryService.java",
            "backend/core/src/main/java/com/careconnect/service/USPSDigestService.java",
            "backend/core/src/main/java/com/careconnect/service/DocumentProcessingService.java",
            "backend/core/src/main/java/com/careconnect/service/MailpieceOcrService.java",
            "backend/core/src/main/java/com/careconnect/service/PatientContextRetrievalService.java",
            "backend/core/src/main/java/com/careconnect/model/MailPiece.java",
        ],
    )

    heading(doc, "13. Conclusion", 1)
    para(
        doc,
        "A single pgvector-backed retrieval_index_chunk table with Bedrock Titan embeddings "
        "supports summaries, mail pieces, and document chunks under one hybrid search path. "
        "Chunking templates—not embedding models—should differ by record type. Summaries need "
        "multi-level chunks (summary + items + derived medication events); mail needs per-piece "
        "explode from digest JSON with enriched OCR text; documents need sliding-window splits "
        "after durable text extraction.",
    )
    para(
        doc,
        "Immediate actions: (1) lock embedding model dimension, (2) add Flyway migration + "
        "EmbeddingService, (3) implement SummaryChunker on unified summary_json, (4) wire RRF "
        "hybrid retrieval before Ask AI gateway goes patient-facing.",
        highlight=True,
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    try:
        doc.save(OUTPUT)
        print(f"Created: {OUTPUT}")
    except PermissionError:
        alt = OUTPUT.with_name(f"{OUTPUT.stem}_refresh{OUTPUT.suffix}")
        doc.save(alt)
        print(f"Created (original locked): {alt}")


if __name__ == "__main__":
    build()
