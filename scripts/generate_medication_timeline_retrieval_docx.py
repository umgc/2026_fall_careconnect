"""Generate Word document: Medication initiation/termination timeline retrieval (FR-AI-11)."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx_codebase_revision import REVISION_BULLETS, REVISION_LABEL, save_document

ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "Medication_Timeline_Retrieval_FR-AI-11.docx"

CALL_SAMPLE = Path(r"C:\Users\ravic\Downloads\sample_summary_response_updated.json")
VISIT_SAMPLE = Path(r"C:\Users\ravic\Downloads\sample_visit_summary_response.json")


def heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def para(doc: Document, text: str, bold: bool = False, highlight: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    if highlight:
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def bullets(doc: Document, items: list[str], highlight_indices: set[int] | None = None) -> None:
    highlight_indices = highlight_indices or set()
    for i, item in enumerate(items):
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        if i in highlight_indices:
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def code(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"


def table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    highlight_cells: set[tuple[int, int]] | None = None,
) -> None:
    highlight_cells = highlight_cells or set()
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for r_idx, row in enumerate(rows, start=1):
        cells = t.add_row().cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = val
            if (r_idx, c_idx) in highlight_cells:
                for p in cells[c_idx].paragraphs:
                    for run in p.runs:
                        run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def build() -> None:
    doc = Document()

    title = doc.add_heading(
        "Medication Initiation / Termination Timeline Retrieval Requirements", 0
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para(doc, "CareConnect Team E — FR-AI-11 Research & Gap Analysis")
    para(
        doc,
        "Synthesizes SRS v2.0, Milestone 2 TDD, Test Plan, unified summary contract samples, "
        "hybrid retrieval scope, RBAC research, and codebase review.",
    )
    para(
        doc,
        "Yellow highlights mark critical gaps, schema misalignments, and recommended changes.",
        highlight=True,
    )
    para(doc, REVISION_LABEL, highlight=True)
    bullets(doc, REVISION_BULLETS, highlight_indices={0, 3, 4})
    doc.add_paragraph()

    # 1 Executive Summary
    heading(doc, "1. Executive Summary", 1)
    para(
        doc,
        "FR-AI-11 requires Ask AI to answer medication, treatment, and care-suggestion questions "
        "with a comprehensive cited timeline of initiation, change, pause, and termination events — "
        "not a single most-recent mention. Team E specifies this as a hybrid-retrieval capability "
        "combining structured SQL date filters, full-text medication-name matching, and semantic search.",
    )
    para(
        doc,
        "GAP: No medication timeline retrieval exists in code. Summary generation produces a reduced "
        "JSON shape without careInstructions.status/effectiveDate. patient_medication has start/end "
        "dates but Ask AI context shows only active med names. retrieval_index_chunk and derived "
        "MEDICATION_TIMELINE_EVENT index are not implemented.",
        highlight=True,
    )

    # 2 Requirement sources
    heading(doc, "2. Requirement Sources", 1)
    table(
        doc,
        ["Source", "Reference", "Relevance"],
        [
            ["SRS v2.0 §3", "FR-AI-11 (High priority)", "Primary functional requirement"],
            ["SRS v2.0 §3", "UC-AI-3 View Medication / Treatment Timeline", "Use case + Fig 3.8–3.9"],
            ["SRS v2.0 §3", "Key capability bullet", "Initiation + termination events, not single mention"],
            ["SRS v2.0 §3.5", "Acceptance highlights", "Cite transcript/summary; disclaimer; no-records path"],
            ["SRS v2.0 Fig 3.3", "Ask AI sequence diagram", "Answer includes medication initiation/termination history"],
            ["Milestone 2 TDD", "Three-layer search model", "Structured SQL best for timelines and dates"],
            ["Milestone 2 TDD", "Ask AI example query", "What did the nurse say about my blood-pressure medication?"],
            ["Hybrid scope doc", "Medication timeline derived index", "record_type + event_type (start/stop/change)"],
            ["Ask AI upstream doc", "careInstructions[] contract", "status + effectiveDate on typed items"],
            ["Test Plan", "TC-E-AI-011 area", "Medication initiation + termination timeline retrieval"],
            ["Test Plan", "TC-E-SC-001", "Medication-change requests → human review hold"],
        ],
    )

    # 3 FR-AI-11 verbatim
    heading(doc, "3. FR-AI-11 — Requirement Text", 1)
    para(
        doc,
        "When a user asks about a medication, treatment, or care suggestion, the system shall search "
        "the patient's records for both initiation and termination events (started, changed, paused, "
        "or stopped) and present a comprehensive, cited timeline rather than a single isolated mention, "
        "surfacing any discontinuation or change prominently.",
        bold=True,
    )
    heading(doc, "3.1 Related SRS capabilities", 2)
    bullets(
        doc,
        [
            "Answers only from authorized, indexed records (FR-AI-1) with citations (FR-AI-2).",
            "Frames responses as records-based, not medical advice; standard disclaimer (FR-AI-3).",
            "Surfaces uncertainty when source material is ambiguous (FR-AI-6 / needsConfirmation).",
            "Voice queries routed through same retrieval pipeline as text (SRS §3.2).",
            "High-risk medication content may trigger confirm-with-provider / HITL (FR-AI-5, REQ-SC-4).",
            "Audit every query, citations, escalation, delivery status (FR-AI-10).",
            "STML cross-session memory may link recurring medication changes across summaries.",
        ],
    )

    # 4 Use case
    heading(doc, "4. Use Case: View Medication / Treatment Timeline (UC-AI-3)", 1)
    para(
        doc,
        "Actors: Patient, Caregiver (with consent), Family Member (read-only where permitted), Admin.",
    )
    heading(doc, "4.1 Primary user flow (target)", 2)
    bullets(
        doc,
        [
            "User asks natural-language question about a specific medication or treatment (text or voice).",
            "AI Gateway applies RBAC + RetrievalScopeService → permitted patient_id and source types.",
            "Retrieval Service detects timeline intent (medication name / treatment entity extraction).",
            "Structured prefilter: all MEDICATION_TIMELINE_EVENT or SUMMARY_CARE_INSTRUCTION rows for entity, ordered by effectiveDate.",
            "FTS + vector search supplement: transcript mentions, SOAP plan strings, clinical notes.",
            "Merge, deduplicate, rank; assemble chronological timeline with citations per event.",
            "Response highlights discontinuations and dose changes prominently at top of timeline section.",
            "Apply safety validation; attach disclaimer; log audit entry.",
        ],
    )
    heading(doc, "4.2 Alternate / error flows", 2)
    bullets(
        doc,
        [
            "No matching records → clear no-records message; offer contact-provider option (SRS §3.5).",
            "Ambiguous medication name (e.g., 'blood pressure pill') → ask clarifying question or show grouped candidates.",
            "needsConfirmation=true on source items → surface uncertainty in answer (FR-AI-6).",
            "Tier-2 safety trigger (medication-change content) → hold for human review before delivery (TC-E-SC-001).",
            "Caregiver without consent → exclude on_consent summary chunks (RBAC gap today).",
        ],
        highlight_indices={4},
    )

    # 5 Example queries
    heading(doc, "5. Example User Queries & Expected Behavior", 1)
    table(
        doc,
        ["Query", "Expected retrieval behavior", "Sample answer shape"],
        [
            [
                "When did I start lisinopril?",
                "Filter careInstructions type=medication, status=started, entity~lisinopril; sort by effectiveDate",
                "Started 2026-06-28 (20 mg daily) — cited call summary careInstruction itemId e5d7…",
            ],
            [
                "Was my aspirin stopped?",
                "Find status=discontinued events for aspirin across all episodes",
                "Yes — discontinued 2026-06-27 per cardiology guidance — citation prominent",
            ],
            [
                "What changed with my water pill?",
                "Entity resolve furosemide; collect started/changed events from visit + call summaries",
                "Increased 20→40 mg for 3 days starting 2026-06-28 — visit summary citation",
            ],
            [
                "History of my blood pressure medications",
                "Multi-entity timeline: all medication-type careInstructions + patient_medication rows",
                "Chronological list with start/stop/change events and episode citations",
            ],
            [
                "What did the nurse say about my pills?",
                "Semantic search over TRANSCRIPT + SUMMARY chunks; timeline assembly if structured events found",
                "Narrative + structured timeline when careInstructions present",
            ],
        ],
    )

    # 6 Data model
    heading(doc, "6. Data Sources for Timeline Events", 1)
    heading(doc, "6.1 Primary — summary careInstructions[] (unified contract)", 2)
    para(doc, "Canonical structured source per Prof directive and sample JSON files.")
    table(
        doc,
        ["Field", "Purpose for FR-AI-11"],
        [
            ["type", "Filter type=medication (exclude procedure/instruction-only rows unless treatment-related)"],
            ["text", "Human-readable event description; FTS indexing; entity extraction input"],
            ["status", "Event type: started | discontinued | changed | paused | active"],
            ["effectiveDate", "Timeline sort key (ISO date); required for initiation/termination queries"],
            ["itemId", "Stable citation ID linking to summary envelope + sourceTurnId"],
            ["sourceTurnId", "Deep link to transcript turn for highlight"],
            ["needsConfirmation", "Propagate uncertainty to Ask AI answer (FR-AI-6)"],
            ["confidence", "Rank tie-breaker; suppress low-confidence in primary timeline"],
        ],
        highlight_cells={(3, 0), (3, 1), (4, 0), (4, 1)},
    )
    heading(doc, "6.2 Sample call summary events (lisinopril + aspirin)", 2)
    para(doc, f"Source: {CALL_SAMPLE}")
    table(
        doc,
        ["Medication (parsed)", "status", "effectiveDate", "Episode"],
        [
            ["Lisinopril 10→20 mg daily", "started", "2026-06-28", "call_2026_06_27_001"],
            ["Aspirin 81 mg daily", "discontinued", "2026-06-27", "call_2026_06_27_001"],
        ],
    )
    heading(doc, "6.3 Sample visit summary events (furosemide)", 2)
    para(doc, f"Source: {VISIT_SAMPLE}")
    table(
        doc,
        ["Treatment (parsed)", "status", "effectiveDate", "Episode"],
        [
            ["Furosemide 20→40 mg × 3 days", "started", "2026-06-28", "visit_2026_06_27_004"],
        ],
    )
    para(
        doc,
        "Visit and call summaries share the same summary_json inner schema — one indexer and one "
        "retrieval path over call_summaries + visit_summaries (planned).",
    )

    heading(doc, "6.4 Secondary sources", 2)
    table(
        doc,
        ["Source", "Timeline fields", "Role in FR-AI-11"],
        [
            [
                "summary.conditions[]",
                "status, effectiveDate, name",
                "Treatment/condition lifecycle; not always medication-specific",
            ],
            [
                "summary.soap.plan.medications[]",
                "Unstructured strings",
                "FTS fallback; lacks status/effectiveDate — derive events via NLP or LLM at index time",
            ],
            [
                "summary.clinicalObservations.medicationRelated[]",
                "Strings",
                "Context only; not structured timeline",
            ],
            [
                "patient_medication table",
                "start_date, end_date, is_active (VARCHAR dates)",
                "Platform med list; approval workflow; snapshot not event history",
            ],
            [
                "call_transcript_segments",
                "transcript_text, occurred_at",
                "Unstructured mentions; cite via sourceTurnId cross-ref",
            ],
            [
                "clinical_notes / notetaker",
                "content, created_at",
                "Supplementary FTS/vector hits",
            ],
            [
                "uploaded documents (Rx, care plans)",
                "OCR text",
                "External prescriber changes if indexed",
            ],
        ],
    )

    # 7 Derived index
    heading(doc, "7. Recommended Derived Index: MEDICATION_TIMELINE_EVENT", 1)
    para(
        doc,
        "Team E hybrid scope defines a derived index row type for medication timeline events. "
        "Populate at summary index time by exploding careInstructions where type=medication.",
    )
    code(
        doc,
        "-- Derived at index time from careInstructions[] + optional NLP on SOAP/transcripts\n"
        "record_type: MEDICATION_TIMELINE_EVENT  (or SUMMARY_CARE_INSTRUCTION subtype)\n"
        "chunk_metadata: {\n"
        "  medication_name_normalized: 'lisinopril',\n"
        "  event_type: 'started' | 'discontinued' | 'changed' | 'paused',\n"
        "  effective_date: '2026-06-28',\n"
        "  dose_from: '10 mg', dose_to: '20 mg',\n"
        "  episode_id: 'call_2026_06_27_001',\n"
        "  episode_type: 'call' | 'visit',\n"
        "  source_item_id: 'e5d7f1a4-...',\n"
        "  source_turn_id: 't_007',\n"
        "  needs_confirmation: true,\n"
        "  confidence: 0.97\n"
        "}",
    )
    heading(doc, "7.1 Index row requirements", 2)
    bullets(
        doc,
        [
            "patient_id — RBAC row-level filter (mandatory)",
            "effective_date — B-tree index for chronological range queries",
            "medication_name_normalized — structured prefilter + FTS",
            "event_type — filter started/stopped/changed",
            "source_record_id — episodeId + itemId for citation deep links",
            "chunk_text — human-readable event for embedding + display excerpt",
            "consent_scope / caregiver_visibility — exclude unauthorized caregiver retrieval",
        ],
        highlight_indices={5},
    )

    # 8 Retrieval algorithm
    heading(doc, "8. Three-Layer Retrieval Strategy for Timeline Queries", 1)
    table(
        doc,
        ["Layer", "When used", "Timeline-specific query"],
        [
            [
                "Structured SQL",
                "Primary for FR-AI-11",
                "SELECT * FROM retrieval_index_chunk WHERE patient_id=? AND record_type IN (...) "
                "AND metadata->>'medication_name_normalized' ILIKE ? ORDER BY metadata->>'effective_date'",
            ],
            [
                "Full-text (FTS)",
                "Medication name variants, brand/generic",
                "plainto_tsquery('lisinopril | prinivil') over chunk_text",
            ],
            [
                "Semantic (pgvector)",
                "Paraphrased queries",
                "'my water pill', 'heart failure diuretic' → furosemide events",
            ],
        ],
    )
    heading(doc, "8.1 Timeline assembly rules", 2)
    bullets(
        doc,
        [
            "Deduplicate events with same itemId; prefer structured careInstruction over SOAP string duplicates.",
            "Sort ascending by effectiveDate; unknown dates sort last with needsConfirmation flag.",
            "Prominently surface most recent discontinued or dose-change event in answer lead paragraph.",
            "Each timeline bullet must include citation: { type, episodeId, itemId, excerpt, sourceTurnId }.",
            "Merge patient_medication active snapshot as 'current state' footer if dates align.",
            "Cross-episode linking supports STML 'recurring medication changes' memory.",
        ],
        highlight_indices={2, 3},
    )

    # 9 Entity resolution
    heading(doc, "9. Medication Entity Resolution", 1)
    bullets(
        doc,
        [
            "Normalize generic names (lisinopril, furosemide, aspirin) at index time via RxNorm or rule list.",
            "Map brand names (Prinivil → lisinopril) in metadata.medication_aliases[].",
            "Colloquial terms ('water pill' → furosemide) handled by vector search + alias table.",
            "Dose-change events: parse '10 mg to 20 mg' into dose_from/dose_to metadata for change detection.",
            "GAP: No medication normalization service exists in codebase today.",
        ],
        highlight_indices={4},
    )

    # 10 Codebase gap analysis
    heading(doc, "10. Codebase Gap Analysis", 1)
    heading(doc, "10.1 Summary generation vs target contract", 2)
    para(doc, "BedrockSentimentService.summarizeTranscript() prompt returns ONLY:")
    bullets(
        doc,
        [
            "headline, overallAssessment, keyConcerns[], recommendedActions[], followUpQuestions[]",
        ],
    )
    para(
        doc,
        "MISSING from live Bedrock output: careInstructions[], conditions[], SOAP, actionItems, "
        "appointments, status, effectiveDate, itemId, sourceTurnId. Sample JSON files represent "
        "TARGET contract — not current production output.",
        highlight=True,
    )
    para(
        doc,
        "parseSummaryResponse() explicitly strips any extended fields even if model returns them.",
        highlight=True,
    )

    heading(doc, "10.2 Platform medication table", 2)
    table(
        doc,
        ["Aspect", "patient_medication today", "FR-AI-11 need"],
        [
            ["start_date / end_date", "VARCHAR(50), optional", "Timeline bounds; weak typing"],
            ["is_active", "Boolean soft-delete", "Current state only — no change history"],
            ["approval_status", "PENDING / APPROVED / REMOVAL_PENDING", "Workflow — not indexed for Ask AI"],
            ["Event history", "Single row per med — updates overwrite", "Multiple initiation/termination events across episodes"],
            ["Ask AI context", "MedicalContextService lists active names only — no dates", "Insufficient for timeline answers"],
        ],
        highlight_cells={(4, 1), (4, 2)},
    )

    heading(doc, "10.3 Retrieval infrastructure", 2)
    bullets(
        doc,
        [
            "retrieval_index_chunk table — not present",
            "PatientContextRetrievalService — in-memory substring stub; no date/entity filtering",
            "BedrockAIChatService — raw message only when careconnect.ai.enabled=true; no timeline assembly",
            "MedicalContextService — lists active medication names/doses but not wired into Bedrock chat path",
            "CallSummaryService — persists summary_json but without structured med events today",
            "visit_summaries table — planned; same contract as call_summaries",
            "No POST /api/ai/ask gateway or timeline-specific endpoint",
        ],
        highlight_indices={0, 1, 2, 3, 6},
    )

    heading(doc, "10.4 RBAC & consent", 2)
    bullets(
        doc,
        [
            "Timeline retrieval must honor RetrievalScopeService (see RBAC doc).",
            "caregiverVisibility: on_consent on sample summaries — not enforced in code.",
            "USE_AI_FEATURES missing for PATIENT role — blocks primary Ask AI users.",
            "Medication timeline answers for caregivers require active link + consent gate.",
        ],
        highlight_indices={1, 2},
    )

    # 11 Safety
    heading(doc, "11. Safety, Consent & Clinical Framing", 1)
    bullets(
        doc,
        [
            "Timeline answers describe what records state — not prescribing advice (SRS out-of-scope).",
            "Discontinuations surfaced prominently reduces missed stop orders (FR-AI-11 intent).",
            "needsConfirmation items: prefix with 'Your records suggest…' not definitive clinical statements.",
            "Medication-change tier-2 content → HITL hold before care recipient sees answer (TC-E-SC-001).",
            "escalation: confirm-with-provider on dosage/clinical interpretation queries.",
            "Citations mandatory — timeline without sources fails FR-AI-2 validation.",
        ],
        highlight_indices={2, 3},
    )

    # 12 Test alignment
    heading(doc, "12. Test Plan Alignment", 1)
    table(
        doc,
        ["Test area", "Scenario", "Pass criteria"],
        [
            [
                "TC-E-AI-011",
                "User asks about med with start + stop in indexed summaries",
                "Answer lists both events chronologically with citations; stop shown prominently",
            ],
            [
                "TC-E-AI-011",
                "Only start event indexed, no stop",
                "Answer states start; does not invent termination",
            ],
            [
                "TC-E-AI-001",
                "Caregiver queries patient without link",
                "403 / empty scope — no cross-patient timeline",
            ],
            [
                "TC-E-AI-002",
                "Timeline response",
                "Every event has citation with record type + ID",
            ],
            [
                "TC-E-AI-006",
                "Ambiguous source (needsConfirmation=true)",
                "Uncertainty language in response",
            ],
            [
                "TC-E-SC-001",
                "Medication-change recommendation detected",
                "Held for human review",
            ],
            [
                "NFR-AI-1",
                "Timeline query latency",
                "≤ 5 s p95 including structured + hybrid search",
            ],
            [
                "NFR-AI-3",
                "New summary with careInstructions indexed",
                "Timeline queryable within 5 minutes",
            ],
        ],
    )

    # 13 Current vs target
    heading(doc, "13. Current vs Target State", 1)
    table(
        doc,
        ["Capability", "Current codebase", "Target (Milestone 3)"],
        [
            ["Structured med events in summaries", "Not generated", "careInstructions with status + effectiveDate"],
            ["Derived timeline index", "None", "MEDICATION_TIMELINE_EVENT rows"],
            ["Timeline SQL prefilter", "None", "Ordered by effective_date metadata"],
            ["Entity normalization", "None", "RxNorm / alias map at index time"],
            ["Ask AI timeline answers", "None (AI chat API opt-in; no context/RAG)", "FR-AI-11 compliant cited timelines"],
            ["Medical context in chat", "MedicalContextService exists but orphaned from Bedrock path", "Inject scoped med list + timeline chunks pre-Bedrock"],
            ["Voice timeline queries", "No Ask AI voice path", "Same retrieval as text"],
            ["Platform med history", "Single-row snapshot", "Supplementary; episodes are authoritative for changes discussed in calls/visits"],
        ],
        highlight_cells={
            (1, 1), (2, 1), (3, 1), (4, 1), (5, 1), (6, 1), (7, 1), (8, 1),
        },
    )

    # 14 Implementation roadmap
    heading(doc, "14. Implementation Roadmap", 1)
    table(
        doc,
        ["Phase", "Deliverable", "Dependency"],
        [
            ["P0", "Upgrade Bedrock summary prompt + parser to full unified contract (careInstructions, conditions, SOAP)", "Team SUM / platform"],
            ["P0", "Validate sample JSON schema in CI; reject summaries missing status/effectiveDate on med items", "Schema tests"],
            ["P1", "Indexer: explode careInstructions type=medication → MEDICATION_TIMELINE_EVENT chunks", "retrieval_index_chunk migration"],
            ["P1", "Medication name normalization at index time", "RxNorm or curated alias list"],
            ["P1", "Timeline query planner in retrieval service (detect med intent → structured prefilter first)", "Hybrid retrieval service"],
            ["P2", "Merge patient_medication snapshot into timeline footer", "MedicationService API"],
            ["P2", "Transcript/sourceTurnId deep links in citations", "Call transcript UI"],
            ["P2", "STML cross-episode medication change linking", "STML memory layer"],
            ["P3", "Voice query → same timeline pipeline", "Voice mic on Ask AI (see voice-query doc)"],
        ],
        highlight_cells={(1, 0), (1, 1), (2, 0), (3, 0), (4, 0)},
    )

    heading(doc, "14.1 Index-time pseudocode", 2)
    code(
        doc,
        "for each summary in (call_summaries UNION visit_summaries):\n"
        "  for item in summary.careInstructions where item.type == 'medication':\n"
        "    if item.status in ('started','discontinued','changed','paused'):\n"
        "      emit chunk(\n"
        "        record_type = 'MEDICATION_TIMELINE_EVENT',\n"
        "        patient_id = resolve_patient(summary.episodeId),\n"
        "        effective_date = parse_date(item.effectiveDate),\n"
        "        metadata = { event_type: item.status, item_id: item.itemId, ... },\n"
        "        chunk_text = item.text,\n"
        "        embedding = embed(item.text)\n"
        "      )",
    )

    # 15 Related docs
    heading(doc, "15. Related Documents & Code References", 1)
    bullets(
        doc,
        [
            "docs/Hybrid_Retrieval_Scope_PostgreSQL_FullText_pgvector.docx",
            "docs/Ask_AI_Upstream_Pipeline_Call_Visit_Summaries.docx",
            "docs/RBAC_Scoped_Retrieval_Source_Types.docx",
            "docs/Voice_Query_Path_and_STT_Framework_Dependencies.docx",
            "C:\\Users\\ravic\\Downloads\\sample_summary_response_updated.json",
            "C:\\Users\\ravic\\Downloads\\sample_visit_summary_response.json",
            "C:\\Users\\ravic\\Downloads\\CareConnect_SRS_Revision 2.0_TEAM E.docx",
            "backend/core/src/main/java/com/careconnect/service/BedrockSentimentService.java",
            "backend/core/src/main/java/com/careconnect/service/CallSummaryService.java",
            "backend/core/src/main/java/com/careconnect/service/MedicalContextService.java",
            "backend/core/src/main/java/com/careconnect/model/Medication.java",
            "backend/core/src/main/java/com/careconnect/service/PatientContextRetrievalService.java",
        ],
    )

    heading(doc, "16. Conclusion", 1)
    para(
        doc,
        "FR-AI-11 is a first-class Ask AI capability requiring structured medication lifecycle events "
        "with effective dates, hybrid retrieval across episodes, and prominent presentation of "
        "discontinuations and dose changes. The unified summary careInstructions[] contract "
        "(status + effectiveDate + itemId) is the authoritative structured source; platform "
        "patient_medication supplements current-state context but cannot alone satisfy multi-event timelines.",
    )
    para(
        doc,
        "Critical path: (1) align live summary generation with target JSON contract, "
        "(2) build MEDICATION_TIMELINE_EVENT derived index, (3) implement structured-then-hybrid "
        "timeline query in retrieval service, (4) assemble cited chronological answers with safety "
        "and RBAC gates, (5) validate with TC-E-AI-011 scenarios.",
        highlight=True,
    )

    save_document(doc, OUTPUT)


if __name__ == "__main__":
    build()
