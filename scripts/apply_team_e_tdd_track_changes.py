"""
Apply Team E TDD revision packets into 01_Technical_Design_Document.docx
with Word track-change insertions (Review → All Markup).

Approach: copy the target TDD, surgically insert NEW subsections and
short EDIT companion notes as w:ins runs. Does not rewrite unrelated body text.
"""

from __future__ import annotations

import copy
import re
import shutil
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

AUTHOR = "Team E — TDD Integration"
DATE = datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")

SOURCE_TDD = Path(r"C:\Users\ravic\Downloads\01_Technical_Design_Document.docx")
TEAM_E_TDD = Path(r"C:\Users\ravic\Downloads\CareConnect_Milestone_2_TDD_TEAM E (1).docx")
PACKETS = Path(
    r"C:\Users\ravic\Downloads\TDD_Revision_Packets_TeamE_into_01_Technical_Design_Document.docx"
)
OUTPUT = Path(
    r"C:\Users\ravic\Downloads\01_Technical_Design_Document_TeamE_tracked.docx"
)
REPO_COPY = (
    Path(__file__).resolve().parents[1]
    / "docs"
    / "01_Technical_Design_Document_TeamE_tracked.docx"
)

_REV_COUNTER = 0


def next_rev_id() -> str:
    global _REV_COUNTER
    _REV_COUNTER += 1
    return str(_REV_COUNTER)


def enable_track_revisions(doc: Document) -> None:
    settings = doc.settings.element
    if settings.find(qn("w:trackRevisions")) is None:
        track = OxmlElement("w:trackRevisions")
        track.set(qn("w:val"), "true")
        settings.append(track)


def heading_level(style_name: str | None) -> int | None:
    if not style_name:
        return None
    m = re.match(r"Heading\s+(\d+)", style_name)
    return int(m.group(1)) if m else None


def paragraph_full_text(p: Paragraph) -> str:
    """Include text inside w:ins / w:del runs (python-docx .text skips these)."""
    return "".join(t.text or "" for t in p._p.iter(qn("w:t")))


def cell_full_text(cell) -> str:
    return "".join(t.text or "" for t in cell._tc.iter(qn("w:t")))


def find_heading(doc: Document, prefix: str) -> Paragraph | None:
    prefix_l = prefix.lower()
    for p in doc.paragraphs:
        style = p.style.name if p.style else ""
        if not style.startswith("Heading"):
            continue
        text = paragraph_full_text(p).strip()
        if text.lower().startswith(prefix_l):
            return p
    return None


def find_section_end(doc: Document, start: Paragraph, max_level: int) -> Paragraph | None:
    """Return the next heading at level <= max_level after start (exclusive)."""
    seen = False
    for p in doc.paragraphs:
        if p._p is start._p:
            seen = True
            continue
        if not seen:
            continue
        lvl = heading_level(p.style.name if p.style else "")
        if lvl is not None and lvl <= max_level:
            return p
    return None


def make_ins_run(text: str) -> OxmlElement:
    ins = OxmlElement("w:ins")
    ins.set(qn("w:id"), next_rev_id())
    ins.set(qn("w:author"), AUTHOR)
    ins.set(qn("w:date"), DATE)
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    ins.append(r)
    return ins


def insert_tracked_paragraph_before(
    doc: Document,
    before: Paragraph | None,
    text: str,
    style: str = "Normal",
) -> Paragraph:
    """Create a paragraph with tracked-insertion runs; insert before `before` or at end."""
    new_p = OxmlElement("w:p")
    # style
    pPr = OxmlElement("w:pPr")
    pStyle = OxmlElement("w:pStyle")
    # resolve style id from doc styles
    style_obj = doc.styles[style] if style in [s.name for s in doc.styles] else doc.styles["Normal"]
    style_id = style_obj.style_id
    pStyle.set(qn("w:val"), style_id)
    pPr.append(pStyle)
    new_p.append(pPr)

    # split long text into chunks Word handles better
    chunk = text if text else " "
    new_p.append(make_ins_run(chunk))

    body = doc.element.body
    if before is not None:
        before._p.addprevious(new_p)
    else:
        body.append(new_p)
    return Paragraph(new_p, doc)


def insert_block_before(
    doc: Document,
    before: Paragraph | None,
    heading_text: str,
    heading_style: str,
    body_paras: list[str],
) -> None:
    # Insert in reverse so order is preserved when using addprevious
    for text in reversed(body_paras):
        if text.strip():
            insert_tracked_paragraph_before(doc, before, text, "Normal")
    insert_tracked_paragraph_before(doc, before, heading_text, heading_style)


def find_first_table_after_heading(doc: Document, heading_prefix: str):
    """Return the first Table element after a heading prefix, or None."""
    from docx.table import Table

    start = find_heading(doc, heading_prefix)
    if start is None:
        return None
    seen = False
    for child in doc.element.body.iterchildren():
        if child is start._p:
            seen = True
            continue
        if not seen:
            continue
        if child.tag == qn("w:p"):
            p = Paragraph(child, doc)
            style = p.style.name if p.style else ""
            text = paragraph_full_text(p).strip()
            lvl = heading_level(style)
            start_lvl = heading_level(start.style.name if start.style else "") or 2
            if lvl is not None and lvl <= start_lvl and text and not text.lower().startswith(
                heading_prefix.lower()[:3]
            ):
                if not text.lower().startswith(heading_prefix.lower()):
                    return None
        if child.tag == qn("w:tbl"):
            return Table(child, doc)
    return None


def set_cell_tracked_text(cell, text: str) -> None:
    """Replace cell content with a single tracked-insertion paragraph."""
    # Keep one paragraph; clear other content nodes except pPr on first p
    while len(cell.paragraphs) > 1:
        p_el = cell.paragraphs[-1]._p
        p_el.getparent().remove(p_el)
    p = cell.paragraphs[0]
    for child_el in list(p._p):
        if child_el.tag != qn("w:pPr"):
            p._p.remove(child_el)
    p._p.append(make_ins_run(text if text else " "))


def append_tracked_rows(table, rows: list[list[str]]) -> int:
    """Append rows by cloning the last row XML (avoids broken tblGrid widths)."""
    added = 0
    if not table.rows:
        return 0
    last_tr = table.rows[-1]._tr
    for values in rows:
        new_tr = copy.deepcopy(last_tr)
        cells = new_tr.findall(qn("w:tc"))
        for i, val in enumerate(values):
            if i >= len(cells):
                break
            tc = cells[i]
            # clear paragraphs in cell
            for p in tc.findall(qn("w:p")):
                tc.remove(p)
            p = OxmlElement("w:p")
            p.append(make_ins_run(val if val else " "))
            tc.append(p)
        last_tr.addnext(new_tr)
        last_tr = new_tr
        added += 1
    return added


def existing_row_keys(table, key_col: int = 0) -> set[str]:
    keys: set[str] = set()
    for i, row in enumerate(table.rows):
        if i == 0:
            continue
        if key_col < len(row.cells):
            keys.add(cell_full_text(row.cells[key_col]).strip().lower())
    return keys


def append_missing_rows(
    doc: Document,
    heading_prefix: str,
    rows: list[list[str]],
    key_col: int = 0,
) -> bool:
    table = find_first_table_after_heading(doc, heading_prefix)
    if table is None:
        print(f"  ! no table after {heading_prefix}")
        return False
    existing = existing_row_keys(table, key_col)
    to_add = []
    for row in rows:
        key = (row[key_col] if row else "").strip().lower()
        if key and key not in existing:
            to_add.append(row)
    if not to_add:
        print(f"  = table after {heading_prefix} already has Team E rows")
        return True
    n = append_tracked_rows(table, to_add)
    print(f"  + TABLE +{n} row(s) after {heading_prefix}")
    return True


def build_table_element(headers: list[str], rows: list[list[str]]) -> OxmlElement:
    """Build a w:tbl with tracked insertion text in every cell."""
    tbl = OxmlElement("w:tbl")
    tblPr = OxmlElement("w:tblPr")
    tblStyle = OxmlElement("w:tblStyle")
    tblStyle.set(qn("w:val"), "TableGrid")
    tblPr.append(tblStyle)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), "0")
    tblW.set(qn("w:type"), "auto")
    tblPr.append(tblW)
    tbl.append(tblPr)

    tblGrid = OxmlElement("w:tblGrid")
    for _ in headers:
        gridCol = OxmlElement("w:gridCol")
        tblGrid.append(gridCol)
    tbl.append(tblGrid)

    def add_row(values: list[str], header: bool = False) -> None:
        tr = OxmlElement("w:tr")
        for val in values:
            tc = OxmlElement("w:tc")
            tcPr = OxmlElement("w:tcPr")
            tcW = OxmlElement("w:tcW")
            tcW.set(qn("w:w"), "0")
            tcW.set(qn("w:type"), "auto")
            tcPr.append(tcW)
            tc.append(tcPr)
            p = OxmlElement("w:p")
            if header:
                pPr = OxmlElement("w:pPr")
                rPr = OxmlElement("w:rPr")
                b = OxmlElement("w:b")
                rPr.append(b)
                # bold inside insertion
                p.append(pPr)
            p.append(make_ins_run(val))
            tc.append(p)
            tr.append(tc)
        tbl.append(tr)

    add_row(headers, header=True)
    for row in rows:
        # pad/truncate to header width
        vals = list(row) + [""] * max(0, len(headers) - len(row))
        add_row(vals[: len(headers)])
    return tbl


def insert_captioned_table_before(
    doc: Document,
    before: Paragraph | None,
    caption_lines: list[str],
    headers: list[str],
    rows: list[list[str]],
) -> None:
    """Insert caption paragraphs then a table immediately before `before` (or at body end)."""
    tbl = build_table_element(headers, rows)
    if before is None:
        for line in caption_lines:
            insert_tracked_paragraph_before(doc, None, line, "Normal")
        doc.element.body.append(tbl)
        return

    # Attach table first, then captions above the table.
    before._p.addprevious(tbl)
    for line in reversed(caption_lines):
        new_p = OxmlElement("w:p")
        pPr = OxmlElement("w:pPr")
        pStyle = OxmlElement("w:pStyle")
        pStyle.set(qn("w:val"), doc.styles["Normal"].style_id)
        pPr.append(pStyle)
        new_p.append(pPr)
        new_p.append(make_ins_run(line))
        tbl.addprevious(new_p)


def insert_table_at_section_end(
    doc: Document,
    heading_prefix: str,
    section_level: int,
    caption_lines: list[str],
    headers: list[str],
    rows: list[list[str]],
) -> bool:
    start = find_heading(doc, heading_prefix)
    if start is None:
        print(f"  ! missing section for table: {heading_prefix}")
        return False
    end = find_section_end(doc, start, section_level)
    insert_captioned_table_before(doc, end, caption_lines, headers, rows)
    print(f"  + TABLE under {heading_prefix}")
    return True


def apply_table_updates(doc: Document) -> None:
    print("TABLE updates...")

    # --- Append rows into existing shared tables ---
    append_missing_rows(
        doc,
        "2.3 Component Responsibilities",
        [
            [
                "AI Gateway",
                "Authentication, RBAC scope filtering, rate limiting, and Ask AI request routing.",
                "Team E (Arturo Santana / Ravichandra Vasireddy)",
                "Cognito/JWT, feature services",
            ],
            [
                "Retrieval Service (Ask AI)",
                "Scoped hybrid search and grounded, cited answer assembly.",
                "Team E (Ravichandra Vasireddy)",
                "Data layer, Bedrock, safety framework",
            ],
            [
                "Summaries Service",
                "Generate structured, cited summaries from diarized transcript text.",
                "Team E (Fon Ade Asaa)",
                "Transcript source, Bedrock, data layer",
            ],
            [
                "STML Service",
                "Daily Memory Brief and recall experiences built on summaries and retrieval.",
                "Team E (John Bui)",
                "Summaries, Retrieval",
            ],
            [
                "USPS Mail Agent",
                "Ingest, OCR, normalize, index, and classify Informed Delivery mail.",
                "Team E (Arturo Santana)",
                "Gmail OAuth, Textract, data layer",
            ],
            [
                "Safety/Consent Layer",
                "Disclaimers, citations, secondary validation, and HITL escalation.",
                "Team E (David Oguh)",
                "All feature services, audit ledger",
            ],
        ],
    )

    append_missing_rows(
        doc,
        "2.6 System Users",
        [
            [
                "Human Reviewer / Clinician",
                "Reviews high-risk (Tier 2) AI output held for HITL.",
                "Requires a review/release queue and clear escalation ownership.",
            ],
            [
                "System Administrator",
                "Configures roles, data-source indexing, retention, and AI access settings.",
                "Requires externalized configuration and auditable admin actions.",
            ],
            [
                "Developer / Tester",
                "Builds and validates Ask AI, summaries, USPS, STML, and safety work.",
                "Needs clear APIs, data models, and testable behavior contracts.",
            ],
        ],
    )

    append_missing_rows(
        doc,
        "3.1 Technology Decisions",
        [
            [
                "Access the LLM through an abstraction layer.",
                "Provider choice and BAA access remain environment-dependent.",
                "Provider can change without reworking Ask AI or summary clients.",
            ],
            [
                "Records-grounded RAG with schema-validated output.",
                "Prevents unsupported medical content and enables provenance.",
                "Every answer is cited or flagged; output stays records-bounded.",
            ],
            [
                "Hybrid retrieval (SQL + full-text + pgvector).",
                "Supports both precise and conversational questions.",
                "Higher recall; semantic matches remain likelihood-ranked, not exact.",
            ],
            [
                "Metadata-first USPS extraction, OCR as fallback.",
                "USPS has no public Informed Delivery API for full image capture.",
                "Mail records populate from metadata even when images are absent.",
            ],
        ],
    )

    append_missing_rows(
        doc,
        "3.2 Technology Risks",
        [
            [
                "Bedrock / LLM access delayed.",
                "Blocks summary generation and Ask AI answers.",
                "Build against a swappable LLM interface; confirm account access early.",
            ],
            [
                "LLM produces inaccurate or hallucinated content.",
                "Unsafe or misleading medical-adjacent output.",
                "Require citations; flag low confidence; apply secondary validation / HITL.",
            ],
            [
                "OCR accuracy insufficient for envelope images.",
                "Missed or wrong mail extraction.",
                "Confidence scoring; route low-confidence pieces to manual review.",
            ],
            [
                "Index refresh lag after transcript or summary save.",
                "Stale Ask AI answers.",
                "Async IndexWorker with lease covering embedding latency.",
            ],
        ],
    )

    append_missing_rows(
        doc,
        "5.5 Transcription and Summary",
        [
            [
                "POST",
                "/api/summaries/{id}/items/{itemId}/confirm",
                "Yes",
                "Confirm or decline a structured summary actionable item after human review.",
            ],
        ],
        key_col=1,
    )

    append_missing_rows(
        doc,
        "6.5.1 Dependency Register",
        [
            [
                "Managed LLM access (Bedrock) for Ask AI embeddings and answers",
                "Team A / Cloud; Team E consumers",
                "Provision endpoint/BAA; keep Titan embed v1 on the task-role allow-list",
                "In progress",
            ],
            [
                "Shared retrieval / summary schema contract",
                "Team E (Ravichandra Vasireddy) + Summaries/STML",
                "Lock versioned fields for transcripts, summaries, embeddings, mail",
                "In progress",
            ],
            [
                "Caregiver visibility policy for AI outputs",
                "Team E Safety (David Oguh) + RBAC owners",
                "Finalize visibility matrix before caregiver-facing Ask AI demo",
                "Open",
            ],
            [
                "Gmail OAuth for USPS Informed Delivery ingest",
                "Team E (Arturo Santana) / Cloud",
                "Complete OAuth setup; replace mock mailbox data",
                "Open",
            ],
        ],
    )

    append_missing_rows(
        doc,
        "7.8 Deployment Environments",
        [
            [
                "Test (Team E AI)",
                "AWS validation for Ask AI, summaries, and mail before class-wide integration",
                "Team E / Planned",
            ],
            [
                "Final Demo (Team E AI)",
                "Milestone demonstration environment for grounded retrieval and mail flows",
                "Team E / Planned",
            ],
        ],
    )

    append_missing_rows(
        doc,
        "11.7 Accessibility Testing",
        [
            [
                "Ask AI answer + citation list",
                "Focus order, screen-reader labels for citations, non-medical disclaimer visibility",
                "Team E Ask AI",
            ],
            [
                "Summary card confirmation",
                "Keyboard-complete confirm/decline; contrast for AI notices",
                "Team E Summaries / STML",
            ],
            [
                "Mail view tiers",
                "Readable importance tier labels; empty and low-OCR states",
                "Team E USPS",
            ],
        ],
    )

    append_missing_rows(
        doc,
        "13.2 Screens and User Workflows Affected",
        [
            [
                "Ask AI",
                "Records-grounded answer with persistent non-medical disclaimer and citations",
                "Care recipient, caregiver",
                "Retrieval Service, Safety layer",
            ],
            [
                "Summary card",
                "AI-generated notice, headline + recap, per-item confirm path",
                "Care recipient, caregiver",
                "Summaries Service, Safety layer",
            ],
            [
                "Mail view",
                "List pieces with sender, date, importance tier, and OCR/empty states",
                "Care recipient, caregiver",
                "USPS Mail Agent",
            ],
            [
                "Daily Memory Brief",
                "Recent prioritized information on app open (STML-first)",
                "Care recipient",
                "STML Service",
            ],
        ],
    )

    # --- New tables under ADD sections (match neighboring feature/entity formats) ---
    insert_table_at_section_end(
        doc,
        "4.9 Mail / USPS Entities (MailPiece)",
        2,
        [
            "Table 4.9a",
            "Mail / USPS entities (Team E)",
        ],
        ["Entity", "Primary Key", "Key Attributes", "Relationships", "Notes"],
        [
            [
                "MailPiece",
                "mail_id (UUID)",
                "patient_id, sender, delivery_date, tier, ocr_confidence, visible_text, image refs",
                "Patient; optional retrieval_index_chunk linkage",
                "PHI-adjacent; caregiver visibility consent-scoped; metadata-first OCR fallback",
            ],
            [
                "MailPiece field dictionary",
                "mail_id",
                "sender (string), delivery_date (date), tier (enum), ocr_confidence (0–1)",
                "USPS Mail Agent, Ask AI indexing",
                "Below OCR threshold → manual review; PK/scoped patient queries",
            ],
        ],
    )

    insert_table_at_section_end(
        doc,
        "5.11 Ask AI Endpoint (POST /api/ai/ask)",
        2,
        [
            "Table 5.11",
            "Ask AI endpoint summary",
        ],
        ["Method", "Path", "Auth Required", "Description"],
        [
            [
                "POST",
                "/api/ai/ask",
                "Yes",
                "Submit an Ask AI question (text or transcribed voice) and receive a grounded answer with citations.",
            ],
        ],
    )
    insert_table_at_section_end(
        doc,
        "5.11 Ask AI Endpoint (POST /api/ai/ask)",
        2,
        [
            "Table 5.12",
            "POST /api/ai/ask parameter reference",
        ],
        ["Parameter", "Location", "Type", "Required?", "Description"],
        [
            ["query", "body", "string", "Yes", "User question (text or transcribed voice)."],
            ["sessionId", "body", "string", "No", "Conversation session for audit correlation."],
            ["patientId", "body", "UUID/long", "Yes*", "Target patient scope (*may be implied by path or token grants)."],
            ["Authorization", "header", "JWT", "Yes", "Cognito-issued bearer token."],
        ],
    )

    feature_tables = [
        (
            "6.12 AI-Assisted Retrieval (Ask AI)",
            "Table 6.12",
            "Ask AI feature specification (Team E)",
            [
                ["Requirement IDs or Names", "FR-AI-1 through FR-AI-10; governed by REQ-SC-1–9"],
                ["Feature Summary", "Lets authorized users ask natural-language questions over scoped care records and receive grounded, cited answers."],
                ["Primary User Flow", "1. User submits a question. 2. System authenticates and authorizes scope. 3. Hybrid retrieval runs. 4. Model assembles a cited answer. 5. Safety rules apply. 6. Client displays answer and sources."],
                ["Main System Behaviors", "Answers only from authorized indexed records; never invents unsupported clinical advice; audit logs asks and 403s."],
                ["Inputs", "Text question or voice-derived text; patient scope; optional source filters."],
                ["Outputs", "Grounded answer; source citations; disclaimer; optional hold/escalation state."],
                ["Alternate / Error Flows", "No relevant records → clear empty-state message. Out-of-scope patient → 403. Bedrock unavailable → graceful degradation."],
            ],
        ),
        (
            "6.13 USPS Informed Delivery Mail Agent",
            "Table 6.13",
            "USPS Informed Delivery mail agent (Team E)",
            [
                ["Requirement IDs or Names", "FR-USPS-1 through FR-USPS-7"],
                ["Feature Summary", "Retrieves, extracts, indexes, classifies, and presents Informed Delivery mail for the patient."],
                ["Primary User Flow", "1. User opts in and authenticates mailbox access. 2. Agent ingests digest/metadata. 3. OCR runs when needed. 4. MailPiece is stored and listed. 5. Optional Ask AI indexing."],
                ["Main System Behaviors", "Metadata-first extraction; OCR/Textract fallback; consent-aware caregiver visibility."],
                ["Inputs", "USPS notification emails; optional mail images; OAuth tokens."],
                ["Outputs", "Normalized MailPiece (sender, delivery date, visible text, tier, OCR status)."],
                ["Alternate / Error Flows", "No mail image → recorded empty-image state. Low OCR confidence → manual review."],
            ],
        ),
        (
            "6.14 Short-Term Memory (STML) Support",
            "Table 6.14",
            "STML support (Team E)",
            [
                ["Requirement IDs or Names", "STML-1 through STML-5; builds on summary and Ask AI outputs"],
                ["Feature Summary", "Memory-recall experiences such as a Daily Memory Brief and short confirmation paths."],
                ["Primary User Flow", "1. On app open, user sees Daily Memory Brief. 2. User may confirm actionable items. 3. Brief updates from recent summaries/retrieval."],
                ["Main System Behaviors", "Consumes summary and retrieval outputs; prefers short, high-signal copy."],
                ["Inputs", "Recent summaries and indexed records; recall queries."],
                ["Outputs", "Daily Memory Brief; plain-language recall answers; confirmation states."],
                ["Alternate / Error Flows", "No recent activity → clear empty-state brief."],
            ],
        ),
        (
            "6.15 Safety, Consent & Clarity",
            "Table 6.15",
            "Safety, consent, and clarity (Team E)",
            [
                ["Requirement IDs or Names", "REQ-SC-1 through REQ-SC-9; SAF-1+"],
                ["Feature Summary", "Cross-cutting governance for AI output: disclaimers, citations, consent, and HITL hold/release."],
                ["Primary User Flow", "1. AI output is generated. 2. Secondary validation runs. 3. Safe output shows with disclaimer/citations, or Tier-2 hold routes to review."],
                ["Main System Behaviors", "Persistent non-dismissible AI notices; default-restricted caregiver visibility; audit of held releases."],
                ["Inputs", "All AI output; user role and consent; configured safety thresholds."],
                ["Outputs", "Validated output with disclaimers and citations; escalation tickets when held."],
                ["Alternate / Error Flows", "Tier-2 medication/emergency cues → withhold and escalate rather than silent failure."],
            ],
        ),
    ]
    for heading, table_no, title, pairs in feature_tables:
        insert_table_at_section_end(
            doc,
            heading,
            2,
            [table_no, title],
            ["Item", "Team Response"],
            pairs,
        )

    insert_table_at_section_end(
        doc,
        "9.4 Bottleneck Mitigation",
        2,
        [
            "Table 9.4",
            "Team E bottleneck mitigation for Ask AI indexing and retrieval",
        ],
        ["Potential Bottleneck", "Impact", "Mitigation Strategy", "Implementation"],
        [
            [
                "LLM inference latency",
                "Slow Ask AI answers",
                "Minimize context; cache embeddings; stream where practical",
                "Gateway minimization; pgvector prefilter",
            ],
            [
                "OCR throughput",
                "Slow mail processing",
                "Metadata-first; async OCR with progress states",
                "Textract fallback only when needed",
            ],
            [
                "Index refresh lag",
                "Stale retrieval",
                "Async indexing within the freshness target",
                "Outbox IndexWorker + claim lease covering embedding",
            ],
            [
                "Cloud cost spikes",
                "Budget risk",
                "Weekly billing review; lifecycle rules; scale-to-zero where possible",
                "AWS budgets/alerts",
            ],
        ],
    )

    insert_table_at_section_end(
        doc,
        "12.9.1 Technical Limitations",
        3,
        [
            "Table 12.9.1",
            "Team E technical limitations",
        ],
        ["ID", "Limitation", "Impact", "Mitigation or Next Step"],
        [
            ["TE-L1", "Output quality depends on transcript quality", "Degraded summaries and Ask AI recall", "Show confidence; allow correction; synthetic fixtures in test"],
            ["TE-L2", "LLMs can produce unsupported content", "Hallucinated medical-adjacent output risk", "Citations, uncertainty flags, secondary validation / HITL"],
            ["TE-L3", "Semantic search returns likely, not exact, matches", "Some relevant items ranked low or missed", "Hybrid FTS + vector retrieval; records-based framing"],
            ["TE-L4", "No public USPS Informed Delivery full-image API", "Incomplete mail capture", "Metadata-first extraction; treat missing image as known state"],
        ],
    )
    insert_table_at_section_end(
        doc,
        "12.9.2 Known Issues",
        3,
        [
            "Table 12.9.2",
            "Team E known issues",
        ],
        ["ID", "Issue", "Impact", "Owner", "Resolution Plan"],
        [
            ["TE-K1", "Bedrock/OAuth not fully provisioned in all envs", "Pipelines may run on mocks", "Team A / Team E", "Confirm access; swap real providers in target env"],
            ["TE-K2", "Caregiver visibility policy not finalized", "Blocks caregiver-facing demo", "Team E Safety", "Lock visibility matrix before caregiver release"],
            ["TE-K3", "App-only redeploy can reset CORS allow-list", "Amplify reports backend unhealthy", "Platform / Team E", "Re-apply Amplify origin after app-only deploys"],
        ],
    )
    insert_table_at_section_end(
        doc,
        "12.9.4 Future Enhancements",
        3,
        [
            "Table 12.9.4",
            "Team E future enhancements",
        ],
        ["Enhancement", "Description", "Priority"],
        [
            ["Voice-Accessible Recall", "Ask follow-up questions by voice and receive spoken/short answers", "High"],
            ["Cross-Session Memory", "Link related items across summaries", "High"],
            ["Full OCR Pipeline", "Complete mail-image extraction beyond metadata-first path", "High"],
            ["Medication Timeline in Retrieval", "Aggregate medication events into Ask AI context", "High"],
            ["Caregiver Digest", "Periodic consent-governed digest of recent sessions", "Medium"],
            ["Multilingual Summaries", "Summary generation beyond English", "Low"],
        ],
    )

    insert_table_at_section_end(
        doc,
        "13.4 Display Language and Localization Impact",
        2,
        [
            "Table 13.4",
            "Team E display text and localization notes",
        ],
        ["Screen / Area", "UI Text or Text Key", "Needs Translation?", "Notes"],
        [
            ["Ask AI / Summary", "ai.disclaimer.notMedicalAdvice", "No (English at launch)", "Persistent, non-dismissible notice"],
            ["Summary care item", "ai.disclaimer.medical", "No", "Shown before care-item confirm"],
            ["Mail view", "mail.tier.critical/high/medium/low", "No", "Importance tier labels"],
            ["Confirmation flow", "action.confirm.once/session/decline", "No", "Multilingual summaries are a future enhancement"],
        ],
    )


def append_edit_note_after_heading(
    doc: Document,
    heading_prefix: str,
    note_paras: list[str],
) -> bool:
    """Insert native-style Normal paragraphs immediately under a heading."""
    start = find_heading(doc, heading_prefix)
    if start is None:
        print(f"  ! missing heading for EDIT: {heading_prefix}")
        return False
    anchor = start._p.getnext()
    before = None
    if anchor is not None and anchor.tag == qn("w:p"):
        before = Paragraph(anchor, doc)
    elif anchor is not None:
        for text in reversed(note_paras):
            new_p = OxmlElement("w:p")
            pPr = OxmlElement("w:pPr")
            pStyle = OxmlElement("w:pStyle")
            pStyle.set(qn("w:val"), doc.styles["Normal"].style_id)
            pPr.append(pStyle)
            new_p.append(pPr)
            new_p.append(make_ins_run(text))
            start._p.addnext(new_p)
        print(f"  + EDIT prose after {heading_prefix}")
        return True

    for text in reversed(note_paras):
        insert_tracked_paragraph_before(doc, before, text, "Normal")
    print(f"  + EDIT prose after {heading_prefix}")
    return True


def insert_before_paragraph_text(
    doc: Document,
    paragraph_prefix: str,
    heading_text: str,
    heading_style: str,
    body_paras: list[str],
) -> bool:
    """Insert a heading + bodies immediately before a Normal/heading paragraph match."""
    for p in doc.paragraphs:
        text = paragraph_full_text(p).strip()
        if text.startswith(paragraph_prefix):
            insert_block_before(doc, p, heading_text, heading_style, body_paras)
            print(f"  + ADD {heading_text} (before '{paragraph_prefix}')")
            return True
    print(f"  ! missing insert-before paragraph: {paragraph_prefix}")
    return False


def append_after_paragraph_text(
    doc: Document,
    paragraph_prefix: str,
    note_paras: list[str],
) -> bool:
    """Insert Normal paragraphs after the first paragraph matching prefix."""
    for p in doc.paragraphs:
        text = paragraph_full_text(p).strip()
        if text.startswith(paragraph_prefix):
            anchor = p._p.getnext()
            before = Paragraph(anchor, doc) if anchor is not None and anchor.tag == qn("w:p") else None
            if before is None and anchor is not None:
                for text_body in reversed(note_paras):
                    new_p = OxmlElement("w:p")
                    pPr = OxmlElement("w:pPr")
                    pStyle = OxmlElement("w:pStyle")
                    pStyle.set(qn("w:val"), doc.styles["Normal"].style_id)
                    pPr.append(pStyle)
                    new_p.append(pPr)
                    new_p.append(make_ins_run(text_body))
                    p._p.addnext(new_p)
                print(f"  + EDIT prose after paragraph '{paragraph_prefix}'")
                return True
            for text_body in reversed(note_paras):
                insert_tracked_paragraph_before(doc, before, text_body, "Normal")
            print(f"  + EDIT prose after paragraph '{paragraph_prefix}'")
            return True
    print(f"  ! missing paragraph: {paragraph_prefix}")
    return False


def insert_new_sections_after(
    doc: Document,
    after_heading_prefix: str,
    after_level: int,
    sections: list[tuple[str, str, list[str]]],
) -> bool:
    """
    sections: list of (heading_text, heading_style, body_paras)
    Inserted in order immediately before the end-of-section boundary after after_heading.
    """
    start = find_heading(doc, after_heading_prefix)
    if start is None:
        print(f"  ! missing anchor heading: {after_heading_prefix}")
        return False
    end = find_section_end(doc, start, after_level)
    # Insert in forward order: each block is placed immediately before `end`,
    # so earlier sections end up above later ones (... s1, s2, s3, end).
    for heading_text, heading_style, bodies in sections:
        insert_block_before(doc, end, heading_text, heading_style, bodies)
        print(f"  + ADD {heading_text}")
    return True


def add_revision_history_row(doc: Document) -> bool:
    """Append a Team E row to the first table after Revision History if possible."""
    start = find_heading(doc, "Revision History")
    if start is None:
        return False
    # find first table after this heading in document body order
    seen = False
    from docx.oxml.table import CT_Tbl
    from docx.table import Table

    for child in doc.element.body.iterchildren():
        if child is start._p:
            seen = True
            continue
        if not seen:
            continue
        if child.tag == qn("w:tbl"):
            table = Table(child, doc)
            # add row
            row = table.add_row()
            values = [
                "1.x",
                date_display(),
                "Integrated Team E Milestone 2 design for Ask AI retrieval, summary confirmation, "
                "USPS mail agent, STML, and safety/consent (new §§4.9, 5.11, 6.12–6.15, 11.19.7, 12.9).",
                "Team E",
            ]
            for i, val in enumerate(values):
                if i >= len(row.cells):
                    break
                cell = row.cells[i]
                # clear and put tracked text in first paragraph
                for p in cell.paragraphs:
                    p._p.clear_content() if False else None
                p = cell.paragraphs[0]
                # clear existing runs
                for child_el in list(p._p):
                    if child_el.tag != qn("w:pPr"):
                        p._p.remove(child_el)
                p._p.append(make_ins_run(val))
            print("  + EDIT Revision History table row")
            return True
        # stop if we hit next Heading 1
        if child.tag == qn("w:p"):
            p = Paragraph(child, doc)
            if (p.style and p.style.name.startswith("Heading 1")
                    and (p.text or "").strip()
                    and not (p.text or "").strip().lower().startswith("revision")):
                break
    print("  ! could not find Revision History table")
    return False


def date_display() -> str:
    return datetime.now().strftime("%B %Y")


# --- Content written to match existing 01_TDD voice (narrative + Design notes: / ID → RESULT) ---

CONTENT = {
    "2.6": [
        "Care recipients use Ask AI, review call and visit summaries, and monitor incoming mail "
        "with short-term-memory-first interaction patterns and accessibility constraints. "
        "Caregivers may use Ask AI and review summaries only for assigned patients under "
        "consent-scoped, default view-only access. Family members receive limited read-only "
        "visibility into selected information; RBAC continues to block restricted clinical detail.",
    ],
    "3.1": [
        "Ask AI and summary enrichment assume AWS Bedrock for chat, structured summary support, "
        "and Titan embeddings; Aurora PostgreSQL with the pgvector extension and full-text search "
        "for hybrid retrieval; a Spring Boot IndexWorker consuming the indexing_outbox for "
        "SUMMARY_CREATED and TRANSCRIPT_INDEXED events; and Flutter clients for patient and "
        "caregiver Ask AI surfaces.",
    ],
    "3.2": [
        "Material risks for the AI retrieval path include Bedrock throttle and embedding latency, "
        "temporary NULL embedding vectors until backfill completes, hybrid-search cost under wide "
        "patient corpora, outbox lease parking while visit-summary indexing remains deferred, and "
        "Amplify CORS drift after app-only backend redeploys.",
    ],
    "5.5": [
        "Summary item confirmation is exposed as POST /api/summaries/{id}/items/{itemId}/confirm. "
        "Authorized reviewers promote a structured summary item after human confirmation. Successful "
        "confirmation pairs with SUMMARY_CREATED indexing so confirmed content becomes eligible for "
        "grounded Ask AI retrieval under the caller’s RBAC scope.",
    ],
    "6.5": [
        "Team E additionally depends on the shared Bedrock model allow-list and ECS task-role IAM "
        "in the platform stack, Chime or transcript upstream emitters for summary indexing, "
        "Cognito/JWT enforcement of USE_AI_FEATURES, Amplify frontend configuration for BACKEND_URL "
        "and CORS, and RDS pgvector enablement from the data stack parameter group.",
    ],
    "7.8": [
        "When 03-platform.yaml changes Bedrock, SSM, or KMS permissions used by Ask AI embeddings, "
        "operators must update the platform stack and force a new ECS deployment. App-only redeploys "
        "refresh the service image without refreshing IAM. After an app-only pass, CorsAllowedList "
        "must still include the Amplify origin or the hosted welcome page reports backend unhealthy "
        "even when curl health checks succeed.",
    ],
    "7.9": [
        "Backend rollback returns ECS to a prior ECR image tag with a forced service deployment. "
        "Platform IAM rollbacks redeploy the previous 03-platform.yaml changeset. Retrieval index "
        "rows already written for a source_record_id are replaced on re-ingest rather than "
        "automatically undone when an image rolls back.",
    ],
    "8.8": [
        "Ask AI audit events record requester identity, patient scope, denied-scope 403 outcomes, "
        "citation source identifiers, and model identifier without persisting full PHI prompts in "
        "clear text. IndexWorker logs warnings or errors when rows dead-letter or when the Bedrock "
        "runtime client is unavailable so embedding gaps remain operable.",
    ],
    "9.4": [
        "Indexing claim leases (default ten minutes) cover Titan embedding work that runs after the "
        "ingest transaction commits. Unimplemented visit-summary work uses a multi-hour no-burn park "
        "so polls do not reclaim the same outbox row every fifteen seconds. Embedding batches retry "
        "on throttle, and hybrid retrieval remains bounded by RBAC scope and top-k reciprocal-rank "
        "fusion.",
    ],
    "11.7": [
        "Ask AI, summary review, and mail surfaces are in scope for the shared VPAT plan: adequate "
        "contrast, keyboard focus order, screen-reader labels for citation lists, and short, "
        "STML-friendly language in patient-facing copy.",
    ],
    "4.9": [
        "The Mail / USPS domain stores Informed Delivery artifacts and OCR-enriched text used for "
        "patient mail recall and optional Ask AI grounding. MailPiece rows link to the patient, "
        "capture delivery or scan timing, sender and subject heuristics, OCR text and status, object-"
        "storage references for images, and consent or visibility flags for caregiver access.",
        "Design notes: Physical columns follow the Team E MailPiece data dictionary. Selected mail "
        "text may also be written to retrieval_index_chunk so hybrid search can cite mail alongside "
        "summaries and transcripts when RBAC permits.",
    ],
    "5.11": [
        "This section specifies the Ask AI gateway endpoint that orchestrates RBAC-scoped hybrid "
        "retrieval and grounded answer assembly on AWS Bedrock. Interfaces follow the conventions "
        "in §5.2 unless noted below.",
        "POST /api/ai/ask accepts an authenticated request identifying the patient scope, the "
        "natural-language query (or voice-derived text after speech-to-text), optional conversation "
        "context, and optional source-type filters (summaries, transcripts, mail, documents) that "
        "remain subject to RBAC. The response returns answer text, citations with source type and "
        "source record identifiers, safety or hold indicators when applicable, and a correlation "
        "identifier for audit. Standard outcomes include 401 or 403 for authentication and scope "
        "failures, 400 for validation errors, 429 for throttle, and 503 when Bedrock or retrieval "
        "dependencies are unavailable, with graceful client messaging.",
    ],
    "6.12": [
        "Ask AI provides grounded, cited answers over a patient’s stored care records. The design "
        "combines RBAC-scoped hybrid retrieval (PostgreSQL full-text search and pgvector similarity) "
        "with a Bedrock assembler that must cite retrieved sources rather than invent unsupported "
        "clinical advice.",
        "Design notes: The runtime pipeline authorizes scope, retrieves and fuses candidates, builds "
        "cited context, invokes the model, applies optional safety or HITL hold rules, and returns "
        "the answer with citations. Voice queries join the same retrieval path after speech-to-text. "
        "Denied-scope asks produce 403 responses that are audit logged with the requester and "
        "patient identifiers.",
    ],
    "6.13": [
        "The USPS Informed Delivery mail agent ingests mail digests and images, applies OCR with "
        "Textract fallback where required, persists MailPiece entities described in §4.9, and "
        "exposes patient-facing mail review flows.",
        "Design notes: Caregiver visibility remains consent-aware. Selected mail text may be indexed "
        "into the shared retrieval corpus so Ask AI can cite mail when the caller’s scope allows.",
    ],
    "6.14": [
        "Short-Term Memory (STML) support shapes patient-facing summaries, reminders, and Ask AI "
        "responses toward brief, high-signal content. Workflows prefer short confirmation paths "
        "over long multi-step forms when an actionable item can be confirmed in place.",
        "Design notes: Confirmed actionable summary items feed reminder cues and remain visible in "
        "STML-oriented presentation without requiring the patient to re-read full transcripts.",
    ],
    "6.15": [
        "Safety, Consent & Clarity defines patient-facing guardrails for Ask AI and generative "
        "summaries. The assistant does not provide diagnosis, treatment, or dosage recommendations. "
        "Caregiver visibility stays consent-scoped, secondary validation or HITL hold-and-release "
        "applies where configured, and AI-generated text carries clear source attribution.",
        "Design notes: Responses that are held show an explicit hold state rather than silent "
        "failure. Patient copy remains short and STML-friendly; caregiver copy may surface "
        "additional consented clinical detail.",
    ],
    "11.19.7": [
        "ASK-001: Authorized patient asks in-scope question via POST /api/ai/ask → SUCCESS",
        "ASK-002: Caregiver asks for assigned patient under consent scope → SUCCESS",
        "ASK-003: User asks for patient outside RBAC scope → BLOCKED",
        "ASK-004: Ask AI response includes source citations for grounded claims → SUCCESS",
        "ASK-005: Bedrock unavailable during ask → GRACEFUL DEGRADATION",
        "SUM-001: Reviewer confirms summary item via confirm endpoint → SUCCESS",
        "SUM-002: SUMMARY_CREATED indexing writes retrieval chunks for confirmed content → SUCCESS",
        "IDX-001: IndexWorker claims outbox row with lease and processes SUMMARY_CREATED → SUCCESS",
        "IDX-002: Deferred visit summary parks without burning attempts → SUCCESS",
        "IDX-003: Missing Bedrock client leaves embedding NULL and logs warning → SUCCESS",
        "MAIL-001: USPS digest ingest creates MailPiece with OCR text → SUCCESS",
        "MAIL-002: OCR fallback to Textract when primary path fails → SUCCESS",
        "STML-001: Patient sees short confirmation path for actionable summary item → SUCCESS",
        "SAFE-001: Ask AI refuses diagnosis or dosing advice → BLOCKED",
        "SAFE-002: Held answer exposes hold state to client → SUCCESS",
    ],
    "11.20_extra": [
        "Team E automation covers JUnit and Mockito suites for retrieval, indexing, and embedding "
        "services; Flutter widget tests for Ask AI screens where present; and CI gates on Team AE "
        "develop branches. Shared pipelines use non-PHI fixtures only.",
    ],
    "12.9": [
        "The following items capture Team E-specific limitations and open work for Ask AI, "
        "summaries, USPS mail, STML, and safety. They complement the shared issues in §12.8.",
    ],
    "12.9.1": [
        "Visit-summary indexing may remain parked until the visit_summaries ingest path lands. "
        "Embedding columns can stay NULL when Bedrock is unavailable until backfill runs. Ask AI "
        "answer quality tracks retrieval-index freshness after transcript and summary emit.",
    ],
    "12.9.2": [
        "App-only redeploys can reset CorsAllowedList to localhost defaults and break Amplify "
        "browser checks. Content-hash short-circuit skips re-chunking when the publisher hash is "
        "unchanged, even if chunker metadata alone would differ.",
    ],
    "12.9.3": [
        "Clinical decision-making, diagnosis, treatment selection, and dosage recommendations remain "
        "outside Ask AI scope.",
    ],
    "12.9.4": [
        "Planned follow-ons include medication-timeline aggregation in retrieval responses, richer "
        "tier-two HITL workflows, speech-to-text hardening for voice asks, and broader mail "
        "classification.",
    ],
    "12.9.5": [
        "Bedrock IAM must keep named foundation-model entries that include Titan embed text v1 "
        "rather than only a wildcard foundation-model resource. Transcript and summary emit "
        "ownership stays with the teams that persist Chime and call-summary records.",
    ],
    "13.4": [
        "Patient and caregiver UI copy defaults to en-US. Patient-mode Ask AI and STML text avoids "
        "dense medical jargon; caregiver mode may present additional clinical detail when consent "
        "allows. Localization follows existing Flutter l10n patterns when those packs are enabled.",
    ],
}


def apply() -> Path:
    if not SOURCE_TDD.exists():
        raise SystemExit(f"Missing source: {SOURCE_TDD}")
    # Prefer canonical output; if Word has it open, write _formatted.docx instead.
    target = OUTPUT
    if target.exists():
        try:
            with open(target, "a"):
                pass
        except PermissionError:
            target = OUTPUT.with_name(OUTPUT.stem + "_formatted.docx")
            print(f"Output locked in Word; writing {target.name}")
    return apply_to(target)


def apply_to(output: Path) -> Path:
    print(f"Copying {SOURCE_TDD.name} -> {output.name}")
    shutil.copy2(SOURCE_TDD, output)

    doc = Document(str(output))
    enable_track_revisions(doc)

    print("Revision History...")
    add_revision_history_row(doc)

    print("EDIT prose...")
    append_edit_note_after_heading(doc, "2.6 System Users", CONTENT["2.6"])
    append_edit_note_after_heading(doc, "3.1 Technology Decisions", CONTENT["3.1"])
    append_edit_note_after_heading(doc, "3.2 Technology Risks", CONTENT["3.2"])
    append_edit_note_after_heading(doc, "5.5 Transcription and Summary", CONTENT["5.5"])
    append_edit_note_after_heading(doc, "6.5 Cross-Team Functional Dependencies", CONTENT["6.5"])
    append_edit_note_after_heading(doc, "7.8 Deployment Environments", CONTENT["7.8"])
    append_edit_note_after_heading(doc, "7.9 Rollback Procedure", CONTENT["7.9"])
    append_edit_note_after_heading(doc, "8.8 Logging and Monitoring", CONTENT["8.8"])
    append_edit_note_after_heading(doc, "9.4 Bottleneck Mitigation", CONTENT["9.4"])
    append_edit_note_after_heading(doc, "11.7 Accessibility Testing", CONTENT["11.7"])

    print("ADD sections...")
    insert_new_sections_after(
        doc,
        "4.8 Administration and Audit Entities",
        2,
        [
            (
                "4.9 Mail / USPS Entities (MailPiece)",
                "Heading 2",
                CONTENT["4.9"],
            )
        ],
    )
    insert_new_sections_after(
        doc,
        "5.10 API Versioning and Backward Compatibility",
        2,
        [
            (
                "5.11 Ask AI Endpoint (POST /api/ai/ask)",
                "Heading 2",
                CONTENT["5.11"],
            )
        ],
    )
    insert_new_sections_after(
        doc,
        "6.11 Screening and Comprehensive Assessment Workflows",
        2,
        [
            ("6.12 AI-Assisted Retrieval (Ask AI)", "Heading 2", CONTENT["6.12"]),
            ("6.13 USPS Informed Delivery Mail Agent", "Heading 2", CONTENT["6.13"]),
            ("6.14 Short-Term Memory (STML) Support", "Heading 2", CONTENT["6.14"]),
            ("6.15 Safety, Consent & Clarity", "Heading 2", CONTENT["6.15"]),
        ],
    )
    # Match §11.19.x pattern (Heading 3 + ID → RESULT lines), placed before existing 11.20 prose.
    insert_before_paragraph_text(
        doc,
        "11.20 Test Automation",
        "11.19.7 Ask AI and Team E Feature Test Coverage",
        "Heading 3",
        CONTENT["11.19.7"],
    )
    append_after_paragraph_text(
        doc,
        "Automated testing will be used whenever practical",
        CONTENT["11.20_extra"],
    )
    insert_new_sections_after(
        doc,
        "12.8 Known Issues and Limitations",
        2,
        [
            ("12.9 Team E Known Issues and Limitations", "Heading 2", CONTENT["12.9"]),
            ("12.9.1 Technical Limitations", "Heading 3", CONTENT["12.9.1"]),
            ("12.9.2 Known Issues", "Heading 3", CONTENT["12.9.2"]),
            ("12.9.3 Exclusions", "Heading 3", CONTENT["12.9.3"]),
            ("12.9.4 Future Enhancements", "Heading 3", CONTENT["12.9.4"]),
            ("12.9.5 Cross-Team Resolution Notes", "Heading 3", CONTENT["12.9.5"]),
        ],
    )
    insert_new_sections_after(
        doc,
        "13.3.6 Screening and Assessment",
        3,
        [
            (
                "13.4 Display Language and Localization Impact",
                "Heading 2",
                CONTENT["13.4"],
            )
        ],
    )

    apply_table_updates(doc)

    # Save
    output.parent.mkdir(parents=True, exist_ok=True)
    try:
        doc.save(output)
        print(f"Wrote {output}")
    except PermissionError:
        alt = output.with_name(output.stem + "_v2.docx")
        doc.save(alt)
        print(f"Wrote {alt} (original locked)")
        output = alt

    try:
        shutil.copy2(output, REPO_COPY)
        print(f"Wrote {REPO_COPY}")
    except Exception as ex:
        print(f"Repo copy skipped: {ex}")

    return output


if __name__ == "__main__":
    apply()
