# -*- coding: utf-8 -*-
"""Populate Section 1.5 (Definitions, Abbreviations, and Acronyms) of the
Deployment and Operations Guide with an intro paragraph + glossary table,
replacing the placeholder text, and matching the visual style of the
existing tables in the document (single black borders, bold header row).
"""
import shutil
import docx
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = r"C:\Users\ravic\Downloads\05_Deployment_and_Operations_Guide.docx"
BACKUP = r"C:\Users\ravic\Downloads\05_Deployment_and_Operations_Guide - backup.docx"

INTRO_TEXT = (
    "To better facilitate the reading of this documentation, special "
    "abbreviations and terms have been placed in one location for the "
    "entire document here. These definitions, acronyms, and abbreviations "
    "are used across this guide:"
)

TABLE_CAPTION_LINES = ["Table 2", "Definitions, Acronyms, and Abbreviations"]

ENTRIES = [
    ("AI", "Artificial Intelligence \u2014 powers the CareConnect chat, sentiment, and summarization features via Amazon Bedrock foundation models (Section 3.3.3)."),
    ("Amplify (AWS Amplify)", "AWS service used to build, host, and configure the CareConnect Flutter web frontend, including branch-based environment variables (Section 3.5.2)."),
    ("API", "Application Programming Interface \u2014 the HTTP endpoints the Spring Boot backend exposes to the frontend and other clients."),
    ("API Gateway", "AWS-managed HTTP API and VPC Link layer that routes browser traffic from Amplify to the ECS Fargate backend (Section 3.5)."),
    ("ARN", "Amazon Resource Name \u2014 the unique identifier format AWS uses for resources such as the GitHub Actions deployment role (AWS_GITHUB_ACTIONS_ROLE_ARN)."),
    ("Arrange\u2013Act\u2013Assert", "A test-structuring pattern (setup, action, verification) used when writing unit, widget, and integration tests (Section 2.5)."),
    ("AVD", "Android Virtual Device \u2014 an emulated Android device created in Android Studio's Device Manager, used to run and test the Flutter app and integration tests."),
    ("AWS", "Amazon Web Services \u2014 the cloud provider hosting CareConnect's backend, database, and frontend infrastructure."),
    ("AWS CLI", "AWS Command Line Interface \u2014 authenticates to and runs deployment scripts against the shared AWS account."),
    ("axe-core", "Open-source accessibility testing engine run in the CI pipeline against the built Flutter web app for WCAG 2.1 Level A and AA compliance."),
    ("Bedrock (Amazon Bedrock)", "Managed AWS service providing access to foundation AI models; powers the CareConnect AI chat feature."),
    ("CI/CD", "Continuous Integration / Continuous Delivery (or Deployment) \u2014 the automated GitHub Actions pipeline that builds, tests, and deploys CareConnect on every change."),
    ("CloudFormation", "AWS infrastructure-as-code service used to provision CareConnect's networking, data, platform, and service stacks (01\u201304-*.yaml)."),
    ("CODEOWNER", "A designated reviewer (initially the team leads) whose approval is required before a pull request can merge into main or develop."),
    ("CORS", "Cross-Origin Resource Sharing \u2014 the browser security mechanism controlled by CorsAllowedList that determines which origins may call the backend API."),
    ("Docker / Docker Desktop", "Containerization platform used to run the local PostgreSQL database and build the backend container image deployed to AWS."),
    ("ECR", "Elastic Container Registry \u2014 the AWS registry that stores the backend's Docker container images."),
    ("ECS / ECS Fargate", "Elastic Container Service \u2014 AWS's serverless container-hosting engine that runs the Spring Boot backend without managing underlying EC2 instances."),
    ("Full-Stack Deploy", "A deployment that provisions all four CloudFormation stacks (networking, data, platform, service) from scratch (Section 3.5)."),
    ("GitHub Actions", "The CI/CD automation platform, defined as workflow files under .github/workflows/, that runs CareConnect's build, test, and deployment pipelines."),
    ("Hotfix", "An urgent, approved fix merged directly into the main production branch outside the normal feature-branch flow."),
    ("HTTPS", "Hypertext Transfer Protocol Secure \u2014 the encrypted protocol required for BACKEND_URL and FRONTEND_URL endpoints."),
    ("IAM", "AWS Identity and Access Management \u2014 provisioned in the platform stack (03-platform.yaml) and assumed by GitHub Actions via OIDC for deployment."),
    ("IDE", "Integrated Development Environment (e.g., VS Code or IntelliJ IDEA) used for editing, debugging, and language support."),
    ("JaCoCo", "Java Code Coverage library that generates the backend's code-coverage report during the CI pipeline."),
    ("JDBC", "Java Database Connectivity \u2014 the API the backend uses to connect to PostgreSQL (see JDBC_URI)."),
    ("JDK", "Java Development Kit \u2014 the Java 17 (LTS) toolchain required to compile and run the Spring Boot backend."),
    ("JWT", "JSON Web Token \u2014 the stateless authentication credential issued on login and validated via SECURITY_JWT_SECRET."),
    ("kebab-case", "A lowercase, hyphen-separated naming convention (e.g., evv-quality-guardrails) used in feature branch names."),
    ("LCOV", "A code-coverage report format produced by the Flutter test suite and consumed by the CI coverage gate."),
    ("LTS", "Long-Term Support \u2014 refers to the supported release lines of the JDK (17) and Node.js (20) used by the project."),
    ("Maven / mvnw", "Apache Maven, the backend's build and dependency-management tool; mvnw is the repository's bundled Maven Wrapper script."),
    ("NGINX", "An example static web server named as an alternative hosting target for compiled Flutter web assets."),
    ("Node.js", "JavaScript runtime required to run the axe-core accessibility tooling used by the CI pipeline."),
    ("OAuth", "Open Authorization \u2014 the protocol used for Google sign-in (GOOGLE_CLIENT_ID / GOOGLE_CLIENT_SECRET)."),
    ("OIDC", "OpenID Connect \u2014 the identity protocol GitHub Actions uses to assume an AWS IAM role for deployment without long-lived credentials."),
    ("pgAdmin", "A graphical PostgreSQL administration client provisioned by the local Docker Compose setup for inspecting data and running manual queries."),
    ("pgvector", "A PostgreSQL extension providing vector similarity search, used by CareConnect's AI retrieval features."),
    ("PostgreSQL", "The relational database (version 15) used by CareConnect in local and AWS environments."),
    ("PR", "Pull Request \u2014 a proposed set of code changes submitted on GitHub for review and merge."),
    ("RDS", "Amazon Relational Database Service \u2014 the managed PostgreSQL hosting used in deployed AWS environments."),
    ("README", "A repository documentation file (e.g., the cloudformation-fargate README) with current setup or deployment procedures."),
    ("S3", "Amazon Simple Storage Service \u2014 example object-storage target for hosting compiled web assets."),
    ("SDK", "Software Development Kit (Android SDK, Flutter SDK, Dart SDK) \u2014 the toolsets used to build the CareConnect frontend and mobile app."),
    ("SendGrid", "Third-party email delivery provider, configured via EMAIL_PROVIDER and SENDGRID_API_KEY."),
    ("Smoke Test", "A quick, high-level check performed after configuration or deployment to confirm a service (health endpoint, OAuth, email, AI) is working."),
    ("SonarCloud", "Cloud-based static code-analysis service that scans pull requests for code-quality issues."),
    ("Spring Boot", "The Java application framework used to build the CareConnect backend REST API and service layer."),
    ("SSM", "AWS Systems Manager Parameter Store \u2014 injects non-secret runtime configuration into the ECS task at startup."),
    ("TBD", "To Be Determined \u2014 marks deliverables in the project-documentation table that don't yet have a published version (Section 1.4)."),
    ("TDD", "Test-Driven Development \u2014 the project's practice of writing tests alongside feature code (Section 2.5)."),
    ("UMGC", "University of Maryland Global Campus \u2014 the academic institution offering the SWEN 670 capstone course under which CareConnect is developed."),
    ("VPC / VPC Link", "Virtual Private Cloud \u2014 the private AWS network, and the API Gateway integration that connects it to the ECS Fargate service."),
    ("WCAG", "Web Content Accessibility Guidelines \u2014 the accessibility standard (2.1 Level A/AA) validated by the CI pipeline's axe-core scan."),
]


def set_cell_borders(cell, sz=8, color="000000"):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), str(sz))
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), color)
        borders.append(el)
    tcPr.append(borders)


def set_cell_margins(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement('w:tcMar')
    for edge, w in (('top', 120), ('left', 180), ('bottom', 120), ('right', 180)):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:w'), str(w))
        el.set(qn('w:type'), 'dxa')
        mar.append(el)
    tcPr.append(mar)


def bold_cell(cell):
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True


def main():
    shutil.copyfile(SRC, BACKUP)
    doc = docx.Document(SRC)

    body_paragraphs = doc.paragraphs

    heading_idx = None
    for i, p in enumerate(body_paragraphs):
        if p.text.strip() == "1.5 Definitions, Abbreviations, and Acronyms":
            heading_idx = i
            break
    if heading_idx is None:
        raise RuntimeError("Could not find 1.5 heading paragraph")

    placeholder1 = body_paragraphs[heading_idx + 1]
    placeholder2 = body_paragraphs[heading_idx + 2]

    if placeholder1.text.strip().startswith("Wait till doc is done"):
        for run in list(placeholder1.runs):
            run.text = ""
        if placeholder1.runs:
            placeholder1.runs[0].text = INTRO_TEXT
        else:
            placeholder1.add_run(INTRO_TEXT)
    else:
        raise RuntimeError(f"Unexpected placeholder1 text: {placeholder1.text!r}")

    if placeholder2.text.strip() == "\u2026":
        for run in list(placeholder2.runs):
            run.text = ""
        if placeholder2.runs:
            placeholder2.runs[0].text = TABLE_CAPTION_LINES[0]
        else:
            placeholder2.add_run(TABLE_CAPTION_LINES[0])
    else:
        raise RuntimeError(f"Unexpected placeholder2 text: {placeholder2.text!r}")

    caption2 = doc.add_paragraph(TABLE_CAPTION_LINES[1])
    placeholder2._p.addnext(caption2._p)

    table = doc.add_table(rows=1, cols=2)
    table.autofit = False

    hdr = table.rows[0].cells
    hdr[0].text = "Term / Abbreviation"
    hdr[1].text = "Definition"
    for c in hdr:
        set_cell_borders(c)
        set_cell_margins(c)
        bold_cell(c)

    for term, definition in ENTRIES:
        row = table.add_row().cells
        row[0].text = term
        row[1].text = definition
        for c in row:
            set_cell_borders(c)
            set_cell_margins(c)

    tbl_element = table._tbl
    caption2._p.addnext(tbl_element)

    trailing_p = doc.add_paragraph("")
    tbl_element.addnext(trailing_p._p)

    doc.save(SRC)
    print("Saved:", SRC)
    print("Backup at:", BACKUP)


if __name__ == "__main__":
    main()
