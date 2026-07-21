"""Generate Word document: PR code review for sentiment clip feature branch."""

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.shared import Pt

OUTPUT = (
    Path(__file__).resolve().parents[1]
    / "docs"
    / "PR_Review_Sentiment_Clip_feature_a-drattray-sentiment.docx"
)


def heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def para(doc: Document, text: str, bold: bool = False, highlight: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    if highlight:
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def bullets(doc: Document, items: list[str], highlight_indices: set[int] | None = None) -> None:
    highlight_indices = highlight_indices or set()
    for i, item in enumerate(items):
        p = doc.add_paragraph(item, style="List Bullet")
        if i in highlight_indices:
            for run in p.runs:
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def code(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    for i, h in enumerate(headers):
        tbl.rows[0].cells[i].text = h
    for row in rows:
        cells = tbl.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val


def build() -> Document:
    doc = Document()
    title = doc.add_heading("PR Code Review", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    para(doc, "Sentiment clip playback — feature/a-drattray-sentiment → team-ae-develop")
    para(doc, f"Review date: {date.today().isoformat()}")
    para(doc, "Reviewer: AI-assisted code review (Cursor)")
    doc.add_paragraph()

    heading(doc, "Review metadata", 1)
    table(
        doc,
        ["Field", "Value"],
        [
            ["Source branch", "feature/a-drattray-sentiment"],
            ["Target branch", "team-ae-develop"],
            ["Scope", "16 files changed (+1,696 / −21 lines)"],
            ["Feature", "Sentiment-linked ±15 s video clip on post-call telemetry timeline"],
            ["WBS refs", "§3.3.2 / M7–M8 (SENT-CLIP-001..006)"],
        ],
    )

    heading(doc, "1. Change summary", 1)
    para(
        doc,
        "This PR enables sentiment-linked video clip playback on the post-call telemetry "
        "screen. When a caregiver taps a sentiment timeline dot and user-initiated recording "
        "is ready, the app loads a presigned S3 URL, computes a ±15 second window relative to "
        "recording start, and plays an inline 16:9 clip. It also configures S3 bucket CORS so "
        "Flutter web can HTTP Range-seek on composited MP4 presigned URLs.",
    )

    heading(doc, "1.1 Backend changes", 2)
    bullets(
        doc,
        [
            "Apply recording bucket CORS at bucket creation (GET/HEAD, expose Accept-Ranges, "
            "Content-Range, Content-Length, ETag).",
            "Add recordingStartedAt to generatePlaybackUrl() response for client clip offset anchor.",
            "Add scripts/recording-bucket-cors.json for manual CORS fallback.",
            "Update TEAM_A_VIDEO_CALL_QUICKSTART.md with s3:PutBucketCors IAM and manual aws s3api steps.",
            "Extend CallRecordingServiceTest with CORS and recordingStartedAt assertions.",
            "Fix CallControllerExtendedTest to expect playbackUrl (not url) JSON field.",
        ],
    )

    heading(doc, "1.2 Frontend changes", 2)
    bullets(
        doc,
        [
            "New sentiment_clip_window.dart — pure clip math (±15 s, UTC normalization, pause logic).",
            "New sentiment_clip_recording_status.dart — recording status copy and gating helpers.",
            "New SentimentClipPlayerWidget — Chewie + video_player with test controllerFactory hook.",
            "PostCallTelemetrySummaryScreen — timeline tap loads clip, snackbar on processing, dismiss UX.",
            "ApiService.getCallRecordingPlaybackData() returns full map; legacy URL helper preserved.",
            "Tests: sentiment_clip_unit_test.dart, sentiment_clip_player_widget_test.dart, "
            "post_call_telemetry_summary_screen_test.dart (~900 lines).",
        ],
    )

    heading(doc, "1.3 Feature commits", 2)
    table(
        doc,
        ["Commit", "Summary"],
        [
            ["fef6f55", "S3 bucket CORS + documentation"],
            ["0f6e2c2", "recordingStartedAt on playback response"],
            ["586df75", "Full playback data map on client"],
            ["b1b65df", "SentimentClipPlayerWidget"],
            ["743f056", "Timeline tap → clip / snackbar / transcript fallback"],
            ["22c40cf", "Dispose video on dismiss"],
            ["b07dae2", "Test coverage gaps filled"],
            ["9f12c92", "Logging revert"],
        ],
    )

    heading(doc, "2. Bug and risk analysis", 1)

    heading(doc, "2.1 High — Timezone mismatch for clip offset", 2)
    para(
        doc,
        "Backend emits recordingStartedAt as ISO_LOCAL_DATE_TIME without timezone offset. "
        "Frontend parses with DateTime.parse() (device-local interpretation) then converts to UTC. "
        "Sentiment occurredAt from telemetry may be UTC/offset-aware. Server/browser timezone "
        "mismatch can shift clip seek by hours.",
        highlight=True,
    )
    para(doc, "Location: CallRecordingService.generatePlaybackUrl() + post_call_telemetry_summary_screen.dart")
    code(
        doc,
        "Backend: rec.getStartedAt().format(DateTimeFormatter.ISO_LOCAL_DATE_TIME)\n"
        "Frontend: DateTime.parse(recordingStartedAtRaw).toUtc()",
    )

    heading(doc, "2.2 Medium — Race on rapid timeline taps", 2)
    para(
        doc,
        "_loadClipForSelectedSentiment() has no request generation token or cancellation. "
        "If the user taps dot A then dot B quickly, a slower response for A can overwrite B's clip state.",
    )

    heading(doc, "2.3 Medium — Silent failure when playback fetch fails", 2)
    para(
        doc,
        "When getCallRecordingPlaybackData() returns null or missing fields, the UI sets "
        "_loadingClip = false with no snackbar or inline error. User sees spinner disappear with no feedback.",
    )

    heading(doc, "2.4 Medium — clipEndSec not clamped to video duration", 2)
    para(
        doc,
        "computeSentimentClipWindow clamps clipStartSec at 0 but not clipEndSec to recording length. "
        "Late-call sentiment samples may seek near EOF; pause logic may behave inconsistently.",
    )

    heading(doc, "2.5 Low — playbackReady semantics differ between endpoints", 2)
    bullets(
        doc,
        [
            "buildRecordingMap: playbackReady requires resolvePlayableVideoKey != null.",
            "generatePlaybackUrl success path: playbackReady = initiatedByUserId != null only "
            "(video key already resolved at that point, so consistent in practice).",
            "Recommend aligning for future API callers.",
        ],
    )

    heading(doc, "2.6 Low — S3 CORS AllowedOrigins wildcard", 2)
    para(
        doc,
        "resolveS3CorsOrigins() collapses Spring CORS patterns to *. Acceptable for dev presigned URLs; "
        "review for production shared recording buckets.",
    )

    heading(doc, "2.7 Low — CORS apply failure is warn-only", 2)
    para(
        doc,
        "If s3:PutBucketCors is missing, web clip seek fails with no user-visible backend signal. "
        "Documentation covers manual fix; consider dev health indicator.",
    )

    heading(doc, "2.8 Positive edge-case handling", 2)
    bullets(
        doc,
        [
            "System-initiated recordings blocked (initiatedByUserId == null).",
            "Processing state → snackbar; transcript scroll still works.",
            "Clip panel dismiss clears selection and disposes controllers.",
            "mounted checks on async paths.",
            "Early sentiment → clipStartSec clamped to 0.",
        ],
    )

    heading(doc, "3. Architecture and style", 1)

    heading(doc, "3.1 Strengths", 2)
    table(
        doc,
        ["Area", "Assessment"],
        [
            ["Separation of concerns", "Clip math and status copy in pure Dart utilities — testable"],
            ["Widget design", "SentimentClipPlayerWidget with controllerFactory test hook — excellent"],
            ["Backward compatibility", "getCallRecordingPlaybackUrl() wrapper preserved"],
            ["Test coverage", "~900 lines unit/widget/screen tests with SENT-CLIP traceability"],
            ["Ops/docs", "IAM + manual CORS fallback in TEAM_A_VIDEO_CALL_QUICKSTART.md"],
            ["Idempotency", "CORS applied at bucket resolve, same pattern as Chime bucket policy"],
        ],
    )

    heading(doc, "3.2 Concerns", 2)
    table(
        doc,
        ["Area", "Assessment"],
        [
            ["CallRecordingService size", "Already large; CORS adds ~90 lines — consider extractor later"],
            ["Screen state", "Clip state fields in PostCallTelemetrySummaryScreen — acceptable for MVP"],
            ["New dependency", "chewie in pubspec.yaml — justified; verify license/size"],
            ["Unrelated diff", "AIServiceFactory log string only — minor PR noise"],
        ],
    )

    para(doc, "Overall: clean, follows project patterns; testability above average for frontend video work.", bold=True)

    heading(doc, "4. Recommendations", 1)

    heading(doc, "R1 — Fix timezone: emit UTC instant from backend (required before merge)", 2)
    para(doc, "Recommended fix:", bold=True)
    code(
        doc,
        "// CallRecordingService.generatePlaybackUrl()\n"
        'playbackResult.put(\n'
        '    "recordingStartedAt",\n'
        '    rec.getStartedAt() != null\n'
        '        ? rec.getStartedAt().atZone(ZoneOffset.UTC).toInstant().toString()\n'
        '        : null);',
    )
    para(doc, "Frontend continues: DateTime.parse(recordingStartedAtRaw).toUtc() when string ends with Z.")
    para(doc, "Add test asserting recordingStartedAt is UTC/offset-aware.", highlight=True)

    heading(doc, "R2 — Guard async clip loads with generation token", 2)
    code(
        doc,
        "int _clipLoadGeneration = 0;\n\n"
        "Future<void> _loadClipForSelectedSentiment() async {\n"
        "  final generation = ++_clipLoadGeneration;\n"
        "  // ... await fetch ...\n"
        "  if (!mounted || generation != _clipLoadGeneration) return;\n"
        "  // apply state\n"
        "}",
    )

    heading(doc, "R3 — User feedback on clip load failure", 2)
    code(
        doc,
        "void _showClipLoadFailedSnackBar() {\n"
        "  if (!mounted) return;\n"
        "  ScaffoldMessenger.of(context).showSnackBar(\n"
        "    const SnackBar(\n"
        "      content: Text(\n"
        "        'Could not load video for this moment. Try again shortly.',\n"
        "      ),\n"
        "    ),\n"
        "  );\n"
        "}",
    )
    para(doc, "Call when data == null or required playback fields are missing.")

    heading(doc, "R4 — Clamp clip end to video duration", 2)
    code(
        doc,
        "final durationSec = controller.value.duration.inMilliseconds / 1000.0;\n"
        "final effectiveEndSec = math.min(widget.clipEndSec, durationSec);",
    )
    para(doc, "Use effectiveEndSec in sentimentClipShouldPause / playback tick handler.")

    heading(doc, "R5 — Align playbackReady in success response", 2)
    code(
        doc,
        'playbackResult.put("playbackReady", true); // video key already resolved',
    )

    heading(doc, "R6 — Tighten S3 CORS in non-dev profiles", 2)
    code(
        doc,
        "@Value(\"${careconnect.recording.cors-allow-wildcard:true}\")\n"
        "private boolean recordingCorsAllowWildcard;",
    )
    para(doc, "Set recording.cors-allow-wildcard=false in production parameter files.")

    heading(doc, "5. Verdict and merge checklist", 1)
    table(
        doc,
        ["Category", "Rating"],
        [
            ["Feature completeness", "Meets sentiment clip MVP"],
            ["Test coverage", "Strong"],
            ["Merge readiness", "Fix R1 (timezone) before merge; R2/R3 recommended"],
            ["Security", "Review S3 CORS * for non-dev environments"],
        ],
    )

    para(doc, "Suggested PR title:", bold=True)
    para(doc, "feat(video): sentiment clip playback with S3 CORS and ±15s seek window")

    heading(doc, "Merge checklist", 2)
    bullets(
        doc,
        [
            "recordingStartedAt emitted as UTC instant (R1)",
            "Flutter tests pass (sentiment_clip_*, post_call_telemetry_summary_screen_test)",
            "Backend tests pass (CallRecordingServiceTest CORS + playback)",
            "Manual web test: timeline dot → clip seeks correctly",
            "IAM docs updated (s3:PutBucketCors) — done in this PR",
        ],
        highlight_indices={0},
    )

    heading(doc, "6. Files changed (reference)", 1)
    bullets(
        doc,
        [
            "backend/core/scripts/recording-bucket-cors.json",
            "backend/core/src/main/java/com/careconnect/ai/AIServiceFactory.java",
            "backend/core/src/main/java/com/careconnect/service/CallRecordingService.java",
            "backend/core/src/test/java/com/careconnect/controller/CallControllerExtendedTest.java",
            "backend/core/src/test/java/com/careconnect/service/CallRecordingServiceTest.java",
            "docs/guides/TEAM_A_VIDEO_CALL_QUICKSTART.md",
            "frontend/lib/services/api_service.dart",
            "frontend/lib/utils/sentiment_clip_recording_status.dart",
            "frontend/lib/utils/sentiment_clip_window.dart",
            "frontend/lib/widgets/post_call_telemetry_summary_screen.dart",
            "frontend/lib/widgets/sentiment_clip_player_widget.dart",
            "frontend/pubspec.yaml",
            "frontend/test/post_call_telemetry_summary_screen_test.dart",
            "frontend/test/sentiment_clip_player_widget_test.dart",
            "frontend/test/sentiment_clip_unit_test.dart",
            "frontend/test/services/api_service_test.dart",
        ],
    )

    return doc


def main() -> None:
    doc = build()
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    main()
