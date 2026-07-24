"""Generate Word document: Retrieval orchestration design (RBAC → hybrid → context → LLM)."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX

ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "Retrieval_Orchestration_RBAC_Hybrid_Context_LLM_Design.docx"


def heading(doc, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def para(doc, text: str, bold: bool = False, highlight: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    if highlight:
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def bullets(doc, items: list[str], highlight_indices: set[int] | None = None) -> None:
    highlight_indices = highlight_indices or set()
    for i, item in enumerate(items):
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        if i in highlight_indices:
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def code(doc, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"


def table(doc, headers, rows, highlight_rows: set[int] | None = None) -> None:
    highlight_rows = highlight_rows or set()
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for r_idx, row in enumerate(rows, start=1):
        cells = t.add_row().cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = val
            if r_idx in highlight_rows:
                for p in cells[c_idx].paragraphs:
                    for run in p.runs:
                        run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def build() -> None:
    doc = Document()

    title = doc.add_heading(
        "Retrieval Orchestration Design: RBAC Filter → Hybrid Search → Context Assembly → LLM",
        0,
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para(doc, "CareConnect Team E — Ask AI Gateway Internal Architecture")
    para(
        doc,
        "Specifies the orchestration pipeline behind POST /api/ai/ask: how authorized retrieval "
        "scope is resolved, hybrid search ranks chunks, minimum-necessary context is assembled, "
        "and Bedrock generates a grounded structured answer. Aligns with SCC-3, hybrid retrieval, "
        "RBAC scope, pgvector embedding, HITL, and API contract research.",
    )
    para(
        doc,
        "Status: DESIGN ONLY — no AiAskOrchestrator exists; BedrockAIChatService sends raw user "
        "message with no retrieval; PatientContextRetrievalService is an in-memory keyword stub.",
        highlight=True,
    )
    doc.add_paragraph()

    # 1 Executive summary
    heading(doc, "1. Executive Summary", 1)
    para(
        doc,
        "Retrieval orchestration is a linear pipeline with short-circuit exits. Every stage runs "
        "inside AiAskOrchestrator (or AiGatewayService), invoked by AiAskController. The pipeline "
        "never calls the LLM until RBAC scope is resolved and hybrid search returns candidate chunks "
        "(or an explicit NO_RECORDS short-circuit). Post-LLM stages (safety, tier, HITL, audit) are "
        "downstream of context assembly but included for end-to-end clarity.",
    )
    bullets(
        doc,
        [
            "Single entry: AiAskOrchestrator.ask(AiAskRequest, User caller)",
            "Fail closed: 403 on scope failure; 503 on index unavailable; no ungrounded LLM fallback",
            "FR-AI-9: Context assembly enforces token budget — only top-K chunk excerpts to LLM",
            "Reuse existing services where possible; replace PatientContextRetrievalService entirely",
        ],
        highlight_indices={1, 3},
    )

    # 2 Architecture diagram
    heading(doc, "2. Pipeline Overview", 1)
    code(
        doc,
        "POST /api/ai/ask\n"
        "    │\n"
        "    ▼\n"
        "[0] Gateway preflight ── rate limit, sanitize query, audit start\n"
        "    │\n"
        "    ▼\n"
        "[1] RBAC filter ──────── RetrievalScopeService → RetrievalScope\n"
        "    │                      (403 if patient/source not permitted)\n"
        "    ▼\n"
        "[2] Query planner ────── intent + structured prefilter hints\n"
        "    │\n"
        "    ▼\n"
        "[3] Hybrid search ────── FTS ∥ pgvector → RRF merge → RankedChunk[]\n"
        "    │                      (NO_RECORDS exit if empty)\n"
        "    ▼\n"
        "[4] Context assembly ─── token budget, prompt blocks, citation map\n"
        "    │\n"
        "    ▼\n"
        "[5] LLM inference ────── LlmRouter → Bedrock structured JSON output\n"
        "    │\n"
        "    ▼\n"
        "[6] Post-process ─────── validate, tier, HITL hold, disclaimer, audit\n"
        "    │\n"
        "    ▼\n"
        "AiAskResponse",
    )

    heading(doc, "2.1 Orchestrator component map", 2)
    table(
        doc,
        ["Stage", "New component", "Existing code to reuse / replace"],
        [
            ["0 Gateway", "AiAskGatewayPreflight", "LangChainGovernanceService, InputSanitizationService"],
            ["1 RBAC", "RetrievalScopeService", "CaregiverService.hasAccessToPatient pattern; FileController scope"],
            ["2 Planner", "RetrievalQueryPlanner", "New — medication timeline rules from FR-AI-11"],
            ["3 Hybrid", "HybridRetrievalService", "Replaces PatientContextRetrievalService stub"],
            ["3 Embed query", "EmbeddingService", "New — Bedrock Titan v2 (pgvector doc)"],
            ["4 Context", "RetrievalContextAssembler", "Do NOT dump MedicalContextService vitals dump"],
            ["5 LLM", "GroundedLlmService", "LlmRouter + structured prompt; not BedrockAIChatService raw"],
            ["6 Safety", "SafetyValidationService + TierClassifier", "Extend GuardrailService concepts"],
            ["6 Audit", "AiAuditService", "Extend ChatAuditService metadata model"],
        ],
        highlight_rows={3, 4, 5, 6},
    )

    # 3 Stage 0
    heading(doc, "3. Stage 0 — Gateway Preflight", 1)
    bullets(
        doc,
        [
            "Authenticate JWT → User + Role (Spring Security).",
            "LangChainGovernanceService.checkRateLimit(userId); validateMessageLength(query).",
            "InputSanitizationService.sanitize(query); GuardrailService.sanitizeRequest for SSN patterns.",
            "Assign requestId + auditId; start timing metrics for NFR-AI-1.",
            "Reject if careconnect.ai.ask.enabled=false or USE_AI_FEATURES permission missing.",
        ],
    )

    # 4 Stage 1 RBAC
    heading(doc, "4. Stage 1 — RBAC Filter (RetrievalScopeService)", 1)
    para(doc, "Input: caller User, patientId from request, optional sourceTypes from request.", bold=True)
    para(doc, "Output: immutable RetrievalScope record:", bold=True)
    code(
        doc,
        "record RetrievalScope(\n"
        "    Long callerUserId,\n"
        "    Role callerRole,\n"
        "    Set<Long> allowedPatientIds,      // usually singleton { patientId }\n"
        "    Set<RetrievalRecordType> allowedSourceTypes,\n"
        "    Set<String> excludedSourceTypes,  // REQ-SC-7 user preferences\n"
        "    CaregiverVisibilityFilter visibilityFilter,\n"
        "    boolean consentGranted\n"
        ")",
    )

    heading(doc, "4.1 Resolution rules by role", 2)
    table(
        doc,
        ["Role", "patientId check", "allowedSourceTypes"],
        [
            ["PATIENT", "patientId must equal caller user id", "All except user-excluded (REQ-SC-7)"],
            ["CAREGIVER", "CaregiverPatientLinkService.hasAccessToPatient", "Exclude patient-excluded types; apply visibility"],
            ["FAMILY_MEMBER", "FamilyMemberService.hasAccessToPatient", "Read-only subset per SRS"],
            ["ADMIN", "Policy-defined", "All types for support/debug"],
        ],
    )

    heading(doc, "4.2 Scope SQL predicate (applied in every search)", 2)
    code(
        doc,
        "WHERE patient_id IN (:allowedPatientIds)\n"
        "  AND record_type IN (:allowedSourceTypes)\n"
        "  AND (consent_scope IS NULL OR consent_scope allows :callerRole)\n"
        "  AND (chunk_metadata->>'caregiver_visibility' IS NULL\n"
        "       OR visibilityFilter permits row for :callerRole)",
    )
    para(
        doc,
        "On failure: throw ForbiddenScopeException → HTTP 403 FORBIDDEN_SCOPE; audit with no retrieval.",
        highlight=True,
    )
    para(
        doc,
        "Gap today: scope checks are scattered (FileController.hasAccessToPatient); no unified "
        "RetrievalScopeService or retrieval_index_chunk table.",
        highlight=True,
    )

    # 5 Stage 2 Query planner
    heading(doc, "5. Stage 2 — Query Planner (RetrievalQueryPlanner)", 1)
    para(
        doc,
        "Lightweight intent detection before hybrid search — not a separate LLM call for MVP. "
        "Rule + keyword patterns; optional Bedrock classify later.",
    )
    table(
        doc,
        ["Intent", "Detection signal", "Hybrid strategy"],
        [
            ["MEDICATION_TIMELINE", "medication names, started/stopped/changed/dose", "Structured prefilter on MEDICATION_TIMELINE_EVENT + RRF"],
            ["USPS_MAIL", "mail, package, delivered, tracking", "Restrict record_type USPS_* ; boost FTS on sender"],
            ["DOCUMENT", "lab, report, uploaded, PDF", "Restrict DOCUMENT; wider chunk excerpts"],
            ["EPISODE_VISIT", "home visit, in-person", "episodeType=visit filter"],
            ["EPISODE_CALL", "phone call, video call", "episodeType=call filter"],
            ["GENERAL", "default", "Full hybrid over all allowed source types"],
        ],
        highlight_rows={1},
    )
    code(
        doc,
        "record RetrievalPlan(\n"
        "    QueryIntent intent,\n"
        "    Set<RetrievalRecordType> recordTypeFilter,\n"
        "    Map<String, Object> structuredPrefilters,  // e.g. medication_normalized\n"
        "    String episodeTypeFilter,                   // call | visit | null\n"
        "    int ftsTopK,\n"
        "    int vectorTopK,\n"
        "    int finalTopK\n"
        ")",
    )

    # 6 Stage 3 Hybrid search
    heading(doc, "6. Stage 3 — Hybrid Search (HybridRetrievalService)", 1)

    heading(doc, "6.1 Parallel retrieval arms", 2)
    table(
        doc,
        ["Arm", "Implementation", "Default top-K"],
        [
            ["Structured", "SQL on chunk_metadata (dates, med name, event_type)", "20 (timeline intent only)"],
            ["FTS", "ts_rank_cd(search_vector, plainto_tsquery(:query))", "20"],
            ["Vector", "ORDER BY embedding <=> :queryEmbedding", "20"],
        ],
    )
    para(doc, "All arms share the same RetrievalScope WHERE clause from Stage 1.", bold=True)

    heading(doc, "6.2 Reciprocal Rank Fusion (RRF)", 2)
    code(
        doc,
        "k = 60\n"
        "For each chunk d appearing in any rank list with rank r:\n"
        "  rrf_score(d) += 1 / (k + r)\n"
        "Sort by rrf_score descending; take finalTopK (default 8–12)\n"
        "Intent weights (optional):\n"
        "  MEDICATION_TIMELINE: structured × 1.5 before RRF\n"
        "  GENERAL: FTS × 1.0, vector × 1.0",
    )

    heading(doc, "6.3 RankedChunk output", 2)
    code(
        doc,
        "record RankedChunk(\n"
        "    UUID chunkId,\n"
        "    Long patientId,\n"
        "    RetrievalRecordType recordType,\n"
        "    String sourceRecordId,\n"
        "    String chunkText,\n"
        "    JsonNode chunkMetadata,\n"
        "    double rrfScore,\n"
        "    Integer ftsRank,\n"
        "    Integer vectorRank\n"
        ")",
    )

    heading(doc, "6.4 NO_RECORDS short-circuit", 2)
    para(
        doc,
        "If finalTopK chunks empty after merge: skip Stages 4–5 LLM grounding; return API NO_RECORDS "
        "response (see POST /api/ai/ask contract). Do not call Bedrock for general medical answer.",
        highlight=True,
    )

    heading(doc, "6.5 Query embedding", 2)
    bullets(
        doc,
        [
            "EmbeddingService.embedQuery(sanitizedQuery) using same model as index (Titan v2 1024).",
            "Cache by hash(sessionId + query) for multi-turn within session.",
            "If embed service down: FTS-only degraded mode with logged warning — not silent vector skip.",
        ],
    )

    # 7 Stage 4 Context assembly
    heading(doc, "7. Stage 4 — Context Assembly (RetrievalContextAssembler)", 1)
    para(
        doc,
        "Builds the LLM prompt context from RankedChunk[] under a strict token budget (FR-AI-9). "
        "Does not include full patient chart dumps from MedicalContextService.",
        highlight=True,
    )

    heading(doc, "7.1 Token budget policy", 2)
    table(
        doc,
        ["Budget slice", "Default tokens", "Content"],
        [
            ["System + safety instructions", "800", "Records-only rules, JSON output schema, citation rules"],
            ["User query", "500", "Sanitized question"],
            ["Retrieved context", "3000–4000", "Chunk excerpts only — largest slice"],
            ["Output reserve", "1000", "max_tokens for model response"],
        ],
    )

    heading(doc, "7.2 Context block format (sent to LLM)", 2)
    code(
        doc,
        "<retrieved_records>\n"
        "<record ref=\"C1\" type=\"MEDICATION_TIMELINE_EVENT\" date=\"2026-03-12\">\n"
        "Started lisinopril 10mg daily.\n"
        "</record>\n"
        "<record ref=\"C2\" type=\"CALL_SUMMARY\" date=\"2026-06-20\">\n"
        "Headline: Follow-up on blood pressure medications\n"
        "Assessment: ...\n"
        "</record>\n"
        "</retrieved_records>\n\n"
        "User question: {query}\n\n"
        "Respond with JSON: { answerText, citationRefs: [\"C1\",\"C2\"] }",
    )
    bullets(
        doc,
        [
            "Assign stable refs C1..Cn mapped to RankedChunk for post-LLM AiCitation assembly.",
            "Truncate chunkText per chunk to ~400 tokens; prefer metadata excerpt over full chunk when long.",
            "Deduplicate overlapping chunks (same sourceRecordId + similar text).",
            "Medication timeline: sort context blocks by effectiveDate ascending before packing.",
            "Never include chunks that failed visibility filter — RBAC re-check optional defense-in-depth.",
        ],
    )

    heading(doc, "7.3 MedicalContextService — explicit non-use", 2)
    para(
        doc,
        "MedicalContextService.buildPatientContext() loads vitals, meds, notes from OLTP tables as a "
        "broad chart dump. For Ask AI MVP, orchestration uses retrieval_index_chunk only. "
        "Future: optionally append active medication list snapshot as a single structured block if "
        "not already in retrieved chunks (FR-AI-11 footer).",
        highlight=True,
    )

    heading(doc, "7.4 GroundedContext output", 2)
    code(
        doc,
        "record GroundedContext(\n"
        "    String promptPayload,\n"
        "    Map<String, RankedChunk> refMap,     // C1 → chunk\n"
        "    int chunksUsed,\n"
        "    int chunksRetrieved,\n"
        "    int estimatedPromptTokens\n"
        ")",
    )

    # 8 Stage 5 LLM
    heading(doc, "8. Stage 5 — LLM Inference (GroundedLlmService)", 1)

    heading(doc, "8.1 Invocation", 2)
    bullets(
        doc,
        [
            "LlmRouter.invoke(GroundedLlmRequest) — Bedrock primary (Nova Lite / Claude profile).",
            "Structured output only: JSON schema { answerText, citationRefs[] } — no free-form prose.",
            "System prompt: answer ONLY from <retrieved_records>; if insufficient say so in answerText.",
            "Temperature low (0.2–0.4) for factual retrieval tasks.",
            "Record model provider + modelId in retrievalMeta for audit.",
        ],
    )

    heading(doc, "8.2 Prompt guardrails (pre-LLM)", 2)
    bullets(
        doc,
        [
            "Forbidden: diagnosis not in records, dosage calculations, treatment recommendations beyond excerpts.",
            "Require citationRefs for every factual claim in answerText.",
            "Locale-aware system copy when request.locale=es-US.",
        ],
    )

    heading(doc, "8.3 GroundedLlmResponse", 2)
    code(
        doc,
        "record GroundedLlmResponse(\n"
        "    String answerText,\n"
        "    List<String> citationRefs,\n"
        "    String rawModelOutput,\n"
        "    LlmProviderInfo providerInfo,\n"
        "    long inferenceLatencyMs\n"
        ")",
    )

    # 9 Stage 6 Post-process
    heading(doc, "9. Stage 6 — Post-Processing (Post-LLM)", 1)
    para(doc, "Not part of retrieval proper but completes orchestration before HTTP response.", bold=True)
    table(
        doc,
        ["Step", "Service", "Action"],
        [
            ["6a", "CitationAssembler", "Map citationRefs → AiCitation[] via refMap + chunk metadata"],
            ["6b", "SafetyValidationService", "Verify claims ⊆ citations; unsupported → 422 or Tier 2"],
            ["6c", "TierClassifier", "Assign tier + triggerCodes (HITL doc Table 8)"],
            ["6d", "HitlService", "If Tier 2: persist held_item; suppress answer to patient"],
            ["6e", "DisclaimerService", "Attach locale disclaimer object (REQ-SC-1)"],
            ["6f", "AiAuditService", "Immutable log: scope, chunk ids, tier, delivery status"],
        ],
        highlight_rows={3, 4},
    )

    # 10 Sequence
    heading(doc, "10. Sequence Diagram (Text)", 1)
    code(
        doc,
        "Client          AiAskController    Orchestrator       ScopeService     HybridService      Assembler       LlmRouter\n"
        "  | POST /ask         |                  |                  |                 |                |              |\n"
        "  |---------------->|                  |                  |                 |                |              |\n"
        "  |                 | ask(req,user)      |                  |                 |                |              |\n"
        "  |                 |----------------->|                  |                 |                |              |\n"
        "  |                 |                  | resolveScope()   |                 |                |              |\n"
        "  |                 |                  |----------------->|                 |                |              |\n"
        "  |                 |                  |<-- RetrievalScope|                 |                |              |\n"
        "  |                 |                  | plan(query)      |                 |                |              |\n"
        "  |                 |                  | search(scope,plan)                |                |              |\n"
        "  |                 |                  |---------------------------------->|                |              |\n"
        "  |                 |                  |<-- RankedChunk[] |                 |                |              |\n"
        "  |                 |                  | assemble(chunks) |                 |                |              |\n"
        "  |                 |                  |-------------------------------------------------->|              |\n"
        "  |                 |                  |<-- GroundedContext                |                |              |\n"
        "  |                 |                  | invoke(context)  |                 |                |              |\n"
        "  |                 |                  |----------------------------------------------------------------->|\n"
        "  |                 |                  |<-- GroundedLlmResponse            |                |              |\n"
        "  |                 |                  | postProcess()    |                 |                |              |\n"
        "  |                 |<-- AiAskResponse |                  |                 |                |              |\n"
        "  |<-- 200 JSON ----|                  |                  |                 |                |              |",
    )

    # 11 Interfaces
    heading(doc, "11. Proposed Java Package Structure", 1)
    code(
        doc,
        "com.careconnect.service.ai.ask\n"
        "  AiAskOrchestrator.java\n"
        "  AiAskGatewayPreflight.java\n"
        "com.careconnect.service.ai.retrieval\n"
        "  RetrievalScopeService.java\n"
        "  RetrievalScope.java\n"
        "  RetrievalQueryPlanner.java\n"
        "  HybridRetrievalService.java\n"
        "  EmbeddingService.java\n"
        "  RetrievalContextAssembler.java\n"
        "  CitationAssembler.java\n"
        "  RankedChunk.java\n"
        "com.careconnect.service.ai.grounded\n"
        "  GroundedLlmService.java\n"
        "com.careconnect.repository\n"
        "  RetrievalIndexChunkRepository.java  // native SQL for FTS + vector",
    )

    heading(doc, "11.1 AiAskOrchestrator pseudocode", 2)
    code(
        doc,
        "public AiAskResponse ask(AiAskRequest req, User caller) {\n"
        "  preflight(req, caller);\n"
        "  RetrievalScope scope = scopeService.resolve(caller, req.patientId(), req.sourceTypes());\n"
        "  RetrievalPlan plan = planner.plan(req.query(), req.episodeType(), scope);\n"
        "  List<RankedChunk> chunks = hybridService.search(scope, plan, req.query());\n"
        "  if (chunks.isEmpty()) return noRecordsResponse(req, scope);\n"
        "  GroundedContext ctx = assembler.assemble(chunks, req.query(), plan);\n"
        "  GroundedLlmResponse llm = groundedLlm.generate(ctx, req.locale());\n"
        "  return postProcessor.buildResponse(req, scope, ctx, llm);\n"
        "}",
    )

    # 12 Config
    heading(doc, "12. Configuration Properties", 1)
    table(
        doc,
        ["Property", "Default", "Stage"],
        [
            ["careconnect.ai.ask.enabled", "false", "0"],
            ["careconnect.ai.ask.retrieval.fts-top-k", "20", "3"],
            ["careconnect.ai.ask.retrieval.vector-top-k", "20", "3"],
            ["careconnect.ai.ask.retrieval.final-top-k", "10", "3"],
            ["careconnect.ai.ask.retrieval.rrf-k", "60", "3"],
            ["careconnect.ai.ask.context.max-tokens", "3500", "4"],
            ["careconnect.ai.ask.context.max-chunk-tokens", "400", "4"],
            ["careconnect.ai.embed.model-id", "amazon.titan-embed-text-v2:0", "3"],
            ["careconnect.ai.ask.llm.model-id", "amazon.nova-lite-v1:0", "5"],
        ],
    )

    # 13 Failure modes
    heading(doc, "13. Failure Modes & Degradation", 1)
    table(
        doc,
        ["Failure", "Behavior", "Rationale"],
        [
            ["Scope denied", "403, no LLM", "FR-AI-1 fail closed"],
            ["Index DB down", "503 RETRIEVAL_UNAVAILABLE", "No ungrounded answers"],
            ["Embed service down", "FTS-only with audit flag", "Partial semantic; logged"],
            ["Bedrock down", "503 after retries", "LlmRouter failover within BAA only"],
            ["Empty retrieval", "NO_RECORDS response", "FR-AI-2 no-records path"],
            ["Safety fail post-LLM", "422 or Tier 2 hold", "REQ-SC-4"],
        ],
        highlight_rows={1, 2, 4},
    )

    # 14 Gap analysis
    heading(doc, "14. Codebase Gap Analysis", 1)
    table(
        doc,
        ["Existing component", "Current behavior", "Orchestration gap"],
        [
            ["PatientContextRetrievalService", "In-memory substring", "Replace with HybridRetrievalService"],
            ["MedicalContextService", "OLTP chart dump for chat", "Not used for Ask AI retrieval path"],
            ["BedrockAIChatService", "Raw message to Bedrock", "Replace with GroundedLlmService path"],
            ["CaregiverService.hasAccessToPatient", "Per-resource checks", "Centralize in RetrievalScopeService"],
            ["GuardrailService", "Phrase block only", "Integrate in pre/post; not retrieval"],
            ["ChatAuditService", "Metadata only", "Extend for chunk ids + tier + scope"],
            ["retrieval_index_chunk", "Not in Flyway", "Required for Stage 3"],
            ["AiAskOrchestrator", "Missing", "Core implementation target"],
        ],
        highlight_rows={1, 3, 7, 8},
    )

    # 15 Testing
    heading(doc, "15. Test Strategy", 1)
    table(
        doc,
        ["Test layer", "Focus"],
        [
            ["Unit: RetrievalScopeService", "Role matrix, REQ-SC-7 exclusions, visibility"],
            ["Unit: HybridRetrievalService", "RRF merge with mock ranks; scope SQL injected"],
            ["Unit: RetrievalContextAssembler", "Token budget, ref mapping, dedup"],
            ["Integration: Orchestrator", "End-to-end with Testcontainers pgvector + mock Bedrock"],
            ["Contract: AiAskResponse", "Citations match refMap; NO_RECORDS skips LLM (verify mock not called)"],
        ],
    )

    # 16 Related
    heading(doc, "16. Related Documents", 1)
    bullets(
        doc,
        [
            "docs/POST_api_ai_ask_Request_Response_Contract_Design.docx",
            "docs/RBAC_Scoped_Retrieval_Source_Types.docx",
            "docs/Hybrid_Retrieval_Scope_PostgreSQL_FullText_pgvector.docx",
            "docs/pgvector_Embedding_Strategy_Summaries_Mail_Documents.docx",
            "docs/HITL_Tier1_Tier2_Escalation_and_Hold_Release_Workflow.docx",
            "docs/Medication_Timeline_Retrieval_FR-AI-11.docx",
            "docs/Storage_Audit_vs_Shared_Index_Schema.docx",
            "docs/Team_E_Implementation_Task_Backlog.docx (tasks 2.1, 5.1–5.4)",
        ],
    )

    heading(doc, "17. Conclusion", 1)
    para(
        doc,
        "Retrieval orchestration is a six-stage pipeline with a single orchestrator entry point. "
        "RBAC filtering happens before any SQL against retrieval_index_chunk. Hybrid search combines "
        "structured, FTS, and vector arms under a unified scope predicate. Context assembly packs "
        "only ranked chunk excerpts into a citation-typed prompt block. The LLM receives grounded "
        "context and returns structured JSON — never the raw user message alone.",
    )
    para(
        doc,
        "Implementation order: RetrievalScopeService + retrieval_index_chunk → HybridRetrievalService "
        "→ RetrievalContextAssembler → GroundedLlmService → AiAskOrchestrator wiring to "
        "POST /api/ai/ask.",
        highlight=True,
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    try:
        doc.save(OUTPUT)
        print(f"Created: {OUTPUT}")
    except PermissionError:
        alt = OUTPUT.with_name(f"{OUTPUT.stem}_refresh{OUTPUT.suffix}")
        doc.save(alt)
        print(f"Created (original locked): {alt}")


if __name__ == "__main__":
    build()
