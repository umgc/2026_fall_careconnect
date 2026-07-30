"""Generate Word document: Hybrid Retrieval Scope (PostgreSQL full-text + pgvector)."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.shared import Pt
from docx_codebase_revision import REVISION_BULLETS, REVISION_LABEL, save_document

OUTPUT = (
    Path(__file__).resolve().parents[1]
    / "docs"
    / "Hybrid_Retrieval_Scope_PostgreSQL_FullText_pgvector.docx"
)

SOURCE_DOCS = [
    "CareConnect_SRS_Revision 2.0_TEAM E.docx",
    "CareConnect_Milestone_2_Software_Test_Plan (2).docx",
    "CareConnect_Milestone_2_TDD_TEAM E.docx",
    "CareConnect_ProjectPlan_Revised (1).docx",
]


def heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def para(doc: Document, text: str, bold: bool = False, highlight: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    if highlight:
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def code(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val


def build() -> None:
    doc = Document()

    title = doc.add_heading(
        "Hybrid Retrieval Scope Documentation", 0
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    para(
        doc,
        "PostgreSQL Full-Text Search + pgvector Semantic Search",
    )
    para(doc, "CareConnect Team E — AI Services & Intelligent Recall")
    para(
        doc,
        "Derived from Team E SRS v2.0, Milestone 2 TDD, Software Test Plan, Project Plan, "
        "and the CareConnect application codebase.",
    )
    para(doc, REVISION_LABEL, highlight=True)
    bullets(doc, REVISION_BULLETS)
    doc.add_paragraph()

    heading(doc, "1. Document Purpose and Sources", 1)
    bullets(
        doc,
        [
            "Purpose: Define the scope, architecture, indexed record types, and upstream "
            "dependencies for hybrid retrieval in CareConnect Ask AI and related workstreams.",
            "Primary requirement source: SRS Section 2.5 (Hybrid retrieval constraint) and "
            "Section 3 (AI-Assisted Retrieval).",
            "Design source: Milestone 2 TDD Sections 3.3, 5.2, 7.x (shared schema, retrieval service).",
            "Validation source: Milestone 2 Software Test Plan (TC-E-AI-* coverage, FR-AI mapping).",
            "Schedule source: Project Plan WBS 3.2 (AI Agent retrieval service & LLM integration).",
        ],
    )
    para(doc, "Reference documents reviewed:", bold=True)
    bullets(doc, [f"C:\\Users\\ravic\\Downloads\\{name}" for name in SOURCE_DOCS])

    heading(doc, "2. Executive Summary", 1)
    para(
        doc,
        "Team E specifies a records-grounded Ask AI capability that answers natural-language "
        "questions strictly from indexed patient records. Retrieval must combine three mechanisms: "
        "(1) structured PostgreSQL queries for precise filters, (2) PostgreSQL full-text search "
        "for keyword and phrase matching, and (3) pgvector semantic search for conversational "
        "and paraphrased questions. Generative output is assembled only after scoped retrieval "
        "and must include citations, disclaimers, and RBAC enforcement on every request.",
    )
    para(
        doc,
        "The shared data layer (Aurora PostgreSQL + pgvector in production design; local "
        "pgvector/pg15 Docker in development) holds metadata, embeddings, and searchable text. "
        "Encrypted S3 holds large artifacts: transcripts, OCR text, and uploaded documents. "
        "Index refresh target: within 5 minutes of a new transcript, summary, or document. "
        "End-to-end Ask AI latency target: 5 seconds p95.",
    )

    heading(doc, "3. Hybrid Retrieval Definition", 1)
    heading(doc, "3.1 SRS Definition", 2)
    para(
        doc,
        "SRS Section 2.5 states: 'Hybrid retrieval — A structured database, full-text search, "
        "and vector search work together so the agent can answer both precise and conversational questions.'",
    )
    heading(doc, "3.2 TDD Retrieval Pattern", 2)
    bullets(
        doc,
        [
            "Retrieve scoped context → call model through abstraction layer → validate JSON schema → "
            "apply safety pass → persist with citations and audit entry.",
            "Orchestration forces the model to confine answers to supplied records only.",
            "Sequence (TDD Figure 3.6): mobile app → JWT-authenticated request → retrieval service "
            "performs hybrid search over Aurora pgvector → Bedrock returns grounded answer.",
            "SCC-3 control sequence: RBAC → consent → retrieval → validate/cite → HITL → audit.",
        ],
    )
    heading(doc, "3.3 Three-Layer Search Model", 2)
    table(
        doc,
        ["Layer", "Technology", "Best for", "Example queries"],
        [
            [
                "Structured SQL",
                "PostgreSQL relational filters + indexes",
                "Exact filters, timelines, IDs, dates, record type",
                "Summaries from last week; medications with status changed",
            ],
            [
                "Full-text (FTS)",
                "PostgreSQL tsvector / tsquery (English + Spanish)",
                "Keywords, sender names, exact phrases, medication names",
                "blood pressure medication; student loan mail",
            ],
            [
                "Semantic (vector)",
                "pgvector cosine similarity on embeddings",
                "Paraphrase, conversational intent, fuzzy recall",
                "What did the nurse say about my pills?",
            ],
        ],
    )
    doc.add_paragraph()
    para(
        doc,
        "Hybrid merge: TDD and SRS describe merging keyword matches over structured fields with "
        "semantic similarity over embeddings, then ranking by combined relevance. Results are "
        "framed as records-based likely matches, not guaranteed exact retrieval (SRS known limitation).",
    )

    heading(doc, "4. Scope Boundaries", 1)
    heading(doc, "4.1 In Scope", 2)
    bullets(
        doc,
        [
            "Ask AI retrieval from indexed patient records (transcripts, summaries, documents, "
            "caregiver notes, care instructions, USPS mail records).",
            "Hybrid keyword + semantic search for USPS mail (FR-USPS-4).",
            "Indexing summary fields for Ask AI (FR-SUM-7).",
            "Async embedding generation within 5-minute index refresh (NFR-AI-3).",
            "RBAC and row-level patient scoping on every retrieval request (FR-AI-1).",
            "Minimum-necessary context sent to Bedrock (FR-AI-9).",
            "STML recall and Daily Memory Brief consuming the same retrieval index.",
        ],
    )
    heading(doc, "4.2 Out of Scope", 2)
    bullets(
        doc,
        [
            "Clinical decision-making, diagnosis, treatment, or dosage recommendations.",
            "Answers from general model knowledge for medical or care-related prompts.",
            "Raw audio persistence in retrieval or summary layers.",
            "Real-time in-call summarization or retrieval (upstream transcript dependency).",
            "EHR, pharmacy, or external provider-system integration.",
            "Automatic writes (calendar, reminders, care plan) from retrieval answers without confirmation flow.",
        ],
    )

    heading(doc, "5. Indexed Record Types", 1)
    para(
        doc,
        "The shared schema (TDD Section 7.1) defines entities indexed for hybrid retrieval. "
        "Each indexed chunk must carry patient scope, record type, source record ID, excerpt text, "
        "optional deep link, and embedding vector for semantic search.",
    )
    table(
        doc,
        ["Record type", "Source workstream", "Structured fields", "FTS target", "Vector target", "SRS/TDD ref"],
        [
            [
                "Transcript segment / chunk",
                "Call/Visit Summaries (upstream)",
                "call_id, speaker, timestamps, session_id",
                "transcript_text",
                "chunk embedding",
                "SRS §3, §4; TDD UC-SUM-5",
            ],
            [
                "Call summary",
                "Summaries",
                "headline, narrative, status, generated_at",
                "summary_json text fields",
                "summary embedding",
                "FR-SUM-7",
            ],
            [
                "Summary item",
                "Summaries",
                "action_items, appointments, care_instructions",
                "item text + citation",
                "item embedding",
                "TDD shared schema",
            ],
            [
                "Uploaded / scanned document",
                "Platform + OCR pipeline",
                "document_id, mime, upload date",
                "OCR text (S3 or inline)",
                "document chunk embeddings",
                "SRS §3 retrieval sources",
            ],
            [
                "Caregiver note / care instruction",
                "Platform",
                "author, patient_id, created_at",
                "note body",
                "note embedding",
                "SRS §3",
            ],
            [
                "USPS mail piece",
                "USPS Mail Agent",
                "sender, delivery_date, tier, image flag",
                "visible_text, sender",
                "mail embedding",
                "FR-USPS-4",
            ],
            [
                "Medication timeline event",
                "Ask AI (derived index)",
                "medication, event_type (start/stop/change)",
                "event description",
                "event embedding",
                "FR-AI-11",
            ],
        ],
    )
    doc.add_paragraph()
    heading(doc, "5.1 Planned Index Table Shape (TDD)", 2)
    code(
        doc,
        "-- Conceptual retrieval index (Milestone 3 target)\n"
        "CREATE TABLE retrieval_index_chunk (\n"
        "  id UUID PRIMARY KEY,\n"
        "  patient_id BIGINT NOT NULL,\n"
        "  record_type VARCHAR(40) NOT NULL,\n"
        "  source_record_id VARCHAR(120) NOT NULL,\n"
        "  chunk_text TEXT NOT NULL,\n"
        "  chunk_metadata JSONB,\n"
        "  search_vector TSVECTOR,\n"
        "  embedding vector(1536),\n"
        "  indexed_at TIMESTAMPTZ NOT NULL DEFAULT now(),\n"
        "  consent_scope VARCHAR(40)\n"
        ");\n"
        "CREATE INDEX idx_retrieval_patient ON retrieval_index_chunk(patient_id);\n"
        "CREATE INDEX idx_retrieval_fts ON retrieval_index_chunk USING GIN(search_vector);\n"
        "CREATE INDEX idx_retrieval_vector ON retrieval_index_chunk USING ivfflat(embedding vector_cosine_ops);",
    )

    heading(doc, "6. Upstream Dependencies", 1)
    table(
        doc,
        ["Upstream producer", "Trigger", "Indexed output", "Refresh SLA"],
        [
            [
                "Transcript ingest (platform / Chime / Transcribe)",
                "Transcript saved post-call",
                "Transcript chunks + FTS + embeddings",
                "≤ 5 min (NFR-AI-3)",
            ],
            [
                "Summary pipeline (Bedrock)",
                "UC-SUM-5 persist + SNS indexing event",
                "Summary fields + item chunks",
                "≤ 5 min",
            ],
            [
                "Document OCR (Textract / platform)",
                "Document uploaded / OCR complete",
                "Document text chunks",
                "≤ 5 min",
            ],
            [
                "USPS mail ingest (Gmail OAuth)",
                "Mail normalized to canonical schema",
                "Mail record + embedding",
                "≤ 5 min",
            ],
            [
                "Outbox poller / embedding pipeline",
                "SNS event from summary/transcript save",
                "Embedding write to pgvector",
                "≤ 5 min",
            ],
        ],
    )
    doc.add_paragraph()
    para(
        doc,
        "Critical upstream note (Project Plan §3.2.1): live in-call summarization, audio capture, "
        "speech-to-text, and diarization are upstream dependencies owned outside Team E. "
        "Retrieval and summary layers operate on transcript text only after conversation ends.",
    )

    heading(doc, "7. Retrieval Service Flow", 1)
    heading(doc, "7.1 Ask AI Request Flow (Planned)", 2)
    bullets(
        doc,
        [
            "1. Client POST /api/ai/ask with { query, sessionId } (JWT authenticated).",
            "2. AI Gateway: rate limit, RBAC scope filter, consent check.",
            "3. Retrieval Service: resolve permitted patient_id set for caller.",
            "4. Parallel retrieval: structured prefilter → FTS rank → pgvector similarity.",
            "5. Merge and rank top-K chunks (hybrid score).",
            "6. Assemble prompt with minimum-necessary context only.",
            "7. Bedrock inference via abstraction layer.",
            "8. Schema + safety validation; withhold on failure (HTTP 422).",
            "9. Return answer + citations[] + disclaimer + escalation flag.",
            "10. Audit log: query, response, sources, escalation, delivery status (FR-AI-10).",
        ],
    )
    heading(doc, "7.2 USPS Mail Hybrid Search (Planned)", 2)
    para(
        doc,
        "SRS FR-USPS-4 and TDD describe GET /api/mail with hybrid keyword + semantic search. "
        "Keyword match runs over structured sender/date/text fields; vector similarity runs over "
        "mail embeddings; results merged and ranked by combined relevance.",
    )

    heading(doc, "8. Security, RBAC, and Consent Scope", 1)
    bullets(
        doc,
        [
            "Every retrieval request enforces RBAC at gateway and service layer (FR-AI-1).",
            "Row-level patient scoping via foreign keys prevents cross-patient retrieval at DB layer (TDD §7.1).",
            "Caregiver access governed by ConsentGrant entity and visibility gates (SRS §7).",
            "Users may exclude source types from indexing (REQ-SC-7).",
            "High-risk queries routed to HITL before delivery when safety rules trigger (FR-AI-5).",
            "Only minimum-necessary retrieved chunks sent to Bedrock (FR-AI-9).",
        ],
    )

    heading(doc, "9. Non-Functional Requirements", 1)
    table(
        doc,
        ["ID", "Requirement", "Target"],
        [
            ["NFR-AI-1", "End-to-end retrieval response latency", "≤ 5 seconds (p95)"],
            ["NFR-AI-2", "Gateway output schema validation", "≤ 500 ms"],
            ["NFR-AI-3", "Index refresh after new transcript/summary/document", "≤ 5 minutes (async OK)"],
            ["NFR-PRV-1", "User-controlled indexing, retention, export, deletion", "Policy-configurable"],
        ],
    )

    heading(doc, "10. Test Plan Alignment", 1)
    para(
        doc,
        "Milestone 2 Software Test Plan defines TC-E-AI-* cases for Ask AI retrieval, mapped to FR-AI-1 "
        "through FR-AI-11. Test areas: SUM, AI, USPS, STML, SC (Safety/Consent), SEC (Security/NFR). "
        "Automated now: unit tests for extraction, classification, schema enforcement. "
        "Integration tests use mock LLM/mail data until Bedrock and Gmail OAuth are provisioned.",
    )
    bullets(
        doc,
        [
            "TC-E-AI-001 area: authorized retrieval only from permitted records.",
            "TC-E-AI-002 area: citations present on every response.",
            "TC-E-AI-003 area: disclaimer and non-medical-advice framing.",
            "TC-E-AI-005 area: high-risk query escalation (HITL).",
            "TC-E-AI-011 area: medication initiation + termination timeline retrieval.",
            "TC-E-USPS-* area: hybrid mail search keyword + semantic.",
        ],
    )

    heading(doc, "11. Current Codebase Alignment (Gap Analysis)", 1)
    para(
        doc,
        "This section maps Team E design documents to the CareConnect repository as of the current codebase review.",
        bold=True,
    )
    heading(doc, "11.1 Infrastructure Ready", 2)
    bullets(
        doc,
        [
            "PostgreSQL Docker image: pgvector/pgvector:pg15 (backend/core/pg_docker/docker-compose.yml).",
            "Init script enables extensions: uuid-ossp and vector (pg_docker/init-scripts/01-init-database.sh).",
            "H2 test profile stubs TSQUERY/TSVECTOR domains for unit tests (application-test.properties).",
        ],
    )
    heading(doc, "11.2 Partially Implemented (Related but Not Hybrid)", 2)
    table(
        doc,
        ["Component", "Location", "Current behavior", "Hybrid gap"],
        [
            [
                "Call transcript storage",
                "CallTranscriptService, call_transcript_segments",
                "Live DB + S3 archive merge for reads",
                "No FTS/pgvector index; not wired to Ask AI",
            ],
            [
                "Call summaries",
                "CallSummaryService, call_summaries",
                "Bedrock summary generation + storage",
                "FR-SUM-7 indexing pipeline not implemented",
            ],
            [
                "USPS digest search",
                "USPSDigestService.search()",
                "Keyword scan of cached digest JSON",
                "No pgvector; no hybrid rank",
            ],
            [
                "AI chat",
                "BedrockAIChatService + BedrockAIChatAdapter",
                "Opt-in (careconnect.ai.enabled); raw message to Bedrock — no medical context, no RAG",
                "No retrieval layer; context injection not wired",
            ],
            [
                "Medical context builder",
                "MedicalContextService",
                "Builds vitals/meds/notes but unused by active Bedrock chat path",
                "Must re-wire for grounded answers",
            ],
            [
                "Patient context retrieval",
                "PatientContextRetrievalService",
                "In-memory substring match only",
                "Placeholder; no PostgreSQL or embeddings",
            ],
        ],
    )
    heading(doc, "11.4 Recent AI / Bedrock Changes (PR #88)", 2)
    bullets(
        doc,
        [
            "careconnect.ai.provider=bedrock; default model amazon.nova-lite-v1:0 (Claude Sonnet 4.5 commented alternate).",
            "BedrockModelSupport: approved model allowlist, Claude inference-profile ID normalization, Nova vs Claude payload builders.",
            "AIServiceFactory selects BedrockAIChatService; DefaultAIChatService (DeepSeek) throws if invoked.",
            "AIChatController @ConditionalOnProperty careconnect.ai.enabled=true — entire /v1/api/ai-chat API absent when disabled.",
            "README documents opt-in via CARECONNECT_AI_ENABLED=true and AWS credentials for local Bedrock testing.",
        ],
    )

    heading(doc, "11.3 Not Yet Implemented (Milestone 3 Target)", 2)
    bullets(
        doc,
        [
            "retrieval_index_chunk table and Flyway migration.",
            "Embedding generation pipeline (SNS/outbox → embedding worker).",
            "PostgreSQL full-text indexing (to_tsvector triggers or batch indexer).",
            "Hybrid rank merge service (FTS score + vector score + structured filters).",
            "POST /api/ai/ask endpoint and retrieval service microservice.",
            "Citation assembly with record type, ID, excerpt, deep link.",
            "Medication timeline derived index (FR-AI-11).",
            "Index refresh monitoring and 5-minute SLA alarms.",
        ],
    )

    heading(doc, "12. Milestone Roadmap", 1)
    table(
        doc,
        ["Milestone", "Hybrid retrieval deliverable", "Status"],
        [
            ["M1 — Initiation", "SRS requirements FR-AI-*, NFR-AI-* defined", "Complete (SRS v2.0)"],
            ["M2 — Design & Test Plan", "TDD architecture, shared schema, test matrix", "Complete (TDD v1.0)"],
            ["M3 — Implementation", "Retrieval service, index pipeline, /api/ai/ask", "Planned (WBS 3.2)"],
            ["M4 — Validation", "E2E retrieval tests, accessibility, safety sign-off", "Planned"],
        ],
    )

    heading(doc, "13. Recommendations for Milestone 3 Implementation", 1)
    bullets(
        doc,
        [
            "Define and lock retrieval_index_chunk shared schema before parallel workstreams diverge (TDD R4 mitigation).",
            "Implement indexing hooks on existing CallTranscriptService.recordSegments and CallSummaryService.persistResponse.",
            "Use Bedrock Titan or configured embedding model; store vectors in pgvector with patient_id partition strategy.",
            "Add GIN index on tsvector column; use plainto_tsquery for user queries with English/Spanish config.",
            "Implement reciprocal rank fusion (RRF) or weighted hybrid scoring for FTS + vector merge.",
            "Replace PatientContextRetrievalService in-memory stub with PostgreSQL-backed hybrid retrieval.",
            "Wire BedrockAIChatService through retrieval layer before model invocation (RAG pattern from TDD).",
            "Extend USPSDigestService.search to query indexed mail_piece table once schema exists.",
            "Add integration tests TC-E-AI-* against mock embeddings before Bedrock provisioning completes.",
        ],
    )

    heading(doc, "14. Key Repository Files", 1)
    bullets(
        doc,
        [
            "backend/core/pg_docker/docker-compose.yml — pgvector-enabled PostgreSQL",
            "backend/core/pg_docker/init-scripts/01-init-database.sh — vector extension",
            "backend/core/src/main/java/com/careconnect/service/PatientContextRetrievalService.java — retrieval stub",
            "backend/core/src/main/java/com/careconnect/service/BedrockAIChatService.java — AI without RAG",
            "backend/core/src/main/java/com/careconnect/service/CallTranscriptService.java — transcript upstream source",
            "backend/core/src/main/java/com/careconnect/service/CallSummaryService.java — summary upstream source",
            "backend/core/src/main/java/com/careconnect/service/USPSDigestService.java — mail keyword search",
            "backend/core/src/main/resources/db/migration/V52/V61 — transcript and summary tables",
            "docs/Call_Transcript_Retrieval_Review.docx — related transcript retrieval review",
        ],
    )

    heading(doc, "15. Conclusion", 1)
    para(
        doc,
        "Hybrid retrieval scope for CareConnect Team E is clearly defined across SRS, TDD, Project Plan, "
        "and Test Plan: PostgreSQL structured queries, full-text search, and pgvector semantic search "
        "must work together to ground Ask AI, STML recall, and USPS mail queries in authorized patient "
        "records. The development environment already provisions pgvector, but the retrieval index schema, "
        "embedding pipeline, hybrid rank merge, and Ask AI API remain Milestone 3 implementation work. "
        "Existing transcript and summary services provide the primary upstream content sources once indexing "
        "hooks and the shared retrieval_index_chunk contract are implemented.",
    )

    save_document(doc, OUTPUT)


if __name__ == "__main__":
    build()
