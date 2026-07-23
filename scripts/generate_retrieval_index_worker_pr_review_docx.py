"""Generate Word document: PR code review for Task 4.1 RetrievalIndexService / IndexWorker."""

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.shared import Pt

OUTPUT = (
    Path(__file__).resolve().parents[1]
    / "docs"
    / "PR_Review_Retrieval_Index_Service_Index_Worker_feature_a-rvasireddy.docx"
)


def heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def para(doc: Document, text: str, bold: bool = False, highlight: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    if highlight:
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def code(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    for i, h in enumerate(headers):
        tbl.rows[0].cells[i].text = h
    for row in rows:
        cells = tbl.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val


def build() -> Document:
    doc = Document()
    title = doc.add_heading("PR Code Review", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    para(
        doc,
        "Task 4.1 — RetrievalIndexService / IndexWorker — "
        "feature/a-rvasireddy-retrieval-index-service-index-worker → team-ae-develop",
    )
    para(doc, f"Review date: {date.today().isoformat()}")
    para(doc, "Reviewer: AI-assisted code review (Cursor)")
    para(
        doc,
        "Note: At review time the Task 4.1 implementation was present on the branch working "
        "tree but not yet committed. Review covers the working-tree delta vs team-ae-develop.",
        highlight=True,
    )
    doc.add_paragraph()

    heading(doc, "Review metadata", 1)
    table(
        doc,
        ["Field", "Value"],
        [
            ["Source branch", "feature/a-rvasireddy-retrieval-index-service-index-worker"],
            ["Target branch", "team-ae-develop"],
            ["Scope", "Task 4.1 indexing pipeline (service, worker, chunkers, tests, config)"],
            ["Commits on branch vs base", "None yet (uncommitted working tree)"],
            ["Primary packages", "com.careconnect.service.ai.indexing (+ chunker)"],
        ],
    )

    # ── 1. Change Summary ─────────────────────────────────────────────────────
    heading(doc, "1. Change Summary", 1)
    para(
        doc,
        "This change implements Team E backlog Task 4.1: consume SUMMARY_CREATED and "
        "TRANSCRIPT_INDEXED rows from the transactional indexing_outbox and write searchable "
        "chunks into retrieval_index_chunk. It introduces RetrievalIndexService (ingest), "
        "IndexWorker (scheduled outbox poller), SummaryChunker / TranscriptSegmentChunker, "
        "IndexingEventType constants, repository deleteBySourceRecordId, and unit tests.",
    )

    heading(doc, "What changed", 2)
    table(
        doc,
        ["File / area", "Change"],
        [
            [
                "RetrievalIndexService",
                "ingestSummaryCreated / ingestTranscriptIndexed: load source rows, chunk, "
                "delete-by-source, saveAll; contentHash skip; patientId required",
            ],
            [
                "IndexWorker",
                "Scheduled poller (EVV-style MVP); dispatch by event type; attempt/dead-letter",
            ],
            [
                "SummaryChunker",
                "Overview CALL_SUMMARY/VISIT_SUMMARY + action/appointment/care/SOAP/observation chunks",
            ],
            [
                "TranscriptSegmentChunker",
                "One segment → one TRANSCRIPT_SEGMENT chunk",
            ],
            [
                "RetrievalIndexChunkRepository",
                "Adds deleteBySourceRecordId for summary re-index across record types",
            ],
            [
                "application.properties",
                "careconnect.indexing.worker/outbox settings; disabled in test profile",
            ],
            [
                "Tests",
                "Chunker, RetrievalIndexService, IndexWorker unit tests",
            ],
        ],
    )

    heading(doc, "Verdict", 2)
    para(
        doc,
        "Request changes before merge. The overall design matches the outbox → worker → "
        "retrieval_index_chunk contract and is well structured, but there is a high-severity "
        "delete-then-empty-draft data-loss bug, fragile contentHash matching, multi-instance "
        "outbox races, and visit-summary / null-patientId edge cases that will burn retries.",
        bold=True,
        highlight=True,
    )

    # ── 2. Bug & Risk Analysis ────────────────────────────────────────────────
    heading(doc, "2. Bug & Risk Analysis", 1)

    heading(doc, "2.1 Strengths", 2)
    bullets(
        doc,
        [
            "Correctly keeps emitters (CallSummaryService / CallTranscriptService) decoupled — "
            "worker consumes outbox asynchronously.",
            "Matches existing EvvOutboxProcessor scheduled-poller pattern for MVP (SNS/SQS later).",
            "Re-index strategy uses source_record_id correlation (summaryId / callId).",
            "FTS left to DB trigger; embeddings deferred to Task 4.3 — appropriate scope split.",
            "Unit tests cover happy path, contentHash skip, patientId failure, dead-letter.",
            "Worker disabled in application-test.properties avoids poller noise in Spring tests.",
        ],
    )

    heading(doc, "2.2 High — Delete before confirming drafts are non-empty", 2)
    para(
        doc,
        "ingestSummaryCreated and ingestTranscriptIndexed delete existing chunks before "
        "persistDrafts. If the chunker returns no drafts (blank JSON, all blank segments, "
        "parse failure already returned empty), prior good index rows are wiped and nothing "
        "is written.",
        highlight=True,
    )
    code(
        doc,
        """// Current (unsafe ordering)
chunkRepository.deleteBySourceRecordId(sourceRecordId);
return persistDrafts(patientId, sourceRecordId, drafts);

// Recommended
if (drafts == null || drafts.isEmpty()) {
    log.warn("No drafts for sourceRecordId={}; leaving existing chunks unchanged", sourceRecordId);
    return 0;
}
chunkRepository.deleteBySourceRecordId(sourceRecordId);
return persistDrafts(patientId, sourceRecordId, drafts);""",
    )

    heading(doc, "2.3 High — Multi-instance outbox race (no claim / SKIP LOCKED)", 2)
    para(
        doc,
        "Every ECS task runs IndexWorker (matchIfMissing=true). findUnprocessedForPolling "
        "has no SELECT FOR UPDATE SKIP LOCKED / lease. Two instances can process the same "
        "row concurrently → duplicate chunks or conflicting delete/insert.",
        highlight=True,
    )
    bullets(
        doc,
        [
            "Short-term: process with pessimistic lock query, or claim via UPDATE … WHERE "
            "processed_at IS NULL AND id=? RETURNING, or run worker on a single instance only.",
            "Longer-term: SNS/SQS with visibility timeout as designed in the indexing contract.",
        ],
    )
    code(
        doc,
        """@Query(value = \"\"\"
    SELECT * FROM indexing_outbox
    WHERE processed_at IS NULL AND attempt_count < :maxAttempts
    ORDER BY id ASC
    FOR UPDATE SKIP LOCKED
    LIMIT :limit
    \"\"\", nativeQuery = true)
List<IndexingOutboxRow> claimUnprocessed(@Param("maxAttempts") int maxAttempts,
                                         @Param("limit") int limit);""",
    )

    heading(doc, "2.4 Medium — contentHash match uses String.contains", 2)
    para(
        doc,
        "hasMatchingContentHash checks chunkMetadata.contains(contentHash). A shorter hash "
        "prefix can false-positive match a longer hash substring, incorrectly skipping re-index.",
        highlight=True,
    )
    code(
        doc,
        """private boolean hasMatchingContentHash(String sourceRecordId, String contentHash) {
    // Parse JSON and compare metadata.contentHash with Objects.equals
    for (RetrievalIndexChunk chunk : loadOverviewChunks(sourceRecordId)) {
        JsonNode meta = objectMapper.readTree(chunk.getChunkMetadata());
        if (contentHash.equals(meta.path("contentHash").asText(null))) {
            return true;
        }
    }
    return false;
}""",
    )

    heading(doc, "2.5 Medium — Null patientId burns attempts then dead-letters", 2)
    para(
        doc,
        "TRANSCRIPT_INDEXED often emits patientId=null (CallTranscriptService). Resolution "
        "via latest CallSummary fails when no summary exists yet. IndexWorker treats this as "
        "a hard failure, increments attemptCount, and eventually dead-letters — even though "
        "a later SUMMARY_CREATED might supply patientId. Transcripts for that call never index.",
    )
    bullets(
        doc,
        [
            "Option A: defer — leave processedAt null without incrementing attempts when "
            "error is 'patientId missing' (or use a next_attempt_at column).",
            "Option B: resolve patient from call telemetry / meeting participants.",
            "Option C: require emitters to populate patientId before emit (stronger contract).",
        ],
    )

    heading(doc, "2.6 Medium — Visit SUMMARY_CREATED always loads CallSummary", 2)
    para(
        doc,
        "Payload allows episodeType=visit / sourceTable=visit_summaries, but ingest always "
        "callSummaryRepository.findById. Visit events will throw 'CallSummary not found' "
        "until visit_summaries exists. Either reject visit early with a clear skip, or "
        "branch on sourceTable.",
    )
    code(
        doc,
        """if ("visit_summaries".equalsIgnoreCase(payload.sourceTable())
        || "visit".equalsIgnoreCase(payload.episodeType())) {
    log.warn("VISIT summary indexing not implemented yet; outboxId will retry/dead-letter");
    throw new UnsupportedOperationException("visit_summaries indexing pending Task 1.4");
}""",
    )

    heading(doc, "2.7 Low — Outbox success stamp not in same TX as ingest", 2)
    para(
        doc,
        "processRow is not @Transactional. Ingest commits, then outbox save stamps processedAt. "
        "If the process crashes between them, the row is reprocessed. Summary contentHash skip "
        "and transcript delete-then-insert mostly make this safe — after fixing 2.2. Still "
        "prefer one transaction or an idempotency key (eventId) table.",
    )

    heading(doc, "2.8 Low — Overview duplicates SOAP fields", 2)
    para(
        doc,
        "SummaryChunker puts subjective/assessment/plan into the overview and also emits "
        "SUMMARY_SOAP. Acceptable for recall; may inflate context tokens in Task 5.1. "
        "Consider overview = headline + overallAssessment only.",
    )

    heading(doc, "2.9 Low — Unknown event types marked processed", 2)
    para(
        doc,
        "Unsupported eventType returns 0 and marks processed. Good for poller liveness; "
        "ensure metrics/alerts exist so new event types are not silently dropped.",
    )

    heading(doc, "2.10 Informational — Uncommitted PR", 2)
    para(
        doc,
        "Commit the Task 4.1 files before opening/updating the GitHub PR so reviewers see "
        "a stable diff against team-ae-develop.",
    )

    # ── 3. Architecture & Style ───────────────────────────────────────────────
    heading(doc, "3. Architecture & Style", 1)

    heading(doc, "3.1 Design patterns", 2)
    bullets(
        doc,
        [
            "Transactional outbox + scheduled worker is the right MVP vs synchronous indexing.",
            "Clear separation: chunkers (pure transform) → RetrievalIndexService (persist) → "
            "IndexWorker (orchestration).",
            "IndexingChunkDraft record is a clean intermediate DTO.",
            "ConditionalOnProperty for worker enablement matches other CareConnect toggles.",
            "Uses RetrievalRecordType enum consistently with RBAC scope work.",
        ],
    )

    heading(doc, "3.2 Cleanliness", 2)
    bullets(
        doc,
        [
            "Constructor injection, explicit logging, package layout under service.ai.indexing "
            "is consistent with RetrievalScopeService.",
            "Prefer static imports of RetrievalIndexSchema constants over FQCN in truncate helpers.",
            "IndexWorker.dispatch could use a small strategy map as more event types arrive "
            "(DOCUMENT_OCR_COMPLETE, USPS_DIGEST_INGESTED).",
        ],
    )

    heading(doc, "3.3 Test quality", 2)
    bullets(
        doc,
        [
            "Good unit coverage for chunkers and worker success/failure/dead-letter.",
            "Missing: test that empty drafts do not delete existing chunks (once fixed).",
            "Missing: concurrent claim / SKIP LOCKED integration test (can be follow-up).",
            "Missing: contentHash exact-match negative test (prefix should not skip).",
        ],
    )

    # ── 4. Recommendations ────────────────────────────────────────────────────
    heading(doc, "4. Recommendations", 1)

    heading(doc, "4.1 [High] Only delete after non-empty drafts", 2)
    code(
        doc,
        """@Transactional
public int ingestSummaryCreated(final SummaryCreatedPayload payload) {
    // ... validate, load summary, resolve patientId ...
    final List<IndexingChunkDraft> drafts = summaryChunker.chunk(...);
    if (drafts.isEmpty()) {
        log.warn("SUMMARY_CREATED produced no drafts for summaryId={}; skipping delete",
                payload.summaryId());
        return 0;
    }
    chunkRepository.deleteBySourceRecordId(String.valueOf(payload.summaryId()));
    return persistDrafts(patientId, String.valueOf(payload.summaryId()), drafts);
}""",
    )

    heading(doc, "4.2 [High] Claim outbox rows safely under multi-instance deploy", 2)
    para(
        doc,
        "Add FOR UPDATE SKIP LOCKED (see §2.3) or temporarily set "
        "careconnect.indexing.worker.enabled=true on only one service instance / use ShedLock.",
    )

    heading(doc, "4.3 [Medium] Exact contentHash comparison", 2)
    para(doc, "See §2.4 code block — parse JSON metadata; do not use String.contains.")

    heading(doc, "4.4 [Medium] Soft-fail missing patientId for transcripts", 2)
    code(
        doc,
        """} catch (IllegalStateException ex) {
    if (ex.getMessage() != null && ex.getMessage().contains("patientId is required")) {
        // Do not increment toward dead-letter; leave unprocessed for a later poll
        // after SUMMARY_CREATED or telemetry backfill supplies patientId.
        log.warn("Deferring outboxId={}: {}", row.getId(), ex.getMessage());
        return;
    }
    row.setAttemptCount(attempts + 1);
    row.setLastError(truncate(ex.getMessage()));
    outboxRepository.save(row);
}""",
    )
    para(
        doc,
        "If deferring without incrementing, also avoid incrementing attemptCount before "
        "dispatch (move increment to success/hard-failure paths only).",
    )

    heading(doc, "4.5 [Low] Tighten overview chunk text", 2)
    code(
        doc,
        """private static String buildOverviewText(final JsonNode root) {
    final StringBuilder sb = new StringBuilder();
    appendLabeled(sb, "Headline", textOrNull(root, "headline"));
    appendLabeled(sb, "Overall assessment", textOrNull(root, "overallAssessment"));
    // Leave SOAP for SUMMARY_SOAP chunk only
    ...
}""",
    )

    heading(doc, "4.6 [Low] Import schema constants", 2)
    code(
        doc,
        """import static com.careconnect.model.retrieval.RetrievalIndexSchema.CONSENT_SCOPE_MAX_LENGTH;
import static com.careconnect.model.retrieval.RetrievalIndexSchema.SOURCE_RECORD_ID_MAX_LENGTH;""",
    )

    heading(doc, "4.7 [Process] Commit before PR", 2)
    para(
        doc,
        "Suggested commit message: "
        "\"Add RetrievalIndexService and IndexWorker to ingest outbox events into retrieval chunks (Task 4.1).\"",
    )

    # ── File-level comments ───────────────────────────────────────────────────
    heading(doc, "File-level review comments", 1)
    table(
        doc,
        ["File", "Severity", "Comment"],
        [
            [
                "RetrievalIndexService.java",
                "High",
                "Delete-then-persist can wipe chunks when drafts are empty.",
            ],
            [
                "RetrievalIndexService.java",
                "Medium",
                "contentHash uses contains(); parse JSON for exact equality.",
            ],
            [
                "RetrievalIndexService.java",
                "Medium",
                "Visit summaries always load CallSummaryRepository.",
            ],
            [
                "IndexWorker.java",
                "High",
                "No row claim — unsafe with multiple ECS tasks.",
            ],
            [
                "IndexWorker.java",
                "Medium",
                "Missing patientId should defer, not dead-letter transcripts.",
            ],
            [
                "SummaryChunker.java",
                "Low",
                "Overview duplicates SOAP; consider slimmer overview.",
            ],
            [
                "TranscriptSegmentChunker.java",
                "Info",
                "Clean 1:1 segment mapping; good metadata.",
            ],
            [
                "RetrievalIndexChunkRepository.java",
                "Info",
                "deleteBySourceRecordId is needed; ensure calling TX clears persistence context if needed.",
            ],
            [
                "application.properties",
                "Info",
                "Sensible defaults; document multi-instance caveat in README/deploy guide.",
            ],
            [
                "*Test.java",
                "Low",
                "Add empty-draft non-delete and exact contentHash tests after fixes.",
            ],
        ],
    )

    heading(doc, "Test plan (recommended before merge)", 1)
    bullets(
        doc,
        [
            "Commit Task 4.1 files; open PR against team-ae-develop.",
            "mvnw test -Dtest=TranscriptSegmentChunkerTest,SummaryChunkerTest,"
            "RetrievalIndexServiceTest,IndexWorkerTest",
            "Manual: emit SUMMARY_CREATED via summary generate; confirm chunks in "
            "retrieval_index_chunk within poll interval.",
            "Manual: emit TRANSCRIPT_INDEXED with patientId set; confirm TRANSCRIPT_SEGMENT rows.",
            "Manual: re-emit same summary contentHash; confirm no duplicate chunks / skip log.",
            "If multi-task ECS: verify no duplicate processing after claim fix.",
        ],
    )

    return doc


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    build().save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    main()
