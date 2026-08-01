"""Generate a Word document from the 3.15.7 live-context wiring PR review."""

from datetime import date
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


def set_cell_shading(cell, hex_color: str) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def add_heading(doc: Document, text: str, level: int = 1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    return heading


def add_para(doc: Document, text: str):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.font.size = Pt(11)
    return paragraph


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(item, style="List Bullet")
        for run in paragraph.runs:
            run.font.size = Pt(11)


def add_code(doc: Document, text: str):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    paragraph.paragraph_format.left_indent = Inches(0.25)
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(6)
    return paragraph


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        header_cells[index].text = header
        set_cell_shading(header_cells[index], "1F3A5F")
        for paragraph in header_cells[index].paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(255, 255, 255)
    for row_index, row in enumerate(rows):
        cells = table.rows[row_index + 1].cells
        for col_index, value in enumerate(row):
            cells[col_index].text = value
            for paragraph in cells[col_index].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()


def build_document() -> Document:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.85)
        section.bottom_margin = Inches(0.85)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading("PR Code Review Findings", 0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

    meta = doc.add_paragraph()
    meta.add_run("Branch: ").bold = True
    meta.add_run("feature/team-e/david-3.15.7-live-context-wiring")

    meta = doc.add_paragraph()
    meta.add_run("Compared to: ").bold = True
    meta.add_run("origin/team-ae-develop")

    meta = doc.add_paragraph()
    meta.add_run("Review date: ").bold = True
    meta.add_run(date.today().strftime("%B %d, %Y"))

    meta = doc.add_paragraph()
    meta.add_run("Diff scope: ").bold = True
    meta.add_run("4 commits ahead / 14 behind; ~12 files, +277 / -8")

    meta = doc.add_paragraph()
    meta.add_run("Verdict: ").bold = True
    meta.add_run(
        "Useful, small, well-tested feature - not merge-ready until patient access is "
        "enforced on the chat path. Wiring context without authz turns a weak endpoint "
        "into a PHI exfiltration path to Bedrock."
    )

    add_heading(doc, "Bugbot Findings", 1)
    add_table(
        doc,
        ["Severity", "Location", "Finding"],
        [
            [
                "High",
                "BedrockAIChatAdapter.java:42",
                "No authz before loading patient PHI into Bedrock context from "
                "client-supplied patientId",
            ],
        ],
    )

    add_heading(doc, "1. Change Summary", 1)
    add_para(
        doc,
        "WBS 3.15.7 wires patient medical context into the live Bedrock chat path and "
        "adds a Documents data-source toggle.",
    )
    add_heading(doc, "What landed", 2)
    add_table(
        doc,
        ["Piece", "Change"],
        [
            [
                "BedrockAIChatAdapter",
                "Before calling the model, build context via MedicalContextService and "
                "prepend it to the user message (toBuilder, no in-place mutation)",
            ],
            [
                "MedicalContextService",
                "shouldIncludeDocuments + gate uploaded files on that flag",
            ],
            [
                "Config/schema",
                "include_documents_by_default on UserAIConfig / DTO / defaults; "
                "Flyway V2607032300 + SchemaPatchRunner mirror",
            ],
            [
                "Tests",
                "Adapter context/fail-open tests; documents gate unit tests",
            ],
        ],
    )
    add_para(
        doc,
        "Goal: Live Bedrock chat finally gets the same patient context + data-source "
        "exclusions that other paths expected, including documents opt-out.",
    )

    add_heading(doc, "2. Bug & Risk Analysis", 1)

    add_heading(doc, "High - IDOR / PHI to Bedrock (new blast radius)", 2)
    add_para(
        doc,
        "AIChatController accepts client patientId / userId with no patient-access "
        "check. This PR makes that fatal: withMedicalContext loads name, DOB, vitals, "
        "meds, notes, allergies, mood/pain (and gated uploads) and sends them to AWS.",
    )
    add_para(
        doc,
        "Previously the Bedrock path sent mostly the raw message. Now any authenticated "
        "caller who can hit /v1/api/ai-chat/chat can point patientId at another patient "
        "and exfiltrate PHI into the model prompt.",
    )
    add_code(
        doc,
        "private ChatRequest withMedicalContext(ChatRequest request) {\n"
        "    Long patientId = request.getPatientId();\n"
        "    // ...\n"
        "    UserAIConfigDTO configDto = userAIConfigService.getUserAIConfig(\n"
        "            request.getUserId(), patientId);\n"
        "    String context = medicalContextService.buildPatientContext(\n"
        "            patientId, request, aiConfig);\n"
        "    return request.toBuilder()\n"
        "            .message(context + \"\\n\\nUSER QUESTION:\\n\" + userMessage)\n"
        "            .build();\n"
        "}",
    )

    add_heading(doc, "Medium - Client can force document inclusion", 2)
    add_para(
        doc,
        "request.getIncludeDocuments() wins over config. A client can set "
        "includeDocuments: true even when the user's default is false (same pattern as "
        "other toggles, but weaker as a privacy control).",
    )

    add_heading(doc, 'Medium - "Documents" only means request uploads', 2)
    add_para(
        doc,
        "The flag gates request.getUploadedFiles() only - not stored patient files in "
        "file management / retrieval index. Product naming may over-promise.",
    )

    add_heading(doc, "Medium - Context stuffed into the user message", 2)
    add_para(
        doc,
        "Prepending PHI into message blurs system vs user roles (weaker prompt hygiene), "
        "may persist/log as the user turn, and grows downstream length/cost. Validation "
        "already ran on the original body (good), but assembled prompt size is larger.",
    )

    add_heading(doc, "Low / other risks", 2)
    add_table(
        doc,
        ["Risk", "Note"],
        [
            [
                "Fail-open on context errors",
                "Intentional; chat continues without context (OK). Authz failures must "
                "not be swallowed the same way.",
            ],
            ["Branch 14 behind", "Rebase before merge"],
            [
                "Flyway CREATE TABLE IF NOT EXISTS user_ai_config",
                "Defensive, but can diverge from Hibernate-shaped tables",
            ],
            [
                "Other shouldInclude*",
                "Null config booleans may NPE; documents path is null-safe",
            ],
            [
                "Controller logging",
                "Logs patientId/userId/uploaded files on every chat",
            ],
        ],
    )

    add_heading(doc, "3. Architecture & Style", 1)
    add_heading(doc, "Strengths", 2)
    add_bullets(
        doc,
        [
            "Context assembled at the single AIChatService entry (BedrockAIChatAdapter) "
            "- right place",
            "toBuilder() avoids mutating the inbound DTO (good fix)",
            "Documents gate mirrors existing vitals/meds/notes pattern",
            "SchemaPatchRunner mirror for prod boot",
            "Focused unit tests for adapter + documents flag",
        ],
    )
    add_heading(doc, "Gaps", 2)
    add_bullets(
        doc,
        [
            "Authz belongs in controller (or a shared patient-access guard) before any "
            "context load",
            "Prefer a dedicated context / system field over mutating the user utterance",
            "Documents naming vs upload-only behavior should be documented or extended",
        ],
    )

    add_heading(doc, "4. Recommendations", 1)

    add_heading(doc, "A. Authorize before assembling context (required)", 2)
    add_para(
        doc,
        "In AIChatController (preferred) or at the start of withMedicalContext:",
    )
    add_code(
        doc,
        "@PostMapping(\"/chat\")\n"
        "public ResponseEntity<ChatResponse> sendMessage("
        "@Valid @RequestBody ChatRequest request) {\n"
        "    User current = securityUtil.resolveCurrentUser();\n"
        "    // Never trust body userId - bind from the session.\n"
        "    request.setUserId(current.getId());\n"
        "\n"
        "    if (request.getPatientId() != null) {\n"
        "        authorizationService.requirePatientAccess("
        "current, request.getPatientId());\n"
        "        // or: caregiver link / patient-self / admin\n"
        "    }\n"
        "    return ResponseEntity.ok(aiChatService.processChat(request));\n"
        "}",
    )
    add_para(
        doc,
        "Defense in depth in the adapter (do not catch authz in the fail-open block):",
    )
    add_code(
        doc,
        "private ChatRequest withMedicalContext(ChatRequest request) {\n"
        "    Long patientId = request.getPatientId();\n"
        "    if (patientId == null) {\n"
        "        return request;\n"
        "    }\n"
        "    patientAccessGuard.requireAccess("
        "request.getUserId(), patientId); // throws 403\n"
        "    // ... build context\n"
        "}",
    )

    add_heading(doc, "B. Don't let the client override exclusion without policy", 2)
    add_code(
        doc,
        "boolean shouldIncludeDocuments(ChatRequest request, UserAIConfig aiConfig) {\n"
        "    Boolean configured = aiConfig.getIncludeDocumentsByDefault();\n"
        "    boolean allowByDefault = configured == null || configured;\n"
        "    if (!allowByDefault) {\n"
        "        return false; // config deny wins\n"
        "    }\n"
        "    return request.getIncludeDocuments() == null "
        "|| request.getIncludeDocuments();\n"
        "}",
    )
    add_para(
        doc,
        "Apply the same policy to other sensitive toggles if product requires it.",
    )

    add_heading(doc, "C. Keep context out of the user utterance when possible", 2)
    add_code(
        doc,
        "// Prefer a dedicated field if ChatRequest / Bedrock client supports it:\n"
        "return request.toBuilder()\n"
        "        .additionalContext(/* or systemContext */ List.of(context))\n"
        "        .message(userMessage)\n"
        "        .build();",
    )
    add_para(
        doc,
        "If the model API only accepts one string, still prefix with a clear delimiter "
        "and avoid writing the full context into persisted chat history as the user "
        "message.",
    )

    add_heading(doc, "D. Clarify documents scope", 2)
    add_para(
        doc,
        "Either rename to includeUploadedFilesByDefault, or also load/gate stored "
        "patient documents when the flag is true.",
    )

    add_heading(doc, "E. Before merge checklist", 2)
    add_bullets(
        doc,
        [
            "Rebase onto latest team-ae-develop (currently 14 behind)",
            "Add an integration/authz test: caregiver without access + foreign "
            "patientId -> 403 and no buildPatientContext",
            "Avoid logging full upload metadata at INFO",
            "Close the IDOR / PHI-to-Bedrock gap (Recommendation A) before merge",
        ],
    )

    add_heading(doc, "Merge Recommendation", 1)
    add_para(
        doc,
        "Do not merge until Recommendation A (authorize before assembling context) is "
        "implemented and covered by a test. The feature itself is otherwise clean, "
        "small, and well unit-tested.",
    )
    return doc


def main() -> None:
    output = Path(__file__).resolve().parents[1] / (
        "PR_Review_Live_Context_Wiring.docx"
    )
    build_document().save(output)
    print(output)
    print(f"size_bytes={output.stat().st_size}")


if __name__ == "__main__":
    main()
