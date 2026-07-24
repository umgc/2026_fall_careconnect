"""Generate Word document: Team E Milestone 2 TDD section outline."""

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

OUTPUT = (
    Path(__file__).resolve().parents[1]
    / "docs"
    / "Team_E_Milestone_2_TDD_Section_Outline.docx"
)

SOURCE = r"C:\Users\ravic\Downloads\CareConnect_Milestone_2_TDD_TEAM E.docx"


def heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def para(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True


def bullets(doc: Document, items: list[str], style: str = "List Bullet") -> None:
    for item in items:
        doc.add_paragraph(item, style=style)


def build() -> Document:
    doc = Document()
    title = doc.add_heading("Team E Milestone 2 TDD — Section Outline", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    para(doc, f"Extracted: {date.today().isoformat()}")
    para(doc, f"Source: {SOURCE}")
    para(
        doc,
        "This document lists the sections present in CareConnect_Milestone_2_TDD_TEAM E.docx.",
    )
    doc.add_paragraph()

    heading(doc, "Front matter", 1)
    bullets(doc, ["Revision History"])

    heading(doc, "1. Introduction", 1)
    bullets(
        doc,
        [
            "1.1 Purpose",
            "1.2 Intended Audience",
            "1.3 Document Scope",
            "1.4 Definitions, Acronyms, and Abbreviations",
            "1.5 References",
        ],
    )

    heading(doc, "2. System Overview", 1)
    bullets(
        doc,
        [
            "2.1 System Description",
            "2.2 System Objectives",
            "2.3 System Users",
            "2.4 Operating Environment",
        ],
    )

    heading(doc, "3. Architecture Diagram", 1)
    bullets(
        doc,
        [
            "3.1 High-Level Architecture",
            "3.2 Component Responsibilities",
            "3.3 Data Flow or Workflow View",
            "3.3.1 Use-Case Diagrams Hosted Here",
            "3.4 Module Architecture",
            "3.5 Deployment View",
        ],
    )

    heading(doc, "4. Technology Stack", 1)
    bullets(
        doc,
        [
            "4.1 Selected Technologies",
            "4.2 Technology Decisions and Assumptions",
            "4.3 Technology Risks",
        ],
    )

    heading(doc, "5. Functional Specifications", 1)
    bullets(
        doc,
        [
            "5.1 Call or Visit Summaries",
            "5.1.1 Implemented Use Cases",
            "5.2 AI-Assisted Retrieval (Ask AI)",
            "5.2.1 Implemented Use Cases",
            "5.3 USPS Informed Delivery Mail Agent",
            "5.3.1 Implemented Use Cases",
            "5.4 Short-Term Memory (STML) Support",
            "5.4.1 Implemented Use Cases",
            "5.5 Safety, Consent & Clarity",
            "5.6 Cross-Team Functional Dependencies",
        ],
    )

    heading(doc, "6. API Design Specification", 1)
    bullets(
        doc,
        [
            "6.1 API Conventions and Standards",
            "6.2 Endpoint Summary",
            "6.2.1 POST /api/ai/ask — Parameter Reference",
            "6.2.2 POST /api/summaries/{id}/items/{itemId}/confirm",
            "6.3 Request and Response Examples",
            "6.4 Error Handling and Status Codes",
        ],
    )

    heading(doc, "7. Data Model Specification", 1)
    bullets(
        doc,
        [
            "7.1 Data Modeling Approach and Assumptions",
            "7.2 Data Models by Assigned Module",
            "7.3 Data Dictionary Alignment",
            "7.3.1 Summary Table — Data Dictionary",
            "7.3.2 MailPiece Table — Data Dictionary",
        ],
    )

    heading(doc, "8. UI / UX Design Specification", 1)
    bullets(
        doc,
        [
            "8.1 UI Design Approach and Principles",
            "8.2 Screens and User Workflows Affected",
            "8.3 Wireframes, Mockups, or Screen Notes",
            "8.4 Accessibility Requirements",
            "8.5 Display Text and Localization Impact",
        ],
    )

    heading(doc, "9. Security Considerations", 1)
    bullets(
        doc,
        [
            "9.1 Data Privacy",
            "9.2 Authentication and Authorization",
            "9.3 Secure Data Transmission",
            "9.4 API Security",
            "9.5 Data Storage and Backup",
            "9.6 Logging and Monitoring",
            "9.7 External Service or AI Tool Security",
            "9.8 Security Testing",
        ],
    )

    heading(doc, "10. Error Handling", 1)
    para(doc, "(Section present; no numbered subsections in the source outline.)")

    heading(doc, "11. Performance & Scalability", 1)
    bullets(
        doc,
        [
            "11.1 Objectives",
            "11.2 Scalability Strategy",
            "11.3 Performance Monitoring",
            "11.4 Bottleneck Mitigation",
        ],
    )

    heading(doc, "12. Testing Strategy", 1)
    bullets(
        doc,
        [
            "12.1 Testing Overview",
            "12.2 Testing Levels",
            "12.3 Feature-Specific Test Cases",
            "12.4 Test Automation",
        ],
    )

    heading(doc, "13. Deployment Strategy", 1)
    bullets(
        doc,
        [
            "13.1 Deployment Environments",
            "13.2 Infrastructure or Hosting",
            "13.3 CI/CD Pipeline",
            "13.4 Rollback Procedure",
        ],
    )

    heading(doc, "14. Known Issues & Limitations", 1)
    bullets(
        doc,
        [
            "14.1 Technical Limitations",
            "14.2 Known Issues",
            "14.3 Exclusions",
            "14.4 Future Enhancements",
            "14.5 Cross-Team Resolution Notes",
        ],
    )

    return doc


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    build().save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    main()
