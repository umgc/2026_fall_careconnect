"""Generate Word document: RBAC-scoped source types for Ask AI retrieval (with highlights)."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.shared import Pt
from docx_codebase_revision import REVISION_BULLETS, REVISION_LABEL, save_document

ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "RBAC_Scoped_Retrieval_Source_Types.docx"


def heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def para(doc: Document, text: str, bold: bool = False, highlight: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    if highlight:
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def bullets(doc: Document, items: list[str], highlight_indices: set[int] | None = None) -> None:
    highlight_indices = highlight_indices or set()
    for i, item in enumerate(items):
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        if i in highlight_indices:
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def table(doc: Document, headers: list[str], rows: list[list[str]], highlight_cells: set[tuple[int, int]] | None = None) -> None:
    highlight_cells = highlight_cells or set()
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for r_idx, row in enumerate(rows, start=1):
        cells = t.add_row().cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = val
            if (r_idx, c_idx) in highlight_cells:
                for p in cells[c_idx].paragraphs:
                    for run in p.runs:
                        run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def build() -> None:
    doc = Document()

    title = doc.add_heading("RBAC-Scoped Source Types for Ask AI Retrieval", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para(doc, "CareConnect Team E — Research & Gap Analysis")
    para(
        doc,
        "Derived from SRS v2.0, Milestone 2 TDD, hybrid retrieval scope, unified summary contract, "
        "and the CareConnect application codebase.",
    )
    para(
        doc,
        "Highlighted text (yellow) marks critical gaps, misalignments, and recommended changes.",
        highlight=True,
    )
    para(doc, REVISION_LABEL, highlight=True)
    bullets(doc, REVISION_BULLETS)
    doc.add_paragraph()

    heading(doc, "1. Executive Summary", 1)
    para(
        doc,
        "Ask AI retrieval must enforce four layers before hybrid search: role permission, patient scope, "
        "source consent/exclusions, and episode scope. Source types should be modeled as a canonical "
        "record_type enum on retrieval_index_chunk rows, with patient_id for row-level filtering.",
    )
    para(
        doc,
        "KEY CHANGE NEEDED: The codebase has no retrieval index, no source-type model, and several RBAC "
        "misalignments with Team E SRS (patient Ask AI permission, call participation vs patient links, "
        "caregiverVisibility not enforced, USPS patient access).",
        highlight=True,
    )

    heading(doc, "2. Four Layers of Access Control", 1)
    table(
        doc,
        ["Layer", "Question", "Design (Team E)", "Codebase today"],
        [
            [
                "Role permission",
                "May this role use Ask AI?",
                "USE_AI_FEATURES at gateway",
                "Caregiver + Admin only — PATIENT missing",
            ],
            [
                "Patient scope",
                "Which patient's records?",
                "RBAC + caregiver/family links",
                "Links exist; calls use participants",
            ],
            [
                "Source consent",
                "Which source families?",
                "REQ-SC-7 exclusions + caregiverVisibility",
                "Not implemented",
            ],
            [
                "Episode scope",
                "Which call/visit/mail row?",
                "Participant + link + owner",
                "Ad hoc per endpoint",
            ],
        ],
        highlight_cells={(1, 3), (2, 3), (3, 3), (4, 3)},
    )

    heading(doc, "3. Canonical Retrieval Source Types", 1)

    heading(doc, "3.1 Tier A — Primary Ask AI Sources (SRS §3)", 2)
    table(
        doc,
        ["Source type", "Origin", "Indexed content"],
        [
            ["TRANSCRIPT_SEGMENT", "call_transcript_segments (+ archive)", "Diarized utterances"],
            ["CALL_SUMMARY", "call_summaries.summary_json", "Headline, SOAP, typed arrays"],
            ["VISIT_SUMMARY", "visit_summaries.summary_json (planned)", "Same contract as call"],
            ["UPLOADED_DOCUMENT", "user_files / S3 + OCR", "Lab, insurance, care plans"],
            ["CLINICAL_NOTE", "clinical_notes", "Caregiver/patient notes"],
            ["USPS_MAIL", "usps_digest_cache", "Sender, visible text, importance tier"],
        ],
    )

    heading(doc, "3.2 Tier B — Derived from Summaries", 2)
    bullets(
        doc,
        [
            "SUMMARY_ACTION_ITEM — from actionItems[]",
            "SUMMARY_APPOINTMENT — from appointments[]",
            "SUMMARY_CARE_INSTRUCTION — from careInstructions[] (status + effectiveDate for FR-AI-11)",
            "SUMMARY_CONDITION — from conditions[]",
            "SUMMARY_SOAP — from soap block",
            "SUMMARY_CLINICAL_OBSERVATION — from clinicalObservations categories",
        ],
    )
    para(
        doc,
        "CHANGE: Use one SUMMARY family with record_type subtypes; same chunker for call and visit "
        "summary_json (unified contract).",
        highlight=True,
    )

    heading(doc, "3.3 Tier C — Platform Data (optional index)", 2)
    bullets(
        doc,
        [
            "MEDICATION — structured med list; supports FR-AI-11 timeline queries",
            "TASK — STML recall of pending items",
            "EVV_RECORD — visit verification; distinct from visit summary",
            "VITAL_SIGN — health metrics via MedicalContextService today",
        ],
    )

    heading(doc, "3.4 Tier D — Exclude from Ask AI Index", 2)
    bullets(
        doc,
        [
            "Raw audio / recording playback blobs",
            "AI_CHAT conversation history (not patient records per SRS)",
            "Audit logs (admin-only)",
            "Billing / subscription data",
        ],
    )

    heading(doc, "4. RBAC Matrix by Role and Source Type", 1)
    para(doc, "Y = retrieve when patient-scoped; Y* = active link required; Y† = consent gate; N = deny; A = admin.")
    table(
        doc,
        ["Source type", "ADMIN", "PATIENT", "CAREGIVER", "FAMILY_MEMBER"],
        [
            ["TRANSCRIPT_SEGMENT", "A", "Y (participant)", "Y* (participant/link†)", "Y* (participant)"],
            ["CALL_SUMMARY", "A", "Y", "Y*†", "Y*† read-only"],
            ["VISIT_SUMMARY", "A", "Y", "Y*†", "Y*† read-only"],
            ["UPLOADED_DOCUMENT", "A", "Y", "Y*", "Y* read-only"],
            ["CLINICAL_NOTE", "A", "Y", "Y*", "Y* read-only"],
            ["USPS_MAIL", "Y", "N (today)", "Y", "N"],
            ["MEDICATION", "A", "Y", "Y*", "Y* read-only"],
        ],
        highlight_cells={
            (6, 2),  # USPS PATIENT N
        },
    )

    heading(doc, "5. Critical Codebase Gaps (Highlighted Changes Required)", 1)

    heading(doc, "5.1 PATIENT lacks USE_AI_FEATURES", 2)
    para(
        doc,
        "RolePermissionService grants USE_AI_FEATURES to Caregiver and Admin only. Patients are the "
        "primary Ask AI users in SRS FR-AI-1 but cannot use AI features per current permissions.",
        highlight=True,
    )
    para(doc, "Recommended change: Add Permission.USE_AI_FEATURES to getPatientPermissions().", highlight=True)
    para(
        doc,
        "Consider scoped USE_AI_FEATURES for FAMILY_MEMBER (read-only retrieval) per STML/SRS.",
        highlight=True,
    )

    heading(doc, "5.2 Call access uses participation, not patient links", 2)
    para(
        doc,
        "CallController grants transcript/summary access via call telemetry participation "
        "(isCallParticipant), not CaregiverPatientLink. Assigned caregivers not on a call are denied "
        "unless admin — misaligned with Ask AI over a patient corpus.",
        highlight=True,
    )
    para(
        doc,
        "Recommended change: Index all chunks with patient_id at ingest; filter retrieval by "
        "allowedPatientIds from links, not callId participation alone.",
        highlight=True,
    )

    heading(doc, "5.3 caregiverVisibility not enforced", 2)
    para(
        doc,
        "Sample summaries include caregiverVisibility: on_consent. CallSummaryService does not filter "
        "by consent when storing or serving. TDD ConsentGrant entity is not in codebase.",
        highlight=True,
    )
    para(
        doc,
        "Recommended change: Store caregiver_visibility on summary rows and index chunks; exclude "
        "from caregiver retrieval when consent absent.",
        highlight=True,
    )

    heading(doc, "5.4 USPS mail: patients blocked", 2)
    para(
        doc,
        "USPSController requires admin or caregiver. SRS FR-USPS-4 implies patients search their own "
        "mail via Ask AI. Patients cannot access mail digest API today.",
        highlight=True,
    )
    para(
        doc,
        "Recommended change: Scope USPS by mailbox owner patient_id; allow patient retrieval of own mail.",
        highlight=True,
    )

    heading(doc, "5.5 AuthorizationService patient access incomplete", 2)
    para(
        doc,
        "AuthorizationService.requirePatientAccess() has TODO for link verification. Controllers "
        "duplicate logic via CaregiverPatientLinkService — inconsistent for centralized retrieval.",
        highlight=True,
    )
    para(
        doc,
        "Recommended change: Implement RetrievalScopeService as single scope resolver for Ask AI.",
        highlight=True,
    )

    heading(doc, "5.7 AI feature gating (careconnect.ai.enabled)", 2)
    para(
        doc,
        "AIChatController and Bedrock beans load only when careconnect.ai.enabled=true. "
        "application-dev.properties sets this to false by default — the entire /v1/api/ai-chat "
        "API is absent in default local dev even for roles with USE_AI_FEATURES.",
        highlight=True,
    )
    para(
        doc,
        "AiSymptomController and AiAllergyController are additionally gated by "
        "careconnect.deepseek.enabled (false in dev) — voice form extraction endpoints unavailable locally.",
        highlight=True,
    )

    heading(doc, "5.8 No retrieval index or source-type model", 2)
    para(
        doc,
        "PatientContextRetrievalService is an in-memory substring stub. No retrieval_index_chunk table, "
        "no record_type enum, no hybrid FTS/pgvector query with scope filters.",
        highlight=True,
    )

    heading(doc, "6. Recommended RetrievalScopeService", 1)
    para(doc, "Single service invoked before hybrid search:")
    bullets(
        doc,
        [
            "resolveRetrievalScope(user, patientId?) → allowedPatientIds, allowedSourceTypes",
            "Apply REQ-SC-7 user source exclusions",
            "Apply caregiverVisibility for summary/mail chunks",
            "Filter retrieval_index_chunk: WHERE patient_id IN (...) AND record_type IN (...)",
        ],
        highlight_indices={0, 3},
    )

    heading(doc, "6.1 Patient ID resolution by role", 2)
    table(
        doc,
        ["Role", "allowedPatientIds"],
        [
            ["ADMIN", "All (or query-selected patient)"],
            ["PATIENT", "{ user.id } only"],
            ["CAREGIVER", "Active caregiver_patient_link targets"],
            ["FAMILY_MEMBER", "Active family_member_link targets"],
        ],
    )

    heading(doc, "7. Index Row Shape (RBAC-Aware)", 1)
    table(
        doc,
        ["Column", "Purpose"],
        [
            ["patient_id", "Row-level scope filter (TDD §7.1)"],
            ["record_type", "Source type enum"],
            ["source_record_id", "episodeId, file id, note id"],
            ["episode_type", "call | visit (optional filter)"],
            ["caregiver_visibility", "on_consent | patient_only | shared"],
            ["consent_scope", "Roles allowed to retrieve chunk"],
            ["search_vector", "PostgreSQL FTS"],
            ["embedding", "pgvector semantic search"],
        ],
        highlight_cells={(5, 0), (5, 1), (6, 0), (6, 1)},
    )
    para(
        doc,
        "CHANGE: Do not index chunks the patient excluded (REQ-SC-7). Do not index summary chunks for "
        "caregiver retrieval when caregiverVisibility denies access.",
        highlight=True,
    )

    heading(doc, "8. Source-Specific Scope Notes", 1)

    heading(doc, "8.1 Transcripts", 2)
    bullets(
        doc,
        [
            "Today: participant via actor_user_id or archive participant list or telemetry",
            "Gap: no patient_id on segments — join via contextPatientUserIds in telemetry",
            "Change: resolve patient_id at ingest; index with patient scope key",
        ],
        highlight_indices={1, 2},
    )

    heading(doc, "8.2 Call / Visit summaries (unified contract)", 2)
    bullets(
        doc,
        [
            "Same summary_json inner structure for both episode types",
            "Separate tables: call_summaries and visit_summaries (same columns)",
            "Single retrieval path over both; optional episodeType filter",
            "Honor caregiverVisibility before returning chunks to caregivers",
        ],
        highlight_indices={2, 3},
    )

    heading(doc, "8.3 Uploaded documents", 2)
    bullets(
        doc,
        [
            "UserFile.FileCategory: MEDICAL_RECORD, CLINICAL_NOTE, PRESCRIPTION, LAB_RESULT, INSURANCE_DOCUMENT, CARE_PLAN, etc.",
            "Scope: patient link rules + file owner",
            "Exclude PROFILE_IMAGE from Ask AI index unless product requires",
        ],
    )

    heading(doc, "8.4 Medication timeline (FR-AI-11)", 2)
    para(
        doc,
        "Index careInstructions and conditions with status (started, discontinued, active) and "
        "effectiveDate to support initiation/termination timeline queries.",
        highlight=True,
    )

    heading(doc, "9. Implementation Priority Catalog", 1)
    table(
        doc,
        ["record_type", "Scope key", "Caregiver gate", "Priority"],
        [
            ["TRANSCRIPT", "patient_id", "optional consent", "P0"],
            ["CALL_SUMMARY", "patient_id", "caregiverVisibility", "P0"],
            ["VISIT_SUMMARY", "patient_id", "caregiverVisibility", "P0"],
            ["DOCUMENT", "patient_id + category", "link", "P1"],
            ["CLINICAL_NOTE", "patient_id", "link", "P1"],
            ["USPS_MAIL", "mailbox owner → patient", "fix patient access", "P1"],
            ["MEDICATION", "patient_id", "link", "P2"],
            ["TASK", "patient_id", "link", "P2"],
        ],
        highlight_cells={(6, 3)},
    )

    heading(doc, "10. Current vs Target Summary", 1)
    table(
        doc,
        ["Area", "Current codebase", "Target (Milestone 3)"],
        [
            ["Source type enum", "None", "RetrievalSourceType catalog"],
            ["RBAC at retrieval", "Per-controller, inconsistent", "RetrievalScopeService"],
            ["Patient scope", "Links exist; calls use participants", "patient_id on all index rows"],
            ["Patient Ask AI permission", "Missing USE_AI_FEATURES", "Add for PATIENT"],
            ["Consent / visibility", "Sample JSON only", "Enforce on index + query"],
            ["Index table", "Not present", "retrieval_index_chunk + FTS + pgvector"],
            ["Hybrid search", "In-memory stub", "PostgreSQL hybrid with scope WHERE"],
            ["AI chat API", "Absent unless careconnect.ai.enabled=true", "Always available to permitted roles when AI on"],
        ],
        highlight_cells={
            (1, 1), (2, 1), (3, 1), (4, 1), (5, 1), (6, 1), (7, 1), (8, 1),
        },
    )

    heading(doc, "11. Conclusion", 1)
    para(
        doc,
        "RBAC-scoped retrieval should use ~8 primary source types with summary subtypes, organized "
        "patient-first with role permission, relationship links, and consent gates. The unified call/visit "
        "summary contract simplifies indexing — one chunker, one retrieval path, episodeType as optional metadata.",
    )
    para(
        doc,
        "Highest-priority changes: (1) RetrievalScopeService + index schema, (2) USE_AI_FEATURES for "
        "patients, (3) patient_id on all chunks, (4) enforce caregiverVisibility, (5) fix USPS patient "
        "scope, (6) replace PatientContextRetrievalService stub with hybrid scoped search.",
        highlight=True,
    )

    heading(doc, "12. Related Documents & Code References", 1)
    bullets(
        doc,
        [
            "docs/Hybrid_Retrieval_Scope_PostgreSQL_FullText_pgvector.docx",
            "docs/Ask_AI_Upstream_Pipeline_Call_Visit_Summaries.docx",
            "docs/Call_Transcript_Retrieval_Review.docx",
            "backend/core/src/main/java/com/careconnect/security/RolePermissionService.java",
            "backend/core/src/main/java/com/careconnect/security/AuthorizationService.java",
            "backend/core/src/main/java/com/careconnect/service/CaregiverPatientLinkService.java",
            "backend/core/src/main/java/com/careconnect/service/CallTranscriptService.java",
            "backend/core/src/main/java/com/careconnect/service/PatientContextRetrievalService.java",
            "backend/core/src/main/java/com/careconnect/controller/CallController.java",
            "backend/core/src/main/java/com/careconnect/controller/USPSController.java",
        ],
    )

    save_document(doc, OUTPUT)


if __name__ == "__main__":
    build()
