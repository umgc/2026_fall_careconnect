"""Generate Word document: PR review v3 for GET summary by id vs team-ae-develop."""

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.shared import Pt

OUTPUT = (
    Path(__file__).resolve().parents[1]
    / "docs"
    / "PR_Review_Get_Summary_By_Id_feature_a-fasaa-get-summary-by-id_v3.docx"
)


def heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def para(doc: Document, text: str, bold: bool = False, highlight: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    if highlight:
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def code(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    for i, h in enumerate(headers):
        tbl.rows[0].cells[i].text = h
    for row in rows:
        cells = tbl.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val


def build() -> Document:
    doc = Document()
    title = doc.add_heading("PR Code Review", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    para(
        doc,
        "WBS 3.11.6 — GET /api/v3/summaries/{id} — "
        "feature/a-fasaa-get-summary-by-id → team-ae-develop",
    )
    para(doc, f"Review date: {date.today().isoformat()}")
    para(doc, "Reviewer: AI-assisted code review (Cursor)")
    doc.add_paragraph()

    heading(doc, "Review metadata", 1)
    table(
        doc,
        ["Field", "Value"],
        [
            ["Source branch", "feature/a-fasaa-get-summary-by-id"],
            ["Target branch", "team-ae-develop"],
            [
                "Commits",
                "11c26ff endpoint; 8480fe1 @PreAuthorize; "
                "33fd262 four-way access check (Dominique)",
            ],
            ["Scope", "6 files (+455 / −4)"],
            [
                "Base note",
                "Branch diverged from 40ce1ca; team-ae-develop tip includes "
                "#297 IndexWorker + @EnableMethodSecurity — rebase before merge",
            ],
            ["Verdict", "Approve with minor changes"],
        ],
    )

    heading(doc, "1. Change Summary", 1)
    para(
        doc,
        "Delivers WBS 3.11.6 so Ask AI / clients holding a SUMMARY_CREATED "
        "summaryId can fetch that exact CallSummary row by primary key, with the "
        "same Map response shape as GET /api/v3/calls/{callId}/summary "
        "(latest-by-call). Authz was iterated: role gate via @PreAuthorize, then "
        "object-level four-way check copied from CallController.getCallSummary "
        "(ADMIN, telemetry participant, transcript access, summary owner) → else "
        "403. GlobalExceptionHandler now maps MethodArgumentTypeMismatchException "
        "to 400 (fixes prior 500 on non-numeric path vars; ScheduledVisitControllerTest "
        "updated to match).",
    )

    heading(doc, "What changed", 2)
    table(
        doc,
        ["File", "Change"],
        [
            [
                "CallSummaryController.java",
                "New REST controller GET /api/v3/summaries/{id} with "
                "@PreAuthorize + four-way object check",
            ],
            [
                "CallSummaryService.java",
                "getSummaryEntityById / getSummaryById (findById + toResponse)",
            ],
            [
                "GlobalExceptionHandler.java",
                "MethodArgumentTypeMismatchException → 400 JSON body",
            ],
            [
                "CallSummaryControllerTest.java",
                "7 cases: 4×200 access paths, 403, 404, 400",
            ],
            [
                "CallSummaryServiceTest.java",
                "Unit coverage for by-id lookups",
            ],
            [
                "ScheduledVisitControllerTest.java",
                "Expect 400 (not 500) on invalid date path param",
            ],
        ],
    )

    heading(doc, "2. Bug & Risk Analysis", 1)

    heading(doc, "Medium — Summary-ID existence oracle (404 before authz)", 2)
    para(
        doc,
        "Controller returns 404 when the row is missing, then runs the four-way "
        "check only when the row exists. An authenticated CAREGIVER/PATIENT who "
        "lacks access therefore sees 403 for real IDs and 404 for unused IDs — "
        "leaking which summary primary keys exist. CallController.getCallSummary "
        "evaluates access against callId first and only then returns the payload, "
        "so unauthorized callers do not get a distinct “no summary” signal in the "
        "same way.",
        highlight=True,
    )
    code(
        doc,
        "final Optional<CallSummary> entity = callSummaryService.getSummaryEntityById(id);\n"
        "if (entity.isEmpty()) {\n"
        "    return ResponseEntity.notFound().build();  // leaks existence\n"
        "}\n"
        "// … four-way check → 403 …",
    )
    para(
        doc,
        "Recommendation: for unauthorized callers, prefer a uniform 403 "
        "(or 404 for both missing and forbidden) so ID enumeration does not "
        "reveal PHI-adjacent inventory. Prefer 404-for-both if you want to hide "
        "existence; prefer always-403 if matching CallController’s explicit deny.",
    )

    heading(doc, "Medium — Double repository round-trip", 2)
    para(
        doc,
        "After authz passes, getSummaryById(id) loads the same row again via "
        "findById + toResponse. Under concurrency the row could disappear "
        "(second Optional empty → 404 after a successful authz), which is "
        "harmless but wasteful. Prefer mapping the already-loaded entity.",
    )

    heading(doc, "Low — PATIENT/CAREGIVER may still get 403 despite role gate", 2)
    para(
        doc,
        "hasTranscriptAccess is actorUserId-on-segments (or archived equivalent). "
        "A PATIENT who is the clinical subject but never an actor on segments, "
        "and who is not telemetry actor/target or generatedByUserId, fails the "
        "four-way check even though @PreAuthorize allows PATIENT. Same model as "
        "CallController — parity is good, but document that role alone is "
        "insufficient. caregiverVisibility / on_consent is also not enforced here "
        "or on CallController (shared gap, not introduced by this PR).",
    )

    heading(doc, "Low — Stale controller javadoc on @PreAuthorize", 2)
    para(
        doc,
        "Comments say @PreAuthorize is a silent no-op until Brandon’s branch. "
        "origin/team-ae-develop already has @EnableMethodSecurity on "
        "SecurityConfig, so after rebase/merge the role gate is live. Update the "
        "javadoc to avoid misleading operators.",
    )

    heading(doc, "Low — Authz cost: load all telemetry events", 2)
    para(
        doc,
        "Inherited from CallController: getTelemetryForCall(callId) materializes "
        "the full event list solely to test actor/target membership. Fine for MVP "
        "volume; consider existsByCallIdAndUserId later.",
    )

    heading(doc, "What looks solid", 2)
    bullets(
        doc,
        [
            "Prior IDOR (patientId-scoped / missing object check) addressed by "
            "mirroring CallController’s four-way gate → 403",
            "Response shape reuses toResponse — interchangeable with latest-by-call",
            "Type-mismatch → 400 is a real global quality fix",
            "Controller tests cover each of the four allow paths plus deny/not-found/bad-id",
            "Service null-id guard returns empty Optional cleanly",
        ],
    )

    heading(doc, "3. Architecture & Style", 1)
    bullets(
        doc,
        [
            "Thin controller + service lookup matches existing CallController patterns",
            "Duplicating the four-way check inline (vs shared helper) preserves "
            "behavioral parity but will drift; extract CallSummaryAccessGuard when "
            "a third reader appears",
            "New dedicated CallSummaryController under /api/v3/summaries is clearer "
            "than overloading CallController",
            "Indentation in CallSummaryService (2-space) matches file; new methods "
            "are consistent",
            "Branch is several commits behind team-ae-develop (#297 IndexWorker); "
            "rebase before merge to avoid surprise conflicts",
        ],
    )

    heading(doc, "4. Recommendations", 1)

    heading(doc, "A. Single load + optional existence-hiding 404", 2)
    code(
        doc,
        "@PreAuthorize(\"hasAnyRole('CAREGIVER', 'PATIENT', 'ADMIN')\")\n"
        "@GetMapping(\"/{id}\")\n"
        "public ResponseEntity<Map<String, Object>> getSummaryById(\n"
        "        @PathVariable(\"id\") final Long id) {\n"
        "    final User currentUser = getCurrentUser();\n"
        "    final Optional<CallSummary> entityOpt =\n"
        "            callSummaryService.getSummaryEntityById(id);\n"
        "\n"
        "    // Hide existence from unauthorized callers (uniform 404).\n"
        "    if (entityOpt.isEmpty() || !hasSummaryAccess(currentUser, entityOpt.get())) {\n"
        "        return ResponseEntity.notFound().build();\n"
        "        // Or: throw new AppException(HttpStatus.FORBIDDEN, MSG_ACCESS_DENIED);\n"
        "        // if matching CallController’s explicit 403 for known resources.\n"
        "    }\n"
        "\n"
        "    return ResponseEntity.ok(callSummaryService.toResponsePublic(entityOpt.get()));\n"
        "}\n"
        "\n"
        "private boolean hasSummaryAccess(final User user, final CallSummary summary) {\n"
        "    if (user.getRole() == Role.ADMIN) {\n"
        "        return true;\n"
        "    }\n"
        "    final Long uid = user.getId();\n"
        "    final String callId = summary.getCallId();\n"
        "    final boolean inTelemetry = callTelemetryService.getTelemetryForCall(callId).stream()\n"
        "            .anyMatch(e -> uid.equals(e.getActorUserId())\n"
        "                    || uid.equals(e.getTargetUserId()));\n"
        "    final boolean inTranscript =\n"
        "            callTranscriptService.hasTranscriptAccess(callId, uid);\n"
        "    final boolean isOwner = uid.equals(summary.getGeneratedByUserId());\n"
        "    return inTelemetry || inTranscript || isOwner;\n"
        "}",
    )
    para(
        doc,
        "Expose a package-visible or public toResponse on the service (or package "
        "helper) so the controller does not call getSummaryById again. If product "
        "wants explicit 403 for known IDs (Dominique’s prior ask), keep 403 for "
        "entity-present + denied, and accept the existence oracle — document it.",
    )

    heading(doc, "B. Include summary id in the response map (optional)", 2)
    code(
        doc,
        "// in CallSummaryService.toResponse\n"
        "response.put(\"id\", summary.getId());\n"
        "response.put(\"callId\", summary.getCallId());\n"
        "// …",
    )
    para(
        doc,
        "Helps clients that landed via SUMMARY_CREATED verify they fetched the "
        "row they asked for; low risk additive field.",
    )

    heading(doc, "C. Fix stale @PreAuthorize javadoc", 2)
    code(
        doc,
        " * <li>{@code @PreAuthorize} requires CAREGIVER, PATIENT, or ADMIN.\n"
        " *     Enforced because {@code SecurityConfig} enables\n"
        " *     {@code @EnableMethodSecurity}.</li>",
    )

    heading(doc, "D. Before merge", 2)
    bullets(
        doc,
        [
            "Rebase/merge origin/team-ae-develop (IndexWorker + CODEOWNERS landed)",
            "Decide 403-vs-404 policy for unauthorized-but-existing IDs and align "
            "tests",
            "Smoke: authenticated caregiver with transcript access GET "
            "/api/v3/summaries/{id} → 200; stranger same id → 403 or 404 per policy; "
            "bogus id → 404; id=abc → 400",
        ],
    )

    heading(doc, "Verdict", 1)
    para(
        doc,
        "Approve with minor changes. Object-level authz correctly closes the prior "
        "IDOR relative to CallController. Remaining items are existence-oracle "
        "policy, avoiding a second findById, stale docs, and rebase onto current "
        "team-ae-develop — none block merge if product accepts 403-on-known-id "
        "as intentional (matching the sibling endpoint’s explicit deny).",
        bold=True,
    )

    return doc


def main() -> None:
    doc = build()
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    main()
