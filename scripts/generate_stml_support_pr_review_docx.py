"""Generate Word document: PR code review for feat/team-ae-stml-support → team-ae-develop."""

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.shared import Pt

OUTPUT = (
    Path(__file__).resolve().parents[1]
    / "docs"
    / "PR_Review_STML_Support_feat_team-ae-stml-support.docx"
)


def heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def para(doc: Document, text: str, bold: bool = False, highlight: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    if highlight:
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def bullets(doc: Document, items: list[str], highlight_indices: set[int] | None = None) -> None:
    highlight_indices = highlight_indices or set()
    for i, item in enumerate(items):
        p = doc.add_paragraph(item, style="List Bullet")
        if i in highlight_indices:
            for run in p.runs:
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def code(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    for i, h in enumerate(headers):
        tbl.rows[0].cells[i].text = h
    for row in rows:
        cells = tbl.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val


def build() -> Document:
    doc = Document()
    title = doc.add_heading("PR Code Review", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    para(doc, "STML Support (STML-1–STML-4) — feat/team-ae-stml-support → team-ae-develop")
    para(doc, f"Review date: {date.today().isoformat()}")
    para(doc, "Reviewer: AI-assisted code review (Cursor)")
    doc.add_paragraph()

    heading(doc, "Review metadata", 1)
    table(
        doc,
        ["Field", "Value"],
        [
            ["Source branch", "feat/team-ae-stml-support"],
            ["Target branch", "team-ae-develop"],
            ["Scope", "11 files, +888 / −0 lines (all new)"],
            ["Commits", "968d5fa feat(stml): STML-1–4 endpoints"],
            ["Commits", "6debb13 fix(stml): revert SecurityConfig (defer to Brandon)"],
            ["Commits", "da8a4c6 fix(stml): RBAC scope gate + AIService"],
            ["Commits", "8fd7f51 Merge origin/team-ae-develop"],
            ["Tests", "None added (0 Stml*Test files)"],
        ],
    )

    # ── 1. Change Summary ─────────────────────────────────────────────────────
    heading(doc, "1. Change Summary", 1)
    para(
        doc,
        "This PR introduces Short-Term Memory Support (STML) APIs for patients with memory "
        "limitations and caregivers preparing check-ins. It adds four endpoints under "
        "/v1/api/stml, DTOs, and four service classes that assemble care-record context "
        "(tasks, notes, medications, allergies, chat) and, for recall, call AIService.",
    )

    heading(doc, "Endpoints", 2)
    table(
        doc,
        ["ID", "Method / path", "Purpose"],
        [
            ["STML-1", "POST /v1/api/stml/patients/{patientId}/recall", "Answer a recall question from care records via AI"],
            ["STML-2", "GET /v1/api/stml/patients/{patientId}/brief", "Daily memory brief from incomplete tasks"],
            ["STML-3", "GET /v1/api/stml/patients/{patientId}/checkin?caregiverId=", "Caregiver check-in prep (consent-gated)"],
            ["STML-4", "POST /v1/api/stml/patients/{patientId}/search", "Keyword/sender/date search across chat, notes, tasks, meds, allergies"],
        ],
    )

    heading(doc, "What changed", 2)
    table(
        doc,
        ["File / area", "Change"],
        [
            ["StmlController.java", "REST controller; RetrievalScopeService gate on every endpoint"],
            ["StmlService.java", "Daily brief from TaskRepository incomplete tasks"],
            ["StmlRecallService.java", "Builds context from notes/meds/tasks/allergies; AIService.processChat"],
            ["StmlCheckInService.java", "Consent via CaregiverPatientLink; returns notes + pending items"],
            ["StmlSearchService.java", "In-memory filter search across five repositories"],
            ["DTOs (6)", "Brief, CheckIn, Recall request/response, Search request/response"],
        ],
    )

    heading(doc, "Verdict", 2)
    para(
        doc,
        "Request changes before merge. The feature direction is sound and RBAC gating via "
        "RetrievalScopeService is the right idea, but STML-3 has a critical ID-space bug "
        "(patientId/caregiverId treated as User IDs), the controller swallows ForbiddenScopeException "
        "(breaking the Task 2.6 audit response contract), caregiverId is client-spoofable, "
        "and there are zero automated tests. Address High findings first.",
        bold=True,
    )

    # ── 2. Bug & Risk Analysis ────────────────────────────────────────────────
    heading(doc, "2. Bug & Risk Analysis", 1)

    heading(doc, "2.1 Strengths", 2)
    bullets(
        doc,
        [
            "Uses RetrievalScopeService (USE_AI_FEATURES + patient access) instead of inventing a parallel auth path.",
            "Recall uses AIService abstraction (not raw ChatModel / Bedrock client) — aligns with LLM provider direction.",
            "Recall prompt constrains the model to care records only and includes a medical disclaimer.",
            "Check-in returns consentGranted=false with a clear message instead of leaking data when consent fails.",
            "SecurityConfig left unchanged intentionally (covered by catch-all /v1/api/** authenticated).",
            "Services are focused by endpoint (brief / recall / checkin / search) — readable structure.",
        ],
    )

    heading(doc, "2.2 High — STML-3 treats patientId/caregiverId as User IDs", 2)
    para(
        doc,
        "StmlCheckInService looks up Users with userRepository.findById(patientId) and "
        "findById(caregiverId). Elsewhere (RetrievalScopeService, path params) patientId is the "
        "Patient entity primary key, not users.id. CaregiverPatientLink.existsActiveNonExpiredLink "
        "requires caregiverUser and patientUser. Using the wrong ID space means consent almost "
        "always fails (or, if IDs collide, wrong users are linked).",
        highlight=True,
    )
    code(
        doc,
        """// Current (incorrect ID space)
User caregiver = userRepository.findById(caregiverId).orElse(null);
User patient = userRepository.findById(patientId).orElse(null);
hasConsent = caregiverPatientLinkRepository
    .existsActiveNonExpiredLink(caregiver, patient, LocalDateTime.now());

// Correct approach: resolve Patient → patient.getUser(); Caregiver → caregiver entity → user
Patient patientEntity = patientRepository.findById(patientId)
    .orElseThrow(...);
User patientUser = patientEntity.getUser();
User caregiverUser = /* from authenticated caller or CaregiverRepository */;""",
    )

    heading(doc, "2.3 High — caregiverId query param is spoofable", 2)
    para(
        doc,
        "GET /checkin takes caregiverId as a request parameter. Any caller who already passes "
        "RetrievalScopeService for that patient can supply an arbitrary caregiverId. Consent "
        "should be evaluated for the authenticated caller (SecurityContext), not a client-supplied ID.",
        highlight=True,
    )
    bullets(
        doc,
        [
            "Attack: authenticated caregiver A with access to patient P passes caregiverId=B; response reflects B's consent state incorrectly, or if IDs are wrong, always denies.",
            "Fix: drop caregiverId param; use getCurrentUser() and verify that user is a caregiver linked to the patient.",
        ],
    )

    heading(doc, "2.4 High — Controller catch(Exception) breaks ForbiddenScopeException contract", 2)
    para(
        doc,
        "Every endpoint wraps resolveRetrievalScope in try/catch Exception and returns "
        "ResponseEntity.status(FORBIDDEN).build() with an empty body. GlobalExceptionHandler "
        "already returns the Task 2.6 audit payload (code, denialReason, auditId, deliveryStatus=WITHHELD). "
        "Swallowing ForbiddenScopeException and UnauthorizedException loses that contract and maps "
        "permission failures to empty 403s.",
        highlight=True,
    )
    code(
        doc,
        """// Prefer: let exceptions propagate to GlobalExceptionHandler
@GetMapping("/patients/{patientId}/brief")
public ResponseEntity<StmlBriefDTO> getDailyBrief(@PathVariable Long patientId) {
    retrievalScopeService.resolveRetrievalScope(getCurrentUser(), patientId);
    return ResponseEntity.ok(stmlService.getDailyBrief(patientId));
}""",
    )

    heading(doc, "2.5 Medium — No input validation on recall/search", 2)
    bullets(
        doc,
        [
            "StmlRecallRequest.question can be null/blank → prompt ends with 'PATIENT QUESTION: null'.",
            "fromDate/toDate parsed with LocalDate.parse without validation → DateTimeParseException → 500.",
            "No @Valid / @NotBlank on DTOs; unused fields userId (recall) and mutable patientId on request body.",
        ],
    )

    heading(doc, "2.6 Medium — Search loads all conversations/messages into memory", 2)
    para(
        doc,
        "StmlSearchService loads every active conversation and all messages, then filters in Java. "
        "For patients with long chat history this is O(conversations × messages) memory and latency. "
        "No pagination, no result limit, no DB-side LIKE/FTS.",
    )

    heading(doc, "2.7 Medium — Broad catch(Exception) in services hides real failures", 2)
    para(
        doc,
        "Recall/search/check-in wrap each repository call in catch(Exception) and log.warn. "
        "A programming error (NPE, ClassCast) is treated like a missing table. Prefer catching "
        "specific data-access exceptions or letting unexpected errors fail the request.",
    )

    heading(doc, "2.8 Medium — STML bypasses retrieval_index_chunk / hybrid path", 2)
    para(
        doc,
        "Team E Ask AI design centers on RetrievalScope + retrieval_index_chunk + hybrid retrieval. "
        "STML re-implements ad-hoc multi-repo assembly. That duplicates logic, ignores source-type "
        "exclusions from RetrievalScope, and will diverge from POST /api/ai/ask. Acceptable as a "
        "short-term prototype if documented; otherwise plan to call HybridRetrievalService when ready.",
    )

    heading(doc, "2.9 Low — Daily brief only surfaces incomplete tasks", 2)
    para(
        doc,
        "StmlBriefDTO comments mention RECALL / APPOINTMENT / MEDICATION card types, but "
        "StmlService only emits ACTION_ITEM from tasks. Timestamp is LocalDateTime.now() instead of "
        "task due date — weak for a 'memory brief'.",
    )

    heading(doc, "2.10 Low — Unused PatientRepository in StmlRecallService", 2)
    para(doc, "Injected but never used — dead dependency; remove or use for patient existence checks.")

    heading(doc, "2.11 Low — No automated tests", 2)
    para(
        doc,
        "Zero unit/WebMvc tests for controller RBAC, check-in consent, recall prompt assembly, "
        "or search filters. High regression risk for a PHI-facing feature.",
    )

    heading(doc, "2.12 Informational — Race / concurrency", 2)
    para(
        doc,
        "No shared mutable state in services; request-scoped. No race conditions identified. "
        "Concurrent AI calls are bounded by AIService/Bedrock capacity only.",
    )

    # ── 3. Architecture & Style ───────────────────────────────────────────────
    heading(doc, "3. Architecture & Style", 1)

    heading(doc, "3.1 Design patterns", 2)
    bullets(
        doc,
        [
            "Controller → Service → Repository is consistent with the codebase.",
            "Lombok @RequiredArgsConstructor / @Builder DTOs match project style.",
            "RBAC via RetrievalScopeService is the correct shared gate for AI features.",
            "Missing: @PreAuthorize or method-level security; relying on manual resolve is fine if exceptions propagate.",
            "Missing: shared context-builder for recall/check-in/search (duplicated med/note/allergy loops).",
        ],
    )

    heading(doc, "3.2 Code organization", 2)
    bullets(
        doc,
        [
            "Four services by use case is clear; consider a package com.careconnect.service.stml for cohesion.",
            "Duplicated disclaimer strings and context-building should be extracted.",
            "Indentation mixes 2-space (controller/search) with 4-space (StmlService) — align with project Checkstyle.",
        ],
    )

    heading(doc, "3.3 API / security style", 2)
    bullets(
        doc,
        [
            "Path under /v1/api/stml is covered by SecurityConfig catch-all authenticated matcher — OK.",
            "Empty 403 bodies are inconsistent with Ask AI forbidden-scope JSON contract.",
            "Prefer not mutating request DTOs (request.setPatientId) — pass path patientId as method arg.",
        ],
    )

    # ── 4. Recommendations ────────────────────────────────────────────────────
    heading(doc, "4. Recommendations", 1)

    heading(doc, "4.1 [High] Fix check-in consent ID resolution + use authenticated caregiver", 2)
    code(
        doc,
        """@GetMapping("/patients/{patientId}/checkin")
public ResponseEntity<StmlCheckInDTO> getCheckInView(@PathVariable Long patientId) {
    User caller = getCurrentUser();
    retrievalScopeService.resolveRetrievalScope(caller, patientId);
    return ResponseEntity.ok(stmlCheckInService.getCheckInView(patientId, caller));
}

// StmlCheckInService
public StmlCheckInDTO getCheckInView(Long patientId, User caregiverUser) {
    Patient patient = patientRepository.findById(patientId)
        .orElseThrow(() -> new AppException(HttpStatus.NOT_FOUND, "Patient not found"));
    User patientUser = patient.getUser();
    boolean hasConsent = caregiverPatientLinkRepository
        .existsActiveNonExpiredLink(caregiverUser, patientUser, LocalDateTime.now());
    // ...
}""",
    )

    heading(doc, "4.2 [High] Propagate scope exceptions to GlobalExceptionHandler", 2)
    code(
        doc,
        """private void requireScope(Long patientId) {
    retrievalScopeService.resolveRetrievalScope(getCurrentUser(), patientId);
}

@PostMapping("/patients/{patientId}/recall")
public ResponseEntity<StmlRecallResponse> recall(
        @PathVariable Long patientId,
        @Valid @RequestBody StmlRecallRequest request) {
    requireScope(patientId);
    return ResponseEntity.ok(stmlRecallService.recall(patientId, request.getQuestion()));
}""",
    )

    heading(doc, "4.3 [Medium] Validate recall question and search dates", 2)
    code(
        doc,
        """@Data
public class StmlRecallRequest {
    @NotBlank
    private String question;
}

@Data
public class StmlSearchRequest {
    private String keyword;
    private String sender;
    @Pattern(regexp = "\\\\d{4}-\\\\d{2}-\\\\d{2}", message = "fromDate must be YYYY-MM-DD")
    private String fromDate;
    @Pattern(regexp = "\\\\d{4}-\\\\d{2}-\\\\d{2}", message = "toDate must be YYYY-MM-DD")
    private String toDate;
}""",
    )

    heading(doc, "4.4 [Medium] Cap search results and avoid full chat scan", 2)
    code(
        doc,
        """private static final int MAX_RESULTS = 50;

// After collecting results:
if (results.size() > MAX_RESULTS) {
    results = results.subList(0, MAX_RESULTS);
}

// Longer-term: repository query
// findByPatientIdAndContentContainingIgnoreCase(..., Pageable.ofSize(50))""",
    )

    heading(doc, "4.5 [Medium] Extract shared care-context builder", 2)
    code(
        doc,
        """@Component
@RequiredArgsConstructor
public class StmlCareContextBuilder {
    private final ClinicalNotesRepository notesRepo;
    private final MedicationRepository medRepo;
    private final AllergyRepository allergyRepo;
    private final TaskRepository taskRepo;

    public CareContext build(Long patientId, int noteLimit) {
        // single place used by StmlRecallService + StmlCheckInService
    }
}""",
    )

    heading(doc, "4.6 [Low] Add unit / WebMvc tests (minimum)", 2)
    bullets(
        doc,
        [
            "StmlControllerTest: ForbiddenScopeException → 403 with audit body (not empty).",
            "StmlCheckInServiceTest: consent true/false with Patient.user vs User IDs.",
            "StmlRecallServiceTest: AIService mocked; empty context → graceful answer.",
            "StmlSearchServiceTest: keyword + date filter; invalid date → 400.",
        ],
    )

    heading(doc, "4.7 [Low] Enrich daily brief beyond tasks", 2)
    para(
        doc,
        "Include upcoming appointments / active medications if product intent matches DTO comments; "
        "use task due date for card timestamp.",
    )

    heading(doc, "4.8 Pre-merge checklist", 2)
    table(
        doc,
        ["Step", "Expected"],
        [
            ["Fix STML-3 User vs Patient ID mapping", "Consent uses patient.getUser()"],
            ["Remove spoofable caregiverId param", "Caller from SecurityContext"],
            ["Remove catch(Exception) around scope resolve", "ForbiddenScopeException → GlobalExceptionHandler body"],
            ["Add @Valid + blank question guard", "400 on bad input"],
            ["Add at least controller + check-in unit tests", "CI green"],
            ["Document STML vs hybrid retrieval relationship", "Follow-up ticket OK"],
        ],
    )

    heading(doc, "Summary table — findings by severity", 2)
    table(
        doc,
        ["Severity", "Finding", "Action"],
        [
            ["High", "patientId/caregiverId used as User IDs in check-in consent", "4.1"],
            ["High", "caregiverId query param spoofable", "4.1"],
            ["High", "catch(Exception) empties ForbiddenScopeException response", "4.2"],
            ["Medium", "No validation on question / dates", "4.3"],
            ["Medium", "Full chat history loaded for search", "4.4"],
            ["Medium", "Broad catch(Exception) in services", "Narrow catches"],
            ["Medium", "Bypasses retrieval_index_chunk / hybrid path", "Document / follow-up"],
            ["Low", "Brief only tasks; unused PatientRepository; no tests", "4.6–4.7"],
        ],
    )

    return doc


def main() -> None:
    doc = build()
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    main()
