"""Generate Word document: Voice-query path and upstream STT framework dependencies."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx_codebase_revision import REVISION_BULLETS, REVISION_LABEL, save_document

ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "Voice_Query_Path_and_STT_Framework_Dependencies.docx"


def heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def para(doc: Document, text: str, bold: bool = False, highlight: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    if highlight:
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def bullets(doc: Document, items: list[str], highlight_indices: set[int] | None = None) -> None:
    highlight_indices = highlight_indices or set()
    for i, item in enumerate(items):
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        if i in highlight_indices:
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def code(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"


def table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    highlight_cells: set[tuple[int, int]] | None = None,
) -> None:
    highlight_cells = highlight_cells or set()
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for r_idx, row in enumerate(rows, start=1):
        cells = t.add_row().cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = val
            if (r_idx, c_idx) in highlight_cells:
                for p in cells[c_idx].paragraphs:
                    for run in p.runs:
                        run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def build() -> None:
    doc = Document()

    title = doc.add_heading("Voice-Query Path & Upstream Speech-to-Text Framework Dependencies", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para(doc, "CareConnect — Research synthesis (Team E Ask AI + codebase analysis)")
    para(
        doc,
        "Synthesizes prior transcript, hybrid retrieval, Ask AI upstream, and RBAC reviews with "
        "a focused investigation of how voice input reaches AI services today and which STT "
        "frameworks sit upstream of retrieval indexing.",
    )
    para(
        doc,
        "Yellow highlights mark gaps, misalignments with Team E design, and recommended changes.",
        highlight=True,
    )
    para(doc, REVISION_LABEL, highlight=True)
    bullets(doc, REVISION_BULLETS, highlight_indices={0, 1, 3, 4})
    doc.add_paragraph()

    # --- 1 Executive Summary ---
    heading(doc, "1. Executive Summary", 1)
    para(
        doc,
        "CareConnect uses five distinct speech-to-text stacks, none of which currently feed the "
        "planned Ask AI retrieval gateway. Voice-assisted AI today is limited to structured "
        "form extraction (symptoms, allergies) and file dictation — not free-form patient Q&A.",
    )
    para(
        doc,
        "GAP: Ask AI chat (AIChatService → BedrockAIChatService) accepts text only. There is no "
        "mic/voice-query path in ai_chat_improved.dart or AIChatController.",
        highlight=True,
    )
    para(
        doc,
        "Telehealth transcripts (Chime + AWS Transcribe) are upstream producers for call summaries "
        "and future retrieval — but they capture call conversation, not the user's Ask AI question.",
    )

    # --- 2 Taxonomy ---
    heading(doc, "2. Voice & Speech Paths — Taxonomy", 1)
    table(
        doc,
        ["Path", "STT engine", "Downstream consumer", "Ask AI relevance"],
        [
            [
                "A — Ask AI chat",
                "None (text input only)",
                "AIChatService / Bedrock",
                "Target voice-query path — NOT implemented",
            ],
            [
                "B — Symptom / allergy voice",
                "speech_to_text (Flutter plugin)",
                "DeepseekService → AiSymptom/AiAllergy",
                "Structured extraction, not retrieval",
            ],
            [
                "C — File dictation",
                "speech_to_text",
                "EnhancedFileService upload (.txt)",
                "Future DOCUMENT index source",
            ],
            [
                "D — Clinical notetaker",
                "Sherpa ONNX (on-device)",
                "Patient notetaker API",
                "Future CLINICAL_NOTE index source",
            ],
            [
                "E — Telehealth call capture",
                "Chime SDK + Web Speech API fallback",
                "CallTranscriptService → summaries",
                "Upstream TRANSCRIPT index (Team E)",
            ],
            [
                "F — Post-call transcription",
                "AWS Transcribe (batch, diarized)",
                "CallTranscriptService",
                "Upstream TRANSCRIPT index (async)",
            ],
            [
                "G — Alexa skill",
                "Amazon Alexa ASR (external)",
                "AlexaController (tasks)",
                "Parallel voice channel; no Ask AI",
            ],
            [
                "H — Voice sentiment (calls)",
                "Audio metrics (not STT)",
                "BedrockSentimentService",
                "Sentiment only; not text retrieval",
            ],
        ],
        highlight_cells={(1, 3), (1, 0)},
    )

    # --- 3 Planned vs Current Ask AI voice query ---
    heading(doc, "3. Planned Ask AI Voice-Query Path (Team E) vs Codebase", 1)
    heading(doc, "3.1 Team E design (SRS / TDD / hybrid scope docs)", 2)
    bullets(
        doc,
        [
            "Client POST /api/ai/ask with { query, sessionId } — JWT authenticated (planned gateway).",
            "AI Gateway: rate limit, RBAC scope (RetrievalScopeService), consent check.",
            "Hybrid retrieval over indexed transcript/summary/document/mail chunks.",
            "Bedrock inference with minimum-necessary context (FR-AI-9).",
            "NFR-AI-1: end-to-end response ≤ 5 s p95 including retrieval.",
            "Project Plan §3.2.1: live in-call STT/diarization owned outside Team E; retrieval operates on text after conversation ends.",
        ],
    )
    heading(doc, "3.2 Current Ask AI implementation", 2)
    bullets(
        doc,
        [
            "Frontend: ai_chat_improved.dart → AIChatService.sendMessage() — text field only.",
            "Backend: POST /v1/api/ai-chat/chat — only when careconnect.ai.enabled=true (false by default in dev).",
            "When enabled: BedrockAIChatAdapter → BedrockAIChatService sends raw message — no MedicalContextService, no hybrid retrieval.",
            "PatientContextRetrievalService: in-memory keyword stub; no pgvector/FTS.",
            "No audio upload, streaming STT, or voice-session endpoint on AI chat API.",
        ],
        highlight_indices={1, 2, 3, 4},
    )
    heading(doc, "3.3 Recommended voice-query architecture (gap closure)", 2)
    code(
        doc,
        "Mic → Client STT (speech_to_text or platform) → transcript text\n"
        "     → POST /v1/api/ai-chat/chat { message: transcript, ... }\n"
        "     → (future) POST /api/ai/ask via AI Gateway + RetrievalScopeService\n"
        "     → Hybrid search → Bedrock → cited answer",
    )
    para(
        doc,
        "RECOMMENDATION: Reuse VoiceCommandAI singleShot pattern for Ask AI mic button; keep STT "
        "on-device/client-side for latency and HIPAA minimization; send text only to backend "
        "(consistent with symptom voice path and FR-AI-9 minimum-necessary principle).",
        highlight=True,
    )

    # --- 4 Voice-assisted form paths ---
    heading(doc, "4. Voice-Assisted AI Form Paths (Implemented)", 1)
    heading(doc, "4.1 Shared component: VoiceCommandAI", 2)
    para(doc, "File: frontend/lib/features/ai/presentation/pages/voice_command_ai.dart")
    table(
        doc,
        ["Layer", "Technology", "Behavior"],
        [
            ["Wake word (native)", "Porcupine (porcupine_flutter)", "BuiltInKeyword.PORCUPINE; disabled on web"],
            ["STT", "speech_to_text ^7.3.0", "12 s listen, 2 s pause, dictation mode, onDevice: false"],
            ["singleShot mode", "Navigator.pop(transcript)", "Returns raw transcript to caller"],
            ["Navigation mode", "Hard-coded phrase routing", "home / calendar / tracker — no router entry found"],
        ],
        highlight_cells={(3, 2)},
    )
    heading(doc, "4.2 Symptom voice path", 2)
    code(
        doc,
        "symptom_input_form.dart → VoiceCommandAI(singleShot: true)\n"
        "  → DeepseekService.extractSymptom(transcript)\n"
        "  → POST /v1/api/ai/analyze/symptom\n"
        "  → AiSymptomService.analyze() [DeepSeek LLM JSON extraction]\n"
        "  → Populate form fields (symptomKey, severity, notes)",
    )
    bullets(
        doc,
        [
            "Backend labels input as 'voice transcript' in AiSymptomService prompt.",
            "Requires CREATE_TASKS permission on controller (likely misaligned — should be health-data permission).",
            "Endpoints disabled when careconnect.deepseek.enabled=false (default dev profile).",
            "Does not persist transcript as retrieval index row.",
        ],
        highlight_indices={1, 2},
    )
    heading(doc, "4.3 Allergy voice path", 2)
    code(
        doc,
        "allergies_input_form.dart → VoiceCommandAI(singleShot: true)\n"
        "  → DeepseekService.extractAllergy(transcript)\n"
        "  → POST /api/ai/analyze/allergy\n"
        "  → AiAllergyService → structured allergen/reaction/severity",
    )

    # --- 5 STT Framework Inventory ---
    heading(doc, "5. Speech-to-Text Framework Inventory", 1)
    heading(doc, "5.1 Client-side STT frameworks", 2)
    table(
        doc,
        ["Framework", "Package / API", "Platform", "Used for"],
        [
            ["Flutter speech_to_text", "speech_to_text ^7.3.0", "iOS, Android, Web (browser STT)", "VoiceCommandAI, SpeechToTextCard"],
            ["Porcupine wake word", "porcupine_flutter ^3.0.5", "Native only (not web)", "Optional wake in VoiceCommandAI"],
            ["Web Speech API", "window.SpeechRecognition in chime_meeting_embed_web.dart", "Web browsers", "Call transcript fallback when Chime STT inactive"],
            ["Sherpa ONNX", "sherpa_onnx ^1.12.14 + bundled ONNX models", "Native (Flutter FFI)", "StreamingAsrAndDiarizationScreen — notetaker"],
        ],
    )
    para(
        doc,
        "speech_to_text uses platform-native engines: iOS Speech framework, Android SpeechRecognizer, "
        "Chrome/WebKit speech recognition on web. onDevice: false in VoiceCommandAI → cloud-assisted "
        "recognition where platform supports it.",
    )
    heading(doc, "5.2 Server-side / AWS STT frameworks", 2)
    table(
        doc,
        ["Framework", "Service class", "Trigger", "Persists to DB?"],
        [
            [
                "AWS Chime Meeting Transcription",
                "ChimeService.startMeetingTranscription()",
                "Meeting start (EngineTranscribeSettings en-US)",
                "No — debug/operational only",
            ],
            [
                "AWS Transcribe (batch + diarization)",
                "PostCallTranscriptionService",
                "After recording concatenation completes",
                "Yes — POST_CALL_TRANSCRIBE segments",
            ],
            [
                "AWS Transcribe (via Chime pipeline)",
                "Implicit in Chime SDK client capture",
                "Live call — transcriptionController events",
                "Yes — CLIENT_TRANSCRIPT via client POST",
            ],
        ],
        highlight_cells={(1, 3), (2, 3), (3, 3)},
    )
    heading(doc, "5.3 External voice platform", 2)
    bullets(
        doc,
        [
            "Amazon Alexa: voice ASR handled by Amazon; CareConnect receives structured intents at /v1/api/alexa.",
            "Alexa endpoints: calendar task list/add — not Ask AI retrieval.",
            "Family members blocked from Alexa features.",
        ],
    )

    # --- 6 Telehealth upstream chain ---
    heading(doc, "6. Telehealth STT → Retrieval Upstream Chain", 1)
    para(
        doc,
        "This is the primary upstream speech pipeline for Ask AI retrieval over call content — "
        "distinct from the user's voice query mic path.",
    )
    heading(doc, "6.1 Live client capture (Producer A — primary)", 2)
    code(
        doc,
        "chime_meeting_embed (web/mobile)\n"
        "  → transcriptionController.subscribeToTranscriptEvent (preferred)\n"
        "  → OR realtimeSubscribeToReceiveDataMessage (multi-topic probe)\n"
        "  → OR Web Speech API fallback (web only)\n"
        "hybrid_video_call_widget._handleTranscriptSample()\n"
        "  → VideoCallService.sendTranscriptSegment()\n"
        "  → POST /api/v3/calls/{callId}/transcript/segments\n"
        "CallTranscriptService.recordSegments() — source CLIENT_TRANSCRIPT",
    )
    bullets(
        doc,
        [
            "Buffered flush every 4 s; max 120 segments; max 1200 chars/segment.",
            "Care-team-only calls disable transcript capture.",
            "actorUserId from JWT poster enables participant-based access control.",
        ],
    )
    heading(doc, "6.2 Post-call AWS Transcribe (Producer B — async)", 2)
    code(
        doc,
        "CallRecordingService → PostCallTranscriptionService.transcribeAndCleanup()\n"
        "  → StartTranscriptionJob (speaker diarization, max 10 labels)\n"
        "  → Poll ≤ 15 min → parse JSON → recordSegments(actorUserId=null)\n"
        "  → Delete concatenated MP4 + transcript JSON from S3",
    )
    bullets(
        doc,
        [
            "Source label: POST_CALL_TRANSCRIBE.",
            "actorUserId always null — RBAC relies on telemetry/archive participant lists.",
            "Often completes AFTER maybeGenerateAndStoreCallSummary() → first summary may be NO_TRANSCRIPT.",
            "Gap: no automatic summary regen when post-call transcription completes.",
        ],
        highlight_indices={2, 3},
    )
    heading(doc, "6.3 Chime backend transcription (Producer C — non-persisting)", 2)
    para(
        doc,
        "ChimeService.ensureMeetingTranscriptionStarted() calls AWS StartMeetingTranscription but "
        "does not write to CallTranscriptService. Client-side capture (6.1) is the actual ingestion path.",
    )
    heading(doc, "6.4 Downstream to Ask AI index (planned)", 2)
    code(
        doc,
        "call_transcript_segments (+ archive)\n"
        "  → CallSummaryService (Bedrock summary_json)\n"
        "  → (planned) retrieval_index_chunk [record_type=TRANSCRIPT | CALL_SUMMARY]\n"
        "  → Hybrid FTS + pgvector search\n"
        "  → Ask AI answer with citations",
    )
    para(
        doc,
        "Team E NFR-AI-3: index refresh ≤ 5 min after new transcript/summary — not implemented.",
        highlight=True,
    )

    # --- 7 Clinical notetaker STT ---
    heading(doc, "7. Clinical Notetaker STT Path (Sherpa ONNX)", 1)
    bullets(
        doc,
        [
            "Entry: notetaker_search.dart → StreamingAsrAndDiarizationScreen.",
            "Models: sherpa-onnx-streaming-zipformer-en (online) + zipformer-gigaspeech (offline).",
            "Diarization: pyannote-segmentation asset + speaker embedding (native FFI).",
            "Output: PatientNote saved via NotetakerConfigService → /v1/api/patient-notetaker.",
            "Web: not supported (native-only ONNX runtime).",
            "Future index: CLINICAL_NOTE / DOCUMENT source type for retrieval.",
        ],
        highlight_indices={4},
    )

    # --- 8 File dictation ---
    heading(doc, "8. File Dictation Path (SpeechToTextCard)", 1)
    bullets(
        doc,
        [
            "Used on file_management_page.dart.",
            "speech_to_text → recognized text saved as .txt via EnhancedFileService.uploadFileWeb().",
            "User selects FileCategory (MEDICAL_RECORD, CLINICAL_NOTE, etc.).",
            "Planned upstream for DOCUMENT retrieval index after OCR/chunk pipeline.",
        ],
    )

    # --- 9 Dependency graph ---
    heading(doc, "9. End-to-End Dependency Graph", 1)
    code(
        doc,
        "VOICE QUERY (user asks Ask AI):\n"
        "  [NOT BUILT] Mic → STT → /ai-chat/chat or /api/ai/ask → retrieval → Bedrock\n\n"
        "VOICE FORM (symptom/allergy):\n"
        "  Mic → speech_to_text → DeepseekService → /api/ai/analyze/* → form UI\n\n"
        "CALL STT (retrieval upstream):\n"
        "  Chime/WebSpeech → POST transcript/segments → CallTranscriptService\n"
        "  AWS Transcribe → POST_CALL_TRANSCRIBE segments\n"
        "  → CallSummaryService → (planned) retrieval_index_chunk\n\n"
        "NOTETaker STT:\n"
        "  Sherpa ONNX → PatientNote API → (planned) CLINICAL_NOTE index",
    )

    # --- 10 Framework comparison ---
    heading(doc, "10. STT Framework Comparison for Ask AI Voice Query", 1)
    table(
        doc,
        ["Option", "Pros", "Cons", "Fit for Ask AI mic"],
        [
            [
                "speech_to_text (existing)",
                "Already integrated; cross-platform; text-only backend",
                "Web/browser variance; no streaming to server",
                "Best near-term — reuse VoiceCommandAI",
            ],
            [
                "Sherpa ONNX",
                "On-device privacy; offline",
                "Heavy assets; native only; separate from chat UI",
                "Secondary — offline mode only",
            ],
            [
                "AWS Transcribe streaming",
                "Consistent cloud quality",
                "Latency, cost, audio leaves device",
                "Not recommended for query mic",
            ],
            [
                "Web Speech API direct",
                "Zero Flutter plugin on web",
                "Already used as call fallback only",
                "Web fallback if speech_to_text fails",
            ],
        ],
        highlight_cells={(1, 3)},
    )

    # --- 11 RBAC & security ---
    heading(doc, "11. RBAC & Security Cross-Cuts", 1)
    bullets(
        doc,
        [
            "USE_AI_FEATURES not granted to PATIENT role — blocks Ask AI even if voice mic added.",
            "Voice form paths use CREATE_TASKS permission on AiSymptomController — permission mismatch.",
            "Call transcript POST requires call participation; caregiver corpus retrieval needs patient_id scope.",
            "Client-side STT keeps raw audio off backend for queries — aligns with FR-AI-9 minimum-necessary.",
            "Alexa token → patient resolution; separate auth model from JWT Ask AI gateway.",
        ],
        highlight_indices={0, 1},
    )

    # --- 12 Gaps ---
    heading(doc, "12. Critical Gaps & Recommendations", 1)
    table(
        doc,
        ["#", "Gap", "Severity", "Recommendation"],
        [
            [
                "1",
                "No voice input on Ask AI chat UI",
                "High",
                "Add mic via VoiceCommandAI; requires careconnect.ai.enabled=true for backend",
            ],
            [
                "2",
                "Ask AI has no retrieval/RAG pipeline; Bedrock path lacks medical context",
                "High",
                "Re-wire MedicalContextService + implement hybrid search",
            ],
            [
                "3",
                "Five STT stacks with no shared abstraction",
                "Medium",
                "Introduce SpeechCaptureService facade in Flutter",
            ],
            [
                "4",
                "VoiceCommandAI navigation mode orphaned",
                "Low",
                "Wire to router or remove dead navigation handlers",
            ],
            [
                "5",
                "Post-call transcript arrives after summary",
                "Medium",
                "Regenerate summary + trigger index on POST_CALL_TRANSCRIBE complete",
            ],
            [
                "6",
                "Chime backend transcription non-persisting",
                "Low",
                "Document as ops-only; rely on client capture",
            ],
            [
                "7",
                "PATIENT lacks USE_AI_FEATURES",
                "High",
                "Fix RolePermissionService before voice Ask AI launch",
            ],
            [
                "8",
                "No voice-query audit field",
                "Medium",
                "Log inputModality: voice|text in FR-AI-10 audit schema",
            ],
        ],
        highlight_cells={
            (1, 1), (1, 3),
            (2, 1), (2, 3),
            (3, 1),
            (5, 1), (5, 3),
            (7, 1), (7, 3),
        },
    )

    # --- 13 Implementation priority ---
    heading(doc, "13. Implementation Priority (Voice + STT + Ask AI)", 1)
    table(
        doc,
        ["Phase", "Deliverable", "STT dependency"],
        [
            ["P0", "Grant USE_AI_FEATURES to PATIENT; fix symptom permission", "None"],
            ["P0", "Mic on Ask AI chat → speech_to_text → existing /ai-chat/chat", "speech_to_text"],
            ["P1", "RetrievalScopeService + retrieval_index_chunk", "Transcript text from 6.1/6.2"],
            ["P1", "Summary regen hook on post-call Transcribe complete", "AWS Transcribe"],
            ["P2", "Shared SpeechCaptureService; Sherpa offline fallback", "Sherpa ONNX"],
            ["P2", "Index notetaker + dictation files", "Sherpa + speech_to_text"],
            ["P3", "Unified /api/ai/ask gateway with inputModality", "Client STT only"],
        ],
    )

    # --- 14 Related ---
    heading(doc, "14. Related Documents & Code References", 1)
    bullets(
        doc,
        [
            "docs/Call_Transcript_Retrieval_Review.docx",
            "docs/Hybrid_Retrieval_Scope_PostgreSQL_FullText_pgvector.docx",
            "docs/Ask_AI_Upstream_Pipeline_Call_Visit_Summaries.docx",
            "docs/RBAC_Scoped_Retrieval_Source_Types.docx",
            "frontend/lib/features/ai/presentation/pages/voice_command_ai.dart",
            "frontend/lib/widgets/ai_chat_improved.dart",
            "frontend/lib/widgets/speech_to_text_widget.dart",
            "frontend/lib/features/streaming_asr_with_diarization/streaming_asr_and_diarization_native.dart",
            "frontend/lib/widgets/chime_meeting_embed_web.dart",
            "frontend/lib/widgets/hybrid_video_call_widget.dart",
            "frontend/lib/services/video_call_service.dart",
            "backend/core/src/main/java/com/careconnect/service/ChimeService.java",
            "backend/core/src/main/java/com/careconnect/service/PostCallTranscriptionService.java",
            "backend/core/src/main/java/com/careconnect/service/CallTranscriptService.java",
            "backend/core/src/main/java/com/careconnect/controller/AIChatController.java",
            "backend/core/src/main/java/com/careconnect/service/AiSymptomService.java",
        ],
    )

    heading(doc, "15. Conclusion", 1)
    para(
        doc,
        "Voice query for Ask AI is architecturally straightforward — client STT producing text that "
        "feeds the same gateway as typed queries — but is not implemented today. The codebase already "
        "proves the pattern via symptom/allergy voice capture (speech_to_text → DeepSeek extraction). "
        "The heavier STT investment is upstream of retrieval: telehealth capture (Chime + AWS Transcribe) "
        "and clinical notetaker (Sherpa ONNX) supply indexed content, not the query interface itself.",
    )
    para(
        doc,
        "Priority order: (1) RBAC fix for patient Ask AI, (2) mic on chat reusing VoiceCommandAI, "
        "(3) hybrid retrieval index over transcript/summary text, (4) post-call transcript/summary "
        "sync, (5) optional SpeechCaptureService unification across five STT stacks.",
        highlight=True,
    )

    save_document(doc, OUTPUT)


if __name__ == "__main__":
    build()
