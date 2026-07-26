"""Generate Word document: PR code review for USPS mail agent branch."""

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.shared import Pt

OUTPUT = (
    Path(__file__).resolve().parents[1]
    / "docs"
    / "PR_Review_USPS_Mail_Agent_feature_usps-mail-agent.docx"
)


def heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def para(doc: Document, text: str, bold: bool = False, highlight: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    if highlight:
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def bullets(doc: Document, items: list[str], highlight_indices: set[int] | None = None) -> None:
    highlight_indices = highlight_indices or set()
    for i, item in enumerate(items):
        p = doc.add_paragraph(item, style="List Bullet")
        if i in highlight_indices:
            for run in p.runs:
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def code(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    for i, h in enumerate(headers):
        tbl.rows[0].cells[i].text = h
    for row in rows:
        cells = tbl.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val


def build() -> Document:
    doc = Document()
    title = doc.add_heading("PR Code Review", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    para(doc, "USPS Informed Delivery mail agent — feature/usps-mail-agent → team-ae-develop")
    para(doc, f"Review date: {date.today().isoformat()}")
    para(doc, "Reviewer: AI-assisted code review (Cursor)")
    doc.add_paragraph()

    heading(doc, "Review metadata", 1)
    table(
        doc,
        ["Field", "Value"],
        [
            ["Source branch", "feature/usps-mail-agent-+-ASantana-+-USPS-Informed-Delivery-mail-agent-sync-logic"],
            ["Target branch", "team-ae-develop"],
            ["Scope", "32 files changed (+1,966 / −701 lines)"],
            ["Commits", "9 commits"],
            ["Feature", "USPS Informed Delivery sync, Gmail OAuth hardening, patient-scoped RBAC"],
        ],
    )

    heading(doc, "1. Change Summary", 1)
    para(
        doc,
        "This PR hardens the USPS mail agent and Gmail OAuth integration so digest and "
        "email-credential access is JWT-authenticated and patient-scoped, replacing earlier "
        "demo-user and role-only (requireAdminOrCaregiver) checks.",
    )

    heading(doc, "Primary goals", 2)
    table(
        doc,
        ["Area", "What changed"],
        [
            ["Auth model", "Removed unauthenticated demo-user fallback; all USPS paths require JWT"],
            [
                "Patient scoping",
                "Introduced UspsPatientResolver + AuthorizationService.requirePatientAccess() "
                "with repo-backed caregiver/family link checks",
            ],
            [
                "Gmail OAuth",
                "Two-step flow: JWT GET /v1/api/email-credentials/gmail/connect-url → "
                "short-lived startToken → public GET /oauth/google/start",
            ],
            [
                "OAuth security",
                "OAuthStateSigner (HMAC state), OAuthRedirectValidator (allowed return hosts), "
                "token refresh/revoke improvements",
            ],
            ["API cleanup", "patientEmail preferred over legacy userId; cache/credentials keyed by patient DB id"],
            [
                "Frontend",
                "usps_test_screen.dart sends JWT via AuthTokenManager, uses patientEmail, "
                "handles OAuth return errors",
            ],
            ["Tests", "New UspsPatientAccessE2ETest, expanded controller/security/OAuth tests"],
        ],
    )

    heading(doc, "Architectural flow (after PR)", 2)
    bullets(
        doc,
        [
            "Flutter app (JWT) calls GET /usps/latest?patientEmail=...",
            "AuthorizationService.requirePatientAccess validates caller vs patient (link/self/admin)",
            "Digest returned or 204 No Content",
            "For Gmail connect: JWT GET /email-credentials/gmail/connect-url issues startToken",
            "Browser opens /oauth/google/start?startToken=... (public, token-bound)",
            "Google OAuth callback exchanges code; user redirected to frontend with success or oauthError",
        ],
    )

    heading(doc, "2. Bug & Risk Analysis", 1)

    heading(doc, "Critical", 2)

    heading(doc, "1. @RequirePermission(VIEW_ASSIGNED_PATIENTS) blocks patients and family members", 3)
    para(
        doc,
        "UspsDigestController and EmailCredentialController gate on VIEW_ASSIGNED_PATIENTS, but "
        "patients and family members only have VIEW_HEALTH_DATA per RolePermissionService.",
    )
    para(doc, "The Flutter screen calls digest endpoints that hit this gate:", bold=True)
    code(doc, "$base/v1/api/usps/latest?patientEmail=$encodedPatient&date=$dateString")
    para(
        doc,
        "USPSController /mail has no @RequirePermission and would work for patients; "
        "/latest, /search, and /clear-cache would not.",
        highlight=True,
    )
    para(doc, "Impact: Patient self-service USPS and family read-only access likely return 403 before requirePatientAccess runs.")

    heading(doc, "2. Inconsistent authorization between the two USPS controllers", 3)
    para(
        doc,
        "UspsDigestController applies @RequirePermission + requirePatientAccess; "
        "USPSController /mail applies only requirePatientAccess. Same domain, different gates — "
        "easy to fix one path and miss the other.",
    )

    heading(doc, "Medium", 2)

    bullets(
        doc,
        [
            "Duplicated patient resolution: UspsPatientResolver exists but EmailCredentialService reimplements it (including demo-user special case).",
            "UspsPatientResolver does not verify resolved user is a PATIENT role — any users row can be resolved.",
            "clearCacheForUser uses cacheRepo.findAll() — O(all rows) per request; race-prone under concurrent clears.",
            "OAuth error details exposed in redirect URL via exception message in oauthError query param.",
            "Missing JWT returns 403 FORBIDDEN, not 401 UNAUTHORIZED (GlobalExceptionHandler maps UnauthorizedException to 403).",
        ],
    )

    heading(doc, "Low", 2)
    bullets(
        doc,
        [
            "Dead code: mockDigest() in USPSDigestService is unused.",
            "Empty-state inconsistency: /mail returns 200 with empty digest; /latest returns 204 No Content.",
            "Swallowed exceptions in cache read/write (catch Exception ignored) hide corruption/token issues.",
            "No E2E test for patient self-access or family member access — only linked/unlinked caregiver scenarios.",
        ],
    )

    heading(doc, "3. Architecture & Style", 1)

    heading(doc, "Strengths", 2)
    bullets(
        doc,
        [
            "Patient-scoped RBAC via requirePatientAccess() with DB link checks is the right model.",
            "OAuth hardening (OAuthStateSigner, OAuthRedirectValidator, short-lived start tokens) is well-designed.",
            "Constructor injection in controllers replaces field @Autowired.",
            "Structured DTOs (EmailConnectionStatus, GmailConnectUrlResponse) improve reconnect UX.",
            "Test coverage improved meaningfully (UspsPatientAccessE2ETest, OAuth signer/validator tests).",
        ],
    )

    heading(doc, "Weaknesses", 2)
    bullets(
        doc,
        [
            "Two-layer auth is redundant and conflicting: @RequirePermission + requirePatientAccess should not fight each other.",
            "Split USPS surface (USPSController vs UspsDigestController on same base path) increases maintenance cost.",
            "Resolver duplication undermines the UspsPatientResolver extraction.",
            "Permission naming (VIEW_ASSIGNED_PATIENTS on patient-health endpoints) does not match RolePermissionService semantics.",
        ],
    )
    para(doc, "Overall: direction is good; authorization layering needs consolidation before merge.")

    heading(doc, "4. Recommendations", 1)

    heading(doc, "R1 — Fix role gate on digest/email endpoints (Critical)", 2)
    para(doc, "Option A (recommended): Remove @RequirePermission from USPS/email-credential endpoints and rely on requirePatientAccess() only (same as /mail).", bold=True)
    code(
        doc,
        """// UspsDigestController — remove @RequirePermission from all methods
@GetMapping("/latest")
public ResponseEntity<USPSDigest> getLatestDigest(...) throws UnauthorizedException {
    requireAuthenticated(jwt);
    User currentUser = securityUtil.resolveCurrentUser();
    User patientUser = patientResolver.resolvePatient(patientEmail, userId, currentUser);
    authorizationService.requirePatientAccess(currentUser, patientUser.getId());
    // ...
}

private static void requireAuthenticated(Jwt jwt) throws UnauthorizedException {
    if (jwt == null) {
        throw new UnauthorizedException("Missing or invalid authentication token");
    }
}""",
    )
    para(doc, "Apply the same pattern to EmailCredentialController.")

    para(doc, "Option B: Add @RequireAnyPermission annotation + aspect for coarse early gate:", bold=True)
    code(
        doc,
        """@RequireAnyPermission({Permission.VIEW_ASSIGNED_PATIENTS, Permission.VIEW_HEALTH_DATA})""",
    )

    heading(doc, "R2 — Consolidate patient resolution (Medium)", 2)
    code(
        doc,
        """// EmailCredentialService — inject UspsPatientResolver
private User resolvePatientUser(String patientIdentifier, User currentUser) throws UnauthorizedException {
    return patientResolver.resolvePatient(patientIdentifier, null, currentUser);
}""",
    )
    para(doc, "Optionally extend resolver to reject non-patient roles when identifier is explicit.")

    heading(doc, "R3 — Fix clearCacheForUser scalability (Medium)", 2)
    code(
        doc,
        """// USPSDigestCacheRepo
List<USPSDigestCache> findByUserId(String userId);
void deleteByUserId(String userId);

// USPSDigestService
@Transactional
public void clearCacheForUser(String userId) {
    cacheRepo.deleteByUserId(userId);
}""",
    )

    heading(doc, "R4 — Sanitize OAuth error redirects (Medium)", 2)
    code(
        doc,
        """private String buildOAuthErrorRedirect(String returnUrl, Exception e) {
    String base = resolveSuccessRedirect(returnUrl);
    String safeCode = "oauth_failed";  // log full exception server-side
    log.warn("Gmail OAuth callback failed", e);
    String separator = base.contains("?") ? "&" : "?";
    return base + separator + OAUTH_ERROR_PARAM + "=" + safeCode;
}""",
    )

    heading(doc, "R5 — Unify empty digest behavior (Low)", 2)
    code(
        doc,
        """return digest
    .map(ResponseEntity::ok)
    .orElseGet(() -> ResponseEntity.noContent().build());""",
    )
    para(doc, "Apply to /mail to match /latest; update Flutter if needed.")

    heading(doc, "R6 — Add missing access tests (Low)", 2)
    bullets(
        doc,
        [
            "Patient self-access on /v1/api/usps/mail and /latest",
            "Family member linked/unlinked access",
            "Verify PermissionAspect does not block VIEW_HEALTH_DATA roles before requirePatientAccess",
        ],
    )

    heading(doc, "Verdict", 1)
    table(
        doc,
        ["Dimension", "Assessment"],
        [
            ["Security intent", "Strong improvement over demo-user / role-only checks"],
            ["OAuth design", "Solid CSRF and redirect controls"],
            ["RBAC implementation", "Correct at requirePatientAccess, undermined by @RequirePermission"],
            ["Merge readiness", "Conditional — fix R1 (role gate) before merge; R2–R4 recommended"],
        ],
    )

    para(
        doc,
        "The PR achieves its security goals in the service layer, but the "
        "@RequirePermission(VIEW_ASSIGNED_PATIENTS) annotation likely breaks the primary Flutter "
        "USPS flow for patients and blocks family members. Aligning controller gates with "
        "requirePatientAccess (as /mail already does) is the highest-priority fix.",
        highlight=True,
    )

    return doc


def main() -> None:
    doc = build()
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    main()
