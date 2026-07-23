"""Generate Word document: Transcript ingest shapes + SUMMARY_CREATED indexing contract.

Supports WBS 3.11.1 (transcript ingest, #186) and 3.11.5 (persistence + SUMMARY_CREATED, #190).
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.shared import Pt
from docx_codebase_revision import REVISION_BULLETS, REVISION_LABEL, save_document

ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "Transcript_Ingest_and_SUMMARY_CREATED_Indexing_Contract.docx"


def heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def para(doc: Document, text: str, bold: bool = False, highlight: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    if highlight:
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def bullets(doc: Document, items: list[str], highlight_indices: set[int] | None = None) -> None:
    highlight_indices = highlight_indices or set()
    for i, item in enumerate(items):
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        if i in highlight_indices:
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def code(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(8)


def table(doc: Document, headers: list[str], rows: list[list[str]], highlight_rows: set[int] | None = None) -> None:
    highlight_rows = highlight_rows or set()
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for r_idx, row in enumerate(rows, start=1):
        cells = t.add_row().cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = val
            if r_idx in highlight_rows:
                for p in cells[c_idx].paragraphs:
                    for run in p.runs:
                        run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def build() -> None:
    doc = Document()

    title = doc.add_heading(
        "Transcript Ingest Shapes & SUMMARY_CREATED Indexing Contract",
        0,
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para(doc, "CareConnect Team E — Integration spec for upstream indexing pipeline")
    para(
        doc,
        "Scope: WBS 3.11.1 transcript ingest (#186) and WBS 3.11.5 summary persistence + "
        "SUMMARY_CREATED (#190). Derived from codebase review (PostCallTranscriptionService, "
        "CallTranscriptService, CallSummaryService) and Team E design docs.",
    )
    para(doc, REVISION_LABEL, highlight=True)
    bullets(doc, REVISION_BULLETS)
    bullets(
        doc,
        [
            "PR #224 (trunk): unified summary schema DTOs (SummaryCareInstruction, etc.) — not indexing event DTOs.",
            "RetrievalScopeService (WBS 3.2.3): RBAC scope resolver implemented; indexing pipeline still absent.",
            "No SUMMARY_CREATED / TRANSCRIPT_INDEXED emitter or consumer exists in codebase today.",
        ],
        highlight_indices={2},
    )
    doc.add_paragraph()

    # 1 Executive summary
    heading(doc, "1. Executive Summary", 1)
    para(
        doc,
        "This document answers two integration questions for the Ask AI indexing pipeline:",
    )
    bullets(
        doc,
        [
            "What is the diarized AWS Transcribe JSON shape, and what normalized form should indexers consume?",
            "What is the SUMMARY_CREATED event contract (payload + transport) after summary persistence?",
        ],
    )
    para(
        doc,
        "Key takeaway: indexers should read normalized call_transcript_segments (and archive S3 fallback), "
        "not raw Transcribe JSON. Summary indexing should be triggered by an outbox → SNS → SQS event "
        "after call_summaries / visit_summaries commit — not implemented yet.",
        highlight=True,
    )

    # 2 Transcript ingest
    heading(doc, "2. Diarized Transcribe JSON Shape (WBS 3.11.1 / #186)", 1)

    heading(doc, "2.1 Canonical implementation pointer", 2)
    bullets(
        doc,
        [
            "Primary parser: backend/core/src/main/java/com/careconnect/service/PostCallTranscriptionService.java",
            "Method: parseTranscriptItems() reads results.items[] from AWS Transcribe output on S3",
            "Job config: showSpeakerLabels=true, maxSpeakerLabels=10, language en-US, media MP4",
            "Output S3 key: {recordingPrefix}/transcripts/{jobName}.json",
            "Source label on persisted rows: POST_CALL_TRANSCRIBE",
        ],
    )

    heading(doc, "2.2 Raw AWS Transcribe JSON — fields consumed", 2)
    para(doc, "Only results.items[] is parsed. Other top-level keys (transcripts, speaker_labels) are ignored today.")
    table(
        doc,
        ["JSON path", "Type values", "Usage"],
        [
            ["results.items[].type", "pronunciation | punctuation", "Group pronunciation tokens; append punctuation inline"],
            ["results.items[].speaker_label", "spk_0, spk_1, …", "Diarization label; mapped to Speaker 1, Speaker 2, …"],
            ["results.items[].start_time", "seconds as string", "Converted to start_ms (rounded ms)"],
            ["results.items[].end_time", "seconds as string", "Converted to end_ms (rounded ms)"],
            ["results.items[].alternatives[0].content", "string", "Word or punctuation character"],
        ],
    )

    heading(doc, "2.3 Minimal sample (raw Transcribe output)", 2)
    code(
        doc,
        '{\n'
        '  "results": {\n'
        '    "transcripts": [{ "transcript": "Hello there." }],\n'
        '    "speaker_labels": { "speakers": 2, "segments": [] },\n'
        '    "items": [\n'
        '      {\n'
        '        "type": "pronunciation",\n'
        '        "speaker_label": "spk_0",\n'
        '        "start_time": "0.64",\n'
        '        "end_time": "1.04",\n'
        '        "alternatives": [{ "confidence": "0.99", "content": "Hello" }]\n'
        '      },\n'
        '      {\n'
        '        "type": "punctuation",\n'
        '        "alternatives": [{ "content": "," }]\n'
        '      },\n'
        '      {\n'
        '        "type": "pronunciation",\n'
        '        "speaker_label": "spk_1",\n'
        '        "start_time": "1.10",\n'
        '        "end_time": "1.50",\n'
        '        "alternatives": [{ "content": "there" }]\n'
        '      },\n'
        '      {\n'
        '        "type": "punctuation",\n'
        '        "alternatives": [{ "content": "." }]\n'
        '      }\n'
        '    ]\n'
        '  }\n'
        '}',
    )
    para(
        doc,
        "Parser behavior: consecutive pronunciation tokens with the same speaker_label are grouped into "
        "one utterance segment. Punctuation items append to the current utterance without resetting speaker.",
    )

    heading(doc, "2.4 Speaker label mapping", 2)
    bullets(
        doc,
        [
            "buildSpeakerRoleMap(callId) counts unique CALL_JOIN actor_user_id values from call telemetry",
            "Participant count defaults to minimum 2 if telemetry unavailable",
            "spk_0 → Speaker 1, spk_1 → Speaker 2, … spk_(N-1) → Speaker N",
            "Labels beyond known participant count → Unidentified Speaker (extra voices / echo)",
            "actor_user_id is null on post-call Transcribe path (speaker is display label only)",
        ],
    )

    heading(doc, "2.5 Normalized persistence shape (index from here)", 2)
    para(
        doc,
        "After parsing, segments are stored via CallTranscriptService.recordSegments() into call_transcript_segments.",
        bold=True,
    )
    table(
        doc,
        ["Field / column", "Post-call Transcribe value", "Notes"],
        [
            ["speaker_label", "Speaker 1, Speaker 2, …", "Mapped from spk_*"],
            ["transcript_text", "Grouped utterance text", "Max 1200 chars per segment at persist"],
            ["start_ms / end_ms", "From Transcribe times", "Milliseconds from recording start"],
            ["source", "POST_CALL_TRANSCRIBE", "Distinct from CLIENT_TRANSCRIPT (live Chime)"],
            ["occurred_at", "recordingStartedAt + start_ms", "UTC timestamp for timeline alignment"],
            ["actor_user_id", "null", "Live client path sets submitting user ID"],
            ["call_id", "Call identifier string", "No patient_id on segment row today"],
        ],
        highlight_rows={7},
    )
    para(
        doc,
        "Internal DTO: CallTranscriptService.TranscriptSegmentInput(speakerLabel, text, startMs, endMs, source, occurredAt)",
    )

    heading(doc, "2.6 Alternate ingest paths (not raw Transcribe)", 2)
    table(
        doc,
        ["Path", "API / service", "Payload shape"],
        [
            [
                "Live Chime (client)",
                "POST /api/v3/calls/{callId}/transcript/segments",
                '{ "segments": [{ "speakerLabel", "text", "startMs", "endMs", "source?" }] } or single-segment body',
            ],
            [
                "Archive S3 (long calls)",
                "CallTranscriptArchiveService",
                "JSON array of ArchivedTranscriptSegment — normalized CareConnect shape, not AWS Transcribe",
            ],
            [
                "Post-call Transcribe",
                "PostCallTranscriptionService → recordSegments",
                "Parsed from AWS results.items[] as above",
            ],
        ],
    )
    code(
        doc,
        '[\n'
        '  {\n'
        '    "speakerLabel": "Speaker 1",\n'
        '    "text": "Hello there.",\n'
        '    "startMs": 640,\n'
        '    "endMs": 1500,\n'
        '    "source": "POST_CALL_TRANSCRIBE",\n'
        '    "actorUserId": null,\n'
        '    "occurredAt": "2026-07-02T14:30:00.640"\n'
        '  }\n'
        ']',
    )
    para(
        doc,
        "Indexer recommendation (#186): chunk call_transcript_segments (record_type TRANSCRIPT_SEGMENT). "
        "Use archive S3 fallback when live rows were purged after archival. Do not depend on raw Transcribe "
        "JSON — it is deleted from S3 after successful post-call processing unless recording was claimed for playback.",
        highlight=True,
    )

    heading(doc, "2.7 Transcript index trigger (companion to SUMMARY_CREATED)", 2)
    para(doc, "Recommended event for transcript indexing pipeline:")
    code(
        doc,
        '{\n'
        '  "eventType": "TRANSCRIPT_INDEXED",\n'
        '  "eventId": "uuid",\n'
        '  "occurredAt": "2026-07-02T18:40:00Z",\n'
        '  "schemaVersion": 1,\n'
        '  "payload": {\n'
        '    "callId": "abc-123",\n'
        '    "patientId": 42,\n'
        '    "segmentCount": 87,\n'
        '    "source": "POST_CALL_TRANSCRIBE"\n'
        '  }\n'
        '}',
    )
    bullets(
        doc,
        [
            "Emit after recordSegments() completes (or debounce per callId)",
            "Alternative name: TRANSCRIPT_SEGMENT_BATCH_SAVED",
            "Resolve patientId from call telemetry CALL_JOIN → context patient (gap: not on segment row)",
        ],
        highlight_indices={2},
    )

    # 3 SUMMARY_CREATED
    heading(doc, "3. SUMMARY_CREATED Event Contract (WBS 3.11.5 / #190)", 1)

    heading(doc, "3.1 Current codebase state", 2)
    bullets(
        doc,
        [
            "CallSummaryService.persistResponse() saves to call_summaries and returns — no event emission",
            "No SNS, EventBridge, outbox table, or IndexWorker consumer in codebase",
            "PR #224 adds summary schema DTOs and call_summaries columns (caregiver_visibility, summarization_engine, …)",
            "Design target: TDD UC-SUM-5 persist + async indexing event; NFR-AI-3 ≤ 5 min refresh",
        ],
        highlight_indices={1},
    )

    heading(doc, "3.2 Recommended transport", 2)
    table(
        doc,
        ["Option", "Pattern", "Recommendation"],
        [
            [
                "Preferred",
                "Transactional outbox in same DB commit as summary insert → poller → SNS topic → SQS IndexWorker",
                "Use this — at-least-once, decouples OLTP from indexer",
            ],
            [
                "Fallback (MVP)",
                "Polling job: SELECT FROM call_summaries WHERE id NOT IN indexed set",
                "Acceptable until outbox lands (backlog task 3.4)",
            ],
            [
                "Not recommended",
                "Direct in-process call from CallSummaryService to indexer",
                "Couples availability; blocks on embed service",
            ],
            [
                "Not recommended",
                "EventBridge-only without outbox",
                "Risk of lost events on commit failure",
            ],
        ],
        highlight_rows={1, 3, 4},
    )

    heading(doc, "3.3 Event envelope — SUMMARY_CREATED", 2)
    code(
        doc,
        '{\n'
        '  "eventType": "SUMMARY_CREATED",\n'
        '  "eventId": "uuid",\n'
        '  "occurredAt": "2026-07-02T18:45:00Z",\n'
        '  "schemaVersion": 1,\n'
        '  "payload": {\n'
        '    "episodeType": "call",\n'
        '    "sourceTable": "call_summaries",\n'
        '    "summaryId": 12345,\n'
        '    "callId": "abc-123",\n'
        '    "patientId": 42,\n'
        '    "status": "SUCCESS",\n'
        '    "generatedAt": "2026-07-02T18:45:00Z",\n'
        '    "transcriptSegmentCount": 87,\n'
        '    "caregiverVisibility": "on_consent",\n'
        '    "summarizationEngine": "aws_bedrock:amazon.nova-pro-v1:0",\n'
        '    "contentHash": "sha256-of-summary_json"\n'
        '  }\n'
        '}',
    )

    heading(doc, "3.4 Payload field reference", 2)
    table(
        doc,
        ["Field", "Required", "Purpose"],
        [
            ["eventType", "Yes", "SUMMARY_CREATED — single handler for call + visit tables"],
            ["eventId", "Yes", "Idempotency key for consumer"],
            ["schemaVersion", "Yes", "Evolve contract without breaking consumers"],
            ["payload.summaryId", "Yes", "PK → source_record_id on retrieval_index_chunk"],
            ["payload.episodeType", "Yes", "call | visit — optional filter at query time"],
            ["payload.sourceTable", "Yes", "call_summaries | visit_summaries"],
            ["payload.callId", "Call only", "Episode identifier; visit uses visitId when table exists"],
            ["payload.patientId", "Yes", "Row-level RBAC scope — resolve at emit if column missing"],
            ["payload.status", "Yes", "Index only SUCCESS; skip NO_TRANSCRIPT and ERROR"],
            ["payload.contentHash", "Yes", "Skip re-embed when summary_json unchanged"],
            ["payload.caregiverVisibility", "Yes", "on_consent | auto | hidden — RBAC at index/query"],
            ["payload.summarizationEngine", "No", "Audit / traceability (PR #224 column)"],
            ["payload.transcriptSegmentCount", "No", "Telemetry; not used for chunking"],
        ],
        highlight_rows={8},
    )

    heading(doc, "3.5 IndexWorker consumption", 2)
    bullets(
        doc,
        [
            "Minimum consumer input: { type, sourceId, patientId } per embedding pipeline design",
            "Load summary row by summaryId; deserialize summary_json (unified contract from PR #224)",
            "SummaryChunker produces: CALL_SUMMARY, SUMMARY_* subtypes, derived MEDICATION_TIMELINE_EVENT",
            "Upsert retrieval_index_chunk with FTS + pgvector; UNIQUE(source_record_id, record_type, chunk_index)",
            "Idempotency: compare contentHash before re-chunking",
        ],
    )

    heading(doc, "3.6 Emit hook point (#190)", 2)
    para(doc, "Insert outbox row immediately after summaryRepository.save(summary) in CallSummaryService.persistResponse() when status == SUCCESS:")
    bullets(
        doc,
        [
            "1. Resolve patientId (telemetry JOIN or new call_summaries.patient_id column)",
            "2. Compute contentHash = SHA-256(summary_json)",
            "3. INSERT indexing_outbox (event_type, payload_json, created_at)",
            "4. Outbox poller publishes to SNS; IndexWorker subscribed via SQS",
            "5. Mirror for visit_summaries when VisitSummaryService lands",
        ],
    )

    # 4 Integration map
    heading(doc, "4. Task Integration Map", 1)
    table(
        doc,
        ["Task", "GitHub", "Read from", "Emit", "Indexer record_type"],
        [
            ["3.11.1 Transcript ingest", "#186", "call_transcript_segments + archive S3", "TRANSCRIPT_INDEXED", "TRANSCRIPT_SEGMENT"],
            ["3.11.5 Summary persist + event", "#190", "call_summaries row after Bedrock", "SUMMARY_CREATED", "CALL_SUMMARY, SUMMARY_*, MEDICATION_TIMELINE_EVENT"],
        ],
    )

    # 5 Gaps
    heading(doc, "5. Known Gaps & Blockers", 1)
    table(
        doc,
        ["Gap", "Impact", "Remediation"],
        [
            ["patient_id not on call_summaries", "SUMMARY_CREATED cannot scope chunks", "Add column or resolve from telemetry at emit"],
            ["patient_id not on transcript segments", "TRANSCRIPT chunks unscoped", "Resolve call_id → patient at index time"],
            ["Raw Transcribe JSON deleted post-ingest", "Cannot re-parse from S3", "Index from DB/archive; optional snapshot before delete"],
            ["No outbox / SNS / consumer", "Events not wired", "Flyway outbox + poller + IndexWorker (backlog 3.4, 1.5)"],
            ["retrieval_index_chunk table absent", "Nowhere to write chunks", "Flyway pgvector migration (backlog 1.5)"],
            ["visit_summaries table absent", "Visit SUMMARY_CREATED deferred", "Mirror call_summaries schema"],
        ],
        highlight_rows={1, 2, 4, 5},
    )

    # 6 References
    heading(doc, "6. Code & Design References", 1)
    bullets(
        doc,
        [
            "backend/core/src/main/java/com/careconnect/service/PostCallTranscriptionService.java",
            "backend/core/src/main/java/com/careconnect/service/CallTranscriptService.java",
            "backend/core/src/main/java/com/careconnect/service/CallTranscriptArchiveService.java",
            "backend/core/src/main/java/com/careconnect/service/CallSummaryService.java",
            "backend/core/src/main/java/com/careconnect/controller/CallController.java (POST transcript/segments)",
            "backend/core/src/main/resources/db/migration/V61__create_call_transcript_and_summary_tables.sql",
            "backend/core/src/main/resources/db/migration/V2607011950__extend_call_summaries_for_soap_and_safety.sql",
            "docs/Call_Transcript_Retrieval_Review.docx",
            "docs/Ask_AI_Upstream_Pipeline_Call_Visit_Summaries.docx",
            "docs/Hybrid_Retrieval_Scope_PostgreSQL_FullText_pgvector.docx",
            "docs/Storage_Audit_vs_Shared_Index_Schema.docx",
        ],
    )

    heading(doc, "7. Conclusion", 1)
    para(
        doc,
        "For #186, treat PostCallTranscriptionService output (call_transcript_segments) as the canonical "
        "transcript ingest surface — not raw AWS Transcribe JSON. For #190, emit SUMMARY_CREATED via "
        "transactional outbox → SNS after successful summary persist; index only SUCCESS rows with "
        "patientId, contentHash, and caregiverVisibility in the payload.",
    )

    save_document(doc, OUTPUT)


if __name__ == "__main__":
    build()
