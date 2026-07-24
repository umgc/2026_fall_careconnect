"""Generate Word document: PR code review for admin analytics dashboard branch."""

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.shared import Pt

OUTPUT = (
    Path(__file__).resolve().parents[1]
    / "docs"
    / "PR_Review_Admin_Analytics_Dashboard_feature_a-dkinchen-admin-analytics-dashboard.docx"
)


def heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def para(doc: Document, text: str, bold: bool = False, highlight: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    if highlight:
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def bullets(doc: Document, items: list[str], highlight_indices: set[int] | None = None) -> None:
    highlight_indices = highlight_indices or set()
    for i, item in enumerate(items):
        p = doc.add_paragraph(item, style="List Bullet")
        if i in highlight_indices:
            for run in p.runs:
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def code(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    for i, h in enumerate(headers):
        tbl.rows[0].cells[i].text = h
    for row in rows:
        cells = tbl.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val


def build() -> Document:
    doc = Document()
    title = doc.add_heading("PR Code Review", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    para(
        doc,
        "Admin Analytics Dashboard — feature/a-dkinchen-admin-analytics-dashboard → team-ae-develop",
    )
    para(doc, f"Review date: {date.today().isoformat()}")
    para(doc, "Reviewer: AI-assisted code review (Cursor)")
    doc.add_paragraph()

    heading(doc, "Review metadata", 1)
    table(
        doc,
        ["Field", "Value"],
        [
            ["Source branch", "feature/a-dkinchen-admin-analytics-dashboard"],
            ["Target branch", "team-ae-develop"],
            ["Feature commit", "8278da8 — admin analytics endpoint + 9 tests (WBS 2.3.3)"],
            ["Also on branch", "eb2e0ac — merge origin/team-ae-develop"],
            ["Feature scope", "17 files in 8278da8 (+796 / −9)"],
            ["Note", "git diff A...B may show unrelated CallSummary/SOAP files due to multiple merge bases; review focuses on 8278da8"],
        ],
    )

    # ── 1. Change Summary ─────────────────────────────────────────────────────
    heading(doc, "1. Change Summary", 1)
    para(
        doc,
        "This PR implements WBS 2.3.3: an admin-only backend endpoint that returns aggregated, "
        "anonymous product telemetry for a dashboard. It exposes GET /v1/api/admin/analytics/summary "
        "with optional days or from/to windowing, aggregates counts from telemetry_events "
        "(totals, sessions, event names, top features, sync metrics, endpoint errors), "
        "and hardens output with bucket sanitization so PII/PHI is not returned.",
    )

    heading(doc, "What changed (feature commit 8278da8)", 2)
    table(
        doc,
        ["File / area", "Change"],
        [
            ["AdminAnalyticsController", "GET /summary; SecurityUtil + AuthorizationService.requireAdmin"],
            ["AdminAnalyticsService", "Range resolve/clamp (1–90 days); aggregations; sanitizeBucket"],
            ["TelemetryEventRepository", "Native SQL aggregates (count, distinct session, feature, sync, errors)"],
            ["Projection interfaces + DTOs", "Event/feature/endpoint/sync/error summary records"],
            ["SecurityConfig", "hasRole(ADMIN) for /v1/api/admin/analytics/**"],
            ["V75__add_session_id_to_telemetry_events.sql", "ADD COLUMN session_id + partial index"],
            ["AdminAnalyticsControllerTest / ServiceTest", "9 unit tests (auth, rates, sanitization, range validation)"],
        ],
    )

    heading(doc, "Verdict", 2)
    para(
        doc,
        "Approve with minor changes. The analytics feature is well-scoped, dual-gated for admin access, "
        "and thoughtfully avoids returning raw telemetry payloads. Address SchemaPatchRunner coverage for "
        "V75 (Flyway off in ECS), inject Clock for testability, and harden JSON→bigint casts in sync sums. "
        "Ignore unrelated CallSummary/SOAP files if they appear only due to merge-base noise.",
        bold=True,
    )

    # ── 2. Bug & Risk Analysis ────────────────────────────────────────────────
    heading(doc, "2. Bug & Risk Analysis", 1)

    heading(doc, "2.1 Strengths", 2)
    bullets(
        doc,
        [
            "Defense in depth: SecurityConfig hasRole(ADMIN) plus controller requireAdmin(currentUser).",
            "Response is aggregates only — no raw details/device_info maps returned (good privacy posture).",
            "sanitizeBucket rejects empty, >64 char, and non [a-zA-Z0-9._-] values for features/endpoints/event names.",
            "SQL also length-limits feature/endpoint JSON keys and LIMITs top-N results.",
            "Time window is half-open [from, to) with max 90 days — prevents unbounded scans.",
            "days parameter is clamped to [1, 90]; invalid from>=to rejected with AppException 400.",
            "Solid unit tests for empty data, sync success rate, error rates, unsafe feature filter, non-admin.",
            "Indexes already exist on event_time and (event_name, event_time); V75 adds session_id index.",
        ],
    )

    heading(doc, "2.2 Medium — V75 may not apply in ECS (Flyway disabled)", 2)
    para(
        doc,
        "Production/ECS uses SchemaPatchRunner + Hibernate ddl-auto=update, not Flyway. "
        "V75__add_session_id_to_telemetry_events.sql will not run in the deploy pipeline. "
        "Hibernate may add session_id from the entity, but the partial index "
        "idx_telemetry_events_session_id_time will not be created unless mirrored in SchemaPatchRunner.",
        highlight=True,
    )
    bullets(
        doc,
        [
            "Risk: COUNT(DISTINCT session_id) over large windows becomes slow without the index.",
            "session_id already exists on TelemetryEvent entity on team-ae-develop — column may already be present via ddl-auto.",
            "Action: add an idempotent SchemaPatchRunner patch for column + index (same as Ask AI pgvector practice).",
        ],
    )

    heading(doc, "2.3 Medium — Unsafe ::bigint cast on sync_completed details", 2)
    para(
        doc,
        "sumSyncCompletedBetween casts details->>'attempted'|succeeded|failed to bigint. "
        "Malformed client telemetry (non-numeric strings) causes PostgreSQL cast errors → 500 for the whole summary.",
        highlight=True,
    )
    code(
        doc,
        """-- Safer pattern (skip bad rows):
COALESCE(SUM(
  CASE WHEN details->>'attempted' ~ '^[0-9]+$'
       THEN (details->>'attempted')::bigint ELSE 0 END
), 0) AS attempted""",
    )

    heading(doc, "2.4 Medium — Clock is not injectable", 2)
    para(
        doc,
        "AdminAnalyticsService uses `private final Clock clock = Clock.systemUTC();` field initializer. "
        "TimeRange.lastDays(clock) exists but is unused. resolveRange tests cannot freeze 'now', so "
        "defaults-to-seven-days only asserts from.isBefore(to).",
    )
    code(
        doc,
        """@Service
@RequiredArgsConstructor
public class AdminAnalyticsService {
  private final TelemetryEventRepository telemetryEventRepository;
  private final Clock clock; // inject Clock.systemUTC() via @Bean, or Optional with default
}""",
    )

    heading(doc, "2.5 Low — resolveRange ignores days when from/to partially set", 2)
    bullets(
        doc,
        [
            "If only `to` is set, from = to − 7 days (DEFAULT_DAYS), not the days param.",
            "If only `from` is set, to = now; days ignored.",
            "Document in OpenAPI/Javadoc or reject mixed usage (days + from/to) with 400.",
        ],
    )

    heading(doc, "2.6 Low — Duration.toDays() truncates sub-day overflow", 2)
    para(
        doc,
        "A window of 90 days + 23 hours still has toDays()==90 and is accepted. Unlikely abuse vector "
        "given admin-only access; optional fix: compare Duration with Duration.ofDays(MAX_DAYS).",
    )

    heading(doc, "2.7 Low — Multiple sequential queries per summary", 2)
    para(
        doc,
        "getSummary issues ~8 repository calls serially. Fine for admin dashboards at current scale; "
        "consider a single SQL CTE or parallel CompletableFuture if latency becomes an issue.",
    )

    heading(doc, "2.8 Low — No WebMvc / Security filter integration test", 2)
    para(
        doc,
        "Controller tests mock AuthorizationService; they do not verify Spring Security rejects "
        "ROLE_CAREGIVER JWT at the filter chain. A @WebMvcTest or @SpringBootTest with mockMvc "
        "would catch SecurityConfig matcher typos.",
    )

    heading(doc, "2.9 Informational — Merge-base noise", 2)
    para(
        doc,
        "Comparing team-ae-develop...HEAD with multiple merge bases can list CallSummary, Summary* DTOs, "
        "CallSummaryItemDecision, BedrockSentimentService, and a V74 SOAP migration that are NOT in "
        "commit 8278da8. Confirm the PR file list on GitHub matches the feature commit before reviewing those.",
    )

    heading(doc, "2.10 No race conditions", 2)
    para(
        doc,
        "Read-only aggregations; no shared mutable service state. Concurrent admin requests are safe. "
        "Telemetry writers vs readers may see slightly different counts mid-window — expected for analytics.",
    )

    # ── 3. Architecture & Style ───────────────────────────────────────────────
    heading(doc, "3. Architecture & Style", 1)

    heading(doc, "3.1 Design patterns", 2)
    bullets(
        doc,
        [
            "Controller → Service → Repository with projection interfaces is clean and idiomatic Spring Data.",
            "Immutable Java records for DTOs match modern CareConnect style.",
            "TimeRange record encapsulates window resolution — good cohesion.",
            "Dual authorization (HTTP matcher + service requireAdmin) is appropriate for admin analytics.",
            "Sanitization at the service boundary (not only SQL) is defense-in-depth for JSON-derived labels.",
        ],
    )

    heading(doc, "3.2 Code quality", 2)
    bullets(
        doc,
        [
            "Javadoc on controller and service is clear about anonymous / no-PII intent.",
            "Native queries use named parameters and half-open intervals consistently.",
            "Tests use AssertJ and cover the important privacy filter (unsafe feature names).",
            "Minor: unused TimeRange.lastDays / ofUtc helpers — either use them in resolveRange or remove until needed.",
        ],
    )

    heading(doc, "3.3 Consistency with platform schema strategy", 2)
    para(
        doc,
        "Adding only a Flyway file for session_id/index is inconsistent with the team's SchemaPatchRunner "
        "production path. Mirror the DDL there for ECS parity.",
    )

    # ── 4. Recommendations ────────────────────────────────────────────────────
    heading(doc, "4. Recommendations", 1)

    heading(doc, "4.1 [Medium] Add SchemaPatchRunner patch for V75", 2)
    code(
        doc,
        """// SchemaPatchRunner.java
applyPatch(
    "V75 – telemetry_events.session_id + index",
    "ALTER TABLE telemetry_events ADD COLUMN IF NOT EXISTS session_id VARCHAR(64);"
    + "CREATE INDEX IF NOT EXISTS idx_telemetry_events_session_id_time "
    + "ON telemetry_events (session_id, event_time DESC) "
    + "WHERE session_id IS NOT NULL"
);""",
    )

    heading(doc, "4.2 [Medium] Harden sync metric JSON casts", 2)
    code(
        doc,
        """SELECT
  COALESCE(SUM(CASE WHEN details->>'attempted' ~ '^[0-9]+$'
    THEN (details->>'attempted')::bigint ELSE 0 END), 0) AS attempted,
  COALESCE(SUM(CASE WHEN details->>'succeeded' ~ '^[0-9]+$'
    THEN (details->>'succeeded')::bigint ELSE 0 END), 0) AS succeeded,
  COALESCE(SUM(CASE WHEN details->>'failed' ~ '^[0-9]+$'
    THEN (details->>'failed')::bigint ELSE 0 END), 0) AS failed
FROM telemetry_events
WHERE event_time >= :from AND event_time < :to
  AND event_name = 'sync_completed'""",
    )

    heading(doc, "4.3 [Medium] Inject Clock", 2)
    code(
        doc,
        """@Bean
Clock utcClock() {
  return Clock.systemUTC();
}

// Test:
service = new AdminAnalyticsService(repo, Clock.fixed(Instant.parse("2026-07-08T00:00:00Z"), ZoneOffset.UTC));
TimeRange range = service.resolveRange(7, null, null);
assertThat(range.from()).isEqualTo(OffsetDateTime.parse("2026-07-01T00:00:00Z"));""",
    )

    heading(doc, "4.4 [Low] Reject ambiguous range params", 2)
    code(
        doc,
        """public TimeRange resolveRange(Integer days, OffsetDateTime from, OffsetDateTime to) {
  if (days != null && (from != null || to != null)) {
    throw new AppException(HttpStatus.BAD_REQUEST,
        "Provide either days or from/to, not both");
  }
  // ... existing logic
}""",
    )

    heading(doc, "4.5 [Low] Add SecurityConfig smoke test", 2)
    code(
        doc,
        """@AutoConfigureMockMvc
@SpringBootTest
class AdminAnalyticsSecurityIT {
  @Test
  void summary_asCaregiver_returns403() throws Exception {
    mockMvc.perform(get("/v1/api/admin/analytics/summary")
            .with(user("c@test.com").roles("CAREGIVER")))
        .andExpect(status().isForbidden());
  }
}""",
    )

    heading(doc, "4.6 Pre-merge checklist", 2)
    table(
        doc,
        ["Step", "Expected"],
        [
            ["Confirm GitHub PR file list == 8278da8 scope", "No unrelated CallSummary/SOAP files"],
            ["Add SchemaPatchRunner for session_id index", "ECS logs show patch applied"],
            ["Harden ::bigint casts", "Bad sync_completed JSON does not 500"],
            ["Optional: inject Clock + freeze in tests", "Deterministic resolveRange(7)"],
            ["Run AdminAnalytics*Test", "9 tests green"],
        ],
    )

    heading(doc, "Summary table — findings by severity", 2)
    table(
        doc,
        ["Severity", "Finding", "Action"],
        [
            ["Medium", "V75 Flyway-only; index missing in ECS SchemaPatchRunner path", "4.1"],
            ["Medium", "Non-numeric sync details → SQL cast 500", "4.2"],
            ["Medium", "Clock not injectable; weak resolveRange time test", "4.3"],
            ["Low", "days ignored when from/to partially set", "4.4"],
            ["Low", "No Security filter integration test", "4.5"],
            ["Low", "8 serial queries; unused TimeRange helpers", "Optional"],
            ["Info", "Three-dot diff may include unrelated merge-base files", "Verify PR files"],
        ],
    )

    return doc


def main() -> None:
    doc = build()
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    main()
