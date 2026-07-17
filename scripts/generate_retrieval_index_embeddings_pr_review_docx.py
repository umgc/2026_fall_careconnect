"""Generate Word document: PR review for Task 4.2+4.3 embeddings branch."""

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.shared import Pt

OUTPUT = (
    Path(__file__).resolve().parents[1]
    / "docs"
    / "PR_Review_Retrieval_Index_Embeddings_feature_a-rvasireddy.docx"
)


def heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def para(doc: Document, text: str, bold: bool = False, highlight: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    if highlight:
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def code(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    for i, h in enumerate(headers):
        tbl.rows[0].cells[i].text = h
    for row in rows:
        cells = tbl.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val


def build() -> Document:
    doc = Document()
    title = doc.add_heading("PR Code Review", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    para(
        doc,
        "Tasks 4.2 + 4.3 (WBS 3.12.1 / 3.16.3) — FTS coverage, IndexWorker lease hardenings, "
        "and Bedrock Titan embeddings — feature/a-rvasireddy-retrieval-index-embeddings "
        "→ team-ae-develop",
    )
    para(doc, f"Review date: {date.today().isoformat()}")
    para(doc, "Reviewer: AI-assisted code review (Cursor)")
    doc.add_paragraph()

    heading(doc, "Review metadata", 1)
    table(
        doc,
        ["Field", "Value"],
        [
            ["Source branch", "feature/a-rvasireddy-retrieval-index-embeddings"],
            ["Target branch", "team-ae-develop"],
            [
                "Commits (ahead)",
                "c8fc9bf FTS 4.2; 071384f claim lease/FTS filter; "
                "8b3faac merge FTS; 977ac1b Titan embeddings 4.3",
            ],
            ["Scope", "31 files (+1151 / −58)"],
            [
                "Verdict",
                "Approve with changes — move Bedrock I/O out of the ingest DB transaction",
            ],
        ],
    )

    heading(doc, "1. Change Summary", 1)
    para(
        doc,
        "Completes the keyword and vector write legs of the Ask AI retrieval index after "
        "Task 4.1 IndexWorker (already on team-ae-develop via #297). Task 4.2 adds "
        "search_vector backfill, FullTextSearchService (patient-scoped plainto_tsquery / "
        "ts_rank_cd), and pushes record_type filters into SQL before LIMIT. Review "
        "hardenings add durable indexing_outbox.claimed_at leasing and visit-summary "
        "deferral without burning attempts. Task 4.3 adds ChunkEmbeddingService: after "
        "chunk saveAll, invoke amazon.titan-embed-text-v1, write 1536-d pgvector literals "
        "via updateEmbedding, with best-effort failure (NULL embedding → Task 4.4). ECS IAM "
        "gains titan-embed-text-v1 InvokeModel.",
    )

    heading(doc, "What changed (high level)", 2)
    table(
        doc,
        ["Area", "Change"],
        [
            ["FTS (4.2)", "Backfill migration, FullTextSearchService, typed FTS query"],
            ["Outbox", "claimed_at lease; visit deferral burnsAttempt=false"],
            ["Embeddings (4.3)", "ChunkEmbeddingService + EmbeddingVectorFormat"],
            ["Wiring", "RetrievalIndexService.persistDrafts → embedAndPersist"],
            ["IAM / config", "Titan v1 ARN; careconnect.embedding.* properties"],
            ["Tests", "FTS/embedding contract + unit/E2E updates"],
        ],
    )

    heading(doc, "2. Bug & Risk Analysis", 1)

    heading(doc, "High — Bedrock calls inside the ingest DB transaction", 2)
    para(
        doc,
        "IndexWorker runs dispatch() inside TransactionTemplate, and "
        "RetrievalIndexService.ingest* is @Transactional. persistDrafts calls "
        "embedAndPersist synchronously, so each Titan InvokeModel holds the "
        "same JDBC transaction/connection open. Under load or throttling this "
        "risks long transactions, connection-pool pressure, and claim-lease "
        "expiry (default 2 minutes) while another ECS task reclaims the same "
        "outbox row — duplicate ingest.",
        highlight=True,
    )
    code(
        doc,
        "// IndexWorker.processRow\n"
        "transactionTemplate.execute(status -> dispatch(row));  // outer TX\n"
        "// → ingestSummaryCreated (@Transactional joins)\n"
        "// → persistDrafts → chunkEmbeddingService.embedAndPersist(saved);  // HTTP in TX",
    )

    heading(doc, "Medium — contentHash short-circuit never retries failed embeds", 2)
    para(
        doc,
        "If text/FTS chunks were written but all Titan calls failed, a later "
        "SUMMARY_CREATED with the same contentHash skips re-index entirely. "
        "Embeddings stay NULL until Task 4.4 backfill (ChunkEmbeddingBackfillWorker; optional "
        "SchemaPatchRunner partial index V2607161317 for large backlogs). Document clearly; consider "
        "skipping the hash short-circuit when count of NULL embeddings for that "
        "source_record_id is > 0.",
    )

    heading(doc, "Medium — “batch” is sequential single-text invokes", 2)
    para(
        doc,
        "batch-size only slices the loop; each chunk is one InvokeModel. Acceptable "
        "for MVP (Titan has no multi-text embed API like some providers), but "
        "latency scales linearly with chunk count. No backoff/retry on throttle "
        "beyond “log and leave NULL”.",
    )

    heading(doc, "Medium — Titan v2 misconfiguration", 2)
    para(
        doc,
        "IAM still allows titan-embed-text-v2:0. If CARECONNECT_EMBEDDING_MODEL_ID "
        "is pointed at v2, responses are ≤1024-d and fail the 1536 length check "
        "(chunks stay NULL). Fail-soft is correct; add startup validation that "
        "rejects v2 while EMBEDDING_DIMENSION=1536.",
    )

    heading(doc, "Low — truncation / precision", 2)
    bullets(
        doc,
        [
            "max-input-chars=8000 truncates mid-string; may bias embeddings for long transcripts",
            "Float.toString for pgvector is usually fine; prefer full precision if ranking is sensitive",
            "No vector search API yet (expected — Task 5.1 HybridRetrievalService)",
        ],
    )

    heading(doc, "What looks solid", 2)
    bullets(
        doc,
        [
            "1536-d lock matches schema + Titan v1; avoids unsafe SchemaPatchRunner DROP/ADD",
            "Best-effort embeds preserve FTS availability when Bedrock is down",
            "Optional BedrockRuntimeClient + embedding.enabled=false in tests",
            "claimed_at lease + visit no-burn address prior IndexWorker review findings",
            "FTS record_type filter applied before LIMIT (RBAC correctness)",
            "Contract/unit tests for format, Titan parse, wiring, FTS SQL",
        ],
    )

    heading(doc, "3. Architecture & Style", 1)
    bullets(
        doc,
        [
            "Clean separation: ChunkEmbeddingService owns Bedrock; repository owns "
            "native vector UPDATE; RetrievalIndexService owns ingest orchestration",
            "Mirrors BedrockSentimentService invoke pattern (InvokeModel + ObjectMapper)",
            "EmbeddingVectorFormat is a focused pure helper — good",
            "PR bundles 4.2 + lease fixes + 4.3 — cohesive for Ask AI index write path, "
            "but large; reviewers should treat FTS and embeddings as two logical slices",
            "Next architectural step (5.1): query-time embed + cosine / RRF — not in scope",
        ],
    )

    heading(doc, "4. Recommendations", 1)

    heading(doc, "A. Persist chunks, commit, then embed (preferred)", 2)
    code(
        doc,
        "// RetrievalIndexService.persistDrafts — stop embedding inside @Transactional ingest\n"
        "final List<RetrievalIndexChunk> saved = chunkRepository.saveAll(entities);\n"
        "return saved;  // caller embeds after TX commits\n"
        "\n"
        "// IndexWorker.processRow\n"
        "final List<UUID> chunkIds = transactionTemplate.execute(status -> {\n"
        "    dispatchAndCollectChunkIds(row);  // ingest only\n"
        "    return ids;\n"
        "});\n"
        "transactionTemplate.executeWithoutResult(status -> markProcessed(...));\n"
        "chunkEmbeddingService.embedByIds(chunkIds);  // outside ingest TX\n"
        "\n"
        "// Or: @Transactional(propagation = REQUIRES_NEW) on updateEmbedding path only,\n"
        "// and call embedAndPersist AFTER the ingest transactionTemplate returns.",
    )
    para(
        doc,
        "Minimal fix without redesigning IndexWorker: move embedAndPersist out of "
        "persistDrafts into IndexWorker after markProcessed, passing saved chunk IDs "
        "from ingest return value (change ingest to return a result object).",
    )

    heading(doc, "B. Reject Titan v2 at construction when dim is 1536", 2)
    code(
        doc,
        "if (this.modelId.contains(\"titan-embed-text-v2\")\n"
        "        && RetrievalIndexSchema.EMBEDDING_DIMENSION > 1024) {\n"
        "    throw new IllegalStateException(\n"
        "            \"Titan Embed Text v2 maxes at 1024-d; schema requires \"\n"
        "                    + RetrievalIndexSchema.EMBEDDING_DIMENSION\n"
        "                    + \". Use amazon.titan-embed-text-v1 or migrate the column.\");\n"
        "}",
    )

    heading(doc, "C. contentHash skip — also require embeddings present", 2)
    code(
        doc,
        "if (payload.contentHash() != null\n"
        "        && hasMatchingContentHash(sourceRecordId, payload.contentHash())\n"
        "        && !hasMissingEmbeddings(sourceRecordId)) {\n"
        "    return 0;  // truly complete\n"
        "}\n"
        "// else re-index or call embedAndPersist on existing NULL rows",
    )

    heading(doc, "D. Ops knobs while embeds stay sync", 2)
    bullets(
        doc,
        [
            "Raise careconnect.indexing.outbox.claim-lease-minutes when embedding.enabled=true "
            "(e.g. 10) so multi-chunk Titan latency cannot outlive the lease",
            "Cap chunks embedded per outbox event or defer remainder to 4.4",
            "Add simple retry (1–2×) on Bedrock ThrottlingException before leaving NULL",
        ],
    )

    heading(doc, "E. Before merge", 2)
    bullets(
        doc,
        [
            "Confirm FTS PR was not already opened separately — this branch supersedes "
            "feature/a-rvasireddy-retrieval-index-fts-coverage content",
            "Redeploy 03-platform.yaml so Titan v1 IAM is live before enabling embeds in ECS",
            "Smoke: IndexWorker writes chunk → embedding non-null; Bedrock down → chunk "
            "still searchable via FTS, countMissingEmbedding > 0",
        ],
    )

    heading(doc, "Verdict", 1)
    para(
        doc,
        "Approve with changes. FTS + lease hardenings are in good shape. Embeddings "
        "are correctly dimension-locked to Titan v1 and fail soft, but synchronous "
        "Bedrock inside the ingest transaction is the main production risk — move "
        "embed I/O after commit (or REQUIRES_NEW + post-ingest) before relying on "
        "this path under multi-ECS load.",
        bold=True,
    )

    return doc


def main() -> None:
    doc = build()
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    main()
