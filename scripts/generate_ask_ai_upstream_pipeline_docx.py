"""Generate Word document: Ask AI upstream data pipeline (call + visit summaries)."""

import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.shared import Pt
from docx_codebase_revision import REVISION_BULLETS, REVISION_LABEL, save_document

ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "Ask_AI_Upstream_Pipeline_Call_Visit_Summaries.docx"

CALL_SAMPLE = Path(r"C:\Users\ravic\Downloads\sample_summary_response_updated.json")
VISIT_SAMPLE = Path(r"C:\Users\ravic\Downloads\sample_visit_summary_response.json")


def heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def para(doc: Document, text: str, bold: bool = False, highlight: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    if highlight:
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def code(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(8)


def table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val


def load_sample(path: Path) -> dict:
    if path.exists():
        return json.loads(path.read_text(encoding="utf-8"))
    return {}


def summarize_top_level(payload: dict) -> list[str]:
    lines = []
    for key in sorted(payload.keys()):
        if key == "summary":
            lines.append(f"summary: object ({len(payload[key])} nested fields)")
        else:
            val = payload[key]
            if isinstance(val, (dict, list)):
                lines.append(f"{key}: {type(val).__name__}")
            else:
                lines.append(f"{key}: {val}")
    return lines


def summarize_summary_object(summary: dict) -> list[str]:
    lines = []
    for key in sorted(summary.keys()):
        val = summary[key]
        if isinstance(val, list):
            lines.append(f"  {key}: array[{len(val)}]")
        elif isinstance(val, dict):
            lines.append(f"  {key}: object ({', '.join(sorted(val.keys())[:6])}...)")
        else:
            preview = str(val)
            if len(preview) > 60:
                preview = preview[:57] + "..."
            lines.append(f"  {key}: {preview}")
    return lines


def build(include_revision: bool = True) -> None:
    call = load_sample(CALL_SAMPLE)
    visit = load_sample(VISIT_SAMPLE)

    doc = Document()
    title = doc.add_heading("Ask AI Upstream Data Pipeline", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para(doc, "Unified Call & Visit Summary Contract for Retrieval")
    para(doc, "CareConnect Team E — AI Services & Intelligent Recall")
    para(
        doc,
        "Documents how Ask AI receives indexed content from upstream summary producers "
        "using a single schema contract across phone calls and in-person visits.",
    )
    if include_revision:
        para(doc, REVISION_LABEL, highlight=True)
        bullets(doc, REVISION_BULLETS)
        doc.add_paragraph()

    heading(doc, "1. Executive Summary", 1)
    para(
        doc,
        "Call summaries and visit summaries share one downstream JSON contract. The retrieval "
        "consumer (Ask AI hybrid search) does not need to branch on episode type for core clinical "
        "artifacts — action items, appointments, care instructions, conditions, SOAP, risk level, "
        "and urgency banner all use identical structure whether the episode was a phone call or "
        "an in-person visit.",
    )
    para(
        doc,
        "Storage follows Dominique's direction: visit_summaries is a separate table from "
        "call_summaries, but both tables use the same column shape and emit the same summary_json "
        "payload structure. Ask AI uses a single retrieval path that reads from both tables, "
        "filters and ranks uniformly, and exposes episodeType (and visit-specific top-level fields) "
        "only when optional consumer logic is needed.",
    )

    heading(doc, "2. Design Decision: One Contract, Two Tables", 1)
    table(
        doc,
        ["Aspect", "Decision", "Rationale"],
        [
            [
                "Schema contract",
                "Single unified summary_json payload",
                "Medication changes, appointments, and care instructions are the same concepts regardless of channel",
            ],
            [
                "Storage",
                "call_summaries + visit_summaries (same columns)",
                "Separate write paths; identical downstream consumption",
            ],
            [
                "Retrieval path",
                "Single path over both tables",
                "No episode-type branching required in hybrid rank/filter logic",
            ],
            [
                "Episode discriminator",
                "episodeType: call | visit at response top level",
                "Optional branching for visitLocation, visitDuration, or future UX",
            ],
        ],
    )

    heading(doc, "3. Unified Schema: What Is the Same", 1)
    para(doc, "These fields are identical for call and visit summaries:", bold=True)
    bullets(
        doc,
        [
            "Top-level clinical narrative: headline, narrative, overallAssessment, riskLevel, urgencyBanner, summaryConfidence",
            "Full SOAP block: subjective, objective[], assessment, plan (emergency, medications, appointments, monitoring, safety)",
            "clinicalObservations: six categories (acuteRedFlags, symptomCharacterization, functionalStatus, cognitiveBehavioral, medicationRelated, caregiverSignals)",
            "Four typed arrays: actionItems, appointments, careInstructions, conditions",
            "Safety per item: confidence, sourceTurnId, needsConfirmation",
            "State fields (Prof directive): status and effectiveDate on careInstructions and conditions",
            "Envelope metadata: status, generatedAt, transcriptSegmentCount, transcriptAvailable, summarizationEngine, caregiverVisibility",
        ],
    )

    heading(doc, "4. Visit-Specific Additions Only", 1)
    bullets(
        doc,
        [
            "episodeType: \"call\" or \"visit\" at the top level of the API response envelope",
            "visitLocation (visits only): home | clinic | telehealth-video",
            "visitDuration (visits only): duration in minutes",
            "clinicalObservations.functionalStatus may contain richer in-person observations (mobility, gait, edema) but uses the same array-of-strings field structure",
        ],
    )
    para(
        doc,
        "Call sample episodeType: \"call\" — no visitLocation or visitDuration.",
    )
    para(
        doc,
        "Visit sample episodeType: \"visit\" — visitLocation: \"home\", visitDuration: 35.",
    )

    heading(doc, "5. End-to-End Upstream Pipeline for Ask AI", 1)
    para(
        doc,
        "The following pipeline describes how summary data reaches the Ask AI retrieval index.",
        bold=True,
    )

    heading(doc, "5.1 Pipeline Stages", 2)
    table(
        doc,
        ["Stage", "Owner", "Input", "Output", "Ask AI impact"],
        [
            [
                "1. Episode capture",
                "Platform / Chime / EVV",
                "Phone call or in-person visit ends",
                "Diarized transcript segments",
                "Transcript chunks indexed separately (upstream of summary)",
            ],
            [
                "2. Summary generation",
                "Summaries service (Bedrock)",
                "Transcript + sentiment context",
                "Validated summary_json (unified schema)",
                "Primary retrieval source for Ask AI",
            ],
            [
                "3. Persist",
                "Summaries service",
                "summary_json + metadata",
                "Row in call_summaries OR visit_summaries",
                "Durable store; same column shape both tables",
            ],
            [
                "4. Index event",
                "Outbox / SNS (TDD UC-SUM-5)",
                "summary.created event",
                "Indexing pipeline triggered",
                "Async; target ≤ 5 min refresh (NFR-AI-3)",
            ],
            [
                "5. Chunk + embed",
                "Indexing / embedding pipeline",
                "summary_json fields + typed arrays",
                "retrieval_index_chunk rows (FTS + pgvector)",
                "Hybrid search consumable by Ask AI",
            ],
            [
                "6. Ask AI query",
                "Retrieval service",
                "User natural-language question + RBAC scope",
                "Ranked chunks from all permitted sources",
                "Grounded answer + citations",
            ],
        ],
    )

    heading(doc, "5.2 Pipeline Flow (Text Diagram)", 2)
    code(
        doc,
        "Phone call ends                          In-person visit ends\n"
        "       |                                          |\n"
        "       v                                          v\n"
        " Transcript stored                         Transcript stored\n"
        "       |                                          |\n"
        "       v                                          v\n"
        " CallSummaryService                        VisitSummaryService\n"
        " (Bedrock + validation)                    (Bedrock + validation)\n"
        "       |                                          |\n"
        "       v                                          v\n"
        " call_summaries.summary_json               visit_summaries.summary_json\n"
        " episodeType=call                          episodeType=visit\n"
        "       \\                                          /\n"
        "        \\                                        /\n"
        "         v                                      v\n"
        "              SNS / outbox indexing event\n"
        "                         |\n"
        "                         v\n"
        "           Indexing pipeline (chunk + FTS + embedding)\n"
        "                         |\n"
        "                         v\n"
        "              retrieval_index_chunk (PostgreSQL)\n"
        "                         |\n"
        "                         v\n"
        "     Ask AI: hybrid search → Bedrock → cited answer",
    )

    heading(doc, "6. Single Retrieval Path (Recommended)", 1)
    para(
        doc,
        "Ask AI retrieval should use one code path that queries both summary tables (or a unified "
        "view) and applies the same filter/rank logic. episodeType is metadata, not a retrieval fork.",
    )
    bullets(
        doc,
        [
            "Query: SELECT from call_summaries UNION ALL visit_summaries WHERE patient_id IN (permitted scope)",
            "Parse summary_json once using the shared schema deserializer",
            "Index chunks from: headline, narrative, overallAssessment, SOAP text, each actionItem/appointment/careInstruction/condition",
            "Hybrid rank: structured filters (date, riskLevel) + FTS + pgvector — identical for both episode types",
            "Citation assembly: record_type = SUMMARY, source_id = episodeId, excerpt from matched chunk, deep link to summary view",
            "Optional filter: episodeType=visit when user asks visit-specific questions (e.g., \"at the home visit\")",
        ],
    )
    para(
        doc,
        "Two separate retrieval paths (one per table) are not required unless operational isolation "
        "is needed for performance tuning. A single path reduces duplication and matches the unified contract.",
    )

    heading(doc, "7. Indexing Strategy for Unified summary_json", 1)
    table(
        doc,
        ["Chunk source", "Example from samples", "Indexed text fields", "Metadata for citations"],
        [
            [
                "Envelope",
                "riskLevel: LOW / MODERATE",
                "riskLevel, summaryConfidence",
                "episodeId, episodeType, generatedAt",
            ],
            [
                "Narrative block",
                "headline, narrative, overallAssessment",
                "Full text of each field",
                "record_type=SUMMARY_NARRATIVE",
            ],
            [
                "SOAP",
                "subjective, objective[], assessment, plan.*",
                "Flattened SOAP sections",
                "sourceTurnId where available",
            ],
            [
                "clinicalObservations",
                "acuteRedFlags, functionalStatus, ...",
                "Each category array joined",
                "episodeType for visit-rich functionalStatus",
            ],
            [
                "actionItems[]",
                "Get basic metabolic panel...",
                "text + actor + dueHint + status",
                "itemId, sourceTurnId, needsConfirmation",
            ],
            [
                "appointments[]",
                "Follow-up call June 26...",
                "date, time, with, purpose",
                "itemId, sourceTurnId",
            ],
            [
                "careInstructions[]",
                "Increase lisinopril / Discontinue aspirin",
                "text + type + status + effectiveDate",
                "itemId — supports FR-AI-11 medication timeline",
            ],
            [
                "conditions[]",
                "Hypertension, Fluid overload",
                "name + description + status + effectiveDate",
                "itemId, status (active/discontinued)",
            ],
        ],
    )

    heading(doc, "8. Sample Payload Walkthrough", 1)

    heading(doc, "8.1 Call Summary Sample", 2)
    para(doc, f"Source: {CALL_SAMPLE}")
    if call:
        para(doc, "Top-level envelope fields:", bold=True)
        bullets(doc, summarize_top_level(call))
        if "summary" in call:
            para(doc, "Nested summary object:", bold=True)
            bullets(doc, summarize_summary_object(call["summary"]))
        para(
            doc,
            "Clinical highlights: lisinopril increased 10→20 mg; aspirin discontinued; "
            "follow-up call June 26; conditions include Hypertension (active) and Afternoon fatigue "
            "(effectiveDate 2026-06-27). careInstructions carry status started/discontinued with effectiveDate.",
        )

    heading(doc, "8.2 Visit Summary Sample", 2)
    para(doc, f"Source: {VISIT_SAMPLE}")
    if visit:
        para(doc, "Top-level envelope fields (includes visit-specific):", bold=True)
        bullets(doc, summarize_top_level(visit))
        if "summary" in visit:
            para(doc, "Nested summary object (same shape as call):", bold=True)
            bullets(doc, summarize_summary_object(visit["summary"]))
        para(
            doc,
            "Clinical highlights: CHF fluid overload at home visit; furosemide 20→40 mg for 3 days; "
            "urgencyBanner.show=true with 48-hour callback; richer functionalStatus and caregiverSignals "
            "from in-person observation. Same typed arrays and SOAP structure as call sample.",
        )

    heading(doc, "8.3 Side-by-Side Comparison", 2)
    table(
        doc,
        ["Field", "Call sample", "Visit sample"],
        [
            ["episodeId", "call_2026_06_27_001", "visit_2026_06_27_004"],
            ["episodeType", "call", "visit"],
            ["visitLocation", "(absent)", "home"],
            ["visitDuration", "(absent)", "35"],
            ["riskLevel", "LOW", "MODERATE"],
            ["urgencyBanner.show", "false", "true"],
            ["actionItems count", "2", "3"],
            ["appointments count", "1", "2"],
            ["careInstructions count", "4", "4"],
            ["conditions count", "2", "3"],
            ["summary.headline", "Present (same schema)", "Present (same schema)"],
            ["summary.soap", "Present (same schema)", "Present (same schema)"],
            ["summary.clinicalObservations", "Present (same schema)", "Present (richer functionalStatus)"],
        ],
    )

    heading(doc, "9. Retrieval Consumer Contract", 1)
    para(doc, "Ask AI retrieval service SHOULD:", bold=True)
    bullets(
        doc,
        [
            "Treat call_summaries and visit_summaries as one logical SUMMARY source family",
            "Deserialize summary_json using one shared schema model (no episode-type switch for core fields)",
            "Include episodeType, visitLocation, visitDuration in chunk metadata for citations and optional filters",
            "Index careInstructions.status and effectiveDate for medication initiation/termination queries (FR-AI-11)",
            "Preserve sourceTurnId in citations for transcript deep-link highlighting",
            "Respect needsConfirmation flags — surface uncertainty in Ask AI responses (FR-AI-6)",
        ],
    )
    para(doc, "Ask AI retrieval service SHOULD NOT:", bold=True)
    bullets(
        doc,
        [
            "Maintain separate ranking algorithms per episode type unless product requirements change",
            "Require different citation parsers for call vs visit summary items",
            "Assume visit summaries omit SOAP or typed arrays — they use the same structure",
        ],
    )

    heading(doc, "10. Current Codebase vs Planned State", 1)
    table(
        doc,
        ["Component", "Current state", "Gap for unified pipeline"],
        [
            [
                "call_summaries table",
                "Exists (V52/V61 migration)",
                "summary_json schema needs upgrade to unified sample shape",
            ],
            [
                "visit_summaries table",
                "Not in codebase yet",
                "New migration; same columns as call_summaries",
            ],
            [
                "CallSummaryService / BedrockSentimentService",
                "Emits headline, assessment, keyConcerns, recommendedActions, followUpQuestions only",
                "Unified sample (careInstructions, SOAP, conditions) not produced",
            ],
            [
                "VisitSummaryService",
                "Not implemented",
                "Mirror CallSummaryService; write to visit_summaries",
            ],
            [
                "Transcript upstream",
                "VideoCallService sends participantUserIds on end; improved multi-party end detection",
                "Summary timing at end-call unchanged; participant tracking more accurate",
            ],
            [
                "Ask AI chat",
                "BedrockAIChatService (opt-in careconnect.ai.enabled); no retrieval or medical context",
                "Must add hybrid retrieval + MedicalContextService re-wire",
            ],
            [
                "Indexing pipeline",
                "Not implemented",
                "Chunk unified summary_json → FTS + pgvector",
            ],
            [
                "Ask AI retrieval",
                "PatientContextRetrievalService stub only",
                "Single path over both tables + hybrid search",
            ],
        ],
    )

    heading(doc, "11. Integration Agreement", 1)
    para(
        doc,
        "The unified schema approach works for Ask AI retrieval. Recommendation: proceed with a single "
        "retrieval path reading from both call_summaries and visit_summaries, identical chunking and "
        "ranking logic, and episodeType available as optional metadata. No need for two separate retrieval "
        "paths unless future performance or isolation requirements emerge.",
    )
    bullets(
        doc,
        [
            "Upstream (Summaries team): emit identical summary_json inner structure; set episodeType at envelope",
            "Storage (Dominique): visit_summaries table with same column shape as call_summaries",
            "Indexing: one pipeline handler for SUMMARY_CREATED events from either table",
            "Ask AI (Retrieval): one deserializer, one chunker, one hybrid ranker, one citation builder",
        ],
    )

    heading(doc, "12. Reference Files", 1)
    bullets(
        doc,
        [
            r"C:\Users\ravic\Downloads\sample_summary_response_updated.json",
            r"C:\Users\ravic\Downloads\sample_visit_summary_response.json",
            "backend/core/src/main/java/com/careconnect/model/CallSummary.java",
            "backend/core/src/main/java/com/careconnect/service/CallSummaryService.java",
            "docs/Hybrid_Retrieval_Scope_PostgreSQL_FullText_pgvector.docx",
            "docs/Call_Transcript_Retrieval_Review.docx",
        ],
    )

    save_document(doc, OUTPUT)


if __name__ == "__main__":
    build()
