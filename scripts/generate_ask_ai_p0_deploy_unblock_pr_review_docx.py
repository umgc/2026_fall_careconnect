"""Generate Word document: PR code review for Ask AI P0 deploy unblock branch."""

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.shared import Pt

OUTPUT = (
    Path(__file__).resolve().parents[1]
    / "docs"
    / "PR_Review_Ask_AI_P0_Deploy_Unblock_feature_a-rvasireddy-ask-ai-p0-deploy-unblock.docx"
)


def heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def para(doc: Document, text: str, bold: bool = False, highlight: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    if highlight:
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def bullets(doc: Document, items: list[str], highlight_indices: set[int] | None = None) -> None:
    highlight_indices = highlight_indices or set()
    for i, item in enumerate(items):
        p = doc.add_paragraph(item, style="List Bullet")
        if i in highlight_indices:
            for run in p.runs:
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def code(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    for i, h in enumerate(headers):
        tbl.rows[0].cells[i].text = h
    for row in rows:
        cells = tbl.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val


def build() -> Document:
    doc = Document()
    title = doc.add_heading("PR Code Review", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    para(
        doc,
        "Ask AI P0 Deploy Unblock — feature/a-rvasireddy-ask-ai-p0-deploy-unblock → team-ae-develop",
    )
    para(doc, f"Review date: {date.today().isoformat()}")
    para(doc, "Reviewer: AI-assisted code review (Cursor)")
    doc.add_paragraph()

    heading(doc, "Review metadata", 1)
    table(
        doc,
        ["Field", "Value"],
        [
            ["Source branch", "feature/a-rvasireddy-ask-ai-p0-deploy-unblock"],
            ["Target branch", "team-ae-develop"],
            ["Scope", "18 tracked files (+211 / −114 lines); 8 new untracked parameter files (staging/prod)"],
            ["Commits", "Uncommitted working tree at review time (no commits ahead of team-ae-develop yet)"],
            ["Primary tasks", "0.1 Bedrock IAM · CARECONNECT_AI_ENABLED · pgvector RDS · prod Spring profile in CF"],
        ],
    )

    # ── 1. Change Summary ─────────────────────────────────────────────────────
    heading(doc, "1. Change Summary", 1)
    para(
        doc,
        "This PR unblocks Ask AI and Bedrock features in ECS/Fargate by wiring infrastructure and runtime "
        "configuration that the application already expects but that was missing or incorrect in deployed "
        "environments. It addresses Team E backlog P0 items: ECS task IAM for bedrock:InvokeModel, "
        "CARECONNECT_AI_ENABLED in ECS, pgvector-capable RDS, and use of the prod Spring profile with SSM/SendGrid.",
    )

    heading(doc, "What changed (by area)", 2)
    table(
        doc,
        ["File / area", "Change"],
        [
            [
                "cloudformation-fargate/templates/03-platform.yaml",
                "ECS task role gains inline policy: SSM read on /careconnect/*; Bedrock InvokeModel for Nova, "
                "Claude, Titan Embed v2, Voxtral, and inference profiles",
            ],
            [
                "cloudformation-fargate/templates/02-data.yaml",
                "RDS DB parameter group with rds.allowed_extensions=vector,uuid-ossp; attached to PostgreSQL 17.6 instance",
            ],
            [
                "cloudformation-fargate/templates/04-service.yaml",
                "New params: SpringProfile (default prod), CareConnectAiEnabled, AiProvider, FromEmail; "
                "ECS env: CARECONNECT_AI_ENABLED, AI_PROVIDER, EMAIL_PROVIDER, FROM_EMAIL, ENVIRONMENT, AWS_DEFAULT_REGION",
            ],
            [
                "parameters/cfdemo-service.json",
                "SpringProfile dev → prod; CareConnectAiEnabled=true",
            ],
            [
                "parameters/dev-service.json",
                "Explicit CareConnectAiEnabled=false (dev profile unchanged)",
            ],
            [
                "parameters/staging-*.json, prod-*.json (untracked)",
                "New parameter sets for staging/prod with prod profile and AI enabled",
            ],
            [
                "SsmPropertySourceInitializer.java",
                "SSM prefix derived from ENVIRONMENT (/careconnect/{env}/) or CARECONNECT_SSM_PREFIX override",
            ],
            [
                "application-prod.properties",
                "careconnect.ai.enabled=${CARECONNECT_AI_ENABLED:true}; Flyway off; schema strategy documented",
            ],
            [
                "application.properties, application-dev.properties, FlywayConfig.java",
                "Flyway disabled globally; FlywayConfig conditional on spring.flyway.enabled=true (local opt-in only)",
            ],
            [
                "SchemaPatchRunner.java, docs (PROGRAMMERS_GUIDE, DEPLOYMENT guide, README)",
                "Comments/docs aligned to SchemaPatchRunner + ddl-auto as production schema path (not Flyway)",
            ],
        ],
    )

    heading(doc, "Behavioral shift vs. base branch", 2)
    bullets(
        doc,
        [
            "Before: ECS task role had no Bedrock or SSM permissions — Bedrock calls fail with AccessDeniedException.",
            "Before: careconnect.ai.enabled=false in dev profile; ECS used SpringProfile=dev — AIChatController never registered.",
            "Before: RDS had no pgvector allowlist — CREATE EXTENSION vector in SchemaPatchRunner may fail on RDS.",
            "Before: SSM initializer hard-coded /careconnect/prod/ — cfdemo/staging could not load env-scoped secrets.",
            "After: cfdemo deploy uses prod profile + AI enabled; expects SSM under /careconnect/cfdemo/ and Bedrock model access.",
            "After: Production schema strategy explicitly documented as SchemaPatchRunner + Hibernate ddl-auto (Flyway removed from deploy path).",
        ],
    )

    heading(doc, "Verdict", 2)
    para(
        doc,
        "Approve with changes requested. The P0 infrastructure direction is correct and aligns with the codebase "
        "(BedrockModelSupport, AIChatController gating, SchemaPatchRunner pgvector patches). Before merge, address "
        "the prod-profile cutover risks for cfdemo (missing WebSocket endpoint env, SSM secret migration), Bedrock "
        "IAM/model parity, and commit the new staging/prod parameter files. Consider splitting Flyway doc/config "
        "changes into a follow-up PR if reviewers want a narrower blast radius.",
        bold=True,
    )

    # ── 2. Bug & Risk Analysis ────────────────────────────────────────────────
    heading(doc, "2. Bug & Risk Analysis", 1)

    heading(doc, "2.1 Strengths", 2)
    bullets(
        doc,
        [
            "Bedrock IAM policy lists concrete foundation-model ARNs plus inference-profile/* — matches how Claude Sonnet 4.x is invoked in code.",
            "Titan Embed Text v2 included — forward-looking for Task 4.3 embeddings (1536-dim contract in RetrievalIndexSchema).",
            "RDS parameter group correctly uses postgres17 family matching EngineVersion 17.6.",
            "CareConnectAiEnabled as explicit CFN parameter avoids magic env-only toggles and documents intent per environment.",
            "SsmPropertySourceInitializer.resolveSsmPrefix() is deterministic and backward-compatible (defaults to prod when ENVIRONMENT unset).",
            "SchemaPatchRunner pgvector patches remain idempotent; RDS allowlist unblocks CREATE EXTENSION vector.",
            "Documentation updates reduce drift between guides and actual ECS deploy behavior.",
        ],
    )

    heading(doc, "2.2 High — cfdemo prod profile cutover may break startup or features", 2)
    para(
        doc,
        "Switching cfdemo-service.json from SpringProfile=dev to prod activates application-prod.properties, including:",
        highlight=True,
    )
    bullets(
        doc,
        [
            "context.initializer.classes=SsmPropertySourceInitializer — requires SSM params under /careconnect/cfdemo/ (not /careconnect/prod/).",
            "careconnect.websocket.enabled=true and careconnect.websocket.mode=aws — requires AWS_WEBSOCKET_API_GATEWAY_ENDPOINT, which is NOT injected in 04-service.yaml.",
            "careconnect.email.provider=sendgrid — requires sendgrid-api-key from SSM or SENDGRID_API_KEY env (not in ECS Secrets block).",
            "springdoc disabled — expected for prod, but changes API discoverability for demo testers.",
        ],
    )
    para(
        doc,
        "Risk: Existing cfdemo stack updates may fail health checks, log WebSocket misconfiguration errors, or send email failures until SSM parameters are migrated and WebSocket endpoint is supplied.",
    )

    heading(doc, "2.3 High — JWT and DB password dual sourcing", 2)
    para(
        doc,
        "ECS injects DB_PASSWORD and SECURITY_JWT_SECRET from Secrets Manager. SsmPropertySourceInitializer also maps "
        "db-password and jwt-secret from SSM and adds them with addFirst() — highest precedence.",
        highlight=True,
    )
    bullets(
        doc,
        [
            "If SSM parameters exist under the new prefix but contain stale/different values, they override ECS secrets silently.",
            "If SSM parameters are missing, fallback to ECS env works — but operators may not realize which source is active.",
            "Recommendation: Either remove db-password/jwt-secret from SSM list when using Secrets Manager, or document single source of truth.",
        ],
    )

    heading(doc, "2.4 Medium — Bedrock IAM vs. approved model list mismatch", 2)
    para(
        doc,
        "BedrockModelSupport.APPROVED_MODEL_IDS includes anthropic.claude-sonnet-4-6 but the IAM policy does not grant "
        "that foundation model ARN. application-dev.properties default model is anthropic.claude-sonnet-4-5-20250929-v1:0 "
        "(covered), but runtime model selection or future config could hit AccessDenied.",
    )
    bullets(
        doc,
        [
            "Claude inference profiles in us-east-1 may require both inference-profile/* and cross-region foundation-model ARNs.",
            "Consider syncing IAM resources with APPROVED_MODEL_IDS or a CFN parameter list to avoid drift.",
        ],
    )

    heading(doc, "2.5 Medium — RDS parameter group update on existing instances", 2)
    para(
        doc,
        "Attaching a new DBParameterGroupName to an existing RDS instance triggers a pending-reboot state. "
        "Until reboot, rds.allowed_extensions may not take effect and SchemaPatchRunner logs only a warning for failed pgvector patch — "
        "retrieval_index_chunk DDL may partially apply or skip ivfflat index creation.",
    )
    bullets(
        doc,
        [
            "No CloudFormation output or doc step automates post-deploy RDS reboot.",
            "ivfflat index creation requires existing rows for optimal build — acceptable at zero rows, but worth noting.",
            "DependsOn: DatabaseParameterGroup does not replace explicit reboot guidance in runbook.",
        ],
    )

    heading(doc, "2.6 Medium — Template default mismatch: SpringProfile=prod, CareConnectAiEnabled=false", 2)
    para(
        doc,
        "04-service.yaml defaults SpringProfile to prod but CareConnectAiEnabled to false. CLI deploys that omit "
        "parameter files get prod profile (SSM, SendGrid, WebSocket AWS) with AI explicitly disabled via env var — "
        "CARECONNECT_AI_ENABLED=false overrides application-prod.properties default true.",
    )

    heading(doc, "2.7 Medium — Missing unit tests for SSM prefix resolution", 2)
    para(
        doc,
        "resolveSsmPrefix() is package-visible static logic with three branches (explicit prefix, ENVIRONMENT, default prod). "
        "No test file added — regression risk if ENVIRONMENT naming changes.",
    )

    heading(doc, "2.8 Low — Untracked staging/prod parameter files", 2)
    para(
        doc,
        "staging-*.json and prod-*.json exist in the working tree but were untracked at review time. "
        "They must be committed or deploy scripts for staging/prod will fail.",
    )

    heading(doc, "2.9 Low — Flyway scope bundled with P0 infra", 2)
    para(
        doc,
        "Disabling FlywayAutoConfiguration globally and rewriting deployment guides is logically related but expands "
        "review surface beyond CloudFormation. Rollback of infra-only changes becomes harder if Flyway docs/config are coupled.",
    )

    heading(doc, "2.10 Low — EMAIL_PROVIDER=sendgrid hardcoded for all profiles", 2)
    para(
        doc,
        "Even dev-service.json deployments receive EMAIL_PROVIDER=sendgrid in the task definition. "
        "Harmless when dev profile overrides provider, but confusing if someone deploys dev profile without checking email config.",
    )

    heading(doc, "2.11 Informational — Pre-existing IAM gaps (not introduced)", 2)
    bullets(
        doc,
        [
            "ECS task role still lacks Chime, S3, and Comprehend permissions documented in application.properties comments.",
            "Video calling and recording may still fail in Fargate until separate IAM work lands.",
        ],
    )

    heading(doc, "2.12 No concurrency / race conditions", 2)
    para(
        doc,
        "Changes are infrastructure and startup configuration only. SchemaPatchRunner remains single-threaded at startup. "
        "No new shared mutable runtime state introduced.",
    )

    # ── 3. Architecture & Style ───────────────────────────────────────────────
    heading(doc, "3. Architecture & Style", 1)

    heading(doc, "3.1 Design patterns", 2)
    bullets(
        doc,
        [
            "Infrastructure-as-code for IAM and env vars follows existing four-stack CloudFormation layout — consistent with repo conventions.",
            "Environment-specific behavior via parameter JSON files (dev vs cfdemo vs staging/prod) matches parallel deployment pattern.",
            "SsmPropertySourceInitializer as ApplicationContextInitializer preserves early secret loading for @Value injection — appropriate for prod.",
            "@ConditionalOnProperty on FlywayConfig aligns with Spring Boot idioms for optional local tooling.",
            "Explicit separation: Secrets Manager for ECS-injected secrets vs SSM for prod profile bulk config — needs clearer ownership doc.",
        ],
    )

    heading(doc, "3.2 CloudFormation quality", 2)
    bullets(
        doc,
        [
            "IAM policy embedded in task role is readable and environment-scoped (careconnect-${Environment}-ecs-task-runtime).",
            "Parameter naming (CareConnectAiEnabled, AiProvider) maps cleanly to Spring relaxed binding env vars.",
            "Missing Conditions block for IsProdProfile — prod-only env vars could be conditional to reduce dev misconfiguration.",
            "No Mappings section for environment defaults — duplicated across four service JSON files (acceptable but DRY opportunity).",
        ],
    )

    heading(doc, "3.3 Application config consistency", 2)
    bullets(
        doc,
        [
            "application-prod.properties now documents the three-layer schema strategy — good operator documentation.",
            "Global spring.flyway.enabled=false matches ECS SPRING_FLYWAY_ENABLED=false — consistent.",
            "careconnect.ai.enabled still false in application-dev.properties; ECS env override works due to Spring property precedence.",
        ],
    )

    heading(doc, "3.4 Documentation", 2)
    bullets(
        doc,
        [
            "parameters/README.md P0 table and runtime notes are accurate and actionable.",
            "DEPLOYMENT_AND_OPERATIONS_GUIDE removal of example flyway:migrate CI job prevents false expectations.",
            "cloudformation-fargate/README.md updated for new env vars — good.",
        ],
    )

    # ── 4. Recommendations ────────────────────────────────────────────────────
    heading(doc, "4. Recommendations", 1)

    heading(doc, "4.1 [High] Add WebSocket API Gateway endpoint to ECS task env", 2)
    para(doc, "Prod profile enables AWS WebSocket mode but CFN does not pass the endpoint:")
    code(
        doc,
        """# 04-service.yaml — add parameter
  WebSocketApiGatewayEndpoint:
    Type: String
    Default: ''

Conditions:
  HasWebSocketEndpoint: !Not [!Equals [!Ref WebSocketApiGatewayEndpoint, '']]

# In ContainerDefinitions Environment (when SpringProfile=prod or always):
            - Name: AWS_WEBSOCKET_API_GATEWAY_ENDPOINT
              Value: !Ref WebSocketApiGatewayEndpoint
            - Name: WEBSOCKET_ENABLED
              Value: 'true'""",
    )

    heading(doc, "4.2 [High] Document and automate cfdemo SSM migration checklist", 2)
    bullets(
        doc,
        [
            "Copy or recreate parameters from /careconnect/prod/* to /careconnect/cfdemo/* before service stack update.",
            "Minimum for smoke test: sendgrid-api-key, google-client-id, google-client-secret, jwt-secret (if not using Secrets Manager only).",
            "Add checklist to parameters/README.md or DEPLOY_2026_SUMMER.md.",
        ],
    )

    heading(doc, "4.3 [High] Resolve JWT/DB secret source of truth", 2)
    para(doc, "Option A — remove overlapping keys from SSM loader when ECS secrets are authoritative:")
    code(
        doc,
        """// SsmPropertySourceInitializer.java — remove from SSM_PARAMETERS:
// "jwt-secret", "db-password", "db-username"
// ECS already injects SECURITY_JWT_SECRET, DB_PASSWORD, DB_USER via Secrets Manager""",
    )
    para(doc, "Option B — remove JWT/DB from ECS Secrets block and load only from SSM in prod profile.")

    heading(doc, "4.4 [Medium] Sync Bedrock IAM with BedrockModelSupport", 2)
    code(
        doc,
        """# 03-platform.yaml — add missing foundation model
                  - !Sub arn:aws:bedrock:${AWS::Region}::foundation-model/anthropic.claude-sonnet-4-6
# Or maintain a single SSM/CFN parameter BedrockModelIds and generate Resource list""",
    )

    heading(doc, "4.5 [Medium] Add unit tests for resolveSsmPrefix", 2)
    code(
        doc,
        """// SsmPropertySourceInitializerTest.java
@Test
void resolveSsmPrefix_usesEnvironment() {
    MockEnvironment env = new MockEnvironment();
    env.setProperty("ENVIRONMENT", "cfdemo");
    assertThat(SsmPropertySourceInitializer.resolveSsmPrefix(env))
        .isEqualTo("/careconnect/cfdemo/");
}

@Test
void resolveSsmPrefix_explicitOverride() {
    MockEnvironment env = new MockEnvironment();
    env.setProperty("CARECONNECT_SSM_PREFIX", "/custom/path");
    assertThat(SsmPropertySourceInitializer.resolveSsmPrefix(env))
        .isEqualTo("/custom/path/");
}""",
    )

    heading(doc, "4.6 [Medium] CloudFormation Conditions for prod-only env vars", 2)
    code(
        doc,
        """Conditions:
  IsProdProfile: !Equals [!Ref SpringProfile, 'prod']

# Use Fn::If for EMAIL_PROVIDER, FROM_EMAIL, or document that dev profile ignores them
# Prefer separate Mappings for EnvironmentDefaults:
Mappings:
  EnvDefaults:
    cfdemo:
      SpringProfile: prod
      CareConnectAiEnabled: 'true'
    dev:
      SpringProfile: dev
      CareConnectAiEnabled: 'false'""",
    )

    heading(doc, "4.7 [Medium] Align CareConnectAiEnabled default with prod profile intent", 2)
    para(
        doc,
        "Either default CareConnectAiEnabled to 'true' when SpringProfile is prod (requires CFN Conditions or macro), "
        "or default SpringProfile back to dev in the template and rely on parameter files only — avoids accidental prod+AI-off deploys."
    )

    heading(doc, "4.8 [Medium] Post-deploy RDS reboot runbook step", 2)
    code(
        doc,
        """# After data stack update — verify parameter group applied, then:
aws rds reboot-db-instance --db-instance-identifier careconnect-cfdemo-db

# Verify in ECS logs after service redeploy:
#   Schema patch applied: V2607071920 – enable pgvector extension""",
    )

    heading(doc, "4.9 [Low] Commit staging/prod parameter files", 2)
    para(
        doc,
        "Add cloudformation-fargate/parameters/staging-*.json and prod-*.json to the PR so CI and deploy scripts can reference them."
    )

    heading(doc, "4.10 [Low] Optional — split Flyway documentation PR", 2)
    para(
        doc,
        "If reviewers prefer minimal infra PRs, move FlywayConfig/application.properties/guide changes to a separate "
        "'codify SchemaPatchRunner prod schema strategy' PR. Functionally fine combined if team agrees."
    )

    heading(doc, "4.11 Pre-merge test plan", 2)
    table(
        doc,
        ["Step", "Expected result"],
        [
            ["Update platform stack (cfdemo)", "Task role policy includes Bedrock + SSM"],
            ["Update data stack", "Parameter group attached; plan for RDS reboot"],
            ["Update service stack with new parameters", "Task env shows SPRING_PROFILES_ACTIVE=prod, CARECONNECT_AI_ENABLED=true"],
            ["Enable Bedrock models in console", "Model access granted for Nova, Claude, Titan Embed"],
            ["Create SSM params under /careconnect/cfdemo/", "Startup log: SSM PropertySource initialized with N parameters"],
            ["Call Bedrock-backed endpoint (symptom analyze or chat)", "No AccessDeniedException in logs"],
            ["Check SchemaPatchRunner logs", "pgvector extension patch applied successfully"],
            ["GET /v1/api/test/health", "200 OK"],
        ],
    )

    heading(doc, "Summary table — findings by severity", 2)
    table(
        doc,
        ["Severity", "Finding", "Action"],
        [
            ["High", "Prod profile on cfdemo without WebSocket endpoint env", "Add CFN parameter + env var (4.1)"],
            ["High", "SSM path change requires secret migration", "Runbook + copy params (4.2)"],
            ["High", "JWT/DB dual sourcing via SSM addFirst()", "Pick single source (4.3)"],
            ["Medium", "Bedrock IAM missing claude-sonnet-4-6", "Extend IAM list (4.4)"],
            ["Medium", "RDS reboot needed after parameter group", "Document reboot step (4.8)"],
            ["Medium", "No tests for resolveSsmPrefix", "Add unit tests (4.5)"],
            ["Low", "Untracked staging/prod JSON", "Commit files (4.9)"],
            ["Low", "Flyway changes broaden PR scope", "Optional split (4.10)"],
        ],
    )

    return doc


def main() -> None:
    doc = build()
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    main()
