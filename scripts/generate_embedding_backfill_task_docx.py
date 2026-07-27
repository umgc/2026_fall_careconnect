"""Generate Word document: Task 4.4 embedding backfill scope (no Flyway migration)."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "Task_4_4_Embedding_Backfill_Scope.docx"


def heading(doc, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def para(doc, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True


def bullets(doc, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def code(doc, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"


def table(doc, headers, rows) -> None:
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for row in rows:
        cells = t.add_row().cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = val


def build() -> None:
    doc = Document()
    title = doc.add_heading("Task 4.4 — Embedding Backfill Scope", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para(doc, "CareConnect Team E — Ask AI retrieval index")
    para(
        doc,
        "Defines the application-only backfill worker for retrieval_index_chunk rows where "
        "embedding IS NULL. Unblocks hybrid search (Task 5.1) after Task 4.3 ingest.",
    )
    doc.add_paragraph()

    heading(doc, "1. Goal", 1)
    para(
        doc,
        "Scheduled worker polls oldest embeddable chunks missing Bedrock vectors (failed 4.3 "
        "calls, pre-4.3 data) and reuses ChunkEmbeddingService.embedAndPersist. Required "
        "deliverable is application-only; optional SchemaPatchRunner partial index DDL "
        "may follow for large backfills (see §3).",
    )

    heading(doc, "2. In scope (application)", 1)
    bullets(
        doc,
        [
            "ChunkEmbeddingBackfillWorker — @Scheduled poll with careconnect.embedding.backfill.*",
            "RetrievalIndexChunkRepository.findMissingEmbeddingsForBackfill(limit)",
            "RetrievalIndexChunkRepository.countMissingEmbeddingsForBackfill()",
            "Aligned embeddable filter on findBySourceRecordIdAndEmbeddingIsNull (ingest retry path)",
            "Dual @ConditionalOnProperty: embedding.enabled + embedding.backfill.enabled",
            "Unit + config + contract tests (no live Bedrock required)",
        ],
    )

    heading(doc, "3. Optional DBA follow-up (SchemaPatchRunner — not required for MVP)", 1)
    para(
        doc,
        "Large backfills may benefit from a partial index on embeddable NULL-embedding rows. "
        "This is applied idempotently at ECS startup via SchemaPatchRunner (not Flyway). "
        "The backfill worker functions correctly without it; add/tune when NULL-embedding "
        "backlog grows or poll latency increases.",
    )
    code(
        doc,
        "-- SchemaPatchRunner patch V2607161317 (mirrors db/migration reference SQL)\n"
        "CREATE INDEX IF NOT EXISTS idx_retrieval_chunk_embedding_null_backfill\n"
        "  ON retrieval_index_chunk (indexed_at ASC NULLS LAST, id ASC)\n"
        "  WHERE embedding IS NULL\n"
        "    AND chunk_text IS NOT NULL\n"
        "    AND TRIM(BOTH FROM chunk_text) <> '';",
    )
    bullets(
        doc,
        [
            "Reference SQL: backend/core/src/main/resources/db/migration/V2607161317__add_retrieval_chunk_embedding_backfill_index.sql",
            "Runtime DDL: SchemaPatchRunner.applyRetrievalIndexChunkPatches() — same pattern as Task 4.2 FTS backfill",
            "Post-backfill: ANALYZE retrieval_index_chunk; tune ivfflat lists on idx_retrieval_chunk_embedding if needed",
        ],
    )

    heading(doc, "4. Out of scope (not required for 4.4)", 1)
    bullets(
        doc,
        [
            "New Flyway migration at ECS deploy (Flyway disabled in production)",
            "Row claim / FOR UPDATE SKIP LOCKED (duplicate Bedrock calls across ECS are idempotent)",
            "HybridRetrievalService (Task 5.1) or POST /api/ai/ask (Task 5.3)",
        ],
    )

    heading(doc, "5. Configuration", 1)
    code(
        doc,
        "careconnect.embedding.backfill.enabled=true\n"
        "careconnect.embedding.backfill.poll-interval-ms=60000\n"
        "careconnect.embedding.backfill.batch-size=50\n"
        "\n"
        "# Requires careconnect.embedding.enabled=true (bean not created when false)\n"
        "# Tests: both backfill and embedding workers disabled in application-test.properties",
    )

    heading(doc, "6. Dependencies", 1)
    table(
        doc,
        ["Task", "Provides"],
        [
            ["4.1", "RetrievalIndexService + indexed chunks"],
            ["4.2", "search_vector FTS leg (orthogonal to embeddings)"],
            ["4.3", "ChunkEmbeddingService, updateEmbedding, Titan v1 1536-d"],
            ["1.5", "retrieval_index_chunk.embedding column (existing schema)"],
        ],
    )

    heading(doc, "7. Production schema note", 1)
    para(
        doc,
        "CareConnect production/ECS does not apply Flyway at deploy time. Schema evolves via "
        "SchemaPatchRunner and Hibernate ddl-auto. Task 4.4 assumes retrieval_index_chunk "
        "already exists from Task 1.5; the worker writes embedding values through native "
        "updateEmbedding SQL. Optional partial index DDL (§3) ships in SchemaPatchRunner "
        "with reference SQL under db/migration for developers.",
    )

    heading(doc, "8. Related code", 1)
    bullets(
        doc,
        [
            "backend/core/src/main/java/com/careconnect/service/ai/embedding/ChunkEmbeddingBackfillWorker.java",
            "backend/core/src/main/java/com/careconnect/service/ai/embedding/ChunkEmbeddingService.java",
            "backend/core/src/main/java/com/careconnect/repository/retrieval/RetrievalIndexChunkRepository.java",
            "backend/core/src/main/java/com/careconnect/config/SchemaPatchRunner.java (V2607161317 optional index)",
            "docs/Team_E_Implementation_Task_Backlog.docx (task 4.4)",
            "docs/pgvector_Embedding_Strategy_Summaries_Mail_Documents.docx (checklist step 8)",
        ],
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    try:
        doc.save(OUTPUT)
        print(f"Created: {OUTPUT}")
    except PermissionError:
        alt = OUTPUT.with_name(f"{OUTPUT.stem}_refresh{OUTPUT.suffix}")
        doc.save(alt)
        print(f"Created (original locked): {alt}")


if __name__ == "__main__":
    build()
