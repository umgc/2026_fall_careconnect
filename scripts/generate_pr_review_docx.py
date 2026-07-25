"""Generate a Word document from the caregiver-visibility PR review findings."""

from datetime import date
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


def set_cell_shading(cell, hex_color: str) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def add_heading(doc: Document, text: str, level: int = 1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    return heading


def add_para(doc: Document, text: str, bold: bool = False):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)
    return paragraph


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(item, style="List Bullet")
        for run in paragraph.runs:
            run.font.size = Pt(11)


def add_code(doc: Document, text: str):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    paragraph.paragraph_format.left_indent = Inches(0.25)
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(6)
    return paragraph


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        header_cells[index].text = header
        set_cell_shading(header_cells[index], "1F3A5F")
        for paragraph in header_cells[index].paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(255, 255, 255)
    for row_index, row in enumerate(rows):
        cells = table.rows[row_index + 1].cells
        for col_index, value in enumerate(row):
            cells[col_index].text = value
            for paragraph in cells[col_index].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()


def build_document() -> Document:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.85)
        section.bottom_margin = Inches(0.85)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading("PR Code Review Findings", 0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

    meta = doc.add_paragraph()
    meta.add_run("Branch: ").bold = True
    meta.add_run("feature/team-e/david-3.15.5-caregiver-visibility")

    meta = doc.add_paragraph()
    meta.add_run("Compared to: ").bold = True
    meta.add_run("origin/team-ae-develop")

    meta = doc.add_paragraph()
    meta.add_run("Review date: ").bold = True
    meta.add_run(date.today().strftime("%B %d, %Y"))

    meta = doc.add_paragraph()
    meta.add_run("Diff scope: ").bold = True
    meta.add_run("14 commits ahead / 14 behind; ~51 files, +4,477 / -148")

    meta = doc.add_paragraph()
    meta.add_run("Verdict: ").bold = True
    meta.add_run(
        "Solid structure with good hardening (self-approval closed, tolerant confirm, "
        "audit redaction). Do not merge until the Ask AI / retrieval consent bypass is "
        "closed; also fix submitForReview demoting GRANTED and dismiss leaving visibility "
        "stuck in PENDING_REVIEW."
    )

    add_heading(doc, "Bugbot Findings", 1)
    add_table(
        doc,
        ["Severity", "Location", "Finding"],
        [
            [
                "High",
                "SummaryConsentGate.java:45 / DefaultRetrievalConsentProvider.java:18",
                "Summary HTTP gate is default-deny; Ask AI retrieval still treats "
                "care-circle link as consent",
            ],
            [
                "Medium",
                "CaregiverVisibilityService.java:57",
                "submitForReview on a GRANTED row silently demotes to PENDING_REVIEW "
                "(revokes access)",
            ],
        ],
    )

    add_heading(doc, "1. Change Summary", 1)
    add_para(
        doc,
        "This PR implements WBS 3.15.1 / 3.15.5 / 3.15.6: default-deny caregiver access "
        "to patient call summaries, with a pre-share review flow and an AI audit ledger.",
    )

    add_heading(doc, "What landed", 2)
    add_table(
        doc,
        ["Area", "What landed"],
        [
            [
                "Visibility",
                "CaregiverSummaryVisibility + grant/revoke/review API; "
                "SummaryConsentGate on summary read paths",
            ],
            [
                "Confirmation",
                "Queue/confirm/dismiss items; visibility confirm -> grant (tolerant)",
            ],
            [
                "Audit",
                "AiAuditLedger with PHI redaction / truncation / nested-depth bound",
            ],
            ["Schema", "Flyway + SchemaPatchRunner mirrors"],
            ["Frontend", "Confirmation provider/API/cache + drawer wiring"],
        ],
    )
    add_para(
        doc,
        "Goal: caregivers cannot see on_consent summaries until a patient/admin grants "
        "after review; changes are audited.",
    )

    add_heading(doc, "2. Bug & Risk Analysis", 1)

    add_heading(doc, "High - Consent hole via Ask AI / retrieval", 2)
    add_para(
        doc,
        "SummaryConsentGate correctly default-denies caregivers without GRANTED. "
        "Retrieval still treats an active caregiver-patient link as sufficient consent "
        "via DefaultRetrievalConsentProvider.",
    )
    add_para(
        doc,
        "Impact: A caregiver blocked on GET .../summary can still pull indexed "
        "on_consent summary chunks through Ask AI. That undercuts the whole visibility "
        "product.",
    )
    add_code(
        doc,
        "public boolean isCaregiverConsentGranted(...) {\n"
        "    return caregiverPatientLinkService.hasAccessToPatient(\n"
        "            caregiverUserId, patientUserId);\n"
        "}",
    )

    add_heading(doc, "Medium - Resubmit revokes an active grant", 2)
    add_para(
        doc,
        "submitForReview always sets PENDING_REVIEW and clears reviewer fields. "
        "Re-calling it while GRANTED immediately makes canViewSummaries false and queues "
        "another confirmation. Either reject when already GRANTED/PENDING_REVIEW, or "
        "require explicit revoke first.",
    )

    add_heading(doc, "Medium - Dismiss does not update visibility", 2)
    add_para(
        doc,
        "Confirm -> approveFromReview. Dismiss only marks the confirmation item "
        "DISMISSED and leaves visibility in PENDING_REVIEW. Stale pending state forever "
        "unless someone uses /grant or /revoke.",
    )

    add_heading(doc, "Medium - Duplicate confirmation spam", 2)
    add_para(
        doc,
        "Each submitForReview creates a new confirmation item with no "
        '"already PENDING for this referenceId" check. Retries / double-taps pile up '
        "review items.",
    )

    add_heading(doc, "Medium - Pending list scalability / privacy filter cost", 2)
    add_para(
        doc,
        "getPendingItemsVisibleTo loads all PENDING rows then filters in memory for "
        "non-admins. Fine for MVP; will not scale and briefly materializes other users' "
        "payloads in the JVM.",
    )

    add_heading(doc, "Other risks", 2)
    add_table(
        doc,
        ["Risk", "Why"],
        [
            [
                "Branch divergence",
                "14 behind team-ae-develop - rebase/merge before PR",
            ],
            [
                "Flyway V300626__...",
                "Odd version vs V260629 / V260713 - ordering/merge collisions likely",
            ],
            [
                "Timestamps",
                "LocalDateTime.now() without UTC in visibility transitions",
            ],
            [
                "Audit semantics",
                "All visibility transitions logged as CONFIRMATION",
            ],
            [
                "FE surface",
                "Provider/cache exist; unclear full patient review UX",
            ],
            [
                "Dirty tree",
                "missing_translations.txt modified locally",
            ],
        ],
    )
    add_para(
        doc,
        "Note: Self-approval is correctly closed for visibility items "
        "(only patient/admin via canAccessItem). Optimistic locking + 409 on races look "
        "solid.",
    )

    add_heading(doc, "3. Architecture & Style", 1)
    add_heading(doc, "Strengths", 2)
    add_bullets(
        doc,
        [
            "Clear split: gate (consent) vs domain service (visibility) vs confirmation "
            "queue vs audit",
            "@Lazy breaks the Confirmation <-> Visibility cycle cleanly",
            "Tolerant approveFromReview avoids wedging confirmations",
            "Fail-closed when patient user id cannot be resolved",
            "Good unit/integration test volume",
            "Audit payload redaction + depth bound is thoughtful",
        ],
    )
    add_heading(doc, "Gaps", 2)
    add_bullets(
        doc,
        [
            "Two consent models (HTTP summary vs retrieval) - must converge",
            "Confirmation is a generic queue but only one side-effect wired",
            "In-memory pending filtering instead of a repository query by "
            "patient/requester",
            "API paths /v1/api/... vs some /api/v3/... summary routes - document "
            "consistency",
        ],
    )

    add_heading(doc, "4. Recommendations", 1)

    add_heading(doc, "A. Wire retrieval consent to the same gate (High)", 2)
    add_para(
        doc,
        "Replace link-as-consent with the visibility gate for on_consent retrieval:",
    )
    add_code(
        doc,
        "@Component\n"
        "@RequiredArgsConstructor\n"
        "public class DefaultRetrievalConsentProvider "
        "implements RetrievalConsentProvider {\n"
        "  private final CaregiverVisibilityGate visibilityGate;\n"
        "\n"
        "  @Override\n"
        "  public boolean isCaregiverConsentGranted("
        "Long caregiverUserId, Long patientUserId) {\n"
        "    if (caregiverUserId == null || patientUserId == null) {\n"
        "      return false;\n"
        "    }\n"
        "    return visibilityGate.canViewSummaries("
        "caregiverUserId, patientUserId);\n"
        "  }\n"
        "}",
    )
    add_para(
        doc,
        'Keep link checks elsewhere for "is this person in the care circle," but not as '
        "a substitute for on_consent grant.",
    )

    add_heading(doc, "B. Make submitForReview idempotent / non-destructive (Medium)", 2)
    add_code(
        doc,
        "@Transactional\n"
        "public VisibilityResponse submitForReview("
        "Long caregiverUserId, Long patientUserId, Long requestedBy) {\n"
        "  requireCaregiverPatientLink(caregiverUserId, patientUserId);\n"
        "  CaregiverSummaryVisibility record = repository\n"
        "      .findByCaregiverUserIdAndPatientUserId("
        "caregiverUserId, patientUserId)\n"
        "      .orElseGet(() -> CaregiverSummaryVisibility.builder()\n"
        "          .caregiverUserId(caregiverUserId)\n"
        "          .patientUserId(patientUserId)\n"
        "          .build());\n"
        "\n"
        "  if (record.getStatus() == VisibilityStatus.GRANTED) {\n"
        "    throw new AppException(HttpStatus.CONFLICT, "
        '"Already granted; revoke before re-review");\n'
        "  }\n"
        "  if (record.getStatus() == VisibilityStatus.PENDING_REVIEW) {\n"
        "    return toResponse(record); // idempotent\n"
        "  }\n"
        "\n"
        "  record.setStatus(VisibilityStatus.PENDING_REVIEW);\n"
        "  // ... set requestedBy, clear reviewed*, save, createItem once, audit\n"
        "}",
    )

    add_heading(doc, "C. Dismiss should close the review gate (Medium)", 2)
    add_code(
        doc,
        "private void applyDismissSideEffect("
        "ConfirmationItem item, Long resolverUserId) {\n"
        "  if (item.getSourceType() != "
        "ConfirmationSourceType.CAREGIVER_VISIBILITY) {\n"
        "    return;\n"
        "  }\n"
        "  Long[] ids = CaregiverVisibilityService.parseVisibilityReference("
        "item.getReferenceId());\n"
        "  if (ids == null) {\n"
        "    return;\n"
        "  }\n"
        "  // transition PENDING_REVIEW -> REVOKED (or DENIED)\n"
        "  caregiverVisibilityService.revoke(ids[0], ids[1], resolverUserId);\n"
        "}",
    )
    add_para(doc, "Call from dismiss(...) after save, mirroring confirm.")

    add_heading(doc, "D. Bound pending queries in the repository (Medium)", 2)
    add_code(
        doc,
        "// Prefer a targeted query instead of findAll PENDING + filter\n"
        "List<ConfirmationItem> findPendingVisibleToPatient(Long patientUserId);",
    )
    add_para(
        doc,
        "Optionally store patientUserId as a first-class column on confirmation_items "
        "for efficient filtering.",
    )

    add_heading(doc, "E. Before merge checklist", 2)
    add_bullets(
        doc,
        [
            "Rebase onto latest team-ae-develop (currently 14 behind)",
            "Rename/renumber V300626__create_confirmation_items_table.sql to the team "
            "VYYMMDDHHmm scheme",
            "Use LocalDateTime.now(ZoneOffset.UTC) on visibility transitions",
            "Do not commit unrelated missing_translations.txt noise",
            "Close Ask AI / retrieval consent bypass (Recommendation A) before merge",
        ],
    )

    add_heading(doc, "Merge Recommendation", 1)
    add_para(
        doc,
        "Do not merge until Recommendation A (retrieval consent) is implemented. Also "
        "address B and C to avoid silent grant revocation and stuck PENDING_REVIEW "
        "state after dismiss. Architecture is otherwise clean and well tested for an "
        "MVP of the visibility/confirmation/audit stack.",
    )
    return doc


def main() -> None:
    output = Path(__file__).resolve().parents[1] / (
        "PR_Review_Caregiver_Visibility.docx"
    )
    build_document().save(output)
    print(output)
    print(f"size_bytes={output.stat().st_size}")


if __name__ == "__main__":
    main()
