"""Generate Word document: PR code review for Task 3.14.4 Textract OCR fallback branch."""

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.shared import Pt

OUTPUT = (
    Path(__file__).resolve().parents[1]
    / "docs"
    / "PR_Review_Textract_OCR_Fallback_Task_3.14.4.docx"
)


def heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def para(doc: Document, text: str, bold: bool = False, highlight: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    if highlight:
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def bullets(doc: Document, items: list[str], highlight_indices: set[int] | None = None) -> None:
    highlight_indices = highlight_indices or set()
    for i, item in enumerate(items):
        p = doc.add_paragraph(item, style="List Bullet")
        if i in highlight_indices:
            for run in p.runs:
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def code(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    for i, h in enumerate(headers):
        tbl.rows[0].cells[i].text = h
    for row in rows:
        cells = tbl.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val


def build() -> Document:
    doc = Document()
    title = doc.add_heading("PR Code Review", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    para(doc, "Task 3.14.4 — Textract OCR fallback — feature/3.14.4-textract-ocr-fallback → team-ae-develop")
    para(doc, f"Review date: {date.today().isoformat()}")
    para(doc, "Reviewer: AI-assisted code review (Cursor)")
    doc.add_paragraph()

    heading(doc, "Review metadata", 1)
    table(
        doc,
        ["Field", "Value"],
        [
            ["Source branch", "feature/3.14.4-textract-ocr-fallback"],
            ["Target branch", "team-ae-develop"],
            ["Scope", "7 files changed (+463 / −55 lines)"],
            ["Commits", "cbac2e5 — Implement Textract OCR fallback (Task 3.14.4)"],
            ["Commits", "ce64ef1 — Add E2E test for full digest parse"],
        ],
    )

    # ── 1. Change Summary ─────────────────────────────────────────────────────
    heading(doc, "1. Change Summary", 1)
    para(
        doc,
        "This PR implements Task 3.14.4: when USPS Informed Delivery digest HTML provides "
        "insufficient mailpiece metadata (missing sender and/or generic summary text such as "
        "'mail', 'image', or 'campaign'), the parser falls back to AWS Textract OCR on inline "
        "data: URL mailpiece images to extract a sender name and optional summary line from the "
        "top 15% of the image.",
    )

    heading(doc, "What changed", 2)
    table(
        doc,
        ["File / area", "Change"],
        [
            [
                "GmailParser.java",
                "Replaces unconditional inferSenderFromImage() with enrichWithOcrFallbackWhenMetadataInsufficient(); "
                "OCR runs only when sender is blank OR summary is generic; decodes data: URLs and calls MailpieceOcrService",
            ],
            [
                "MailpieceOcrService.java",
                "Adds extractMailpieceMetadata() returning sender + optional summaryLine; refactors top-region "
                "line collection; extractTopLeftLabel() delegates to new method",
            ],
            [
                "MailpieceOcrResult.java",
                "New record: (sender, summaryLine)",
            ],
            [
                "GmailParserOcrFallbackTest.java",
                "Unit tests: sender fill, skip when metadata present, summary enrichment, CID campaign path, "
                "skip non-data URLs",
            ],
            [
                "GmailParserOcrFallbackE2ETest.java",
                "E2E: real parser + real OCR service, mocked TextractClient, 3-mailpiece HTML fixture",
            ],
            [
                "MailpieceOcrServiceTest.java",
                "New tests for extractMailpieceMetadata sender+summary extraction",
            ],
            [
                "gmail-digest-ocr-fallback.html",
                "Fixture: 3 mailpieces (no metadata, generic alt, complete metadata)",
            ],
        ],
    )

    heading(doc, "Behavioral shift vs. base branch", 2)
    bullets(
        doc,
        [
            "Before: OCR (extractTopLeftLabel) ran whenever HTML sender extraction failed, regardless of summary quality.",
            "After: OCR runs only when isMetadataInsufficientForOcr() is true — blank sender OR generic/blank summary.",
            "After: OCR can enrich summary (subject) from a second top-region line, not just sender.",
            "After: OCR is limited to data: inline images; remote https:// thumbnails are not processed.",
        ],
    )

    heading(doc, "Verdict", 2)
    para(
        doc,
        "Approve with minor recommendations. The change is focused, well-tested, and correctly gates expensive "
        "Textract calls behind insufficient-metadata detection. No blocking security or correctness issues were found.",
        bold=True,
    )

    # ── 2. Bug & Risk Analysis ────────────────────────────────────────────────
    heading(doc, "2. Bug & Risk Analysis", 1)

    heading(doc, "2.1 Strengths", 2)
    bullets(
        doc,
        [
            "Conditional OCR reduces cost vs. always-on sender inference on every blank-sender mailpiece.",
            "Graceful degradation: @Autowired(required=false) + @ConditionalOnProperty(careconnect.aws.enabled=true) "
            "means OCR is silently skipped when AWS/Textract is unavailable.",
            "Textract failures are caught and logged; parser returns original metadata unchanged.",
            "CID inline images resolved via payload.inlineCidData() before OCR (campaign path tested).",
            "E2E fixture validates the important negative case: complete HTML metadata must not be overridden by OCR.",
        ],
    )

    heading(doc, "2.2 Medium — HTTPS / non-inline images skip OCR", 2)
    para(
        doc,
        "runOcrOnImage() only decodes data: URLs. Mailpieces whose thumbnails are hosted at https:// URLs "
        "(common in some digest templates) will never receive OCR even when metadata is insufficient.",
        highlight=True,
    )
    bullets(
        doc,
        [
            "GmailParserOcrFallbackTest.toDomain_ocrSkipped_forNonDataImageUrls() documents this limitation.",
            "Risk: production digests with remote-only images retain null/generic metadata.",
            "Mitigation options: fetch remote image bytes (with size/timeout limits) or pre-download during Gmail sync.",
        ],
    )

    heading(doc, "2.3 Medium — Sequential Textract calls during digest parse", 2)
    para(
        doc,
        "Each insufficient mailpiece triggers a synchronous detectDocumentText API call during toDomain(). "
        "A digest with 10+ mailpieces needing OCR could add significant latency and AWS cost.",
    )
    bullets(
        doc,
        [
            "No caching of OCR results by image hash or mailpiece ID.",
            "No concurrency limit or batch API usage.",
            "Consider async enrichment post-parse if digest sync latency becomes a user-visible issue.",
        ],
    )

    heading(doc, "2.4 Low — OCR result discarded when sender line fails validation", 2)
    para(
        doc,
        "MailpieceOcrService.extractMailpieceMetadata() returns Optional.empty() if no line passes looksLikeSender(), "
        "even when a valid summaryLine was detected. This affects the case where HTML sender is present but summary "
        "is generic: OCR must find a sender line to return any enrichment.",
    )
    code(
        doc,
        """String sender = topRegionLines.stream()
        .filter(this::looksLikeSender)
        .findFirst()
        .orElse(null);
if (sender == null) {
    return Optional.empty();  // summaryLine is lost
}""",
    )

    heading(doc, "2.5 Low — Wrong but present sender is never corrected", 2)
    para(
        doc,
        "If HTML provides a non-blank sender that is incorrect (e.g., marketing boilerplate that passes sanitizeSender) "
        "and summary is adequate, OCR will not run. This is a pre-existing class of parsing error, slightly narrowed "
        "by the new gate but not introduced by this PR.",
    )

    heading(doc, "2.6 Low — Summary fallback may duplicate sender", 2)
    para(
        doc,
        "resolveSummaryWithOcrFallback() uses firstNonBlank(ocr.summaryLine(), ocr.sender()). When Textract finds "
        "only one usable line, subject may equal sender — acceptable but redundant in the UI.",
    )

    heading(doc, "2.7 Low — looksLikeSummary is less strict than looksLikeSender", 2)
    para(
        doc,
        "looksLikeSummary does not filter generic tokens 'mail', 'campaign', or 'image' that looksLikeSender rejects. "
        "An OCR summary line equal to 'mail' could slip through if Textract reads it from the image.",
    )

    heading(doc, "2.8 Low — E2E test does not assert Textract call count", 2)
    para(
        doc,
        "GmailParserOcrFallbackE2ETest verifies output correctness but does not verify detectDocumentText was called "
        "exactly twice (mailpieces 1 and 2). A regression that over-invokes Textract would not be caught.",
    )

    heading(doc, "2.9 Informational — Debug logging via System.out.println", 2)
    para(
        doc,
        "decodeDataUrlImage() logs decode failures with System.out.println rather than SLF4J. This matches existing "
        "GmailParser debug patterns but is noisy in production when OCR triggers frequently.",
    )

    heading(doc, "2.10 No race conditions", 2)
    para(
        doc,
        "Parsing is synchronous and stateless per call. MailpieceOcrService is a stateless Spring bean. "
        "No concurrency concerns identified.",
    )

    # ── 3. Architecture & Style ───────────────────────────────────────────────
    heading(doc, "3. Architecture & Style", 1)

    heading(doc, "3.1 Design patterns", 2)
    bullets(
        doc,
        [
            "Optional service injection (@Autowired(required=false)) cleanly supports AWS-off dev environments.",
            "MailpieceOcrResult record is a clear, immutable DTO between OCR service and parser.",
            "extractTopLeftLabel() preserved as a thin delegate — backward compatible for any existing callers.",
            "Private MailpieceFields record in GmailParser groups sender/summary through enrichment pipeline.",
            "Test seam setMailpieceOcrService() avoids Spring context in unit tests — good pattern.",
        ],
    )

    heading(doc, "3.2 Code organization", 2)
    bullets(
        doc,
        [
            "OCR trigger logic (isMetadataInsufficientForOcr) lives in GmailParser; heuristics (looksLikeSender/Summary) "
            "live in MailpieceOcrService — reasonable separation.",
            "collectTopRegionLines() extraction improves readability vs. inline stream in original extractTopLeftLabel.",
            "Three parse paths updated consistently: toMailPiece(), standalone img path, and campaign table path.",
        ],
    )

    heading(doc, "3.3 Minor inconsistency", 2)
    para(
        doc,
        "Campaign path (parseCampaignMailPieces) assigns summary = fields.summary() then later normalizeSummary(), "
        "while toMailPiece() calls normalizeSummary(fields.summary(), sender) inline. Functionally equivalent but "
        "slightly inconsistent ordering.",
    )

    heading(doc, "3.4 Test quality", 2)
    bullets(
        doc,
        [
            "Unit tests cover positive, negative, CID, and non-data URL cases.",
            "E2E test uses real GmailParser + real MailpieceOcrService with mocked Textract boundary — excellent integration coverage.",
            "HTML fixture clearly documents intent per mailpiece with comments.",
            "MailpieceOcrServiceTest extended without breaking existing extractTopLeftLabel coverage.",
        ],
    )

    # ── 4. Recommendations ────────────────────────────────────────────────────
    heading(doc, "4. Recommendations", 1)

    heading(doc, "4.1 [Medium] Assert Textract invocation count in E2E test", 2)
    para(doc, "Prevent over-calling Textract for mailpieces with complete metadata:")
    code(
        doc,
        """import static org.mockito.Mockito.times;
import static org.mockito.Mockito.verify;

// ... after parser.toDomain(payload):
verify(textractClient, times(2)).detectDocumentText(any(DetectDocumentTextRequest.class));""",
    )

    heading(doc, "4.2 [Medium] Align looksLikeSummary with looksLikeSender generic filters", 2)
    para(doc, "Reject generic OCR summary candidates:")
    code(
        doc,
        """private boolean looksLikeSummary(String text) {
    if (text == null || text.isBlank() || text.length() < 3) {
        return false;
    }
    String lower = text.strip().toLowerCase(Locale.ROOT);
    if (lower.startsWith("learn more") || lower.contains("click") || lower.contains("visit")) {
        return false;
    }
    if (lower.contains("ridealong") || lower.contains("ride along")) {
        return false;
    }
    if (lower.equals("campaign") || lower.equals("mail") || lower.equals("image")) {
        return false;
    }
    return lower.matches(".*[a-z].*");
}""",
    )

    heading(doc, "4.3 [Low] Return partial OCR result when only summary is usable", 2)
    para(
        doc,
        "For sender-present / summary-generic cases, allow summary enrichment even if OCR sender line fails:"
    )
    code(
        doc,
        """String sender = topRegionLines.stream()
        .filter(this::looksLikeSender)
        .findFirst()
        .orElse(null);

String summaryLine = topRegionLines.stream()
        .filter(this::looksLikeSummary)
        .filter(line -> sender == null || !line.equalsIgnoreCase(sender))
        .findFirst()
        .orElse(null);

if (sender == null && summaryLine == null) {
    return Optional.empty();
}
return Optional.of(new MailpieceOcrResult(sender, summaryLine));""",
    )
    para(doc, "Update resolveSummaryWithOcrFallback() to handle null ocr.sender() when only summaryLine is present.")

    heading(doc, "4.4 [Low] Replace System.out.println with structured logging", 2)
    code(
        doc,
        """// In GmailParser — add @Slf4j or static logger
log.debug("Failed to decode inline image for OCR ({}): {}", metadata, ex.getMessage());""",
    )

    heading(doc, "4.5 [Low] Guard against oversized inline images before Textract", 2)
    code(
        doc,
        """private static final int MAX_OCR_IMAGE_BYTES = 5 * 1024 * 1024; // 5 MB

private Optional<MailpieceOcrResult> runOcrOnImage(String imageSrc) {
    if (mailpieceOcrService == null || isBlank(imageSrc)) {
        return Optional.empty();
    }
    return decodeDataUrlImage(imageSrc)
        .filter(decoded -> decoded.bytes().length <= MAX_OCR_IMAGE_BYTES)
        .flatMap(decoded -> mailpieceOcrService.extractMailpieceMetadata(
            decoded.bytes(), decoded.metadata()));
}""",
    )

    heading(doc, "4.6 [Future] Remote image OCR support", 2)
    para(
        doc,
        "If production digests rely on https:// thumbnails, add a bounded HTTP fetch (timeout, max bytes, "
        "allowlisted hosts) before calling extractMailpieceMetadata(). Track as follow-up unless Task 3.14.4 "
        "scope explicitly covers inline-only images.",
    )

    heading(doc, "4.7 [Future] OCR result caching", 2)
    para(
        doc,
        "Cache Textract results keyed by SHA-256 of image bytes (in-memory or Redis) to avoid duplicate API calls "
        "when the same digest is re-parsed or mailpieces share thumbnails.",
    )

    # ── File-level review comments ────────────────────────────────────────────
    heading(doc, "File-level review comments", 1)
    table(
        doc,
        ["File", "Severity", "Comment"],
        [
            [
                "GmailParser.java",
                "Info",
                "Good refactor from inferSenderFromImage to gated enrichWithOcrFallbackWhenMetadataInsufficient.",
            ],
            [
                "GmailParser.java",
                "Low",
                "System.out.println for decode failures; prefer SLF4J.",
            ],
            [
                "GmailParser.java",
                "Medium",
                "data:-only OCR leaves https thumbnail mailpieces unenriched.",
            ],
            [
                "MailpieceOcrService.java",
                "Low",
                "All-or-nothing Optional when sender line missing discards valid summaryLine.",
            ],
            [
                "MailpieceOcrService.java",
                "Low",
                "looksLikeSummary should reject generic tokens like looksLikeSender.",
            ],
            [
                "MailpieceOcrResult.java",
                "Info",
                "Clean record; consider @Nullable on summaryLine in Javadoc if sender-only results allowed later.",
            ],
            [
                "GmailParserOcrFallbackE2ETest.java",
                "Low",
                "Add verify(textractClient, times(2)) to lock negative path for mailpiece 3.",
            ],
            [
                "GmailParserOcrFallbackTest.java",
                "Info",
                "Solid unit coverage including CID campaign scenario.",
            ],
            [
                "gmail-digest-ocr-fallback.html",
                "Info",
                "Well-documented fixture; mailpiece 3 correctly has complete metadata despite generic alt text.",
            ],
        ],
    )

    heading(doc, "Test plan (recommended before merge)", 1)
    bullets(
        doc,
        [
            "Run: mvnw test -Dtest=GmailParserOcrFallbackTest,GmailParserOcrFallbackE2ETest,MailpieceOcrServiceTest",
            "Verify careconnect.aws.enabled=true in staging with real Textract on a sample digest.",
            "Confirm digest parse latency acceptable with 5–10 OCR-eligible mailpieces.",
            "Validate mailpieces with https-only thumbnails behave as expected (no OCR, metadata unchanged).",
        ],
    )

    return doc


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    build().save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    main()
